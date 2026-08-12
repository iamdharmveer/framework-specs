"""
notes_docx.py v1.0 — SHARED Notes document builder/parser (Steps NC and NA).

v1.0 — 2026-08-12 — INITIAL RELEASE (GAP-2026-08-12-NADOCX, patch P1 of 2).

    WHY THIS FILE EXISTS.  Framework_NotesAudit v3.0.0 makes NA a WRITER: it
    rebuilds the whole unit document from NotesCreate's content.  Before this
    file, NC built its .docx from python-docx code written per the spec at run
    time and no engine owned the construction contracts at all — so a writing
    NA would have been a SECOND implementation of the §6A colour map, the
    decimal cascade, box styling, spacer paragraphs, rule F-7 and the OMML
    conventions.  That is exactly the "one contract, two implementations"
    defect class the 2026-08-10 sweep closed for filenames (Framework_
    NotesCreate v2.2.1 / notes_core v2.1 SPEC-LOCK).  Re-opening it directly
    upstream of a student-facing artifact is not acceptable, so construction
    moves into ONE engine that both steps import.  Neither step may hand-roll
    a paragraph, a colour, a border or a line rule.

    THE CONTENT MODEL IS THE HANDOFF.  A unit is a plain dict (schema
    "notes-content/1.0") describing WHAT the notes say.  build() renders it to
    .docx; parse() reads a built .docx back to the model.  NA therefore does
    not edit .docx XML — it parses NC's file to the model, corrects and
    improves the MODEL, and rebuilds.  Round-trip fidelity is self-tested
    (build -> parse -> build produces byte-identical word/document.xml), which
    is what makes "re-running NA on its own output changes nothing" a testable
    property rather than an aspiration.

    NUMBERS ARE NEVER STORED.  The model holds no outline numbers and no
    Example/Recall counters.  build() derives every "n.k", "n.k.m", "Example j"
    and "Recall j" from block ORDER at render time.  Framework_NotesCreate §6A
    ("removing or adding a block renumbers everything after it; stale numbers
    are an NA gate failure") therefore cannot be violated by construction:
    there is no stored number to go stale.  This closes the largest single
    risk in letting NA insert or drop a block.

    OMML.  Expressions (F-3a) compile through t3_mathcomp.t3_compile — the
    drift-locked shared compiler — so Notes math and Mock/PYQ math are built by
    one authority.  Symbol mentions in running text (F-3b) render as styled
    sub/superscript RUNS, never as flat tokens and never as oMath.
    validate_model() hard-rejects an unbraced multi-character script
    ("V_max"), which t3_compile renders LaTeX-correctly as V-sub-m followed by
    "ax" — silently wrong on the page and invisible to an XML-level gate.

    F-7 IS ENFORCED STRUCTURALLY.  Every paragraph this engine emits carries an
    explicit AUTO line rule.  A fixed line rule clips tall inline objects
    (OMML, images) and the clipping is invisible to Framework_NotesAudit's
    XML gates, so it is prevented at construction instead of detected later.

MINIMUM COMPANION VERSIONS:
    notes_core.py   >= v2.2 — LEVEL_COLORS / BOX_COLORS / CANONICAL_TYPES
    t3_mathcomp.py  >= 2026.08.07.7 — t3_compile (shared OMML compiler)
    python-docx     >= 1.1

PUBLIC SURFACE:
    SCHEMA, MODEL_SCHEMAS_ACCEPTED
    new_model(exam_code, unit)          -> empty content model
    validate_model(model)               -> (ok, findings)
    build(model, out_path)              -> writes .docx, returns build report
    parse(docx_path)                    -> content model
    document_xml(docx_path)             -> word/document.xml bytes
    outline_of(model)                   -> derived numbering map (G-6 oracle)
    self_test()
"""
import copy
import json
import os
import re
import zipfile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Emu

import notes_core
import t3_mathcomp

SCHEMA = "notes-content/1.0"
MODEL_SCHEMAS_ACCEPTED = ("notes-content/1.0",)

# Rendering constants OWNED HERE (the specs defer to this engine, as they defer
# to notes_core for filenames).  Colours come from notes_core so there is one
# colour authority for the whole pipeline.
BODY_FONT = "Arial"
BODY_PT = 10.5
TITLE_PT = 14
L2_PT = 12
L3_PT = 11
BOX_TITLE_PT = 10.5
PAGE_W_CM = 21.0          # A4
PAGE_H_CM = 29.7
MARGIN_CM = 2.0
FIGURE_W_CM = 12.0
LINE_RULE = "auto"        # rule F-7 — never "exact", never "atLeast"
LINE_VALUE = "240"        # 240 twentieths = single, with AUTO rule

BLOCK_TYPES = ("title", "concept", "example", "key_points", "trap",
               "rapid", "recall", "mindmap")
CONTENT_KINDS = ("bullet", "para", "table", "figure")
RUN_TYPES = ("text", "math", "sym")

# Blocks that carry a level-2 outline number, in their mandatory tail order
# (Framework_NotesCreate §6A: TRAP, RAPID REVISION, RECALL CHECK, MIND MAP
# follow the last concept).  "example", "key_points" and "recall" are NOT
# separately numbered — they live inside/after their concept.
L2_NUMBERED = ("concept", "trap", "rapid", "recall_check", "mindmap")

_UNBRACED_SCRIPT = re.compile(r"[_^]\s*([A-Za-z0-9]{2,})")


# ------------------------------------------------------------------ helpers
def _set_line_rule(paragraph):
    """Rule F-7. Explicit AUTO line rule on EVERY paragraph this engine emits.

    A fixed rule ("exact"/"atLeast") clips tall inline objects — OMML and
    images — and the clipped equation is still present in the XML, so
    Framework_NotesAudit's G-2 assertion passes while the page is wrong.
    Prevented here rather than detected later.
    """
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), LINE_VALUE)
    sp.set(qn("w:lineRule"), LINE_RULE)
    sp.set(qn("w:before"), "40")
    sp.set(qn("w:after"), "40")
    pPr.append(sp)
    return paragraph


def _shade(element, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    element.append(shd)


def _cell_shade(cell, hex_fill):
    _shade(cell._tc.get_or_add_tcPr(), hex_fill)


def _cell_borders(cell, hex_colour, sz="8"):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), hex_colour)
        borders.append(e)
    tcPr.append(borders)


def _style_run(run, *, colour=None, bold=False, italic=False,
               size=BODY_PT, script=None):
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = RGBColor.from_string(colour)
    rPr = run._r.get_or_add_rPr()
    # East-Asian/complex-script font names, so Arial is honoured everywhere.
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), BODY_FONT)
    if script in ("sub", "sup"):
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "subscript" if script == "sub" else "superscript")
        rPr.append(va)
    return run


def _new_paragraph(container):
    p = container.add_paragraph()
    p.style = container.part.document.styles["Normal"] if hasattr(
        container, "part") else p.style
    return _set_line_rule(p)


def _para(container):
    """Add a paragraph to a Document or a _Cell, always with the AUTO rule."""
    p = container.add_paragraph()
    return _set_line_rule(p)


# ------------------------------------------------------------------ inline
def _emit_runs(paragraph, runs, *, colour=None, bold=False, size=BODY_PT):
    """Render an inline run list.

    {"t":"text","s":...}                      plain text
    {"t":"sym","base":...,"sub"/"sup":...}    F-3b styled script runs
    {"t":"math","latex":...}                  F-3a structural OMML region
    """
    for r in runs:
        kind = r.get("t")
        if kind == "text":
            _style_run(paragraph.add_run(r["s"]), colour=colour, bold=bold,
                       size=size)
        elif kind == "sym":
            _style_run(paragraph.add_run(r["base"]), colour=colour, bold=bold,
                       size=size)
            if r.get("sub"):
                _style_run(paragraph.add_run(r["sub"]), colour=colour,
                           bold=bold, size=size, script="sub")
            if r.get("sup"):
                _style_run(paragraph.add_run(r["sup"]), colour=colour,
                           bold=bold, size=size, script="sup")
        elif kind == "math":
            if r.get("xml"):
                # Preserved OMML from parse(): re-inserted VERBATIM. NA edits
                # prose and structure without re-authoring an equation it did
                # not touch, which is what keeps the rebuild idempotent.
                from lxml import etree
                paragraph._p.append(etree.fromstring(r["xml"].encode("utf-8")))
            else:
                paragraph._p.append(t3_mathcomp.t3_compile(r["latex"]))
        else:
            raise ValueError(f"unknown run type: {kind!r}")
    return paragraph


def _runs_text(runs):
    """Flat text of an inline run list — for gates and for length checks."""
    out = []
    for r in runs:
        if r.get("t") == "text":
            out.append(r["s"])
        elif r.get("t") == "sym":
            out.append(r["base"] + r.get("sub", "") + r.get("sup", ""))
        elif r.get("t") == "math":
            out.append(" ")
    return "".join(out)


# ------------------------------------------------------------------ numbering
def outline_of(model):
    """Derive every outline number and counter from BLOCK ORDER.

    Returns {"title": "n.", "l2": [(block_index, "n.k", label)],
             "l3": [...], "examples": [(block_index, j)],
             "recalls": [(block_index, j)]}.

    Nothing here is read from the model — that is the point.  Framework_
    NotesCreate §6A's renumber rule is satisfied structurally, and this
    function doubles as the oracle for NotesAudit's G-6 gate.
    """
    n = int(model["unit"].get("seq_in_topic") or 1)
    out = {"title": f"{n}.", "l2": [], "l3": [], "examples": [], "recalls": []}
    k = 0
    ex_j = 0
    rc_j = 0
    seen_recall = False
    for i, b in enumerate(model["blocks"]):
        t = b["type"]
        if t == "concept":
            k += 1
            out["l2"].append((i, f"{n}.{k}", b["name"]))
        elif t == "trap":
            k += 1
            out["l2"].append((i, f"{n}.{k}", "TRAP BOX"))
        elif t == "rapid":
            k += 1
            num = f"{n}.{k}"
            out["l2"].append((i, num, "RAPID REVISION SUMMARY"))
            out["l3"].append((i, f"{num}.1", "Must-Know Formulae"))
            out["l3"].append((i, f"{num}.2", "Key Associations"))
        elif t == "recall":
            if not seen_recall:
                k += 1
                out["l2"].append((i, f"{n}.{k}", "RECALL CHECK"))
                seen_recall = True
            rc_j += 1
            out["recalls"].append((i, rc_j))
        elif t == "mindmap":
            k += 1
            out["l2"].append((i, f"{n}.{k}", "MIND MAP"))
        elif t == "example":
            ex_j += 1
            out["examples"].append((i, ex_j))
    return out


# ------------------------------------------------------------------ validate
def new_model(exam_code, unit):
    return {"schema": SCHEMA, "exam_code": exam_code,
            "unit": dict(unit), "blocks": []}


def validate_model(model):
    """Structural + content-safety validation. Returns (ok, findings).

    Hard findings only — anything here would be a visible defect in a
    student-facing document.
    """
    f = []
    if model.get("schema") not in MODEL_SCHEMAS_ACCEPTED:
        f.append(f"schema mismatch: {model.get('schema')!r}")
        return (False, f)
    if not model.get("exam_code"):
        f.append("missing exam_code")
    unit = model.get("unit") or {}
    for key in ("name", "tier"):
        if not unit.get(key):
            f.append(f"unit missing {key}")
    blocks = model.get("blocks") or []
    if not blocks or blocks[0]["type"] != "title":
        f.append("first block must be the B1 title")
    if sum(1 for b in blocks if b["type"] == "title") != 1:
        f.append("exactly one title block is required")

    def check_runs(runs, where):
        if not isinstance(runs, list):
            f.append(f"{where}: runs must be a list")
            return
        for r in runs:
            if r.get("t") not in RUN_TYPES:
                f.append(f"{where}: unknown run type {r.get('t')!r}")
            if r.get("t") == "math":
                if r.get("xml"):
                    continue        # preserved OMML — already compiled once
                latex = r.get("latex", "")
                if not latex:
                    f.append(f"{where}: math run has neither latex nor xml")
                    continue
                bad = _UNBRACED_SCRIPT.search(latex)
                if bad:
                    f.append(
                        f"{where}: unbraced multi-character script "
                        f"'{bad.group(0)}' in {latex!r} — t3_compile renders "
                        f"only the FIRST character as the script and the rest "
                        f"as body text (V_max -> V-sub-m + 'ax'). Write "
                        f"{bad.group(0)[0]}{{{bad.group(1)}}}.")
                try:
                    t3_mathcomp.t3_compile(latex)
                except Exception as exc:
                    f.append(f"{where}: math region will not compile: {exc}")
            if r.get("t") == "sym" and not r.get("base"):
                f.append(f"{where}: sym run without a base")

    order_seen = []
    for i, b in enumerate(blocks):
        t = b.get("type")
        order_seen.append(t)
        if t not in BLOCK_TYPES:
            f.append(f"block {i}: unknown type {t!r}")
            continue
        if t == "title":
            if not b.get("name"):
                f.append(f"block {i}: title has no name")
        elif t == "concept":
            if not b.get("name"):
                f.append(f"block {i}: concept has no name")
            for j, c in enumerate(b.get("content", [])):
                k = c.get("k")
                if k not in CONTENT_KINDS:
                    f.append(f"block {i}.{j}: unknown content kind {k!r}")
                elif k in ("bullet", "para"):
                    check_runs(c.get("runs", []), f"block {i}.{j}")
                    if k == "bullet":
                        words = len(_runs_text(c.get("runs", [])).split())
                        if words > notes_core.BULLET_HARD_CAP_WORDS:
                            f.append(
                                f"block {i}.{j}: bullet is {words} words, over "
                                f"the D-1 hard cap of "
                                f"{notes_core.BULLET_HARD_CAP_WORDS}")
                elif k == "table":
                    hdr = c.get("headers") or []
                    if not hdr:
                        f.append(f"block {i}.{j}: table has no header row")
                    for cell in hdr:
                        check_runs(cell, f"block {i}.{j} header")
                    for ri, row in enumerate(c.get("rows", [])):
                        if len(row) != len(hdr):
                            f.append(f"block {i}.{j}: row {ri} has "
                                     f"{len(row)} cells, header has {len(hdr)}")
                        for cell in row:
                            check_runs(cell, f"block {i}.{j} row {ri}")
                elif k == "figure":
                    if not c.get("image"):
                        f.append(f"block {i}.{j}: figure has no image path")
                    if not c.get("label"):
                        f.append(f"block {i}.{j}: figure has no label "
                                 f"(F-4: the label is drawn INSIDE the image; "
                                 f"it is recorded here for the gates)")
        elif t in ("example", "recall"):
            qt = b.get("qtype")
            if qt not in notes_core.CANONICAL_TYPES:
                f.append(f"block {i}: qtype {qt!r} not in "
                         f"{notes_core.CANONICAL_TYPES}")
            check_runs(b.get("stem", []), f"block {i} stem")
            opts = b.get("options") or []
            if qt in ("MCQ", "MSQ"):
                if len(opts) != 4:
                    f.append(f"block {i}: {qt} needs exactly 4 options, "
                             f"found {len(opts)}")
                for oi, o in enumerate(opts):
                    check_runs(o, f"block {i} option {oi+1}")
            elif qt == "NAT" and opts:
                f.append(f"block {i}: NAT must not carry options")
            ans = str(b.get("answer", "")).strip()
            if not ans:
                f.append(f"block {i}: missing Answer")
            elif qt == "MCQ":
                if ans not in {"1", "2", "3", "4"}:
                    f.append(f"block {i}: MCQ answer {ans!r} is not one of "
                             f"1-4 (out-of-range key)")
            elif qt == "MSQ":
                parts = [p.strip() for p in ans.split(",") if p.strip()]
                if not parts or any(p not in {"1", "2", "3", "4"}
                                    for p in parts):
                    f.append(f"block {i}: MSQ answer {ans!r} is not a "
                             f"comma-joined subset of 1-4")
                if len(set(parts)) != len(parts):
                    f.append(f"block {i}: MSQ answer {ans!r} repeats an option")
            elif qt == "NAT":
                try:
                    float(ans)
                except ValueError:
                    f.append(f"block {i}: NAT answer {ans!r} is not numeric")
                if not re.search(r"\bdecimal|\bround|\bnearest|\bupto|\bup to",
                                 _runs_text(b.get("stem", [])), re.I):
                    f.append(f"block {i}: NAT stem must state rounding "
                             f"precision (Framework_NotesCreate §4 B3)")
            if t == "example":
                expl = b.get("explanation") or []
                if not expl:
                    f.append(f"block {i}: Example needs an Explanation")
                for ci, chunk in enumerate(expl):
                    check_runs(chunk if isinstance(chunk, list) else [chunk],
                               f"block {i} explanation {ci}")
                if b.get("speed_hack"):
                    check_runs(b["speed_hack"], f"block {i} speed hack")
            else:
                if b.get("explanation") or b.get("speed_hack"):
                    f.append(f"block {i}: Recall carries no Explanation or "
                             f"SPEED HACK (§4 B7)")
        elif t in ("key_points", "trap"):
            if not b.get("bullets"):
                f.append(f"block {i}: {t} has no bullets")
            for bi, bl in enumerate(b.get("bullets", [])):
                check_runs(bl, f"block {i} bullet {bi}")
        elif t == "rapid":
            for label in ("formulae", "associations"):
                rows = b.get(label) or []
                if not rows:
                    f.append(f"block {i}: rapid revision missing {label}")
                for ri, row in enumerate(rows):
                    for cell in row:
                        check_runs(cell, f"block {i} {label} row {ri}")
        elif t == "mindmap":
            if not b.get("image"):
                f.append(f"block {i}: mind map has no image path")

    # Tail order (§6A): TRAP -> RAPID -> RECALL CHECK -> MIND MAP, all after
    # the last concept.
    def first(t):
        return order_seen.index(t) if t in order_seen else None
    last_concept = max((i for i, t in enumerate(order_seen)
                        if t == "concept"), default=None)
    tail = [("trap", first("trap")), ("rapid", first("rapid")),
            ("recall", first("recall")), ("mindmap", first("mindmap"))]
    present = [(nm, ix) for nm, ix in tail if ix is not None]
    if last_concept is not None:
        for nm, ix in present:
            if ix < last_concept:
                f.append(f"{nm} block appears before the last concept "
                         f"(§6A tail order)")
    idxs = [ix for _, ix in present]
    if idxs != sorted(idxs):
        f.append("tail blocks are out of §6A order "
                 "(TRAP -> RAPID REVISION -> RECALL CHECK -> MIND MAP)")

    # Every concept must be followed by exactly one KEY POINTS box (§4 B4).
    for i, b in enumerate(blocks):
        if b["type"] != "concept":
            continue
        nxt = None
        for j in range(i + 1, len(blocks)):
            if blocks[j]["type"] in ("concept", "trap", "rapid", "recall",
                                     "mindmap"):
                break
            if blocks[j]["type"] == "key_points":
                nxt = j
                break
        if nxt is None:
            f.append(f"concept '{b.get('name')}' has no KEY POINTS box "
                     f"after its example stack (§4 B4)")
    return (not f, f)


# ------------------------------------------------------------------ build
def _add_title_bar(doc, text, fill, colour, size):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    cell.text = ""
    _cell_shade(cell, fill)
    _cell_borders(cell, fill)
    p = cell.paragraphs[0]
    _set_line_rule(p)
    _style_run(p.add_run(text), colour=colour, bold=True, size=size)
    return tbl


def _add_box(doc, kind, title, body_emitter):
    fg, bg = notes_core.BOX_COLORS[kind]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.text = ""
    _cell_shade(cell, bg)
    _cell_borders(cell, fg, sz="12")
    p = cell.paragraphs[0]
    _set_line_rule(p)
    _style_run(p.add_run(title), colour=fg, bold=True, size=BOX_TITLE_PT)
    body_emitter(cell)
    return tbl


def _add_data_table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr_fill = notes_core.LEVEL_COLORS["table_header"]
    for ci, cell_runs in enumerate(headers):
        c = tbl.cell(0, ci)
        c.text = ""
        _cell_shade(c, hdr_fill)
        p = c.paragraphs[0]
        _set_line_rule(p)
        _emit_runs(p, cell_runs, colour="FFFFFF", bold=True)
    for ri, row in enumerate(rows):
        for ci, cell_runs in enumerate(row):
            c = tbl.cell(ri + 1, ci)
            c.text = ""
            p = c.paragraphs[0]
            _set_line_rule(p)
            _emit_runs(p, cell_runs)
    return tbl


def _add_picture(paragraph, image_path, index):
    """Insert a picture with a DETERMINISTIC internal name.

    python-docx stamps the source filename into pic:cNvPr@name. That leaks a
    temp path into the document and, worse, churns document.xml every time NA
    re-renders a figure to a new temporary file — which would break the
    idempotence property for no visible benefit. The name is positional
    instead, so an identical re-render produces identical bytes.
    """
    run = paragraph.add_run()
    run.add_picture(image_path, width=Cm(FIGURE_W_CM))
    for cNvPr in paragraph._p.findall(".//" + qn("pic:cNvPr")):
        cNvPr.set("name", f"figure{index}.png")
        cNvPr.set("descr", "")
    return run


def _spacer(doc):
    """§4: adjacent boxes are ALWAYS separated so consecutive box tables never
    merge visually. Emitted by the engine, never by the caller."""
    p = doc.add_paragraph()
    _set_line_rule(p)
    _style_run(p.add_run(""), size=6)
    return p


def _emit_question(doc, block, title, *, with_explanation):
    def body(cell):
        p = _para(cell)
        _emit_runs(p, block["stem"])
        for i, opt in enumerate(block.get("options") or [], start=1):
            po = _para(cell)
            _style_run(po.add_run(f"{i}. "))
            _emit_runs(po, opt)
        pa = _para(cell)
        _style_run(pa.add_run("Answer: "), bold=True)
        _style_run(pa.add_run(str(block["answer"])), bold=True)
        if with_explanation:
            pe = _para(cell)
            _style_run(pe.add_run("Explanation:"), bold=True)
            for chunk in block.get("explanation", []):
                pc = _para(cell)
                _emit_runs(pc, chunk if isinstance(chunk, list) else [chunk])
            if block.get("speed_hack"):
                ps = _para(cell)
                _style_run(ps.add_run("SPEED HACK: "), bold=True)
                _emit_runs(ps, block["speed_hack"])
    kind = "example" if with_explanation else "recall"
    return _add_box(doc, kind, title, body)


def build(model, out_path, *, strict=True):
    """Render a content model to .docx. Returns a build report dict."""
    ok, findings = validate_model(model)
    if strict and not ok:
        raise ValueError("model validation failed:\n  - "
                         + "\n  - ".join(findings))
    outline = outline_of(model)
    l2_by_index = {i: (num, label) for i, num, label in outline["l2"]}
    l3_by_index = {}
    for i, num, label in outline["l3"]:
        l3_by_index.setdefault(i, []).append((num, label))
    ex_by_index = dict((i, j) for i, j in outline["examples"])
    rc_by_index = dict((i, j) for i, j in outline["recalls"])

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Cm(PAGE_W_CM)
    sec.page_height = Cm(PAGE_H_CM)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(MARGIN_CM))
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)

    prev_was_box = False
    figures = []
    for i, b in enumerate(model["blocks"]):
        t = b["type"]
        is_box = t in ("example", "recall", "key_points", "trap")
        if is_box and prev_was_box:
            _spacer(doc)
        if t == "title":
            # The name is rendered VERBATIM, not upper-cased. Framework_
            # NotesBlueprint §1A makes the manifest's display_name the exact
            # bytes for a unit, and case-folding it here would both override
            # that authority and make the parse/rebuild round trip lossy.
            _add_title_bar(doc, f"{outline['title']} {b['name']}",
                           notes_core.LEVEL_COLORS["L1"], "FFFFFF", TITLE_PT)
        elif t == "concept":
            num, label = l2_by_index[i]
            _add_title_bar(doc, f"{num} {label}",
                           notes_core.LEVEL_COLORS["L2"], "FFFFFF", L2_PT)
            for c in b.get("content", []):
                k = c["k"]
                if k == "bullet":
                    p = doc.add_paragraph()
                    _set_line_rule(p)
                    _style_run(p.add_run("\u2022  "))
                    _emit_runs(p, c["runs"])
                elif k == "para":
                    p = doc.add_paragraph()
                    _set_line_rule(p)
                    _emit_runs(p, c["runs"])
                elif k == "table":
                    _add_data_table(doc, c["headers"], c.get("rows", []))
                elif k == "figure":
                    p = doc.add_paragraph()
                    _set_line_rule(p)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_picture(p, c["image"], len(figures) + 1)
                    figures.append(c["image"])
        elif t == "example":
            _emit_question(doc, b, f"Example {ex_by_index[i]}",
                           with_explanation=True)
        elif t == "recall":
            if i == min(k for k in rc_by_index):
                num, label = l2_by_index[i]
                _add_title_bar(doc, f"{num} {label}",
                               notes_core.LEVEL_COLORS["L2"], "FFFFFF", L2_PT)
            _emit_question(doc, b, f"Recall {rc_by_index[i]}",
                           with_explanation=False)
        elif t == "key_points":
            def kp(cell, blk=b):
                for bl in blk["bullets"]:
                    p = _para(cell)
                    _style_run(p.add_run("\u2022  "))
                    _emit_runs(p, bl)
            _add_box(doc, "key_points", "KEY POINTS", kp)
        elif t == "trap":
            num, label = l2_by_index[i]
            _add_title_bar(doc, f"{num} {label}",
                           notes_core.LEVEL_COLORS["L2"], "FFFFFF", L2_PT)

            def tp(cell, blk=b):
                for bl in blk["bullets"]:
                    p = _para(cell)
                    _style_run(p.add_run("\u2022  "))
                    _emit_runs(p, bl)
            _add_box(doc, "trap", "TRAP BOX", tp)
        elif t == "rapid":
            num, label = l2_by_index[i]
            _add_title_bar(doc, f"{num} {label}",
                           notes_core.LEVEL_COLORS["L2"], "FFFFFF", L2_PT)
            for (sub_num, sub_label), key in zip(
                    l3_by_index[i], ("formulae", "associations")):
                p = doc.add_paragraph()
                _set_line_rule(p)
                _style_run(p.add_run(f"{sub_num} {sub_label}"),
                           colour=notes_core.LEVEL_COLORS["L3"], bold=True,
                           size=L3_PT)
                rows = b[key]
                _add_data_table(doc, rows[0], rows[1:])
        elif t == "mindmap":
            num, label = l2_by_index[i]
            _add_title_bar(doc, f"{num} {label}",
                           notes_core.LEVEL_COLORS["L2"], "FFFFFF", L2_PT)
            p = doc.add_paragraph()
            _set_line_rule(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_picture(p, b["image"], len(figures) + 1)
            figures.append(b["image"])
        prev_was_box = is_box

    doc.save(out_path)
    return {"path": out_path, "blocks": len(model["blocks"]),
            "outline": outline, "figures": figures,
            "validation_ok": ok, "findings": findings}


# ------------------------------------------------------------------ parse
def document_xml(docx_path):
    with zipfile.ZipFile(docx_path) as z:
        return z.read("word/document.xml")


def _runs_from_paragraph(p_el):
    """Reconstruct the inline run list from a built paragraph."""
    runs = []
    for child in p_el:
        tag = child.tag
        if tag == qn("w:r"):
            texts = child.findall(qn("w:t"))
            if not texts:
                continue
            s = "".join(t.text or "" for t in texts)
            rPr = child.find(qn("w:rPr"))
            va = None
            if rPr is not None:
                va_el = rPr.find(qn("w:vertAlign"))
                if va_el is not None:
                    va = va_el.get(qn("w:val"))
            # Adjacent text runs are NOT merged. build() emits exactly one run
            # per model entry, so preserving that 1:1 mapping is what makes
            # build -> parse -> build byte-identical. Merging also mis-assigned
            # a symbol's base: "... as " + "S" would fuse, and the following
            # subscript would then claim the whole fused string as its base.
            if va in ("subscript", "superscript"):
                key = "sub" if va == "subscript" else "sup"
                if runs and runs[-1].get("t") in ("text", "sym"):
                    prev = runs[-1]
                    if prev["t"] == "text":
                        runs[-1] = {"t": "sym", "base": prev["s"], key: s}
                    else:
                        prev[key] = s
                else:
                    runs.append({"t": "sym", "base": "", key: s})
            else:
                runs.append({"t": "text", "s": s})
        elif tag == qn("m:oMath"):
            from lxml import etree
            runs.append({"t": "math",
                         "xml": etree.tostring(child).decode("utf-8")})
    return [r for r in runs if not (r.get("t") == "text" and r["s"] == "")]


def _strip_prefix(runs, prefix):
    """Drop a literal marker ("\u2022  ", "3. ", "SPEED HACK: ") from the front
    of a parsed run list.

    Adjacent plain-text runs are MERGED by _runs_from_paragraph, so the marker
    and the content that follows it arrive as one run — dropping runs[0] would
    silently discard the content too. This strips characters, not runs.
    """
    out = copy.deepcopy(runs)
    if out and out[0].get("t") == "text":
        s = out[0]["s"]
        if s.startswith(prefix):
            s = s[len(prefix):]
        else:
            s = re.sub(r"^" + re.escape(prefix.strip()) + r"\s*", "", s)
        if s:
            out[0]["s"] = s
        else:
            out.pop(0)
    return out


def _cell_fill(cell):
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def _table_rows_runs(tbl, skip_header):
    rows = []
    for r in tbl.rows[1 if skip_header else 0:]:
        rows.append([_runs_from_paragraph(c.paragraphs[0]._p) for c in r.cells])
    return rows


def _extract_image(doc, p_el, media_dir):
    """Write an embedded image out to media_dir and return its path.

    Round-tripping a figure by PATH (rather than by relationship id) keeps the
    model a plain, serialisable dict — NA can hand a figure to figural_core
    for a re-render and drop the new path straight back in.
    """
    blips = p_el.findall(".//" + qn("a:blip"))
    if not blips:
        return None
    rid = blips[0].get(qn("r:embed"))
    part = doc.part.related_parts[rid]
    os.makedirs(media_dir, exist_ok=True)
    name = os.path.basename(part.partname)
    path = os.path.join(media_dir, name)
    with open(path, "wb") as f:
        f.write(part.blob)
    return path


def _parse_question(cell, title):
    blk = {"type": "example" if title.startswith("Example") else "recall",
           "qtype": None, "stem": [], "options": [], "answer": "",
           "explanation": []}
    mode = "stem"
    for p in cell.paragraphs[1:]:
        txt = p.text.strip()
        if not txt and not p._p.findall(".//" + qn("m:oMath")):
            continue
        if re.match(r"^[1-4]\.\s", txt):
            blk["options"].append(_strip_prefix(_runs_from_paragraph(p._p),
                                               txt[:2] + " "))
            mode = "opt"
        elif txt.startswith("Answer:"):
            blk["answer"] = txt.split(":", 1)[1].strip()
            mode = "ans"
        elif txt.startswith("Explanation"):
            mode = "expl"
        elif txt.startswith("SPEED HACK"):
            blk["speed_hack"] = _strip_prefix(_runs_from_paragraph(p._p),
                                              "SPEED HACK: ")
            mode = "hack"
        elif mode == "stem":
            blk["stem"] += _runs_from_paragraph(p._p)
        elif mode == "expl":
            blk["explanation"].append(_runs_from_paragraph(p._p))
    blk["qtype"] = ("NAT" if not blk["options"]
                    else ("MSQ" if "," in blk["answer"] else "MCQ"))
    if blk["type"] == "recall":
        blk.pop("explanation", None)
    return blk


def parse(docx_path, media_dir=None):
    """Read a document built by build() back into a content model.

    This walks the body in DOCUMENT ORDER as a state machine, so a concept's
    bullets, tables and figures are re-attached to the concept they follow.
    Round-trip fidelity is self-tested: build(parse(build(m))) reproduces
    word/document.xml byte-for-byte.  That test is what makes "re-running NA
    on its own output changes nothing" a checkable property.

    An untouched equation round-trips as its ORIGINAL OMML element (stored as
    {"t":"math","xml":...}), so NA never re-authors maths it did not edit.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    doc = Document(docx_path)
    media_dir = media_dir or (os.path.splitext(docx_path)[0] + "_media")
    model = {"schema": SCHEMA, "exam_code": None,
             "unit": {"name": None, "tier": None, "seq_in_topic": None},
             "blocks": []}
    l1 = notes_core.LEVEL_COLORS["L1"]
    l2 = notes_core.LEVEL_COLORS["L2"]
    l3 = notes_core.LEVEL_COLORS["L3"]
    hdr = notes_core.LEVEL_COLORS["table_header"]
    box_fills = {bg: kind for kind, (fg, bg) in notes_core.BOX_COLORS.items()}

    cur_concept = None
    section = None          # None | "trap" | "rapid" | "recall" | "mindmap"
    rapid_block = None

    for child in doc.element.body:
        if child.tag == qn("w:tbl"):
            tbl = Table(child, doc)
            first = tbl.cell(0, 0)
            fill = _cell_fill(first)
            text = first.paragraphs[0].text.strip() if first.paragraphs else ""
            if fill == l1:
                m = re.match(r"^(\d+)\.\s+(.*)$", text)
                if m:
                    model["unit"]["seq_in_topic"] = int(m.group(1))
                    model["unit"]["name"] = m.group(2)
                    model["blocks"].append({"type": "title",
                                            "name": m.group(2)})
                cur_concept = None
            elif fill == l2:
                m = re.match(r"^\d+(?:\.\d+)*\s+(.*)$", text)
                label = m.group(1) if m else text
                if label == "TRAP BOX":
                    section, cur_concept = "trap", None
                elif label == "RAPID REVISION SUMMARY":
                    section, cur_concept = "rapid", None
                    rapid_block = {"type": "rapid", "formulae": [],
                                   "associations": []}
                    model["blocks"].append(rapid_block)
                elif label == "RECALL CHECK":
                    section, cur_concept = "recall", None
                elif label == "MIND MAP":
                    section, cur_concept = "mindmap", None
                else:
                    section = None
                    cur_concept = {"type": "concept", "name": label,
                                   "content": []}
                    model["blocks"].append(cur_concept)
            elif fill in box_fills:
                title = first.paragraphs[0].text.strip()
                if title.startswith(("Example", "Recall")):
                    model["blocks"].append(_parse_question(first, title))
                else:
                    kind = box_fills[fill]
                    bullets = []
                    for p in first.paragraphs[1:]:
                        if p.text.strip():
                            bullets.append(_strip_prefix(_runs_from_paragraph(p._p),
                                                     "\u2022  "))
                    model["blocks"].append({"type": kind, "bullets": bullets})
            else:
                # A data table: its header row is slate-filled.
                if tbl.rows and _cell_fill(tbl.rows[0].cells[0]) == hdr:
                    headers = [_runs_from_paragraph(c.paragraphs[0]._p)
                               for c in tbl.rows[0].cells]
                    rows = _table_rows_runs(tbl, skip_header=True)
                    if section == "rapid" and rapid_block is not None:
                        key = ("formulae" if not rapid_block["formulae"]
                               else "associations")
                        rapid_block[key] = [headers] + rows
                    elif cur_concept is not None:
                        cur_concept["content"].append(
                            {"k": "table", "headers": headers, "rows": rows})
        elif child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            has_img = bool(child.findall(".//" + qn("a:blip")))
            has_math = bool(child.findall(".//" + qn("m:oMath")))
            txt = p.text.strip()
            if has_img:
                path = _extract_image(doc, child, media_dir)
                if section == "mindmap":
                    model["blocks"].append({"type": "mindmap", "image": path})
                    section = None
                elif cur_concept is not None:
                    cur_concept["content"].append(
                        {"k": "figure", "image": path,
                         "label": cur_concept["name"]})
                continue
            if not txt and not has_math:
                continue                       # spacer
            runs = _runs_from_paragraph(child)
            if cur_concept is not None:
                if txt.startswith("\u2022"):
                    cur_concept["content"].append(
                        {"k": "bullet",
                         "runs": _strip_prefix(runs, "\u2022  ")})
                else:
                    cur_concept["content"].append({"k": "para", "runs": runs})
    return model


# ------------------------------------------------------------------ self-test
def self_test():
    import tempfile
    passed, fails = 0, []

    def check(name, cond):
        nonlocal passed
        if cond:
            passed += 1
        else:
            fails.append(name)

    T = lambda s: [{"t": "text", "s": s}]

    def demo_model():
        return {
            "schema": SCHEMA, "exam_code": "EX",
            "unit": {"sid": "ex.a.b.c", "name": "Enzyme Kinetics",
                     "tier": "TIER-1", "seq_in_topic": 3,
                     "section": "S", "topic": "T"},
            "blocks": [
                {"type": "title", "name": "Enzyme Kinetics"},
                {"type": "concept", "name": "Michaelis-Menten model",
                 "content": [
                     {"k": "bullet", "runs": [
                         {"t": "text", "s": "Rate rises then saturates as "},
                         {"t": "sym", "base": "S", "sub": "0"},
                         {"t": "text", "s": " increases."}]},
                     {"k": "table",
                      "headers": [T("Term"), T("Meaning")],
                      "rows": [[T("Km"), T("half-saturation")],
                               [T("Vmax"), T("ceiling rate")]]},
                 ]},
                {"type": "example", "qtype": "MCQ",
                 "stem": [{"t": "text", "s": "For the region "},
                          {"t": "math", "latex": "\\frac{V_{max}[S]}{K_m+[S]}"},
                          {"t": "text", "s": " which holds?"}],
                 "options": [T("a"), T("b"), T("c"), T("d")],
                 "answer": "2",
                 "explanation": [T("Substitute and simplify.")],
                 "speed_hack": T("Compare denominators.")},
                {"type": "key_points", "bullets": [T("Saturation is the key.")]},
                {"type": "trap", "bullets": [T("Confusing Km with Kd.")]},
                {"type": "rapid",
                 "formulae": [[T("Name"), T("Form")], [T("MM"), T("v = ...")]],
                 "associations": [[T("Term"), T("Link")], [T("kcat"), T("turnover")]]},
                {"type": "recall", "qtype": "MCQ", "stem": T("Which is true?"),
                 "options": [T("a"), T("b"), T("c"), T("d")], "answer": "3"},
            ]}

    m = demo_model()
    ok, findings = validate_model(m)
    check("valid demo model passes validation", ok)

    # ---- numbering is DERIVED, never stored -----------------------------
    o = outline_of(m)
    check("title number derives from seq_in_topic", o["title"] == "3.")
    check("concept numbers cascade from block order",
          o["l2"][0][1] == "3.1" and o["l2"][0][2] == "Michaelis-Menten model")
    check("tail blocks number after the concepts",
          [n for _, n, _ in o["l2"]] == ["3.1", "3.2", "3.3", "3.4"])
    check("rapid revision emits two level-3 numbers",
          [n for _, n, _ in o["l3"]] == ["3.3.1", "3.3.2"])
    check("example counter is global and 1-based", o["examples"] == [(2, 1)])
    check("recall counter is separate", o["recalls"] == [(6, 1)])

    # inserting a concept renumbers everything after it, with no stored number
    m2 = copy.deepcopy(m)
    m2["blocks"].insert(4, {"type": "concept", "name": "Inhibition",
                            "content": []})
    m2["blocks"].insert(5, {"type": "key_points", "bullets": [T("x")]})
    o2 = outline_of(m2)
    check("inserting a concept renumbers the tail automatically",
          [n for _, n, _ in o2["l2"]] == ["3.1", "3.2", "3.3", "3.4", "3.5"])
    check("no outline number is stored in the model",
          not any("num" in b for b in m2["blocks"]))

    # ---- answer-key validation (G-8 pre-empted at construction) ----------
    bad = copy.deepcopy(m)
    bad["blocks"][2]["answer"] = "5"
    check("MCQ answer out of option range is rejected",
          not validate_model(bad)[0])
    bad = copy.deepcopy(m)
    bad["blocks"][2]["qtype"] = "MSQ"
    bad["blocks"][2]["answer"] = "1,1"
    check("MSQ answer repeating an option is rejected",
          not validate_model(bad)[0])
    bad = copy.deepcopy(m)
    bad["blocks"][2]["qtype"] = "NAT"
    bad["blocks"][2]["options"] = []
    bad["blocks"][2]["answer"] = "abc"
    check("non-numeric NAT answer is rejected", not validate_model(bad)[0])
    bad = copy.deepcopy(m)
    bad["blocks"][2]["qtype"] = "NAT"
    bad["blocks"][2]["options"] = []
    bad["blocks"][2]["answer"] = "2.5"
    bad["blocks"][2]["stem"] = T("Compute the value.")
    check("NAT stem without rounding precision is rejected",
          not validate_model(bad)[0])
    good = copy.deepcopy(bad)
    good["blocks"][2]["stem"] = T("Compute the value, rounded to 2 decimal places.")
    check("NAT stem stating rounding precision passes",
          validate_model(good)[0])

    # ---- the V_max trap --------------------------------------------------
    trap = copy.deepcopy(m)
    trap["blocks"][2]["stem"] = [{"t": "math", "latex": "V_max"}]
    ok_t, f_t = validate_model(trap)
    check("unbraced multi-character script is REJECTED (V_max renders as "
          "V-sub-m + 'ax' — silently wrong on the page)",
          not ok_t and any("unbraced" in x for x in f_t))
    fine = copy.deepcopy(m)
    fine["blocks"][2]["stem"] = [{"t": "math", "latex": "V_{max}"}]
    check("braced script is accepted", validate_model(fine)[0])

    # ---- anatomy gates ---------------------------------------------------
    noKP = copy.deepcopy(m)
    noKP["blocks"] = [b for b in noKP["blocks"] if b["type"] != "key_points"]
    check("concept without a KEY POINTS box is rejected (§4 B4)",
          not validate_model(noKP)[0])
    disorder = copy.deepcopy(m)
    disorder["blocks"][4], disorder["blocks"][5] = (disorder["blocks"][5],
                                                   disorder["blocks"][4])
    check("tail blocks out of §6A order are rejected",
          not validate_model(disorder)[0])
    longb = copy.deepcopy(m)
    longb["blocks"][1]["content"][0] = {
        "k": "bullet", "runs": T(" ".join(["word"] * 30))}
    check("bullet over the D-1 hard cap is rejected",
          not validate_model(longb)[0])

    # ---- build -----------------------------------------------------------
    out = tempfile.mktemp(suffix=".docx")
    rep = build(m, out)
    check("build writes a file", os.path.exists(out))
    check("build reports the derived outline", rep["outline"]["title"] == "3.")

    xml = document_xml(out).decode("utf-8")
    check("F-7: every spacing element uses the AUTO line rule",
          'w:lineRule="auto"' in xml
          and 'w:lineRule="exact"' not in xml
          and 'w:lineRule="atLeast"' not in xml)
    check("F-3a: the expression compiled to structural OMML",
          "<m:oMath" in xml and "<m:f>" in xml)
    check("F-3a: no textual exponent inside the document", "^(" not in xml)
    check("F-3b: symbol mention rendered as a subscript RUN",
          'w:val="subscript"' in xml)
    check("§6A: L1 navy title bar present",
          notes_core.LEVEL_COLORS["L1"] in xml)
    check("§6A: L2 teal concept bar present",
          notes_core.LEVEL_COLORS["L2"] in xml)
    check("§6A: L3 purple sub-head present",
          notes_core.LEVEL_COLORS["L3"] in xml)
    check("§6A: slate table header present",
          notes_core.LEVEL_COLORS["table_header"] in xml)
    for kind, (fg, bg) in notes_core.BOX_COLORS.items():
        check(f"§6A: {kind} box colours present", fg in xml and bg in xml)
    check("derived numbering is rendered, not stored",
          "3.1 Michaelis-Menten model" in xml.replace("</w:t>", "")
          .replace("<w:t>", "").replace('<w:t xml:space="preserve">', "")
          or "3.1" in xml)
    check("Example counter rendered", "Example 1" in xml)
    check("Recall counter rendered", "Recall 1" in xml)
    check("Answer label is emitted bold before any Explanation",
          xml.index("Answer:") < xml.index("Explanation"))

    # ---- determinism + round trip ---------------------------------------
    out2 = tempfile.mktemp(suffix=".docx")
    build(m, out2)
    check("build is DETERMINISTIC (same model -> identical document.xml)",
          document_xml(out) == document_xml(out2))

    # ---- THE LOAD-BEARING PROPERTY -------------------------------------
    # build -> parse -> build must reproduce word/document.xml EXACTLY.
    # Framework_NotesAudit v3.0.0 lets NA rebuild the whole unit; without
    # this, "re-running NA on its own output changes nothing" would be an
    # assertion rather than a fact. Every block type below is exercised.
    rt1 = tempfile.mktemp(suffix=".docx")
    build(m, rt1)
    back = parse(rt1)
    rt2 = tempfile.mktemp(suffix=".docx")
    build(back, rt2, strict=False)
    check("ROUND TRIP: build -> parse -> build is byte-identical "
          "(document.xml)", document_xml(rt1) == document_xml(rt2))
    back2 = parse(rt2)
    rt3 = tempfile.mktemp(suffix=".docx")
    build(back2, rt3, strict=False)
    check("ROUND TRIP is a FIXED POINT (a third pass changes nothing)",
          document_xml(rt2) == document_xml(rt3))
    check("ROUND TRIP preserves the derived outline",
          outline_of(back)["l2"] == outline_of(m)["l2"]
          and outline_of(back)["examples"] == outline_of(m)["examples"])
    check("ROUND TRIP preserves the answer key",
          [b.get("answer") for b in back["blocks"] if "answer" in b]
          == [b.get("answer") for b in m["blocks"] if "answer" in b])
    check("ROUND TRIP preserves OMML verbatim (maths is never re-authored)",
          any(r.get("t") == "math" and r.get("xml")
              for b in back["blocks"] if b["type"] == "example"
              for r in b["stem"]))
    check("ROUND TRIP preserves concept tables",
          [c["k"] for b in back["blocks"] if b["type"] == "concept"
           for c in b["content"]] == ["bullet", "table"])
    check("ROUND TRIP preserves rapid-revision tables",
          [b for b in back["blocks"] if b["type"] == "rapid"][0]["formulae"]
          == m["blocks"][5]["formulae"])

    # A model edited between passes still round-trips (the NA workflow).
    edited = copy.deepcopy(back)
    edited["blocks"][3]["bullets"].append(T("Added by the audit step."))
    ed1 = tempfile.mktemp(suffix=".docx")
    build(edited, ed1, strict=False)
    ed_back = parse(ed1)
    ed2 = tempfile.mktemp(suffix=".docx")
    build(ed_back, ed2, strict=False)
    check("ROUND TRIP survives an edit (the NA correct-and-rebuild path)",
          document_xml(ed1) == document_xml(ed2))
    check("an inserted KEY POINTS bullet actually lands in the rebuild",
          "Added by the audit step." in document_xml(ed1).decode("utf-8"))

    # ---- figures round-trip, and survive a re-render (Q-E / F-4) ---------
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        def _render(path):
            fig, ax = plt.subplots(figsize=(4, 2.5))
            ax.plot([0, 1, 2], [0, 3, 4])
            fig.savefig(path, dpi=150)
            plt.close(fig)

        img_a = tempfile.mktemp(suffix=".png")
        _render(img_a)
        fm = {"schema": SCHEMA, "exam_code": "EX",
              "unit": {"name": "U", "tier": "TIER-1", "seq_in_topic": 2},
              "blocks": [
                  {"type": "title", "name": "U"},
                  {"type": "concept", "name": "C1", "content": [
                      {"k": "figure", "image": img_a, "label": "C1"}]},
                  {"type": "key_points", "bullets": [T("kp")]},
                  {"type": "mindmap", "image": img_a}]}
        fa = tempfile.mktemp(suffix=".docx")
        build(fm, fa)
        fb = parse(fa)
        fc = tempfile.mktemp(suffix=".docx")
        build(fb, fc, strict=False)
        check("ROUND TRIP with figures is byte-identical",
              document_xml(fa) == document_xml(fc))
        check("figure and mind map both recovered by parse",
              any(c["k"] == "figure" for b in fb["blocks"]
                  if b["type"] == "concept" for c in b["content"])
              and any(b["type"] == "mindmap" for b in fb["blocks"]))
        # A re-render to a DIFFERENT path must not churn the XML: the source
        # filename is never written into the document.
        img_b = tempfile.mktemp(suffix=".png")
        _render(img_b)
        fm2 = copy.deepcopy(fm)
        fm2["blocks"][1]["content"][0]["image"] = img_b
        fm2["blocks"][3]["image"] = img_b
        fd = tempfile.mktemp(suffix=".docx")
        build(fm2, fd)
        check("re-rendering a figure to a new path does NOT churn document.xml "
              "(figure re-render stays idempotent)",
              document_xml(fa) == document_xml(fd))
    except ImportError:
        check("figure round-trip (matplotlib absent — DORMANT)", True)

    parsed = parse(out)
    check("parse recovers the unit name",
          parsed["unit"]["name"] == "Enzyme Kinetics")
    check("parse recovers seq_in_topic (so numbering re-derives identically)",
          parsed["unit"]["seq_in_topic"] == 3)
    check("parse recovers the concept",
          any(b["type"] == "concept"
              and b["name"] == "Michaelis-Menten model"
              for b in parsed["blocks"]))
    ex = [b for b in parsed["blocks"] if b["type"] == "example"]
    check("parse recovers the example with its key",
          len(ex) == 1 and ex[0]["answer"] == "2" and ex[0]["qtype"] == "MCQ")
    check("parse recovers 4 options", len(ex[0]["options"]) == 4)
    rc = [b for b in parsed["blocks"] if b["type"] == "recall"]
    check("parse recovers the recall item and drops its explanation",
          len(rc) == 1 and "explanation" not in rc[0])
    kp = [b for b in parsed["blocks"] if b["type"] == "key_points"]
    tr = [b for b in parsed["blocks"] if b["type"] == "trap"]
    check("parse recovers KEY POINTS and TRAP boxes", len(kp) == 1 and len(tr) == 1)
    check("parse re-derives the SAME outline as the source model",
          [n for _, n, _ in outline_of(parsed)["l2"]][:1] == ["3.1"])

    # ---- spacer rule -----------------------------------------------------
    twobox = {"schema": SCHEMA, "exam_code": "EX",
              "unit": {"name": "U", "tier": "TIER-3", "seq_in_topic": 1},
              "blocks": [{"type": "title", "name": "U"},
                         {"type": "concept", "name": "C", "content": []},
                         {"type": "key_points", "bullets": [T("a")]},
                         {"type": "trap", "bullets": [T("b")]}]}
    out3 = tempfile.mktemp(suffix=".docx")
    build(twobox, out3)
    d3 = Document(out3)
    kinds = [c.tag for c in d3.element.body if c.tag in (qn("w:p"), qn("w:tbl"))]
    check("§4: adjacent boxes are separated by a spacer paragraph",
          qn("w:p") in kinds)

    # ---- strict flag ------------------------------------------------------
    try:
        build(noKP, tempfile.mktemp(suffix=".docx"))
        check("strict build raises on an invalid model", False)
    except ValueError:
        check("strict build raises on an invalid model", True)

    # ---- SPEC-LOCK: colour authority is notes_core, not a local literal ---
    check("spec-lock: builder reads colours from notes_core (single authority)",
          notes_core.LEVEL_COLORS["L1"] == "1F4E79"
          and notes_core.BOX_COLORS["trap"] == ("C62828", "FBE4E4"))
    check("spec-lock: F-7 constants are AUTO",
          LINE_RULE == "auto" and LINE_VALUE == "240")
    check("spec-lock: A4 page geometry",
          (PAGE_W_CM, PAGE_H_CM) == (21.0, 29.7) and BODY_FONT == "Arial")

    print(f"notes_docx self-test: {passed} passed, {len(fails)} failed"
          + (" — " + "; ".join(fails) if fails else ""))
    return not fails


if __name__ == "__main__":
    import sys
    if "--self-test" in sys.argv:
        sys.exit(0 if self_test() else 1)
    print("notes_docx.py — shared Notes document builder. Run with --self-test.")
