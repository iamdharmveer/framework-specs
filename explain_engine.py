#!/usr/bin/env python3
# explain_engine.py — UNIVERSAL, EXAM-AGNOSTIC explanation engine for Step 9
# (MockExplain). The ONLY path by which an explanation may enter the docx.
#
# Zero hardcoded exam values: question/option regex, option count, section
# labels, language and block-header markers are all supplied at runtime by an
# EngineConfig built from blueprint.json + section_rules.md (CATEGORY C). The
# same engine drives SSC, GATE, NEET, IBPS, UPSC, CAT and any exam with valid
# Step 5/6/7/8 outputs.
#
# Public surface (called by the Step-4 spec):
#   EngineConfig                       — runtime exam parameters (no defaults baked as exam facts)
#   ExplanationBlock(...).validate()   — validated container; raises on any structural breach
#   add_math_text(p, text)             — THE prose emitter; auto-OMMLs digit/digit; raises on bad math
#   add_math(p, omath_xml) + frac/sup/sqrt/nary/omath/_r  — explicit OMML route
#   parse_paper(path, cfg)             — read the Step-3 paper → question map (P3 validation)
#   build_interleaved_docx(src, blocks, out, cfg)  — seed WHOLE source, append blocks (append-only)
#   verify_fidelity(out, src, cfg)     — question region byte-identical to source; rIds resolve (every batch)
#   verify_structure(out, blocks, cfg) — block model + coverage + CA-binding (every batch)
#   verify_explanations(out, blocks, cfg) — INDEPENDENT post-render re-audit of the rendered docx (every batch); v1.17 adds NAT Correct-Answer portal-charset validation (RXA-CHARSET)
#   strip_solutions(out, stripped, cfg)— questions-only copy for the Step-2 re-audit
#   self_test()                        — N/N PASS gate; run with --self-test
#   parse_solution_blocks(path, cfg)   — [STEP 5] read a Solutions docx back into blocks (inverse of build)
#   self_test_audit()                  — [STEP 5] reader round-trip gate; run with --self-test-audit
#   parse_learnings(path)              — [STEP 4] read an EXPLAIN(_AUDIT)_LEARNINGS md into structured rules (P10)
#   build_report_docx(out, title, sections) / read_report_docx(path) — [v2.9] §20 END-OF-MOCK REPORT as a docx
#
# This file is embedded verbatim in Appendix A of Framework_MockTestExplain.md.
# It is the canonical copy; never patch it by hand — regenerate from the spec.
#
# v2.9 — 2026-08-26 — GAP-2026-08-26-REGISTRY-HANDOFF-SEAM (paired with
#   MockTestExplain v1.46.0 / MockTestCreate v5.73 / MockDeliver v1.16.0 /
#   DeliveryFooter v1.27 / paper_pipeline v5.74 Cluster RH / final_assembly
#   v5.60). ADDITIVE: build_report_docx / read_report_docx write and read the
#   Step-9 END-OF-MOCK REPORT (§20) as [ExamCode]_[slug]_Explain_Report.docx,
#   delivered at the final batch beside the Explanation docx and the updated
#   registry. MANDATE 0 is enforced at the writer (an explicit answer
#   declaration raises before any byte is written). No existing gate, block
#   field or docx output changes; Step 11 never reads the report.
# v2.8 — 2026-08-21 — GAP-2026-08-21-EXPLANATION-PROVENANCE (paired with
#   MockTestExplain v1.37.0 / PYQExplain v2.15 / MockTestCreate v5.59 /
#   paper_pipeline v5.39 / final_assembly v5.55). A delivered 60-question paper
#   passed every v2.7 gate and still (1) PUBLISHED A WRONG KEY on a figural
#   item — the structure was misread, both derive-twice routes shared the
#   misread, and the step that CREATED the question (which held the opposite
#   key) was never consulted; (2) carried 24 hedged WHY WRONG / COMMON
#   PITFALLS lines on 15 questions ("or otherwise …", "perhaps by …", "or a
#   similar …") of which at least nine asserted arithmetic that is FALSE —
#   §15-2's "a real path always exists" forced a path to be invented where
#   none existed; (3) reported 0 NARROWED transfer claims across 60 AXIOMs
#   while the loaded subject library already named the families that were
#   overgeneralised — the library was loaded and never consulted, because the
#   lookup was by discipline; (4) rendered every chemical formula in prose as
#   ASCII ("[V(CO)6]-", "sp3d", "pi*") — zero sub/superscript runs in 71
#   pages — because nothing gated formula typography. What the engine can
#   prove is now gated here:
#   (a) ERROR PROVENANCE (§15-2 rewritten). Every WHY WRONG option and every
#       COMMON PITFALL value carries an error_provenance record: mode
#       VERIFIED_ERROR_PATH (the wrong operation is RECOMPUTED by the engine
#       from an arithmetic expression and must reproduce the target at the
#       target's own precision) or DIRECT_CONTRADICTION (no path is claimed;
#       the line states why the option contradicts the correct relation). A
#       claimed path that does not reproduce its target raises
#       DST_UNVERIFIED_NUMERICAL_ORIGIN; a missing record raises
#       DST_NO_PROVENANCE. HEDGED provenance language is banned in those two
#       sections (DST_HEDGED_PROVENANCE) — measured on the reference paper:
#       24 hits, 0 false positives. No pitfall quota: >=1 stays, a second is
#       never required.
#   (b) transfer_record is MANDATORY for an authored block (a read-back
#       block, _preserved, is exempt as it always was). Each entry may name
#       neighbour_source CURATED:<rule-code> | GENERATED. EngineConfig carries
#       learnings_triggers (built from the loaded learnings' **Triggers:**
#       field by triggers_from_learnings); when a trigger fires on an AXIOM /
#       SPEED HACK sentence and no entry for that section cites that rule,
#       construction raises GEN_CANONICAL_EXCEPTION_MISSED. A GENERATED
#       neighbour is legal only where no trigger fires. transfer_tripwire()
#       reports a run whose AXIOM claims were all SAFE (the self-attestation
#       shape) so the spec can demand the recorded second pass.
#   (c) FORMULA TYPOGRAPHY. normalise_formula_text() rewrites ASCII chemical
#       / orbital notation to Unicode sub- and superscripts at construction
#       (⟦MATH:⟧ regions and ⟦M:⟧ tokens are never touched; single
#       letter+digit locants such as C2, C3 are deliberately left alone); the
#       residual gate raises FMT_UNFORMATTED_FORMULA on what the normaliser
#       could not safely rewrite. Per-exam switch: EngineConfig(
#       formula_typography=False).
#   (d) reconcile_key_commitments(): Step 7 (MockTestCreate v5.59) commits a
#       salted hash of every canonical answer into registry.key_commitments;
#       Step 9 recomputes from its OWN derived answers and compares. Step 9
#       never sees a plaintext key; a mismatch is a KEY_CONFLICT that the spec
#       resolves IN-RUN (§17 v1.37.0), never a halt.
#   (e) scan_risk_markers() / --scan-risk: marks an existing Explanation docx
#       by provenance-hedge, absolute-term and unformatted-formula markers so
#       earlier papers can be queued for regeneration (never regex-patched).
#   Switches: EngineConfig(provenance_gates=False) is for LEGACY fixtures and
#   read-back only; a production EngineConfig never sets it. NO existing gate
#   is weakened; every new gate ships with fixtures below.
# v2.7 — 2026-08-20 — GAP-2026-08-20-TRANSFER-SAFE-EXPLANATIONS (paired with
#   MockTestExplain v1.36.0 / PYQExplain v2.14). A delivered 60-question paper
#   was answer-correct on every item and still carried ~17 sentences a learner
#   could memorise and be failed by on the next related item: AXIOMs stating a
#   class-level rule the question's own neighbours break, SPEED HACKs that work
#   only for the options shown, unjustified absolutes, and learner-psychology
#   boilerplate mandated by the old §15-3 MSQ wording. Root cause: §8-0b and
#   §14-3b ADVISED what only a gate can enforce. What the engine can prove is
#   now gated here; the rest is the §7-7 protocol with a recorded artefact.
#   (a) ABSOLUTE-TERM GATE — an unqualified universal (always/never/cannot/
#       at all/regardless of/impossible/…) in AXIOM, SPEED HACK, WHY WRONG or COMMON
#       PITFALLS raises at construction unless the block DECLARES that sentence in
#       absolutes_justified with a reason (definition / conservation law /
#       mathematical property / symmetry-forbidden). Plain quantifiers ('only
#       two ions', 'every formula unit', 'exactly 208') are NOT gated — measured
#       on a delivered paper they were 80 percent false positives.
#       Configurable per language via EngineConfig(absolute_terms_re=...).
#   (b) LEARNER-PSYCHOLOGY TEMPLATES banned in student text ('the seductive half',
#       'the student thinks', 'a hasty solver', 'invites the reading', …): a
#       distractor is refuted chemically/mathematically, never by guessing at
#       the learner's mind (§15-3 rewritten; §9 DST).
#   (c) transfer_record — OPTIONAL per-block metadata (the §7-7 artefact): one
#       entry per transferable claim {section, claim, epistemic_type, scope,
#       neighbour_tested, outcome}. Shape-validated when supplied: every AXIOM
#       needs >=1 AXIOM entry, a SPEED HACK needs >=1 SPEED_HACK entry, a
#       QUESTION_SPECIFIC claim may never sit in AXIOM, an OPTION_SET_SHORTCUT
#       may sit only in SPEED_HACK, every claim names its neighbour and outcome.
#   (d) CONFORMER added to the §6A-2 vocabulary (visual verdict; requires a
#       figure, exactly like STRUCTURE_GRAPH) — closes Step-9 run-report F3.
#   (e) strip_solutions now garbage-collects the media parts its removed figure
#       paragraphs referenced, so the stripped copy is byte-comparable, not
#       merely text-comparable, to the Step-7 source (run-report F2).
#   NO existing guard is weakened; every new gate ships with fixtures below.

import re, sys, hashlib, base64
from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from lxml import etree

# ───────────────────────────── namespaces ──────────────────────────────────
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

# ───────────────────────────── exam config ─────────────────────────────────
def _int_to_roman(n):
    out, vals = '', [(10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')]
    for v, s in vals:
        while n >= v:
            out += s; n -= v
    return out

class EngineConfig:
    DEFAULT_PROVENANCE_GATES = True   # v2.8 — flipped ONLY inside self_test for legacy fixtures
    """Runtime exam parameters. NOTHING here is an exam fact baked into code —
    every value is passed in by the spec from blueprint.json / section_rules.md.
    The constructor defaults are STRUCTURAL fallbacks (English labels, 4 numeric
    options, single-correct MCQ, Latin-script terminators); the spec always passes
    the real values read at runtime (RE-9). Supports, per question, any of three
    question types — MCQ (one correct option), MSQ (a set of correct options) and
    NAT (a typed numerical answer, no options) — and per-section option counts,
    alphabetic/roman/custom option labels, and language-specific sentence
    terminators (e.g. the Devanagari danda)."""
    def __init__(self, q_pattern, opt_pattern, options_count,
                 labels=None, markers=None, colors=None, language='en',
                 label_scheme='numeric', sentence_terminators='.!?',
                 options_by_q=None,
                 banned_blocks=None, banned_templates=None,
                 banned_fakecites=None, metacommentary_re=None,
                 absolute_terms_re=None,
                 provenance_gates=None, learnings_triggers=None,
                 formula_typography=True):
        # v2.8 — provenance_gates: the authoring-time mandatory gates (error
        # provenance, transfer_record, formula residual). None -> the module
        # default (True). Only legacy self-test fixtures and read-back paths
        # turn it off; a production config never does.
        self.provenance_gates = (EngineConfig.DEFAULT_PROVENANCE_GATES
                                 if provenance_gates is None else bool(provenance_gates))
        # v2.8 — [(rule_code, compiled_regex), ...] from triggers_from_learnings().
        self.learnings_triggers = list(learnings_triggers or [])
        # v2.8 — formula typography normaliser + residual gate (per-exam switch).
        self.formula_typography = bool(formula_typography)
        self.q_re = re.compile(q_pattern)            # e.g. r'^Q\.?\s*(\d+)' or r'^Q(\d+)\.'
        self.opt_re = re.compile(opt_pattern)        # e.g. r'^([1-9])[.\)]' or r'^([A-D])[.\)]'
        self._uniform = int(options_count) if options_count is not None else None
        # v1.3.1 ROOT-CAUSE FIX (registry→engine str/int boundary, ND6): the per-question
        # option-count map round-trips through registry.json, so its keys arrive as JSON
        # STRINGS ("3") even though expected_options() is queried with an INT q
        # (ExplanationBlock.q = int(q)). Normalise keys to int on construction so a NAT
        # question (value 0) resolves regardless of how the caller keyed the map; without
        # this, int-vs-str misses fall through to the uniform count and every NAT question
        # is silently mis-typed as MCQ. Accepts int- or str-keyed input identically.
        self.options_by_q = ({int(k): v for k, v in options_by_q.items()}
                             if options_by_q else None)
        self.options_count = self._uniform           # back-compat (uniform exams)
        self.language = language
        self.label_scheme = label_scheme             # 'numeric'|'alpha_upper'|'alpha_lower'|'roman_lower'|'roman_upper'|list
        self.sentence_terminators = sentence_terminators or '.!?'
        self.labels = labels or {
            'correct_answer': 'Correct Answer',
            'axiom': 'AXIOM', 'deduction': 'DEDUCTION',
            'speed_hack': 'SPEED HACK', 'why_wrong': 'WHY WRONG?',
            'common_pitfalls': 'COMMON PITFALLS',
            'option': 'Option', 'solution_ref': 'Solution',
            'accepted_range': 'accepted range',
        }
        self.labels.setdefault('common_pitfalls', 'COMMON PITFALLS')
        self.labels.setdefault('accepted_range', 'accepted range')
        # v2.4 — COVERAGE BANNER (GAP-2026-08-19-INTERIM-ARTEFACT-UNLABELLED).
        # Language-configurable like every other label. The banner is a
        # DOCUMENT-LEVEL line stating how much of the paper is explained; it is
        # NOT part of any question region and NOT part of any explanation block.
        self.labels.setdefault('coverage_banner', 'COVERAGE')
        self.markers = markers or {
            'axiom': '\u2b1b', 'deduction': '\u2b1b',
            'speed_hack': '\u26a1', 'why_wrong': '\u274c', 'common_pitfalls': '\u274c',
        }
        self.markers.setdefault('common_pitfalls', '\u274c')
        self.colors = colors or {
            'ca': '003366', 'sub': '000000', 'sent': '000000',
            'hdr': '000000', 'sol_ref': '000000',
        }
        # v1.9 EXAM-AGNOSTIC: language-specific prose guards. Defaults are English
        # (the AI's primary output language); non-English exams pass their own patterns
        # via config so the guards catch language-specific template/metacommentary.
        # Pass None to use the built-in English defaults; pass an explicit list/pattern
        # to REPLACE the defaults for that exam's language.
        # NOTE on metacommentary_re: Python's \b word boundary is ASCII-only and does
        # NOT match Devanagari/CJK/Arabic word boundaries. For non-Latin scripts, omit
        # \b and use plain substring patterns (e.g. 'मुझे लगता है' not r'\bमुझे लगता है\b').
        self.banned_blocks = tuple(banned_blocks) if banned_blocks is not None else _BANNED_BLOCKS
        self.banned_templates = tuple(banned_templates) if banned_templates is not None else _BANNED_TEMPLATE
        self.banned_fakecites = tuple(banned_fakecites) if banned_fakecites is not None else _BANNED_FAKECITE
        self.metacommentary_re = (re.compile(metacommentary_re, re.I)
                                  if metacommentary_re is not None else _META_RE)
        # v2.7 — language-configurable absolute-term pattern (§8-0b gate). Non-
        # English exams pass their own; None = the English default.
        self.absolute_terms_re = (re.compile(absolute_terms_re, re.I)
                                  if absolute_terms_re is not None else _ABSOLUTE_TERMS_RE)

    def expected_options(self, q):
        """Option count expected for question q. 0 means NAT (no options).
        Reads the per-question map when present (mixed-section / NAT papers),
        else the uniform count."""
        if self.options_by_q is not None and q in self.options_by_q:
            return self.options_by_q[q]
        return self._uniform

    def option_label(self, i):
        """Map a 1-based option index to the paper's displayed label
        (1/2/3, A/B/C, a/b/c, i/ii/iii, or a custom list).
        Raises ValueError on out-of-bounds for custom/alpha schemes."""
        sch = self.label_scheme
        if isinstance(sch, (list, tuple)):
            if i < 1 or i > len(sch):
                raise ValueError(f'option_label({i}): out of range for custom scheme '
                                 f'of length {len(sch)} — check label_scheme vs option count')
            return str(sch[i - 1])
        if sch == 'alpha_upper':
            if i < 1 or i > 26:
                raise ValueError(f'option_label({i}): alpha_upper supports 1–26 only')
            return chr(ord('A') + i - 1)
        if sch == 'alpha_lower':
            if i < 1 or i > 26:
                raise ValueError(f'option_label({i}): alpha_lower supports 1–26 only')
            return chr(ord('a') + i - 1)
        if sch == 'roman_lower':
            return _int_to_roman(i).lower()
        if sch == 'roman_upper':
            return _int_to_roman(i)
        return str(i)                                # 'numeric' default

# ───────────────────────────── prose guards ────────────────────────────────
# Banned glyphs in student text (EX1). The block-header markers are whitelisted.
_BANNED_GLYPHS = ('\u2713', '\u2714', '\u2717', '\u2718', '\u2611', '\u2612')  # ✓✔✗✘☑☒
_BANNED_LATEX = ('\\frac', '\\sqrt', '\\left', '\\right', '$')
_DIALECT_BANS = (
    (re.compile(r'[\w)\]²³⟧]\s*÷|÷\s*[\w(⟦]'), 'division-sign fraction', r'\frac{a}{b}'),
    (re.compile(r'\^'), 'caret exponent', 'x^{n}'),
    (re.compile(r'(?<![A-Za-z0-9_])[A-Za-zψφχε]_(?=[A-Za-z0-9{(])'), 'underscore subscript', 'V_{B}'),
    (re.compile(r'√\s*\(|√\s*[A-Za-zπλωεℏ]'), 'flat radical', r'\sqrt{…}'),
    (re.compile(r'[\u0300-\u036f\u20d0-\u20ff]'), 'combining-character accent', r'\bar{A} or \vec{E}'),
)
_BANNED_BLOCKS = ('REMEMBER', 'EXAM CONNECTION')
_BANNED_TEMPLATE = (
    'this option is wrong', 'this is a common misconception', "doesn't match the answer",
    'is incorrect as it doesn', 'this option is incorrect', 'simply wrong',
    # v2.7 (b) — learner-psychology speculation. A WHY WRONG line proves the
    # chemical / mathematical contradiction; it does not narrate the learner's
    # mind. The old §15-3 MSQ wording ('lead with the SEDUCTIVE HALF') produced
    # this phrase on 10 of 10 MSQ blocks in the reference paper.
    'the seductive half', 'the seductive part', 'the student thinks',
    'the student assumes', 'a student might think', 'the solver assumes',
    'a solver assumes', 'the solver thinks', 'a hasty solver', 'a careless solver',
    'invites the reading', 'the reader who remembers', 'the reader who',
    'tempts the student', 'tempts a student', 'the trap here is that the student',
)

# v2.7 (a) — ABSOLUTE TERMS. §8-0b reserves absolutes for claims that are
# absolute in the subject's own terms (a definition, a conservation law, a
# mathematical impossibility). Everything else is a TENDENCY, and a tendency
# stated as an absolute is the single most transferable way to teach a false
# rule. The gate is a DECLARATION gate, not a ban: the author may keep an
# absolute by listing the sentence in ExplanationBlock(absolutes_justified=
# {sentence: reason}) — which forces the judgement to be made consciously and
# leaves it on the record. Scoping clauses are excluded because they are the
# form §14-3b ASKS for: 'only when X' / 'only under Y' / 'only if Z' state a
# validity domain, they do not overstate one. 'exactly' followed by a number
# or a math placeholder is arithmetic, not a universal claim.
_ABSOLUTE_TERMS_RE = re.compile(
    r'\b(?:always|never|cannot|impossible|universally|without exception|'
    r'in all cases|in every case|at all|regardless of|irrespective of|'
    r'no matter (?:how|what|which|where)|whatever the|whichever the)\b', re.I)
# NOT in the default set, by measurement: 'only', 'every', 'exactly'. Run over
# a delivered 60-question paper, those three flagged 81 sentences of which ~78
# were plain quantifiers ("only two ions", "for every formula unit", "exactly
# 208") — an 80-percent false-positive rate that would have been answered with
# declaration spam, the exact failure §8-0b exists to prevent. The universal-
# claim shape above ("always", "never", "regardless of", "at all") flagged 21
# sentences of which every one was either a transfer-unsafe universal or a
# genuine absolute worth declaring. A per-exam override may widen the set.

def find_absolute(text, pattern=None):
    """Return the first unqualified absolute term in a student-facing sentence,
    or None. ⟦MATH:…⟧ regions and preserved-OMML tokens are masked first so a
    symbol body can never trip a prose rule."""
    pat = pattern if pattern is not None else _ABSOLUTE_TERMS_RE
    s = _OPAQUE_MATH_RE.sub('\u2202M\u2202', str(text))
    s = T3_REGION_RE.sub('\u2202M\u2202', s)
    m = pat.search(s)
    return m.group(0) if m else None

# v2.7 (c) — epistemic types a transferable claim may carry (§8-2 / §7-7).
TRANSFER_EPISTEMIC_TYPES = ('SCIENTIFIC_GENERAL_RULE', 'MODEL_DEPENDENT_RULE',
                            'EXAM_CONVENTION', 'QUESTION_SPECIFIC_INFERENCE',
                            'OPTION_SET_SHORTCUT')
TRANSFER_OUTCOMES = ('SAFE', 'NARROWED', 'MOVED_TO_DEDUCTION', 'OMITTED')
TRANSFER_SECTIONS = ('AXIOM', 'SPEED_HACK', 'WHY_WRONG', 'COMMON_PITFALLS', 'DEDUCTION')

# v2.8 — neighbour provenance for a §7-7 claim.
TRANSFER_NEIGHBOUR_SOURCE_RE = re.compile(r'^(?:GENERATED|CURATED:[A-Za-z0-9][A-Za-z0-9_.-]*)$')

# v2.8 (a) — HEDGED PROVENANCE. A WHY WRONG / COMMON PITFALLS line either
# states ONE verified wrong path or states the contradiction; it never offers
# a menu of guesses. These phrases are the fingerprint of an invented path.
# Measured on the reference paper: 24 hits on 15 questions, 0 false positives.
_HEDGE_RE = re.compile(
    r'\b(?:or otherwise|otherwise mis\w*|or a similar|perhaps|some other|something|'
    r'or mistakenly|mishandl\w*|miscombination|or by some|by an incomplete|'
    r'in some way|one way or another|or the like)\b|-style\b', re.I)

def find_hedge(text):
    """First hedged-provenance phrase in a student-facing sentence, or None."""
    s = _OPAQUE_MATH_RE.sub('\u2202M\u2202', str(text))
    s = T3_REGION_RE.sub('\u2202M\u2202', s)
    m = _HEDGE_RE.search(s)
    return m.group(0) if m else None

PROVENANCE_MODES = ('VERIFIED_ERROR_PATH', 'DIRECT_CONTRADICTION')

_SAFE_EVAL_NAMES = None
def _safe_eval(expr):
    """Evaluate a pure arithmetic expression (numbers, + - * / ** %, parentheses,
    math.* functions and constants). Anything else raises ValueError."""
    global _SAFE_EVAL_NAMES
    import math, ast
    if _SAFE_EVAL_NAMES is None:
        _SAFE_EVAL_NAMES = {k: getattr(math, k) for k in dir(math) if not k.startswith('_')}
        _SAFE_EVAL_NAMES.update({'abs': abs, 'round': round, 'min': min, 'max': max})
    s = str(expr).strip().replace('^', '**').replace('\u00d7', '*').replace('\u2212', '-')
    try:
        tree = ast.parse(s, mode='eval')
    except SyntaxError as e:
        raise ValueError(f'recompute expression does not parse: {s!r} ({e})')
    ok = (ast.Expression, ast.BinOp, ast.UnaryOp, ast.Constant, ast.Call,
          ast.Name, ast.Load, ast.Add, ast.Sub, ast.Mult, ast.Div, ast.Pow, ast.Mod,
          ast.USub, ast.UAdd, ast.FloorDiv, ast.Tuple)
    for node in ast.walk(tree):
        if not isinstance(node, ok):
            raise ValueError(f'recompute expression uses a disallowed construct '
                             f'{type(node).__name__}: {s!r}')
        if isinstance(node, ast.Name) and node.id not in _SAFE_EVAL_NAMES:
            raise ValueError(f'recompute expression names an unknown symbol {node.id!r}: {s!r}')
        if isinstance(node, ast.Call) and not (isinstance(node.func, ast.Name)
                                               and node.func.id in _SAFE_EVAL_NAMES):
            raise ValueError(f'recompute expression calls a disallowed function: {s!r}')
    try:
        return float(eval(compile(tree, '<recompute>', 'eval'), {'__builtins__': {}},
                          dict(_SAFE_EVAL_NAMES)))
    except Exception as e:
        raise ValueError(f'recompute expression failed to evaluate: {s!r} ({e})')

_NUM_RE = re.compile(r'[-+]?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?')
def _parse_target_number(target):
    """The numeric value of a target and its displayed decimal precision, or
    (None, None) for a non-numeric target."""
    s = str(target).strip().replace('\u2212', '-').replace(',', '')
    m = _NUM_RE.fullmatch(s)
    if not m:
        return None, None
    prec = len(s.split('.')[1]) if '.' in s and 'e' not in s.lower() else 0
    return float(s), prec

def numbers_match(recomputed, target):
    """The reproduce test (§15-2 v1.37.0): equal after rounding to the target's
    own displayed precision, or within 0.5 percent when the target is an
    integer-looking value that the question itself rounded."""
    tv, prec = _parse_target_number(target)
    if tv is None:
        return False
    r = float(recomputed)
    if round(r, prec) == round(tv, prec):
        return True
    return tv != 0 and abs(r - tv) / abs(tv) <= 0.005

def validate_error_provenance(record, keys, section, ctx=''):
    """Shape + reproduce-validate the v2.8 error_provenance mapping for one
    block. `keys` are the WHY WRONG option indices (mcq/msq) or the COMMON
    PITFALLS value strings (nat). Raises ValueError on breach."""
    if not isinstance(record, dict):
        raise ValueError(f'{ctx}: error_provenance must be a dict keyed by wrong option / value')
    want = {str(k) for k in keys}
    have = {str(k) for k in record}
    if have != want:
        raise ValueError(f'{ctx}: DST_NO_PROVENANCE — error_provenance keys {sorted(have)} '
                         f'!= {section} keys {sorted(want)} (one record per wrong '
                         f'option / value, no more, no fewer)')
    for k in keys:
        e = record.get(k, record.get(str(k)))
        if not isinstance(e, dict):
            raise ValueError(f'{ctx}: error_provenance[{k!r}] is not a dict')
        mode = e.get('mode')
        if mode not in PROVENANCE_MODES:
            raise ValueError(f'{ctx}: error_provenance[{k!r}].mode {mode!r} not in {PROVENANCE_MODES}')
        if mode == 'DIRECT_CONTRADICTION':
            if not str(e.get('contradiction', '')).strip():
                raise ValueError(f'{ctx}: error_provenance[{k!r}] DIRECT_CONTRADICTION needs '
                                 f'a non-empty contradiction (why the option/value is '
                                 f'inconsistent with the correct relation)')
            for bad in ('recompute', 'wrong_operation'):
                if str(e.get(bad, '')).strip():
                    raise ValueError(f'{ctx}: error_provenance[{k!r}] DIRECT_CONTRADICTION '
                                     f'must not claim a {bad} — if a path is claimed, '
                                     f'verify it (VERIFIED_ERROR_PATH)')
            continue
        # VERIFIED_ERROR_PATH
        if not str(e.get('wrong_operation', '')).strip():
            raise ValueError(f'{ctx}: error_provenance[{k!r}] VERIFIED_ERROR_PATH needs '
                             f'wrong_operation (the mistake, in words)')
        target = e.get('target', k if section == 'COMMON PITFALLS' else None)
        if target is None or str(target).strip() == '':
            raise ValueError(f'{ctx}: error_provenance[{k!r}] VERIFIED_ERROR_PATH needs target '
                             f'(the option value / content the path must reproduce)')
        tv, _prec = _parse_target_number(target)
        if section == 'COMMON PITFALLS':
            kv, _ = _parse_target_number(k)
            if kv is not None and tv is not None and not numbers_match(tv, k):
                raise ValueError(f'{ctx}: error_provenance[{k!r}].target {target!r} is not '
                                 f'the pitfall value {k!r}')
        if tv is not None:
            if not str(e.get('recompute', '')).strip():
                raise ValueError(f'{ctx}: error_provenance[{k!r}] — numeric target {target!r} '
                                 f'needs recompute (an arithmetic expression the engine '
                                 f'evaluates); a path that is not recomputed is a guess')
            val = _safe_eval(e['recompute'])
            if not numbers_match(val, target):
                raise ValueError(f'{ctx}: DST_UNVERIFIED_NUMERICAL_ORIGIN — '
                                 f'error_provenance[{k!r}] recompute {e["recompute"]!r} = '
                                 f'{val:.6g} does not reproduce target {target!r}; either '
                                 f'find the operation that does, or write the line as a '
                                 f'DIRECT_CONTRADICTION')
            e['recomputed'] = val
        else:
            if not str(e.get('recomputed', '')).strip():
                raise ValueError(f'{ctx}: error_provenance[{k!r}] — non-numeric target '
                                 f'{str(target)[:40]!r} needs recomputed (the wrong content '
                                 f'the stated operation produces, in words)')
            if e.get('matches_target') is not True:
                raise ValueError(f'{ctx}: error_provenance[{k!r}] — non-numeric path must '
                                 f'declare matches_target=True after checking the produced '
                                 f'content IS this option; otherwise use DIRECT_CONTRADICTION')
    return True

# v2.8 (c) — FORMULA TYPOGRAPHY. Deliberately conservative: rewrite only the
# shapes that are unambiguous in any subject, leave locants (C2, C3, N1) and
# bare single-letter+digit tokens alone, never touch ⟦MATH:⟧ / ⟦M:⟧.
_SUB = str.maketrans('0123456789', '\u2080\u2081\u2082\u2083\u2084\u2085\u2086\u2087\u2088\u2089')
_SUP = str.maketrans('0123456789+-', '\u2070\u00b9\u00b2\u00b3\u2074\u2075\u2076\u2077\u2078\u2079\u207a\u207b')
_ELEMENTS = set("""H He Li Be B C N O F Ne Na Mg Al Si P S Cl Ar K Ca Sc Ti V Cr Mn Fe Co Ni Cu
Zn Ga Ge As Se Br Kr Rb Sr Y Zr Nb Mo Tc Ru Rh Pd Ag Cd In Sn Sb Te I Xe Cs Ba La Ce Pr Nd
Pm Sm Eu Gd Tb Dy Ho Er Tm Yb Lu Hf Ta W Re Os Ir Pt Au Hg Tl Pb Bi Po At Rn Fr Ra Ac Th
Pa U Np Pu Am Cm Bk Cf Es Fm Md No Lr Rf Db Sg Bh Hs Mt Ds Rg Cn Nh Fl Mc Lv Ts Og""".split())
# a chemical token: brackets/parens/elements/digits/charge, >=2 element symbols
_CHEM_TOKEN_RE = re.compile(r'(?<![\w\u2080-\u2089\u2070-\u2079\-])'
                            r'([\[(]?[A-Z][A-Za-z0-9()\[\]]*(?:[+\-\u2212])?)'
                            r'(?![\w\u2080-\u2089])')
_DIATOMIC = {'H2', 'O2', 'N2', 'F2', 'Cl2', 'Br2', 'I2', 'O3', 'P4', 'S8'}
_ELEM_DIGIT_RE = re.compile(r'([A-Z][a-z]?|\))(\d+)')
_CHARGE_TAIL_RE = re.compile(r'(\]|\)|[A-Za-z\u2080-\u2089])(\d*)([+\-\u2212])$')
_ION_RE = re.compile(r'(?<![\w\u2080-\u2089])([A-Z][a-z]?)(\d?)([+\-\u2212])(?![\w\u2080-\u2089])')  # Fe3+, Cl-
_HYBRID_RE = re.compile(r'\bsp(\d)(?:d(\d)?)?\b')
_SUBSHELL_RE = re.compile(r'\b([1-7])([spdf])(\d{1,2})\b')
_T2G_RE = re.compile(r'\bt2g(\d{0,2})(?=eg|\b)')
_EG_RE = re.compile(r'(?<![A-Za-z])eg(\d{1,2})\b')
_DCOUNT_RE = re.compile(r'\b([df])(\d{1,2})\b(?=\s+(?:ion|metal|configuration|system|complex|centre|center|case|species|electron|count|cation|state))')
_GREEK = (('pi*', '\u03c0*'), ('sigma*', '\u03c3*'), ('pi-', '\u03c0-'), ('sigma-', '\u03c3-'))

def _fmt_chem_token(tok):
    syms = re.findall(r'[A-Z][a-z]?', tok)
    if not syms or not all(s in _ELEMENTS for s in syms):
        return tok
    if not re.search(r'\d|[+\-\u2212]', tok):
        return tok                      # no digit and no charge — nothing to format
    multi = tok[0] in '[(' or len(syms) >= 2
    if not multi and tok not in _DIATOMIC and not re.search(r'[+\-\u2212]$', tok):
        return tok                      # single element + digit = a locant (C2) — leave
    if tok in _DIATOMIC:
        return tok[:-1] + tok[-1].translate(_SUB)
    m = _CHARGE_TAIL_RE.search(tok)
    charge, body = '', tok
    if m:
        # after ] or ) or in a single-element ion the digits are the CHARGE
        # magnitude (]4-, Fe3+); after an element in a multi-element ion they
        # are a SUBSCRIPT (NH4+ -> NH4 + charge '+').
        if m.group(1) in '])' or not multi:
            charge = (m.group(2) + m.group(3)).replace('\u2212', '-').translate(_SUP)
            body = tok[:m.start(2)]
        else:
            charge = m.group(3).replace('\u2212', '-').translate(_SUP)
            body = tok[:m.start(3)]
    body = _ELEM_DIGIT_RE.sub(lambda mm: mm.group(1) + mm.group(2).translate(_SUB), body)
    return body + charge

def normalise_formula_text(text):
    """Rewrite ASCII chemical / orbital notation to Unicode sub/superscripts.
    Idempotent; ⟦MATH:⟧ regions and preserved tokens are masked; single-letter
    locants (C2, C3) are left alone."""
    s = str(text)
    masks = []
    def _mask(m):
        masks.append(m.group(0)); return f'\u2202{len(masks)-1}\u2202'
    s = T3_REGION_RE.sub(_mask, s)
    s = _OPAQUE_MATH_RE.sub(_mask, s)
    for a, b in _GREEK:
        s = s.replace(a, b)
    s = _CHEM_TOKEN_RE.sub(lambda m: _fmt_chem_token(m.group(1)), s)
    s = _ION_RE.sub(lambda m: (m.group(1) + (m.group(2) + m.group(3)).replace('\u2212', '-')
                               .translate(_SUP)) if m.group(1) in _ELEMENTS else m.group(0), s)
    s = _HYBRID_RE.sub(lambda m: 'sp' + m.group(1).translate(_SUP)
                       + (('d' + (m.group(2) or '').translate(_SUP)) if 'd' in m.group(0) else ''), s)
    s = _SUBSHELL_RE.sub(lambda m: m.group(1) + m.group(2) + m.group(3).translate(_SUP), s)
    s = _T2G_RE.sub(lambda m: 't\u2082g' + m.group(1).translate(_SUP), s)
    s = _EG_RE.sub(lambda m: 'eg' + m.group(1).translate(_SUP), s)
    s = _DCOUNT_RE.sub(lambda m: m.group(1) + m.group(2).translate(_SUP), s)
    def _unmask(m):
        return masks[int(m.group(1))]
    return re.sub(r'\u2202(\d+)\u2202', _unmask, s)

# residual shapes the normaliser refused to touch but which are unmistakably
# formulae: a bracketed/parenthesised formula with a plain digit, an ion with
# a plain charge, an sp-hybrid or orbital label with plain digits.
_UNFORMATTED_RE = re.compile(
    '\\[[A-Za-z()0-9\\-\u03b7\u00b9\u00b2\u00b3\u2070-\u2079\u2080-\u2089]*[A-Z][a-z]?\\)?\\d[A-Za-z()0-9\\-]*\\]|'
    r'\b[A-Z][a-z]?\d[+-](?![\w])|\bsp\d|\bt2g\b|\b[1-7][spdf]\d\b')

def find_unformatted_formula(text):
    s = _OPAQUE_MATH_RE.sub('\u2202M\u2202', str(text))
    s = T3_REGION_RE.sub('\u2202M\u2202', s)
    m = _UNFORMATTED_RE.search(s)
    return m.group(0) if m else None

# v2.8 (b) — curated-neighbour triggers.
def triggers_from_learnings(parsed):
    """[(rule_code, compiled_regex)] from a parse_learnings() result (or a list
    of them). A rule contributes when it carries a **Triggers:** field: comma-
    separated terms; a term that starts with 're:' is a raw regex, anything
    else is matched as a case-insensitive phrase."""
    out = []
    srcs = parsed if isinstance(parsed, (list, tuple)) else [parsed]
    for p in srcs:
        for r in (p or {}).get('rules', []):
            t = r.get('triggers')
            if not t or r.get('superseded'):
                continue
            pats = []
            for term in re.split(r'\s*[,;\n]\s*', str(t).strip()):
                if not term:
                    continue
                pats.append(term[3:] if term.startswith('re:') else
                            r'(?<!\w)' + re.escape(term).replace(r'\ ', r'\s+') + r'(?!\w)')
            if pats:
                out.append((r['code'], re.compile('|'.join(pats), re.I)))
    return out

def find_trigger_hits(sentences, triggers):
    """{rule_code: first_matching_phrase} over a list of sentences."""
    hits = {}
    for s in sentences or []:
        s2 = T3_REGION_RE.sub(' ', _OPAQUE_MATH_RE.sub(' ', str(s)))
        for code, rx in triggers or []:
            if code in hits:
                continue
            m = rx.search(s2)
            if m:
                hits[code] = m.group(0)
    return hits

def transfer_tripwire(blocks, min_claims=20):
    """§7-7 v1.37.0 tripwire: a run whose AXIOM claims are ALL SAFE is the
    self-attestation shape. Returns (fired: bool, summary: dict)."""
    n = safe = narrowed = moved = 0
    for b in (blocks.values() if isinstance(blocks, dict) else blocks):
        for e in (getattr(b, 'transfer_record', None) or []):
            if e.get('section') != 'AXIOM':
                continue
            n += 1
            oc = e.get('outcome')
            safe += oc == 'SAFE'; narrowed += oc == 'NARROWED'; moved += oc == 'MOVED_TO_DEDUCTION'
    fired = n >= min_claims and (narrowed + moved) == 0
    return fired, {'axiom_claims': n, 'safe': safe, 'narrowed': narrowed,
                   'moved_to_deduction': moved, 'fired': fired}

def canonical_answer_of_block(blk):
    """The canonical answer string Step 7 committed (paper_pipeline.canonical_answer):
    mcq -> '2'; msq -> '2,3'; nat -> the grading string (lo-hi for a range)."""
    if blk.qtype == 'nat':
        if blk.ca_range is not None and blk.ca is None:
            return format_nat_range(blk.ca_range[0], blk.ca_range[1])
        return str(blk.ca).strip()
    if blk.qtype == 'msq':
        return ','.join(str(i) for i in sorted(blk.ca_set()))
    return str(blk.ca)

def canonical_structure(smiles):
    """Canonical isomeric SMILES via rdkit, with valence sanitisation. Returns
    (canonical|None, reason); reason == 'ok' | 'rdkit_unavailable' | a parse
    error. NEVER raises — the caller decides what an unavailable rdkit means."""
    try:
        from rdkit import Chem
        from rdkit import RDLogger
        RDLogger.DisableLog('rdApp.*')
    except Exception:
        return None, 'rdkit_unavailable'
    try:
        mol = Chem.MolFromSmiles(str(smiles), sanitize=True)
    except Exception as e:
        return None, f'parse_error: {e}'
    if mol is None:
        return None, 'parse_error: rdkit rejected the SMILES (valence / syntax)'
    return Chem.MolToSmiles(mol, isomericSmiles=True, canonical=True), 'ok'

def semantic_objects_agree(mine, theirs):
    """v2.8 — §13-2b: paper_pipeline.semantic_objects_agree with this module's
    rdkit canonicaliser injected (the Explain-route home of the impure dep)."""
    import paper_pipeline as pp
    return pp.semantic_objects_agree(mine, theirs, canon=canonical_structure)

def reconcile_key_commitments(blocks, registry, paper_id):
    """v2.8 (d) — compare Step 9's derived answers against Step 7's committed
    hashes (registry['key_commitments'][paper_id]). Returns a dict:
    {'available': bool, 'matched': [q], 'mismatched': [q], 'uncommitted': [q],
     'candidates': {q: <which candidate canonical the commitment matches, if any>}}.
    Never raises on a mismatch — §17 v1.37.0 resolves in-run."""
    import paper_pipeline as pp
    com = (registry or {}).get('key_commitments', {}).get(paper_id)
    out = {'available': bool(com), 'matched': [], 'mismatched': [], 'uncommitted': [],
           'candidates': {}}
    if not com:
        return out
    derived = {}
    for q, b in (blocks.items() if isinstance(blocks, dict) else ((b.q, b) for b in blocks)):
        if getattr(b, 'anomaly', None) is None:
            derived[int(q)] = canonical_answer_of_block(b)
    res = pp.verify_key_commitments(com, derived, paper_id)
    out['matched'] = res['matched']; out['mismatched'] = res['mismatched']
    out['uncommitted'] = res['missing']
    for q in res['mismatched']:
        b = blocks[q] if isinstance(blocks, dict) else next(x for x in blocks if x.q == q)
        n = b.cfg.expected_options(q) or 0
        cands = [str(i) for i in range(1, n + 1)] if b.qtype == 'mcq' else []
        out['candidates'][q] = pp.resolve_commitment(com, q, cands, paper_id)
    return out

_RISK_ABS_RE = _ABSOLUTE_TERMS_RE
def scan_risk_markers(path, cfg):
    """v2.8 (e) — migration scan over an existing Explanation docx. Returns
    {q: {'hedge': [...], 'absolute': [...], 'formula': [...]}} for questions
    carrying any marker; a question with none is omitted."""
    import copy
    _cfg = copy.copy(cfg)
    _cfg.formula_typography = False          # read the text AS SHIPPED, un-normalised
    _cfg.provenance_gates = False
    blocks = parse_solution_blocks(path, _cfg)
    out = {}
    for q, b in sorted(blocks.items()):
        rec = {'hedge': [], 'absolute': [], 'formula': []}
        ww = [s for v in b.why_wrong.values() for s in v] + \
             [s for v in b.common_pitfalls.values() for s in v]
        for s in ww:
            h = find_hedge(s)
            if h: rec['hedge'].append(h)
        for s in b.axiom + (b.speed_hack or []) + ww:
            a = find_absolute(s, getattr(cfg, 'absolute_terms_re', None))
            if a: rec['absolute'].append(a)
        for s in b.axiom + b.deduction + (b.speed_hack or []) + ww:
            f = find_unformatted_formula(s)
            if f: rec['formula'].append(f)
        if any(rec.values()):
            out[q] = rec
    return out

def validate_transfer_record(record, axiom_present=True, speed_hack_present=False,
                             ctx=''):
    """Shape-validate a §7-7 transfer record. Raises ValueError on breach.
    A record is a list of claim dicts. This proves the PROTOCOL WAS RUN and its
    outcome recorded; it cannot prove the neighbour test was judged correctly —
    that stays with the §7-7 discipline, exactly as §6A-5 proves the drawn
    artefact and not the request."""
    if not isinstance(record, (list, tuple)) or not record:
        raise ValueError(f'{ctx}: transfer_record must be a non-empty list of claims')
    need = ('section', 'claim', 'epistemic_type', 'scope', 'neighbour_tested', 'outcome')
    seen = set()
    for i, c in enumerate(record):
        if not isinstance(c, dict):
            raise ValueError(f'{ctx}: transfer_record[{i}] is not a dict')
        for k in need:
            if not str(c.get(k, '')).strip():
                raise ValueError(f'{ctx}: transfer_record[{i}] missing {k!r}')
        sec = c['section']; et = c['epistemic_type']; oc = c['outcome']
        if sec not in TRANSFER_SECTIONS:
            raise ValueError(f'{ctx}: transfer_record[{i}] section {sec!r} not in {TRANSFER_SECTIONS}')
        if et not in TRANSFER_EPISTEMIC_TYPES:
            raise ValueError(f'{ctx}: transfer_record[{i}] epistemic_type {et!r} not in {TRANSFER_EPISTEMIC_TYPES}')
        if oc not in TRANSFER_OUTCOMES:
            raise ValueError(f'{ctx}: transfer_record[{i}] outcome {oc!r} not in {TRANSFER_OUTCOMES}')
        _ns = c.get('neighbour_source')
        if _ns is not None and not TRANSFER_NEIGHBOUR_SOURCE_RE.match(str(_ns)):
            raise ValueError(f'{ctx}: transfer_record[{i}] neighbour_source {_ns!r} must be '
                             f'GENERATED or CURATED:<rule-code>')
        if sec == 'AXIOM' and et == 'QUESTION_SPECIFIC_INFERENCE' and oc not in ('MOVED_TO_DEDUCTION', 'OMITTED'):
            raise ValueError(f'{ctx}: transfer_record[{i}] — a QUESTION_SPECIFIC_INFERENCE '
                             f'may not stand in AXIOM; move it to DEDUCTION (§8-2)')
        if et == 'OPTION_SET_SHORTCUT' and sec != 'SPEED_HACK' and oc not in ('OMITTED',):
            raise ValueError(f'{ctx}: transfer_record[{i}] — an OPTION_SET_SHORTCUT may '
                             f'stand only in SPEED HACK (§8-2 / §14-1 part 3)')
        seen.add(sec)
    if axiom_present and 'AXIOM' not in seen:
        raise ValueError(f'{ctx}: transfer_record carries no AXIOM claim — every AXIOM '
                         f'is tested for transfer safety (§7-7)')
    if speed_hack_present and 'SPEED_HACK' not in seen:
        raise ValueError(f'{ctx}: SPEED HACK present but transfer_record carries no '
                         f'SPEED_HACK claim (§14-1 part 3)')
    return True
_BANNED_FAKECITE = (
    'official key', 'official answer', 'official solution', 'the answer key says',
    'as per the official', 'per the official key',
)

# Metacommentary as a WORD-BOUNDED regex (substring matching gave false positives:
# 'hmm' hit "ohmmeter", 'wait,' hit "await,", 'actually,' hit "factually,",
# 'as an ai' hit "has an air"). The boundaries make detection precise.
_META_RE = re.compile(
    r'\bre-reading\b|\blet me reconsider\b|\bactually,|\bhmm\b|\bwait\s*[,\u2014]|'
    r'\bas an ai\b|\blet me re-check\b|\blet me check again\b|\bi think we\b|'
    r'\bon reflection\b',
    re.I)

# v2.6 (GAP-2026-08-19-EXPLANATION-EXECUTION-INTEGRITY, D1) — INTERNAL ERROR-
# TAXONOMY TOKENS ARE METADATA, NEVER STUDENT TEXT. The §9 error type exists so a
# WHY WRONG line is a real diagnosis; it was never meant to be RENDERED. In the
# reference incident a delivered 60-question paper opened every WHY WRONG entry
# with the raw snake_case token ('regiochemistry_error: the para phenol ...'),
# because the spec asked the first line to "name" the type and nothing separated
# the internal name from the visible sentence. The diagnosis is recorded in
# progress state; the visible line states the same content in natural language.
# Any taxonomy token in a student-facing sentence now raises at write time.
_INTERNAL_TAG_RE = re.compile(
    r'\b(?:value_swap|sign_error|unit_error|off_by_one|partial_truth|'
    r'process_confusion|reversed_relationship|name_swap|formula_error|'
    r'rounding_trap|polarity_flip|wrong_condition|regiochemistry_error|'
    r'stereochemistry_error|mechanism_confusion|electron_count_error|'
    r'symmetry_error|overgeneralised_rule|concept_reversal)\b')

# digit/digit fraction with a negative lookahead so a trailing '.' (decimal) does
# NOT match — that leftover then trips has_inline_fraction (EX13, intentional).
_SIMPLE_FRAC = re.compile(r'(?<![\w/.])(\d+)\s*/\s*(\d+)(?![\d.\w/])')
# letter/letter units that are legitimately NOT fractions (km/h, m/s, w/o).
_UNIT_SLASH = re.compile(r'\b[a-zA-Z]{1,3}\s*/\s*[a-zA-Z]{1,3}\b')
# Vulgar-fraction glyphs and the Unicode fraction slash (U+2044): never allowed in
# prose — algebra/fractions must be real OMML, so these raise (A1, parity with T2).
VULGAR = '\u00bd\u2153\u2154\u00bc\u00be\u2155\u2156\u2157\u2158\u2159\u215a' \
         '\u2150\u215b\u215c\u215d\u215e\u2151\u2152\u2044'
# Inline-fraction detectors. Beyond digit/digit, catch a/letter, a/(, a/√, x²/2,
# (a+b)/c — non-convertible forms that must be built explicitly via frac() (A4).
# Units (km/h, m/s) are letter/letter and are masked out before these run.
_FRAC_PATS = (
    re.compile(r'\d\s*/\s*\d'),                  # 3/4, 2025/26
    re.compile(r'\d\s*/\s*[A-Za-z(\u221a]'),     # 1/x, 1/(x+1), 1/√2
    re.compile(r'[\u00b2\u00b3)]\s*/\s*[\dA-Za-z(]'),  # x²/2, (a+b)/c
)
# Consecutive-year slash backstop (1947/48, 2025/26, 2024/2025) — a year range must
# never silently render as a stacked fraction. Only flags when the second part is
# actually (year+1), so a genuine n/(n+1) telescoping fraction is NOT false-flagged (A6).
_YEAR_RANGE = re.compile(r'\b(1[89]\d{2}|20\d{2})\s*/\s*(\d{2}|\d{4})\b')

def _year_range_hit(text):
    """Return the matched 'YYYY/NN' string if it is a consecutive year range
    (1947/48, 2025/26, 2024/2025), else None. A real fraction whose denominator
    is not year+1 passes through untouched."""
    for m in _YEAR_RANGE.finditer(text):
        y, nxt = int(m.group(1)), m.group(2)
        if (len(nxt) == 2 and int(nxt) == (y + 1) % 100) or (len(nxt) == 4 and int(nxt) == y + 1):
            return m.group(0)
    return None

_SENT = '\u0001'   # transient sentinel standing in for a non-terminal dot

def _abbrev_safe(text):
    """Replace abbreviations / initials / decimals so the sentence counter does
    not treat their dots as sentence ends (EX2, EX12)."""
    t = text.replace('...', _SENT * 3)
    for ab in ('Dr.', 'Mr.', 'Mrs.', 'Ms.', 'Smt.', 'Shri.', 'Sh.', 'Pt.', 'Mt.',
               'Lt.', 'Rs.', 'No.', 'Nos.', 'vs.', 'etc.', 'approx.', 'viz.', 'ft.',
               'Govt.', 'i.e.', 'e.g.', 'Col.', 'Gen.', 'Capt.', 'Maj.', 'Hon.',
               'Prof.', 'Jr.', 'Sr.', 'Corp.', 'Inc.', 'Ltd.', 'Rev.', 'Md.',
               'sq.', 'cu.', 'Art.', 'Sec.', 'Vol.', 'Ed.', 'Fig.', 'St.'):
        t = t.replace(ab, ab.replace('.', _SENT))
    # lowercase dotted acronyms: u.s, a.m, p.m, e.g, i.e (any single-letter . chain)
    t = re.sub(r'\b([a-z](?:\.[a-z])+)\.?',
               lambda m: m.group(0).replace('.', _SENT), t)
    t = re.sub(r'\b([A-Z])\.', lambda m: m.group(1) + _SENT, t)   # initials
    t = re.sub(r'(\d)\.(\d)', lambda m: m.group(1) + _SENT + m.group(2), t)  # decimals
    # v2.1: a decimal whose fractional digits sit inside an adjacent math token
    # renders as 'N.' + ∂M∂ (the ∂M∂ standing for ⟦MATH:…⟧ or a preserved ⟦M:…⟧);
    # that dot is a decimal point, never a sentence end.
    t = re.sub(r'(\d)\.(?=\u2202M\u2202)', lambda m: m.group(1) + _SENT, t)
    return t

def sentence_count(text, terminators='.!?'):
    """Abbreviation-aware sentence counter. Semicolon/comma joins count as ONE.
    `terminators` is language-configurable (e.g. include the Devanagari danda
    '\u0964' and double-danda '\u0965' for Hindi/Sanskrit/Marathi)."""
    t = _abbrev_safe(text.strip())
    cls = re.escape(''.join(dict.fromkeys(terminators)))
    parts = [p for p in re.split(f'[{cls}]+', t) if p.strip()]
    return max(1, len(parts))

_NAT_CHARSET_ALLOWED = frozenset('0123456789.-')
_NAT_POINT_RE = re.compile(r'^-?\d+(?:\.\d+)?$')

def validate_nat_grading_value(s, ctx=''):
    """Portal-safe grading-value check for a NAT point value (locked, DJ rule,
    v1.16 hardening). ALLOWLIST-only — 0-9 . - and NOTHING else — never a
    blacklist of known-bad patterns, because a blacklist only catches the
    patterns someone thought to ban (scientific notation, e.g.) and silently
    lets a stray unit letter, en-dash, space, or paren through. Raises
    ValueError naming the exact bad character(s) so the failure is traceable
    to Step 7/9 authoring, never silently coerced/stripped.
    `ctx` is prepended to the error (e.g. 'Q7' or 'Q7 ca_range.lo')."""
    if s is None:
        raise ValueError(f'{ctx}: NAT grading value missing')
    s = str(s).strip()
    if s == '':
        raise ValueError(f'{ctx}: NAT grading value empty')
    bad = sorted(set(s) - _NAT_CHARSET_ALLOWED)
    if bad:
        raise ValueError(f'{ctx}: NAT grading value {s!r} contains banned '
                          f'character(s) {bad} — portal charset is exactly '
                          f'"0123456789.-", nothing else (no scientific notation, '
                          f'no units, no spaces, no en-dash, no parentheses). The '
                          f'value must already be pre-scaled to the units the '
                          f'question stem states before it reaches this engine.')
    if not _NAT_POINT_RE.match(s):
        raise ValueError(f'{ctx}: NAT grading value {s!r} is not a well-formed '
                          f'plain number (expected: optional leading "-", digits, '
                          f'optional single "." followed by digits — e.g. "3", '
                          f'"-47", "3.00")')
    return s


def format_nat_range(lo, hi, ctx=''):
    """Build the portal Range-type grading string '{lo}-{hi}' (locked, DJ rule,
    v1.16 hardening) — no space, no parentheses, no words, no en-dash. Each
    bound is independently charset-validated first.
    NEGATIVE BOUNDS ARE NOT SUPPORTED (explicit DJ decision, not an inferred
    default): the '-' join delimiter is structurally ambiguous with a bound's
    own leading minus sign (e.g. is '-5-7' the range [-5,7] or a malformed
    single value?). Rather than guess a convention the portal was never
    confirmed to parse, this raises and requires escalation — matching the
    project's existing halt-on-genuine-ambiguity pattern (G-NAT-ANSWER /
    A-NAT-ANSWER), not a silent workaround."""
    lo_s = validate_nat_grading_value(lo, ctx=f'{ctx} ca_range.lo')
    hi_s = validate_nat_grading_value(hi, ctx=f'{ctx} ca_range.hi')
    if lo_s.startswith('-') or hi_s.startswith('-'):
        raise ValueError(f'{ctx}: NAT range with a negative bound '
                          f'(lo={lo_s!r}, hi={hi_s!r}) is NOT SUPPORTED — the '
                          f'"-" delimiter between lo/hi is ambiguous with a '
                          f'bound\'s own sign. Escalate to Step 7/8: rework this '
                          f'question so both bounds are non-negative, or confirm '
                          f'an unambiguous portal-side range convention before '
                          f'this gate is ever lifted.')
    if float(lo_s) > float(hi_s):
        raise ValueError(f'{ctx}: NAT range lo>hi ({lo_s!r} > {hi_s!r})')
    return f'{lo_s}-{hi_s}'


def has_inline_fraction(text):
    """True if a vulgar-fraction glyph, the Unicode fraction slash, or any bare
    inline fraction (digit/digit, a/letter, a/(, a/√, x²/2, (a+b)/c) remains after
    units are masked. Such forms must be real OMML, so the prose emitter raises on
    any that survive auto-conversion (EX13, A1/A4)."""
    if any(ch in text for ch in VULGAR):
        return True
    masked = _UNIT_SLASH.sub('', text)
    return any(p.search(masked) for p in _FRAC_PATS)

# ── v2.2 (GAP-2026-08-19-EXPLAIN-MATH-NOTATION) — TIER-3 NOTATION GUARD ────
# THE HOLE THIS CLOSES: guard_sentence deliberately exempts every ⟦MATH:…⟧ body
# from every prose guard, on the stated grounds that "the compiler validates
# them". The compiler validates GRAMMAR, never NOTATION. t3_compile binds a bare
# _ / ^ to exactly ONE preceding character and passes unknown words through as
# literal text, so it SILENTLY mis-renders ordinary scientific notation and
# NEVER raises — the output ships clean and reads wrong:
#     Delta_o   -> "Delt" + a_o   (reads "Delt aₒ", not Δₒ)
#     K_sp      -> K_s  + "p"     (reads "Kₛp",     not K_sp)
#     E_cell^0  -> E_c + "el" + l^0
#     sqrt(x)   -> literal text, no radical
# Each has a spelling the SHIPPED compiler already handles perfectly —
# \Delta_{o}, K_{sp}, E_{cell}^{0}, \sqrt{x} — verified against t3_mathcomp.
# This guard rejects the mis-compiling spelling at AUTHORING time and names the
# exact remedy. It is a NOTATION check only: it adds no grammar, changes no
# compiler behaviour, and t3_mathcomp.py is NOT touched — that file is byte-
# locked to Framework_PYQPrepare §S3-5b by the self-test drift lock and must
# never be edited from this side.
_T3_GREEK = ('alpha','beta','gamma','delta','epsilon','zeta','eta','theta',
             'iota','kappa','lambda','mu','nu','xi','omicron','rho','sigma',
             'tau','upsilon','phi','chi','psi','omega')
# A Greek NAME spelled without its backslash. NOTE: \b is unusable here because
# '_' is a regex word character, so \bDelta\b does NOT match inside 'Delta_o' —
# the exact bug this guard exists to catch. Explicit letter lookarounds instead.
_T3_BARE_GREEK = re.compile(
    r'(?<![A-Za-z\\])(' + '|'.join(_T3_GREEK) + r')(?![A-Za-z])', re.I)
# EVERY sub/superscript must be braced. A bare script is ambiguous by nature:
# '_2SO' in H_2SO_4 is correct (one-char sub) while '_2g' in t_2g is wrong
# (two-char sub) and the two are indistinguishable mechanically. Requiring
# braces everywhere removes the ambiguity instead of guessing at intent.
_T3_BARE_SUBSUP = re.compile(r'[_^](?!\{)')
# Function words that only compile with a backslash and braces.
_T3_BARE_FUNC = re.compile(
    r'(?<!\\)\b(sqrt|int|sum|prod|lim|frac|bar|vec|hat)\s*[({_^]')

def t3_notation_guard(body, ctx=''):
    """Reject ⟦MATH:…⟧ notation that t3_compile mis-renders SILENTLY.
    Raises ValueError naming the exact remedy; returns body unchanged on success."""
    where = f' in {ctx}' if ctx else ''
    m = _T3_BARE_GREEK.search(body)
    if m:
        raise ValueError(
            f'⟦MATH:⟧ bare Greek name {m.group(1)!r}{where} compiles to literal '
            f'letters, not a symbol — write \\{m.group(1)}: {body[:60]!r}')
    m = _T3_BARE_FUNC.search(body)
    if m:
        raise ValueError(
            f'⟦MATH:⟧ function {m.group(1)!r}{where} without a backslash compiles '
            f'to literal text — write \\{m.group(1)}{{…}}: {body[:60]!r}')
    if _T3_BARE_SUBSUP.search(body):
        raise ValueError(
            f'⟦MATH:⟧ unbraced sub/superscript{where} binds only ONE character in '
            f'the compiler — brace every script, e.g. K_{{sp}} not K_sp: {body[:60]!r}')
    return body

def guard_sentence(text, cfg=None):
    """Validate one student-facing sentence. Raises ValueError on any breach.
    Returns the text unchanged on success. `cfg` (optional) supplies language-
    specific sentence terminators and configurable banned patterns; without it
    the English/Latin defaults are used."""
    if text is None or not str(text).strip():
        raise ValueError('empty sentence')
    s_full = str(text)
    # v2.1 — opaque preserved-OMML tokens (⟦M:<b64>⟧, and legacy bodiless ⟦M⟧) are
    # RESOLVED math citizens: collapse them to the delimiter-free ∂M∂ placeholder
    # FIRST, so their closing bracket is never read as an unbalanced ⟦MATH:⟧
    # delimiter and their base64 body never reaches the ASCII-dialect / year-range
    # guards below. This runs before — and independently of — the Tier-3 region
    # handling, and weakens no existing check.
    s_full = _OPAQUE_MATH_RE.sub('\u2202M\u2202', s_full)
    # v2.0 — Tier-3 region awareness: ⟦MATH:…⟧ bodies are validated by the
    # compiler, not by prose guards; every check below sees regions as ⟦M⟧.
    _stripped = T3_REGION_RE.sub('', s_full)
    if T3_OPEN in _stripped or T3_CLOSE in _stripped:
        raise ValueError(f'unbalanced ⟦MATH:⟧ delimiters in: {s_full[:60]!r}')
    # v2.2 — the compiler checks GRAMMAR but not NOTATION: every ⟦MATH:⟧ body is
    # checked here for spellings that compile silently to the WRONG symbol.
    for _t3m in T3_REGION_RE.finditer(s_full):
        t3_notation_guard(_t3m.group(1))
    s = T3_REGION_RE.sub('\u2202M\u2202', s_full)   # ∂M∂ placeholder — delimiter-free
    # v2.0 — ASCII-math dialect is BANNED in prose. The v1 guard only understood
    # digit/digit slashes, so generation evaded it with ÷, carets, V_B
    # subscripts, √( … ) and combining overbars/arrows (measured: 234 ÷-fractions
    # in one paper). Every construct now has a legal spelling: a ⟦MATH:…⟧ region.
    for pat, what, fixit in _DIALECT_BANS:
        mban = pat.search(s)
        if mban:
            raise ValueError(f'{what} {mban.group(0)!r} — write it as a '
                             f'⟦MATH:{fixit}⟧ region: {s_full[:60]!r}')
    low = s.lower()
    terminators = cfg.sentence_terminators if cfg is not None else '.!?'
    meta_re = cfg.metacommentary_re if cfg is not None else _META_RE
    banned_templates = cfg.banned_templates if cfg is not None else _BANNED_TEMPLATE
    banned_fakecites = cfg.banned_fakecites if cfg is not None else _BANNED_FAKECITE
    banned_blocks = cfg.banned_blocks if cfg is not None else _BANNED_BLOCKS
    for g in _BANNED_GLYPHS:
        if g in s:
            raise ValueError(f'banned glyph {g!r} in: {s[:60]!r}')
    for v in VULGAR:
        if v in s:
            raise ValueError(f'vulgar fraction glyph {v!r} — build it as OMML via frac(): {s[:60]!r}')
    for lx in _BANNED_LATEX:
        if lx in s:
            raise ValueError(f'banned LaTeX {lx!r} in: {s[:60]!r}')
    m_meta = meta_re.search(s)
    if m_meta:
        raise ValueError(f'metacommentary {m_meta.group(0)!r} in: {s[:60]!r}')
    # v2.6 D1 — the §9 diagnosis is INTERNAL metadata: record the token in
    # progress state; write the visible line in natural language (§9/§15).
    m_tag = _INTERNAL_TAG_RE.search(s)
    if m_tag:
        raise ValueError(
            f'internal error-taxonomy token {m_tag.group(0)!r} in student-facing '
            f'text — the diagnosis is metadata, never rendered; state the error '
            f'in natural language and record the token internally (§9/§15): '
            f'{s[:60]!r}')
    for tp in banned_templates:
        if tp in low:
            raise ValueError(f'template sentence {tp!r} in: {s[:60]!r}')
    for fc in banned_fakecites:
        if fc in low:
            raise ValueError(f'fake citation {fc!r} in: {s[:60]!r}')
    for bb in banned_blocks:
        if bb in s:
            raise ValueError(f'banned block {bb!r} in: {s[:60]!r}')
    yr = _year_range_hit(_UNIT_SLASH.sub('', s))
    if yr:
        raise ValueError(f'year-range slash {yr!r} (use en-dash) in: {s[:60]!r}')
    if sentence_count(s, terminators) != 1:
        raise ValueError(f'not exactly one sentence: {s[:60]!r}')
    return s

# ───────────────────────────── OMML helpers ────────────────────────────────
# v2.0 (GAP-2026-08-07-EXPLAIN-OMML): shared Tier-3 compiler — SINGLE SOURCE with
# Framework_PYQPrepare §S3-5b; the self-test drift lock keeps them byte-identical.
from t3_mathcomp import (t3_compile, MathCompileError, count_math_regions,
                         MATH_OPEN as T3_OPEN, MATH_CLOSE as T3_CLOSE,
                         _REGION_RE as T3_REGION_RE, _T3_STATS as T3_STATS)

# v2.1 (GAP-2026-08-07-EXPLAIN-OMML-ROUNDTRIP) — PRESERVE-AND-REEMIT.
# The v2.0 strict reader collapsed every non-digit/digit OMML structure to a
# BODILESS token ⟦M⟧, which (a) tripped guard_sentence's ⟦MATH:⟧ balance check
# on its lone ⟧ and (b) could not be rebuilt, so build_interleaved_docx and
# verify_structure failed on ANY math-bearing explanation (all pre-v2.0 papers,
# and any physics/quant paper). v2.1 makes the strict reader LOSS-LESS: each such
# structure is serialised, wrapped in a standalone <m:oMath>, base64-encoded and
# carried inside a self-delimiting token ⟦M:<base64>⟧. base64's alphabet
# (A-Za-z0-9+/=) contains no ⟦, ⟧ or the literal "⟦MATH:" sequence, so the token
# is opaque to the Tier-3 delimiter guard. add_math_text() re-emits the preserved
# OMML byte-faithfully; guard_sentence() collapses the token to the same ∂M∂
# placeholder it already uses for ⟦MATH:…⟧ regions. Legacy bodiless ⟦M⟧ is still
# tolerated by the guard (it only ever appears in the non-strict verifier text,
# never on the writer path). ADDITIVE: no existing guard is weakened, and the
# digit/digit auto-fraction path is byte-unchanged.
_OPAQUE_MATH_RE = re.compile(r'\u27e6M(?::[A-Za-z0-9+/=]*)?\u27e7')
_MATH_TOKEN_RE  = re.compile(
    r'\u27e6MATH:(?P<t3>.*?)\u27e7'
    r'|\u27e6M(?::(?P<op>[A-Za-z0-9+/=]*))?\u27e7')

def _encode_opaque(node):
    """v2.1 — serialise ONE non-fraction OMML child, wrapped in a standalone
    <m:oMath>, into the loss-less, guard-safe token ⟦M:<base64>⟧ that
    add_math_text() re-emits verbatim. The wrapper carries the m: namespace so the
    payload parses standalone; a doubled xmlns:m (the child already declares it) is
    harmless and idempotent under parse_xml."""
    raw = etree.tostring(node)                       # bytes; declares the m: ns
    xml = b'<m:oMath xmlns:m="' + M.encode('utf-8') + b'">' + raw + b'</m:oMath>'
    return '\u27e6M:' + base64.b64encode(xml).decode('ascii') + '\u27e7'

def _esc(t):
    return (str(t).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))

def _r(t):     return f'<m:r xmlns:m="{M}"><m:t xml:space="preserve">{_esc(t)}</m:t></m:r>'

def _r_wrap(x):
    """Builder-argument normaliser (v2.0). The v1 builders interpolated RAW text
    into <m:num>/<m:den>/<m:e> — schema-invalid OMML that every Word engine
    renders as an EMPTY placeholder while itertext()-based checks still read it
    (the GAP-2026-08-07-EXPLAIN-OMML defect, measured: 12 destroyed fractions in
    one delivered paper). Every builder argument now passes through here: raw
    OMML markup ('<m:…') is kept, anything else becomes a proper <m:r><m:t> run."""
    x = '' if x is None else str(x)
    return x if x.lstrip().startswith('<m:') else _r(x)
def frac(n, d):return f'<m:f xmlns:m="{M}"><m:num>{_r_wrap(n)}</m:num><m:den>{_r_wrap(d)}</m:den></m:f>'
def sup(b, e): return f'<m:sSup xmlns:m="{M}"><m:e>{_r_wrap(b)}</m:e><m:sup>{_r_wrap(e)}</m:sup></m:sSup>'
def sqrt(x):   return (f'<m:rad xmlns:m="{M}"><m:radPr><m:degHide m:val="1"/></m:radPr>'
                       f'<m:deg/><m:e>{_r_wrap(x)}</m:e></m:rad>')
def nary(op, sub, sup_, body):
    return (f'<m:nary xmlns:m="{M}"><m:naryPr><m:chr m:val="{op}"/></m:naryPr>'
            f'<m:sub>{_r_wrap(sub)}</m:sub><m:sup>{_r_wrap(sup_)}</m:sup>'
            f'<m:e>{_r_wrap(body)}</m:e></m:nary>')
def omath(*parts):
    return f'<m:oMath xmlns:m="{M}">{"".join(parts)}</m:oMath>'

def add_math(paragraph, omath_xml):
    """Insert an explicit OMML node into a paragraph."""
    paragraph._p.append(parse_xml(omath_xml))

def _run(paragraph, text, bold=False, color=None):
    r = paragraph.add_run(text)
    r.bold = bool(bold)
    if color:
        rpr = r._r.get_or_add_rPr()
        c = OxmlElement('w:color'); c.set(qn('w:val'), color); rpr.append(c)
    return r

def add_math_text(paragraph, text, bold=False, color=None, cfg=None, preserve=False):
    """THE prose emitter. Auto-converts every digit/digit fraction to stacked
    OMML; raises ValueError on any inline fraction it cannot convert, on a
    year-range slash, or on any guard breach. This is the ONLY sanctioned way to
    put explanation prose into the docx."""
    guard_sentence(text, cfg)                          # full guard first (region-aware)
    # v2.0/v2.1 — UNIFIED REGION WALK. Two kinds of math token are handled in one
    # ordered pass, with plain segments between them keeping the v1 digit/digit
    # auto-fraction path byte-compatibly:
    #   • ⟦MATH:…⟧  (Tier-3 DISPATCH) compiles through the shared t3_mathcomp
    #     compiler (one <m:oMath> each). A region the compiler rejects NEVER raises
    #     here and NEVER ships silently: it degrades to ordinary plain text (no
    #     colour, no markup), is recorded in T3_STATS['failed'], and
    #     verify_explanations() quotes it verbatim so the author can Ctrl+F to it.
    #   • ⟦M:<b64>⟧ (v2.1 PRESERVE-AND-REEMIT) decodes to the original <m:oMath>
    #     serialised by the strict reader and is re-emitted byte-faithfully.
    # (Authoring-time strictness lives in guard_sentence; this render boundary is
    # the no-halt safety net.)
    if T3_OPEN in text or _OPAQUE_MATH_RE.search(text):
        pos = 0
        for mr in _MATH_TOKEN_RE.finditer(text):
            if mr.start() > pos:
                _emit_plain(paragraph, text[pos:mr.start()], bold, color, preserve)
            if mr.group('t3') is not None:             # ⟦MATH:…⟧ — Tier-3 compile
                body = mr.group('t3')
                try:
                    node = t3_compile(body)
                    paragraph._p.append(node)
                except MathCompileError as t3err:
                    T3_STATS['failed'].append((body, str(t3err)))
                    _run(paragraph, body, bold, color)
            else:                                      # ⟦M:<b64>⟧ / ⟦M⟧ — preserved
                payload = mr.group('op')
                if payload:
                    add_math(paragraph,
                             base64.b64decode(payload).decode('utf-8'))
                else:
                    # A bodiless ⟦M⟧ carries no math to re-emit; it must never
                    # reach the writer (the strict reader emits ⟦M:<b64>⟧). Fail
                    # loud rather than silently drop the math.
                    raise ValueError(
                        'bodiless opaque math token ⟦M⟧ reached the writer — '
                        f'preservation was bypassed: {text[:60]!r}')
            pos = mr.end()
        if pos < len(text):
            _emit_plain(paragraph, text[pos:], bold, color, preserve)
        return paragraph
    _emit_plain(paragraph, text, bold, color, preserve)
    return paragraph

def _emit_plain(paragraph, text, bold=False, color=None, preserve=False):
    """The v1.x plain-segment emitter, verbatim: digit/digit fractions to stacked
    OMML; raises on any inline fraction it cannot convert.

    v2.1 — preserve=True (re-emitting content the reader lifted verbatim from an
    already-shipped doc) emits an unconvertible inline fraction as the LITERAL text
    it already was, rather than raising. The raise is an AUTHORING guard that forces
    NEW math to be real OMML; it must not fire when faithfully re-emitting existing
    content (e.g. a pre-v2.0 paper that typed '1/2k' as plain text)."""
    pos = 0
    for m in _SIMPLE_FRAC.finditer(text):
        if text[pos:m.start()]:
            _run(paragraph, text[pos:m.start()], bold, color)
        add_math(paragraph, omath(frac(m.group(1), m.group(2))))
        pos = m.end()
    tail = text[pos:]
    if has_inline_fraction(tail) and not preserve:
        raise ValueError(f'unconvertible inline fraction in: {tail[:60]!r}')
    if tail:
        _run(paragraph, tail, bold, color)

# ───────────────────────────── block model ─────────────────────────────────
# ── v2.3 (GAP-2026-08-19-EXPLAIN-REPRESENTATION-EMISSION) — FIGURES ────────
# A validated container for ONE explanation-side figure. The §6A-5 contract in
# Framework_MockTestExplain applies verbatim: a figure is PROVED, not trusted.
#   path        — rendered PNG on disk (the renderer's output).
#   width_in    — display width in inches (0.5..7.0; the printable column).
#   validation  — the §6A-5 record, REQUIRED, a dict carrying at least:
#                   renderer  : str  (which renderer produced it)
#                   intended  : str  (canonical identifier of what was requested)
#                   derived   : str  (canonical identifier re-derived from output)
#                   match     : True (derived == intended; False NEVER ships)
#                 A structural renderer puts canonical SMILES in intended/derived;
#                 a level-diagram renderer puts the computed occupancy string.
#                 Molecular formula alone is INSUFFICIENT (two different answers
#                 commonly share one formula) — the identifier must be canonical.
#   after_step  — 0-based count of DEDUCTION sentences rendered BEFORE this
#                 figure (clamped to the deduction length). Default 1: the
#                 figure follows the first deduction sentence, which names what
#                 the reader is about to see.
# DESIGN INVARIANT — THE FIGURE PARAGRAPH CARRIES NO TEXT. The strict reader
# keeps an explanation paragraph only when it has display text or math source, so
# a text-free drawing paragraph is INVISIBLE to parse_solution_blocks and the
# round-trip needs no changes. There is deliberately NO caption paragraph: the
# surrounding DEDUCTION prose describes the figure (it must — §6A-1 requires the
# figure to be non-redundant WITH that prose, and prose readable without it).
# v2.6 D3 — the §6A router's requirement vocabulary. A block may carry its
# routed verdict; a VISUAL verdict with no figure raises (routing without
# emission is the defect §6A exists to remove — a §6A-4 degrade must record
# the DEGRADED requirement, never the original).
# v2.7 (d) — CONFORMER: a projection (Newman / sawhorse / chair) depicts HOW
# atoms are arranged at a given rotation, which a constitution renderer cannot
# express (run-report F3). Visual: it requires its figure like the other three.
REPRESENTATION_VERDICTS = ('PROSE', 'EQUATION', 'TABLE',
                           'STRUCTURE_GRAPH', 'LEVEL_DIAGRAM', 'DATA_PLOT', 'CONFORMER')
_VISUAL_VERDICTS = ('STRUCTURE_GRAPH', 'LEVEL_DIAGRAM', 'DATA_PLOT', 'CONFORMER')

class RepresentationFigure:
    def __init__(self, path, width_in=6.0, validation=None, after_step=1):
        self.path = str(path)
        self.width_in = float(width_in)
        self.validation = dict(validation or {})
        self.after_step = int(after_step)

    def validate(self, ctx=''):
        import os as _os
        w = f'{ctx}: ' if ctx else ''
        if not _os.path.isfile(self.path):
            raise ValueError(f'{w}figure file missing: {self.path!r}')
        if not (0.5 <= self.width_in <= 7.0):
            raise ValueError(f'{w}figure width_in {self.width_in} outside 0.5..7.0')
        v = self.validation
        for k in ('renderer', 'intended', 'derived', 'match'):
            if k not in v:
                raise ValueError(f'{w}figure validation record missing {k!r} '
                                 f'(§6A-5: a figure is proved, not trusted)')
        if v['match'] is not True:
            raise ValueError(f'{w}figure validation match is not True — a figure '
                             f'whose derived identifier differs from the intended '
                             f'one NEVER ships (§6A-5); degrade per §6A-4 instead')
        if not str(v['intended']).strip() or not str(v['derived']).strip():
            raise ValueError(f'{w}figure validation intended/derived empty')
        if str(v['intended']) != str(v['derived']):
            raise ValueError(f'{w}figure validation intended != derived but match '
                             f'claims True — inconsistent record')
        if self.after_step < 0:
            raise ValueError(f'{w}figure after_step must be >= 0')
        return True

class ExplanationBlock:
    """Validated container for one question's explanation. validate() raises on
    any structural breach BEFORE the block can be written (fail-at-construction).

    Three question types (auto-inferred or set via `qtype`):
      • 'mcq' — single correct option. ca = a 1-based int. why_wrong keys = the
                other options. (default; the common case.)
      • 'msq' — multiple correct options. ca = a set/list of 1-based ints. The last
                DEDUCTION step binds EVERY selected option; why_wrong keys = the
                NON-selected options.
      • 'nat' — numerical-answer-type, NO options. ca = the answer value (number or
                string); optional ca_range = (lo, hi) tolerance. The last DEDUCTION
                step contains the value; common_pitfalls (value -> sentences, ≥1)
                replaces why_wrong (there are no options to reject)."""
    def __init__(self, q, ca=None, axiom=None, deduction=None, speed_hack=None,
                 why_wrong=None, anomaly=None, cfg=None,
                 qtype=None, ca_range=None, common_pitfalls=None, figures=None,
                 representation_verdict=None, absolutes_justified=None,
                 transfer_record=None, error_provenance=None):
        self.q = int(q)
        self.cfg = cfg or EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4)
        self.ca = ca
        self.ca_range = ca_range
        # v2.8 (c) — formula typography is applied at construction to every
        # student-facing sentence (idempotent; math regions masked), so an
        # ASCII formula can never reach the renderer. absolutes_justified keys
        # are normalised the same way so declarations still match.
        _nf = (normalise_formula_text if (self.cfg.formula_typography) else (lambda s: s))
        self.axiom = [_nf(s) for s in (axiom or [])]
        self.deduction = [_nf(s) for s in (deduction or [])]
        self.speed_hack = [_nf(s) for s in speed_hack] if speed_hack else None
        self.why_wrong = {k: [_nf(s) for s in v] for k, v in dict(why_wrong or {}).items()}
        self.common_pitfalls = {k: [_nf(s) for s in v]
                                for k, v in dict(common_pitfalls or {}).items()}
        absolutes_justified = {_nf(k): v for k, v in dict(absolutes_justified or {}).items()}
        # v2.8 (a) — {wrong option index | pitfall value: provenance record}.
        self.error_provenance = dict(error_provenance or {})
        # v2.3 — explanation-side figures (list[RepresentationFigure], may be []).
        self.figures = list(figures or [])
        # v2.6 D3 — the §6A router verdict for this question (optional; None for
        # a legacy caller). When set, validate() enforces verdict↔emission
        # coherence: a visual verdict requires >=1 figure on this block.
        self.representation_verdict = representation_verdict
        # v2.7 (a) — {sentence: reason} declaring each absolute the author KEEPS.
        self.absolutes_justified = dict(absolutes_justified or {})
        # v2.7 (c) — the §7-7 transfer-safety record (optional; shape-validated).
        self.transfer_record = list(transfer_record) if transfer_record else None
        self.anomaly = anomaly
        # infer question type if not given
        if qtype is None:
            if self.cfg.expected_options(self.q) == 0:
                qtype = 'nat'
            elif isinstance(ca, (set, list, tuple)):
                qtype = 'msq'
            else:
                qtype = 'mcq'
        self.qtype = qtype

    def ca_set(self):
        """The selected option indices as a set (mcq -> {ca}; msq -> set(ca))."""
        if self.qtype == 'mcq':
            return {self.ca}
        return set(self.ca)

    def validate(self):
        cfg = self.cfg

        def g(s):
            # Prose guards (banned glyphs/templates/metacommentary/one-sentence/…).
            guard_sentence(s, cfg)
            # AUTHORING-TIME TIER-3 COMPILE GATE (2026.08.10.3).
            # Before this gate, a ⟦MATH:…⟧ region was compiled ONLY at render
            # (add_math_text): a region t3_compile rejected did NOT raise there —
            # by the no-halt render contract it DEGRADED to raw plain text, was
            # recorded in T3_STATS['failed'], and shipped as literal LaTeX unless
            # the producer separately consumed verify_explanations()'s RETURNED
            # (ok, problems) ledger. That returned-not-raised signal was easy to
            # drop (a harness that checks only "did the call raise" passes a
            # degraded document). validate() is the ONE universal chokepoint —
            # called on every block, every step, both pipelines, all exams — and
            # it RAISES, so gating regions here makes a malformed region fail at
            # construction, before any docx exists, and it can never reach the
            # renderer. This mirrors the NAT grading-value posture ("Fail-at-
            # construction: the primary gate; render() re-checks as defense-in-
            # depth"). Exam-agnostic: pure Tier-3 grammar, no exam value. Only
            # ⟦MATH:…⟧ (Tier-3) is compiled here; ⟦M:<b64>⟧ preserve tokens are
            # never matched by T3_REGION_RE and are untouched.
            for _body in T3_REGION_RE.findall(s):
                try:
                    t3_compile(_body)
                except MathCompileError as _t3err:
                    raise ValueError(
                        f'Q{self.q}: ⟦MATH:…⟧ region does not compile — {_t3err}. '
                        f'Region body: {_body!r}. Fix the Tier-3 syntax (§11); an '
                        f'uncompilable region would otherwise degrade to raw text '
                        f'at render.')
        if self.anomaly is not None:
            # anomaly is an INTERNAL escalation signal in Step 9 (never published).
            # v2.6: figures are student content too — an anomaly block carries none.
            if any([self.axiom, self.deduction, self.why_wrong,
                    self.speed_hack, self.common_pitfalls, self.figures]):
                raise ValueError(f'Q{self.q}: anomaly block must carry no student content')
            return True
        # v2.6 D3 — ROUTER↔EMISSION COHERENCE (§6A-3/§6A-4). The reference
        # incident: a structure-heavy paper carried 46 stem images and only 2
        # explanation figures — the router classified, the renderer shipped
        # prose. A recorded VISUAL verdict now requires the figure it promised;
        # a degrade records the degraded requirement instead.
        if self.representation_verdict is not None:
            _rv = str(self.representation_verdict)
            if _rv not in REPRESENTATION_VERDICTS:
                raise ValueError(f'Q{self.q}: representation_verdict {_rv!r} not '
                                 f'in {REPRESENTATION_VERDICTS}')
            if _rv in _VISUAL_VERDICTS and not self.figures:
                raise ValueError(
                    f'Q{self.q}: representation_verdict {_rv} declared but the '
                    f'block carries no figure — emit the routed visual or record '
                    f'the DEGRADED requirement as the verdict (§6A-3/§6A-4); '
                    f'routing without emission is the defect this gate removes')
        # AXIOM + DEDUCTION are common to all three types.
        if len(self.axiom) < 1:
            raise ValueError(f'Q{self.q}: AXIOM empty')
        for s in self.axiom:
            g(s)
        if len(self.deduction) < 2:
            raise ValueError(f'Q{self.q}: DEDUCTION needs >=2 steps')
        for s in self.deduction:
            g(s)
        if self.speed_hack is not None:
            if len(self.speed_hack) < 1:
                raise ValueError(f'Q{self.q}: SPEED HACK present but empty')
            for s in self.speed_hack:
                g(s)
        # v2.6 D4 — figures are validated for EVERY question type. Since v2.3
        # this loop sat AFTER the NAT branch's return, so the "every question
        # type" its comment promised excluded NAT: a NAT block's figures were
        # never validated at construction. Moved above the type split.
        for i, fg in enumerate(self.figures):
            if not isinstance(fg, RepresentationFigure):
                raise ValueError(f'Q{self.q}: figures[{i}] is not a '
                                 f'RepresentationFigure')
            fg.validate(ctx=f'Q{self.q} figures[{i}]')

        # v2.7 (a) — ABSOLUTE-TERM DECLARATION GATE (§8-0b made enforceable).
        # Covers the sections a learner MEMORISES (AXIOM, SPEED HACK) and the
        # ones they REMEMBER as reasons (WHY WRONG, COMMON PITFALLS). DEDUCTION
        # is item-specific working and is governed by the §7-7 protocol instead.
        _abs_re = getattr(cfg, 'absolute_terms_re', _ABSOLUTE_TERMS_RE)
        def _abs_gate(section, sents):
            for _s in sents:
                _hit = find_absolute(_s, _abs_re)
                if _hit is None:
                    continue
                _why = self.absolutes_justified.get(_s)
                if not (_why and str(_why).strip()):
                    raise ValueError(
                        f'Q{self.q}: {section} states an unqualified absolute '
                        f'{_hit!r} — a tendency written as an absolute teaches a '
                        f'false rule (§8-0b / §7-7). Either narrow it (name the '
                        f'mechanism or the condition: "only when …", "under these '
                        f'conditions …"), or, if it IS absolute in the subject\'s own '
                        f'terms (definition / conservation law / mathematical '
                        f'property), declare it in absolutes_justified with that '
                        f'reason: {_s[:70]!r}')
        _abs_gate('AXIOM', self.axiom)
        if self.speed_hack:
            _abs_gate('SPEED HACK', self.speed_hack)
        for _k, _ss in self.why_wrong.items():
            _abs_gate(f'WHY WRONG {_k}', _ss)
        for _k, _ss in self.common_pitfalls.items():
            _abs_gate(f'COMMON PITFALLS {_k}', _ss)
        for _decl in self.absolutes_justified:
            if _decl not in self.axiom and _decl not in self.deduction and \
               _decl not in (self.speed_hack or []) and \
               not any(_decl in _v for _v in self.why_wrong.values()) and \
               not any(_decl in _v for _v in self.common_pitfalls.values()):
                raise ValueError(f'Q{self.q}: absolutes_justified declares a sentence '
                                 f'that is not in the block: {_decl[:70]!r}')
        # v2.7 (c) — §7-7 transfer record, shape-validated when supplied.
        # v2.8 (b) — MANDATORY for an authored block under provenance_gates; a
        # read-back block (_preserved) carries no metadata and is exempt.
        _authoring = bool(getattr(cfg, 'provenance_gates', True)) and \
            not getattr(self, '_preserved', False)
        if self.transfer_record is None and _authoring:
            raise ValueError(
                f'Q{self.q}: transfer_record missing — the §7-7 transfer-safety '
                f'protocol must RUN and be RECORDED for every authored block '
                f'(one entry per AXIOM sentence, one per SPEED HACK); an '
                f'explanation with no record has not been tested on its neighbour')
        if self.transfer_record is not None:
            validate_transfer_record(self.transfer_record,
                                     axiom_present=bool(self.axiom),
                                     speed_hack_present=bool(self.speed_hack),
                                     ctx=f'Q{self.q}')
        if _authoring and getattr(cfg, 'learnings_triggers', None):
            # v2.8 (b) — a curated-library family named in AXIOM / SPEED HACK
            # text must be the neighbour actually tested; GENERATED is legal
            # only where no trigger fires.
            for _sec, _sents in (('AXIOM', self.axiom), ('SPEED_HACK', self.speed_hack or [])):
                _hits = find_trigger_hits(_sents, cfg.learnings_triggers)
                if not _hits:
                    continue
                _cited = {str(e.get('neighbour_source', ''))[len('CURATED:'):]
                          for e in (self.transfer_record or [])
                          if e.get('section') == _sec and
                          str(e.get('neighbour_source', '')).startswith('CURATED:')}
                _missed = [c for c in _hits if c not in _cited]
                if _missed:
                    raise ValueError(
                        f'Q{self.q}: GEN_CANONICAL_EXCEPTION_MISSED — {_sec} text matches the '
                        f'curated library family {_missed[0]} (trigger {_hits[_missed[0]]!r}) '
                        f'but no transfer_record entry for {_sec} cites '
                        f'neighbour_source CURATED:{_missed[0]}; test the claim on that '
                        f'rule\'s canonical neighbours and record it (§7-7 step 3)')
        # v2.8 (c) — residual formula gate (what the normaliser would not touch).
        if _authoring and getattr(cfg, 'formula_typography', True):
            for _sec, _sents in (('AXIOM', self.axiom), ('DEDUCTION', self.deduction),
                                 ('SPEED HACK', self.speed_hack or []),
                                 *[(f'WHY WRONG {k}', v) for k, v in self.why_wrong.items()],
                                 *[(f'COMMON PITFALLS {k}', v) for k, v in self.common_pitfalls.items()]):
                for _s in _sents:
                    _bad = find_unformatted_formula(_s)
                    if _bad:
                        raise ValueError(
                            f'Q{self.q}: FMT_UNFORMATTED_FORMULA — {_sec} carries an ASCII '
                            f'formula {_bad!r} the typography normaliser could not rewrite '
                            f'safely; write it with Unicode sub/superscripts or as a '
                            f'⟦MATH:⟧ region: {_s[:70]!r}')
        # v2.8 (a) — HEDGED PROVENANCE is banned where a wrong path is named.
        for _k, _ss in (list(self.why_wrong.items()) + list(self.common_pitfalls.items())
                        if _authoring else []):
            for _s in _ss:
                _h = find_hedge(_s)
                if _h:
                    raise ValueError(
                        f'Q{self.q}: DST_HEDGED_PROVENANCE — a WHY WRONG / COMMON PITFALLS '
                        f'line offers an unverified alternative path ({_h!r}); state ONE '
                        f'verified path or the direct contradiction, never a menu of '
                        f'guesses (§15-2): {_s[:70]!r}')

        last = self.deduction[-1]
        opt_label = cfg.labels["option"]

        if self.qtype == 'nat':
            # NAT: no options, a typed value, value-bound deduction, common_pitfalls.
            if cfg.expected_options(self.q) not in (0, None):
                raise ValueError(f'Q{self.q}: NAT must have 0 expected options, '
                                 f'got {cfg.expected_options(self.q)}')
            # v1.16: ca may be None ONLY when ca_range is present — that is the
            # reader-reconstruction case (parse_solution_blocks rebuilding a
            # block from a rendered 'lo-hi' Correct-Answer line, where the
            # single point value is no longer recoverable from that line by
            # design: the range format replaces it entirely, it doesn't append
            # to it). Author-constructed blocks (Step 7/9, building fresh) must
            # still always supply ca — this does not relax authoring.
            if self.ca_range is None:
                if self.ca is None or str(self.ca).strip() == '':
                    raise ValueError(f'Q{self.q}: NAT answer value missing')
                # v1.16: portal-safe grading charset (0-9 . - only), independent
                # of nat_answer_type — catches scientific notation, unit
                # suffixes, en-dashes, spaces, parentheses BEFORE the block can
                # be written. Fail-at-construction: this is the primary gate;
                # render() re-checks as defense-in-depth, not the only line
                # of defense.
                validate_nat_grading_value(self.ca, ctx=f'Q{self.q}')
                if str(self.ca) not in last:
                    raise ValueError(f'Q{self.q}: DEDUCTION last step must contain '
                                     f'the answer value {str(self.ca)!r}')
            else:
                lo, hi = self.ca_range
                # raises on bad charset / negative bound (NOT SUPPORTED) / lo>hi
                format_nat_range(lo, hi, ctx=f'Q{self.q}')
                if self.ca is not None:
                    # author-constructed range block: ca is the underlying point
                    # value the range was built around (V ± tolerance) — still
                    # charset-validated and still must be traceable in DEDUCTION,
                    # exactly as before this change, unaffected by the display
                    # format switch.
                    validate_nat_grading_value(self.ca, ctx=f'Q{self.q}')
                    if str(self.ca) not in last:
                        raise ValueError(f'Q{self.q}: DEDUCTION last step must '
                                         f'contain the answer value {str(self.ca)!r}')
            if self.why_wrong:
                raise ValueError(f'Q{self.q}: NAT uses common_pitfalls, not why_wrong')
            if len(self.common_pitfalls) < 1:
                raise ValueError(f'Q{self.q}: NAT needs >=1 common pitfall')
            for v, sents in self.common_pitfalls.items():
                if len(sents) < 1:
                    raise ValueError(f'Q{self.q}: common pitfall {v!r} empty')
                for s in sents:
                    g(s)
            if self.ca_range is not None:
                lo, hi = self.ca_range
                if not (lo <= hi):
                    raise ValueError(f'Q{self.q}: ca_range {self.ca_range} not lo<=hi')
            if _authoring:
                validate_error_provenance(self.error_provenance,
                                          [str(k) for k in self.common_pitfalls],
                                          'COMMON PITFALLS', ctx=f'Q{self.q}')
            return True

        # MCQ / MSQ share the option machinery.
        n = cfg.expected_options(self.q)
        if not n or n < 1:
            raise ValueError(f'Q{self.q}: {self.qtype} needs >=1 option, expected={n}')
        if self.common_pitfalls:
            raise ValueError(f'Q{self.q}: {self.qtype} uses why_wrong, not common_pitfalls')
        sel = self.ca_set()
        if not sel or not sel <= set(range(1, n + 1)):
            raise ValueError(f'Q{self.q}: selected {sel} not a non-empty subset of 1..{n}')
        if self.qtype == 'mcq' and len(sel) != 1:
            raise ValueError(f'Q{self.q}: mcq must select exactly one option, got {sel}')
        # DEDUCTION last step must bind EVERY selected option (label-bounded).
        # v2.5 (GAP-2026-08-19-UNSATISFIABLE-LABEL-BINDING): the trailing guard was
        # \b, which requires a WORD character on one side of the boundary. A label
        # ENDING in a non-word character — '(A)', '[A]', 'A)', any parenthesised or
        # bracketed custom scheme — can therefore NEVER satisfy it, whatever the
        # author writes. The check was UNSATISFIABLE for such an exam: every block
        # raised at construction and no paper could be produced at all. This spec
        # advertises custom label schemes, and real section_rules files declare
        # option_label_format '(A)/(B)/(C)/(D)', so the case is reachable, not
        # theoretical. Found by executing the pipeline across an exam-shape matrix;
        # no amount of reading the code or the spec surfaced it.
        # FIX: (?!\w) — "the label is not immediately continued by another word
        # character". Identical behaviour for word-final labels (D, iv, 1) and it
        # still prevents the Option 1 / Option 10 confusion the self-test locks,
        # because there the next character IS a word character. For labels ending
        # in punctuation it is now satisfiable, which \b never was.
        for i in sorted(sel):
            lab = cfg.option_label(i)
            if not re.search(rf'\b{re.escape(opt_label)}\s+{re.escape(lab)}(?!\w)', last, re.I):
                raise ValueError(f'Q{self.q}: DEDUCTION last step must bind '
                                 f'{opt_label!r} {lab} (word-bounded)')
        # v2.6 D2 — the AXIOM states the transferable principle; the answer
        # binding belongs to the DEDUCTION's last step (§8-2). An AXIOM that
        # names an option has leaked the conclusion into the principle.
        for _s_ax in self.axiom:
            for _i in range(1, n + 1):
                _lab = cfg.option_label(_i)
                if re.search(rf'\b{re.escape(opt_label)}\s+{re.escape(_lab)}(?!\w)',
                             _s_ax, re.I):
                    raise ValueError(
                        f'Q{self.q}: AXIOM names {opt_label!r} {_lab} — the '
                        f'AXIOM states the transferable principle; answer '
                        f'binding belongs to the DEDUCTION last step (§8-2)')
        # WHY WRONG: keys == exactly the NON-selected options.
        expected = set(range(1, n + 1)) - sel
        if set(self.why_wrong.keys()) != expected:
            raise ValueError(f'Q{self.q}: WHY WRONG keys {set(self.why_wrong)} '
                             f'!= wrong options {expected}')
        for k, sents in self.why_wrong.items():
            if len(sents) < 1:
                raise ValueError(f'Q{self.q}: WHY WRONG option {k} empty')
            for s in sents:
                g(s)
        if _authoring:
            validate_error_provenance(self.error_provenance, list(self.why_wrong),
                                      'WHY WRONG', ctx=f'Q{self.q}')
        return True

# ───────────────────────── paragraph construction ──────────────────────────
def _new_para(cfg, kind, text=None, runs=None, before=0, after=120,
              bold=False, color=None, math=True, preserve=False):
    """Create a standalone <w:p> element (not yet attached)."""
    p = OxmlElement('w:p')
    ppr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before)); sp.set(qn('w:after'), str(after))
    ppr.append(sp); p.append(ppr)
    # Wrap in a transient Paragraph proxy to reuse run helpers.
    from docx.text.paragraph import Paragraph
    proxy = Paragraph(p, None)
    if text is not None:
        if math:
            add_math_text(proxy, text, bold=bold, color=color, cfg=cfg,
                          preserve=preserve)
        else:
            _run(proxy, text, bold=bold, color=color)
    return p

def _header_para(cfg, key, before=240, after=120):
    marker = cfg.markers.get(key, '')
    label = cfg.labels.get(key, key.upper())
    txt = f'{marker} {label}'.strip()
    return _new_para(cfg, 'hdr', text=txt, before=before, after=after,
                     bold=True, color=cfg.colors['hdr'], math=False)

def _figure_para(doc_part, fig):
    """v2.3 — build a standalone, CENTERED <w:p> holding ONE inline picture and
    NO text (the invariant that keeps figures invisible to the strict reader).
    The image part + relationship are registered on doc_part, so verify_fidelity's
    dangling-rId check (A3) resolves. EMU: 914400 per inch."""
    rid, image = doc_part.get_or_add_image(fig.path)
    cx = int(fig.width_in * 914400)
    # preserve aspect ratio from the image's native size
    try:
        cy = int(cx * (image.px_height / image.px_width))
    except Exception:
        cy = cx
    p = OxmlElement('w:p')
    ppr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '80'); sp.set(qn('w:after'), '80')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    ppr.append(sp); ppr.append(jc); p.append(ppr)
    did = abs(hash((fig.path, rid))) % 2147483647 or 1
    drawing_xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:drawing xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:docPr id="{did}" name="RepresentationFigure"/>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:nvPicPr><pic:cNvPr id="{did}" name="RepresentationFigure"/>'
        '<pic:cNvPicPr/></pic:nvPicPr>'
        '<pic:blipFill>'
        f'<a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:embed="{rid}"/>'
        '<a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        '<pic:spPr><a:xfrm><a:off x="0" y="0"/>'
        f'<a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        '</pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r>')
    p.append(parse_xml(drawing_xml))
    return p

def _block_paragraphs(cfg, blk, doc_part=None):
    """Render one ExplanationBlock to an ordered list of <w:p> elements, shaped to
    the question type (mcq / msq / nat). v2.3: when doc_part is given and the
    block carries figures, each figure is emitted as a text-free centred picture
    paragraph after its after_step-th DEDUCTION sentence; without doc_part
    (legacy callers) figures are SKIPPED, never half-rendered."""
    # v2.1 — blocks lifted by the reader from an already-shipped doc re-emit their
    # verbatim content faithfully (a pre-v2.0 plain-text fraction is kept as text,
    # not raised on). Author-constructed blocks (Step 9 / self-test) keep the
    # default strict authoring behaviour; their new prose carries proper OMML and
    # is unaffected by this flag.
    pres = getattr(blk, '_preserved', False)
    out = []
    if blk.anomaly is not None:
        # Should never be written in Step 9 (halt-and-escalate); guarded here.
        raise ValueError(f'Q{blk.q}: anomaly must not be rendered in Step 9')
    ca_label = cfg.labels['correct_answer']
    opt_word = cfg.labels['option']
    # ── Correct Answer line (shaped by type; index/label only, no option text) ──
    if blk.qtype == 'nat':
        # v1.16: portal grading value must be EXACTLY '0-9 . -', nothing else.
        # A range REPLACES the point display entirely ('lo-hi', no parenthetical
        # 'accepted range' wording, no en-dash) — the portal's Range validator
        # reads the whole field as one delimited pair, not a value + a note.
        if blk.ca_range is not None:
            lo, hi = blk.ca_range
            val = format_nat_range(lo, hi, ctx=f'Q{blk.q}')
        else:
            val = validate_nat_grading_value(blk.ca, ctx=f'Q{blk.q}')
        # math=True so a fractional answer renders as OMML
        out.append(_new_para(cfg, 'ca', text=f'{ca_label}: {val}', before=120,
                             after=120, bold=True, color=cfg.colors['ca'], math=True, preserve=pres))
    else:
        disp = ', '.join(cfg.option_label(i) for i in sorted(blk.ca_set()))
        out.append(_new_para(cfg, 'ca', text=f'{ca_label}: {disp}', before=120,
                             after=120, bold=True, color=cfg.colors['ca'], math=False))
    out.append(_header_para(cfg, 'axiom'))
    for s in blk.axiom:
        out.append(_new_para(cfg, 'sent', text=s, before=0, after=120,
                             color=cfg.colors['sent'], preserve=pres))
    out.append(_header_para(cfg, 'deduction'))
    # v2.3 — figures interleave with DEDUCTION sentences at their after_step
    # positions (clamped). Figure paragraphs carry NO text by invariant.
    _figs_at = {}
    if doc_part is not None and getattr(blk, 'figures', None):
        for _fg in blk.figures:
            _pos = min(max(_fg.after_step, 0), len(blk.deduction))
            _figs_at.setdefault(_pos, []).append(_fg)
    for _fg in _figs_at.get(0, ()):
        out.append(_figure_para(doc_part, _fg))
    for _si, s in enumerate(blk.deduction, start=1):
        out.append(_new_para(cfg, 'sent', text=s, before=0, after=120,
                             color=cfg.colors['sent'], preserve=pres))
        for _fg in _figs_at.get(_si, ()):
            out.append(_figure_para(doc_part, _fg))
    if blk.speed_hack:
        out.append(_header_para(cfg, 'speed_hack'))
        for s in blk.speed_hack:
            out.append(_new_para(cfg, 'sent', text=s, before=0, after=120,
                                 color=cfg.colors['sent'], preserve=pres))
    if blk.qtype == 'nat':
        # COMMON PITFALLS (wrong values), the NAT analogue of WHY WRONG
        out.append(_header_para(cfg, 'common_pitfalls'))
        for v in blk.common_pitfalls:
            out.append(_new_para(cfg, 'sub', text=str(v), before=160, after=40,
                                 bold=True, color=cfg.colors['sub'], math=True, preserve=pres))
            for s in blk.common_pitfalls[v]:
                out.append(_new_para(cfg, 'sent', text=s, before=0, after=120,
                                     color=cfg.colors['sent'], preserve=pres))
    else:
        out.append(_header_para(cfg, 'why_wrong'))
        for k in sorted(blk.why_wrong):
            out.append(_new_para(cfg, 'sub', text=f'{opt_word} {cfg.option_label(k)}',
                                 before=160, after=40, bold=True,
                                 color=cfg.colors['sub'], math=False))
            for s in blk.why_wrong[k]:
                out.append(_new_para(cfg, 'sent', text=s, before=0, after=120,
                                     color=cfg.colors['sent'], preserve=pres))
    # trailing blank separator
    out.append(_new_para(cfg, 'blank', text=None, before=0, after=0))
    return out

# ───────────────────────────── paper parsing ───────────────────────────────
def parse_paper(path, cfg):
    """Read the Step-3 paper. Returns (doc, qmap) where qmap[q] = {
       'q_para': idx, 'opt_paras': [idx...], 'last_anchor': idx }.
    Validates: questions ascending+contiguous from 1, and each question carries the
    option count expected for it (cfg.expected_options(q) — a single uniform count,
    a per-section map, or 0 for NAT questions that have no options)."""
    doc = Document(path)
    paras = doc.paragraphs
    qmap = {}
    cur = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        mq = cfg.q_re.match(t)
        if mq:
            cur = int(mq.group(1))
            qmap[cur] = {'q_para': i, 'opt_paras': [], 'last_anchor': i}
            continue
        if cur is not None and cfg.opt_re.match(t):
            qmap[cur]['opt_paras'].append(i)
            qmap[cur]['last_anchor'] = i
    if not qmap:
        raise ValueError('no questions matched q_re — check exam config')
    nums = sorted(qmap)
    if nums != list(range(1, len(nums) + 1)):
        raise ValueError(f'question numbers not contiguous from 1: {nums[:5]}...')
    for q in nums:
        oc = len(qmap[q]['opt_paras'])
        exp = cfg.expected_options(q)
        if exp is None:
            continue                      # no expectation supplied → skip count check
        if oc != exp:
            raise ValueError(f'Q{q} has {oc} options, expected {exp}')
    return doc, qmap

# ─────────────────────── interleaved (append-only) build ────────────────────
def build_interleaved_docx(source_path, blocks, out_path, cfg):
    """Seed the WHOLE source paper, then APPEND each block's paragraphs after that
    question's ENTIRE region — i.e. after the last body element (paragraph OR table)
    before the next question's stem (or end of document). Anchoring at the region
    end (not the last option label) is essential: figural questions place an option
    figure in a paragraph AFTER its 'N.' label, and DI questions end in a table —
    both must stay inside the question region, ahead of the appended explanation.
    Question regions are NEVER modified. Returns the count of questions explained."""
    # Validate the paper (option counts, contiguity) up front.
    parse_paper(source_path, cfg)
    for _w in source_math_health(source_path, cfg):
        print('  ⚠️ ' + _w)
    doc = Document(source_path)
    body = doc.element.body
    children = [c for c in body.iterchildren()
                if c.tag in (qn('w:p'), qn('w:tbl'))]
    q_start = {}
    for idx, c in enumerate(children):
        if c.tag == qn('w:p'):
            txt = ''.join(t.text or '' for t in c.iter(qn('w:t'))).strip()
            mq = cfg.q_re.match(txt)
            if mq:
                q_start[int(mq.group(1))] = idx
    nums = sorted(q_start)
    by_q = {}
    for b in blocks:
        b.validate()
        by_q[b.q] = b
    for q, blk in by_q.items():
        if q not in q_start:
            raise ValueError(f'block for Q{q} but Q{q} not in paper')
        start = q_start[q]
        later = [q_start[n] for n in nums if q_start[n] > start]
        end = (min(later) - 1) if later else (len(children) - 1)
        cursor = children[end]               # last element of this question's region
        for pel in _block_paragraphs(cfg, blk, doc_part=doc.part):
            cursor.addnext(pel)
            cursor = pel
    doc.save(out_path)
    return len(by_q)

# ───────────────────────── content-fidelity verifier ───────────────────────
def _qregion_signature(path, cfg):
    """Per-question signature of the QUESTION REGION only (stem+options, in
    document order up to but excluding any appended Correct-Answer line):
       text lines, OMML <m:t> sequence, drawing count, table-cell grids."""
    doc = Document(path)
    body = doc.element.body
    ca_label = cfg.labels['correct_answer'].lower()
    sig = {}
    cur = None
    region = None
    def flush():
        if cur is not None:
            sig[cur] = region
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn('w:p'):
            txt = ''.join(t.text or '' for t in child.iter(qn('w:t')))
            mt = [t.text or '' for t in child.iter(qn('m:t'))]
            ndraw = (len(child.findall('.//' + qn('w:drawing'))) +
                     len(child.findall('.//' + qn('w:pict'))))
            st = txt.strip()
            mq = cfg.q_re.match(st)
            if mq:
                flush()
                cur = int(mq.group(1)); region = {'lines': [], 'omml': [], 'draw': 0, 'tables': []}
            if cur is None:
                continue
            if st.lower().startswith(ca_label + ':'):
                # appended explanation begins → stop collecting this region
                # (mark closed by switching cur off until next question)
                flush(); cur = None; region = None
                continue
            region['lines'].append(txt)
            region['omml'].extend(mt)
            region['draw'] += ndraw
        elif tag == qn('w:tbl') and cur is not None and region is not None:
            grid = []
            for row in child.findall(qn('w:tr')):
                cells = []
                for cell in row.findall(qn('w:tc')):
                    cells.append(''.join(t.text or '' for t in cell.iter(qn('w:t'))))
                grid.append(cells)
            region['tables'].append(grid)
    flush()
    return sig

def _media_md5(path):
    import zipfile
    out = {}
    with zipfile.ZipFile(path) as z:
        for n in z.namelist():
            if n.startswith('word/media/'):
                out[n] = hashlib.md5(z.read(n)).hexdigest()
    return out

def _rel_ids(path):
    """Set of relationship Ids declared in word/_rels/document.xml.rels."""
    import zipfile
    with zipfile.ZipFile(path) as z:
        try:
            root = parse_xml(z.read('word/_rels/document.xml.rels'))
        except KeyError:
            return set()
    return {r.get('Id') for r in root if r.get('Id')}

def _embed_ids(path):
    """Every r:embed / r:id image reference in word/document.xml."""
    import zipfile
    with zipfile.ZipFile(path) as z:
        root = parse_xml(z.read('word/document.xml'))
    ids = set()
    for el in root.iter():
        for a in (qn('r:embed'), qn('r:id'), qn('r:link')):
            v = el.get(a)
            if v:
                ids.add(v)
    return ids

def verify_fidelity(out_path, source_path, cfg):
    """Confirm every QUESTION REGION in the output is byte-identical to the
    source: stem/option text, OMML m:t sequence, drawing counts, table grids,
    that every source media part survives MD5-identical, and that every image
    relationship referenced in the body actually resolves (no dangling rId, A3).
    Returns (ok: bool, problems: list[str])."""
    problems = []
    src = _qregion_signature(source_path, cfg)
    got = _qregion_signature(out_path, cfg)
    def _rstrip_blanks(lines):
        out = list(lines)
        while out and not out[-1].strip():
            out.pop()
        return out
    for q in src:
        if q not in got:
            problems.append(f'Q{q}: region missing in output'); continue
        a, b = src[q], got[q]
        if _rstrip_blanks(a['lines']) != _rstrip_blanks(b['lines']):
            problems.append(f'Q{q}: stem/option text changed')
        if a['omml'] != b['omml']:
            problems.append(f'Q{q}: OMML math content changed')
        if a['draw'] != b['draw']:
            problems.append(f'Q{q}: drawing count changed {a["draw"]}->{b["draw"]}')
        if a['tables'] != b['tables']:
            problems.append(f'Q{q}: table cell grid changed')
    sm, gm = _media_md5(source_path), _media_md5(out_path)
    for n, h in sm.items():
        if gm.get(n) != h:
            problems.append(f'media {n}: MD5 changed or dropped')
    dangling = _embed_ids(out_path) - _rel_ids(out_path)
    if dangling:
        problems.append(f'dangling image relationship ids (no rel): {sorted(dangling)[:6]}')
    return (not problems), problems

# ───────────────────────── structure / coverage verifier ───────────────────
def verify_structure(out_path, blocks, cfg, expected_qs=None):
    """Confirm coverage and block integrity. Doc-level: every expected question
    carries a Correct-Answer line, and NO question outside expected_qs does (no
    look-ahead). Object-level: every expected block re-passes ExplanationBlock
    .validate() (header order, the CA three-way binding, WHY-WRONG key set), which
    — combined with the deterministic _block_paragraphs build — guarantees the
    written structure. Returns (ok: bool, problems: list[str])."""
    problems = []
    by_q = {b.q: b for b in blocks}
    if expected_qs is None:
        expected_qs = sorted(by_q)
    expected = set(expected_qs)
    doc = Document(out_path)
    paras = [p.text.strip() for p in doc.paragraphs]
    ca_label = cfg.labels['correct_answer'].lower()
    explained = set()
    cur = None
    for t in paras:
        mq = cfg.q_re.match(t)
        if mq:
            cur = int(mq.group(1)); continue
        if cur is not None and t.lower().startswith(ca_label + ':'):
            explained.add(cur)
    missing = expected - explained
    extra = explained - expected
    if missing:
        problems.append(f'missing explanations for {sorted(missing)}')
    if extra:
        problems.append(f'look-ahead: explained {sorted(extra)} beyond batch')
    # per-block structural re-validation (defence in depth)
    for q in sorted(expected):
        if q in by_q:
            try:
                by_q[q].validate()
            except ValueError as e:
                problems.append(str(e))
    return (not problems), problems

# ────────────────── post-render explanation audit (independent) ─────────────
def _para_mtext(p_el):
    return ''.join(t.text or '' for t in p_el.iter(qn('m:t')))

def _para_prose(p_el):
    """v2.1 — MATH-AWARE prose projection: literal w:r text verbatim, every OMML
    node collapsed to a single ∂M∂ placeholder. verify_explanations() runs its
    inline-fraction and one-sentence prose guards on THIS, not on p.text: a slash
    or decimal point that is part of a text+OMML expression ('10' sup '/' '10' sup,
    or '1.' + an OMML fraction-part) is no longer mis-read as a literal inline
    fraction or a sentence break once the OMML it belongs to is visible as ∂M∂ —
    while a genuinely literal inline fraction (no adjacent OMML, e.g. '1/2k') is
    still caught. This mirrors guard_sentence's region-aware handling; p.text,
    which silently drops OMML, false-positived on every mixed expression."""
    out = []
    for c in p_el:
        if c.tag == qn('m:oMath'):
            out.append('\u2202M\u2202')
        elif c.tag == qn('w:pPr'):
            continue
        else:
            out.append(''.join(t.text or '' for t in c.iter(qn('w:t'))))
    return ''.join(out)

def verify_explanations(out_path, blocks, cfg, expected_qs=None):
    """Independent POST-RENDER audit (A2, parity with T2 verify_master 10-17): it
    re-parses the RENDERED docx — NOT the in-memory blocks — and re-checks the
    explanation region from the bytes that were actually written. Catches anything
    a build bug or future renderer change could let slip past the
    construction-time guards. Returns (ok, problems).

    Per question it confirms: header order AXIOM → DEDUCTION →
    (SPEED HACK) → WHY WRONG/COMMON PITFALLS; the last DEDUCTION line binds the
    answer (option / set / value, type-aware); WHY WRONG covers exactly the wrong options (mcq/msq)
    or COMMON PITFALLS the wrong values (nat); zero banned glyphs/metacommentary/
    templates/fake-cites/banned-blocks and zero inline or vulgar fractions in the
    rendered prose; one sentence per rendered prose paragraph. Document-wide it
    confirms every OMML fraction has a non-empty numerator AND denominator built
    from RUN-LEVEL children (bare text = schema-invalid, v2.0), that no fraction
    is a consecutive-year artifact, Tier-3 structural integrity (sSubSup/rad/
    nary/limLow complete, matrices rectangular), zero ASCII-dialect residue in
    rendered explanation prose, and quotes every degraded ⟦MATH:⟧ region
    verbatim with its remedy (v2.0)."""
    problems = []
    by_q = {b.q: b for b in blocks}
    if expected_qs is None:
        expected_qs = sorted(by_q)
    expected = sorted(set(expected_qs))
    doc = Document(out_path)

    # v2.3 — count drawings per EXPLANATION region (question regions excluded):
    # walked once here so the figure-landing check below is O(doc), not O(q*doc).
    _expl_drawings = {}
    _cur = 0; _in_expl = False
    _ca_low = cfg.labels['correct_answer'].lower()
    for _p in doc.paragraphs:
        _t = _p.text.strip()
        _mq = cfg.q_re.match(_t)
        if _mq:
            _cur = int(_mq.group(1)); _in_expl = False; continue
        if _cur and _t.lower().startswith(_ca_low + ':'):
            _in_expl = True; continue
        if _cur and _in_expl:
            _nd = sum(1 for _ in _p._p.iter(qn('w:drawing')))
            if _nd:
                _expl_drawings[_cur] = _expl_drawings.get(_cur, 0) + _nd

    ca_label = cfg.labels['correct_answer']
    opt_word = cfg.labels['option']
    H = {k: f"{cfg.markers.get(k, '')} {cfg.labels.get(k, k)}".strip()
         for k in ('axiom', 'deduction', 'speed_hack', 'why_wrong', 'common_pitfalls')}
    HEADERS = set(H.values())

    # segment rendered paragraphs into per-question explanation line lists
    segs = {}
    ca_texts = {}   # v1.17: the ACTUAL rendered 'Correct Answer: X' value per question,
                     # captured (not just consumed as a boundary marker) so it can be
                     # independently re-validated against the portal charset below —
                     # this checks what was WRITTEN, not what the in-memory block claims.
    cur = None
    in_expl = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if cfg.q_re.match(t):
            cur = int(cfg.q_re.match(t).group(1)); in_expl = False; continue
        if cur is None:
            continue
        if t.startswith(ca_label + ':'):
            in_expl = True; segs[cur] = []
            ca_texts[cur] = t[len(ca_label) + 1:].strip()
            continue
        if in_expl:
            _s = _para_source(p._p, cur, strict=False).strip()
            if t or _s:                          # keep OMML-only paras (fraction value headers)
                segs[cur].append((t, _para_mtext(p._p), _s, _para_prose(p._p)))

    for q in expected:
        b = by_q.get(q)
        if q not in segs:
            problems.append(f'Q{q}: no rendered explanation found'); continue
        # v2.3 — FIGURE LANDING (GAP-2026-08-19-EXPLAIN-REPRESENTATION-EMISSION):
        # a block that declares figures must have EXACTLY that many drawing
        # paragraphs rendered in its explanation region. A missing drawing means
        # the render path silently skipped a figure (e.g. a legacy caller without
        # doc_part) — the §6A-4 rule is degrade LOUDLY, so silence is a FAIL here.
        _blk_figs = len(getattr(by_q[q], 'figures', []) or [])
        _got_figs = _expl_drawings.get(q, 0)
        if _blk_figs != _got_figs:
            problems.append(f'Q{q}: figure landing mismatch — block declares '
                            f'{_blk_figs} figure(s), explanation region renders '
                            f'{_got_figs}')
        lines = segs[q]
        texts = [t for t, _, _, _ in lines]
        srcs = [s for _, _, s, _ in lines]
        # v1.17: independent portal-charset check on the RENDERED Correct-Answer text
        # (RXA-2 / RXA-CHARSET) — runs regardless of whether any other lane flags this
        # question for rectification, so a bad grading string can never silently pass
        # through unexamined just because the surrounding prose looked fine.
        if b and b.qtype == 'nat':
            ca_rendered = ca_texts.get(q)
            if ca_rendered is None:
                problems.append(f'Q{q}: no rendered Correct Answer text captured')
            else:
                bad_chars = sorted(set(ca_rendered) - _NAT_CHARSET_ALLOWED)
                is_range = bool(re.fullmatch(r'\d+(?:\.\d+)?-\d+(?:\.\d+)?', ca_rendered))
                is_point = bool(_NAT_POINT_RE.match(ca_rendered))
                if bad_chars:
                    problems.append(f'Q{q}: rendered Correct Answer {ca_rendered!r} has '
                                     f'banned character(s) {bad_chars} — portal charset is '
                                     f'exactly "0123456789.-"')
                elif not (is_range or is_point):
                    problems.append(f'Q{q}: rendered Correct Answer {ca_rendered!r} is not '
                                     f'a well-formed plain number or lo-hi range')
        seq = [t for t in texts if t in HEADERS]
        # 1. header order (SPEED HACK is positional/optional; core = the rest)
        core = [h for h in seq if h != H['speed_hack']]
        wrong_hdr = H['common_pitfalls'] if (b and b.qtype == 'nat') else H['why_wrong']
        if core != [H['axiom'], H['deduction'], wrong_hdr]:
            problems.append(f'Q{q}: header order/presence wrong: {core}')
        # 2. DEDUCTION last line binds the answer (type-aware), read from render
        if H['deduction'] in texts:
            di = texts.index(H['deduction'])
            stops = [i for i, t in enumerate(texts)
                     if i > di and t in (H['speed_hack'], wrong_hdr)]
            last_i = (stops[0] - 1) if stops else len(texts) - 1
            last = texts[last_i] if last_i > di else ''
            last_src = srcs[last_i] if last_i > di else ''
            if b is not None:
                if b.qtype == 'nat':
                    # v1.16: a range-type block with no point value (ca is None
                    # — the reader-reconstruction case; see ExplanationBlock.
                    # validate()) has nothing to bind-check here. The band
                    # itself round-trips via ca_range, checked elsewhere.
                    if b.ca is not None and str(b.ca) not in last_src:
                        problems.append(f'Q{q}: DEDUCTION last line does not bind value {b.ca!r}')
                else:
                    need = [f'{opt_word} {cfg.option_label(i)}' for i in sorted(b.ca_set())]
                    # v2.5 — SAME FIX AS validate() (GAP-2026-08-19-UNSATISFIABLE-
                    # LABEL-BINDING), applied at the SECOND site. The construction-time
                    # check and this post-render verifier each carried their own copy of
                    # the trailing \b, so fixing one left the other unsatisfiable and the
                    # exam still could not ship — construction passed, verification
                    # failed. Two sites, one defect: the exam-shape matrix caught the
                    # second only because it runs the whole pipeline, not just validate().
                    miss = [n for n in need if not re.search(rf'{re.escape(n)}(?!\w)', last)]
                    if miss:
                        problems.append(f'Q{q}: DEDUCTION last line missing binding {miss}')
        # 3. wrong-section coverage from render
        if b is not None:
            if b.qtype == 'nat':
                subs = {s for s in srcs if s in {str(v) for v in b.common_pitfalls}}
                if subs != {str(v) for v in b.common_pitfalls}:
                    problems.append(f'Q{q}: COMMON PITFALLS render coverage mismatch')
            else:
                subs = set()
                for t in texts:
                    mo = re.match(rf'^{re.escape(opt_word)}\s+(\S+)', t)
                    if mo and t not in HEADERS:
                        subs.add(mo.group(1))
                need = {cfg.option_label(k) for k in b.why_wrong}
                if subs != need:
                    problems.append(f'Q{q}: WHY WRONG render coverage {sorted(subs)} != {sorted(need)}')
        # 4. prose guards on rendered sentence lines (exclude headers + sub-headers)
        sub_pat = re.compile(rf'^{re.escape(opt_word)}\s+\S')
        nat_keys = {str(v) for v in (b.common_pitfalls if b else {})}
        for t, _mt, _s, _mp in lines:
            if t in HEADERS or sub_pat.match(t) or _s in nat_keys or not t:
                continue
            for g in _BANNED_GLYPHS:
                if g in t:
                    problems.append(f'Q{q}: rendered banned glyph {g!r}')
            for v in VULGAR:
                if v in t:
                    problems.append(f'Q{q}: rendered vulgar fraction {v!r}')
            lw = t.lower()
            meta_re = cfg.metacommentary_re if hasattr(cfg, 'metacommentary_re') else _META_RE
            b_templates = cfg.banned_templates if hasattr(cfg, 'banned_templates') else _BANNED_TEMPLATE
            b_fakecites = cfg.banned_fakecites if hasattr(cfg, 'banned_fakecites') else _BANNED_FAKECITE
            if meta_re.search(t) or any(x in lw for x in b_templates) \
               or any(x in lw for x in b_fakecites):
                problems.append(f'Q{q}: rendered banned phrase in {t[:40]!r}')
            # v2.6 D1 read-back — the write-time guard is the primary gate; this
            # re-scan of the RENDERED bytes is the belt to that braces.
            _mtag = _INTERNAL_TAG_RE.search(t)
            if _mtag:
                problems.append(f'Q{q}: rendered internal error-taxonomy token '
                                f'{_mtag.group(0)!r} in {t[:40]!r}')
            if has_inline_fraction(_mp):
                problems.append(f'Q{q}: rendered inline fraction in {t[:40]!r}')
            if sentence_count(_mp, cfg.sentence_terminators) != 1:
                problems.append(f'Q{q}: rendered multi-sentence paragraph {t[:40]!r}')

    # 5. document-wide OMML fraction well-formedness + year-range artifact
    import zipfile
    root = parse_xml(zipfile.ZipFile(out_path).read('word/document.xml'))
    for f in root.iter(qn('m:f')):
        num = f.find(qn('m:num')); den = f.find(qn('m:den'))
        nt = ''.join(num.itertext()) if num is not None else ''
        dt = ''.join(den.itertext()) if den is not None else ''
        if not nt.strip() or not dt.strip():
            problems.append(f'malformed OMML fraction (empty num/den): {nt!r}/{dt!r}')
        elif _year_range_hit(f'{nt}/{dt}'):
            problems.append(f'year-range rendered as OMML fraction: {nt}/{dt}')
        # v2.0 STRICT SHAPE: itertext() reads bare text nodes that Word CANNOT
        # render — the exact loosening that let 12 destroyed fractions ship under
        # green checks. num/den content must live in run-level children.
        for part, name in ((num, 'num'), (den, 'den')):
            if part is None:
                continue
            if (part.text or '').strip() or not len(part):
                problems.append(
                    f'SCHEMA-INVALID fraction m:{name} carries bare text '
                    f'{(part.text or "").strip()!r} — Word renders it EMPTY; '
                    f'content must be <m:r><m:t> runs (builders now enforce this)')

    # 6. v2.0 — Tier-3 structural integrity (parity with PYQPrepare CHECK 10 v2)
    for el in root.iter(qn('m:sSubSup')):
        if any(el.find(qn(t)) is None for t in ('m:e', 'm:sub', 'm:sup')):
            problems.append('broken m:sSubSup (missing e/sub/sup)')
    for tag in ('m:rad', 'm:nary', 'm:limLow'):
        for el in root.iter(qn(tag)):
            if el.find(qn('m:e')) is None:
                problems.append(f'broken {tag} (missing m:e)')
    for el in root.iter(qn('m:m')):
        widths = {len(mr_.findall(qn('m:e'))) for mr_ in el.findall(qn('m:mr'))}
        if len(widths) != 1:
            problems.append('ragged OMML matrix (unequal row widths)')

    # 7. v2.0 — ASCII-math dialect residue in RENDERED prose <w:t> runs of the
    # explanation regions (the evasion channel, now detected post-render too).
    _failed_bodies = {b for b, _ in T3_STATS.get('failed', [])}
    for q, seg in _expl_segments(out_path, cfg).items():
        for t in seg:
            if any(fb and fb in t for fb in _failed_bodies):
                continue      # degraded region — reported by check 8 verbatim
            for pat, what, fixit in _DIALECT_BANS:
                if pat.search(t):
                    problems.append(f'Q{q}: rendered {what} in {t[:50]!r} — '
                                    f'should be a ⟦MATH:…⟧ region')
                    break

    # 8. v2.0 — Tier-3 degrade report (plain operator words). The ledger is the
    # renderer itself: every ⟦MATH:⟧ region either compiled (t3_compile) or was
    # recorded here — add_math_text has no third path, and guard_sentence rejects
    # unbalanced delimiters at authoring time. T3_STATS is session-cumulative:
    # entries are quoted, never counted against a single build.
    for body, reason in T3_STATS.get('failed', []):
        snippet = body if len(body) <= 60 else body[:57] + '…'
        problems.append(
            'one maths expression could not be structured and was delivered as '
            f'plain text: "{snippet}" (reason: {reason}). Remedy: Ctrl+F the '
            'quoted text, fix that ⟦MATH:⟧ spelling in the block, rebuild.')
    return (not problems), problems

def _expl_segments(out_path, cfg):
    """Rendered prose lines of each question's EXPLANATION region only (after the
    Correct Answer line, question regions excluded — those belong to the source
    paper and are audited by source_math_health() instead)."""
    doc = Document(out_path)
    out = {}
    cur = 0; in_expl = False
    for p in doc.paragraphs:
        t = p.text.strip()
        mq = cfg.q_re.match(t)
        if mq:
            cur = int(mq.group(1)); in_expl = False; continue
        if cur and t.startswith('Correct Answer'):
            in_expl = True; continue
        if cur and in_expl and t:
            out.setdefault(cur, []).append(t)
    return out

def source_math_health(path, cfg):
    """v2.0 INPUT-SIDE HEALTH CHECK (advisory, never halts). Scans the SOURCE
    paper's question regions for upstream math loss so PYQ-1 never silently
    launders damage it inherited: (a) gap signatures where a symbol vanished
    ("= " with nothing after, multi-space holes before punctuation), (b) OMML
    islands whose math text is empty, (c) ASCII-math dialect already present in
    stems/options. Any hit means the paper predates PYQPrepare v2.0 — the remedy
    is to re-run Step 1 v2.0 on the source PDF FIRST, then regenerate. Returns a
    list of plain-language warnings; build_interleaved_docx prints them."""
    doc = Document(path)
    warns = []
    _gap = re.compile(r'(=\s*$|=\s{2}|\s{2,}[,.);]|\(\s+\)|\s{3,}\S)')
    cur = 0
    for p in doc.paragraphs:
        t = p.text
        mq = cfg.q_re.match(t.strip())
        if mq:
            cur = int(mq.group(1))
        if not cur:
            continue
        _has_math = any(''.join(x.text or '' for x in om.iter(qn('m:t'))).strip()
                        for om in p._element.iter(qn('m:oMath')))
        if len(t.strip()) > 2 and not _has_math and _gap.search(t):
            warns.append(f'Q{cur}: probable missing math symbol (gap) in source: {t.strip()[:70]!r}')
        for om in p._element.iter(qn('m:oMath')):
            if not ''.join(x.text or '' for x in om.iter(qn('m:t'))).strip():
                warns.append(f'Q{cur}: EMPTY OMML island in source (content already lost upstream)')
        for pat, what, _fx in _DIALECT_BANS:
            if pat.search(t):
                warns.append(f'Q{cur}: source carries {what} in {t.strip()[:60]!r}')
                break
    if warns:
        warns.append('SOURCE MATH HEALTH: the input paper predates PYQPrepare v2.0 — '
                     're-run Step 1 v2.0 on the source PDF, then regenerate explanations.')
    return warns

# ───────────────────────────── strip for re-audit ──────────────────────────
def set_coverage_banner(out_path, cfg, text):
    """v2.4 — write (or replace) the DOCUMENT-LEVEL coverage banner as the first
    body paragraph, so an interim artefact declares its own state wherever it
    travels. GAP-2026-08-19-INTERIM-ARTEFACT-UNLABELLED: coverage was announced
    only in chat, so a partially-explained paper forwarded to a third party was
    indistinguishable from a finished one.

    Safe by construction against all three gates:
      • verify_fidelity  — the banner is outside every question region, and
        question regions are what that gate compares.
      • strip_solutions  — removes it (below), so the questions-only copy stays
        byte-equal to the Step-7 source and the §12-3 re-audit still passes.
      • MANDATE 0        — the CALLER supplies the text and must keep it
        content-free (counts and Q-ranges only, never stem/answer text).
    Idempotent: re-running replaces the existing banner rather than stacking."""
    doc = Document(out_path)
    body = doc.element.body
    prefix = cfg.labels['coverage_banner']
    for child in list(body.iterchildren()):
        if child.tag == qn('w:p'):
            t = ''.join(x.text or '' for x in child.iter(qn('w:t'))).strip()
            if t.startswith(prefix):
                child.getparent().remove(child)
    if text:
        p = _new_para(cfg, 'banner', text=f'{prefix}: {text}', before=0, after=160,
                      bold=True, color=cfg.colors.get('hdr'), math=False)
        body.insert(0, p)
    doc.save(out_path)
    return out_path

# ── v2.9 (GAP-2026-08-26-REGISTRY-HANDOFF-SEAM) — END-OF-MOCK REPORT AS A DOCX ──
REPORT_MANDATE0_RE = re.compile(r'(?i)\b(answer\s*[:=]|option\s+[A-D1-4]\s+is\s+correct)')
REPORT_FONT = 'Arial'   # plain operator document; no exam value, no content of any question

def build_report_docx(out_path, title, sections, meta_lines=()):
    """Write the Step-9 END-OF-MOCK REPORT (§20) as a standalone .docx so the
    operator keeps it beside the Explanation docx instead of scrolling chat.

    title       document heading, e.g. 'END-OF-MOCK REPORT — EX Mock01'
    meta_lines  optional short lines under the title (provenance, timestamp)
    sections    ordered list of (heading, lines) — e.g. ('§R3 COVERAGE', [...]).
                Every line is one paragraph. Lines are MANDATE-0 content: Q-numbers,
                codes, counts, URLs — never a stem, option, answer or solution
                sentence. A line matching REPORT_MANDATE0_RE (an explicit answer
                declaration) raises ValueError BEFORE anything is written.

    The file is INERT downstream — Step 11 reads only registry.json and the
    Explanation docx (its filename ends _Explanation.docx; this one ends
    _Explain_Report.docx, so the S1-2 attachment gate can never confuse them).
    Returns out_path."""
    from docx.shared import Pt
    for hd, lines in sections:
        for ln in lines:
            if REPORT_MANDATE0_RE.search(str(ln)):
                raise ValueError(f'MANDATE 0: report line under {hd!r} declares an answer '
                                 f'— the report is content-free by contract')
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = REPORT_FONT
    st.font.size = Pt(10)
    doc.add_heading(title, level=1)
    for ml in meta_lines:
        doc.add_paragraph(str(ml))
    for hd, lines in sections:
        doc.add_heading(str(hd), level=2)
        if not lines:
            doc.add_paragraph('(empty — stated explicitly, never omitted)')
        for ln in lines:
            doc.add_paragraph(str(ln))
    doc.save(out_path)
    return out_path


def read_report_docx(path):
    """Inverse of build_report_docx for the self-test and any later audit:
    returns {'title': str, 'sections': [(heading, [lines])]}."""
    doc = Document(path)
    title, secs, cur = None, [], None
    for p in doc.paragraphs:
        sty = p.style.name if p.style is not None else ''
        if sty.startswith('Heading 1'):
            title = p.text
        elif sty.startswith('Heading 2'):
            cur = (p.text, [])
            secs.append(cur)
        elif cur is not None:
            cur[1].append(p.text)
    return {'title': title, 'sections': secs}


def strip_solutions(out_path, stripped_path, cfg):
    """Produce a questions-only copy (every appended explanation paragraph
    removed) so the Step-2 paper auditor sees ONLY the paper (Conflict-3).
    v2.4: the document-level coverage banner is removed here too — it is
    framework-added, not paper content, so leaving it would make the stripped
    copy differ from the Step-7 source and fail the §12-3 re-audit."""
    doc = Document(out_path)
    body = doc.element.body
    ca_label = cfg.labels['correct_answer'].lower()
    banner_prefix = cfg.labels['coverage_banner'].lower()
    in_expl = False
    to_remove = []
    for child in list(body.iterchildren()):
        if child.tag == qn('w:p'):
            txt = ''.join(t.text or '' for t in child.iter(qn('w:t'))).strip()
            if txt.lower().startswith(banner_prefix):
                to_remove.append(child); continue
            if cfg.q_re.match(txt):
                in_expl = False
            elif txt.lower().startswith(ca_label + ':'):
                in_expl = True
            if in_expl:
                to_remove.append(child)
        elif child.tag == qn('w:tbl') and in_expl:
            to_remove.append(child)
    for el in to_remove:
        el.getparent().remove(el)
    # v2.7 (e) — garbage-collect media parts that the removed explanation
    # figures referenced. Before this, the stripped copy held every explanation
    # image as an orphaned part (39 media against the source's 36 in the
    # reference run): text-identical to the source, byte-different from it.
    _still = set()
    for _blip in doc.element.body.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
        _rid = _blip.get(qn('r:embed')) or _blip.get(qn('r:link'))
        if _rid:
            _still.add(_rid)
    _IMG = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    for _rid, _rel in list(doc.part.rels.items()):
        if _rel.reltype == _IMG and _rid not in _still:
            doc.part.drop_rel(_rid)
    doc.save(stripped_path)
    return stripped_path

# ───────────────────────────── self-test ───────────────────────────────────
def _tiny_png():
    """Minimal valid 1x1 PNG (no external deps)."""
    import struct, zlib, io
    sig = b'\x89PNG\r\n\x1a\n'
    def chunk(typ, data):
        return (struct.pack('>I', len(data)) + typ + data +
                struct.pack('>I', zlib.crc32(typ + data) & 0xffffffff))
    ihdr = struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b'\x00\xff\x00\x00')        # one red pixel row
    return io.BytesIO(sig + chunk(b'IHDR', ihdr) + chunk(b'IDAT', idat) + chunk(b'IEND', b''))

def _make_sample_paper(path, cfg, nq=6):
    """Build a synthetic paper that EXERCISES the hard fidelity cases:
       Q2 carries an option figure in a paragraph AFTER its last option label
       (the anchor-bug case), and Q4 ends its region with a table (DI case)."""
    from docx.shared import Inches
    doc = Document()
    for q in range(1, nq + 1):
        doc.add_paragraph(f'Q.{q}  Sample stem number {q}.')
        for o in range(1, cfg.options_count + 1):
            doc.add_paragraph(f'{o}.  opt{o}')
        if q == 2:
            doc.add_paragraph().add_run().add_picture(_tiny_png(), width=Inches(0.3))
        if q == 4:
            t = doc.add_table(rows=2, cols=2)
            t.cell(0, 0).text = 'a'; t.cell(0, 1).text = 'b'
            t.cell(1, 0).text = 'c'; t.cell(1, 1).text = 'd'
        doc.add_paragraph('')
    doc.save(path)
    return path

def self_test():
    results = []
    def check(name, cond):
        results.append((name, bool(cond)))
    def _raises(fn):
        try:
            fn(); return False
        except ValueError:
            return True
    # v2.8 — the pre-v2.8 fixtures below construct blocks without the now-
    # mandatory provenance metadata; they test OTHER gates and stay byte-
    # unchanged. The module default is flipped for their duration only and
    # restored before the v2.8 fixtures, which run under the production default.
    EngineConfig.DEFAULT_PROVENANCE_GATES = False
    cfg = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4)

    # 1 guard: banned glyph
    try: guard_sentence('value is 21 \u2713'); check('G-GLYPH', False)
    except ValueError: check('G-GLYPH', True)
    # 2 guard: LaTeX
    try: guard_sentence('use \\frac here'); check('G-LATEX', False)
    except ValueError: check('G-LATEX', True)
    # 3 guard: metacommentary
    try: guard_sentence('wait, let me reconsider this'); check('G-META', False)
    except ValueError: check('G-META', True)
    # 4 guard: template sentence
    try: guard_sentence('This option is wrong here'); check('G-TEMPLATE', False)
    except ValueError: check('G-TEMPLATE', True)
    # 5 guard: fake citation
    try: guard_sentence('per the official key it is 3'); check('G-FAKECITE', False)
    except ValueError: check('G-FAKECITE', True)
    # 6 guard: two sentences
    try: guard_sentence('First fact. Second fact.'); check('G-2SENT', False)
    except ValueError: check('G-2SENT', True)
    # 7 guard: one sentence OK
    try: check('G-1SENT', bool(guard_sentence('A single clean factual sentence here.')))
    except ValueError: check('G-1SENT', False)
    # 8 guard: year-range slash
    try: guard_sentence('in 2025/26 the budget rose.'); check('G-YEAR', False)
    except ValueError: check('G-YEAR', True)
    # 9 units are NOT fractions
    check('G-UNIT', not has_inline_fraction('speed is 60 km/h here'))
    # 10 inline fraction detected
    check('G-INLINEFRAC', has_inline_fraction('value 3/4 remains'))
    # 11 add_math_text converts a fraction without raising
    from docx.text.paragraph import Paragraph
    p = OxmlElement('w:p'); proxy = Paragraph(p, None)
    add_math_text(proxy, 'the share is 2/3 of total.')
    check('G-OMMLCONV', len(p.findall('.//' + qn('m:oMath'))) == 1)
    # 12 end-of-sentence fraction raises (EX13)
    try:
        p2 = OxmlElement('w:p'); add_math_text(Paragraph(p2, None), 'result = 3/4.')
        check('G-ENDFRAC', False)
    except ValueError: check('G-ENDFRAC', True)
    # 13 ExplanationBlock validates a good block
    good = ExplanationBlock(q=1, ca=1,
        axiom=['The sum of a group equals its average times its count.'],
        deduction=['Total is 9 times 43 = 387.',
                   'Remaining average is 235 over 5 = 47, which is Option 1.'],
        why_wrong={2: ['Option 2 uses 12 not 13, one count short.'],
                   3: ['Option 3 shifts by four not three places.'],
                   4: ['Option 4 drops the letter rule midway.']},
        cfg=cfg)
    check('B-VALID', good.validate())
    # 14 DEDUCTION must bind Option N
    try:
        bad = ExplanationBlock(q=1, ca=1, axiom=['x.'],
            deduction=['step one.', 'step two without binding.'],
            why_wrong={2:['a.'],3:['b.'],4:['c.']}, cfg=cfg); bad.validate()
        check('B-BIND', False)
    except ValueError: check('B-BIND', True)
    # 15 WHY WRONG keys must equal wrong options
    try:
        bad = ExplanationBlock(q=1, ca=1, axiom=['x.'],
            deduction=['s.', 'final is Option 1.'],
            why_wrong={2:['a.'],3:['b.']}, cfg=cfg); bad.validate()
        check('B-WWKEYS', False)
    except ValueError: check('B-WWKEYS', True)
    # 13a NAT charset guard — full portal scenario matrix (v1.16)
    # positive integer
    check('NAT-CS-POSINT', validate_nat_grading_value('3') == '3')
    # signed integer
    check('NAT-CS-NEGINT', validate_nat_grading_value(-47) == '-47')
    # plain decimal, decimal-capable (no forced point)
    check('NAT-CS-DEC', validate_nat_grading_value('3.5') == '3.5')
    # fixed-precision decimal (padding is Step 7's job; engine just accepts a
    # well-formed already-padded string)
    check('NAT-CS-DECFIX', validate_nat_grading_value('3.00') == '3.00')
    # leading zero required under 1 — engine does not coerce, only rejects what
    # is actually malformed; '.5' has no digit before the point
    check('NAT-CS-LEADZERO-REJECT', _raises(lambda: validate_nat_grading_value('.5')))
    # scientific notation — the original bug class ('3e-9')
    check('NAT-CS-SCINOT-REJECT', _raises(lambda: validate_nat_grading_value('3e-9')))
    # unicode multiplication sign / superscript exponent
    check('NAT-CS-UNICODE-SCINOT-REJECT',
          _raises(lambda: validate_nat_grading_value('3\u00d710\u207b\u2079')))
    # unit suffix leaking into the grading field
    check('NAT-CS-UNIT-REJECT', _raises(lambda: validate_nat_grading_value('47 nm')))
    # stray plus sign (not in the allowed charset)
    check('NAT-CS-PLUS-REJECT', _raises(lambda: validate_nat_grading_value('+3')))
    # valid range, both bounds non-negative
    check('NAT-RANGE-VALID', format_nat_range('46.50', '47.50', ctx='Q1') == '46.50-47.50')
    check('NAT-RANGE-INT', format_nat_range('5', '7', ctx='Q1') == '5-7')
    # old broken format (parenthetical + en-dash + word) must be gone for good
    check('NAT-RANGE-NO-ENDASH',
          _raises(lambda: validate_nat_grading_value('47 (accepted range 46.5\u201347.5)')))
    # negative-bounded range — explicit NOT SUPPORTED gate (locked decision)
    check('NAT-RANGE-NEGATIVE-NOTSUPPORTED',
          _raises(lambda: format_nat_range('-5', '7', ctx='Q1')))
    check('NAT-RANGE-BOTH-NEGATIVE-NOTSUPPORTED',
          _raises(lambda: format_nat_range('-10', '-5', ctx='Q1')))
    # lo>hi still hard-fails (pre-existing invariant, must survive the rewrite)
    check('NAT-RANGE-LOGT-HI-REJECT', _raises(lambda: format_nat_range('9', '3', ctx='Q1')))
    # ExplanationBlock end-to-end: a NAT block with a bad ca must fail AT
    # CONSTRUCTION (validate()), never silently reach render()
    nat_cfg = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4, options_by_q={9: 0})
    try:
        bad_nat = ExplanationBlock(q=9, ca='3e-9', qtype='nat',
            axiom=['A rate under neutral theory equals mutation rate times neutral fraction.'],
            deduction=['5e-9 times 0.60 gives the neutral rate.', '3e-9 per site per generation.'],
            common_pitfalls={'5e-9': ['Uses the total rate, ignoring the neutral fraction.']},
            cfg=nat_cfg)
        bad_nat.validate()
        check('NAT-BLOCK-SCINOT-REJECT', False)
    except ValueError:
        check('NAT-BLOCK-SCINOT-REJECT', True)
    # ExplanationBlock end-to-end: a correctly pre-scaled NAT value passes clean
    good_nat = ExplanationBlock(q=9, ca='3', qtype='nat',
        axiom=['A rate under neutral theory equals mutation rate times neutral fraction.'],
        deduction=['Total rate is 5, in units of 10 to the -9, times 0.60.',
                   'The neutral substitution rate is 3, in units of 10 to the -9.'],
        common_pitfalls={'5': ['Uses the total rate, ignoring the neutral fraction.']},
        cfg=nat_cfg)
    check('NAT-BLOCK-VALID', good_nat.validate())
    # ExplanationBlock end-to-end: an UNCOMPILABLE ⟦MATH:…⟧ region must fail AT
    # CONSTRUCTION (validate()), never degrade to raw LaTeX at render
    # (GAP-2026-08-10-EXPLAIN-MATH-DEGRADE-SILENT). The body below is chosen so the
    # ONLY thing that can reject it is the Tier-3 compile gate: \tfrac survives
    # guard_sentence (verified — the LaTeX guard does not fire on it) and fails
    # t3_compile as an unknown command. A body the prose guard already rejects would
    # make this fixture pass with the gate removed, which is no test at all.
    try:
        ExplanationBlock(q=1, ca=1,
            axiom=['The ratio of the two quantities is fixed by the definition.'],
            deduction=['Substituting gives ⟦MATH:\\tfrac{1}{2}⟧ for the ratio.',
                       'That value settles option 1.'],
            why_wrong={2: ['a.'], 3: ['b.'], 4: ['c.']}, cfg=cfg).validate()
        check('T3-GATE-UNCOMPILABLE-REJECT', False)
    except ValueError:
        check('T3-GATE-UNCOMPILABLE-REJECT', True)
    # …and a well-formed region still passes, so the gate rejects bad math rather
    # than all math.
    check('T3-GATE-COMPILABLE-ACCEPT', ExplanationBlock(q=1, ca=1,
        axiom=['The ratio of the two quantities is fixed by the definition.'],
        deduction=['Substituting gives ⟦MATH:\\frac{1}{2}⟧ for the ratio.',
                   'That value settles option 1.'],
        why_wrong={2: ['a.'], 3: ['b.'], 4: ['c.']}, cfg=cfg).validate())

    # 16 build + fidelity + structure round-trip (Q2 has a trailing option figure,
    #    Q4 ends in a table — both must survive byte-identical)
    import tempfile, os
    d = tempfile.mkdtemp()
    src = os.path.join(d, 'paper.docx'); out = os.path.join(d, 'sol.docx')
    _make_sample_paper(src, cfg, nq=6)
    blocks = []
    for q in range(1, 5):  # batch 1 = Q1..Q4 (covers the figure + table cases)
        blocks.append(ExplanationBlock(q=q, ca=1,
            axiom=['A reusable principle stated as a truth here.'],
            deduction=['First derived step gives a value.',
                       f'Final value maps to Option 1 for question {q}.'],
            speed_hack=['A genuinely shorter elimination route here.'],
            why_wrong={2:['Option 2 uses the right value for the wrong quantity.'],
                       3:['Option 3 subtracts the term that must be added.'],
                       4:['Option 4 satisfies one condition and misses the other.']},
            cfg=cfg))
    n = build_interleaved_docx(src, blocks, out, cfg)
    check('BUILD-N', n == 4)
    okf, pf = verify_fidelity(out, src, cfg)
    check('FIDELITY', okf)
    oks, ps = verify_structure(out, blocks, cfg, expected_qs=[1, 2, 3, 4])
    check('STRUCT', oks)
    # 17 look-ahead is caught: claim batch is only Q1..Q3 but doc has Q4 explained
    okl, pl = verify_structure(out, blocks, cfg, expected_qs=[1, 2, 3])
    check('NO-LOOKAHEAD', (not okl))
    # 18 strip_solutions removes explanations (questions-only copy parses clean)
    stripped = os.path.join(d, 'stripped.docx')
    strip_solutions(out, stripped, cfg)
    _, qmap = parse_paper(stripped, cfg)
    check('STRIP', len(qmap) == 6 and all(len(qmap[q]['opt_paras']) == 4 for q in qmap))

    # ── edge-case guards (added after a forensic audit) ──────────────────────
    # 19 meta false-positives are GONE: 'ohmmeter'/'await'/'factually'/'air' pass
    def _passes(t):
        try: guard_sentence(t); return True
        except ValueError: return False
    check('G-META-FP', all(_passes(t) for t in [
        'The ohmmeter reads two ohms across the resistor here.',
        'Candidates await, then the second method confirms it.',
        'The statement is factually, not morally, the weaker claim.',
        'The curve has an air of symmetry about the origin here.']))
    # 20 fake-citation guard is EXAM-AGNOSTIC: a former hardcoded token now passes,
    #    while a generic fabricated-authority phrase is still caught
    check('G-NOEXAM', _passes('The SSC pattern places this topic in the first section.')
          and not _passes('The official answer key says the value is three here.'))
    # 21 CA binding is word-bounded: 12 options, ca=1 binds 'Option 1' (not 'Option 10')
    cfg12 = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9]|1[0-2])[.\)]', 12)
    okbind = ExplanationBlock(q=1, ca=1, cfg=cfg12,
        axiom=['A principle stated as a truth here.'],
        deduction=['Eliminating Option 10 and Option 11 narrows the field.',
                   'Only the first choice survives, which is Option 1.'],
        why_wrong={k: [f'Option {k} fails one stated condition here.']
                   for k in range(2, 13)}).validate()
    check('B-BIND-WB-OK', okbind)
    def _bind_bad():
        try:
            ExplanationBlock(q=1, ca=1, cfg=cfg12,
                axiom=['A principle stated as a truth here.'],
                deduction=['A first step yields a value.',
                           'The survivor is Option 10 only here.'],   # ca=1 but binds Option 10
                why_wrong={k: [f'Option {k} fails one stated condition here.']
                           for k in range(2, 13)}).validate()
            return False
        except ValueError:
            return True
    check('B-BIND-WB-FAIL', _bind_bad())
    # 22 five-option config builds + validates + round-trips fidelity
    cfg5 = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-5])[.\)]', 5)
    d5 = tempfile.mkdtemp(); s5 = os.path.join(d5, 'p.docx'); o5 = os.path.join(d5, 's.docx')
    _make_sample_paper(s5, cfg5, nq=3)
    b5 = [ExplanationBlock(q=q, ca=2, cfg=cfg5,
        axiom=['A principle stated as a truth here.'],
        deduction=['A first step yields a value.', f'The value maps to Option 2 for item {q}.'],
        why_wrong={1:['Option 1 uses the right value for the wrong quantity.'],
                   3:['Option 3 subtracts the term that must be added.'],
                   4:['Option 4 satisfies one condition and misses the other.'],
                   5:['Option 5 rounds at an intermediate step.']}) for q in [1, 2]]
    n5 = build_interleaved_docx(s5, b5, o5, cfg5)
    okf5, _ = verify_fidelity(o5, s5, cfg5)
    oks5, _ = verify_structure(o5, b5, cfg5, expected_qs=[1, 2])
    check('CFG-5OPT', n5 == 2 and okf5 and oks5)
    # 23 non-English labels flow through build AND coverage detection
    hi = {'correct_answer': '\u0938\u0939\u0940 \u0909\u0924\u094d\u0924\u0930',
          'axiom': 'AXIOM', 'deduction': 'DEDUCTION', 'speed_hack': 'SPEED HACK',
          'why_wrong': 'WHY WRONG?', 'option': 'Option', 'solution_ref': 'Solution'}
    cfgL = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4, labels=hi)
    dL = tempfile.mkdtemp(); sL = os.path.join(dL, 'p.docx'); oL = os.path.join(dL, 's.docx')
    _make_sample_paper(sL, cfgL, nq=2)
    bL = [ExplanationBlock(q=1, ca=1, cfg=cfgL,
        axiom=['A principle stated as a truth here.'],
        deduction=['A first step yields a value.', 'The value maps to Option 1 here.'],
        why_wrong={2:['Option 2 uses the right value for the wrong quantity.'],
                   3:['Option 3 subtracts the term that must be added.'],
                   4:['Option 4 satisfies one condition and misses the other.']})]
    build_interleaved_docx(sL, bL, oL, cfgL)
    oksL, _ = verify_structure(oL, bL, cfgL, expected_qs=[1])
    okfL, _ = verify_fidelity(oL, sL, cfgL)
    check('CFG-LABELS', oksL and okfL)
    # 24 an anomaly (escalation) block validates but MUST NOT render/ship
    anom = ExplanationBlock(q=1, anomaly='no single defensible answer', cfg=cfg)
    anom_ok = anom.validate()
    def _anom_blocked():
        try:
            build_interleaved_docx(src, [anom], os.path.join(d, 'x.docx'), cfg)
            return False
        except ValueError:
            return True
    check('ANOMALY-NORENDER', anom_ok and _anom_blocked())
    # 25 alternate question/option format parses (Q1. stem, (1) option)
    cfgA = EngineConfig(r'^Q(\d+)\.', r'^\(([1-9])\)', 4)
    dA = tempfile.mkdtemp(); sA = os.path.join(dA, 'p.docx')
    docA = Document()
    for q in range(1, 4):
        docA.add_paragraph(f'Q{q}. stem {q}')
        for o in range(1, 5):
            docA.add_paragraph(f'({o}) opt{o}')
        docA.add_paragraph('')
    docA.save(sA)
    _, qmA = parse_paper(sA, cfgA)
    check('CFG-ALTFMT', len(qmA) == 3 and all(len(qmA[q]['opt_paras']) == 4 for q in qmA))

    # ── NAT / MSQ / mixed-section / label-scheme / terminator tests ──────────
    # 26 NAT: a question with NO options; value answer; value-bound deduction;
    #    common pitfalls instead of why_wrong; parse allows 0 options
    cfgN = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', None,
                        options_by_q={1: 0, 2: 0})
    dN = tempfile.mkdtemp(); sN = os.path.join(dN, 'p.docx'); oN = os.path.join(dN, 's.docx')
    docN = Document()
    for q in [1, 2]:
        docN.add_paragraph(f'Q.{q} compute the numerical value and enter it')
        docN.add_paragraph('')                       # NO options
    docN.save(sN)
    _, qmN = parse_paper(sN, cfgN)
    natblk = ExplanationBlock(q=1, ca='47', cfg=cfgN, ca_range=(46.5, 47.5),
        axiom=['The mean equals the total over the count here.'],
        deduction=['The total is 235 over a count of five.',
                   'Dividing gives the value 47 as the answer.'],
        common_pitfalls={'235': ['Forgetting to divide leaves 235 unchanged.'],
                         '9.4': ['Dividing by the wrong count gives 9.4 instead.']})
    nat_ok = natblk.validate()
    nN = build_interleaved_docx(sN, [natblk], oN, cfgN)
    okfN, _ = verify_fidelity(oN, sN, cfgN)
    oksN, _ = verify_structure(oN, [natblk], cfgN, expected_qs=[1])
    # NAT CA line shows the VALUE (not an option index). v1.16: when ca_range
    # is present the rendered line is the portal Range format 'lo-hi' ONLY —
    # it replaces the point value + parenthetical wording entirely (the old
    # 'Correct Answer: 47 (accepted range 46.5-47.5)' format is banned: it
    # violates the portal's 0-9.- grading charset on five different counts).
    docNo = Document(oN)
    ca_line_ok = any(p.text.strip() == 'Correct Answer: 46.5-47.5' for p in docNo.paragraphs)
    check('NAT', nat_ok and nN == 1 and okfN and oksN and ca_line_ok)

    # 26b STR-KEYED options_by_q (registry.json boundary): keys arriving as JSON strings
    #     must resolve identically to int keys. Regression lock for the str/int ND6 fix —
    #     a NAT question (0) keyed "3" must still type as nat when queried with int 3.
    cfgS = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4,
                        options_by_q={'1': 4, '2': 4, '3': 0})       # STRING keys
    check('NAT-STRKEY', cfgS.expected_options(3) == 0
                        and cfgS.expected_options(1) == 4
                        and cfgS.expected_options(2) == 4)
    # 27 NAT rejects option-style why_wrong and an unbound value
    def _nat_bad():
        try:
            ExplanationBlock(q=1, ca='47', cfg=cfgN,
                axiom=['x truth here.'], deduction=['a.', 'b.'],
                why_wrong={2: ['no.']}).validate()           # why_wrong illegal for NAT
            return False
        except ValueError:
            try:
                ExplanationBlock(q=1, ca='47', cfg=cfgN, axiom=['x truth here.'],
                    deduction=['a.', 'final value is 99 here.'],   # 47 not bound
                    common_pitfalls={'5': ['p.']}).validate()
                return False
            except ValueError:
                return True
    check('NAT-GUARDS', _nat_bad())
    # 28 MSQ: multiple correct; ca is a set; last step binds all selected;
    #    why_wrong = the non-selected options
    cfgM = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4)
    dM = tempfile.mkdtemp(); sM = os.path.join(dM, 'p.docx'); oM = os.path.join(dM, 's.docx')
    _make_sample_paper(sM, cfgM, nq=2)
    msq = ExplanationBlock(q=1, ca={1, 3}, cfg=cfgM, qtype='msq',
        axiom=['Each statement is judged independently here.'],
        deduction=['Statement one holds and statement three holds.',
                   'The correct set is Option 1 and Option 3 here.'],
        why_wrong={2: ['Option 2 fails one stated condition here.'],
                   4: ['Option 4 inverts the relation being tested.']})
    msq_ok = msq.validate()
    nM = build_interleaved_docx(sM, [msq], oM, cfgM)
    okfM, _ = verify_fidelity(oM, sM, cfgM)
    oksM, _ = verify_structure(oM, [msq], cfgM, expected_qs=[1])
    docMo = Document(oM)
    msq_ca_ok = any(p.text.strip() == 'Correct Answer: 1, 3' for p in docMo.paragraphs)
    check('MSQ', msq_ok and nM == 1 and okfM and oksM and msq_ca_ok)
    # 29 MSQ rejects: last step not binding all selected; wrong why_wrong keys
    def _msq_bad():
        try:
            ExplanationBlock(q=1, ca={1, 3}, cfg=cfgM, qtype='msq',
                axiom=['x truth here.'],
                deduction=['a.', 'only Option 1 here.'],      # Option 3 not bound
                why_wrong={2: ['p.'], 4: ['q.']}).validate()
            return False
        except ValueError:
            try:
                ExplanationBlock(q=1, ca={1, 3}, cfg=cfgM, qtype='msq',
                    axiom=['x truth here.'],
                    deduction=['a.', 'set is Option 1 and Option 3 here.'],
                    why_wrong={2: ['p.']}).validate()         # missing key 4
                return False
            except ValueError:
                return True
    check('MSQ-GUARDS', _msq_bad())
    # 30 mixed-section option counts: Q1 has 4, Q2 has 5 — passes with the map,
    #    and the same paper FAILS under a single uniform count
    cfgX = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', None,
                        options_by_q={1: 4, 2: 5})
    dX = tempfile.mkdtemp(); sX = os.path.join(dX, 'p.docx')
    docX = Document()
    docX.add_paragraph('Q.1 stem'); [docX.add_paragraph(f'{o}. o') for o in range(1, 5)]; docX.add_paragraph('')
    docX.add_paragraph('Q.2 stem'); [docX.add_paragraph(f'{o}. o') for o in range(1, 6)]; docX.add_paragraph('')
    docX.save(sX)
    mix_ok = False
    try:
        parse_paper(sX, cfgX); mix_ok = True
    except ValueError:
        mix_ok = False
    uni_fail = False
    try:
        parse_paper(sX, EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4))
    except ValueError:
        uni_fail = True
    check('MIXED-OPTS', mix_ok and uni_fail)
    # 31 alphabetic option labels: CA line + Option refs show A/B/C/D
    cfgAl = EngineConfig(r'^Q\.?\s*(\d+)', r'^([A-D])[.\)]', 4, label_scheme='alpha_upper')
    dAl = tempfile.mkdtemp(); sAl = os.path.join(dAl, 'p.docx'); oAl = os.path.join(dAl, 's.docx')
    docAl = Document()
    docAl.add_paragraph('Q.1 stem')
    for L in ['A', 'B', 'C', 'D']:
        docAl.add_paragraph(f'{L}. choice')
    docAl.add_paragraph(''); docAl.save(sAl)
    albl = ExplanationBlock(q=1, ca=1, cfg=cfgAl,
        axiom=['A principle stated as a truth here.'],
        deduction=['A first step yields a value.', 'The value maps to Option A here.'],
        why_wrong={2: ['Option B uses the right value for the wrong quantity.'],
                   3: ['Option C subtracts the term that must be added.'],
                   4: ['Option D satisfies one condition and misses the other.']})
    al_ok = albl.validate()
    build_interleaved_docx(sAl, [albl], oAl, cfgAl)
    docAlo = Document(oAl)
    alpha_ca = any(p.text.strip() == 'Correct Answer: A' for p in docAlo.paragraphs)
    alpha_sub = any(p.text.strip() == 'Option B' for p in docAlo.paragraphs)
    check('LABEL-ALPHA', al_ok and alpha_ca and alpha_sub)
    # 32 Devanagari sentence terminator (danda): a 3-sentence paragraph now FAILS
    cfgHi = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4,
                         sentence_terminators='.!?\u0964\u0965')
    three = '\u092a\u0939\u0932\u093e\u0964 \u0926\u0942\u0938\u0930\u093e\u0964 \u0924\u0940\u0938\u0930\u093e\u0964'
    danda_caught = False
    try:
        guard_sentence(three, cfgHi)
    except ValueError:
        danda_caught = True
    danda_default_misses = (sentence_count(three) == 1)   # default '.!?' misses danda
    check('TERMINATOR-DANDA', danda_caught and danda_default_misses)
    # 33 OMML helpers all produce well-formed math nodes (the author-facing API)
    omml_ok = True
    for x in (sup('x', '2'), sqrt('2'), nary('\u2211', 'i', 'n', 'i'), _r('5'),
              frac('1', '2'), omath(frac('1', '2'))):
        try:
            parse_xml(x)
        except Exception:
            omml_ok = False
    check('OMML-HELPERS', omml_ok)
    # 34 figural question: NO FIGURE section is rendered (removed by design), yet
    #    the image inside the question region survives byte-identical and the
    #    post-render audit passes. Regression-lock for the FIGURE-section removal.
    cfgF = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4)
    dF = tempfile.mkdtemp(); sF = os.path.join(dF, 'p.docx'); oF = os.path.join(dF, 's.docx')
    docF = Document()
    docF.add_paragraph('Q.1 which figure continues the series')
    from docx.shared import Inches
    docF.add_paragraph().add_run().add_picture(_tiny_png(), width=Inches(0.3))
    for o in range(1, 5):
        docF.add_paragraph(f'{o}. figure')
    docF.add_paragraph(''); docF.save(sF)
    figblk = ExplanationBlock(q=1, ca=1, cfg=cfgF,
        axiom=['A figural series applies one fixed transformation at every step.'],
        absolutes_justified={'A figural series applies one fixed transformation at every step.':
                             'definition of a figural series'},
        deduction=['Tracing the rotation forward predicts the next figure.',
                   'That predicted figure matches Option 1 here.'],
        why_wrong={2:['Option 2 rotates the opposite way.'],
                   3:['Option 3 over-rotates by one full step.'],
                   4:['Option 4 changes the wrong element of the pair.']})
    figblk.validate()
    build_interleaved_docx(sF, [figblk], oF, cfgF)
    okfF, _ = verify_fidelity(oF, sF, cfgF)                 # question-region image preserved
    fig_paras = [p.text.strip() for p in Document(oF).paragraphs]
    no_figure_section = not any('FIGURE' in t for t in fig_paras)   # section fully gone
    okeF, _ = verify_explanations(oF, [figblk], cfgF)       # audit passes with no FIGURE header
    check('FIGURAL-NO-FIGURE-SECTION', okfF and no_figure_section and okeF)
    # 35 vulgar-fraction glyphs and U+2044 are rejected (A1)
    vulgar_caught = all(_raises(lambda s=s: guard_sentence(s))
                        for s in ('The share is \u00bd of the total here.',
                                  'About \u2154 of the class passed today.',
                                  'The ratio 3\u20444 appears in the stem here.'))
    check('VULGAR-FRACTION', vulgar_caught and has_inline_fraction('a \u00bd here'))
    # 36 wider inline-fraction forms (1/x, x²/2, (a+b)/c) are caught; units still pass (A4)
    wide = (has_inline_fraction('the term 1/x grows here')
            and has_inline_fraction('x\u00b2/2 is the area here')
            and has_inline_fraction('(a+b)/c equals the mean')
            and not has_inline_fraction('the speed is 60 km/h here'))
    check('WIDE-FRACTION', wide)
    # 37 year-range precision: 2024/25 flagged, genuine n/(n+1) like 2024/26 NOT (A6)
    yr = (_year_range_hit('growth in 2024/25 was steady') is not None
          and _year_range_hit('the fraction 2024/26 reduces') is None)
    check('YEAR-RANGE-PRECISION', yr)
    # 38 richer sentence counter: abbreviations / acronyms / initials stay one sentence (A5)
    one = (sentence_count('Govt. data shows the approx. value here.') == 1
           and sentence_count('The meeting is at 10 a.m. sharp today.') == 1
           and sentence_count('C. V. Raman won the prize in physics.') == 1
           and sentence_count('First fact here. Second fact here.') == 2)
    check('SENTENCE-COUNTER', one)
    # 39 post-render audit passes a clean build and CATCHES a tampered explanation (A2)
    cfgP = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4)
    dP = tempfile.mkdtemp(); sP = os.path.join(dP, 'p.docx'); oP = os.path.join(dP, 's.docx')
    dpz = Document(); dpz.add_paragraph('Q.1 stem')
    for o in range(1, 5):
        dpz.add_paragraph(f'{o}. opt')
    dpz.add_paragraph(''); dpz.save(sP)
    pblk = ExplanationBlock(q=1, ca=1, cfg=cfgP,
        axiom=['A truth with its reason stated here.'],
        deduction=['A first step yields a value here.', 'That value is Option 1 here.'],
        why_wrong={2: ['Option 2 uses the right value for the wrong quantity.'],
                   3: ['Option 3 subtracts the term that must be added.'],
                   4: ['Option 4 satisfies one condition and misses the other.']})
    build_interleaved_docx(sP, [pblk], oP, cfgP)
    clean_ok, _ = verify_explanations(oP, [pblk], cfgP)
    # tamper: inject a banned glyph into a rendered explanation paragraph
    dtam = Document(oP)
    for p in dtam.paragraphs:
        if p.text.strip().startswith('A first step'):
            if p.runs:
                p.runs[0].text = p.runs[0].text + ' \u2713'
            break
    tP = os.path.join(dP, 't.docx'); dtam.save(tP)
    tam_caught, _ = verify_explanations(tP, [pblk], cfgP)
    check('POST-RENDER-AUDIT', clean_ok and not tam_caught)
    # v1.17: independent Correct-Answer charset audit (RXA-CHARSET) — a NAT block that
    # is perfectly clean AT CONSTRUCTION (passes validate()) but whose RENDERED docx is
    # hand-tampered afterward (simulating a future renderer bug or a hand-edited file)
    # must still be caught by verify_explanations, since it re-reads the actual bytes on
    # disk rather than trusting the in-memory block that produced them.
    cfgNG = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', None, options_by_q={1: 0})
    dNG = tempfile.mkdtemp(); sNG = os.path.join(dNG, 'p.docx'); oNG = os.path.join(dNG, 's.docx')
    dNGdoc = Document(); dNGdoc.add_paragraph('Q.1 stem'); dNGdoc.add_paragraph(''); dNGdoc.save(sNG)
    ngblk = ExplanationBlock(q=1, ca='3', cfg=cfgNG,
        axiom=['A rate under neutral theory equals mutation rate times neutral fraction.'],
        deduction=['Total rate is 5, in units of 10 to the -9, times 0.60.',
                   'The neutral substitution rate is 3, in units of 10 to the -9.'],
        common_pitfalls={'5': ['Uses the total rate, ignoring the neutral fraction.']})
    build_interleaved_docx(sNG, [ngblk], oNG, cfgNG)
    ng_clean_ok, _ = verify_explanations(oNG, [ngblk], cfgNG)
    dtamNG = Document(oNG)
    for p in dtamNG.paragraphs:
        if p.text.strip().startswith('Correct Answer:'):
            for r in list(p.runs):
                r.text = ''
            p.runs[0].text = 'Correct Answer: 3e-9'   # the original bug, hand-injected
            break
    tNG = os.path.join(dNG, 't.docx'); dtamNG.save(tNG)
    ng_tam_caught, ng_probs = verify_explanations(tNG, [ngblk], cfgNG)
    check('RXA-CHARSET-CATCH', ng_clean_ok and not ng_tam_caught
          and any('banned character' in p for p in ng_probs))
    # 40 rels resolution: clean build has no dangling image rIds (A3)
    okrel, _ = verify_fidelity(oF, sF, cfgF)
    check('RELS-RESOLVE', okrel and not (_embed_ids(oF) - _rel_ids(oF)))

    # ══ v2.5 LABEL BINDING (GAP-2026-08-19-UNSATISFIABLE-LABEL-BINDING) ═══════
    # A custom label_scheme ending in a NON-WORD character must be BINDABLE. The
    # old trailing \b made these schemes unsatisfiable — every block raised at
    # construction, so the exam could not be explained at all.
    _cfgL = EngineConfig(r'^Q\.?\s*(\d+)', r'^(\([A-Z]\))', 4,
                         label_scheme=['(A)', '(B)', '(C)', '(D)'])
    def _lblblk(last):
        return ExplanationBlock(q=1, ca=4, cfg=_cfgL,
            axiom=['A governing principle.'],
            deduction=['A first step.', last],
            why_wrong={1: ['A value swapped - a.'], 2: ['A sign flipped - b.'],
                       3: ['A unit slipped - c.']})
    check('LABEL-PAREN-BINDS', _lblblk('The answer is Option (D).').validate() is True)
    check('LABEL-PAREN-MIDSENTENCE',
          _lblblk('That makes Option (D) the answer.').validate() is True)
    # A WRONG label must still fail — the fix must not make the check vacuous.
    check('LABEL-PAREN-WRONG-STILL-FAILS',
          _raises(lambda: _lblblk('The answer is Option (B).').validate()))
    # The prefix-confusion lock must survive for numeric schemes (1 vs 10).
    _cfg10 = EngineConfig(r'^Q\.?\s*(\d+)', r'^(\d+)[.\)]', 12)
    def _blk10(last):
        return ExplanationBlock(q=1, ca=1, cfg=_cfg10,
            axiom=['A governing principle.'],
            deduction=['A first step.', last],
            why_wrong={k: [f'A value swapped - {k}.'] for k in range(2, 13)})
    check('LABEL-NUMERIC-PREFIX-LOCK-HELD',
          _blk10('The answer is Option 1.').validate() is True
          and _raises(lambda: _blk10('The answer is Option 10.').validate()))
    # Bracketed and trailing-paren schemes bind too.
    for _sch, _lab in ((['[A]', '[B]', '[C]', '[D]'], '[D]'),
                       (['A)', 'B)', 'C)', 'D)'], 'D)')):
        _c = EngineConfig(r'^Q\.?\s*(\d+)', r'^(.)', 4, label_scheme=_sch)
        ExplanationBlock(q=1, ca=4, cfg=_c, axiom=['A principle.'],
            deduction=['A step.', f'The answer is Option {_lab}.'],
            why_wrong={1: ['A value swapped - a.'], 2: ['A sign flipped - b.'],
                       3: ['A unit slipped - c.']}).validate()
    check('LABEL-BRACKET-AND-TRAILING-PAREN', True)
    # END-TO-END: a paren-scheme paper must survive BUILD + all three verifiers.
    # The construction fix alone was not enough — verify_explanations carried its own
    # copy of the same trailing-\b test, so the exam still could not ship.
    _lsrc = os.path.join(tempfile.gettempdir(), 'st_lblsrc.docx')
    _lout = os.path.join(tempfile.gettempdir(), 'st_lblout.docx')
    _ld = Document()
    _ld.add_paragraph('Q.1  Which option is correct')
    for _lb in ('(A)', '(B)', '(C)', '(D)'):
        _ld.add_paragraph(f'{_lb} choice')
    _ld.add_paragraph(''); _ld.save(_lsrc)
    _lb1 = _lblblk('The answer is Option (D).')
    build_interleaved_docx(_lsrc, [_lb1], _lout, _cfgL)
    _lf, _ = verify_fidelity(_lout, _lsrc, _cfgL)
    _ls, _ = verify_structure(_lout, [_lb1], _cfgL)
    _le, _lep = verify_explanations(_lout, [_lb1], _cfgL)
    check('LABEL-PAREN-E2E-ALL-VERIFIERS', _lf and _ls and _le)


    # ══ v2.4 COVERAGE BANNER (GAP-2026-08-19-INTERIM-ARTEFACT-UNLABELLED) ═════
    _bsrc = os.path.join(tempfile.gettempdir(), 'st_bansrc.docx')
    _bout = os.path.join(tempfile.gettempdir(), 'st_banout.docx')
    _bstr = os.path.join(tempfile.gettempdir(), 'st_banstr.docx')
    _cfgB = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4)
    _make_sample_paper(_bsrc, _cfgB, nq=6)
    _bb = ExplanationBlock(q=1, ca=4, cfg=_cfgB,
        axiom=['A governing principle stated once.'],
        deduction=['The first step.', 'The second step gives Option 4.'],
        why_wrong={1: ['A value swapped - a.'], 2: ['A sign flipped - b.'], 3: ['A unit slipped - c.']})
    build_interleaved_docx(_bsrc, [_bb], _bout, _cfgB)
    set_coverage_banner(_bout, _cfgB, 'Batch 1 of 6 - Q1..Q1 explained of 6. NOT FINAL.')
    _btxt = [p.text for p in Document(_bout).paragraphs]
    check('BANNER-PRESENT-FIRST', _btxt and _btxt[0].startswith(
        _cfgB.labels['coverage_banner']))
    # The banner must not disturb the gates the paper already passes.
    _okbf, _ = verify_fidelity(_bout, _bsrc, _cfgB)
    _okbs, _ = verify_structure(_bout, [_bb], _cfgB, expected_qs=[1])
    _okbe, _ = verify_explanations(_bout, [_bb], _cfgB, expected_qs=[1])
    check('BANNER-GATES-UNAFFECTED', _okbf and _okbs and _okbe)
    # STRIP MUST REMOVE IT — otherwise the questions-only copy differs from the
    # Step-7 source and the §12-3 re-audit fails. This is the whole reason the
    # banner needed engine support rather than spec text alone.
    strip_solutions(_bout, _bstr, _cfgB)
    _sp = [p.text for p in Document(_bstr).paragraphs]
    _op = [p.text for p in Document(_bsrc).paragraphs]
    check('BANNER-STRIPPED-CLEAN', _sp == _op)
    # Idempotent: re-setting replaces, never stacks.
    set_coverage_banner(_bout, _cfgB, 'Batch 2 of 6 - Q1..Q2 explained of 6. NOT FINAL.')
    _b2 = [p.text for p in Document(_bout).paragraphs]
    check('BANNER-IDEMPOTENT',
          sum(1 for t in _b2 if t.startswith(_cfgB.labels['coverage_banner'])) == 1
          and 'Batch 2 of 6' in _b2[0])
    # Removable (final delivery may drop it).
    set_coverage_banner(_bout, _cfgB, None)
    _b3 = [p.text for p in Document(_bout).paragraphs]
    check('BANNER-REMOVABLE',
          not any(t.startswith(_cfgB.labels['coverage_banner']) for t in _b3))

    # ══ v2.3 FIGURE EMISSION (GAP-2026-08-19-EXPLAIN-REPRESENTATION-EMISSION) ═
    _figpng = os.path.join(tempfile.gettempdir(), 'st_fig.png')
    open(_figpng, 'wb').write(_tiny_png().getvalue())
    _VREC = {'renderer': 'selftest', 'intended': 'C=C(C)CC',
             'derived': 'C=C(C)CC', 'match': True}
    def _figblk(figs, q=1):
        return ExplanationBlock(q=q, ca=4, cfg=cfg,
            axiom=['A governing principle stated once.'],
            deduction=['The first step names what the figure shows.',
                       'The second step completes it, giving Option 4.'],
            why_wrong={1: ['A value swapped - a.'], 2: ['A sign flipped - b.'],
                       3: ['A unit slipped - c.']},
            figures=figs)
    # FIG-VALIDATE-GOOD: a proved figure constructs and validates.
    _fg = RepresentationFigure(_figpng, 4.0, dict(_VREC))
    check('FIG-VALIDATE-GOOD', _figblk([_fg]).validate() is True)
    # FIG-REJECTS: every §6A-5 breach fails AT CONSTRUCTION (validate), each for
    # its own reason — no record, match False, inconsistent record, missing file,
    # bad width, and a non-RepresentationFigure object.
    check('FIG-REJECTS-NO-RECORD', _raises(
        lambda: _figblk([RepresentationFigure(_figpng, 4.0, {})]).validate()))
    check('FIG-REJECTS-MATCH-FALSE', _raises(
        lambda: _figblk([RepresentationFigure(_figpng, 4.0,
            {'renderer': 't', 'intended': 'A', 'derived': 'B',
             'match': False})]).validate()))
    check('FIG-REJECTS-INCONSISTENT', _raises(
        lambda: _figblk([RepresentationFigure(_figpng, 4.0,
            {'renderer': 't', 'intended': 'A', 'derived': 'B',
             'match': True})]).validate()))
    check('FIG-REJECTS-MISSING-FILE', _raises(
        lambda: _figblk([RepresentationFigure(
            _figpng + '.absent', 4.0, dict(_VREC))]).validate()))
    check('FIG-REJECTS-BAD-WIDTH', _raises(
        lambda: _figblk([RepresentationFigure(_figpng, 9.0,
                                              dict(_VREC))]).validate()))
    check('FIG-REJECTS-WRONG-TYPE', _raises(
        lambda: _figblk(['not-a-figure']).validate()))
    # FIG-E2E on a FIGURAL SOURCE: the sample paper already carries images inside
    # question regions; a declared explanation figure must land, all three gates
    # must pass, fidelity must still hold (question-region media untouched), and
    # the region drawing counter must NOT confuse stem images with explanation
    # figures.
    _fsrc = os.path.join(tempfile.gettempdir(), 'st_figsrc.docx')
    _fout = os.path.join(tempfile.gettempdir(), 'st_figout.docx')
    _cfgG = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4)
    _make_sample_paper(_fsrc, _cfgG, nq=6)
    _gb = _figblk([RepresentationFigure(_figpng, 4.0, dict(_VREC),
                                        after_step=1)], q=1)
    build_interleaved_docx(_fsrc, [_gb], _fout, _cfgG)
    _okf, _ = verify_fidelity(_fout, _fsrc, _cfgG)
    _oks, _ = verify_structure(_fout, [_gb], _cfgG, expected_qs=[1])
    _oke, _pe = verify_explanations(_fout, [_gb], _cfgG, expected_qs=[1])
    check('FIG-E2E-GATES', _okf and _oks and _oke)
    # FIG-READER-INVISIBLE: the figure paragraph carries no text, so the strict
    # reader reproduces the block WITHOUT it and deduction count is unchanged.
    _rb = parse_solution_blocks(_fout, _cfgG, expected_qs=[1])
    check('FIG-READER-INVISIBLE',
          len(_rb[1].deduction) == 2 and not getattr(_rb[1], 'figures', []))
    # FIG-SILENT-SKIP-CAUGHT: a doc built WITHOUT the declared figure fails the
    # landing check LOUDLY (§6A-4: never silent).
    _nb = _figblk([], q=1)
    build_interleaved_docx(_fsrc, [_nb], _fout, _cfgG)
    _nb.figures = [RepresentationFigure(_figpng, 4.0, dict(_VREC))]
    _okn, _pn = verify_explanations(_fout, [_nb], _cfgG, expected_qs=[1])
    check('FIG-SILENT-SKIP-CAUGHT',
          not _okn and any('figure landing mismatch' in p for p in _pn))

    # ══ v2.2 TIER-3 NOTATION GUARD (GAP-2026-08-19-EXPLAIN-MATH-NOTATION) ═════
    # Each MIS-COMPILING spelling must RAISE; each CORRECT spelling must PASS.
    # These lock the guard against regression in BOTH directions — a guard that
    # only ever raises is as broken as one that never does.
    _T3_BAD = ('Delta_{o}', 'K_sp', 'E_cell^0', 'sqrt(3RT/M)', 'H_2SO_4',
               'theta', 'x^10', 'int_{0}^{1}')
    _T3_GOOD = ('\\Delta_{o}', 'K_{sp}', 'E_{cell}^{0}', '\\sqrt{3RT/M}',
                'H_{2}SO_{4}', '\\theta', 'x^{10}', '\\frac{0.05912}{2}',
                'V_{B} = \\frac{\u210f}{2m\u0394x}', '\\bar{A}\\vec{E}', '10^{-2}')
    check('T3-NOTATION-REJECTS-BAD',
          all(_raises(lambda b=b: guard_sentence(f'Value \u27e6MATH:{b}\u27e7 here.'))
              for b in _T3_BAD))
    check('T3-NOTATION-ACCEPTS-GOOD',
          all(guard_sentence(f'Value \u27e6MATH:{g}\u27e7 here.') for g in _T3_GOOD))
    # The guard must fire on the REGION BODY only — never on ordinary prose that
    # merely contains a Greek word or an underscore outside a math region.
    check('T3-NOTATION-PROSE-UNTOUCHED',
          guard_sentence('The delta between the two theta values is small.')
          and guard_sentence('Beta decay converts a neutron into a proton.'))
    # A notation breach must be caught at AUTHORING time (block construction),
    # not merely degraded at render time.
    def _notation_block():
        ExplanationBlock(
            q=1, ca=1, cfg=cfg,
            axiom=['The splitting is \u27e6MATH:Delta_o\u27e7 for this ion.'],
            deduction=['First step here.', 'The answer is Option A.'],
            why_wrong={2: ['A value swapped - wrong.'], 3: ['A sign flipped - wrong.'],
                       4: ['A unit slipped - wrong.']}).validate()
    check('T3-NOTATION-FAILS-AT-CONSTRUCTION', _raises(_notation_block))

    # ══ v2.6 EXECUTION-INTEGRITY GUARDS (GAP-2026-08-19-EXPLANATION-EXECUTION-
    #    INTEGRITY) — each guard locked in BOTH directions. ════════════════════
    # D1 — an internal error-taxonomy token in student text RAISES ...
    check('TAG-LEAK-RAISES', _raises(
        lambda: guard_sentence('regiochemistry_error: the para product forms here.')))
    check('TAG-LEAK-RAISES-PARENTHETICAL', _raises(
        lambda: guard_sentence('Option 2 swaps a value (value_swap) here.')))
    # ... while natural-language diagnosis PASSES (no token, same content).
    check('TAG-FREE-PASSES', bool(
        guard_sentence('Option 2 uses the right value for the wrong quantity.')))
    # ... and the block-level chokepoint fires too (validate(), not just the
    # bare guard) — the path every real explanation takes.
    check('TAG-LEAK-FAILS-AT-CONSTRUCTION', _raises(
        lambda: ExplanationBlock(q=1, ca=1, cfg=cfg,
            axiom=['The rule under test is stated here.'],
            deduction=['A first step yields a value.',
                       'The value maps to Option 1 here.'],
            why_wrong={2: ['partial_truth: this misses one condition.'],
                       3: ['Option 3 subtracts the term that must be added.'],
                       4: ['Option 4 satisfies one condition and misses the other.']}
            ).validate()))
    # D1 read-back — a token in the RENDERED bytes is caught by
    # verify_explanations even if the writer path were bypassed.
    _tagdir = tempfile.mkdtemp()
    _tagsrc = os.path.join(_tagdir, 's.docx'); _tagout = os.path.join(_tagdir, 'o.docx')
    _make_sample_paper(_tagsrc, cfg, nq=1)
    _tagblk = ExplanationBlock(q=1, ca=1, cfg=cfg,
        axiom=['The rule under test is stated here.'],
        deduction=['A first step yields a value.', 'The value maps to Option 1 here.'],
        why_wrong={2: ['Option 2 uses the right value for the wrong quantity.'],
                   3: ['Option 3 subtracts the term that must be added.'],
                   4: ['Option 4 satisfies one condition and misses the other.']})
    build_interleaved_docx(_tagsrc, [_tagblk], _tagout, cfg)
    _tagdoc = Document(_tagout)
    for _tp in _tagdoc.paragraphs:
        if _tp.text.startswith('Option 2 uses the right value'):
            _tp.runs[0].text = 'sign_error: ' + _tp.runs[0].text
            break
    _tagdoc.save(_tagout)
    _tok, _tpr = verify_explanations(_tagout, [_tagblk], cfg, expected_qs=[1])
    check('TAG-READBACK-CAUGHT',
          not _tok and any('error-taxonomy token' in p for p in _tpr))
    # D2 — an AXIOM naming an option RAISES; the same AXIOM without it passes
    # (the clean shape is every other fixture in this file).
    check('AXIOM-OPTION-LEAK-RAISES', _raises(
        lambda: ExplanationBlock(q=1, ca=1, cfg=cfg,
            axiom=['The rule always selects Option 1 in such cases.'],
            deduction=['A first step yields a value.',
                       'The value maps to Option 1 here.'],
            why_wrong={2: ['Option 2 uses the right value for the wrong quantity.'],
                       3: ['Option 3 subtracts the term that must be added.'],
                       4: ['Option 4 satisfies one condition and misses the other.']}
            ).validate()))
    # D3 — a VISUAL verdict with no figure RAISES; the same verdict with its
    # figure passes; PROSE with no figure passes; an unknown verdict RAISES.
    def _vblk(verdict, figs):
        return ExplanationBlock(q=1, ca=4, cfg=cfg,
            axiom=['A governing principle stated once.'],
            deduction=['The first step names what the figure shows.',
                       'The second step completes it, giving Option 4.'],
            why_wrong={1: ['A value swapped - a.'], 2: ['A sign flipped - b.'],
                       3: ['A unit slipped - c.']},
            figures=figs, representation_verdict=verdict)
    check('VERDICT-VISUAL-NO-FIG-RAISES',
          _raises(lambda: _vblk('STRUCTURE_GRAPH', []).validate()))
    check('VERDICT-VISUAL-WITH-FIG-PASSES',
          _vblk('STRUCTURE_GRAPH',
                [RepresentationFigure(_figpng, 4.0, dict(_VREC))]).validate() is True)
    check('VERDICT-PROSE-NO-FIG-PASSES', _vblk('PROSE', []).validate() is True)
    check('VERDICT-UNKNOWN-RAISES',
          _raises(lambda: _vblk('DIAGRAM', []).validate()))
    # D4 — a NAT block's bad figure now fails at construction (before v2.6 the
    # figure loop sat after NAT's return and never ran for NAT).
    _cfgNF = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4,
                          options_by_q={1: 0})
    check('NAT-FIG-VALIDATED', _raises(
        lambda: ExplanationBlock(q=1, ca='47', cfg=_cfgNF, qtype='nat',
            axiom=['The mean equals the sum over the count here.'],
            deduction=['Divide the total by the count here.',
                       'The value is 47 here.'],
            common_pitfalls={'235': ['Forgetting to divide leaves 235 unchanged.']},
            figures=[RepresentationFigure(_figpng, 4.0, {})]).validate()))
    # ... and a NAT block with a PROVED figure validates (the positive lock).
    check('NAT-FIG-GOOD-PASSES',
          ExplanationBlock(q=1, ca='47', cfg=_cfgNF, qtype='nat',
            axiom=['The mean equals the sum over the count here.'],
            deduction=['Divide the total by the count here.',
                       'The value is 47 here.'],
            common_pitfalls={'235': ['Forgetting to divide leaves 235 unchanged.']},
            figures=[RepresentationFigure(_figpng, 4.0, dict(_VREC))]
            ).validate() is True)
    # ... and an anomaly block carrying a figure RAISES (figures are content).
    check('ANOMALY-FIG-RAISES', _raises(
        lambda: ExplanationBlock(q=1, anomaly='no defensible answer', cfg=cfg,
            figures=[RepresentationFigure(_figpng, 4.0, dict(_VREC))]).validate()))

    # ── v2.7 GAP-2026-08-20-TRANSFER-SAFE-EXPLANATIONS fixtures ───────────
    def _blk27(**kw):
        base = dict(q=1, ca=1, cfg=cfg,
            axiom=['A governing relation links the two quantities here.'],
            deduction=['A first step yields a value.', 'The value maps to Option 1 here.'],
            why_wrong={2: ['Option 2 uses the right value for the wrong quantity.'],
                       3: ['Option 3 subtracts the term that is added.'],
                       4: ['Option 4 satisfies one condition and misses the other.']})
        base.update(kw)
        return ExplanationBlock(**base)
    # (a) absolute gate — AXIOM
    check('ABS-AXIOM-RAISES', _raises(lambda: _blk27(
        axiom=['Electron-withdrawing groups always direct the attack to the meta position.']).validate()))
    check('ABS-AXIOM-DECLARED-OK', _blk27(
        axiom=['Mass number never changes in a beta emission.'],
        absolutes_justified={'Mass number never changes in a beta emission.':
                             'conservation of nucleon number'}).validate() is True)
    check('ABS-SCOPING-CLAUSE-OK', _blk27(
        axiom=['The rule applies only when the two beta sites compete and the base is hindered.']).validate() is True)
    check('ABS-QUANTIFIER-NOT-GATED', _blk27(
        axiom=['The salt releases exactly three ions, only two of which are chloride, for every formula unit.']).validate() is True)
    check('ABS-REGARDLESS-RAISES', _raises(lambda: _blk27(
        axiom=['Translation carries three quadratic terms regardless of how the molecule is built.']).validate()))
    check('ABS-CANNOT-RAISES', _raises(lambda: _blk27(
        axiom=['A primary alcohol cannot react under these conditions.']).validate()))
    check('ABS-AT-ALL-RAISES', _raises(lambda: _blk27(
        why_wrong={2: ['Dissolved oxygen cannot be titrated directly at all.'],
                   3: ['Option 3 subtracts the term that is added.'],
                   4: ['Option 4 satisfies one condition and misses the other.']}).validate()))
    check('ABS-SPEEDHACK-RAISES', _raises(lambda: _blk27(
        speed_hack=['A molecule with identical halves always collapses one pair into a meso form, giving Option 1.']).validate()))
    check('ABS-DEDUCTION-NOT-GATED', _blk27(
        deduction=['Beta emission never changes the mass number here.',
                   'The value maps to Option 1 here.']).validate() is True)
    check('ABS-DECLARED-ABSENT-RAISES', _raises(lambda: _blk27(
        absolutes_justified={'A sentence not in the block never appears.': 'definition'}).validate()))
    check('ABS-MATH-REGION-MASKED', find_absolute('The ratio is ⟦MATH:\\frac{a}{b}⟧ under these conditions.') is None)
    check('ABS-NAT-PITFALL-RAISES', _raises(lambda: ExplanationBlock(q=1, ca='6', cfg=EngineConfig(
        r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 0), qtype='nat',
        axiom=['Terminal and bridging roles are counted separately here.'],
        deduction=['Four hydrogens bridge.', 'The remaining count is 6.'],
        common_pitfalls={'4': ['A primary alcohol never clouds the reagent at all.']}).validate()))
    # per-language override replaces the default pattern
    _cfg_abs = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4, absolute_terms_re=r'\bsiempre\b')
    check('ABS-CONFIG-OVERRIDE', _raises(lambda: _blk27(cfg=_cfg_abs,
        axiom=['La regla siempre se cumple aqui.']).validate()) and
        _blk27(cfg=_cfg_abs, axiom=['This rule always holds here.']).validate() is True)
    # (b) learner-psychology templates
    check('PSYCH-TEMPLATE-RAISES', _raises(lambda: _blk27(
        why_wrong={2: ['The seductive half is that both share one metal ion.'],
                   3: ['Option 3 subtracts the term that is added.'],
                   4: ['Option 4 satisfies one condition and misses the other.']}).validate()))
    check('PSYCH-TEMPLATE-GUARD', _raises(lambda: guard_sentence('A hasty solver picks this because it looks symmetric.')))
    # (c) transfer record
    _TR_OK = [{'section': 'AXIOM', 'claim': 'back-donation lowers the C-O stretch',
               'epistemic_type': 'SCIENTIFIC_GENERAL_RULE', 'scope': 'isoelectronic carbonyls',
               'neighbour_tested': 'a carbonyl with a different ligand set', 'outcome': 'NARROWED'}]
    check('TR-OK', _blk27(transfer_record=_TR_OK).validate() is True)
    check('TR-EMPTY-RAISES', _raises(lambda: _blk27(transfer_record=[]).validate() if False else
                                     validate_transfer_record([], ctx='Q1')))
    check('TR-MISSING-FIELD-RAISES', _raises(lambda: _blk27(transfer_record=[
        {'section': 'AXIOM', 'claim': 'x', 'epistemic_type': 'SCIENTIFIC_GENERAL_RULE',
         'scope': 's', 'outcome': 'SAFE'}]).validate()))
    check('TR-BAD-TYPE-RAISES', _raises(lambda: _blk27(transfer_record=[
        dict(_TR_OK[0], epistemic_type='FOLKLORE')]).validate()))
    check('TR-QSPEC-IN-AXIOM-RAISES', _raises(lambda: _blk27(transfer_record=[
        dict(_TR_OK[0], epistemic_type='QUESTION_SPECIFIC_INFERENCE', outcome='SAFE')]).validate()))
    check('TR-QSPEC-MOVED-OK', _blk27(transfer_record=[
        dict(_TR_OK[0], epistemic_type='QUESTION_SPECIFIC_INFERENCE', outcome='MOVED_TO_DEDUCTION')]).validate() is True)
    check('TR-OPTSET-OUTSIDE-SH-RAISES', _raises(lambda: _blk27(transfer_record=[
        dict(_TR_OK[0], epistemic_type='OPTION_SET_SHORTCUT', outcome='SAFE')]).validate()))
    check('TR-NO-AXIOM-CLAIM-RAISES', _raises(lambda: _blk27(transfer_record=[
        dict(_TR_OK[0], section='WHY_WRONG')]).validate()))
    check('TR-SH-WITHOUT-CLAIM-RAISES', _raises(lambda: _blk27(
        speed_hack=['Check the sign first, a move that suits any cell quoted as a diagram, giving Option 1.'],
        transfer_record=_TR_OK).validate()))
    check('TR-SH-WITH-CLAIM-OK', _blk27(
        speed_hack=['Check the sign first, a move that suits any cell quoted as a diagram, giving Option 1.'],
        transfer_record=_TR_OK + [{'section': 'SPEED_HACK', 'claim': 'sign first',
            'epistemic_type': 'OPTION_SET_SHORTCUT', 'scope': 'when both signs are offered',
            'neighbour_tested': 'a cell whose options share a sign', 'outcome': 'SAFE'}]).validate() is True)
    # (d) CONFORMER verdict: visual, requires a figure
    check('CONFORMER-NO-FIG-RAISES', _raises(lambda: _blk27(representation_verdict='CONFORMER').validate()))
    check('CONFORMER-WITH-FIG-OK', _blk27(representation_verdict='CONFORMER', ca=4,
        deduction=['The first step names what the figure shows.', 'The second step completes it, giving Option 4.'],
        why_wrong={1: ['A value swapped - a.'], 2: ['A sign flipped - b.'], 3: ['A unit slipped - c.']},
        figures=[RepresentationFigure(_figpng, 4.0, dict(_VREC))]).validate() is True)
    # (e) strip_solutions media GC: a figure-bearing paper strips back to the
    # source's media count, not source + explanation figures.
    _gcdir = tempfile.mkdtemp()
    _gcsrc = os.path.join(_gcdir, 's.docx'); _gcout = os.path.join(_gcdir, 'o.docx')
    _gcstr = os.path.join(_gcdir, 'st.docx')
    _make_sample_paper(_gcsrc, cfg, nq=1)
    build_interleaved_docx(_gcsrc, [_blk27(ca=4,
        deduction=['The first step names what the figure shows.', 'The second step completes it, giving Option 4.'],
        why_wrong={1: ['A value swapped - a.'], 2: ['A sign flipped - b.'], 3: ['A unit slipped - c.']},
        figures=[RepresentationFigure(_figpng, 4.0, dict(_VREC))],
        representation_verdict='STRUCTURE_GRAPH')], _gcout, cfg)
    strip_solutions(_gcout, _gcstr, cfg)
    def _n_media(pth):
        import zipfile
        with zipfile.ZipFile(pth) as _z:
            return sum(1 for _n in _z.namelist() if _n.startswith('word/media/'))
    check('STRIP-MEDIA-GC', _n_media(_gcout) == _n_media(_gcsrc) + 1 and
                            _n_media(_gcstr) == _n_media(_gcsrc))
    _okgc, _ = verify_fidelity(_gcout, _gcsrc, cfg)
    check('STRIP-MEDIA-GC-FIDELITY-KEPT', _okgc)

    # ───────────── v2.8 — GAP-2026-08-21-EXPLANATION-PROVENANCE fixtures ─────────────
    EngineConfig.DEFAULT_PROVENANCE_GATES = True
    cfg8 = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4)
    v8_tr = [{'section': 'AXIOM', 'claim': 'c', 'epistemic_type': 'SCIENTIFIC_GENERAL_RULE',
            'scope': 's', 'neighbour_tested': 'n', 'outcome': 'SAFE',
            'neighbour_source': 'GENERATED'}]
    def v8_mcq(**kw):
        base = dict(q=1, ca=2, cfg=cfg8,
                    axiom=['A first-order half-life is ln 2 over k.'],
                    deduction=['Substituting k gives 100 s.', 'So the answer is Option 2.'],
                    why_wrong={1: ['Halving k doubles the half-life to 200 s, not 100 s.'],
                               3: ['Reading 693 as the half-life reports the digits of ln 2 scaled, not 100 s.'],
                               4: ['A value of 50 s would need a rate constant twice the stated one.']},
                    transfer_record=v8_tr,
                    error_provenance={
                        1: {'mode': 'VERIFIED_ERROR_PATH', 'wrong_operation': 'used k/2',
                            'recompute': '0.693/(6.93e-3/2)', 'target': '200'},
                        3: {'mode': 'DIRECT_CONTRADICTION',
                            'contradiction': '693 is neither ln2/k nor 1/k at this k'},
                        4: {'mode': 'VERIFIED_ERROR_PATH', 'wrong_operation': 'used 2k',
                            'recompute': '0.693/(2*6.93e-3)', 'target': '50'}})
        base.update(kw); return ExplanationBlock(**base)
    # 1 a fully-provenanced block validates
    try: check('V28-PROV-OK', v8_mcq().validate())
    except ValueError as e: check('V28-PROV-OK', False); print('   V28-PROV-OK:', e)
    # 2 a claimed path that does not reproduce its target raises (the Q17/Q23/Q52 shape)
    check('V28-PROV-FALSE-PATH', _raises(lambda: v8_mcq(error_provenance={
        1: {'mode': 'VERIFIED_ERROR_PATH', 'wrong_operation': 'added without halving',
            'recompute': '2.2+9.4', 'target': '7.2'},
        3: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'},
        4: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'}}).validate()))
    # 3 missing record raises
    check('V28-PROV-MISSING', _raises(lambda: v8_mcq(error_provenance={}).validate()))
    # 4 a record for an option that is not wrong raises
    check('V28-PROV-EXTRAKEY', _raises(lambda: v8_mcq(error_provenance={
        1: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'},
        2: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'},
        3: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'},
        4: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'}}).validate()))
    # 5 DIRECT_CONTRADICTION that smuggles a recompute raises
    check('V28-PROV-DIRECT-NOPATH', _raises(lambda: v8_mcq(error_provenance={
        1: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x', 'recompute': '1+1'},
        3: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'},
        4: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'}}).validate()))
    # 6 recompute with a disallowed construct is refused (no code execution)
    check('V28-PROV-SAFEEVAL', _raises(lambda: _safe_eval('__import__("os").system("x")')))
    check('V28-PROV-SAFEEVAL2', abs(_safe_eval('sqrt(3/2)*395') - 483.77) < 0.01)
    # 7 precision: target 484 accepts 483.77 (integer-rounded target)
    check('V28-NUMMATCH', numbers_match(483.77, '484') and not numbers_match(558, '484')
          and numbers_match(3.085, '3.09') and not numbers_match(11.6, '7.2'))
    # 8 hedged provenance language raises (24 hits on the reference paper)
    for v8_ph in ('or otherwise mishandling the power of ten', 'perhaps by pairing early',
                'or a similar miscombination', 'n(n+2)+something style over-counting'):
        check(f'V28-HEDGE[{v8_ph[:12]}]', _raises(lambda v8_ph=v8_ph: v8_mcq(
            why_wrong={1: [f'Using k/2 {v8_ph} gives 200 s, not 100 s.'],
                       3: ['Reading 693 reports ln 2 scaled, not 100 s.'],
                       4: ['A value of 50 s would need twice the stated k.']}).validate()))
    check('V28-HEDGE-CLEAN', find_hedge('Using n = 1 instead of 2 doubles the term to 0.2368 V.') is None)
    # 9 NAT pitfalls: key IS the target; one pitfall is enough; false arithmetic raises
    def v8_nat(**kw):
        cfg_n = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4, options_by_q={5: 0})
        base = dict(q=5, ca='484', cfg=cfg_n, qtype='nat',
                    axiom=['The rms to most-probable speed ratio is a fixed number.'],
                    deduction=['Multiplying 395 by 1.2247 gives 483.77.', 'Rounded, the answer is 484.'],
                    common_pitfalls={'323': ['Dividing by 1.2247 instead of multiplying gives 323, not 484.']},
                    transfer_record=v8_tr,
                    error_provenance={'323': {'mode': 'VERIFIED_ERROR_PATH',
                                              'wrong_operation': 'divided by sqrt(3/2)',
                                              'recompute': '395/sqrt(3/2)'}})
        base.update(kw); return ExplanationBlock(**base)
    try: check('V28-NAT-PROV-OK', v8_nat().validate())
    except ValueError as e: check('V28-NAT-PROV-OK', False); print('   V28-NAT-PROV-OK:', e)
    check('V28-NAT-PROV-FALSE', _raises(lambda: v8_nat(
        common_pitfalls={'558': ['Using the rms to average ratio near 1.414 gives 558, not 484.']},
        error_provenance={'558': {'mode': 'VERIFIED_ERROR_PATH',
                                  'wrong_operation': 'used rms/avg ratio',
                                  'recompute': '395*sqrt(3*pi/8)'}}).validate()))
    check('V28-NAT-PROV-TARGETMISMATCH', _raises(lambda: v8_nat(
        error_provenance={'323': {'mode': 'VERIFIED_ERROR_PATH', 'wrong_operation': 'x',
                                  'recompute': '395/sqrt(3/2)', 'target': '330'}}).validate()))
    # 10 transfer_record mandatory for an authored block; read-back exempt
    check('V28-TR-MANDATORY', _raises(lambda: v8_mcq(transfer_record=None).validate()))
    v8_rb = v8_mcq(transfer_record=None, error_provenance=None); v8_rb._preserved = True
    try: check('V28-TR-READBACK-EXEMPT', v8_rb.validate())
    except ValueError as e: check('V28-TR-READBACK-EXEMPT', False); print('   readback:', e)
    check('V28-TR-SOURCE-SHAPE', _raises(lambda: v8_mcq(transfer_record=[dict(v8_tr[0],
        neighbour_source='FROM MEMORY')]).validate()))
    # 11 curated triggers: a family named in the AXIOM must be the neighbour cited
    v8_lrn = parse_learnings('# X\n\n## EX-CHEM-007 — Buffer\n\n**Defect code:** NEIGHBOUR-LIBRARY\n\n'
                           '**Pattern:** p\n\n**Prevention rule:** r\n\n**Triggers:** buffer, '
                           'Henderson, re:pH\\s+depends\\s+only\n\n**Verification:** v\n')
    v8_trig = triggers_from_learnings(v8_lrn)
    check('V28-TRIG-PARSE', v8_lrn['rules'][0]['triggers'].startswith('buffer') and len(v8_trig) == 1)
    cfg_t = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4, learnings_triggers=v8_trig)
    v8_ax_b = ['A buffer resists pH change because the ratio sets the pH.']
    check('V28-TRIG-MISSED', _raises(lambda: v8_mcq(cfg=cfg_t, axiom=v8_ax_b).validate()))
    try:
        check('V28-TRIG-CITED', v8_mcq(cfg=cfg_t, axiom=v8_ax_b, transfer_record=[dict(v8_tr[0],
            neighbour_source='CURATED:EX-CHEM-007', outcome='NARROWED')]).validate())
    except ValueError as e: check('V28-TRIG-CITED', False); print('   cited:', e)
    try: check('V28-TRIG-NOHIT-GENERATED', v8_mcq(cfg=cfg_t).validate())
    except ValueError as e: check('V28-TRIG-NOHIT-GENERATED', False); print('   nohit:', e)
    # 12 tripwire
    v8_fired, v8_sm = transfer_tripwire([v8_mcq(q=i) for i in range(1, 25)])
    check('V28-TRIPWIRE-FIRES', v8_fired and v8_sm['axiom_claims'] == 24)
    check('V28-TRIPWIRE-QUIET', not transfer_tripwire([v8_mcq(q=i) for i in range(1, 5)])[0])
    # 13 formula typography: normaliser + residual gate + idempotence + locants untouched
    v8_n = normalise_formula_text('Across [V(CO)6]-, Cr(CO)6 and [Mn(CO)6]+ the CO pi* orbitals; '
                                'Fe3+ is d5 ion, sp3d2, t2g3eg2, NH4+, C2-C3 bond, H2O, Option 2.')
    check('V28-FMT-NORM', '[V(CO)\u2086]\u207b' in v8_n and 'Fe\u00b3\u207a' in v8_n and 'sp\u00b3d\u00b2' in v8_n
          and 't\u2082g\u00b3eg\u00b2' in v8_n and 'NH\u2084\u207a' in v8_n and 'C2-C3' in v8_n
          and 'H\u2082O' in v8_n and 'Option 2' in v8_n and '\u03c0*' in v8_n and 'd\u2075 ion' in v8_n)
    check('V28-FMT-IDEMPOTENT', normalise_formula_text(v8_n) == v8_n)
    check('V28-FMT-MATHMASK', normalise_formula_text('see \u27e6MATH:H2O\u27e7 here') == 'see \u27e6MATH:H2O\u27e7 here')
    v8_blk = v8_mcq(deduction=['The ion [Fe(CN)6]4- is low-spin.', 'So the answer is Option 2.'])
    check('V28-FMT-APPLIED-AT-INIT', '[Fe(CN)\u2086]\u2074\u207b' in v8_blk.deduction[0])
    check('V28-FMT-RESIDUAL', find_unformatted_formula('[(eta5-C5H5)Mn(CO)3] is 18e') is not None
          and find_unformatted_formula('[Fe(CN)\u2086]\u2074\u207b') is None)
    cfg_f = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4, formula_typography=False)
    check('V28-FMT-SWITCH', '[V(CO)6]-' in v8_mcq(cfg=cfg_f, deduction=['Take [V(CO)6]-.',
                                                                       'Answer Option 2.']).deduction[0])
    # 14 key reconciliation round-trip against paper_pipeline
    try:
        import paper_pipeline as _pp
        v8_com = _pp.seal_key_commitments('MOCK:M01', {1: '2', 5: '484', 7: '2,3'})
        v8_b1 = v8_mcq(); v8_b5 = v8_nat()
        v8_b7 = v8_mcq(q=7, ca={2, 3}, qtype='msq',
                   deduction=['x.', 'So the set is Option 2, Option 3.'],
                   why_wrong={1: ['a.'], 4: ['b.']},
                   error_provenance={1: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'},
                                     4: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'}})
        v8_r = reconcile_key_commitments({1: v8_b1, 5: v8_b5, 7: v8_b7}, {'key_commitments': {'MOCK:M01': v8_com}}, 'MOCK:M01')
        check('V28-KEY-MATCH', v8_r['available'] and v8_r['matched'] == [1, 5, 7] and not v8_r['mismatched'])
        v8_b1b = v8_mcq(ca=3, deduction=['x.', 'So the answer is Option 3.'],
                    why_wrong={1: ['a.'], 2: ['b.'], 4: ['c.']},
                    error_provenance={k: {'mode': 'DIRECT_CONTRADICTION', 'contradiction': 'x'} for k in (1, 2, 4)})
        v8_r2 = reconcile_key_commitments({1: v8_b1b}, {'key_commitments': {'MOCK:M01': v8_com}}, 'MOCK:M01')
        check('V28-KEY-MISMATCH', v8_r2['mismatched'] == [1] and v8_r2['candidates'][1] == '2')
        check('V28-KEY-ABSENT', not reconcile_key_commitments({1: v8_b1}, {}, 'MOCK:M01')['available'])
    except Exception as e:
        check('V28-KEY-MATCH', False); print('   key reconcile:', repr(e))
    # 14b semantic-object agreement through the injected canonicaliser
    v8_so = {'role': 'problem', 'kind': 'STRUCTURE', 'name': 'salicylic acid',
             'canonical': 'OC(=O)c1ccccc1O', 'descriptor': {}}
    if canonical_structure('C')[1] == 'ok':
        check('V28-SEM-AGREE', semantic_objects_agree(v8_so, dict(v8_so, canonical='Oc1ccccc1C(O)=O'))[0]
              and not semantic_objects_agree(v8_so, dict(v8_so, canonical='O=C(C)c1ccccc1O'))[0]
              and canonical_structure('C(C)(C)(C)(C)C')[0] is None)
    else:
        check('V28-SEM-AGREE', semantic_objects_agree(v8_so, dict(v8_so))[0])
    # 15 scan-risk over a rendered doc finds the markers
    try:
        import tempfile, os
        v8_d = tempfile.mkdtemp(); v8_s = os.path.join(v8_d, 's.docx'); v8_o = os.path.join(v8_d, 'o.docx')
        v8_doc = Document(); v8_doc.add_paragraph('Q.1 stem')
        for v8_i in range(1, 5): v8_doc.add_paragraph(f'{v8_i}. opt')
        v8_doc.save(v8_s)
        EngineConfig.DEFAULT_PROVENANCE_GATES = False
        v8_legacy = ExplanationBlock(q=1, ca=2, cfg=EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4,
                                                              formula_typography=False),
            axiom=['Ions always pair in [Fe(CN)6]4- complexes.'],
            deduction=['x.', 'So the answer is Option 2.'],
            why_wrong={1: ['Perhaps by halving k this gives 200 s.'], 3: ['b.'], 4: ['c.']},
            absolutes_justified={'Ions always pair in [Fe(CN)6]4- complexes.': 'fixture'})
        build_interleaved_docx(v8_s, [v8_legacy], v8_o, v8_legacy.cfg)
        EngineConfig.DEFAULT_PROVENANCE_GATES = True
        v8_sc = scan_risk_markers(v8_o, EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4))
        check('V28-SCANRISK', 1 in v8_sc and v8_sc[1]['hedge'] and v8_sc[1]['absolute'] and v8_sc[1]['formula'])
    except Exception as e:
        EngineConfig.DEFAULT_PROVENANCE_GATES = True
        check('V28-SCANRISK', False); print('   scan-risk:', repr(e))

    # v2.9 — END-OF-MOCK REPORT DOCX (GAP-2026-08-26-REGISTRY-HANDOFF-SEAM)
    try:
        import tempfile as _tf, os as _os
        _rp = _os.path.join(_tf.mkdtemp(), 'EX_Mock01_Explain_Report.docx')
        build_report_docx(_rp, 'END-OF-MOCK REPORT — EX Mock01',
                          [('§R2 VERDICT', ['SHIP']),
                           ('§R3 COVERAGE', ['60/60 explained', 'SPEED HACK on Q.3 Q.9']),
                           ('§R5 DERIVATION-CONFIDENCE', [])],
                          meta_lines=['Framework 2026.08.26.3'])
        _rr = read_report_docx(_rp)
        check('V29-REPORT-ROUNDTRIP', _rr['title'] == 'END-OF-MOCK REPORT — EX Mock01'
              and [h for h, _ in _rr['sections']] == ['§R2 VERDICT', '§R3 COVERAGE',
                                                     '§R5 DERIVATION-CONFIDENCE']
              and _rr['sections'][1][1] == ['60/60 explained', 'SPEED HACK on Q.3 Q.9']
              and _rr['sections'][2][1] == ['(empty — stated explicitly, never omitted)'])
        check('V29-REPORT-NAME-DISJOINT', not _rp.endswith('_Explanation.docx'))
        check('V29-REPORT-MANDATE0', _raises(lambda: build_report_docx(
            _rp, 't', [('§R7', ['Q.4 answer: 3'])])))
    except Exception as e:
        check('V29-REPORT-ROUNDTRIP', False); print('   report:', repr(e))

    passed = sum(1 for _, ok in results if ok)
    total = len(results)
    for name, ok in results:
        if not ok:
            print(f'  FAIL: {name}')
    print(f'SELF-TEST: {passed}/{total} PASS')
    return passed == total


# ══════════════════════════════════════════════════════════════════════════════
# EXPLANATION-AUDITOR READER — the docx->ExplanationBlock reader.
# (Added for canonical Step 10, retired 2026.08.03.5; its later consumer PYQExplainAudit
#  was retired 2026.08.09.1. UNCHANGED and still live/self-tested, kept for any legacy
#  _Complete.docx round-trip; no active step now calls it.)
# Additive: no existing write/verify path changes; --self-test stays 44/44. The
# EXACT INVERSE of _block_paragraphs, driven by the SAME cfg, so a rectified block
# is structurally identical to how Step 9 would have written it correctly.
# ══════════════════════════════════════════════════════════════════════════════
def _omml_to_source(node, q, strict=True):
    """SOURCE-TEXT form of one OMML child (v2.1). Explanation prose contains
    digit/digit auto-fractions AND arbitrary OMML structures. strict=True (the
    Step-5 READER): a bare digit/digit m:f round-trips as 'num/den' (adjacency is
    resolved by the caller _para_source); every other OMML structure is preserved
    LOSS-LESSLY as ⟦M:<b64>⟧ via _encode_opaque, so build re-emits its exact
    <m:oMath> body (a byte-faithful round-trip, body included — not a bodiless
    token). strict=False (the VERIFIER): unknown OMML degrades to its m:t text so a
    post-render audit never crashes."""
    tag = node.tag
    if tag == qn('m:r'):
        return ''.join(t.text or '' for t in node.iter(qn('m:t')))
    if tag == qn('m:f'):
        num = node.find(qn('m:num')); den = node.find(qn('m:den'))
        nt = ''.join(num.itertext()) if num is not None else ''
        dt = ''.join(den.itertext()) if den is not None else ''
        if strict and not (nt.strip().isdigit() and dt.strip().isdigit()):
            return _encode_opaque(node)   # v2.1: loss-less preserve-and-reemit
        return f'{nt}/{dt}'
    # v2.1: a non-fraction OMML structure is a FIRST-CLASS citizen of explanation
    # prose. The strict Step-5 reader preserves it LOSS-LESSLY as ⟦M:<b64>⟧ so
    # build_interleaved_docx re-emits the exact <m:oMath> (byte-faithful
    # round-trip, body included). The non-strict verifier keeps degrading unknown
    # OMML to its m:t text so a post-render audit never crashes.
    return _encode_opaque(node) if strict else ''.join(t.text or '' for t in node.iter(qn('m:t')))

def _is_digit_frac(mf):
    """True iff an m:f is a bare digit/digit fraction (num and den are integers)."""
    n = mf.find(qn('m:num')); d = mf.find(qn('m:den'))
    nt = ''.join(n.itertext()) if n is not None else ''
    dt = ''.join(d.itertext()) if d is not None else ''
    return nt.strip().isdigit() and dt.strip().isdigit()

def _para_source(p_el, q, strict=True):
    """Rebuild the source string add_math_text was given: text runs verbatim,
    m:f as 'num/den', in document order. strict is passed to _omml_to_source.

    v2.1 — adjacency-safe fraction round-trip. A bare digit/digit fraction flattens
    to 'num/den' ONLY when that survives the authoring emitter in its actual
    textual context (standalone, e.g. '235/5'). When it abuts a word character —
    '1/2k', 'k1/2', two adjacent fractions — the flattened form would NOT re-parse
    (_SIMPLE_FRAC refuses it and add_math_text would raise), so the fraction is
    preserved LOSS-LESSLY as an opaque ⟦M:<b64>⟧ token instead. Non-strict
    (verifier) output is unchanged."""
    if not strict:
        out = []
        for child in p_el:
            tag = child.tag
            if tag == qn('w:pPr'):
                continue
            if tag == qn('m:oMath'):
                for m in child:
                    out.append(_omml_to_source(m, q, strict))
            else:
                out.append(''.join(t.text or '' for t in child.iter(qn('w:t'))))
        return ''.join(out)
    # strict: build typed segments so a fraction's neighbours can be inspected.
    segs = []                      # ('text', s) | ('frac', 'n/d', node)
    for child in p_el:
        tag = child.tag
        if tag == qn('w:pPr'):
            continue
        if tag == qn('m:oMath'):
            for m in child:
                if m.tag == qn('m:f') and _is_digit_frac(m):
                    num = m.find(qn('m:num')); den = m.find(qn('m:den'))
                    nd = f"{''.join(num.itertext())}/{''.join(den.itertext())}"
                    segs.append(('frac', nd, m))
                else:
                    segs.append(('text', _omml_to_source(m, q, strict)))
        else:
            segs.append(('text', ''.join(t.text or '' for t in child.iter(qn('w:t')))))
    raw = [s[1] for s in segs]
    out = []
    for i, s in enumerate(segs):
        if s[0] != 'frac':
            out.append(s[1]); continue
        left = ''.join(out)
        right = ''.join(raw[i + 1:])
        probe = (left[-1:] if left else '') + s[1] + (right[:1] if right else '')
        out.append(s[1] if _SIMPLE_FRAC.search(probe) else _encode_opaque(s[2]))
    return ''.join(out)

def _label_to_index(cfg, n):
    """Inverse of cfg.option_label over 1..n."""
    return {cfg.option_label(i): i for i in range(1, n + 1)}

def _para_spacing(p_el):
    """(before, after) twips from the paragraph's spacing, or None if not present."""
    pr = p_el.find(qn('w:pPr'))
    sp = pr.find(qn('w:spacing')) if pr is not None else None
    if sp is None:
        return None
    b, a = sp.get(qn('w:before')), sp.get(qn('w:after'))
    try:
        return (int(b), int(a)) if (b is not None and a is not None) else None
    except ValueError:
        return None

def _is_subheader(p_el, terminators='.!?'):
    """True if the paragraph is a rendered sub-header (an 'Option X' key or a
    COMMON-PITFALLS value header). The writer gives every sub-header before>after
    spacing while sentences get before<after, so this is structural and works even
    when the value is pure OMML (a fraction) with empty display text. Falls back to
    a text heuristic only when spacing is absent (a non-engine-written doc).
    `terminators` is language-configurable (pass cfg.sentence_terminators)."""
    sp = _para_spacing(p_el)
    if sp is not None:
        return sp[0] > sp[1]
    s = _para_source(p_el, 0, strict=False).strip()  # fallback for spacing-less docs
    return bool(s) and sentence_count(s, terminators) <= 1 and s[-1] not in terminators \
        and len(s) <= 24 and s.count(' ') <= 3

def parse_solution_blocks(path, cfg, expected_qs=None):
    """Read a Step-4 Solutions docx back into {q: ExplanationBlock}. Inverse of
    build_interleaved_docx's per-block render, driven by the same cfg (labels /
    markers / label-scheme / options). MCQ / MSQ / NAT, any label scheme, any
    language - all via cfg, nothing hardcoded. Was used by the explanation auditor
    (PYQExplainAudit, retired 2026.08.09.1) to reason about each
    explanation and to rebuild corrected blocks; kept live for legacy round-trips."""
    doc = Document(path)
    ca_label = cfg.labels['correct_answer'].lower(); opt_word = cfg.labels['option']
    H = {k: f"{cfg.markers.get(k, '')} {cfg.labels.get(k, k)}".strip()
         for k in ('axiom', 'deduction', 'speed_hack', 'why_wrong', 'common_pitfalls')}
    HREV = {v: k for k, v in H.items()}
    ca_prefix = ca_label + ':'
    segs = {}; cur = None; in_expl = False
    for p in doc.paragraphs:
        t = p.text.strip()
        mq = cfg.q_re.match(t)
        if mq:
            cur = int(mq.group(1)); in_expl = False; continue
        if cur is None:
            continue
        if t.lower().startswith(ca_prefix):
            in_expl = True
            segs[cur] = {'ca': _para_source(p._p, cur).strip(), 'paras': []}
            continue
        if in_expl:
            _src = _para_source(p._p, cur).strip()
            if t or _src:                       # keep OMML-only paras; drop blank separators
                segs[cur]['paras'].append((t, _src, p._p))
    want = sorted(segs) if expected_qs is None else sorted(set(expected_qs))
    blocks = {}
    for q in want:
        if q not in segs:
            raise ValueError(f'Q{q}: no explanation region found in {path}')
        ca_src = segs[q]['ca']; body = segs[q]['paras']; n = cfg.expected_options(q)
        disp = ca_src[len(ca_prefix):].strip(); ca_range = None
        if n == 0:
            qtype = 'nat'
            # v1.16: portal grading format is EXACTLY '{lo}-{hi}' for a range, or
            # a plain point value otherwise — no parentheses, no 'accepted range'
            # wording, no en-dash (that format is retired; the charset guard in
            # ExplanationBlock.validate()/render would now reject it outright).
            # A range never has a leading '-' (negative-bounded ranges are a
            # locked NOT-SUPPORTED case — see format_nat_range), so a leading
            # '-' unambiguously means a signed point value, never a range.
            _rng_m = re.match(r'^(\d+(?:\.\d+)?)-(\d+(?:\.\d+)?)$', disp)
            if _rng_m:
                ca_range = (float(_rng_m.group(1)), float(_rng_m.group(2)))
                # the point value is not recoverable from a range-format line
                # by design (see ExplanationBlock.validate() ca_range branch)
                ca = None
            else:
                ca = disp
        else:
            l2i = _label_to_index(cfg, n)
            idxs = []
            for lb in [x.strip() for x in disp.split(',') if x.strip()]:
                if lb not in l2i:
                    raise ValueError(f'Q{q}: CA label {lb!r} not in scheme 1..{n}')
                idxs.append(l2i[lb])
            if len(idxs) == 1:
                qtype = 'mcq'; ca = idxs[0]
            else:
                qtype = 'msq'; ca = set(idxs)
        axiom, deduction, speed = [], [], []
        why_wrong, pitfalls = {}, {}
        mode = None; key = None
        sub_re = re.compile(rf'^{re.escape(opt_word)}\s+(\S+)\s*$')
        for disp, s, p_el in body:
            if disp in HREV:                     # block headers are plain text
                mode = HREV[disp]; key = None; continue
            if mode == 'axiom':
                axiom.append(s)
            elif mode == 'deduction':
                deduction.append(s)
            elif mode == 'speed_hack':
                speed.append(s)
            elif mode == 'why_wrong':
                m = sub_re.match(disp)           # 'Option X' sub-headers are plain text
                if m and _is_subheader(p_el, cfg.sentence_terminators):
                    l2i = _label_to_index(cfg, n)
                    if m.group(1) not in l2i:
                        raise ValueError(f'Q{q}: WHY WRONG label {m.group(1)!r} not in 1..{n}')
                    key = l2i[m.group(1)]; why_wrong[key] = []
                else:
                    if key is None:
                        raise ValueError(f'Q{q}: WHY WRONG sentence before any Option sub-header')
                    why_wrong[key].append(s)
            elif mode == 'common_pitfalls':
                # value headers carry the writer's sub-header spacing (before > after) and
                # may be pure OMML (a fractional wrong value) with empty display text, so
                # detect structurally, not by a text heuristic.
                if _is_subheader(p_el, cfg.sentence_terminators):
                    key = s; pitfalls[key] = []
                else:
                    if key is None:
                        raise ValueError(f'Q{q}: COMMON PITFALLS sentence before any value header')
                    pitfalls[key].append(s)
        kwargs = dict(q=q, cfg=cfg, qtype=qtype, ca=ca, ca_range=ca_range,
                      axiom=axiom, deduction=deduction,
                      speed_hack=(speed or None))
        kwargs['common_pitfalls' if qtype == 'nat' else 'why_wrong'] = (
            pitfalls if qtype == 'nat' else why_wrong)
        blocks[q] = ExplanationBlock(**kwargs)
        blocks[q]._preserved = True   # v2.1: content lifted verbatim from a shipped
                                      # doc re-emits faithfully (see _block_paragraphs)
    return blocks

def parse_learnings(path):
    """Read an EXPLAIN_AUDIT_LEARNINGS / EXPLAIN_LEARNINGS markdown into structured
    rules for the Step-4 consumer (P10 load / §24 apply). Exam-agnostic: rules
    are indexed by their defect_code (the universal-taxonomy routing key), never by
    an exam section name. Accepts audit-learning (AL-*) and explain-guardrail (EX-*)
    rule headers. `path` may be a filepath or the raw markdown text. Returns
    {'rules':[{code,title,defect_code,occurrences,first_seen,pattern,prevention,
    verification,superseded}], 'by_defect':{defect_code:[rule_code,...]}}."""
    import os
    text = open(path, encoding='utf-8').read() if (isinstance(path, str)
            and len(path) < 4096 and os.path.exists(path)) else path
    rules = []
    for b in re.split(r'\n(?=## (?:AL|EX)[-\w]* \u2014 )', text):
        m = re.match(r'## ((?:AL|EX)[-\w]*) \u2014 (.+)', b)
        if not m:
            continue
        code, title = m.group(1), m.group(2).strip()
        def field(name):
            fm = re.search(rf'\*\*{name}:\*\*\s*(.+?)(?=\n\*\*|\n## |\n# |\Z)', b, re.S)
            return fm.group(1).strip() if fm else None
        rules.append({
            'code': code, 'title': title,
            'defect_code': field('Defect code'),
            'occurrences': field(r'Occurrences[^:]*'),
            'first_seen': field('First seen'),
            'pattern': field('Pattern'),
            'prevention': field('Prevention rule') or field('Rule'),
            'verification': field('Verification'),
            'triggers': field('Triggers'),          # v2.8 — curated-neighbour trigger terms
            'superseded': bool(field('Supersedes')),
        })
    by_defect = {}
    for r in rules:
        dc = r['defect_code']
        if dc:
            by_defect.setdefault(dc, []).append(r['code'])
    return {'rules': rules, 'by_defect': by_defect}

def self_test_audit():
    """Round-trip gate for parse_solution_blocks: write -> read -> rebuild -> assert
    the read-back block reproduces the source, across mcq/msq/nat, numeric/alpha/
    roman labels, OMML fractions. Run with --self-test-audit."""
    import tempfile, os
    # v2.8 — reader round-trips carry no authoring metadata by design; the
    # fixtures below are legacy-shaped and run with the authoring gates off.
    EngineConfig.DEFAULT_PROVENANCE_GATES = False
    res = []
    def chk(name, cond): res.append((name, bool(cond)))
    def roundtrip(cfg, blocks, nq, tag):
        d = tempfile.mkdtemp(); src = os.path.join(d, 's.docx'); out = os.path.join(d, 'o.docx')
        doc = Document()
        for q in range(1, nq + 1):
            doc.add_paragraph(f'Q.{q} stem')
            for o in range(1, (cfg.expected_options(q) or 0) + 1):
                doc.add_paragraph(f'{cfg.option_label(o)}. opt')
            doc.add_paragraph('')
        doc.save(src)
        build_interleaved_docx(src, blocks, out, cfg)
        got = parse_solution_blocks(out, cfg); ok = True
        for b in blocks:
            r = got.get(b.q)
            if r is None: ok = False; break
            same = (r.qtype == b.qtype
                    and (b.qtype == 'nat' or r.ca_set() == b.ca_set())
                    and [s.strip() for s in r.axiom] == [s.strip() for s in b.axiom]
                    and [s.strip() for s in r.deduction] == [s.strip() for s in b.deduction]
                    and set(r.why_wrong) == set(b.why_wrong)
                    and set(map(str, r.common_pitfalls)) == set(map(str, b.common_pitfalls)))
            if b.qtype == 'nat':
                if b.ca_range is not None:
                    # point value is not recoverable from a rendered range
                    # line by design — only the band round-trips
                    same = same and r.ca_range == b.ca_range
                else:
                    same = same and str(r.ca) == str(b.ca) and r.ca_range == b.ca_range
            r.validate(); ok = ok and same
        out2 = os.path.join(d, 'o2.docx')
        build_interleaved_docx(src, list(got.values()), out2, cfg)
        ve, _ = verify_explanations(out2, list(got.values()), cfg)
        chk(tag, ok and ve)
    cfg4 = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4)
    roundtrip(cfg4, [ExplanationBlock(q=1, ca=3, cfg=cfg4,
        axiom=['The mean equals the sum over the count here.'],
        deduction=['Divide 235/5 to get the value here.', 'The result is Option 3 here.'],
        why_wrong={1: ['Option 1 uses the right value for the wrong quantity.'],
                   2: ['Option 2 subtracts the term that must be added.'],
                   4: ['Option 4 satisfies one condition and misses the other.']})], 1, 'RT-MCQ-FRAC')
    cfgA = EngineConfig(r'^Q\.?\s*(\d+)', r'^([A-D])[.\)]', 4, label_scheme='alpha_upper')
    roundtrip(cfgA, [ExplanationBlock(q=1, ca={1, 3}, cfg=cfgA,
        axiom=['A statement is valid when both halves hold here.'],
        deduction=['Statement one holds and three holds here.', 'So Option A and Option C are correct here.'],
        why_wrong={2: ['Option B passes one test but fails another test.'],
                   4: ['Option D fails the parity test outright.']})], 1, 'RT-MSQ-ALPHA')
    cfgN = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', None, options_by_q={1: 0})
    roundtrip(cfgN, [ExplanationBlock(q=1, ca='47', cfg=cfgN, ca_range=(46.5, 47.5),
        axiom=['The rate equals distance over time here.'],
        deduction=['Compute the quotient step here.', 'The value is 47 here.'],
        common_pitfalls={'235': ['Forgetting to divide leaves 235 unchanged.'],
                         '9': ['Dividing by the wrong count gives 9 instead.']})], 1, 'RT-NAT-RANGE-PITFALL')
    cfgR = EngineConfig(r'^Q\.?\s*(\d+)', r'^(i{1,3}|iv)[.\)]', 4, label_scheme='roman_lower')
    roundtrip(cfgR, [ExplanationBlock(q=1, ca=2, cfg=cfgR,
        axiom=['A tangent meets the radius at a right angle here.'],
        deduction=['Apply the perpendicular property here.', 'This gives Option ii here.'],
        why_wrong={1: ['Option i confuses radius with diameter outright.'],
                   3: ['Option iii mis-reads the chord length.'],
                   4: ['Option iv drops a factor of two entirely.']})], 1, 'RT-MCQ-ROMAN')
    # regression lock for the m:num/m:den itertext fix: a REAL fraction must pass
    import tempfile as _t, os as _o
    _c = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', 4)
    _b = ExplanationBlock(q=1, ca=1, cfg=_c,
        axiom=['The mean equals the sum over the count here.'],
        deduction=['Divide 235/5 to reach the value here.', 'The result is Option 1 here.'],
        why_wrong={2: ['Option 2 uses the right value for the wrong quantity.'],
                   3: ['Option 3 subtracts the term that must be added.'],
                   4: ['Option 4 satisfies one condition and misses the other.']})
    _d = _t.mkdtemp(); _s = _o.path.join(_d, 's.docx'); _oo = _o.path.join(_d, 'o.docx')
    _dd = Document(); _dd.add_paragraph('Q.1 stem')
    for _x in range(1, 5): _dd.add_paragraph(f'{_x}. opt')
    _dd.add_paragraph(''); _dd.save(_s)
    build_interleaved_docx(_s, [_b], _oo, _c)
    _ok, _ = verify_explanations(_oo, [_b], _c); chk('RT-FRAC-VERIFY', _ok)

    # ══ v2.0 (GAP-2026-08-07-EXPLAIN-OMML) locks ══════════════════════════════
    # DRIFT LOCK — t3_mathcomp.py body must be byte-identical to the S3-5b embed
    # in Framework_PYQPrepare.md (single source, two consumers).
    import os as _os, re as _re2
    _here = _os.path.dirname(_os.path.abspath(__file__))
    _spec = open(_os.path.join(_here, 'Framework_PYQPrepare.md'), encoding='utf-8').read()
    _mm = _re2.search(r'### S3-5b — Tier-3 structured math compiler \(v2\.0\)\n.*?```python\n(.*?)```', _spec, _re2.S)
    _mod = open(_os.path.join(_here, 't3_mathcomp.py'), encoding='utf-8').read()
    _body = _mod.split('# ── end shim; everything below is the verbatim S3-5b embed ──\n', 1)[1]
    chk('T3-DRIFT-LOCK', _mm is not None and _body == _mm.group(1))

    # BUILDER STRICTNESS — the v1 bare-text fraction shape must now FAIL verify.
    _bad = Document()
    _bp = _bad.add_paragraph()
    _bp._p.append(parse_xml(f'<m:oMath xmlns:m="{M}"><m:f><m:num>5</m:num><m:den>4</m:den></m:f></m:oMath>'))
    _bo = _os.path.join(_t.gettempdir(), 't3_badfrac.docx'); _bad.save(_bo)
    import zipfile as _zf
    _root = parse_xml(_zf.ZipFile(_bo).read('word/document.xml'))
    _bare = 0
    for _f in _root.iter(qn('m:f')):
        for _part in (_f.find(qn('m:num')), _f.find(qn('m:den'))):
            if _part is not None and ((_part.text or '').strip() or not len(_part)):
                _bare += 1
    chk('T3-BARE-TEXT-DETECTED', _bare == 2)
    # …and the FIXED builder emits run-wrapped, digit-preserving XML.
    _gd = Document(); _gp = _gd.add_paragraph()
    add_math(_gp, omath(frac('5', '4')))
    _gf = next(_gp._p.iter(qn('m:f')))
    chk('T3-BUILDER-RUNWRAPPED',
        _gf.find(qn('m:num')).find(qn('m:r')) is not None
        and ''.join(_gf.find(qn('m:num')).itertext()) == '5'
        and not (_gf.find(qn('m:num')).text or '').strip())
    chk('T3-ESCAPE', '<m:t' in frac('a<b', '&') and '&lt;b' in frac('a<b', '&') and '&amp;' in frac('a<b', '&'))

    # TIER-3 REGION RENDER — fraction, subscript, bar, vector through add_math_text.
    _rd = Document(); _rp = _rd.add_paragraph()
    add_math_text(_rp, 'Here ⟦MATH:V_{B} = \\frac{ℏ}{2mΔx}⟧ and ⟦MATH:\\bar{A}\\vec{E}⟧ close the loop.')
    chk('T3-REGION-RENDER',
        sum(1 for _ in _rp._p.iter(qn('m:oMath'))) == 2
        and next(_rp._p.iter(qn('m:sSub'))) is not None
        and next(_rp._p.iter(qn('m:acc'))) is not None
        and '⟦' not in _rp.text)

    # DIALECT GUARD — every evasion spelling now raises with the region remedy.
    for _bad_s, _lbl in (('The value is ℏ ÷ 2 here.', 'DIV'),
                         ('Take x^2 as the term here.', 'CARET'),
                         ('Set V_B to zero here.', 'SUB'),
                         ('Use √(2I) as the field here.', 'SQRT'),
                         ('Factor A\u0305B out here.', 'COMBINING')):
        try:
            guard_sentence(_bad_s); chk('T3-GUARD-' + _lbl, False)
        except ValueError as _e:
            chk('T3-GUARD-' + _lbl, '⟦MATH:' in str(_e))
    # …and the SAME constructs inside a region are legal.
    chk('T3-GUARD-REGION-OK',
        guard_sentence('Then ⟦MATH:V_{B} = \\sqrt{x^{2}}⟧ holds here.') is not None)

    # GRACEFUL DEGRADATION — a bad region never raises at render; it ships as
    # plain text and the ledger records it for the verifier's verbatim quote.
    _n0 = len(T3_STATS['failed'])
    _dd = Document(); _dp = _dd.add_paragraph()
    add_math_text(_dp, 'Bad ⟦MATH:\\frobnicate{x}⟧ region here.')
    chk('T3-DEGRADE-NO-RAISE', len(T3_STATS['failed']) == _n0 + 1
        and '\\frobnicate{x}' in _dp.text and '⟦' not in _dp.text)
    T3_STATS['failed'].pop()      # keep the shared ledger clean for later locks

    # SOURCE MATH HEALTH — a gap-bearing input paper is named in plain words.
    _sh = Document()
    _sh.add_paragraph('Q.1  The voltage  , in Volts, equals what value?')
    for _i in (1, 2, 3, 4): _sh.add_paragraph(f'{_i}. {_i}')
    _so = _os.path.join(_t.gettempdir(), 't3_gapsrc.docx'); _sh.save(_so)
    _w = source_math_health(_so, _c)
    chk('T3-SOURCE-HEALTH', any('missing math symbol' in x for x in _w)
        and any('re-run Step 1 v2.0' in x for x in _w))
    chk('T3-SOURCE-HEALTH-CLEAN', source_math_health(_oo, _c) == [])

    # (audit-hardening) degraded body inside a larger paragraph must be MASKED
    # from the dialect scan by substring containment, not exact match.
    T3_STATS['failed'].append(('V_{B} = \\frac{1}{2}', 'probe'))
    _fb = {b for b, _ in T3_STATS['failed']}
    chk('T3-MASK-CONTAINMENT',
        any(fb and fb in 'prefix V_{B} = \\frac{1}{2} suffix.' for fb in _fb))
    T3_STATS['failed'].pop()
    # (audit-hardening) a HEALTHY paragraph whose math follows as trailing OMML
    # ("V = <oMath>") must NOT be flagged as a gap.
    _hd = Document()
    _hd.add_paragraph('Q.1  What does the following equal?')
    _hp = _hd.add_paragraph('The value V = ')
    add_math(_hp, omath(frac('1', '2')))
    for _i in (1, 2, 3, 4): _hd.add_paragraph(f'{_i}. {_i}')
    _ho = _os.path.join(_t.gettempdir(), 't3_healthy.docx'); _hd.save(_ho)
    chk('T3-SOURCE-HEALTH-NO-FALSEPOS',
        not any('missing math symbol' in x for x in source_math_health(_ho, _c)))
    # regression: a NAT with a FRACTION answer AND a fraction pitfall value must
    # round-trip AND pass verify_explanations (both were OMML-blind before the fix).
    _cn = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', None, options_by_q={1: 0})
    # v1.16: ca (the GRADED value) must be portal-charset-pure — a fraction
    # answer is pre-converted to its decimal equivalent before reaching the
    # engine, same as any other NAT value. common_pitfalls KEYS are prose
    # labels only (never graded, never sent to the portal), so fraction
    # notation there is unaffected and still exercises the OMML-fraction
    # round-trip path this fixture exists to protect.
    _bn = ExplanationBlock(q=1, ca='0.75', cfg=_cn,
        axiom=['The probability equals favourable over total here.'],
        deduction=['Count 3 favourable of 4 here.', 'The value is 0.75 here.'],
        common_pitfalls={'1/2': ['An even split gives 1/2 here instead.'],
                         '4/3': ['Inverting gives 4/3 here by mistake.']})
    _dn = _t.mkdtemp(); _sn = _o.path.join(_dn, 's.docx'); _on = _o.path.join(_dn, 'o.docx')
    _dc = Document(); _dc.add_paragraph('Q.1 stem'); _dc.add_paragraph(''); _dc.save(_sn)
    build_interleaved_docx(_sn, [_bn], _on, _cn)
    _ven, _ = verify_explanations(_on, [_bn], _cn)
    _gn = parse_solution_blocks(_on, _cn)[1]
    chk('RT-NAT-FRAC', _ven and set(map(str, _gn.common_pitfalls)) == {'1/2', '4/3'}
        and str(_gn.ca) == '0.75')
    # parse_learnings round-trip: a synthetic AUDIT_LEARNINGS parses to structured
    # rules indexed by defect_code (the Step-4 P10 consumer contract).
    _lm = (
        "# X_EXPLAIN_AUDIT_LEARNINGS_v1.md\n\n## AL-1 \u2014 TRACE MUST PRODUCE THE WRONG VALUE\n\n"
        "**Defect code:** WHY-WRONG-DIAG\n**Occurrences in M1:** 6 of 20\n"
        "**Pattern:** the named error does not reproduce the option value.\n"
        "**Prevention rule:** execute the claimed error and confirm it yields the option.\n"
        "**Verification:** each wrong value traces to a named mistake.\n\n"
        "## AL-2 \u2014 SHORTCUT MUST BE FASTER\n\n**Defect code:** FAKE-SPEED-HACK\n"
        "**Pattern:** the speed hack restates the deduction.\n"
        "**Prevention rule:** require a distinct route with fewer steps, else omit.\n\n"
        "## EX-9 \u2014 VIEW EVERY IMAGE\n\n**Rule:** derive figural answers from the viewed image.\n")
    _pl = parse_learnings(_lm)
    _codes = {r['code'] for r in _pl['rules']}
    chk('LEARN-PARSE-COUNT', _codes == {'AL-1', 'AL-2', 'EX-9'})
    chk('LEARN-PARSE-INDEX', _pl['by_defect'].get('FAKE-SPEED-HACK') == ['AL-2']
        and _pl['by_defect'].get('WHY-WRONG-DIAG') == ['AL-1'])
    _al1 = next(r for r in _pl['rules'] if r['code'] == 'AL-1')
    chk('LEARN-PARSE-FIELDS', bool(_al1['prevention']) and '6 of 20' in (_al1['occurrences'] or '')
        and next(r for r in _pl['rules'] if r['code'] == 'EX-9')['prevention'] is not None)
    _lm2 = ("## AL-5 \u2014 NEW RULE\n\n**Defect code:** CA-WRONG-FACTUAL\n"
            "**Supersedes:** AL-1\n**Prevention rule:** web-verify every fact first.\n")
    _pl2 = parse_learnings(_lm2)
    chk('LEARN-SUPERSEDE', _pl2['rules'][0]['superseded'] is True
        and _pl2['by_defect'].get('CA-WRONG-FACTUAL') == ['AL-5'])
    passed = sum(1 for _, ok in res if ok); total = len(res)
    for nm, ok in res:
        if not ok: print(f'  AUDIT-FAIL: {nm}')
    EngineConfig.DEFAULT_PROVENANCE_GATES = True
    print(f'AUDIT-SELF-TEST: {passed}/{total} PASS')
    return passed == total

if __name__ == '__main__':
    if '--self-test-audit' in sys.argv:
        sys.exit(0 if self_test_audit() else 1)
    if '--self-test' in sys.argv:
        sys.exit(0 if self_test() else 1)
    if len(sys.argv) >= 3 and sys.argv[1] == '--scan-risk':
        # v2.8 (e): python3 explain_engine.py --scan-risk <Explanation.docx> [opt_count]
        import json as _j
        _n = int(sys.argv[3]) if len(sys.argv) > 3 else 4
        _cfg = EngineConfig(r'^Q\.?\s*(\d+)', r'^([1-9])[.\)]', _n, provenance_gates=False)
        _res = scan_risk_markers(sys.argv[2], _cfg)
        print(_j.dumps({'questions_flagged': len(_res), 'detail': _res}, indent=1,
                       ensure_ascii=False))
        sys.exit(0)
    print('explain_engine.py — universal exam-agnostic. '
          '--self-test (core) or --self-test-audit (Step-5 reader round-trip); '
          '--scan-risk <docx> [options] marks an existing Explanation for regeneration.')
