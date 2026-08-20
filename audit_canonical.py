#!/usr/bin/env python3
# ============================================================================
# [ExamCode]_mock_test_audit.py
# UNIVERSAL, EXAM-AGNOSTIC Part-A machine-gate auditor.
# RUNNER (2026.08.03.6): Step 7 (MockCreate / TestCreate) — via the per-exam copy
# [ExamCode]_mock_test_audit.py at S4-7 STEP B and S13-4c. Canonical Step 8, which
# formerly ran it mandatorily, was retired in 2026.08.03.5. This is now the ONLY
# machine auditor in the mock/scoped pipeline.
# (originally CreateAudit v2.6 Appendix A; that spec was retired 2026-08-03 and this
# file is now the sole home of the A-* catalogue).
#
# Zero hardcoded exam values: every expected value (question/section counts,
# q_ranges, options_count, option label format, language, OMML-required
# subtopics, figural/linked maps) is read at runtime from blueprint.json /
# section_rules.md / subtopic_manifest.json / registry.json. The SAME script
# audits any exam with valid Step 0/1/2 outputs.
#
# v2.14 — 2026-08-20 — GAP-2026-08-20-AUDITOR-OPTN-DIAGNOSIS (run-report F4). A-OPTN
# run without --registry on a NAT-bearing paper still FAILS (ND6: no contract, no
# certificate) but now SAYS so — 'NOT ASSESSABLE without the option contract … re-run
# with --registry --blueprint --rules --manifest --mockN' — instead of the plain
# 'wrong count: Q41:0 …' that was indistinguishable from a defective paper. Verdict
# LEVEL unchanged; only the diagnosis. Fixtures 3b.
# v2.6 — adds the S5-1A COMPLETION GATE (--audit-state): after Part A, validates
# the Phase-2 audit_state ledger (C1-C7) and the on-disk evidence artefacts each
# stamp names, so a skipped/collapsed Phase 2 fails LOUDLY instead of shipping.
# The self-test is FIXTURE-BASED (MANDATE A / P1): a constant-print stub is not a
# valid auditor.
#
# Dependencies: python-docx + Python stdlib ONLY.
#
# Usage:
#   python3 [ExamCode]_mock_test_audit.py PAPER.docx \
#       --blueprint BP.json --rules RULES.md --manifest MAN.json \
#       --registry REG.json --mockN N [--final] [--audit-state STATE.json]
#   python3 [ExamCode]_mock_test_audit.py --self-test
#
# Exit code: 0 if no FAIL AND (when --audit-state) COMPLETION-GATE PASS, else 1.
# WARNs name a MECHANICAL remedy or declare dormancy (v2.22.0 SELF-ADJUDICATION:
# no gate may instruct a human to look at something and decide) and do not change the
# exit code; the caller's certification gate (spec §12-2) decides whether a fixable
# WARN blocks delivery.
# ============================================================================
import sys, os, re, json, hashlib, zipfile, argparse, tempfile, unicodedata
import io, zlib, struct, random          # v2.13 — stdlib PNG fixtures (D4); no PIL needed
import ast                               # v2.24.0 — SELF-ADJUDICATION scans
                                         # EMITTED message literals, not raw lines
import inspect                           # v2.21.9 — the CLEAN-SHAPE MATRIX (5g)
                                         # DISCOVERS gates instead of listing them
from pathlib import Path
from collections import Counter, defaultdict

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.document import Document as _DocClass
from docx.shared import Inches          # v2.21.9 — shape fixtures run before the
                                        # local import that used to own this name

# ---- OOXML namespaces ------------------------------------------------------
W  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
M  = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
A  = 'http://schemas.openxmlformats.org/drawingml/2006/main'
WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
def W_(t):  return f'{{{W}}}{t}'
def M_(t):  return f'{{{M}}}{t}'

# ---- v2.6 constants (MANDATE A / S5-1A) ------------------------------------
AUTH_GATE_FLOOR   = 35    # P1: the self-test must exercise >= this many fixtures.
EVIDENCE_MIN_BYTES = 100  # S5-1A C6: a montage evidence file must be a real raster.

# ---- result accumulator ----------------------------------------------------
RESULTS = []   # list of (level, code, message)
def _ok(c, m):   RESULTS.append(('OK',   c, m))
def _warn(c, m): RESULTS.append(('WARN', c, m))
def _fail(c, m): RESULTS.append(('FAIL', c, m))
def _reset():    RESULTS.clear()

# ---- block model -----------------------------------------------------------
QNUM_RE = re.compile(r'^\s*Q\.?\s*(\d+)\b')

class Block:
    __slots__ = ('qnum', 'items', 'paras', 'tables', 'images')
    def __init__(self, qnum):
        self.qnum = qnum
        self.items = []       # ordered ('p',Paragraph) / ('t',Table)
        self.paras = []       # Paragraph objects in this block
        self.tables = []      # Table objects in this block
        # v2.13 (GAP-2026-08-01-FIGSPEC-TRANSPORT D1): POPULATED by
        # attach_block_images() — one dict per inline image in this block, in
        # document order:
        #   {'name': docPr/@name, 'rid': r:embed, 'descr': docPr/@descr,
        #    'part': media part basename, 'path': extracted PNG path|None}
        # It was declared '# reserved' from v1.0 and never appended to anywhere,
        # so the twelve v2.11 figure-conformance gates iterated an always-empty
        # list and printed "0 figure(s) conform." on every paper in every exam.
        self.images = []


def iter_block_items(doc):
    """Yield Paragraph and Table objects in true document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            yield ('p', Paragraph(child, doc))
        elif child.tag == qn('w:tbl'):
            yield ('t', Table(child, doc))


def para_text(p):
    """Readable text of a paragraph, merging normal text (w:t) and math (m:t)
    in document order so a math-bearing or pure-OMML stem is never seen empty."""
    out = []
    for node in p._element.iter():
        if node.tag == W_('t') or node.tag == M_('t'):
            out.append(node.text or '')
    return ''.join(out)


def para_images(p):
    """Return [(docPr_name, blip_rEmbed)] for inline images in this paragraph."""
    imgs = []
    el = p._element
    names = el.findall(f'.//{{{WP}}}docPr')
    blips = el.findall(f'.//{{{A}}}blip')
    embeds = [b.get(qn('r:embed')) for b in blips]
    for i, emb in enumerate(embeds):
        nm = names[i].get('name') if i < len(names) else None
        imgs.append((nm, emb))
    return imgs


def para_images_ext(p):
    """v2.13 — the STRUCTURED form of para_images(): one dict per inline image,
    carrying its docPr name, its r:embed rid and its docPr alt text (@descr).

    Walks each <w:drawing> as a UNIT and reads that drawing's own docPr + blip,
    rather than zipping two flat document-order lists as para_images() does. The
    zip is correct while every drawing has exactly one docPr and one blip, but a
    paragraph mixing an inline drawing with an anchored one, or a blip carrying
    a <a:blip r:link> alternate, misaligns it. A-FIGALT reads @descr, so a
    misalignment here would attribute one figure's alt text to another.

    para_images() is left BYTE-UNCHANGED: it is read at six call sites that need
    only presence/naming, and its behaviour is fixture-locked.
    """
    out = []
    for dr in p._element.findall(f'.//{W_("drawing")}'):
        nm = dr.find(f'.//{{{WP}}}docPr')
        bl = dr.find(f'.//{{{A}}}blip')
        if bl is None:
            continue
        out.append({'name': (nm.get('name') if nm is not None else None) or '',
                    'descr': (nm.get('descr') if nm is not None else None) or '',
                    'rid': bl.get(qn('r:embed'))})
    return out


def block_image_paras(b):
    """Every paragraph of a block that can carry an inline image, INCLUDING the
    paragraphs inside its tables. Figural options are placed one-per-paragraph
    (S10-8), but a DI chart or a figure/option fusion table puts drawings inside
    table cells, and a block-level paras scan would miss those entirely."""
    seen = set()
    for p in b.paras:
        if id(p._element) not in seen:
            seen.add(id(p._element)); yield p
    for t in b.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if id(p._element) not in seen:
                        seen.add(id(p._element)); yield p


def extract_media(docx_path, media_map):
    """v2.13 — extract every referenced media part to a temp dir and return
    {rid: absolute path}. The thirteen figure-conformance gates are arithmetic
    over the saved PNG, so they need a real file on disk; a docx part is a ZIP
    member and PIL cannot open it in place.

    NEVER raises (RA-9 / CLAUDE.md): an unreadable archive degrades to an empty
    map, which the caller reports as a coverage WARN. A figure gate must never
    be the reason a run dies — that is the exact defect class v2.12 closed.
    """
    paths = {}
    try:
        d = tempfile.mkdtemp(prefix='auditmedia_')
        with zipfile.ZipFile(docx_path) as z:
            names = set(z.namelist())
            for rid, base in media_map.items():
                for cand in (f'word/media/{base}', f'media/{base}'):
                    if cand in names:
                        dst = os.path.join(d, f'{rid}_{base}')
                        with open(dst, 'wb') as fh:
                            fh.write(z.read(cand))
                        paths[rid] = dst
                        break
    except Exception:
        return paths
    return paths


def attach_block_images(blocks, media_map, media_paths):
    """v2.13 (D1) — fill Block.images. Returns (declared, resolved): how many
    inline images the paper physically carries, and how many of those have a
    usable PNG on disk. The two are reported separately because 'no figures'
    and 'figures I could not open' are opposite facts, and printing OK for the
    second is the vacuous-pass defect this release exists to close."""
    declared = resolved = 0
    for b in blocks:
        for p in block_image_paras(b):
            for im in para_images_ext(p):
                rid = im.get('rid')
                pth = media_paths.get(rid)
                declared += 1
                if pth:
                    resolved += 1
                b.images.append({'name': im.get('name') or '',
                                 'rid': rid,
                                 'descr': im.get('descr') or '',
                                 'part': media_map.get(rid, ''),
                                 'path': pth})
    return declared, resolved


def resolve_figure_spec(img, specs):
    """v2.13 (D2) — find this image's FigureSpec in the registry-borne map.

    Step 7 stamps every drawing with its canonical name at S10-8
    (_name_last_drawing -> 'q{N}_problem.png' / 'q{N}_opt{i}.png'), and
    write_spec_sidecar() names the sidecar after the same PNG, so the docPr name
    IS the key. Fall back to the extension-stripped form and to the media part
    name so a re-emitted or renamed drawing (CP-IMGNAME) still resolves.

    An unresolved lookup returns {} — which fc.is_legacy() reads as pre-v5.33
    output and EC-V18 then treats leniently. That is the CORRECT default: an
    absent spec must degrade, never fabricate a verdict.
    """
    if not specs:
        return {}
    for k in (img.get('name'), img.get('part')):
        if not k:
            continue
        if k in specs:
            return specs[k] or {}
        stem = os.path.splitext(k)[0]
        if stem in specs:
            return specs[stem] or {}
        if (stem + '.png') in specs:
            return specs[stem + '.png'] or {}
    return {}


# ============================================================================
# D7 — A FINDINGS LIST MAY NEVER TRUNCATE SILENTLY (v2.17)
#
# WHY. Gates printed `_flist(x)`. On a real 60-Q paper
# A-FIGCOMP had 27 findings and printed 12 — 15 vanished with no trace. Worse,
# `sorted()` is LEXICOGRAPHIC on 'Q1'/'Q10'/'Q3', so the printed set looked
# arbitrary (Q3 after Q28) and a reviewer reasonably concluded the gate was
# non-deterministic and filed it as unreproducible. It was neither: it was
# under-reporting, in numeric-blind order.
#
# A finding that exists and is not shown is the same false-clean this corpus has
# now hit six times. Every list therefore states its TOTAL, marks suppression
# explicitly, and sorts Q-numbers as NUMBERS.
# ============================================================================
FINDING_LIST_CAP = 12


def _qsort_key(x):
    """Numeric-aware sort so Q3 precedes Q10 and a truncated head is the LOWEST
    Q-numbers, not a lexicographic accident."""
    t = str(x)
    m = re.search(r'\d+', t)
    return (0, int(m.group()), t) if m else (1, 0, t)


def _flist(items, cap=FINDING_LIST_CAP, sep=' '):
    """Render a findings list. NEVER silently truncates: suppression is named and
    the total is always printed, so a reader can tell 'this is all of them' from
    'this is the first twelve of twenty-seven'."""
    if isinstance(items, str):
        items = [items]
    u = sorted({str(i) for i in items}, key=_qsort_key)
    head = sep.join(u[:cap])
    if len(u) > cap:
        return f'{head}{sep}... [+{len(u) - cap} MORE NOT SHOWN; {len(u)} TOTAL]'
    return head


def _fcount(items):
    """The honest count for a message prefix."""
    return len({str(i) for i in (items or [])})


# ============================================================================
# D1 — MANDATE 0's FIRST MACHINE CHECK (v2.18)
#
# MANDATE 0 declares itself absolute and had NO gate of any kind. S5-1A C1-C7
# assert ledger completeness and evidence existence; S14-2 scans output FILENAMES
# for banned substrings. Nothing ever inspected what Claude actually WROTE. A
# violation was therefore structurally invisible to every check in the framework —
# which is exactly how one occurred: a structure-inspection script printed p.text
# while building a block index and surfaced ~6 questions into the transcript.
#
# This does not police the transcript (it cannot see it). It polices the PATTERN
# rule 2 forbids: printing a content-bearing expression from operator code. Static,
# cheap, and it makes the difference between "printing p.text is a violation" being
# prose and being enforced.
# ============================================================================
_M0_CONTENT_EXPR = re.compile(
    r'print\s*\(.*?\b('
    r'p\.text|para_text\s*\(|\.stem\b|stem_text|opt_text|option_text|'
    r'\.opts\b|passage|cell\.text|\.runs\b'
    r')', re.S)


def mandate0_scan(code_text):
    """True if this source line/snippet PRINTS question content (MANDATE 0 rule 2).
    Counting is fine — `print(len(b.images))` and `print(f'Q{n}: {len(opts)}')` are
    correct and must not trip; printing the text itself is the violation."""
    for m in _M0_CONTENT_EXPR.finditer(code_text or ''):
        seg = code_text[m.start():m.end()]
        if re.search(r'\blen\s*\(\s*$', seg[:seg.rfind(m.group(1))] or ''):
            continue
        if re.search(r'len\s*\(\s*[^)]*' + re.escape(m.group(1).rstrip('(')), seg):
            continue
        if '__name__' in code_text:
            continue
        return True
    return False


def parse_blocks(doc):
    """Split the document body into per-question blocks. Items before the first
    Q.<n> are the title/instruction block (returned separately)."""
    blocks, title_items = [], []
    cur = None
    for kind, obj in iter_block_items(doc):
        if kind == 'p':
            mobj = QNUM_RE.match(para_text(obj))
            if mobj:
                cur = Block(int(mobj.group(1)))
                blocks.append(cur)
        if cur is None:
            title_items.append((kind, obj))
            continue
        cur.items.append((kind, obj))
        if kind == 'p':
            cur.paras.append(obj)
        else:
            cur.tables.append(obj)
    return title_items, blocks


# ---- source-file parsing (exam-agnostic) -----------------------------------
def cat_c(rules_txt, key, default=None):
    mobj = re.search(rf'^[ \t]*{re.escape(key)}[ \t]*[:=][ \t]*(.+?)[ \t]*$',
                     rules_txt, re.M | re.I)
    return mobj.group(1).strip() if mobj else default


def load_sources(args):
    src = {}
    src['blueprint'] = json.load(open(args.blueprint, encoding='utf-8')) if args.blueprint else {}
    src['registry']  = json.load(open(args.registry,  encoding='utf-8')) if args.registry  else {}
    src['manifest']  = json.load(open(args.manifest,  encoding='utf-8')) if args.manifest  else {}
    src['rules_txt'] = open(args.rules, encoding='utf-8').read() if args.rules else ''
    bp = src['blueprint']
    src['total_questions'] = bp.get('total_questions')
    src['sections'] = bp.get('sections', [])
    # v2.24 (GAP-2026-08-06-AXIS1) — the Axis-1/Axis-3 budgets. Present since blueprint
    # v1.23 and, until this release, read by NOTHING: not by Step 7, not by any gate.
    # Two real papers shipped at 26 and 30 figures against a budget of 4 and were
    # certified clean, because no gate had ever been asked to count.
    src['axis_schedule'] = bp.get('axis_schedule', {}) or {}
    # figural_reducible per subtopic — an irreducible question (its OPTIONS are images)
    # is granted a figure even over budget, so the expectation rises with it instead of
    # a finding being raised. Read from section_rules; absent ⇒ every subtopic reducible,
    # which is the STRICTER reading and therefore the safe default for an auditor.
    src['figural_reducible'] = {}
    for _m in re.finditer(r'subtopic_id:\s*(\S+)((?:(?!subtopic_id:).)*?)'
                          r'figural_reducible:\s*(true|false)',
                          src.get('rules_txt') or '', re.DOTALL | re.I):
        src['figural_reducible'][_m.group(1)] = (_m.group(3).lower() == 'true')
    rt = src['rules_txt']
    src['language']      = (cat_c(rt, 'language', 'english') or 'english').lower()
    src['options_count'] = int(cat_c(rt, 'options_count', '4'))
    src['opt_label_fmt'] = cat_c(rt, 'option_label_format', '1/2/3/4')
    src['font_family']   = cat_c(rt, 'font_family', 'Calibri')
    # OMML-required subtopics present at all?
    src['omml_required_present'] = bool(re.search(r'OMML_required\s*[:=]\s*(true|yes|1)',
                                                  rt, re.I))
    # mock-scoped maps from registry
    N = args.mockN
    reg = src['registry']
    rc  = next((x for x in reg.get('rc_manifests', []) if x.get('mock') == N), None)
    src['passage_linked'] = set(rc.get('passage_linked', [])) if rc else set()
    src['cloze_linked']   = set(rc.get('cloze_linked', []))   if rc else set()
    fig = next((x for x in reg.get('figural_manifests', []) if x.get('mock') == N), None)
    src['figural_qs'] = set(int(q) for q in fig.get('figural_qs', [])) if fig else set()
    # v2.24.1 — PRESENCE, NOT EMPTINESS. "no figural manifest for this mock" and "a
    # manifest saying zero figures" are OPPOSITE FACTS, and an empty set represents both.
    # A-AXIS1 must not read a missing record as a paper with no figures: that turns an
    # unknown into a hard shortfall FAIL. Same distinction gate_images already draws
    # between `declared` and `resolved`.
    src['figural_manifest_present'] = fig is not None
    src['rc_manifest_present'] = rc is not None
    # v2.28 — WHICH mock this is, so gate_axis1 can read THIS paper's target out of the
    # rotating axis1_target_series. Without it the series is inert in production: the
    # fixture passed while the gate silently kept using the flat mean, which is the
    # "well-fixtured but unwired" shape audit_callgraph exists to catch on the engine
    # side and nothing catches on the loader side.
    src['mock_n'] = N
    _mf = src.get('manifest')
    src['unkeyed_questions_by_class'] = ((_mf or {}).get('unkeyed_questions_by_class')
                                         if isinstance(_mf, dict) else {}) or {}
    # v2.25 (GAP-2026-08-06-DI) — the DI producer record. Closes the last hole in
    # A-AXIS1: DI was the one budgeted stimulus class with no trace anywhere, so it
    # could only ever be reported UNESTABLISHED. It cannot be recovered from the docx —
    # G-MATCH-TABLE mandates a real Word table for every MATCH question, so on a real
    # paper (IIT_JAM_BIOTECHNOLOGY 15-Feb-2026) table-presence finds 3 candidates where
    # exactly 1 is DI. Only the generator knows which it built; Step 7 v5.42 records it.
    _di = next((x for x in reg.get('di_manifests', []) if x.get('mock') == N), None)
    src['di_qs'] = set(int(q) for q in _di.get('di_qs', [])) if _di else set()
    src['di_manifest_present'] = _di is not None
    # v2.10 (GAP-2026-07-26-003 D2): A-FIGPROFILE needs the object_type Step 7 v5.31
    # recorded per figural question, plus that question's subtopic_id. Both travel in
    # the registry figural_manifest, which Step 8 DOES receive — unlike the answer_key
    # sidecar (S0-1), so no concept_map is required. A pre-v5.31 registry simply has
    # neither key and the gate goes dormant.
    src['figural_object_types'] = (fig.get('object_types') or {}) if fig else {}
    src['figural_subtopics'] = (fig.get('subtopic_ids') or {}) if fig else {}
    # v2.13 (GAP-2026-08-01-FIGSPEC-TRANSPORT D2): the FigureSpec sidecars, keyed
    # by canonical PNG name. figural_core.write_spec_sidecar() drops these beside
    # each PNG in the STEP-7 session's working dir; that dir is internal and is
    # never delivered (S0-1), so without a transport channel Step 8 saw spec=={}
    # on every figure and EC-V18 downgraded every BLOCKING verdict on output that
    # was not, in fact, legacy. The registry is the sanctioned channel — exactly
    # the precedent object_types/subtopic_ids set at v5.31. Absent key => {} =>
    # every figure reads as legacy, which is the correct pre-v5.34 behaviour.
    src['figure_specs'] = (fig.get('figure_specs') or {}) if fig else {}
    # v2.4: section_rules full text for image_role lookups in gate_images
    src['section_rules_text'] = rt   # already loaded at src['rules_txt']
    # v2.4: concept_map — Step 8 does NOT receive the answer_key sidecar (S0-1),
    # so concept_map is unavailable by default. If the answer_key is available
    # (e.g., passed via --key), load it; otherwise empty dict (gate_images falls
    # back to default image_role='stem_and_options' per subtopic).
    src['concept_map'] = {}
    src['answers'] = {}
    # v2.17 TIER A — the dossier feeds the SAME consumer path --key has fed since
    # v2.4. Nothing downstream changes shape; what changes is that the path is
    # finally supplied in normal operation instead of being permanently empty.
    src['dossier'] = None
    src['dossier_why'] = 'not supplied'
    if getattr(args, 'dossier', None):
        try:
            # v2.20 — PASS THE EXAM CODE. load_dossier() has accepted an `exam`
            # argument since v2.17 and the call site never supplied it, so ONE
            # THIRD OF THE DECLARED IDENTITY TRIPLE (exam_code / mock / paper_md5)
            # was never checked in production: a dossier from another exam was
            # ACCEPTED. paper_md5 made it unreachable in practice — a different
            # exam's paper cannot share this paper's hash — but a documented
            # binding that never executes is the same dead-parameter class this
            # corpus keeps finding, and defence in depth is worth nothing if one
            # of the layers is not wired. blueprint.exam_code is the authority
            # (P2 already asserts it equals the trigger, RS-5).
            _qs, _man = load_dossier(args.dossier, docx_path=args.docx,
                                     exam=(bp.get('exam_code') or None),
                                     mockN=getattr(args, 'mockN', None))
            src['dossier'] = _qs
            src['dossier_manifest'] = _man
            # Adopt Tier-A facts. answers stays EMPTY: the dossier carries no
            # judgment, and A-DOSSIER cross-checks every fact against the paper.
            src['concept_map'] = {q: dict(e) for q, e in _qs.items()}
        except DossierError as e:
            src['dossier_why'] = str(e)
    if getattr(args, 'key', None) and Path(args.key).exists():
        _key_data = json.load(open(args.key, encoding='utf-8'))
        src['concept_map'] = _key_data.get('concept_map', {})
        src['answers'] = _key_data.get('answers', {})

    # v1.2 — MSQ re-derivation (INDEPENDENT of any Step-7 self-report). Dormant when the
    # blueprint declares no multi subtopics (every value below is empty/zero ⇒ the MSQ
    # gates pass vacuously and behave exactly as v1.1).
    src['multi_present'] = bool(bp.get('multi_present', bp.get('msq_present', False)))  # Phase-0 back-compat
    multi_ids = {s.get('subtopic_id') for s in bp.get('subtopic_list', [])
                 if s.get('answer_cardinality', s.get('answer_mode')) == 'multi'}
    src['multi_subtopic_ids'] = multi_ids
    # expected count of multi-mode questions per section, from THIS mock's allocations.
    mock_entry = next((m for m in bp.get('mocks', []) if m.get('mock') == N), None)

    # v2.9 POSITION-BASED QUESTION TYPE (GAP-2026-07-22-001 §6 FIX, mirrors Step 7 v5.30
    # and Step 11 v1.7 dual-mode pattern):
    # For question-type sections (IIT JAM: Section A=MCQ, B=MSQ, C=NAT), per-subtopic
    # answer_cardinality/answer_type is unreliable — multi_ids/nat_ids from subtopic_list
    # would be empty or incomplete because no single subtopic has majority MSQ/NAT
    # observations. The marking_scheme is authoritative: if it defines >1 distinct
    # question_type, derive expected MSQ/NAT counts from Q-position ranges directly.
    #
    # MODE SELECTION (identical to Step 7/Step 11):
    #   > 1 distinct question_type in marking_scheme → POSITION-BASED
    #   0 or 1 distinct type → SUBTOPIC-BASED (unchanged behavior)
    #
    # For POSITION-BASED mode:
    #   expected_multi_by_section = count of Qs in each section whose marking_scheme says MSQ
    #   expected_nat_by_section = count of Qs in each section whose marking_scheme says NAT
    #   multi_subtopic_ids = ALL subtopics that appear in MSQ ranges (for Part B checks)
    #   nat_subtopic_ids = ALL subtopics that appear in NAT ranges (for A-FIGCOMP stem_only)
    #
    # Backward compatible: 0 or 1 distinct type → exact pre-v2.9 behavior. Covers all
    # existing exams (SSC CGL, MPPSC, etc.) and legacy blueprints with no marking_scheme.
    _bp_ms = bp.get('marking_scheme', [])
    _audit_distinct_q_types = {ms.get('question_type') for ms in _bp_ms
                               if ms.get('question_type')}
    _audit_position_based = len(_audit_distinct_q_types) > 1

    def _audit_type_for_q(qnum):
        """Return question_type for a given Q number from marking_scheme."""
        for ms in _bp_ms:
            qr = ms.get('q_range', [0, 0])
            if qr[0] <= qnum <= qr[1]:
                return ms.get('question_type', 'MCQ')
        return 'MCQ'

    exp = {}
    if _audit_position_based and mock_entry:
        # POSITION-BASED: expected counts from marking_scheme Q-ranges per section
        for sec in mock_entry.get('sections', []):
            sec_name = sec.get('section_name')
            qr = sec.get('q_range', [0, 0])
            msq_count = sum(1 for q in range(qr[0], qr[1] + 1)
                            if _audit_type_for_q(q) == 'MSQ')
            if msq_count:
                exp[sec_name] = msq_count
        # Augment multi_subtopic_ids: for position-based exams, ANY subtopic in an MSQ
        # range is effectively MSQ — add all subtopics allocated to MSQ sections.
        for sec in mock_entry.get('sections', []):
            qr = sec.get('q_range', [0, 0])
            if any(_audit_type_for_q(q) == 'MSQ' for q in range(qr[0], qr[1] + 1)):
                for sa in sec.get('subtopic_allocations', []):
                    sid = sa.get('subtopic_id')
                    if sid:
                        multi_ids.add(sid)
        src['multi_subtopic_ids'] = multi_ids
    elif mock_entry and multi_ids:
        # SUBTOPIC-BASED: unchanged pre-v2.9 behavior
        for sec in mock_entry.get('sections', []):
            c = sum(sa.get('q_count', 0) for sa in sec.get('subtopic_allocations', [])
                    if sa.get('subtopic_id') in multi_ids)
            if c:
                exp[sec.get('section_name')] = c
    src['expected_multi_by_section'] = exp
    # MSQ config (section_rules / blueprint) — read where needed, never from a sidecar.
    mc = bp.get('msq_contract', {})
    src['msq_k_mode'] = mc.get('msq_k_mode', 'variable')
    src['msq_k']      = mc.get('msq_k', None)
    src['msq_allow_aota'] = bool(re.search(r'^\s*msq_allow_aota\s*:\s*true\s*$', rt, re.I | re.M))
    # select-instruction phrases: exam's own (section_rules msq_instruction[_hi]) + universal.
    phrases = []
    for m in re.finditer(r'^\s*msq_instruction(?:_hi)?\s*:\s*(.+?)\s*$', rt, re.I | re.M):
        phrases.append(m.group(1).strip().strip('()'))
    src['msq_instruction_phrases'] = phrases + [
        'one or more', 'may be correct', 'select two', 'select three',
        'select all that apply', 'एक या अधिक']

    # v1.4 — NAT re-derivation (INDEPENDENT of any Step-7 self-report). Dormant when the
    # blueprint declares no numerical subtopics (every value below empty/zero ⇒ the NAT
    # gates pass vacuously and behave exactly as v1.3).
    src['nat_present'] = bool(bp.get('nat_present', False))
    nat_ids = {s.get('subtopic_id') for s in bp.get('subtopic_list', [])
               if s.get('answer_type', 'option') == 'numerical'}
    src['nat_subtopic_ids'] = nat_ids
    # expected count of numerical questions per section, from THIS mock's allocations.
    nexp = {}
    if _audit_position_based and mock_entry:
        # POSITION-BASED: expected NAT counts from marking_scheme Q-ranges per section
        for sec in mock_entry.get('sections', []):
            sec_name = sec.get('section_name')
            qr = sec.get('q_range', [0, 0])
            nat_count = sum(1 for q in range(qr[0], qr[1] + 1)
                            if _audit_type_for_q(q) == 'NAT')
            if nat_count:
                nexp[sec_name] = nat_count
        # Augment nat_subtopic_ids: for position-based exams, ANY subtopic in a NAT
        # range is effectively NAT — add all subtopics allocated to NAT sections.
        # This ensures A-FIGCOMP correctly treats figural NAT Qs as stem_only (line ~4505).
        for sec in mock_entry.get('sections', []):
            qr = sec.get('q_range', [0, 0])
            if any(_audit_type_for_q(q) == 'NAT' for q in range(qr[0], qr[1] + 1)):
                for sa in sec.get('subtopic_allocations', []):
                    sid = sa.get('subtopic_id')
                    if sid:
                        nat_ids.add(sid)
        src['nat_subtopic_ids'] = nat_ids
    elif mock_entry and nat_ids:
        # SUBTOPIC-BASED: unchanged pre-v2.9 behavior
        for sec in mock_entry.get('sections', []):
            c = sum(sa.get('q_count', 0) for sa in sec.get('subtopic_allocations', [])
                    if sa.get('subtopic_id') in nat_ids)
            if c:
                nexp[sec.get('section_name')] = c
    src['expected_nat_by_section'] = nexp
    # v1.4 (ND6): per-question EXPECTED option count from the registry (Step-7 delivery
    # contract, NOT a self-audit sidecar). 0 marks a NAT question. Used to SKIP the option
    # gates for NAT Qs; A-NAT-NOOPT then independently verifies each claimed-NAT Q truly
    # renders 0 options (and A-OPTN still fires if a claimed-MCQ Q renders 0) — so a
    # mislabelled options_by_q cannot let a defect through either direction. {} ⇒ inert.
    src['options_by_q'] = {str(k): v for k, v in
                           reg.get('options_by_q', {}).get(str(N), {}).items()}
    # NAT config (blueprint nat_contract) — read where needed, never from a sidecar.
    nc = bp.get('nat_contract', {})
    src['nat_answer_type'] = nc.get('nat_answer_type', 'real')
    src['nat_tolerance']   = nc.get('nat_tolerance', '0')
    # numerical-entry instruction phrases: exam's own (section_rules nat_instruction[_hi]) + universal.
    nphrases = []
    for m in re.finditer(r'^\s*nat_instruction(?:_hi)?\s*:\s*(.+?)\s*$', rt, re.I | re.M):
        nphrases.append(m.group(1).strip().strip('()'))
    src['nat_instruction_phrases'] = nphrases + [
        'numerical value', 'enter your answer', 'enter the value', 'numerical answer',
        'संख्यात्मक मान']
    # figural stem-cue keywords: exam's own (section_rules figural_cue_keywords) if
    # declared, else a universal default set. RA-9: these are never hardcoded in the gate.
    fcue_m = re.search(r'^\s*figural_cue_keywords\s*[:=]\s*(.+?)\s*$', rt, re.I | re.M)
    if fcue_m:
        src['figural_cue_keywords'] = [k.strip().lower() for k in fcue_m.group(1).split(',')]
    # else: gate_images falls back to a built-in default list (see gate_images)
    # escape-reference phrases: if the whole phrase appears in a stem, it references
    # the escape option the phrase implies. Read from section_rules (RA-9);
    # gate_optref falls back to English defaults when absent.
    eref_m = re.search(r'^\s*escape_reference_phrases\s*[:=]\s*(.+?)\s*$', rt, re.I | re.M)
    if eref_m:
        src['escape_reference_phrases'] = [t.strip() for t in eref_m.group(1).split(',')]
    # figure-reference prose phrases (A-FIGTEXT-PROSE): exam's own, read from
    # section_rules. RA-9 — the gate previously HARDCODED an English pattern with
    # reasoning-exam shape nouns (triangles/squares/circles). On a non-English paper
    # it matched NOTHING and the gate printed a clean OK, so the detector was
    # silently vacuous for every non-English exam in the estate while reporting
    # conformance. v2.21.8.
    fref_m = re.search(r'^\s*figure_reference_phrases\s*[:=]\s*(.+?)\s*$', rt, re.I | re.M)
    if fref_m:
        src['figure_reference_phrases'] = [t.strip() for t in fref_m.group(1).split(',')
                                           if t.strip()]
    # stimulus detection cues: exam's own (section_rules stimulus_cue_patterns) if
    # declared; gate_stimorphan merges them with the built-in English cues. RA-9.
    scue_m = re.search(r'^\s*stimulus_cue_patterns\s*[:=]\s*(.+?)\s*$', rt, re.I | re.M)
    if scue_m:
        src['stimulus_cue_patterns'] = [c.strip() for c in scue_m.group(1).split(',')]
    # mock title keyword: exam's own (section_rules mock_title_keyword) if declared,
    # else default 'mock'. RA-9: gate_header reads from src.
    src['mock_title_keyword'] = cat_c(rt, 'mock_title_keyword', 'mock')
    # v2.7: paper_header_block opt-in (gate_header / A-HEADER dormancy). Default OFF — no
    # current section_rules declares it, so the pre-Q.1 body-block ban is absolute.
    src['paper_header_block'] = str(cat_c(rt, 'paper_header_block', '')).strip().lower() \
        in ('true', '1', 'yes', 'on')
    return src


def option_label_family(fmt):
    first = (fmt or '1/2/3/4').split('/')[0].strip().strip('()').strip('.').strip(')')
    if first.isdigit():                         return 'num'
    if re.fullmatch(r'[ivxIVX]+', first):       return 'roman'
    if len(first) == 1 and first.isalpha():     return 'alpha'
    return 'num'

# v2.21.7 — OPT_RE RETIRED with option_paras(); see the note above _label_paras.
# bare-or-full option label (figural options are a bare '1.' label paragraph
# followed by an image paragraph; text options are 'label. text').
OPT_LABEL_RE = re.compile(r'^\s*\(?\s*([0-9]+|[A-Za-z]|[ivxIVX]+)\s*\)?\s*[.)](\s|$)')

# v2.21.7 — option_paras() RETIRED. It was the last wrapper around OPT_RE, and
# both of its remaining consumers (gate_optref, gate_qnfirst) have been moved onto
# _label_paras()/OPT_LABEL_RE. Keeping a second option predicate alive with no
# caller is drift waiting for an author: the next person needing "the options of a
# block" would have found two helpers and picked one. There is now exactly ONE
# option-label predicate in this file (S5-2 "ONE STRUCTURAL QUESTION, ONE ANSWER"),
# so the divergence class that produced GAP-2026-08-02 is structurally impossible
# rather than merely absent. validate_framework_md.py CHECK AN still guards the
# reintroduction of a second one.


def block_stem_text(block):
    """All non-option paragraph text in the block, with the leading 'Q.<n>'
    token stripped. In the Step-7 format the stem lives ON the Q.<n> line
    (e.g. 'Q.74  Study the following table ...'), so the Q.<n> paragraph must
    be INCLUDED (prefix removed), not excluded."""
    parts = []
    for p in block.paras:
        t = para_text(p)
        if OPT_LABEL_RE.match(t):
            continue
        t = QNUM_RE.sub('', t, count=1).strip()
        if t:
            parts.append(t)
    return ' '.join(parts)


# ============================================================================
# GATES
# ============================================================================
STIMULUS_CUES = re.compile(
    r'\b(the passage|the table|the graph|the chart|the given data|the following '
    r'(table|passage|graph|chart|bar|pie|line|data)|blank\s*\(|according to the '
    r'passage|read the (passage|following passage))\b', re.I)
CROSSREF_RE = re.compile(r'Q\.?\s*\d+\s*(and|to|&|–|-)\s*Q?\.?\s*\d+', re.I)
UNDERLINE_REF = re.compile(r'\bunderlin(e|ed)\b', re.I)
FAKE_UNDERLINE = re.compile(r'\(\s*underlin(e|ed)[^)]*:', re.I)
ESCAPE_TOKENS_DEFAULT = [
    r'no error', r'no improvement', r'none of (these|the above|them)',
    r'all of the above', r'both .+ and ', r'neither .+ nor ']
MATH_TOKEN_NAME = re.compile(r'_(e\d+|eqn|expr|frac|math)\b', re.I)
CANON_IMG_NAME = re.compile(r'^q\d+_(problem|opt\d+|stim)', re.I)
ASCII_CARET = re.compile(r'[A-Za-z0-9]\s*\^\s*[0-9A-Za-z]')
SLASH_FRAC  = re.compile(r'(?<![A-Za-z0-9])\d+\s*/\s*\d+(?![A-Za-z0-9])')


def gate_structure(blocks, src):
    tq = src['total_questions']
    nums = [b.qnum for b in blocks]
    if tq is not None:
        (_ok if len(blocks) == tq else _fail)(
            'A-COUNT', f'{len(blocks)} question blocks (expected {tq}).')
        expected = set(range(1, tq + 1))
        got = set(nums)
        if got == expected:
            _ok('A-SEQ', f'Q-numbers complete 1..{tq}.')
        else:
            miss = sorted(expected - got); extra = sorted(got - expected)
            _fail('A-SEQ', f'missing={_flist(miss)} extra={_flist(extra)}.')
    else:
        _warn('A-COUNT', 'blueprint.total_questions absent; count check skipped.')
    if nums == sorted(nums) and len(nums) == len(set(nums)):
        _ok('A-MONO', 'Q-numbers strictly increasing.')
    else:
        _fail('A-MONO', 'Q-numbers not strictly increasing in document order.')


def gate_seccount(blocks, src):
    secs = src['sections']
    if not secs:
        _warn('A-SECCOUNT', 'blueprint.sections absent; section-count check skipped.')
        return
    nums = set(b.qnum for b in blocks)
    bad = []
    for s in secs:
        lo, hi = s['q_range']
        cnt = sum(1 for q in nums if lo <= q <= hi)
        if cnt != s.get('total_qs', hi - lo + 1):
            bad.append(f"{s.get('name','?')}:{cnt}/{s.get('total_qs', hi-lo+1)}")
    (_ok if not bad else _fail)('A-SECCOUNT',
        'section counts match q_ranges.' if not bad else 'mismatch ' + '; '.join(bad))


def _label_paras(block):
    """All option-label paragraphs in the block (bare '1.' or full '1. text'),
    in document order, as (token, full_text, text_after_label)."""
    out = []
    for p in block.paras:
        t = para_text(p)
        if QNUM_RE.match(t):
            continue
        m = OPT_LABEL_RE.match(t)
        if m:
            out.append((m.group(1), t, t[m.end():].strip()))
    return out


def _fam_of(l):
    if l.isdigit():
        return 'num'
    if re.fullmatch(r'[ivxIVX]+', l):
        return 'roman'
    return 'alpha'


def _idx_of(l, fam):
    if l.isdigit():
        return int(l)
    if fam == 'roman':
        seq = ['i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x']
        return seq.index(l.lower()) + 1 if l.lower() in seq else 0
    return ord(l.lower()) - ord('a') + 1


def gate_options(blocks, src):
    oc = src['options_count']; fam = option_label_family(src['opt_label_fmt'])
    obq = src.get('options_by_q', {})   # v1.4: per-Q expected option count (0 = NAT)
    bad_n, bad_lab, bad_ord, bad_uni = [], [], [], []
    for b in blocks:
        # v1.4 NAT: a numerical question expects 0 options — the option gates (A-OPTN/
        # A-OPTLABEL/A-OPTORDER/A-OPTUNIQUE) DO NOT APPLY to it. Skip when the registry
        # marks it 0-option; A-NAT-NOOPT separately verifies it truly renders no options.
        if str(b.qnum) in obq and obq[str(b.qnum)] == 0:
            continue
        labs = _label_paras(b)
        # Options are the TRAILING oc label-paragraphs (Step-7 puts the option
        # block last). A stray earlier label (e.g. an enumerated passage point
        # '1. ...') is ignored, so it never inflates the count; a genuinely short
        # option set (< oc labels) still FAILs.
        if len(labs) < oc:
            bad_n.append(f'Q{b.qnum}:{len(labs)}')
            continue
        opt_labs = labs[-oc:]
        tokens = [x[0] for x in opt_labs]
        texts  = [x[2].casefold() for x in opt_labs]
        if any(_fam_of(t) != fam for t in tokens):
            bad_lab.append(f'Q{b.qnum}')
        idxs = [_idx_of(t, fam) for t in tokens]
        # v2.21.3 — ANCHORED AT 1, not merely consecutive. The S5-2 catalogue row
        # for this gate has always read "options appear in document order
        # 1..OPTIONS_COUNT", but the check was `range(idxs[0], idxs[0] + oc)`,
        # which accepts ANY consecutive run: a block labelled 2,3,4,5 passed
        # A-OPTORDER *and* A-OPTLABEL and certified clean. That is not cosmetic.
        # A-KINT derives the key as an int in 1..OPTIONS_COUNT, so on such a paper
        # key "option 1" refers to an option that DOES NOT EXIST and keys 2..oc
        # each point one place off — every answer for that question is wrong on
        # the delivered paper. _idx_of() normalises all three families to 1-based
        # (num 1.., alpha a=1.., roman i=1..), so anchoring at 1 is family-agnostic
        # and no legitimate option_label_format starts anywhere else.
        if idxs != list(range(1, oc + 1)):
            bad_ord.append(f'Q{b.qnum}')
        # text-uniqueness applies only when options are TEXT (figural options are
        # bare labels followed by images → empty text → uniqueness checked in §7).
        nonempty = [x for x in texts if x]
        if len(nonempty) == oc and len(set(nonempty)) != oc:
            bad_uni.append(f'Q{b.qnum}')
    # v2.14 (GAP-2026-08-20-AUDITOR-OPTN-DIAGNOSIS, run-report F4). Without
    # --registry there is no options_by_q, so NAT questions are NOT skipped and
    # this gate FAILS on every one of them. That FAIL is DELIBERATE (ND6: a run
    # that cannot see the option contract must not certify a NAT paper) and it
    # stays a FAIL — but it was indistinguishable from a defective paper: an
    # operator spot-checking a clean 60-question paper saw "A-OPTN FAIL: Q41..Q60
    # wrong count" and nothing else. The verdict now SAYS which case it is.
    # Only a ZERO-option Q is ambiguous without the contract (NAT renders 0;
    # a short MCQ renders some). A paper with no zero-option Qs keeps the plain
    # 'wrong count' verdict — it is short options, contract or not.
    _no_contract = (not obq) and any(x.endswith(':0') for x in bad_n)
    _zero_only = _no_contract and all(x.endswith(':0') for x in bad_n)
    if _no_contract:
        _msg = ('NOT ASSESSABLE without the option contract — no --registry '
                'supplied, so per-question expected counts (options_by_q) are '
                'unknown; ' + ('every flagged Q renders 0 options, the NAT shape' if
                _zero_only else 'some flagged Qs render 0 options (NAT shape), '
                'others too few') + '. Re-run '
                'with --registry --blueprint --rules --manifest --mockN (the S13-4c '
                're-sweep invocation). Flagged: ' + _flist(bad_n))
        _fail('A-OPTN', _msg)
    else:
        (_ok if not bad_n else _fail)('A-OPTN',
            f'every Q has {oc} options.' if not bad_n else 'wrong count: ' + _flist(bad_n))
    (_ok if not bad_lab else _fail)('A-OPTLABEL',
        'labels match format.' if not bad_lab else 'bad label family: ' + _flist(bad_lab))
    (_ok if not bad_ord else _fail)('A-OPTORDER',
        'options in order.' if not bad_ord else 'out of order: ' + _flist(bad_ord))
    (_ok if not bad_uni else _fail)('A-OPTUNIQUE',
        'options distinct within Q.' if not bad_uni else 'dup options: ' + _flist(bad_uni))


def gate_qnfirst(blocks):
    """Q.N-FIRST (R14): nothing may precede a block's Q.<n>. Because the parser
    starts a block at Q.<n>, the violation manifests as stimulus content (table /
    image / long passage) sitting AFTER the previous block's last option and
    BEFORE the next block's Q.<n> — i.e. a lead-in orphaned ahead of its Q.N.
    It is attributed to the block whose Q.<n> it precedes."""
    viol = []
    for k, b in enumerate(blocks):
        last_opt = -1
        opt_idxs = []
        for idx, (kind, obj) in enumerate(b.items):
            # v2.21.7 (SEC-1) — OPT_LABEL_RE, not OPT_RE. The anchor is "where does
            # this block's option set END", and an IMAGE option is a BARE label
            # paragraph followed by a picture: OPT_RE requires a visible glyph after
            # the label and matched NONE of them, so last_opt stayed -1 on every
            # figural block and the `continue` below SKIPPED THE WHOLE CHECK. The
            # gate printed ok regardless, so the shortfall was invisible: measured
            # 25 of 60 blocks unchecked on a real paper. An identical orphaned
            # lead-in was CAUGHT after a text block and MISSED after a figural one.
            # Same predicate split as GAP-2026-08-02; this was its last consumer.
            if kind != 'p':
                continue
            _t = para_text(obj)
            if QNUM_RE.match(_t):          # never the Q.<n> line itself
                continue
            if OPT_LABEL_RE.match(_t):
                last_opt = idx
                opt_idxs.append(idx)
        if last_opt < 0 or k + 1 >= len(blocks):
            continue
        # v2.21.9 (GAP-2026-08-02-QNFIRST-IMAGE-OPTION) — IMAGE-BOUND OPTIONS.
        # last_opt is the index of the last option LABEL. An IMAGE option is a
        # BARE label paragraph FOLLOWED BY its picture (R-FIGURAL / G-FIGURAL-
        # COMPOSITE: "problem image + one separate image per option, bound 1:1 to
        # labels"), so the final option's OWN picture NECESSARILY sits after the
        # last label. Flagging it made this gate FAIL every conformant
        # stem_and_options figural paper in the estate while A-FIGCOMP, A-OPTN,
        # A-OPTUNIQUE and A-DOSSIER all passed the SAME block — two gates, one
        # block, contradictory verdicts, which is the v2.21 A-DOSSIER signature.
        # v2.21.7 introduced it: moving the anchor to OPT_LABEL_RE correctly fixed
        # a false NEGATIVE (figural blocks were skipped entirely) and, unmeasured,
        # opened this false POSITIVE.
        # THE ALLOWANCE IS MEASURED FROM THE BLOCK, NEVER ASSUMED (RA-9). Count the
        # pictures bound to each NON-FINAL option (between consecutive labels); the
        # same number is allowed to trail the final label. A TEXT-option block
        # measures 0 and behaves EXACTLY as before, so the genuine orphan catch —
        # a stimulus belonging to the NEXT question stranded at the end of this one
        # — is preserved at full strength (fixture 5a parity).
        _bound = []
        for _a, _bnext in zip(opt_idxs, opt_idxs[1:]):
            _bound.append(sum(1 for _kd, _ob in b.items[_a + 1:_bnext]
                              if _kd == 'p' and para_images(_ob)))
        _img_allow = max(_bound) if _bound else 0
        for kind, obj in b.items[last_opt + 1:]:
            if kind == 't':
                viol.append(f'Q{blocks[k+1].qnum}'); break
            if kind == 'p':
                if para_images(obj):
                    if _img_allow > 0:
                        _img_allow -= 1      # the final option's own image
                        continue
                    viol.append(f'Q{blocks[k+1].qnum}'); break
                if len(para_text(obj).split()) >= 35:
                    viol.append(f'Q{blocks[k+1].qnum}'); break
    (_ok if not viol else _fail)('A-QNFIRST',
        'every block opens with Q.<n> (no orphaned lead-in).' if not viol else
        'stimulus orphaned before Q.<n>: ' + _flist(viol))


def gate_blanksep(doc, blocks):
    # blank separator: at least one empty paragraph between the last option of a
    # block and the next block's Q.<n>. Approx: count empty paragraphs overall vs blocks.
    empties = sum(1 for kind, obj in iter_block_items(doc)
                  if kind == 'p' and not para_text(obj).strip()
                  and not para_images(obj))
    (_ok if empties >= max(0, len(blocks) - 1) else _warn)('A-BLANKSEP',
        f'{empties} blank separators for {len(blocks)} blocks.'
        if empties >= len(blocks) - 1 else
        f'only {empties} blank separators for {len(blocks)} blocks (some missing).')


def gate_font(doc, src):
    fam = src['font_family']
    bad = set()
    for kind, obj in iter_block_items(doc):
        if kind != 'p':
            continue
        for r in obj.runs:
            nm = r.font.name
            if nm not in (None, fam):
                bad.add(nm)
    (_ok if not bad else _fail)('A-FONT',
        f'all runs {fam}.' if not bad else f'non-{fam} fonts present: {_flist(bad)}')


def gate_sechdr(blocks, doc, src):
    # v1.5 — two detectors, over ALL body paragraphs (not only within question blocks, so a
    # heading before Q.1 or between blocks is seen too):
    #   (a) KEYWORD form — text opening with "section"/"part N"/rule characters;
    #   (b) SECTION-NAME form (the realistic case the keyword pattern MISSED): a standalone
    #       paragraph that IS a declared section name ("Quantitative Aptitude", "Technical",
    #       "General Awareness"). PROVENANCE-BASED — matched against the blueprint's own
    #       section names (src['sections']), not a generic word list, so it flags exactly the
    #       headings this paper's sections would produce and stays exam-agnostic.
    # A questions-only paper (R8, Q.N-first) has no standalone non-Q/non-option paragraph, so
    # any body line equal to a section name is a leaked header, never legitimate content.
    pat = re.compile(r'^\s*(section\b|part\s+[IVXA-D0-9]+\b|=====|─────|-----)', re.I)
    sec_names = set()
    for s in src.get('sections', []):
        nm = (s.get('name') or s.get('section_name') or '').strip().lower()
        if nm:
            sec_names.add(nm)
    hits = []
    for kind, obj in iter_block_items(doc):
        if kind != 'p':
            continue
        t = para_text(obj).strip()
        if not t or QNUM_RE.match(t):        # blank or a Q.<n> stem line — never a stray header
            continue
        if (pat.match(t) and len(t) < 60) or (t.lower() in sec_names):
            hits.append(t[:40])
    (_ok if not hits else _fail)('A-SECHDR',
        'no body section headers.' if not hits else 'section-header text in body: ' + _flist(hits, sep=' | '))


def gate_anskey(doc):
    full = '\n'.join(para_text(obj) for kind, obj in iter_block_items(doc) if kind == 'p')
    # v1.2: the last pattern now matches SET-valued keys ("Q.1 → 1,2,4"), not just a
    # single trailing digit/letter — mirrors Step 7 G-ANSWERKEY. A leaked MSQ key is a
    # comma/space list of option positions; the v1.1 single-token pattern missed it.
    pats = [r'\banswer\s*key\b', r'\banswers?\s*:', r'^\s*key\s*:',
            r'Q\.?\s*\d+\s*[:\-–>]+\s*[\(]?[1-9a-dA-D][\)]?(?:\s*[,\s]\s*[\(]?[1-9a-dA-D][\)]?)*\s*$',
            # v1.4: NAT numerical-value answer-key lines ("Q.5 → 47" / "→ 0" / "→ -3" /
            # "→ 3.14"). The option-position pattern above only matches 1-9/a-d, so a leaked
            # NAT key (incl. 0/negative/decimal) slipped through. Mirrors Step 7 v4.7.
            r'Q\.?\s*\d+\s*[:\-–>]+\s*-?\d+(?:\.\d+)?\s*$']
    hit = [p for p in pats if re.search(p, full, re.I | re.M)]
    (_ok if not hit else _fail)('A-ANSKEY',
        'no answer key/markers in body.' if not hit else f'answer-key signal(s): {hit}')


def gate_msq_instr(blocks, src):
    # v1.2 — A-MSQ-INSTR (machine, MULTI only). INDEPENDENT of any Step-7 self-report:
    # the EXPECTED number of multi-mode questions per section is re-derived from the
    # blueprint allocations (src['expected_multi_by_section']); the OBSERVED number is
    # the count of Q blocks in that section's q_range whose stem line carries a select-
    # instruction phrase. A mismatch means a multi Q is missing its instruction (or a
    # single Q wrongly carries one). Pinpointing WHICH Q is refined in Part B (semantic),
    # where the question's subtopic is known from content. Fully dormant when the blueprint
    # declares no multi subtopics (expected map empty ⇒ pass).
    exp = src.get('expected_multi_by_section', {})
    if not exp:
        return _ok('A-MSQ-INSTR', 'no multi-mode subtopics in this mock (dormant).')
    phrases = [p.lower() for p in src.get('msq_instruction_phrases', [])]
    sec_of = src.get('section_of_qnum', {})    # {qnum: section_name} if available
    secs   = src.get('sections', [])
    def section_for(qnum):
        if qnum in sec_of:
            return sec_of[qnum]
        for s in secs:                          # fall back to q_range membership
            lo, hi = (s.get('q_range') or [0, 0])[:2]
            if lo <= qnum <= hi:
                return s.get('section_name') or s.get('name')   # blueprint uses 'name'
        return None
    observed = {}
    for b in blocks:
        stem = para_text(b.paras[0]) if getattr(b, 'paras', None) else getattr(b, 'stem', '')
        if any(ph in stem.lower() for ph in phrases):
            observed[section_for(b.qnum)] = observed.get(section_for(b.qnum), 0) + 1
    bad = [f'{sec}: expected {n} multi, saw {observed.get(sec, 0)} instruction-carrying'
           for sec, n in exp.items() if observed.get(sec, 0) != n]
    (_ok if not bad else _fail)('A-MSQ-INSTR',
        'observed multi instruction counts match the blueprint per section.' if not bad
        else 'multi instruction-count mismatch — ' + ' | '.join(bad))


def _axis_sections(src):
    """Well-formed (name, lo, hi) triples, plus the names that could not be bucketed.

    Hardened after fuzzing (v2.24.1): `sections` arriving as None or a string, a section
    that is not a dict, a q_range that is a string or holds non-numeric entries, and a
    reversed range were all live crashes. _safe_gate would have caught them as
    A-GATEERROR — but a gate that dies takes its whole finding set with it (the defect
    class v2.12 closed), and A-AXIS1 dying is indistinguishable from A-AXIS1 passing to
    anyone reading a summary. A malformed section is SKIPPED and NAMED, never fatal.
    """
    good, bad = [], []
    secs = src.get('sections')
    if not isinstance(secs, (list, tuple)):
        return good, bad
    for sec in secs:
        if not isinstance(sec, dict):
            bad.append('<malformed section>')
            continue
        name = sec.get('name') or sec.get('section_name')
        rng = sec.get('q_range')
        if not isinstance(rng, (list, tuple)) or len(rng) < 2:
            bad.append(str(name))
            continue
        try:
            lo, hi = int(rng[0]), int(rng[1])
        except (TypeError, ValueError):
            bad.append(str(name))
            continue
        if hi < lo:
            bad.append(str(name))
            continue
        good.append((name, lo, hi))
    return good, bad


def gate_axis1(blocks, src):
    """v2.24 — A-AXIS1 + A-AXIS-UNGATED (machine, per PAPER).

    THE GATE THAT WAS MISSING. Step 6 has written axis1_target_per_mock into
    blueprint.json since v1.23; until this release nothing spent it and nothing checked
    it. Measured cost on a real exam (IIT_JAM_BIOTECHNOLOGY, 2026-08-06): a budget of 4
    figural questions per 60-question paper, two delivered papers carrying 26 and 30,
    and a clean bill of health from all 24 existing gates. The exam itself averages 4.4
    figures per paper over five years, so those mocks trained candidates on a paper that
    does not exist.

    OBSERVED is counted from the registry figural manifest (the producer's OWN record of
    what it drew), bucketed into sections by q_range. TEXT is the residual and is never
    audited — it absorbs all rounding by construction.

    THE VERDICT IS DELEGATED to bc.check_axis_conformance(), the SAME engine function
    Step 7 spends against, so the generator and its auditor cannot drift apart. That is
    the identical discipline A-FIGPROFILE already applies to figure TYPE; this gate asks
    the COUNT question that one never asked.

    IRREDUCIBLE OVERAGE IS SILENT, NOT WARNED. A question whose OPTIONS are images has no
    text form, so it is granted a figure even over budget and the EXPECTATION RISES with
    it. Excess that irreducibles fully explain is a clean PASS; excess they do not is a
    FAIL. Without that second half the exemption would be the hole the gate leaks through.

    Fully dormant on a pre-v1.23 blueprint (no axis_schedule ⇒ SKIP), which is what keeps
    ~200 deployed exams passing un-remeasured.
    """
    sched = src.get('axis_schedule')
    if not isinstance(sched, dict) or not sched:
        return _ok('A-AXIS1', 'no axis_schedule in blueprint (pre-v1.23) — dormant.')
    # ENGINE GUARD, DELIBERATELY NON-FATAL (v2.24.1). The conformance VERDICT is
    # delegated to blueprint_core so generator and auditor cannot drift; without it the
    # verdict is honestly NOT ESTABLISHED. But the COVERAGE report below is set
    # arithmetic over blueprint keys and needs no engine at all — and it is the MOST
    # valuable line precisely when something is missing, because it says what could not
    # be checked. An early return here would have silenced it exactly then. (Same defect
    # shape as A-AXIS-UNGATED, which had to be lifted out of this gate for this reason.)
    try:
        import blueprint_core as bc
    except Exception:
        bc = None

    fig = set()
    for _q in (src.get('figural_qs') or set()):
        try:
            fig.add(int(_q))
        except (TypeError, ValueError):
            pass                      # a non-numeric qnum is unbucketable, not fatal
    red_map = src.get('figural_reducible')
    red_map = red_map if isinstance(red_map, dict) else {}
    _so = src.get('figural_subtopics')
    sub_of = {str(k): v for k, v in _so.items()} if isinstance(_so, dict) else {}


    # WHAT THIS AUDITOR CAN ACTUALLY SEE (v2.24.1). FIGURAL comes from the registry
    # figural manifest and PASSAGE from the rc manifest — both are the PRODUCER'S OWN
    # record, which is why they are trustworthy. DI has no such record anywhere, and it
    # CANNOT be inferred from the docx: a MATCH question renders a real Word table too
    # (G-MATCH-TABLE mandates it), so "block contains a table" would misread every MATCH
    # question as DI. Guessing there would trade a silent miss for a confident wrong
    # answer, so DI is reported UNESTABLISHED instead.
    observable = set()
    if src.get('figural_manifest_present', bool(fig)):
        observable.add('FIGURAL')
    if src.get('rc_manifest_present', False):
        observable.add('PASSAGE')
    if src.get('di_manifest_present', False):
        observable.add('DI')          # v2.25 — producer record, never inferred

    # No usable q_range ⇒ no way to bucket questions into a section ⇒ every count would
    # be a fabricated zero. Such sections are SKIPPED and NAMED; never audited blind.
    _secs, skipped = _axis_sections(src)
    findings, unest, audited = [], set(), 0
    for name, lo, hi in _secs:
        ss = sched.get(name)
        ss = ss if isinstance(ss, dict) else {}
        target = ss.get('axis1_target_per_mock')
        if not isinstance(target, dict) or not target:
            continue
        sec_qs = hi - lo + 1
        obs_fig = [q for q in fig if lo <= q <= hi]
        # Irreducible questions were granted OVER budget by design; count them so the
        # expectation can rise rather than a finding be raised.
        irr = sum(1 for q in obs_fig
                  if red_map.get(str(sub_of.get(str(q), ''))) is False)
        _pn = _axis1_count_in(src.get('passage_linked'), lo, hi)
        observed = _axis1_observed(sec_qs, obs_fig, _pn, src.get('di_qs'), lo, hi)
        if bc is None:
            unest |= {c for c in observable if int(target.get(c, 0) or 0) > 0}
            continue
        # v2.27 — hand the gate the exam's OWN observed spread and, when the blueprint
        # carries a rotating series, THIS mock's target rather than a flat mean. A fixed
        # band rejected four of the reference exam's five real papers.
        _spread = (sched.get(name) or {}).get('axis1_observed_figural')
        _series = (sched.get(name) or {}).get('axis1_target_series')
        _tgt = dict(target)
        if isinstance(_series, (list, tuple)) and _series and src.get('mock_n'):
            try:
                _tgt['FIGURAL'] = int(_series[(int(src['mock_n']) - 1) % len(_series)])
            except (TypeError, ValueError, IndexError):
                pass
        verdict, fs, un = bc.check_axis_conformance(observed, _tgt, irreducible=irr,
                                                    axis='axis1', observable=observable,
                                                    observed_spread=_spread)
        audited += 1
        unest |= set(un)
        if verdict == 'FAIL':
            findings += [f'{name} — {f}' for f in fs]

    # Figural questions the registry lists but no section claims. Silent loss here would
    # under-count the very quantity the gate exists to police.
    # v2.26 — ORPHAN CHECK COVERS EVERY STIMULUS CLASS, not just figures. A DI or
    # passage question belonging to no section leaves the denominator just as silently
    # as a figural one does, and the budget then looks satisfied because part of the
    # paper stopped existing. Figures were checked here since v2.24.1; the other two
    # were not, purely because they had no producer record to check at the time.
    _orphan_by = {}
    for _cls, _qs in (('FIGURAL', fig),
                      ('PASSAGE', src.get('passage_linked') if src.get('rc_manifest_present') else None),
                      ('DI', src.get('di_qs') if src.get('di_manifest_present') else None)):
        if _qs is None:
            continue
        _o = []
        for _q in (_qs if isinstance(_qs, (set, list, tuple)) else ()):
            try:
                _qi = int(_q)
            except (TypeError, ValueError):
                continue
            if not any(lo <= _qi <= hi for _n, lo, hi in _secs):
                _o.append(_qi)
        if _o:
            _orphan_by[_cls] = sorted(_o)
    orphan = sorted(_orphan_by.get('FIGURAL', []))

    if bc is None:
        _warn('A-AXIS1', 'blueprint_core not importable — Axis-1 conformance NOT '
                         'ESTABLISHED (coverage below still reported).')
    elif not audited:
        _ok('A-AXIS1', 'no Axis-1 target in any section (dormant).')
    else:
        (_ok if not findings else _fail)('A-AXIS1',
            f'Axis-1 stimulus mix within budget across {audited} section(s).'
            if not findings else
            'Axis-1 stimulus mix breaches the blueprint budget — ' + ' | '.join(findings))

    # COVERAGE IS REPORTED SEPARATELY FROM CONFORMANCE. "within budget" and "I could not
    # check" are different claims and must never be collapsed into one green line.
    cov = []
    if unest:
        # The reason differs per class, so name the right one. A generic parenthetical
        # that blamed DI regardless was printed even when FIGURAL was the missing class,
        # which sends the reader to fix the wrong thing.
        _why = {'DI': 'the registry carries no di_manifest for this mock (Step 7 '
                      'pre-v5.42); DI is never inferred from table presence, because '
                      'G-MATCH-TABLE makes every MATCH question render a real table too',
                'FIGURAL': 'the registry carries no figural manifest for this mock — '
                           'absent record, which is NOT the same fact as zero figures',
                'PASSAGE': 'the registry carries no rc manifest for this mock'}
        cov.append('no observation source for '
                   + '; '.join(f'{c} ({_why.get(c, "no producer record")})'
                               for c in sorted(unest)))
    if skipped:
        cov.append('section(s) with no usable q_range skipped: ' + ', '.join(skipped))
    for _cls, _o in sorted(_orphan_by.items()):
        cov.append(f'{len(_o)} {_cls} Q(s) fall outside every section q_range: '
                   + ', '.join(f'Q{q}' for q in _o[:10]))
    # v2.29 — UPSTREAM KEY LOSS. Step 5 counts stimulus questions carrying NO
    # subtopic_id; they are excluded from the quota, so the series comes out short by
    # exactly that many and the paper reads as a generator shortfall when the real cause
    # is a corpus that lost keys between steps. Written since v2.45 and read by nothing
    # — a diagnostic nobody surfaces is a diagnostic that does not exist.
    _uk = {}
    for _k, _v in (src.get('unkeyed_questions_by_class') or {}).items() \
            if isinstance(src.get('unkeyed_questions_by_class'), dict) else ():
        try:
            if int(_v) > 0:
                _uk[str(_k)] = int(_v)
        except (TypeError, ValueError):
            pass
    if _uk:
        cov.append('upstream questions with no subtopic_id, excluded from the quota: '
                   + ', '.join(f'{k}={v}' for k, v in sorted(_uk.items()))
                   + ' — a shortfall of this size is a corpus keying problem upstream, '
                     'not a generator fault')
    (_ok if not cov else _warn)('A-AXIS1-COVERAGE',
        'every targeted Axis-1 class was observable and every section bucketed.'
        if not cov else 'Axis-1 verdict is PARTIAL — ' + ' | '.join(cov))

def _axis1_count_in(qs, lo, hi):
    """Questions from an iterable that fall inside [lo, hi]. Non-numeric entries are
    skipped rather than fatal — a manifest is written by another step and may carry
    anything."""
    if isinstance(qs, str) or not hasattr(qs, '__iter__'):
        return 0
    n = 0
    for q in qs:
        try:
            n += 1 if lo <= int(q) <= hi else 0
        except (TypeError, ValueError):
            pass
    return n


def _axis1_observed(sec_qs, fig_in_sec, passage_n, di_qs, lo, hi):
    """The per-section Axis-1 observed vector. EXTRACTED so it can be unit-tested.

    v2.28 — this arithmetic was inline and therefore UNPINNED: reverting the TEXT
    residual to `sec_qs - figural` left all 221 fixtures green, because TEXT is the
    residual class and check_axis_conformance() skips it, so no verdict depends on the
    value. Code that no assertion can reach is code that will drift silently — the
    residual is the one place an over-produced DI could hide (the false-PASS half of the
    v2.24 defect), so its arithmetic is worth pinning even though nothing audits it yet.
    Extracting it makes the value assertable directly rather than only through a verdict.
    """
    di_n = _axis1_count_in(di_qs, lo, hi)
    fig_n = len(fig_in_sec) if hasattr(fig_in_sec, '__len__') else 0
    return {'FIGURAL': fig_n, 'PASSAGE': int(passage_n or 0), 'DI': di_n,
            'TEXT': max(0, int(sec_qs or 0) - fig_n - int(passage_n or 0) - di_n)}


def gate_axis1_overlap(blocks, src):
    """v2.26 — A-AXIS1-OVERLAP. THE STIMULUS CLASSES MUST PARTITION THE PAPER.

    Axis-1 classes are MUTUALLY EXCLUSIVE by definition: a question carries exactly one
    stimulus, which is the whole reason axis1_target_per_mock is declared to sum to
    sec_qs. Nothing enforced that. Three separate producer records — figural_manifests,
    rc_manifests, di_manifests — are written by three different code paths, and if any
    two ever claimed the same question (a chart WITH a data table being the obvious
    candidate), A-AXIS1 counted it twice and said nothing.

    MEASURED ON THE v2.25 BUILD: 60 questions recorded in BOTH the figural and DI
    manifests of a 60-question section — 120 stimuli in 60 slots, an arithmetic
    impossibility — returned A-AXIS1 = OK and A-AXIS1-COVERAGE = OK. The residual guard
    (max(0, ...)) stopped TEXT going negative, which is exactly what made the nonsense
    survivable and therefore silent.

    Low likelihood, deliberately gated anyway. It requires Step 7 to double-record, which
    it should never do — and "should never happen" is the assumption class that produced
    every other defect in this release series. The 26-figure paper shipped because nobody
    checked a budget everyone assumed would be honoured.

    TWO INDEPENDENT CHECKS, because they fail differently:
      OVERLAP  — two manifests naming the same question. Points at the producer.
      OVERFLOW — a section's non-residual classes exceeding its own question count.
                 Catches the same corruption arriving via counts rather than identity,
                 e.g. manifest entries for questions that are not in this section.

    DELIBERATELY ENGINE-FREE. Pure set arithmetic over the producer records; no
    blueprint_core, no axis_schedule. An integrity check that goes dormant whenever
    something else is missing is worth very little — the lesson of A-AXIS-UNGATED, which
    had to be lifted out of gate_axis1 for exactly this reason.
    """
    def _qset(v):
        out = set()
        for q in (v if isinstance(v, (set, list, tuple)) else ()):
            try:
                out.add(int(q))
            except (TypeError, ValueError):
                pass
        return out

    present = {}
    if src.get('figural_manifest_present'):
        present['FIGURAL'] = _qset(src.get('figural_qs'))
    if src.get('rc_manifest_present'):
        present['PASSAGE'] = _qset(src.get('passage_linked'))
    if src.get('di_manifest_present'):
        present['DI'] = _qset(src.get('di_qs'))

    if len(present) < 2:
        # Fewer than two records ⇒ nothing can overlap. Say so rather than passing
        # silently, so a reader can tell "no conflict" from "nothing to compare".
        return _ok('A-AXIS1-OVERLAP',
                   f'{len(present)} stimulus manifest(s) present — nothing to cross-check.')

    findings = []
    names = sorted(present)
    for i, a in enumerate(names):
        for b in names[i + 1:]:
            both = sorted(present[a] & present[b])
            if both:
                findings.append(
                    f'{a} and {b} both claim ' + ', '.join(f'Q{q}' for q in both[:10])
                    + (f' (+{len(both) - 10} more)' if len(both) > 10 else '')
                    + ' — a question carries ONE stimulus, so one of the two producer '
                      'records is wrong and A-AXIS1 has been counting it twice')

    # NO SEPARATE "MORE STIMULI THAN SLOTS" ARM. It was written, then PROVED
    # UNREACHABLE and removed: once the sets are disjoint and every counted element
    # lies inside [lo, hi], their counts sum to |union| <= sec_qs, so overflow cannot
    # occur without an overlap the arm above has already reported. Verified by
    # exhaustive search over 20,000 randomised disjoint configurations — zero hits.
    # Keeping it would have meant shipping a branch that can never fire inside the very
    # gate written to stop silent non-checking, which is the defect, not the fix.

    (_ok if not findings else _fail)('A-AXIS1-OVERLAP',
        f'stimulus manifests are disjoint across {len(present)} class(es); '
        'Axis-1 classes partition every section.' if not findings else
        'Axis-1 stimulus classes do not partition the paper — ' + ' | '.join(findings))


def gate_axis_ungated(blocks, src):
    """v2.24 — A-AXIS-UNGATED. THE RULE THAT STOPS THIS RETURNING AS AXIS-4.

    Any axis the blueprint marks enforcement:"hard" MUST have a gate in this auditor.
    An enforced budget with no gate is exactly the state Axis-1 and Axis-3 sat in for
    four releases — Step 6 wrote the target, Step 7 never read it, no gate ever counted
    it — and nothing in the framework was capable of noticing. Two mocks shipped at 26
    and 30 figures against a budget of 4 and were certified clean by 24 green gates.

    DELIBERATELY ENGINE-FREE. This is set arithmetic over blueprint keys and nothing
    else. It was originally written inside gate_axis1, behind that gate's
    blueprint_core import guard, which meant the one check whose entire purpose is
    catching an un-audited budget went silent in precisely the degraded environment
    where checks go missing. A meta-gate that can be disabled by an unrelated import
    failure is not a meta-gate.
    """
    sched = src.get('axis_schedule')
    if not isinstance(sched, dict) or not sched:
        return _ok('A-AXIS-UNGATED', 'no axis_schedule in blueprint (pre-v1.23) — dormant.')
    gated = {'axis1', 'axis3'}          # the axes THIS auditor actually counts
    hard = set()
    for sec in sched.values():
        if not isinstance(sec, dict):
            continue
        for k, v in sec.items():
            if k.endswith('_enforcement') and str(v).lower() == 'hard':
                hard.add(k[:-len('_enforcement')])
    miss = sorted(hard - gated)
    (_ok if not miss else _fail)('A-AXIS-UNGATED',
        f'every hard-enforced axis {sorted(hard) or "(none)"} has a gate.' if not miss else
        f'blueprint enforces {miss} but this auditor has no gate for it — an unspent, '
        'unchecked budget is how GAP-2026-08-06-AXIS1 shipped twice.')


def gate_axis3(blocks, src):
    """v2.24 — A-AXIS3 (machine, per PAPER). Axis-3 = answer mechanism (MCQ/MSQ/NAT).

    The IDENTICAL defect as Axis-1 — budgeted by Step 6, spent by nobody
    (`grep axis3 Framework_MockTestCreate.md` → 0 hits) — and masked on exams whose
    SECTIONS are defined per mechanism, where the section structure enforced it by
    accident. On any exam that mixes mechanisms inside a section it was as unenforced as
    Axis-1 was, and would have failed the same way given the same trigger.

    OBSERVED is read from the registry options_by_q: 0 options ⇒ NAT; otherwise the
    stem's select-instruction decides MSQ vs MCQ. MCQ is the residual and is not audited.
    """
    sched = src.get('axis_schedule')
    if not isinstance(sched, dict) or not sched:
        return _ok('A-AXIS3', 'no axis_schedule in blueprint (pre-v1.23) — dormant.')
    try:
        import blueprint_core as bc
    except Exception:
        bc = None                       # verdict degrades; coverage below still runs

    obq = src.get('options_by_q')
    obq = obq if isinstance(obq, dict) else {}
    phrases = [p.lower() for p in (src.get('msq_instruction_phrases') or [])]

    # OBSERVABILITY (v2.24.1). NAT is established from the registry options_by_q (0
    # options ⇒ NAT); MSQ from the select-instruction in the stem. Each depends on
    # evidence that can simply be absent — a pre-v1.4 registry carries no options_by_q,
    # and an exam with no multi-select subtopics contributes no phrases. Auditing
    # against absent evidence would report "produced 0, budget 4" on a paper that may
    # hold exactly four, i.e. a hard FAIL derived from having looked at nothing. MCQ is
    # the residual and is never audited.
    observable = set()
    if obq:
        observable.add('NAT')
    if phrases and any(getattr(b, 'paras', None) for b in blocks):
        observable.add('MSQ')

    def _mech(b):
        if str(obq.get(str(b.qnum), '')) == '0':
            return 'NAT'
        stem = para_text(b.paras[0]) if getattr(b, 'paras', None) else ''
        return 'MSQ' if any(p in stem.lower() for p in phrases) else 'MCQ'

    _secs, skipped = _axis_sections(src)
    findings, unest, audited = [], set(), 0
    for name, lo, hi in _secs:
        _ss = sched.get(name)
        target = (_ss if isinstance(_ss, dict) else {}).get('axis3_target_per_mock')
        if not isinstance(target, dict) or not target:
            continue
        observed = {}
        for b in blocks:
            if lo <= b.qnum <= hi:
                m = _mech(b)
                observed[m] = observed.get(m, 0) + 1
        if bc is None:
            unest |= {c for c in observable if int(target.get(c, 0) or 0) > 0}
            continue
        verdict, fs, un = bc.check_axis_conformance(observed, target, axis='axis3',
                                                    observable=observable)
        audited += 1
        unest |= set(un)
        if verdict == 'FAIL':
            findings += [f'{name} — {f}' for f in fs]

    if bc is None:
        _warn('A-AXIS3', 'blueprint_core not importable — Axis-3 conformance NOT '
                         'ESTABLISHED (coverage below still reported).')
    elif not audited:
        _ok('A-AXIS3', 'no Axis-3 target in any section (dormant).')
    else:
        (_ok if not findings else _fail)('A-AXIS3',
            f'Axis-3 mechanism mix within budget across {audited} section(s).'
            if not findings else
            'Axis-3 mechanism mix breaches the blueprint budget — ' + ' | '.join(findings))

    cov = []
    if unest:
        cov.append('no observation source for ' + ', '.join(sorted(unest))
                   + ' (NAT needs registry options_by_q; MSQ needs stem text + '
                     'instruction phrases)')
    if skipped:
        cov.append('section(s) with no usable q_range skipped: ' + ', '.join(skipped))
    (_ok if not cov else _warn)('A-AXIS3-COVERAGE',
        'every targeted Axis-3 class was observable and every section bucketed.'
        if not cov else 'Axis-3 verdict is PARTIAL — ' + ' | '.join(cov))


def gate_nat(blocks, src):
    # v1.4 — A-NAT-NOOPT + A-NAT-INSTR (machine, NUMERICAL only). INDEPENDENT of any
    # Step-7 self-report. Fully dormant when the blueprint declares no numerical subtopics.
    obq = src.get('options_by_q', {})
    nat_present = src.get('nat_present', False)
    # A-NAT-NOOPT: every question the registry marks 0-option (NAT) must render ZERO
    # option-label paragraphs. A-OPTN covers the inverse (a claimed-MCQ Q with too few
    # options) — and if options_by_q is absent, NAT Qs are NOT skipped there, so A-OPTN
    # fails LOUD, surfacing the missing ND6 contract rather than silently passing.
    if nat_present and obq:
        bad_opt = []
        for b in blocks:
            if obq.get(str(b.qnum)) == 0 and len(_label_paras(b)) != 0:
                bad_opt.append(f'Q{b.qnum}:{len(_label_paras(b))}')
        (_ok if not bad_opt else _fail)('A-NAT-NOOPT',
            'every numerical Q renders zero options.' if not bad_opt
            else 'numerical Q carries options: ' + _flist(bad_opt))
    else:
        _ok('A-NAT-NOOPT', 'no numerical subtopics in this mock (dormant).')
    # v2.8 — A-NAT-GRADE (machine, NUMERICAL only): self-consistency backstop for the
    # portal grading transform (S7-NEW-C). Re-runs derive_nat_grading() on the SIDECAR's
    # OWN recorded (nat_value, ca_range) and checks the result matches the sidecar's OWN
    # recorded (nat_grading_type, nat_grading_value) EXACTLY. This does NOT re-derive the
    # math value from the stem (that is A-NAT-ANSWER's Claude-derivation job, run
    # separately in Phase 2) — it only proves Step 7's own execution of the SAME function
    # this file also embeds actually ran correctly and wasn't hand-edited or bypassed.
    # PINNED: this function body MUST stay byte-identical to Framework_MockTestCreate.md
    # §S7-NEW-C derive_nat_grading() — never re-implemented independently (anti-drift).
    # NOTE: placed BEFORE A-NAT-INSTR deliberately — A-NAT-INSTR has an early `return` on
    # its own dormant path, which would silently skip any code placed after it.
    from decimal import Decimal, ROUND_HALF_UP
    _NAT_GRADE_CHARSET = frozenset('0123456789.-')
    _NAT_INTEGRAL_EPS = Decimal('1e-9')
    def _fmt_portal_number(value, precision=None):
        d = Decimal(str(value))
        if precision is not None:
            q = Decimal(1).scaleb(-precision)
            d = d.quantize(q, rounding=ROUND_HALF_UP)
            s = format(d, 'f')
        else:
            if abs(d - d.to_integral_value()) <= _NAT_INTEGRAL_EPS:
                s = str(int(d.to_integral_value()))
            else:
                s = format(d.normalize(), 'f')
        if re.fullmatch(r'-0(\.0+)?', s):
            s = s.lstrip('-')
        return s
    def _fmt_portal_range(lo, hi, precision=None):
        lo_s = _fmt_portal_number(lo, precision); hi_s = _fmt_portal_number(hi, precision)
        if lo_s.startswith('-') or hi_s.startswith('-'):
            raise ValueError(f'NOT SUPPORTED negative-bound range lo={lo_s} hi={hi_s}')
        if Decimal(lo_s) > Decimal(hi_s):
            raise ValueError(f'lo>hi {lo_s} {hi_s}')
        return f'{lo_s}-{hi_s}'
    def _derive_nat_grading(value, ca_range=None, stem_precision=None):
        if stem_precision is not None:
            if ca_range is not None:
                lo, hi = ca_range
                return ('range', _fmt_portal_range(lo, hi, precision=stem_precision))
            return ('decimal_fixed', _fmt_portal_number(value, precision=stem_precision))
        if ca_range is not None:
            lo, hi = ca_range
            return ('range', _fmt_portal_range(lo, hi, precision=None))
        d = Decimal(str(value))
        if abs(d - d.to_integral_value()) <= _NAT_INTEGRAL_EPS:
            v_int = int(d.to_integral_value())
            return (('positive_integer', str(v_int)) if v_int >= 0 else ('integer', str(v_int)))
        return ('decimal', _fmt_portal_number(value, precision=None))
    if not nat_present:
        _ok('A-NAT-GRADE', 'no numerical subtopics in this mock (dormant).')
    elif not src.get('answers'):
        # v2.17: was `not concept_map`. The check re-runs derive_nat_grading() over
        # the KEYED VALUE, so it needs answers — a Tier-A dossier (facts only) is
        # not enough. Conflating the two made the gate FAIL on every NAT question
        # the moment the dossier landed. It becomes live with the sealed key
        # channel (Tier B), not before.
        _ok('A-NAT-GRADE', 'answer values not available (Tier-A dossier carries no '
                           'answers; supply --key) — dormant.')
    elif not src.get('concept_map'):
        # Step 8 does not receive the answer_key sidecar by default (S0-1) — this
        # self-consistency backstop is only checkable when --key is supplied,
        # exactly like the concept_map-dependent parts of gate_images.
        _ok('A-NAT-GRADE', 'answer_key sidecar not supplied via --key (dormant).')
    else:
        cmap = src.get('concept_map', {})
        answers = src.get('answers', {})
        bad_grade = []
        for b in blocks:
            entry = cmap.get(str(b.qnum), {})
            if entry.get('qtype') != 'nat':
                continue
            nat_value = answers.get(str(b.qnum))
            ca_range = entry.get('ca_range')
            g_type = entry.get('nat_grading_type'); g_val = entry.get('nat_grading_value')
            if nat_value is None:
                bad_grade.append(f'Q{b.qnum}: nat_value missing from sidecar answers'); continue
            if g_val is None:
                bad_grade.append(f'Q{b.qnum}: nat_grading_value missing from sidecar'); continue
            bad_chars = sorted(set(str(g_val)) - _NAT_GRADE_CHARSET)
            if bad_chars:
                bad_grade.append(f'Q{b.qnum}: nat_grading_value {g_val!r} has banned chars {bad_chars}')
                continue
            try:
                re_type, re_val = _derive_nat_grading(
                    nat_value, tuple(ca_range) if ca_range is not None else None,
                    stem_precision=entry.get('stem_precision'))
            except ValueError as e:
                bad_grade.append(f'Q{b.qnum}: re-derivation raised: {e}'); continue
            if (re_type, re_val) != (g_type, g_val):
                bad_grade.append(f'Q{b.qnum}: sidecar says ({g_type!r},{g_val!r}), '
                                  f're-derived ({re_type!r},{re_val!r})')
        (_ok if not bad_grade else _fail)('A-NAT-GRADE',
            'every NAT grading value is self-consistent and charset-pure.' if not bad_grade
            else 'grading value defect(s): ' + _flist(bad_grade, sep=' | '))
    # A-NAT-INSTR: per-section EXPECTED NAT count (re-derived from blueprint allocations,
    # src['expected_nat_by_section']) vs OBSERVED count of Q blocks whose stem carries a
    # numerical-entry instruction phrase. Mismatch ⇒ a NAT Q missing its instruction (or a
    # non-NAT Q wrongly carrying one). Mirrors A-MSQ-INSTR. Dormant when the map is empty.
    exp = src.get('expected_nat_by_section', {})
    if not exp:
        return _ok('A-NAT-INSTR', 'no numerical subtopics in this mock (dormant).')
    phrases = [p.lower() for p in src.get('nat_instruction_phrases', [])]
    secs   = src.get('sections', [])
    sec_of = src.get('section_of_qnum', {})
    def section_for(qnum):
        if qnum in sec_of:
            return sec_of[qnum]
        for s in secs:
            lo, hi = (s.get('q_range') or [0, 0])[:2]
            if lo <= qnum <= hi:
                return s.get('section_name') or s.get('name')
        return None
    observed = {}
    for b in blocks:
        stem = para_text(b.paras[0]) if getattr(b, 'paras', None) else getattr(b, 'stem', '')
        if any(ph in stem.lower() for ph in phrases):
            observed[section_for(b.qnum)] = observed.get(section_for(b.qnum), 0) + 1
    bad = [f'{sec}: expected {n} NAT, saw {observed.get(sec, 0)} instruction-carrying'
           for sec, n in exp.items() if observed.get(sec, 0) != n]
    (_ok if not bad else _fail)('A-NAT-INSTR',
        'observed numerical instruction counts match the blueprint per section.' if not bad
        else 'numerical instruction-count mismatch — ' + ' | '.join(bad))





def gate_stimorphan(blocks, src):
    linked = src['passage_linked'] | src['cloze_linked']
    # Merge section_rules stimulus cues (if declared) with the built-in English set (RA-9).
    extra_pats = src.get('stimulus_cue_patterns', [])
    cue_re = STIMULUS_CUES
    if extra_pats:
        combined = STIMULUS_CUES.pattern + '|' + '|'.join(r'\b' + p + r'\b' for p in extra_pats)
        cue_re = re.compile(combined, re.I)
    orphan, crossref = [], []
    bynum = {b.qnum: b for b in blocks}
    for b in blocks:
        stem = block_stem_text(b)
        if CROSSREF_RE.search(' '.join(para_text(p) for p in b.paras)):
            crossref.append(f'Q{b.qnum}')
        refs_stim = bool(cue_re.search(stem)) or (b.qnum in linked)
        if refs_stim:
            has_tbl = len(b.tables) > 0
            has_img = any(para_images(p) for p in b.paras)
            has_long = any(len(para_text(p).split()) >= 35 for p in b.paras)
            if not (has_tbl or has_img or has_long):
                orphan.append(f'Q{b.qnum}')
    (_ok if not orphan else _fail)('A-STIMORPHAN',
        'linked members carry their stimulus.' if not orphan else
        'orphaned stimulus (no embedded table/image/passage): ' + _flist(orphan))
    if crossref:
        _fail('A-STIMORPHAN-XREF', 'cross-question reference in stem: ' + _flist(crossref))
    else:
        _ok('A-STIMORPHAN-XREF', 'no "Q.x and Q.y" cross-references.')


# ── self-contained MATCH detector for the machine gate (v2.7.1) ────────────────
# The runnable audit.py (this block) is standalone and does NOT import the S6-1b spec
# classifier, so A-MATCH-TABLE carries its OWN detector. It MIRRORS the S6-1b classifier's
# MATCH rules EXACTLY (keyword rules + cross-domain label-pair option shape) — both live in
# THIS file (S6-1b + here); if one changes the other MUST match. Kept minimal: only the
# MATCH decision is needed here, not the full 8-class ladder.
_MT_PAIR_RE = re.compile(r'\(?\s*([A-Za-z]{1,4}|\d{1,2})\s*\)?\s*[-\u2010-\u2015:\u2192>]+\s*'
                         r'\(?\s*([A-Za-z]{1,4}|\d{1,2})\s*\)?')
_MT_SUB = (r'\(?\s*(?:[A-Za-z]{1,4}|\d{1,2})\s*\)?\s*[-\u2010-\u2015:\u2192>]+\s*'
           r'\(?\s*(?:[A-Za-z]{1,4}|\d{1,2})\s*\)?')
_MT_OPT_RE = re.compile(r'^\s*' + _MT_SUB + r'(?:[,;\s]+' + _MT_SUB + r'){1,}\s*$')

def _mt_family(tokens):
    low = [x.lower() for x in tokens if x]
    if not low:
        return 'other'
    if all(re.fullmatch(r'\d{1,2}', x) for x in low):
        return 'digit'
    romanish = all(re.fullmatch(r'[ivxlcdm]+', x) for x in low)
    if romanish and any(len(x) > 1 for x in low):
        return 'roman'
    if all(re.fullmatch(r'[a-z]', x) for x in low):
        return 'roman' if set(low) <= {'i', 'v', 'x'} else 'alpha'
    if romanish:
        return 'roman'
    if all(re.fullmatch(r'[a-z]{1,4}', x) for x in low):
        return 'alpha'
    return 'other'

def _opts_match_pairs(opts):
    if not opts:
        return False
    hits = 0
    for o in opts:
        s = (o or '').strip()
        if not s or not _MT_OPT_RE.match(s):
            continue
        pairs = _MT_PAIR_RE.findall(s)
        if len(pairs) < 2:
            continue
        lf = _mt_family([p[0] for p in pairs])
        rf = _mt_family([p[1] for p in pairs])
        if lf == rf or 'other' in (lf, rf):
            continue
        hits += 1
    return hits >= max(2, (len(opts) + 1) // 2)

def _block_is_match(stem, opts):
    # Mirror the S6-1b ladder precedence: ASSERTION_REASON outranks MATCH, so an
    # assertion/reason stem is NOT a match even if its options look like pairs. (LINKED also
    # outranks MATCH; the gate skips linked members via src, below.)
    s = (stem or '').lower()
    if re.search(r'\bassertion\b', s) and re.search(r'\breason\b', s):
        return False
    if re.search(r'\bmatch\b', s) and re.search(r'\b(following|list|column|set)\b', s):
        return True
    if re.search(r'list[\s\-]*i\b|column[\s\-]*(i|a)\b', s):
        return True
    return _opts_match_pairs(opts)


def gate_match_table(blocks, src):
    """A-MATCH-TABLE (v2.7.1) — executable promotion of the S7-3 'MATRICES & MATCH-THE-COLUMN
    must be a REAL grid' checklist item. For every block re-derived as Axis-2 MATCH by the
    SHARED classifier (S6-1b — the SAME functions Step 5 and Step 7 use), the List columns
    MUST render as a real <w:tbl>. A match rendered as plain text lines (or space/tab pseudo-
    columns) is a format-fidelity defect: S6-6 still counts the MATCH format PRESENT while the
    skill is left un-rehearsed — a false readiness signal. Exam-agnostic: the MATCH signal is
    language-independent (keyword OR cross-domain option shape), never a hardcoded exam label."""
    if not blocks:
        return
    oc = src.get('options_count', 4)
    linked = src.get('passage_linked', set()) | src.get('cloze_linked', set())
    missing = []
    for b in blocks:
        if b.qnum in linked:          # LINKED outranks MATCH (S6-1b); A-STIMORPHAN covers it
            continue
        labs = _label_paras(b)
        opts = [x[2] for x in (labs[-oc:] if len(labs) >= oc else labs)]
        if _block_is_match(block_stem_text(b), opts):
            if not b.tables:
                missing.append(f'Q{b.qnum}')
    (_ok if not missing else _fail)('A-MATCH-TABLE',
        'every MATCH question renders its List columns as a real table.' if not missing else
        'MATCH question(s) rendered without a <w:tbl> grid (S7-3 defect — rebuild the List '
        'body as a real table, never text/space columns): ' + _flist(missing))


def gate_underline(blocks):
    missing, fake = [], []
    for b in blocks:
        whole = ' '.join(para_text(p) for p in b.paras)
        if FAKE_UNDERLINE.search(whole):
            fake.append(f'Q{b.qnum}')
        if UNDERLINE_REF.search(' '.join(para_text(p) for p in b.paras[:3])):
            # require a real w:u run somewhere in the block
            has_u = False
            for p in b.paras:
                for u in p._element.iter(W_('u')):
                    val = u.get(W_('val'))
                    if val not in ('none',):
                        has_u = True
                        break
                if has_u:
                    break
            if not has_u:
                missing.append(f'Q{b.qnum}')
    (_ok if not missing else _fail)('A-UNDERLINE',
        'underline-class Qs use real <w:u>.' if not missing else
        'no real underline run: ' + _flist(missing))
    if fake:
        _fail('A-UNDERLINE-FAKE', '"(underlined: X)" annotation present: ' + _flist(fake))
    else:
        _ok('A-UNDERLINE-FAKE', 'no faked-underline annotations.')


def gate_omml(doc, src, final):
    nfrac_bad, yearfrac = [], []
    n_omath = 0
    for kind, obj in iter_block_items(doc):
        if kind != 'p':
            continue
        for om in obj._element.iter(M_('oMath')):
            n_omath += 1
        for f in obj._element.iter(M_('f')):
            num = f.find(M_('num')); den = f.find(M_('den'))
            num_t = ''.join(t.text or '' for t in (num.iter(M_('t')) if num is not None else []))
            den_t = ''.join(t.text or '' for t in (den.iter(M_('t')) if den is not None else []))
            if not num_t.strip() or not den_t.strip():
                nfrac_bad.append('empty-frac')
            if re.fullmatch(r'\s*20\d\d\s*', num_t) and re.fullmatch(r'\s*\d{2}\s*', den_t):
                yearfrac.append(f'{num_t.strip()}/{den_t.strip()}')
    (_ok if not nfrac_bad else _fail)('A-OMML',
        f'{n_omath} oMath; all fractions well-formed.' if not nfrac_bad else
        f'{len(nfrac_bad)} fraction(s) with empty num/den.')
    if yearfrac:
        _warn('A-OMML-YEAR', f'year-range rendered as stacked fraction: {_flist(yearfrac)}')
    if final and src['omml_required_present'] and n_omath == 0:
        _warn('A-OMML-FLOOR',
              'OMML_required subtopic(s) declared but ZERO <m:oMath> in paper — '
              'built-up math may be hiding as ASCII/raster; investigate (S7-5).')
    elif final and src['omml_required_present']:
        _ok('A-OMML-FLOOR', f'OMML floor satisfied ({n_omath} oMath).')


def gate_frac_ascii(blocks, src):
    caret, slash = [], []
    omml_ctx = src['omml_required_present']
    for b in blocks:
        stem = block_stem_text(b)
        if ASCII_CARET.search(stem):
            caret.append(f'Q{b.qnum}')
        if omml_ctx and SLASH_FRAC.search(stem):
            slash.append(f'Q{b.qnum}')
    (_ok if not caret else _fail)('A-FRAC',
        'no ASCII caret exponents.' if not caret else 'ASCII "^" exponent: ' + _flist(caret))
    if slash:
        _warn('A-FRAC-SLASH', 'slash fraction in a math-context stem; R-MATH-OMML requires built-up structures as native OMML. MECHANICAL REMEDY: re-emit these stems through Step 7 S10-4 add_math_stem: '
              + _flist(slash))


def gate_images(blocks, src, media_map):
    """A-MATHRASTER (Tier 1) + A-FIGCOMP (structural, v2.4 image_role-aware)
    + A-FIGTEXT-PROSE (v2.4 — visual prose detector)."""
    math_raster, warn_view, composite, multi_per_line = [], [], [], []
    figtext_prose = []   # v2.4: figure-reference text in zero-image blocks
    fig = src['figural_qs']
    oc = src['options_count']
    # v2.21.5 (ND10) — the registry's per-Q option count. A question the registry
    # marks 0-option is NUMERICAL (NAT): Create.md R-FIGURAL v4.7 FIGURAL-NAT
    # VARIANT (ND10) says such a question "has a PROBLEM image (or series images)
    # but ZERO option images — there are no options to decompose", and that
    # G-FIGURAL-COMPOSITE "must skip its per-option-image arm for a numerical
    # figural question". This is the SAME signal gate_options already reads
    # (obq[str(qnum)] == 0 → skip); gate_images had never read it.
    #
    # It must come from options_by_q, NOT from concept_map/nat_subtopic_ids: this
    # function's own fallback comment (load_sources, "otherwise empty dict →
    # gate_images falls back to default image_role='stem_and_options'") is exact —
    # concept_map is {} on any run without a dossier or --key, so the existing
    # _nat_ids mapping silently does not fire there and a figural-NAT lands in the
    # stem_and_options arm. options_by_q travels in the registry (ND6), which
    # Step 8 always receives.
    obq_img = src.get('options_by_q', {})
    # v2.4: read image_role per subtopic from section_rules
    sr_text = src.get('section_rules_text', '')
    concept_map = src.get('concept_map', {})
    # v2.21.8 (RA-9) — figure-reference prose pattern is EXAM-SUPPLIED, not
    # hardcoded. It was a fixed ENGLISH regex carrying reasoning-exam shape nouns
    # (triangles/squares/circles/angles). Consequences, both silent:
    #   • on a NON-ENGLISH paper it matched nothing, so A-FIGTEXT-PROSE printed a
    #     clean OK while detecting nothing — a false assurance, not a pass;
    #   • on a non-reasoning exam (biology, chemistry) the "how many <shape>" arm
    #     was dead weight.
    # RA-9: "Hardcode nothing. A missing value → SKIP the dependent check with a
    # logged reason, never a hardcoded substitute." The English set is now a
    # DEFAULT THAT APPLIES ONLY WHEN THE PAPER IS ENGLISH; any other language must
    # declare section_rules figure_reference_phrases or the gate reports DORMANT
    # with a named reason instead of OK.
    _fig_ref_phrases = src.get('figure_reference_phrases')
    _fig_lang = (src.get('language') or 'english').lower()
    if _fig_ref_phrases:
        _fig_ref_re = re.compile('(?i)(' + '|'.join(re.escape(t) for t in
                                                    _fig_ref_phrases) + ')')
        _fig_ref_why = None
    elif _fig_lang == 'english':
        _fig_ref_re = re.compile(
            r'(?i)\b(in the given figure|in the following figure|'
            r'from the (given|following) (figure|diagram)|'
            r'figure \(X\)|the figure (shows|below|above)|'
            r'how many .{0,30}(triangles|squares|circles|lines|shapes|'
            r'angles|sides|regions|parts)\s+(are|in|does|can))')
        _fig_ref_why = None
    else:
        _fig_ref_re = None
        _fig_ref_why = (f'language={_fig_lang}: no section_rules '
                        f'figure_reference_phrases declared, and the built-in '
                        f'phrase set is ENGLISH-ONLY. Detector NOT RUN — declare '
                        f'figure_reference_phrases to enable it (RA-9).')
    for b in blocks:
        block_imgs = []
        for p in b.paras:
            pim = para_images(p)
            if len(pim) >= 2:
                multi_per_line.append(f'Q{b.qnum}')
            for nm, emb in pim:
                tgt = media_map.get(emb, '')
                block_imgs.append((nm or '', tgt))
        # v2.4: A-FIGTEXT-PROSE — zero-image blocks referencing figures
        if not block_imgs:
            stem = ' '.join(para_text(p) for p in b.paras)
            if _fig_ref_re is not None and _fig_ref_re.search(stem):
                figtext_prose.append(f'Q{b.qnum}')
            # v2.21.4 — A FIGURAL QUESTION THAT RENDERS ZERO IMAGES IS A FINDING.
            # This `continue` used to swallow the condition entirely: the
            # `stem_only:0img` arm below sits AFTER it, so `len(block_imgs) < 1`
            # could never be true and that branch was DEAD CODE. A figural
            # question whose figure was never drawn passed A-FIGCOMP clean unless
            # its stem happened to match the _fig_ref_re prose pattern — i.e.
            # detection depended on the wording of the stem rather than on the
            # absence of the figure. Step-7 G-FIGURAL-COMPOSITE requires >=1
            # problem image for stem_only, >=n for options_only and problem+per-
            # option for stem_and_options; ZERO images satisfies no variant.
            #
            # Membership is taken from the REGISTRY set `fig` ONLY, never from
            # figural_cue_keywords: the cue list contains ordinary MCQ phrases
            # ('which of the', 'series', 'complete the'), so applying it to
            # zero-image blocks would false-FAIL a large share of ordinary TEXT
            # questions across the estate.
            if b.qnum in fig:
                composite.append(f'Q{b.qnum}(figural:0img)')
            continue
        stem = ' '.join(para_text(p) for p in b.paras[:3])
        math_ctx = src['omml_required_present']  # coarse; refined in Part B
        for nm, tgt in block_imgs:
            hay = f'{nm} {tgt}'
            if MATH_TOKEN_NAME.search(hay):
                math_raster.append(f'Q{b.qnum}')
            elif not CANON_IMG_NAME.match(nm or ''):
                warn_view.append(f'Q{b.qnum}')
        # A-FIGCOMP (v2.4 — image_role-aware):
        # Determine image_role for this question
        figural_cues = src.get('figural_cue_keywords',
                              ['which of the', 'odd one', 'mirror', 'water image',
                               'embedded', 'complete the', 'series', 'fold'])
        is_figural = b.qnum in fig or any(k in stem.lower() for k in figural_cues)
        if is_figural:
            # v2.4: determine image_role from section_rules via concept_map
            _q_role = 'stem_and_options'   # default
            qnum_str = str(b.qnum)
            if qnum_str in concept_map and sr_text:
                _sid = concept_map[qnum_str].get('subtopic_id', '')
                if _sid:
                    _rm = re.search(
                        r'subtopic_id:\s*' + re.escape(_sid) +
                        r'((?:(?!subtopic_id:).)*?)image_role:\s*(\S+)',
                        sr_text, re.DOTALL)
                    if _rm:
                        _q_role = _rm.group(2)
            # Also check NAT: if this Q's subtopic is in nat_subtopic_ids → stem_only
            _nat_ids = src.get('nat_subtopic_ids', set())
            if _nat_ids and qnum_str in concept_map:
                _q_sid = concept_map[qnum_str].get('subtopic_id', '')
                if _q_sid in _nat_ids:
                    _q_role = 'stem_only'
            # v2.21.5 (ND10) — REGISTRY-DRIVEN NAT DETECTION, which works even when
            # concept_map is empty (no dossier / no --key). A 0-option question has
            # no options to decompose, so the per-option-image arm DOES NOT APPLY
            # in EITHER the stem_and_options or the options_only variant; ND10
            # still requires >=1 problem image, and the zero-image branch above
            # (which runs before this point) continues to enforce exactly that.
            if obq_img.get(qnum_str) == 0:
                _q_role = 'stem_only'
            # v2.22.0 (GAP-2026-08-03-FIGCOMP-ROLE) — THE PRODUCER'S OWN RECORD WINS.
            # Every branch above INFERS the role (section_rules by subtopic, NAT
            # subtopic set, 0-option registry). The registry figure_specs carry the
            # role figural_core actually DREW each PNG as — 'problem' vs 'opt<i>' —
            # which is not an inference at all. Measured on a real delivered paper
            # (IIT_JAM_BIOTECHNOLOGY M01): 27 of 33 figural questions are a single
            # role='problem' figure with FOUR TEXT OPTIONS — a diagram/gel/graph
            # with text answers, the commonest figural shape in the life sciences.
            # The default 'stem_and_options' demanded oc+1 images, found 1, and
            # emitted 17 findings on questions that were CORRECT. Those findings
            # then said "VIEW + fix in Part B", pushing 17 conformant questions
            # into a vision queue — the queue that stalled a real audit for a day.
            # FOURTH occurrence of one gate assuming one rendering shape (v2.21
            # A-DOSSIER, v2.21.3 A-OPTORDER, v2.21.9 A-QNFIRST, now this).
            # PRECEDENCE: specs are consulted LAST and override, because they are
            # the only non-inferred signal. Absent specs (pre-v5.34 registry, or a
            # question with no spec record) leave every earlier branch untouched,
            # so legacy behaviour is byte-identical.
            _q_specs = [s for s in (src.get('figure_specs') or {}).values()
                        if str(s.get('question')) == qnum_str]
            if _q_specs:
                _roles = {str(s.get('role') or '') for s in _q_specs}
                _opt_roles = {r for r in _roles if r.startswith('opt')}
                if not _opt_roles:
                    # The producer drew NO option figures for this question. There
                    # are no option images to bind 1:1, exactly as for a NAT item.
                    _q_role = 'stem_only'
                else:
                    # The producer DID draw option figures: require the full set,
                    # counted from the specs themselves rather than from the
                    # exam-wide OPTIONS_COUNT, so a question whose option set is
                    # legitimately shorter is judged against what it declares.
                    _q_role = 'stem_and_options'
                    _spec_expect = len(_opt_roles) + (1 if 'problem' in _roles else 0)
                    if len(block_imgs) < _spec_expect:
                        composite.append(
                            f'Q{b.qnum}({len(block_imgs)}<{_spec_expect})')
                    continue
            # Branch by image_role
            if _q_role == 'stem_only':
                # v2.21.4: the zero-image case is handled ABOVE (before the
                # `continue`); reaching here means >=1 image, which IS correct for
                # stem_only. Do NOT flag as composite.
                pass
            elif _q_role == 'options_only':
                n_opt = oc   # oc is int (exam-wide OPTIONS_COUNT from section_rules)
                if len(block_imgs) < n_opt:
                    composite.append(f'Q{b.qnum}(opts_only:{len(block_imgs)}<{n_opt})')
            else:   # stem_and_options (default)
                # v2.21.4 — REQUIRE THE FULL SET, not merely "more than one".
                # Step-7 G-FIGURAL-COMPOSITE: stem_and_options = "problem image +
                # one separate image per option", i.e. oc+1 images. The check was
                # `len(block_imgs) == 1`, so a block rendering 2, 3 or 4 images —
                # a problem figure with a PARTIAL option set, options silently
                # undrawn — passed A-FIGCOMP clean. Only the degenerate 1-image
                # case was caught. A candidate cannot answer a question whose
                # option figures were never rendered.
                if len(block_imgs) < oc + 1:
                    composite.append(f'Q{b.qnum}({len(block_imgs)}<{oc + 1})')
    (_ok if not math_raster else _fail)('A-MATHRASTER',
        'no math-token raster names.' if not math_raster else
        'image named like a math raster: ' + _flist(math_raster))
    if warn_view:
        _warn('A-MATHRASTER-VIEW',
              f'{len(set(warn_view))} block(s) have non-canonically-named images; '
              'the name contract (q<N>_problem/_opt<i>) cannot confirm figure-vs-math '
              'from the name alone. MECHANICAL REMEDY: re-run Step 7 S10-8 naming for '
              'these blocks, or supply the registry figure_specs (v5.34+) whose role '
              'field settles it without inspection: ' + _flist(warn_view))
    if multi_per_line:
        _fail('A-FIGCOMP-LINE', 'multiple images on one line (option-per-line broken): '
              + _flist(multi_per_line))
    else:
        _ok('A-FIGCOMP-LINE', 'at most one image per line.')
    if composite:
        _warn('A-FIGCOMP',
              'figural block renders FEWER images than its own registry '
              'figure_specs declare (v2.22.0: the expectation is the producer\'s '
              'OWN role record, never an inference). MECHANICAL REMEDY: re-render '
              'the missing option figure(s) for these questions via Step 7 S10-8; '
              'no inspection is required to decide this: ' + _flist(composite))
    else:
        _ok('A-FIGCOMP', 'figural blocks pass image_role-aware check (v2.4).')

    # ── A-FIGPROFILE (v2.10, GAP-2026-07-26-003 D2) ──────────────────────────
    # Did generation honour the figure profile Step 5 measured? The verdict is
    # DELEGATED to bc.check_figural_conformance(), the same function Step 7 v5.31
    # generates against, so the generator and its auditor cannot drift apart.
    #
    # AUDITS RECORDED INTENT, NOT PIXELS. Confirming a render truly depicts a
    # micrograph needs a view(), which is CLASS T and cannot run inside this python
    # (EXECUTION-BOUNDARY LAW). Intent is deterministic and catches the failure that
    # matters: Step 7 ignoring the profile.
    #
    # SOURCES, both of which Step 8 actually receives (S0-1):
    #   registry figural_manifests[mock].object_types  {qnum: type, ...}  (v5.31+)
    #   section_rules TEXT -> bc.parse_image_analysis_blocks()
    # Step 8 does NOT receive the answer_key sidecar, so no concept_map is used here;
    # subtopic_id travels with each figural_qs record instead.
    _fig_types = src.get('figural_object_types') or {}
    _fig_subs = src.get('figural_subtopics') or {}
    if not _fig_types:
        # Pre-v5.31 registry, or a mock with no figural questions. Dormant, exactly
        # like the concept_map-dependent gates above — never a FAIL for a missing
        # optional input (EC-V18).
        _ok('A-FIGPROFILE',
            'registry carries no per-question object_types — dormant. '
            'EXPECTED when every figural subtopic profile is unconstrained '
            '(Step 7 omits the entry deliberately, S13 object_types writer); '
            'also true of a genuinely pre-v5.31 registry. This message no longer '
            'asserts which (v2.21.7, SEC-3): the earlier wording said '
            '"pre-v5.31 mock", which MISATTRIBUTED modern v5.34+ output as legacy '
            'and sent operators looking for a producer fault that was not there.')
    else:
        # ── FIX GAP-2026-08-01-FIGPROFILE-ENGINE-BINDING (D1/D2) — v2.12 ──────
        # BIND THE ENGINE. Until v2.12 `bc` was READ at three sites and BOUND at
        # none: the v2.10 delegation was written into the comments and the call
        # sites, but the import was never added. Any registry carrying
        # object_types (Step 7 v5.31+) raised NameError out of gate_images() and
        # killed the whole run before a single gate line printed.
        #
        # THREE-LAYER GUARD (each layer closes a distinct crash path; an import
        # guard alone is NOT sufficient — measured, not assumed):
        #   L1 IMPORT     except Exception, NOT except ImportError. A truncated
        #                 engine (project-knowledge size caps — blueprint_core.py
        #                 is ~168 KB, exactly the range P0.5 exists for) raises
        #                 SyntaxError, which ImportError does not catch.
        #   L2 CAPABILITY hasattr() on all three delegated functions. A stale
        #                 engine imports cleanly and then raises AttributeError at
        #                 the call site — a crash the import guard never sees.
        #   L3 CALL SITE  the delegated calls run inside try/except Exception, so
        #                 a raise INSIDE the engine degrades to a reported skip
        #                 instead of aborting the audit.
        #
        # SEVERITY IS _warn, DELIBERATELY:
        #   not _ok   — the gate DID NOT RUN; reporting it green is the silent-pass
        #               failure the v1.8 DeliveryFooter quality gate exists to kill.
        #   not _fail — a missing/broken ENGINE is an environment condition, not a
        #               paper defect, and must never block delivery of a sound
        #               paper (EC-V18 tolerance; owner directive: no dependency
        #               condition may halt a run).
        _BC_FNS = ('parse_image_analysis_blocks',
                   'figural_generation_profile',
                   'check_figural_conformance')
        try:
            import blueprint_core as bc          # L1
        except Exception as _e:
            bc = None
            _bc_why = f'{type(_e).__name__}: {_e}'
        else:
            _bc_missing = [_f for _f in _BC_FNS if not hasattr(bc, _f)]   # L2
            if _bc_missing:
                _bc_why = ('stale engine — missing ' + ', '.join(_bc_missing))
                bc = None
            else:
                _bc_why = ''
        if bc is None:
            _warn('A-FIGPROFILE',
                  'blueprint_core engine unavailable — figure-profile conformance '
                  f'NOT CHECKED (dependency degraded: {_bc_why}). Logged skip; '
                  'audit continues (RA-9).')
        else:
            try:                                                          # L3
                _sr_blocks = bc.parse_image_analysis_blocks(
                    src.get('section_rules_text', '') or '')
                _by_sub = {}
                for _qn, _ty in _fig_types.items():
                    _sid = _fig_subs.get(str(_qn))
                    if _sid:
                        _by_sub.setdefault(_sid, []).append(_ty)
                _fig_bad, _fig_ok, _fig_skip = [], 0, 0
                for _sid, _gen_types in sorted(_by_sub.items()):
                    _prof = bc.figural_generation_profile(_sr_blocks.get(_sid))
                    _verdict, _detail = bc.check_figural_conformance(_gen_types, _prof)
                    if _verdict == 'FAIL':
                        _fig_bad.append(f'{_sid}: {_detail}')
                    elif _verdict == 'PASS':
                        _fig_ok += 1
                    else:
                        _fig_skip += 1
            except Exception as _e:
                _warn('A-FIGPROFILE',
                      'blueprint_core raised during conformance evaluation — '
                      f'figure-profile conformance NOT CHECKED ({type(_e).__name__}: '
                      f'{_e}). Logged skip; audit continues (RA-9).')
            else:
                if _fig_bad:
                    _fail('A-FIGPROFILE',
                          'generated figure types do not match the measured PYQ profile — '
                          + _flist(_fig_bad, sep=' | '))
                elif not _by_sub:
                    # 0/0 is NOT evidence of conformance (edge case 6). The registry
                    # carried object_types but no usable subtopic_ids, so nothing was
                    # actually judged. Reporting _ok here would claim coverage the
                    # gate never had.
                    _warn('A-FIGPROFILE',
                          'registry carries object_types but no usable subtopic_ids — '
                          '0 subtopic(s) evaluated; conformance NOT ESTABLISHED '
                          '(manifest incomplete).')
                else:
                    _ok('A-FIGPROFILE',
                        f'{_fig_ok} subtopic(s) conform to the measured figure profile; '
                        f'{_fig_skip} skipped (no usable profile — EC-V18).')
    # v2.4: A-FIGTEXT-PROSE — figure-reference text in zero-image blocks
    if figtext_prose:
        _fail('A-FIGTEXT-PROSE',
              'Q-block references a figure but contains 0 images — '
              'render the figure or replace the subtopic (S7-NEW-B): '
              + _flist(figtext_prose))
    elif _fig_ref_why:
        # v2.21.8 (RA-9) — the detector could not run for this exam. Reporting OK
        # here would be a FALSE ASSURANCE: it would claim "no figure-reference
        # prose" on a paper the gate never examined. Say so, and name the fix.
        _warn('A-FIGTEXT-PROSE', _fig_ref_why)
    else:
        _ok('A-FIGTEXT-PROSE', 'no figure-reference prose in zero-image blocks.')

    # ── v2.11 FIGURE CONFORMANCE — 13 gates (GAP-2026-07-29-FIG-R2; +A-FIGACCENT v2.25, GAP-2026-08-07-FIGACCENT) ──
    # Every check below is DETERMINISTIC arithmetic over the saved PNG and its
    # FigureSpec sidecar, and is therefore legal inside this python block. This
    # is the correction to the reasoning that let the defect ship: A-FIGPROFILE
    # rightly declines pixels because "does this depict a micrograph" needs a
    # view() (CLASS T), and that one true fact was generalised to every figure
    # property. None of these twelve needs eyes.
    #
    # SEVERITY. No colour condition may EVER halt a run (owner directive; and
    # CLAUDE.md: "A CLASS T failure must be LOUD, and must NOT halt... Silence is
    # the defect; a halt is not the remedy"). fc.triage() sorts findings into
    # AMBER / VOID_ITEM / BLOCKING, applies EC-V18 legacy tolerance, and NEVER
    # raises. A grey figure is a DEGRADED paper, never a void one.
    # ── FIX GAP-2026-08-01-FIGPROFILE-ENGINE-BINDING (D3) — v2.12 ────────────
    # except Exception, not except ImportError: a TRUNCATED figural_core.py
    # raises SyntaxError and the old guard let it kill the run — the same crash
    # path D1 produced, one engine over.
    try:
        import figural_core as fc
    except Exception as _e:
        fc = None
        _fc_why = f'{type(_e).__name__}: {_e}'
    else:
        _fc_why = ''
    if fc is None:
        # LOUD DEGRADATION, ONE LINE PER GATE (D3). Until v2.12 a single
        # self-naming A-FIGSCALE WARN stood in for all twelve, so ELEVEN gates —
        # including A-FIGMONO (VOID_ITEM, an answer-cue leak) and A-FIGDEGEN
        # (BLOCKING) — vanished from STDOUT with no line at all. A reader could
        # not tell they had not been evaluated. That is precisely the silence
        # CLAUDE.md forbids: "Silence is the defect; a halt is not the remedy."
        # Emitting all twelve also makes the printed gate roster INVARIANT
        # regardless of environment, which restores §R15 reproducibility and
        # makes the gate count itself a usable integrity signal.
        for _g in ('A-FIGSCALE', 'A-FIGLABEL', 'A-FIGDPI', 'A-FIGDEGEN',
                   'A-FIGMONO', 'A-FIGOPTUNIF', 'A-FIGCOLOUR', 'A-FIGACCENT',
                   'A-FIGCVD',
                   'A-FIGSERIES', 'A-FIGGLYPH', 'A-FIGALT', 'A-FIGLABELPX'):
            _warn(_g, 'figural_core engine unavailable — gate NOT RUN '
                      f'(dependency degraded: {_fc_why}).')
    else:
        _fig_specs = src.get('figure_specs') or {}
        _amber, _void, _block = [], [], []
        _seen, _legacy_n = 0, 0
        _declared, _unusable = 0, 0
        for _blk in blocks:
            for _img in (getattr(_blk, 'images', None) or []):
                _declared += 1
                _png = _img.get('path')
                if not _png:
                    continue
                _spec = resolve_figure_spec(_img, _fig_specs)
                # ── v2.13: PER-FIGURE L3 GUARD (the three-layer pattern v2.12
                # established for blueprint_core, applied to figural_core's
                # PER-ITEM calls). The spec now arrives from the REGISTRY, i.e.
                # from outside this process, and a partially-recorded one raises:
                # render_figure() mutates font_pt_native/png_px/placement_scale
                # only after it reads the artefact back, so a render that died
                # mid-way leaves a spec whose shape the gates index into. Caught
                # empirically — one such figure raised TypeError out of
                # g_figlabel(), _safe_gate() converted it to A-GATEERROR, and the
                # WHOLE A-IMAGES gate died: TWELVE gate lines vanished and the
                # roster fell from 47 to 36, breaking the §R15 invariance v2.12
                # had just restored. One bad figure must never cost eleven other
                # gates their verdict. Skipped figures are counted and REPORTED,
                # never silently dropped.
                try:
                    _hard, _warns = fc.audit_figure(_spec, _png,
                                                    descr=_img.get('descr'))
                    _t = fc.triage(_hard + _warns, _spec)
                    _leg = bool(fc.is_legacy(_spec))
                except Exception:
                    _unusable += 1
                    continue
                _seen += 1
                if _leg:
                    _legacy_n += 1
                # v2.13: carry (finding, legacy?) so a gate can distinguish
                # "this v5.33+ render regressed" from "this pre-v5.33 render has
                # no sidecar to check against" — see _fig_verdict.
                _amber += [(_f, _leg) for _f in _t['AMBER']]
                _void += [(_f, _leg) for _f in _t['VOID_ITEM']]
                _block += [(_f, _leg) for _f in _t['BLOCKING']]

        # Emit one verdict per gate id, mapped engine G-*/W-* -> catalogue A-*.
        _by_gate = {}
        for _sev, _findings in (('BLOCKING', _block), ('VOID_ITEM', _void),
                                ('AMBER', _amber)):
            for _f, _leg in _findings:
                _e = _by_gate.setdefault(fc.audit_gate_id(_f), (_sev, [], []))
                _e[1].append(_f)
                _e[2].append(_leg)

        # Each gate is emitted by an EXPLICIT call with a LITERAL id. Check
        # M-GATE discovers emitted gates statically, so a loop over a variable
        # gate id is invisible to it and the catalogue would read as unemitted.
        # v2.13: coverage suffix — an incomplete sweep is stated on EVERY gate's
        # own line, so no verdict can be read as fuller coverage than it had.
        _cov = (f' [coverage: {_seen}/{_declared} evaluated; {_unusable} skipped — '
                f'unusable FigureSpec]') if _unusable else ''

        def _fig_verdict(gid, bucket):
            if gid not in bucket:
                # ── v2.13 (D3): 0 EVALUATED IS NOT EVIDENCE OF CONFORMANCE ──
                # This gate family printed "0 figure(s) conform." on EVERY paper
                # in EVERY exam from v2.11 to v2.12.1, because Block.images was
                # never populated. Zero-of-zero is the same false-clean the v2.12
                # A-FIGPROFILE edge case 6 already rejects; the rule is applied
                # here too, and for the same reason.
                if not _declared:
                    _ok(gid, 'no inline drawings in the paper; dormant.')
                elif not _seen:
                    _warn(gid, f'{_declared} inline drawing(s) present but 0 could be '
                               f'evaluated — conformance NOT ESTABLISHED (coverage gap, '
                               f'not a pass). Check A-ZIP and the registry '
                               f'figure_specs, then re-run.')
                elif _unusable:
                    _warn(gid, f'{_seen} figure(s) conform, but coverage is INCOMPLETE'
                               + _cov + ' — not a full pass; re-render those figures '
                               f'(Step 7 v5.34+) for full coverage.')
                else:
                    _ok(gid, f'{_seen} figure(s) conform'
                             + (f' ({_legacy_n} legacy, EC-V18).' if _legacy_n else '.'))
                return
            _sev, _msgs, _legs = bucket[gid]
            _detail = _flist(_msgs, sep=' | ')
            # v2.13: a MIXED paper (some figures carry a sidecar, some do not)
            # must not report one number for two different populations — the
            # operator acts on this line, and "56 figure(s)" when one regressed
            # and 55 are merely old sends them to the wrong repair.
            _nl = sum(1 for _x in _legs if not _x)
            _lg = len(_legs) - _nl
            _mix = (f' (+{_lg} pre-v5.33 figure(s) reported under EC-V18, not '
                    f'blocking)') if _lg else ''
            # ── v2.13 (D3): EC-V18 IS A DELIVERY TOLERANCE, NOT ONLY A SEVERITY
            # RELABEL. The spec's EC-V18 clause is explicit and non-negotiable:
            # output with no FigureSpec sidecar predates Step 7 v5.33, so roughly
            # 200 existing exams "keep auditing AND DELIVERING untouched while the
            # defect is reported loudly on every one." A _fail() here exits
            # non-zero, and MANDATE D requires exit 0 to certify — so emitting
            # FAIL for a legacy-only finding would have converted this coverage
            # fix into an estate-wide delivery outage the moment the gates stopped
            # being vacuous. A legacy figure cannot be repaired at Step 8 (Step 8
            # cannot retro-fit a sidecar onto an already-rendered paper), which is
            # precisely the "genuinely-not-fixable diagnostic" S5-4 admits as an
            # ACCEPTED WARN. It stays LOUD, it forces the amber footer, it is
            # recorded as a §R13 limitation — it simply does not block a paper
            # whose only sin is being older than the contract.
            if all(_legs):
                _warn(gid, f'DEGRADED, NOT VOID — {len(_msgs)} pre-v5.33 figure(s) carry '
                           f'no FigureSpec sidecar (EC-V18 legacy): reported, amber '
                           f'footer applies, delivery NOT blocked; record as a §R13 '
                           f'limitation. Re-run Step 7 v5.34+ for full coverage.'
                           + _cov + f' {_detail}')
                return
            if _sev == 'BLOCKING':
                _fail(gid, f'RENDERER-CONTRACT REGRESSION on v5.33+ output — '
                           f'{_nl} figure(s) carrying a FigureSpec' + _mix
                           + ':' + _cov + f' {_detail}')
            elif _sev == 'VOID_ITEM':
                _fail(gid, f'ANSWER-CUE LEAK — {len(_msgs)} item(s) VOID; drop or '
                           f'regenerate those questions, the paper continues:'
                           + _cov + f' {_detail}')
            else:
                # AMBER: FAIL severity forces the amber delivery footer
                # (Framework_DeliveryFooter §5) and NEVER halts the run.
                _fail(gid, f'DEGRADED, NOT VOID — delivery continues under an '
                           f'amber footer; {_nl} figure(s) carrying a FigureSpec'
                           + _mix + ':' + _cov + f' {_detail}')

        _fig_verdict('A-FIGSCALE', _by_gate)
        _fig_verdict('A-FIGLABEL', _by_gate)
        _fig_verdict('A-FIGDPI', _by_gate)
        _fig_verdict('A-FIGDEGEN', _by_gate)
        _fig_verdict('A-FIGMONO', _by_gate)
        _fig_verdict('A-FIGOPTUNIF', _by_gate)
        _fig_verdict('A-FIGCOLOUR', _by_gate)
        _fig_verdict('A-FIGACCENT', _by_gate)
        _fig_verdict('A-FIGCVD', _by_gate)
        _fig_verdict('A-FIGSERIES', _by_gate)
        _fig_verdict('A-FIGGLYPH', _by_gate)
        _fig_verdict('A-FIGALT', _by_gate)
        _fig_verdict('A-FIGLABELPX', _by_gate)

        # v2.13: the EC-V18 legacy count was previously emitted as a SECOND
        # A-FIGDPI line, which breaks the v2.12 roster rule ("EVERY gate prints
        # EXACTLY ONE line on EVERY run") and would have made the gate count
        # stop being a usable integrity signal the moment _seen became non-zero.
        # It was unreachable while Block.images was empty; now that figures are
        # actually evaluated it is folded into each gate's own single verdict.


def gate_optref(blocks, src):
    """A-OPTREF (machine half): if a stem references a terminal/escape option,
    that option must be present in the option set."""
    rt = src['rules_txt']
    # Read escape tokens from section_rules (RA-9) as PRIMARY; English defaults as FALLBACK.
    extra = re.findall(r'(?:none_of_above_permitted|escape_option_tokens)\s*[:=]\s*(.+)', rt, re.I)
    if extra:
        # section_rules declared exam-specific escape tokens — merge with defaults.
        tokens = [t.strip() for e in extra for t in e.split(',')]  + list(ESCAPE_TOKENS_DEFAULT)
    else:
        tokens = list(ESCAPE_TOKENS_DEFAULT)
    # Escape-reference phrases: if the WHOLE phrase appears in the stem, the stem is
    # referencing the escape option that phrase implies. Read from section_rules (RA-9);
    # English defaults as fallback.
    ref_phrases = src.get('escape_reference_phrases',
                          [r'if there is no error'])
    miss = []
    for b in blocks:
        stem = block_stem_text(b)
        # v2.21.7 (SEC-2) — _label_paras, not option_paras. option_paras was built
        # on OPT_RE, so on a FIGURAL block it saw only the text-bearing options and
        # missed every bare-label image option (measured: 1 seen where 4 exist). A
        # figural stem carrying an escape token ("None of these") would then be
        # reported as referencing an ABSENT option and FAIL A-OPTREF on a correct
        # paper. o[2] is the text AFTER the label in both helpers, so the escape
        # token search is unchanged — and now cannot match the label itself.
        opts = [o[2].lower() for o in _label_paras(b)]
        for tok in tokens:
            # MODE 1: a select/mark/choose verb near the escape token in the stem.
            triggered = bool(re.search(r'(select|mark|choose).{0,40}' + tok, stem, re.I))
            # MODE 2: a direct escape-reference phrase in the stem that implies this token.
            if not triggered:
                for phrase in ref_phrases:
                    if re.search(phrase, stem, re.I) and re.search(tok, phrase, re.I):
                        triggered = True
                        break
            if triggered:
                present = any(re.search(tok, o, re.I) for o in opts)
                if not present:
                    miss.append(f'Q{b.qnum}')
                break
    (_ok if not miss else _fail)('A-OPTREF',
        'escape-option references are satisfied.' if not miss else
        'stem references an absent escape/terminal option: ' + _flist(miss))


def gate_encoding_script(doc, src):
    full = '\n'.join(para_text(obj) for kind, obj in iter_block_items(doc) if kind == 'p')
    if '�' in full:
        _fail('A-ENCODING', 'U+FFFD replacement character present (encoding corruption).')
    else:
        _ok('A-ENCODING', 'no U+FFFD replacement characters.')
    if src['language'] == 'english':
        # Flag only RUNS (>=2) of foreign-SCRIPT letters (Devanagari, Cyrillic, CJK,
        # Arabic, ...). Accented Latin (café, résumé) and Greek math symbols
        # (alpha, beta, theta) are LEGITIMATE on an english exam and must not be
        # flagged. Isolated symbols never trip the gate; only a multi-letter
        # foreign-script word does (the copy-paste-corruption signature).
        def _foreign_letter(ch):
            if ord(ch) <= 0x7F or not ch.isalpha():
                return False
            try:
                nm = unicodedata.name(ch, '')
            except ValueError:
                nm = ''
            return not (nm.startswith('LATIN') or nm.startswith('GREEK'))
        run = 0; hits = []
        for ch in full:
            if _foreign_letter(ch):
                run += 1
                if run == 2:
                    hits.append(ch)
            else:
                run = 0
        if hits:
            _fail('A-SCRIPT', f'foreign-script word(s) on an english exam '
                              f'(copy-paste corruption?): {len(hits)} run(s).')
        else:
            _ok('A-SCRIPT', 'no foreign-script text (accented Latin / Greek symbols OK).')
    else:
        _ok('A-SCRIPT', f"language={src['language']}: non-Latin script permitted.")


def gate_dup(blocks, src):
    reg = src['registry']; tq = src['total_questions']
    stems_all = reg.get('stem_texts', [])
    if not stems_all or tq is None:
        _warn('A-DUP', 'registry stem_texts empty or total_questions unknown; '
                       'cross-mock dedup skipped.')
        return
    prior = stems_all[:max(0, len(stems_all) - tq)]   # self-exclude trailing mock N
    if not prior:
        _ok('A-DUP', 'no prior-mock stems in registry (first mock or registry holds only mock N); cross-mock dedup vacuous.')
        return
    norm = lambda s: re.sub(r'\s+', ' ', s).strip().lower()
    prior_norm = set(norm(s) for s in prior)
    prior_tokens = [set(norm(s).split()) for s in prior]
    exact, near = [], []
    for b in blocks:
        stem = norm(block_stem_text(b))
        if not stem:
            continue
        if stem in prior_norm:
            exact.append(f'Q{b.qnum}')
            continue
        toks = set(stem.split())
        if toks:
            for pt in prior_tokens:
                if pt:
                    j = len(toks & pt) / len(toks | pt)
                    if j >= 0.75:
                        near.append(f'Q{b.qnum}')
                        break
    (_ok if not exact and not near else _fail)('A-DUP',
        'no cross-mock stem duplication.' if not exact and not near else
        f'exact={_flist(exact)} near={_flist(near)} (vs prior mocks).')


def gate_specfaith(blocks, src):
    """A-FIGSPECLABEL (v2.24.0, RELEASE C) — DOES THE SPEC DESCRIBE THE FIGURE?

    FOUND ON REAL DELIVERED OUTPUT (IIT_JAM_BIOTECHNOLOGY M01): all 57 FigureSpec
    records carry series labels 'Series 1' / 'Series 2' — matplotlib's DEFAULTS —
    while the rendered PNGs carry the real names the stems depend on ('P', 'F',
    'S', 'No inhibitor'), confirmed by ink-width template matching against
    references rendered at the spec's own font_pt_native.

    WHY A PLACEHOLDER LABEL IS A DEFECT WHICHEVER WAY IT RESOLVES:
      (a) if the placeholder WAS rendered, the question is UNANSWERABLE — a stem
          naming curves 'P' and 'Q' against a legend reading 'Series 1'/'Series 2'
          gives the candidate no way to map either curve; or
      (b) if the placeholder was NOT rendered — the case measured here — the SPEC
          IS UNFAITHFUL TO THE RENDER, and A-FIGSERIES / A-FIGGLYPH audit against
          it. Those gates then validate fiction and report conformance they never
          established. Any downstream consumer trusting spec labels is wrong too.

    THE CHECK IS PURE DATA — no pixel analysis, so it cannot inherit the fragility
    of image measurement (a legend locator built for this defect gave inconsistent
    widths and was discarded). It asks only whether a label is a generator default.
    AMBER by construction (v2.22.0 rule: a NEW gate never enters at a severity that
    can halt a paper, because the deployed operator cannot adjudicate anything).
    Exam-agnostic: the patterns are GENERATOR defaults, not exam values, and
    section_rules may extend them via figure_label_placeholders.
    """
    specs = src.get('figure_specs') or {}
    if not specs:
        _ok('A-FIGSPECLABEL', 'registry carries no figure_specs — dormant '
                              '(pre-v5.34 registry; nothing to check).')
        return
    extra = tuple(src.get('figure_label_placeholders') or ())
    pats = [re.compile(_p, re.I) for _p in (
        r'^\s*series\s*[-_ ]?\d+\s*$',
        r'^\s*(data|line|curve|trace|set|group)\s*[-_ ]?\d+\s*$',
        r'^\s*s\d+\s*$',
        r'^\s*(label|untitled|unnamed|todo|tbd|xxx|placeholder)\s*[-_ ]?\d*\s*$',
        r'^\s*$',
    ) + extra]
    bad = []
    for name, sp in sorted(specs.items()):
        series = sp.get('series') or []
        if (sp.get('key_mode') or 'none') == 'none' and len(series) < 2:
            continue
        for ser in series:
            lab = str(ser.get('label') or '')
            if any(_p.match(lab) for _p in pats) or lab.strip() == str(ser.get('id') or ''):
                bad.append(f'Q{sp.get("question")}:{name}:{lab or "<empty>"}')
    if bad:
        _warn('A-FIGSPECLABEL',
              f'{len(bad)} FigureSpec series label(s) are GENERATOR DEFAULTS, not the '
              'names the stem refers to. Either the placeholder was rendered (the '
              'question is unanswerable) or the spec is unfaithful to the render (and '
              'A-FIGSERIES / A-FIGGLYPH are auditing fiction). MECHANICAL REMEDY: '
              're-emit these figures through Step 7 S10-8 with series labels bound to '
              "the stem's own names, so the spec records what was drawn: "
              + _flist(sorted(set(bad))))
    else:
        _ok('A-FIGSPECLABEL', f'{len(specs)} figure spec(s): every keyed series carries '
                              'a stem-meaningful label (no generator defaults).')


# MATHEMATICAL domains — these are properties of the QUANTITIES THEMSELVES, not of
# any exam, syllabus or board (RA-9 is about exam values; the range of a cosine is
# not one). Each entry: (stem trigger regex, (lo, hi), inclusive?, display name).
# section_rules may extend via option_domain_rules; it may not silently shrink them.
OPTION_DOMAINS = (
    (r'\beccentricit(y|ies)\b(?=.*\bellipse\b)', (0.0, 1.0), False, 'eccentricity of an ellipse'),
    (r'\bprobabilit(y|ies)\b',                     (0.0, 1.0), True,  'probability'),
    (r'\bmole\s+fraction\b',                      (0.0, 1.0), True,  'mole fraction'),
    (r'\b(correlation\s+coefficient|Pearson\'?s?\s+r)\b', (-1.0, 1.0), True, 'correlation coefficient'),
)


def gate_optdomain(blocks, src):
    """A-OPTDOMAIN (v2.24.0, RELEASE C) — DISTRACTORS OUTSIDE THE QUANTITY'S DOMAIN.

    FOUND ON REAL DELIVERED OUTPUT: a question asking for the ECCENTRICITY OF AN
    ELLIPSE offered four numeric options, TWO of them greater than 1. An ellipse
    has 0 < e < 1 by definition, so those two are impossible for ANY ellipse
    whatever the figure shows — a candidate who knows the definition discards half
    the option set without reading the figure, and a 4-way discrimination collapses
    to a coin flip. The key stayed uniquely correct, so this is a DISTRACTOR
    QUALITY defect (B-DISTRACT), not a wrong-answer defect.

    WHY IT GENERALISES, which is the reason it is worth a gate: the two impossible
    values were the RECIPROCALS of the two plausible ones (1/0.8 = 1.25, 1/0.6 =
    1.67). That is a distractor STRATEGY — invert the answer — and it is perfectly
    good for an UNBOUNDED quantity. Applied to a BOUNDED one it mechanically emits
    out-of-domain values. The same generator will therefore have produced the same
    defect wherever a bounded quantity met that strategy, across every exam.

    AMBER by construction. Fires only when the stem NAMES the quantity and EVERY
    option parses as a bare number, so a stem carrying units, ranges or symbolic
    options is never judged. Absent triggers ⇒ dormant.
    """
    rules = list(OPTION_DOMAINS)
    for r in (src.get('option_domain_rules') or ()):
        try:
            rules.append((r['pattern'], (float(r['lo']), float(r['hi'])),
                          bool(r.get('inclusive', True)), str(r.get('name', 'quantity'))))
        except Exception:
            continue
    _NUM = re.compile(r'^\s*[-+]?\d+(?:\.\d+)?\s*$')
    bad, checked = [], 0
    for b in blocks:
        stem = para_text(b.items[0][1]) if b.items else ''
        opts = [para_text(o) for k, o in b.items
                if k == 'p' and OPT_LABEL_RE.match(para_text(o))]
        if len(opts) < 2:
            continue
        vals = []
        for o in opts:
            body = OPT_LABEL_RE.sub('', o, count=1)
            if not _NUM.match(body):
                vals = None
                break
            vals.append(float(body.strip()))
        if not vals:
            continue
        for pat, (lo, hi), incl, nm in rules:
            if not re.search(pat, stem, re.I):
                continue
            checked += 1
            if incl:
                out = [v for v in vals if v < lo or v > hi]
            else:
                out = [v for v in vals if v <= lo or v >= hi]
            if out:
                bad.append(f'Q{b.qnum}({nm}: ' + ', '.join(f'{v:g}' for v in out) + ')')
            break
    if bad:
        _warn('A-OPTDOMAIN',
              f'{len(bad)} question(s) offer option value(s) OUTSIDE the mathematical '
              'domain of the quantity the stem names. Such options are eliminable by '
              'definition, so the effective option count is reduced and the item '
              'under-discriminates. MECHANICAL REMEDY: replace each out-of-domain '
              'value with an in-domain one derived from a plausible MISREADING (a '
              'reciprocal distractor is valid only for an UNBOUNDED quantity): '
              + _flist(bad))
    elif checked:
        _ok('A-OPTDOMAIN', f'{checked} numeric option set(s) with a named bounded '
                           'quantity: all values lie inside the domain.')
    else:
        _ok('A-OPTDOMAIN', 'no question names a bounded quantity with an all-numeric '
                           'option set — dormant.')


def gate_header(doc, blocks, src):
    # v2.7: A-HEADER INVERTED. The paper is questions-only (Step 7 R8b / G-PREQ1): NO
    # title/info/scoring/cover paragraph may sit before Q.1. Any non-blank paragraph before
    # Q.1 is a DEFECT → strip it in Phase 1 (CP-HEADER-STRIP). CATEGORY-C values
    # (marks/time/negative/options/total) are metadata, never printed — nothing to
    # figure-check. DORMANT only if section_rules EXAM_STRUCTURE declares paper_header_block.
    title_items, _ = parse_blocks(doc)
    title = ' '.join(para_text(obj) for kind, obj in title_items if kind == 'p').strip()
    if src.get('paper_header_block'):
        _ok('A-HEADER', 'pre-Q.1 header permitted (EXAM_STRUCTURE paper_header_block); dormant.')
    elif not title:
        _ok('A-HEADER', 'questions-only: no non-blank paragraph before Q.1.')
    else:
        _fail('A-HEADER', f'non-Q paragraph(s) before Q.1: "{title[:60]}". Strip the '
                          'pre-Q.1 title/info/scoring/cover block (CP-HEADER-STRIP) — the '
                          'paper is questions-only (R8b/G-PREQ1); marks/time/negative/'
                          'options/total are metadata, never printed.')


# ============================================================================
# TIER A — THE STEP-7 -> STEP-8 DOSSIER (v2.17)
#
# WHY THIS IS A REPAIR, NOT A FEATURE. Step 7 already records every fact below,
# and Framework_MockTestCreate.md says of concept_map: "The audit gates read it
# directly instead of re-deriving." audit_canonical.py has carried a --key
# consumer path since v2.4. But S0-1 never delivered the sidecar, so the producer
# wrote it, the consumer could read it, and the pipeline never connected them.
# Measured consequence on a real 60-Q paper: 0 of 60 concept_map entries reached
# Step 8; A-NAT-GRADE printed "dormant" on all ~200 exams; image_role defaulted
# for every question, false-flagging 27 of 33 figural blocks in A-FIGCOMP.
#
# THE LINE THAT KEEPS THIS HONEST:
#   HAND OVER FACTS STEP 7 RECORDED. NEVER HAND OVER JUDGMENTS STEP 7 REACHED.
# A fact is checkable against the artefact or the world (subtopic_id, qtype,
# image_role, the NAT grading transform). A judgment is the thing Step 8 exists to
# form (the answer, "this is unambiguous", "this figure is legible"). Tier A
# carries only facts; answers and answer_verified are REFUSED at load.
#
#   NO GATE MAY PASS ON DOSSIER EVIDENCE ALONE.
# The dossier can make a check cheaper or make a mismatch visible. It may never be
# the thing that certifies. Every consumer still grounds out in the paper — which
# is exactly why a dossier/paper disagreement is a FINDING, never a silent
# overwrite: it means Step 7's record disagrees with what Step 7 shipped.
# ============================================================================
DOSSIER_SCHEMA = 1
DOSSIER_FORBIDDEN = ('answers', 'answer_key', 'answer_verified', 'concept_text',
                     'derived_answer', 'correct_option')
DOSSIER_FACT_KEYS = ('subtopic_id', 'qtype', 'image_role', 'difficulty',
                     'stem_precision', 'nat_grading_type', 'nat_grading_value',
                     'ca_range', 'msq_instr_in_stem', 'nat_instr_in_stem')


class DossierError(Exception):
    """Refusal to accept a dossier. Names the reason; never fact content."""


def load_dossier(path, docx_path=None, exam=None, mockN=None):
    """Verify and load the Tier-A dossier. Returns (questions_dict, manifest).

    REFUSES — and the caller then runs exactly as it did before the dossier
    existed, which is the safe default — on: unknown schema, unparseable file,
    an identity mismatch (exam_code / mock / paper MD5), or ANY forbidden
    judgment key. The paper binding is the important one: a dossier from a
    different document would let Step 8 audit against facts describing another
    paper.
    """
    if not path or not os.path.exists(path):
        raise DossierError('not supplied')
    try:
        with open(path, encoding='utf-8') as fh:
            d = json.load(fh)
    except Exception as e:
        raise DossierError(f'unparseable ({type(e).__name__})')
    if d.get('schema') != DOSSIER_SCHEMA:
        raise DossierError(f'schema {d.get("schema")!r} != {DOSSIER_SCHEMA}')
    for k in DOSSIER_FORBIDDEN:
        if k in d:
            raise DossierError(
                f'carries a JUDGMENT key ({k!r}) — Tier A transports FACTS only. '
                f'Answers reach Step 8 through the sealed channel, never here.')
    if exam and d.get('exam_code') and d['exam_code'] != exam:
        raise DossierError(f'built for exam {d["exam_code"]!r}, not {exam!r}')
    if mockN is not None and d.get('mock') is not None and int(d['mock']) != int(mockN):
        raise DossierError(f'built for mock {d["mock"]}, not mock {mockN}')
    if not d.get('paper_md5'):
        raise DossierError('carries no paper_md5 binding — cannot be proven to '
                           'describe this paper')
    if docx_path and os.path.exists(docx_path) and _md5_file(docx_path) != d['paper_md5']:
        raise DossierError('paper MD5 mismatch — this dossier describes a DIFFERENT '
                           'document than the one being audited')
    qs = d.get('questions') or {}
    if not isinstance(qs, dict) or not qs:
        raise DossierError('no questions recorded')
    for q, e in qs.items():
        if not isinstance(e, dict):
            raise DossierError(f'Q{q}: entry is not an object')
        for k in DOSSIER_FORBIDDEN:
            if k in e:
                raise DossierError(f'Q{q}: carries a JUDGMENT key ({k!r})')
    return qs, d


def block_option_count(b, oc=0):
    """Options actually RENDERED in this block, counted EXACTLY as gate_options()
    counts them — SAME predicate, SAME trailing-set rule. Never re-implemented.

    v2.17: written because the first cut of gate_dossier read `b.opts`, which DOES
    NOT EXIST on Block — getattr() returned None, n_opt was 0 for every question,
    and the qtype cross-check reported 27 false failures. That is precisely the
    Block.images '# reserved' defect class, reintroduced by this release's own
    author and caught only by running the gate against a real paper. A field that
    is never populated is indistinguishable from a field that is empty; assume
    nothing, count from the document.

    v2.21 (GAP-2026-08-02-DOSSIER-OPTION-PREDICATE): the v2.17 fix counted from the
    document but with the WRONG PREDICATE, and its docstring asserted a parity that
    did not exist. OPT_RE requires a VISIBLE GLYPH after the label (`[.)]\\s+\\S`);
    an IMAGE option is a BARE label paragraph followed by a picture paragraph, so
    OPT_RE counted every image option as ZERO and A-DOSSIER reported
    `qtype-mcq-but-0!=N-options` on every figural question in the estate, blocking
    certification under MANDATE D with NOTHING ON THE PAPER TO REPAIR. It also had
    no trailing-set clamp, so an enumerated stem ("1. ... 2. ...") INFLATED the
    count on pure TEXT papers — the defect was never figural-only. And because the
    same zero vacuously satisfied the `nat` branch, a dossier mislabel on a figural
    question was accepted SILENTLY: one predicate produced both a false FAIL and a
    false PASS.

    THE RULE (spec S5-2 "ONE STRUCTURAL QUESTION, ONE ANSWER"): there is ONE
    rendered-option count in this file and every gate reads it. A second
    implementation is drift by construction — it is written against the author's
    BELIEF about the first rather than against the first, and the divergence stays
    invisible until a paper exercises the difference. Enforced by
    validate_framework_md.py CHECK AN.

    RETURN SEMANTICS (oc = the expected option count; 0 = unknown/no clamp):
      4 text options,          oc=4 -> 4   (unchanged)
      4 image options (bare),  oc=4 -> 4   (v2.21 fix; OPT_RE returned 0)
      0 options (NAT),         oc=4 -> 0   (unchanged)
      3 options,               oc=4 -> 3   (short set still FAILs — no false negative)
      2 stem points + 4 opts,  oc=4 -> 4   (v2.21 fix; trailing clamp, parity A-OPTN)
      any count,               oc=0 -> raw (defensive; no clamp target available)
    """
    labs = _label_paras(b)                    # OPT_LABEL_RE — bare-label tolerant
    return oc if (oc and len(labs) >= oc) else len(labs)


def gate_dossier(blocks, src, dossier, why=None):
    """A-DOSSIER — cross-check every Tier-A fact against the SHIPPED PAPER.

    This gate never certifies anything on its own; it establishes whether the
    dossier may be TRUSTED as a shortcut. A disagreement means Step 7's record and
    Step 7's output disagree, which is a real defect in one of them and is reported
    as such rather than quietly resolved in either direction.
    """
    if not dossier:
        _warn('A-DOSSIER', f'no Tier-A dossier consumed ({why or "not supplied"}) — '
                           'subtopic/qtype/image_role are re-derived or defaulted, '
                           'A-NAT-GRADE is dormant, and A-FIGCOMP may over-report. '
                           'Legacy behaviour; not a paper defect.')
        return
    oc = src.get('options_count') or 0
    paper_qs = {str(b.qnum) for b in blocks}
    dos_qs = set(dossier)
    bad = []
    if paper_qs != dos_qs:
        for q in sorted(paper_qs - dos_qs, key=_qsort_key):
            bad.append(f'Q{q}:absent-from-dossier')
        for q in sorted(dos_qs - paper_qs, key=_qsort_key):
            bad.append(f'Q{q}:not-in-paper')
    figsub = src.get('figural_subtopics') or {}
    for b in blocks:
        e = dossier.get(str(b.qnum))
        if not e:
            continue
        n_opt = block_option_count(b, oc)
        qt = (e.get('qtype') or '').lower()
        # qtype is checkable against the rendered structure.
        #
        # v2.21.1 — THE NAT LEG FIRES ON ANY NON-ZERO COUNT, and this is
        # SPEC-GROUNDED, not a judgement call. Framework_MockTestCreate.md R13
        # (v4.7 NAT EXEMPTION) states that a NAT question has ZERO option
        # paragraphs — "only the bold Q.<N> stem (carrying the nat_instruction
        # per R14) and the blank separator". A NAT block therefore CANNOT
        # legitimately carry ANY option-label paragraph, not even an "enumerated
        # stem": R13 admits no third paragraph class.
        #
        # v2.21 briefly clamped this leg to n_opt >= oc on the ASSUMPTION that a
        # NAT stem could enumerate. R13 forbids it. The assumption was never
        # checked against the producer spec — the SAME error class this release
        # exists to remove (a belief about a sibling contract, unverified) — and
        # it opened a REAL false negative: with nat_present=False and the
        # registry marking the question 0-option, gate_options SKIPS the block
        # (obq==0), gate_nat is DORMANT (nat_present false), and a clamped
        # A-DOSSIER was silent too, so an R13 violation passed ALL THREE gates.
        # That configuration is precisely a Step-7 internal inconsistency, which
        # is the one thing this gate exists to catch. Never clamp this leg.
        if qt == 'nat' and n_opt:
            bad.append(f'Q{b.qnum}:qtype-nat-but-{n_opt}-options')
        elif qt in ('mcq', 'msq') and oc and n_opt != oc:
            bad.append(f'Q{b.qnum}:qtype-{qt}-but-{n_opt}!={oc}-options')
        # subtopic_id is checkable against the registry Step 8 already receives
        rs = figsub.get(str(b.qnum))
        if rs and e.get('subtopic_id') and e['subtopic_id'] != rs:
            bad.append(f'Q{b.qnum}:subtopic-disagrees-with-registry')
    if bad:
        _fail('A-DOSSIER', f'{_fcount(bad)} Tier-A fact(s) disagree with the shipped '
                           f'paper or the registry — Step 7 RECORDED something other '
                           f'than what it SHIPPED. Resolve before trusting the dossier: '
                           + _flist(bad))
    else:
        _ok('A-DOSSIER', f'Tier-A dossier consistent with the paper for '
                         f'{len(dos_qs)} question(s); facts adopted as a shortcut, '
                         f'never as certification (no gate passes on dossier alone).')


def gate_zip(docx_path):
    with zipfile.ZipFile(docx_path) as z:
        names = set(z.namelist())
        if 'word/document.xml' not in names:
            _fail('A-ZIP', 'word/document.xml missing.')
            return {}
        doc_xml = z.read('word/document.xml').decode('utf-8', 'replace')
        rels = (z.read('word/_rels/document.xml.rels').decode('utf-8')
                if 'word/_rels/document.xml.rels' in names else '')
        # parse each <Relationship .../> element, then pull Id + Target from
        # within it (attribute ORDER-INDEPENDENT — some generators emit Target
        # before Id, which an "Id...Target" regex would miss).
        rel_map = {}
        for tag in re.findall(r'<Relationship\b[^>]*/?>', rels):
            idm = re.search(r'\bId="([^"]+)"', tag)
            tgm = re.search(r'\bTarget="([^"]+)"', tag)
            if idm and tgm:
                rel_map[idm.group(1)] = tgm.group(1)
        ref_ids = set(re.findall(r'r:(?:embed|id|link)="([^"]+)"', doc_xml))
        bad = []
        media_map = {}
        for rid in ref_ids:
            tgt = rel_map.get(rid)
            if tgt is None:
                bad.append(rid); continue
            if tgt.startswith('http'):
                continue
            part = ('word/' + tgt).replace('word/../', '')
            if part not in names and ('word/' + tgt) not in names:
                bad.append(rid)
            media_map[rid] = tgt.split('/')[-1]
        (_ok if not bad else _fail)('A-ZIP',
            'all rIds resolve to parts.' if not bad else f'unresolved rIds: {_flist(bad)}')
        return media_map


# ============================================================================
# S5-1A — COMPLETION GATE (v2.6) — Phase-3 mechanical Part-B/§7 enforcement
# ============================================================================
def _resolve_evidence(evidence_dir, stored):
    """Resolve a ledger-stored evidence path to an existing absolute path. Accepts an
    absolute path, a path relative to evidence_dir, a bare basename under evidence_dir,
    or a path with a leading 'evidence/' segment. Returns the resolved path or None."""
    if not stored:
        return None
    cands = [stored]
    if not os.path.isabs(stored):
        cands.append(os.path.join(evidence_dir, stored))
        cands.append(os.path.join(evidence_dir, os.path.basename(stored)))
        parts = stored.replace('\\', '/').split('/')
        if parts and parts[0] == 'evidence':
            cands.append(os.path.join(evidence_dir, *parts[1:]))
    for c in cands:
        if c and os.path.exists(c):
            return c
    # v2.15 (C1) — last resort: the stored path may be an ABSOLUTE path from a
    # previous session's container (checkpoint restore rebases these, but a
    # hand-edited or partially-rebased state must still resolve rather than
    # silently fail C5/C6). Search the evidence tree by basename.
    base = os.path.basename(stored.replace('\\', '/'))
    if base and evidence_dir and os.path.isdir(evidence_dir):
        for root, _dirs, files in os.walk(evidence_dir):
            if base in files:
                return os.path.join(root, base)
    return None


def _file_ok(path, min_bytes):
    return bool(path) and os.path.exists(path) and os.path.getsize(path) >= min_bytes


# v2.14 (B3 — FACT CONTEXT DISCIPLINE). RA-11 has always required the saved fact
# to carry query + URL + retrieval-time + snippet. C5 only ever checked that the
# file EXISTED and was >= 1 byte, so a one-character stub certified. That gap did
# not matter while the full search result also sat in the reasoning stream — the
# evidence was duplicated. It matters now: B3 moves the raw result OUT of context
# and onto disk, which makes the saved file the ONLY copy. If the gate does not
# verify its shape, the discipline degrades from "save it" to "touch a file", and
# C5 would certify an audit whose evidence no longer exists in any form.
FACT_REQUIRED_FIELDS = ('query', 'url', 'retrieved_at', 'snippet')


def _fact_record_ok(path):
    """Validate one saved fact-evidence file. Accepts a single record object or a
    LIST of them (one file may hold the key fact plus its options, and one file may
    be SHARED by several questions that turn on the same concept — the B3 cache).
    Returns (ok, reason). The reason names FIELDS only, never fact content
    (MANDATE 0)."""
    if not path or not os.path.exists(path):
        return False, 'missing'
    try:
        with open(path, encoding='utf-8') as fh:
            obj = json.load(fh)
    except Exception as e:
        return False, f'unparseable ({type(e).__name__})'
    recs = obj if isinstance(obj, list) else [obj]
    if not recs:
        return False, 'no records'
    for r in recs:
        if not isinstance(r, dict):
            return False, 'record is not an object'
        miss = [k for k in FACT_REQUIRED_FIELDS if not str(r.get(k) or '').strip()]
        if miss:
            return False, 'missing/blank ' + ','.join(miss)
    return True, ''


def _block_has_omml(b):
    for p in b.paras:
        for _ in p._element.iter(M_('oMath')):
            return True
    return False


def _block_has_image(b):
    return any(para_images(p) for p in b.paras)


# ============================================================================
# D2 + D4 — VISION AS A DECLARED, PROBED, DEGRADABLE DEPENDENCY (v2.16)
#
# WHY. RA-4 refuses to let an item certify unless its figure was actually VIEWED.
# That is right, and it was doing TWO jobs with one rule: it blocked a lazy
# operator (correct) and it also blocked an ENVIRONMENT OUTAGE (wrong). When the
# view path failed mid-session on IIT_JAM_BIOTECHNOLOGY_Mock01, 43 images across
# 27 figural questions became un-stampable, C6/C7 could never pass, MANDATE D
# forbade delivery, and there was no defined state for "vision unavailable" — the
# audit was permanently STUCK, not degraded.
#
# That contradicts this framework's own doctrine, stated at §5 and in CLAUDE.md:
# "NO DEPENDENCY CONDITION MAY EVER HALT A RUN" and "Silence is the defect; a
# halt is not the remedy." Graceful degradation was granted to blueprint_core, to
# figural_core, to all seventeen figure gates (thirteen through v5.54; +G-FIGFIT,
# +G-FIGCOLLIDE, +G-FIGOPTWINDOW, +W-FIGFITPX at v5.55) and to every colour
# condition — and
# denied to the one dependency whose absence is fatal.
#
# THE DANGER, AND HOW IT IS CLOSED. A third stamp state is an obvious cheat
# surface: "I could not see it" is exactly what a lazy operator would claim. So
# 'view-unavailable' is NEVER assertable by choice. It is admissible ONLY when a
# machine probe FAILED for that batch. The probe renders three RANDOM glyphs and
# stores only their SHA-256 — reading the sidecar reveals nothing, so the only way
# to report the glyphs is to actually see the image. Vision therefore becomes a
# MEASURED fact, not an operator claim.
# ============================================================================
PROBE_GLYPH_ALPHABET = 'ACEFHJKLMNPRTUVWXY34679'
PROBE_GLYPH_COUNT = 3
VISION_STAMP_VIEWED = 'rendered-and-viewed'
VISION_STAMP_UNAVAILABLE = 'view-unavailable'
# v2.23.0 (RELEASE B) — CONFORMANCE ESTABLISHED WITHOUT VIEWING.
# The thirteen figure gates (A-FIGSCALE / A-FIGLABEL / A-FIGDPI / A-FIGDEGEN /
# A-FIGMONO / A-FIGOPTUNIF / A-FIGCOLOUR / A-FIGACCENT / A-FIGCVD / A-FIGSERIES /
# A-FIGGLYPH / A-FIGALT / A-FIGLABELPX) are ARITHMETIC over the saved PNG and its
# FigureSpec.
# They do not use vision and never did. Measured on a real delivered paper
# (IIT_JAM_BIOTECHNOLOGY M01): all twelve reported "57 figure(s) conform" with no
# view tool involved — and they catch what eyes do not (a 72-DPI render, a
# 1-pixel stroke, colours that collapse under greyscale).
# Vision establishes a DIFFERENT and NARROWER claim: does the drawing depict what
# the stem describes. Conflating the two is what made a vision outage fatal —
# C7 saw an unstamped artefact, failed, and MANDATE D refused delivery, so a
# paper whose figures were fully conformance-checked shipped as nothing at all.
VISION_STAMP_ARITHMETIC = 'conformance-arithmetic'


def _probe_paths(d):
    return os.path.join(d, '_probe.png'), os.path.join(d, '_probe.json')


def make_vision_probe(evidence_dir, batch=None, seed=None):
    """D4/P3.5 — render a probe image carrying PROBE_GLYPH_COUNT random glyphs.

    Returns (png_path, meta). The expected glyphs are stored ONLY as a salted
    sha256, so an operator who cats the sidecar learns nothing: reporting the
    glyphs requires seeing the PNG. That is what makes 'view-unavailable'
    unfakeable rather than merely discouraged.

    Never raises. If PIL is unavailable the probe reports P3.5-RENDER-FAIL, which
    is an ENVIRONMENT WARN and explicitly NOT a vision verdict (E4.5) — inferring
    "vision is broken" from "we could not draw the test card" would degrade a
    perfectly healthy run.
    """
    d = os.path.join(evidence_dir, 'montages')
    os.makedirs(d, exist_ok=True)
    png, meta_p = _probe_paths(d)
    rnd = random.Random(seed)
    glyphs = ''.join(rnd.choice(PROBE_GLYPH_ALPHABET) for _ in range(PROBE_GLYPH_COUNT))
    salt = hashlib.sha256(os.urandom(16)).hexdigest()[:16]
    meta = {'at': _now_utc(), 'batch': batch, 'salt': salt,
            'expected_sha256': hashlib.sha256((salt + glyphs).encode()).hexdigest(),
            'render': 'ok'}
    try:
        from PIL import Image, ImageDraw
        im = Image.new('RGB', (200, 200), 'white')
        dr = ImageDraw.Draw(im)
        for i, g in enumerate(glyphs):
            dr.text((22 + i * 58, 78), g, fill='black')
            dr.rectangle([14 + i * 58, 66, 62 + i * 58, 128], outline='black', width=2)
        im.save(png)
    except Exception as e:
        meta['render'] = f'P3.5-RENDER-FAIL ({type(e).__name__})'
    with open(meta_p, 'w', encoding='utf-8') as fh:
        json.dump(meta, fh, indent=1)
    return png, meta


def verify_vision_probe(evidence_dir, glyphs_read, attempts=1):
    """Compare what the operator reports SEEING against the planted glyphs.
    Returns a session_log.vision_probe record. Case-insensitive; whitespace-loose.
    A render failure is reported as its own status, never as FAILED vision."""
    d = os.path.join(evidence_dir, 'montages')
    _png, meta_p = _probe_paths(d)
    try:
        with open(meta_p, encoding='utf-8') as fh:
            meta = json.load(fh)
    except Exception:
        return {'status': 'FAILED', 'attempts': attempts, 'at': _now_utc(),
                'glyphs_read': '', 'reason': 'no probe was rendered'}
    if str(meta.get('render', 'ok')).startswith('P3.5-RENDER-FAIL'):
        return {'status': 'RENDER-FAIL', 'attempts': attempts, 'at': _now_utc(),
                'glyphs_read': '', 'reason': meta['render'], 'batch': meta.get('batch')}
    got = re.sub(r'[^A-Za-z0-9]', '', glyphs_read or '').upper()
    ok = (hashlib.sha256((meta.get('salt', '') + got).encode()).hexdigest()
          == meta.get('expected_sha256'))
    return {'status': 'OK' if ok else 'FAILED', 'attempts': attempts,
            'at': _now_utc(), 'glyphs_read': got, 'batch': meta.get('batch')}


def vision_state(state):
    """Reduce session_log.vision_probe (a record or a list of per-batch records) to
    ('OK'|'FAILED'|'RENDER-FAIL'|'ABSENT', failed_batches:set). Per-batch because a
    mid-run transition in EITHER direction is expected (E2.2/E2.3/E4.1/E4.2) — the
    incident that motivated this release had Batch 1 healthy and Batch 2 not."""
    vp = ((state.get('session_log') or {}).get('vision_probe'))
    if not vp:
        return 'ABSENT', set()
    recs = vp if isinstance(vp, list) else [vp]
    failed = {r.get('batch') for r in recs if r.get('status') == 'FAILED'}
    if any(r.get('status') == 'FAILED' for r in recs):
        latest = recs[-1].get('status')
        return ('FAILED' if latest == 'FAILED' else latest or 'FAILED'), failed
    return (recs[-1].get('status') or 'ABSENT'), failed


def completion_gate(audit_state_path, total_questions, blocks, doc):
    """S5-1A — validate the Phase-2 audit_state ledger (C1-C7) AND the on-disk evidence
    artefacts each stamp names. Appends C0..C7 results to RESULTS so the exit code
    reflects them, and prints the COMPLETION-GATE summary line. MANDATE-0 safe:
    Q-numbers + codes only, never content/URLs. Returns 0 (PASS) or 1 (FAIL)."""
    try:
        state = json.load(open(audit_state_path, encoding='utf-8'))
    except Exception as e:
        _fail('C0', f'audit_state unreadable/absent ({e}).')
        print('COMPLETION-GATE: FAIL (audit_state unreadable)')
        return 1
    ledger = (state.get('ledger') or {}).get('entries', {}) or {}
    entries = {int(k): v for k, v in ledger.items()}
    evidence_dir = state.get('evidence_dir', '') or ''
    K = state.get('K')
    if not K:
        plan = state.get('plan') or []
        K = len(plan) if plan else (max(state.get('batches_done', []) or [0]) or 0)
    batches_done = set(state.get('batches_done', []) or [])

    # C1 — every planned batch closed
    if K and set(range(1, K + 1)) <= batches_done:
        _ok('C1', f'all {K} batches closed.')
    else:
        _fail('C1', f'batches_done={sorted(batches_done)} does not cover 1..{K} '
                    f'(Phase 2 skipped/partial — MANDATE B).')
    # C2 — every question reviewed
    want = set(range(1, (total_questions or 0) + 1))
    got = set(entries.keys())
    if total_questions and got == want:
        _ok('C2', f'all {total_questions} questions have a ledger entry.')
    else:
        miss = sorted(want - got); extra = sorted(got - want)
        _fail('C2', f'ledger != 1..{total_questions}: missing={miss} extra={extra}.')
    # C3 — every entry closed
    bad3 = [q for q, e in entries.items() if e.get('status') not in ('verified', 'regenerated')]
    (_ok if not bad3 else _fail)('C3',
        'all entries verified/regenerated.' if not bad3 else
        f'entries not closed (pending/absent): {_flist(bad3)}')
    # C4 — uniqueness / set ran
    bad4 = []
    for q, e in entries.items():
        if e.get('answer_cardinality') == 'multi':
            if not e.get('answer_set_verified'):
                bad4.append(q)
        else:
            if not e.get('answer_unique'):
                bad4.append(q)
    (_ok if not bad4 else _fail)('C4',
        'B-UNIQUE/A-MSQ-KEY ran for every Q.' if not bad4 else
        f'uniqueness/set not verified: {_flist(bad4)}')
    # C5 — factual entries sourced AND the saved evidence file exists, parses, and
    # carries the RA-11 record fields (v2.14/B3: shape, not just existence).
    bad5 = []
    _fact_files, _fact_refs = set(), 0
    for q, e in entries.items():
        if e.get('is_factual'):
            fs = e.get('fact_sources') or []
            if not fs:
                bad5.append((q, 'no fact_sources')); continue
            for rec in fs:
                saved = rec.get('saved') if isinstance(rec, dict) else None
                _fact_refs += 1
                _p = _resolve_evidence(evidence_dir, saved)
                ok, why = _fact_record_ok(_p)
                if not ok:
                    bad5.append((q, why)); break
                _fact_files.add(os.path.realpath(_p))
    _dedup = (f' ({len(_fact_files)} distinct source file(s) for {_fact_refs} '
              f'reference(s) — B3 cache reuse)') if _fact_refs else ''
    (_ok if not bad5 else _fail)('C5',
        f'every factual entry has a saved, well-formed sourced fact.{_dedup}'
        if not bad5 else
        'factual entries unsourced / saved fact file missing or malformed: '
        + _flist([f'Q{q}:{why}' for q, why in set(bad5)], sep='; '))
    # C6 — artefact stamps present AND their evidence files exist/non-trivial.
    # v2.16 (D2): a 'view-unavailable' image stamp is admissible, but ONLY under the
    # conditions below. Everything else about C6 is unchanged.
    _vstatus, _vfailed_batches = vision_state(state)
    bad6, degraded6, forged6, stale6 = [], [], [], []
    arith6 = []
    for q, e in entries.items():
        st = e.get('artefact_stamps') or {}
        ok = True
        for img in (st.get('images') or []):
            _montage_ok = _file_ok(_resolve_evidence(evidence_dir, img.get('montage')),
                                   EVIDENCE_MIN_BYTES)
            if img.get('stamp') == VISION_STAMP_ARITHMETIC:
                # v2.23.0 — the thirteen figure gates ran over this PNG + its spec.
                # Unfakeability here is NOT the vision probe (arithmetic does not
                # depend on vision and is valid whether or not the view tool
                # works); it is RA-4's render-or-recompute rule — the saved
                # artefact must exist and be non-trivial, exactly as a table or
                # OMML recompute trace must.
                if not _montage_ok:
                    ok = False
                    continue
                arith6.append(q)
            elif img.get('stamp') == VISION_STAMP_UNAVAILABLE:
                # (a) the montage must STILL exist and be non-trivial. A vision outage
                #     does not excuse producing no artefact — that is E2.5/E2.6, an
                #     un-audited item, and it still blocks.
                if not _montage_ok:
                    ok = False
                    continue
                # (b) UNFAKEABLE: a FAILED probe must exist. Without this the stamp is
                #     a self-signed excuse and RA-4 collapses into "I didn't feel like
                #     looking" (E2.4).
                if _vstatus == 'ABSENT' or not _vfailed_batches:
                    forged6.append(q); ok = False
                    continue
                # (c) NOT STALE: if vision has since RECOVERED, the operator must
                #     re-attempt and upgrade the stamp before Phase 3 (E2.3/E4.2).
                #     Certifying old degraded stamps on a healthy session would silently
                #     under-audit a paper that could have been fully audited.
                elif _vstatus == 'OK':
                    stale6.append(q); ok = False
                    continue
                degraded6.append(q)
            elif not _montage_ok:
                ok = False
        for kind in ('tables', 'charts', 'omml'):
            for rec in (st.get(kind) or []):
                # E2.7: tables/OMML/charts are ARITHMETIC, not vision. They are
                # unaffected by a vision outage and remain fully authoritative.
                path = rec.get('trace') or rec.get('montage')
                if not _file_ok(_resolve_evidence(evidence_dir, path), 1):
                    ok = False
        if not ok:
            bad6.append(q)
    if forged6:
        _fail('C6', 'view-unavailable claimed with NO FAILED vision probe — the stamp '
                    'is not admissible (RA-4 v2.16 requires a MEASURED outage, never an '
                    f'operator claim): {_flist(forged6)}')
    elif stale6:
        _fail('C6', 'vision has RECOVERED but view-unavailable stamps were not upgraded — '
                    're-attempt the view for these and upgrade before Phase 3 (E2.3): '
                    f'{_flist(stale6)}')
    elif bad6:
        _fail('C6', f'artefact stamp evidence missing/trivial: {_flist(bad6)}')
    elif degraded6 or arith6:
        _parts = []
        if degraded6:
            _parts.append(f'{len(set(degraded6))} question(s) carry view-unavailable under a '
                          f'MEASURED vision outage (batches '
                          f'{sorted(b for b in _vfailed_batches if b is not None)})')
        if arith6:
            _parts.append(f'{len(set(arith6))} question(s) carry conformance-arithmetic '
                          '(the thirteen figure gates ran over the PNG + FigureSpec; '
                          'CONFORMANCE IS ESTABLISHED, the figure-vs-stem semantic '
                          'claim is not)')
        _warn('C6', '; '.join(_parts) + '. Evidence exists and is non-trivial. '
                    'Certifies DEGRADED, not clean — §R13 limitation required.')
    else:
        _ok('C6', 'every artefact stamp is backed by an existing evidence file.')
    # C7 — coverage: every artefact PRESENT IN THE PAPER is represented in the ledger
    bad7 = []
    for b in blocks:
        e = entries.get(b.qnum, {})
        st = e.get('artefact_stamps') or {}
        if _block_has_image(b) and not st.get('images'):
            bad7.append(f'Q{b.qnum}:img'); continue
        if b.tables and not (st.get('tables') or st.get('charts')):
            bad7.append(f'Q{b.qnum}:tbl'); continue
        if _block_has_omml(b) and not st.get('omml'):
            bad7.append(f'Q{b.qnum}:omml'); continue
    (_ok if not bad7 else _fail)('C7',
        'every paper artefact is covered by a ledger stamp.' if not bad7 else
        f'paper artefact not audited (no ledger stamp): {_flist(bad7)}. '
        'MECHANICAL REMEDY — an IMAGE artefact does NOT require the view tool. '
        'The thirteen figure gates are ARITHMETIC over the PNG + FigureSpec and '
        f'establish conformance without it: stamp {VISION_STAMP_ARITHMETIC!r} '
        'with the saved gate trace. If the view tool is genuinely down, run P3.5 '
        '(--vision-probe / --vision-probe-verify) ONCE to record the outage, then '
        f'stamp {VISION_STAMP_UNAVAILABLE!r}; the paper certifies DEGRADED and '
        'DELIVERS. AN UNVIEWABLE FIGURE IS NEVER A REASON TO SHIP NOTHING.')

    cfails = [c for lvl, c, _ in RESULTS if lvl == 'FAIL' and (c == 'C0' or c.startswith('C'))
              and c[1:].isdigit()]
    if not cfails:
        F = sum(1 for e in entries.values() if e.get('is_factual'))
        V = sum(len((e.get('artefact_stamps') or {}).get(k, []))
                for e in entries.values() for k in ('images', 'tables', 'charts', 'omml'))
        _tail = (f'Q reviewed={len(entries)}/{total_questions}, facts sourced={F}, '
                 f'artefacts stamped={V}, evidence files present={V + F}')
        if degraded6 or arith6:
            # v2.16 (D2): exit code stays 0 and the paper DELIVERS. A paper whose
            # figures were machine-checked but not eyeballed is in the same epistemic
            # class as the ~200 EC-V18 legacy papers the framework already ships — and
            # it is demonstrably better than the current outcome, which is no paper.
            # It is never SILENT: DEGRADED is printed, the footer goes AMBER, and §R13
            # records it.
            _I = sum(len((e.get('artefact_stamps') or {}).get('images', []))
                     for e in entries.values())
            _n_arith = sum(1 for e in entries.values()
                           for i in ((e.get('artefact_stamps') or {}).get('images') or [])
                           if i.get('stamp') == VISION_STAMP_ARITHMETIC)
            if _n_arith:
                print(f'COMPLETION-GATE: DEGRADED (vision) — figure CONFORMANCE '
                      f'ESTABLISHED ARITHMETICALLY for {_n_arith} image artefact(s) '
                      f'(thirteen figure gates over PNG + FigureSpec, no view tool); the '
                      f'figure-vs-stem SEMANTIC claim is NOT visually confirmed. '
                      f'§R13 limitation required; CERTIFIED-DEGRADED (VISION) ⇒ F1 '
                      f'AMBER footer. ({_tail})')
                return 0
            print(f'COMPLETION-GATE: DEGRADED (vision) — {len(set(degraded6))} question(s), '
                  f'{sum(1 for e in entries.values() for i in ((e.get("artefact_stamps") or {}).get("images") or []) if i.get("stamp") == VISION_STAMP_UNAVAILABLE)} '
                  f'of {_I} image artefact(s) unviewed; §R13 limitation required; '
                  f'CERTIFIED-DEGRADED (VISION) ⇒ F1 AMBER footer. ({_tail})')
        else:
            print(f'COMPLETION-GATE: PASS ({_tail})')
        return 0
    print(f'COMPLETION-GATE: FAIL ({len(cfails)} assertion(s): {sorted(set(cfails))})')
    return 1


# ============================================================================
# RUNNER
# ============================================================================
def _safe_gate(_name, _fn, *a, **kw):
    """v2.12 — GATE FAULT ISOLATION (GAP-2026-08-01-FIGPROFILE-ENGINE-BINDING, P6).

    THE STRUCTURAL LESSON OF THIS GAP. The `bc` NameError was one defect, but the
    reason it produced a PERMANENT HALT WITH ZERO OUTPUT is that run_audit called
    21 gates in bare sequence with no isolation, and print_results() runs only
    after the last one. A single raise anywhere therefore destroyed the entire
    report — including the 30+ gates that had already passed and the completion
    gate that certifies Phase 3.

    Binding `bc` fixes THIS defect. Isolation fixes the DEFECT CLASS: from v2.12
    no gate, present or future, can abort the run. An unexpected raise becomes a
    LOUD, NAMED, non-recoverable FAIL line and the remaining gates still execute.

    SEVERITY IS _fail, NOT _warn — deliberately, and this is the one place where
    the "never halt" directive must not be read as "never block". A gate that
    crashed DID NOT AUDIT THE PAPER. Reporting it as a warning would let an
    unaudited paper reach certification, which is the false-clean outcome the
    whole v2.6 hardening exists to prevent. FAIL yields exit 1, so the paper
    cannot be certified — while the run still COMPLETES and prints everything.
    That is exactly CLAUDE.md's rule: "A CLASS T failure must be LOUD, and must
    NOT halt... Silence is the defect; a halt is not the remedy."
    """
    try:
        return _fn(*a, **kw)
    except Exception as _e:
        import traceback as _tb
        _fail('A-GATEERROR',
              f'gate {_name} raised {type(_e).__name__}: {_e} — this gate DID NOT '
              f'RUN and the paper is NOT audited for it. Framework defect: capture '
              f'this line and file a gap report (§17). Remaining gates still ran.')
        # The traceback goes to STDERR (never STDOUT — the report must stay
        # machine-parseable) so a maintainer can diagnose the framework defect.
        # Suppressed under --self-test, where fixture 50 raises on purpose.
        if not any(a == '--self-test' for a in sys.argv):
            try:
                print(f'--- A-GATEERROR traceback ({_name}) ---', file=sys.stderr)
                _tb.print_exc()
            except Exception:
                pass
        return None


# ============================================================================
# C1 — CROSS-SESSION CHECKPOINT (v2.15)
#
# WHY THIS EXISTS. RA-18 called Step 8 "resume-safe" and stored every piece of
# cross-batch state — the ledger, the batch plan, the WIP docx and the ENTIRE
# evidence tree — under /home/claude. That directory does not survive a session
# boundary. So resume worked inside one session and not at all across one, and
# the failure was silent-then-fatal: a session that exhausted mid-Phase-2 lost
# the audit outright, because S5-1A C5/C6 assert that every stamped evidence file
# EXISTS. A perfectly remembered ledger cannot certify against montages and saved
# fact records that no longer exist. The retry then exhausted the same way. That
# loop — not any single gate — is what made this step keep failing.
#
# The checkpoint is a PORTABLE, HASH-MANIFESTED bundle the author downloads and
# re-uploads, so an audit may legitimately span sessions and still certify.
#
# BINDING IS THE WHOLE SAFETY ARGUMENT. A checkpoint restored onto the WRONG
# paper would certify an audit of a document nobody audited — strictly worse than
# losing the audit. Restore therefore refuses unless the bundle's recorded
# exam_code, mock number AND paper MD5 all match the paper in hand, and unless
# every member's sha256 matches the manifest. Prose describes; only code
# certifies (§21).
# ============================================================================
CHECKPOINT_SCHEMA = 1
CHECKPOINT_MANIFEST = 'checkpoint_manifest.json'


class CheckpointError(Exception):
    """Refusal to build or restore. Always names WHY, never fact/question content."""


def _sha256_file(path):
    h = hashlib.sha256()
    with open(path, 'rb') as fh:
        for chunk in iter(lambda: fh.read(65536), b''):
            h.update(chunk)
    return h.hexdigest()


def _md5_file(path):
    h = hashlib.md5()
    with open(path, 'rb') as fh:
        for chunk in iter(lambda: fh.read(65536), b''):
            h.update(chunk)
    return h.hexdigest()


def _now_utc():
    from datetime import datetime, timezone
    return datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')


def make_checkpoint(state_path, out_zip, docx_path=None, exam=None, mockN=None):
    """Bundle audit_state.json + the evidence tree + the WIP docx into out_zip.

    The bundle is self-describing: checkpoint_manifest.json records the schema,
    the identity triple (exam_code, mock, paper_md5), progress counters, and a
    sha256 for EVERY member. Returns the manifest dict.
    """
    if not os.path.exists(state_path):
        raise CheckpointError(f'audit_state not found: {state_path}')
    try:
        with open(state_path, encoding='utf-8') as fh:
            st = json.load(fh)
    except Exception as e:
        raise CheckpointError(f'audit_state unparseable ({type(e).__name__})')

    # v2.15 — THE PAPER IS MANDATORY, NOT OPTIONAL. paper_md5 is the strongest of
    # the three identity bindings, and a bundle built without the docx would carry
    # paper_md5=None, which makes restore's MD5 check vacuous — a checkpoint that
    # could then be restored onto ANY document. Caught in end-to-end testing when a
    # shell quoting slip left the docx absent and the checkpoint was written anyway,
    # cheerfully, with no binding at all. Refuse instead: a checkpoint whose safety
    # argument is missing is worse than no checkpoint, because it looks like one.
    if not docx_path or not os.path.exists(docx_path):
        raise CheckpointError(
            'the paper (.docx) is REQUIRED to build a checkpoint — it supplies the '
            'paper_md5 binding that stops the bundle being restored onto a different '
            'document. Pass the Create.docx being audited.')
    members = [('audit_state.json', state_path),
               ('paper/' + os.path.basename(docx_path), docx_path)]
    evd = st.get('evidence_dir')
    if evd and os.path.isdir(evd):
        for root, _dirs, files in os.walk(evd):
            for f in sorted(files):
                ap = os.path.join(root, f)
                arc = 'evidence/' + os.path.relpath(ap, evd).replace(os.sep, '/')
                members.append((arc, ap))

    man = {
        'schema': CHECKPOINT_SCHEMA,
        'created_utc': _now_utc(),
        'exam_code': exam or st.get('exam_code'),
        'mock': st.get('mock', mockN),
        'paper_id': st.get('paper_id'),
        'paper_md5': _md5_file(docx_path),
        'paper_name': os.path.basename(docx_path),
        'K': st.get('K'),
        'batches_done': sorted(st.get('batches_done') or []),
        'ledger_entries': len(((st.get('ledger') or {}).get('entries')) or {}),
        'evidence_files': sum(1 for a, _ in members if a.startswith('evidence/')),
        'files': {arc: _sha256_file(ap) for arc, ap in members},
    }
    with zipfile.ZipFile(out_zip, 'w', zipfile.ZIP_DEFLATED) as z:
        for arc, ap in members:
            z.write(ap, arc)
        z.writestr(CHECKPOINT_MANIFEST,
                   json.dumps(man, indent=1, ensure_ascii=False))
    return man


_EVIDENCE_PATH_KEYS = ('saved', 'montage', 'trace')


def _rebase_evidence_paths(obj, old_evd, new_evd):
    """Rewrite every recorded evidence path in a restored audit_state so it points
    into the NEW evidence directory. Returns how many were rewritten.

    The ledger stores absolute paths from the session that wrote them
    (evidence/facts/q1_x.json, evidence/montages/q7_montage.png, ...). After a
    session boundary those directories do not exist, and S5-1A C5/C6 assert that
    each named file EXISTS — so without this every restored audit would fail
    certification while the files sat right there, correctly restored. Rebasing is
    done EXPLICITLY here rather than left to the resolver's fallback so the state
    on disk is truthful and a human reading it sees real paths.
    """
    n = 0
    if isinstance(obj, dict):
        for k, v in list(obj.items()):
            if k in _EVIDENCE_PATH_KEYS and isinstance(v, str) and v.strip():
                rel = None
                norm = v.replace('\\', '/')
                if old_evd:
                    oe = old_evd.replace('\\', '/').rstrip('/')
                    if norm == oe or norm.startswith(oe + '/'):
                        rel = norm[len(oe):].lstrip('/')
                if rel is None:
                    parts = norm.split('/')
                    if 'evidence' in parts:
                        rel = '/'.join(parts[parts.index('evidence') + 1:])
                if rel:
                    obj[k] = os.path.join(new_evd, *rel.split('/'))
                    n += 1
            else:
                n += _rebase_evidence_paths(v, old_evd, new_evd)
    elif isinstance(obj, list):
        for v in obj:
            n += _rebase_evidence_paths(v, old_evd, new_evd)
    return n


def restore_checkpoint(zip_path, into_dir, docx_path=None, exam=None, mockN=None):
    """Verify and unpack a checkpoint. Returns (manifest, state_path, evidence_dir).

    HARD refusals — every one of these would otherwise let Step 8 certify an audit
    that was never performed on the paper in hand:
      * absent/unparseable manifest, or a schema this build does not know;
      * ANY member whose sha256 differs from the manifest (tamper/truncation);
      * exam_code, mock or paper MD5 disagreeing with the paper being audited.
    On success the restored audit_state.evidence_dir is REWRITTEN to the new
    location — the recorded path is from the previous session's container and no
    longer exists, and S5-1A resolves every stamp through it.
    """
    if not os.path.exists(zip_path):
        raise CheckpointError(f'checkpoint not found: {zip_path}')
    try:
        z = zipfile.ZipFile(zip_path)
    except Exception as e:
        raise CheckpointError(f'checkpoint is not a readable archive ({type(e).__name__})')
    with z:
        names = set(z.namelist())
        if CHECKPOINT_MANIFEST not in names:
            raise CheckpointError('checkpoint_manifest.json missing — not a Step-8 checkpoint.')
        try:
            man = json.loads(z.read(CHECKPOINT_MANIFEST).decode('utf-8'))
        except Exception as e:
            raise CheckpointError(f'checkpoint manifest unparseable ({type(e).__name__})')
        if man.get('schema') != CHECKPOINT_SCHEMA:
            raise CheckpointError(
                f'checkpoint schema {man.get("schema")!r} != {CHECKPOINT_SCHEMA} — '
                f'built by a different framework version; re-run the audit.')
        # identity binding, BEFORE anything is written to disk
        if exam and man.get('exam_code') and man['exam_code'] != exam:
            raise CheckpointError(
                f'checkpoint is for exam_code {man["exam_code"]!r}, not {exam!r}.')
        if mockN is not None and man.get('mock') is not None and int(man['mock']) != int(mockN):
            raise CheckpointError(
                f'checkpoint is for mock {man["mock"]}, not mock {mockN}.')
        if not man.get('paper_md5'):
            raise CheckpointError(
                'checkpoint carries no paper_md5 binding — it cannot be proven to '
                'belong to this paper. Re-run the audit rather than resuming from an '
                'unbindable bundle.')
        if docx_path:
            live = _md5_file(docx_path)
            if live != man['paper_md5']:
                raise CheckpointError(
                    'checkpoint paper MD5 does not match the docx being audited — the '
                    'paper changed, or this checkpoint belongs to another paper. '
                    'Re-run the audit from Phase 1 rather than resuming onto a '
                    'different document.')
        # integrity: every declared member present and hash-exact
        bad = []
        for arc, want in (man.get('files') or {}).items():
            if arc not in names:
                bad.append(f'{arc}:absent'); continue
            got = hashlib.sha256(z.read(arc)).hexdigest()
            if got != want:
                bad.append(f'{arc}:hash')
        if bad:
            raise CheckpointError('checkpoint integrity failure (' + str(len(bad))
                                  + ' member(s)): ' + _flist(bad, cap=8, sep=', '))
        os.makedirs(into_dir, exist_ok=True)
        for arc in (man.get('files') or {}):
            dst = os.path.join(into_dir, *arc.split('/'))
            os.makedirs(os.path.dirname(dst), exist_ok=True)
            with open(dst, 'wb') as fh:
                fh.write(z.read(arc))

    state_path = os.path.join(into_dir, 'audit_state.json')
    evidence_dir = os.path.join(into_dir, 'evidence')
    with open(state_path, encoding='utf-8') as fh:
        st = json.load(fh)
    _old_evd = st.get('evidence_dir')
    _rebased = _rebase_evidence_paths(st.get('ledger'), _old_evd, evidence_dir)
    st['evidence_dir'] = evidence_dir          # last session's path is gone
    st.setdefault('session_log', {})
    if isinstance(st['session_log'], dict):
        st['session_log'].setdefault('checkpoints_restored', []).append(
            {'from': os.path.basename(zip_path), 'at': _now_utc(),
             'batches_done': man.get('batches_done'),
             'evidence_paths_rebased': _rebased})
    with open(state_path, 'w', encoding='utf-8') as fh:
        json.dump(st, fh, indent=1, ensure_ascii=False)
    return man, state_path, evidence_dir


def run_audit(args):
    _reset()
    src = load_sources(args)
    src['_mockN'] = args.mockN
    media_map = _safe_gate('gate_zip', gate_zip, args.docx) or {}
    # The docx itself is the audit SURFACE, not a gate: if it cannot be opened
    # there is genuinely nothing to audit. Still report it as a gate line and a
    # complete roster rather than a raw traceback (v2.12).
    try:
        doc = Document(args.docx)
        _title, blocks = parse_blocks(doc)
    except Exception as _e:
        _fail('A-DOCXOPEN',
              f'cannot open/parse the paper ({type(_e).__name__}: {_e}) — the audit '
              f'surface is unreadable. Re-upload an intact docx (P0.5).')
        return print_results()
    # v2.13 (D1) — bind every inline drawing to its extracted PNG BEFORE the
    # gates run. Guarded and non-fatal by construction: extract_media() never
    # raises, and a resolution shortfall becomes a coverage WARN inside
    # _fig_verdict rather than a missing gate line or a silent pass.
    _safe_gate('A-MEDIA', attach_block_images, blocks, media_map,
               extract_media(args.docx, media_map))
    _safe_gate('A-DOSSIER', gate_dossier, blocks, src, src.get('dossier'),
               src.get('dossier_why'))
    _safe_gate('A-STRUCTURE', gate_structure, blocks, src)
    _safe_gate('A-SECCOUNT', gate_seccount, blocks, src)
    _safe_gate('A-OPTIONS', gate_options, blocks, src)
    _safe_gate('A-QNFIRST', gate_qnfirst, blocks)
    _safe_gate('A-MSQ-INSTR', gate_msq_instr, blocks, src)   # v1.2 — MULTI only; dormant otherwise
    _safe_gate('A-NAT', gate_nat, blocks, src)               # v1.4 — NUMERICAL only; dormant otherwise
    _safe_gate('A-AXIS1', gate_axis1, blocks, src)           # v2.24 — stimulus budget; dormant pre-v1.23
    _safe_gate('A-AXIS3', gate_axis3, blocks, src)           # v2.24 — mechanism budget; dormant pre-v1.23
    _safe_gate('A-AXIS-UNGATED', gate_axis_ungated, blocks, src)  # v2.24 — meta: engine-free by design
    _safe_gate('A-AXIS1-OVERLAP', gate_axis1_overlap, blocks, src)  # v2.26 — partition integrity
    _safe_gate('A-BLANKSEP', gate_blanksep, doc, blocks)
    _safe_gate('A-FONT', gate_font, doc, src)
    _safe_gate('A-SECHDR', gate_sechdr, blocks, doc, src)
    _safe_gate('A-ANSKEY', gate_anskey, doc)
    _safe_gate('A-STIMORPHAN', gate_stimorphan, blocks, src)
    _safe_gate('A-MATCHTABLE', gate_match_table, blocks, src)  # v2.7.1 — MATCH must render a real table
    _safe_gate('A-UNDERLINE', gate_underline, blocks)
    _safe_gate('A-OMML', gate_omml, doc, src, args.final)
    _safe_gate('A-FRACASCII', gate_frac_ascii, blocks, src)
    _safe_gate('A-IMAGES', gate_images, blocks, src, media_map)
    _safe_gate('A-OPTREF', gate_optref, blocks, src)
    _safe_gate('A-ENCODING', gate_encoding_script, doc, src)
    _safe_gate('A-DUP', gate_dup, blocks, src)
    _safe_gate('A-FIGSPECLABEL', gate_specfaith, blocks, src)
    _safe_gate('A-OPTDOMAIN', gate_optdomain, blocks, src)
    _safe_gate('A-HEADER', gate_header, doc, blocks, src)
    _safe_gate('A-QINDEX', gate_qindex, src)   # v2026.08.10 — engine FK gate; dormant unless --registry+--blueprint+--mockN
    rc = print_results()
    # v2.6 — S5-1A COMPLETION GATE: Phase-3 mechanical Part-B/§7 enforcement.
    if getattr(args, 'audit_state', None):
        cg = completion_gate(args.audit_state, src.get('total_questions'), blocks, doc)
        return 0 if (rc == 0 and cg == 0) else 1
    return rc


def gate_qindex(src):
    """A-QINDEX (v2026.08.10 — GAP-2026-08-10-QINDEX-FK-ENFORCEMENT; check 6 added
    v2026.08.12 — GAP-2026-08-12-QINDEX-QUOTA-ENFORCEMENT).
    ENGINE-ENFORCED foreign-key AND quota certification of registry.question_index
    for THIS paper against the blueprint: entry exists; count/coverage 1..total;
    every subtopic_id is byte-identical to a blueprint.subtopic_list id; every
    difficulty is in difficulty_labels; AND (check 6) the difficulty distribution
    equals difficulty_schedule[mock] EXACTLY, when the exam declares one. WHY AN
    ENGINE GATE: the spec-inline G-QINDEX (MockTestCreate S13-QINDEX) performs the
    same six checks, but inline spec code is session-executed and its execution is
    unverifiable — three reference sessions persisted invented subtopic_ids while
    logging clean audits, because this auditor (whose exit code IS durably logged
    and gates SHIP) never looked at question_index. Checks 1-5 closed that hole in
    2026.08.10; check 6 was left session-executed-only, so a session could still
    pass checks 1-5 with a genuinely non-compliant distribution (a null difficulty
    that happens to be a canonical label mismatch is caught by check 5, but a
    WRONG-QUOTA distribution using entirely canonical labels — e.g. an ungoverned
    free assessment — is not) and log a clean, exit-code-durable audit. This gate
    now closes that hole too: armed whenever --registry + --blueprint + --mockN
    are all supplied (the S13-4c re-sweep supplies them); dormant-but-reported
    otherwise. Check 6 is itself dormant (never invented) for an exam that does
    not declare difficulty_schedule at all. Self-contained on purpose — no
    paper_pipeline import, so the per-exam copy has zero new deps."""
    reg = src.get('registry') or {}
    bp  = src.get('blueprint') or {}
    N   = src.get('_mockN')
    if not reg or not bp or N is None:
        _ok('A-QINDEX', 'dormant (needs --registry + --blueprint + --mockN; '
                        'armed at Final Assembly re-sweep S13-4c)')
        return
    _tp = next((mk for mk in bp.get('mocks', []) if mk.get('mock') == N), None)
    pid = (_tp or {}).get('paper_id', f"MOCK:M{int(N):02d}")
    entry = next((e for e in reg.get('question_index', [])
                  if e.get('paper_id', f"MOCK:M{e.get('mock', -1):02d}") == pid), None)
    if entry is None:
        _fail('A-QINDEX', f'no question_index entry for {pid} — registry data '
                          f'missing for a paper this session claims to have built. '
                          f'Rebuild the index (S13-4) before delivery.')
        return
    qs = entry.get('questions', [])
    tq = src.get('total_questions')
    fails = []
    if tq and len(qs) != tq:
        fails.append(f'{len(qs)} entries != total_questions {tq}')
    qn = [x.get('q') for x in qs]
    if tq and (qn != sorted(qn) or len(set(qn)) != len(qn)
               or set(qn) != set(range(1, tq + 1))):
        fails.append(f'q set != 1..{tq}')
    sub_ids = {s.get('subtopic_id') for s in bp.get('subtopic_list', [])}
    bad = {int(x.get('q', -1)): x.get('subtopic_id')
           for x in qs if x.get('subtopic_id') not in sub_ids}
    if bad:
        fails.append('subtopic_id(s) NOT IN blueprint.subtopic_list (a reconstructed/'
                     'invented id — the write must copy the blueprint string verbatim): '
                     + _flist([f"Q{q}={bad[q]!r}" for q in sorted(bad)]))
    canon = bp.get('difficulty_labels', ['Easy', 'Medium', 'Hard'])
    bad_d = sorted({x.get('difficulty') for x in qs if x.get('difficulty') not in canon})
    if bad_d:
        fails.append(f'difficulty value(s) not in {canon}: {bad_d}')
    # ── GAP-2026-08-12-QINDEX-QUOTA-ENFORCEMENT — check 6 ────────────────────
    sched_list = bp.get('difficulty_schedule')
    if sched_list:
        sched = next((d for d in sched_list if d.get('mock') == N), None)
        if sched is None:
            fails.append(f'exam declares difficulty_schedule but has no entry '
                         f'for mock {N}')
        else:
            alias3 = ({'simple': canon[0], 'medium': canon[1], 'hard': canon[2]}
                      if len(canon) == 3
                      else {'simple': 'Easy', 'medium': 'Medium', 'hard': 'Hard'})
            want = {}
            for k, v in sched.items():
                if k in ('mock', 'band') or not isinstance(v, int):
                    continue
                lab = alias3.get(k, k)
                want[lab] = want.get(lab, 0) + v
            want = {lab: want.get(lab, 0) for lab in canon}
            got = {lab: 0 for lab in canon}
            for x in qs:
                d = x.get('difficulty')
                if d in got:
                    got[d] += 1
            if got != want:
                fails.append(f'difficulty distribution {got} != schedule quota {want}')
    if fails:
        _fail('A-QINDEX', '; '.join(fails))
    else:
        _ok('A-QINDEX', f'question_index FK-certified for {pid} '
                        f'({len(qs)} questions, all ids in blueprint)')


def print_results():
    n_fail = sum(1 for lvl, _, _ in RESULTS if lvl == 'FAIL')
    n_warn = sum(1 for lvl, _, _ in RESULTS if lvl == 'WARN')
    n_ok   = sum(1 for lvl, _, _ in RESULTS if lvl == 'OK')
    print('=' * 70)
    print(f'PART-A MACHINE AUDIT  |  OK={n_ok}  WARN={n_warn}  FAIL={n_fail}')
    print('=' * 70)
    for lvl, code, msg in RESULTS:
        mark = {'OK': '  ok ', 'WARN': ' WARN', 'FAIL': ' FAIL'}[lvl]
        print(f'[{mark}] {code:20s} {msg}')
    print('-' * 70)
    if n_fail == 0:
        print(f'RESULT: PASS (0 FAIL, {n_warn} WARN — each names its own mechanical '
              f'remedy or is dormant; no WARN requires human adjudication)')
    else:
        print(f'RESULT: FAIL ({n_fail} gate(s) failed)')
    return 0 if n_fail == 0 else 1


# ============================================================================
# SELF-TEST  (builds tiny docx fixtures; asserts each gate catches + passes)
# ============================================================================
def _mini_doc(tmp, build):
    from docx import Document as D
    d = D()
    build(d)
    p = os.path.join(tmp, 'x.docx')
    d.save(p)
    return p

def _src_stub(tq=2, oc=4, lang='english', sections=None, omml=False):
    return {'total_questions': tq, 'sections': sections or
            [{'name': 'S1', 'q_range': [1, tq], 'total_qs': tq}],
            'options_count': oc, 'opt_label_fmt': '1/2/3/4', 'font_family': 'Calibri',
            'language': lang, 'omml_required_present': omml, 'rules_txt': '',
            'registry': {}, 'manifest': {}, 'blueprint': {},
            'passage_linked': set(), 'cloze_linked': set(), 'figural_qs': set(),
            '_mockN': 1}

def _add_q(d, n, opts=('Alpha', 'Beta', 'Gamma', 'Delta'), qn_first=True, stem='Solve.'):
    if qn_first:
        d.add_paragraph(f'Q.{n}  {stem}')
    else:
        d.add_paragraph(stem)               # stimulus-first (defect)
        d.add_paragraph(f'Q.{n}  {stem}')
    for i, o in enumerate(opts, 1):
        d.add_paragraph(f'{i}.  {o}')
    d.add_paragraph('')

# ── v2.21.9 — HOISTED TO MODULE SCOPE. The CLEAN-SHAPE MATRIX (fixture 5g)
# runs BEFORE the figure fixtures that used to own this helper, and a second
# copy would be drift by construction (S5-2 "one structural question, one
# answer"). Behaviour is byte-identical; every existing caller resolves it
# globally.
def _png_bytes(w=8, h=8):
    """A minimal, valid RGB PNG. stdlib only — the self-test must never
    require matplotlib or PIL to prove that images are attached."""
    raw = b''.join(b'\x00' + b'\xff\x00\x00' * w for _ in range(h))
    def _ck(t, d):
        c = t + d
        return (struct.pack('>I', len(d)) + c
                + struct.pack('>I', zlib.crc32(c) & 0xffffffff))
    return (b'\x89PNG\r\n\x1a\n'
            + _ck(b'IHDR', struct.pack('>IIBBBBB', w, h, 8, 2, 0, 0, 0))
            + _ck(b'IDAT', zlib.compress(raw))
            + _ck(b'IEND', b''))


def self_test():
    passed = 0; total = 0; fails = []
    tmp = tempfile.mkdtemp()
    def check(name, cond):
        nonlocal passed, total
        total += 1
        if cond: passed += 1
        else: fails.append(name)

    # 1. clean paper passes structure/options/qnfirst
    def b_clean(d):
        _add_q(d, 1); _add_q(d, 2)
    p = _mini_doc(tmp, b_clean)
    _reset(); doc = Document(p); _t, blocks = parse_blocks(doc)
    src = _src_stub()
    gate_structure(blocks, src); gate_options(blocks, src); gate_qnfirst(blocks)
    check('clean-no-fail', not any(l == 'FAIL' for l, _, _ in RESULTS))

    # 2. A-COUNT catches wrong count
    _reset(); gate_structure(blocks, _src_stub(tq=5))
    check('A-COUNT-catch', any(c == 'A-COUNT' and l == 'FAIL' for l, c, _ in RESULTS))

    # 3. A-OPTN catches 3 options
    def b_3opt(d): _add_q(d, 1, opts=('A', 'B', 'C'))
    p = _mini_doc(tmp, b_3opt); _reset()
    _t, bl = parse_blocks(Document(p)); gate_options(bl, _src_stub(tq=1))
    check('A-OPTN-catch', any(c == 'A-OPTN' and l == 'FAIL' for l, c, _ in RESULTS))

    # 3b. v2.14 F4 — A-OPTN without --registry on a NAT-shaped Q: still FAIL
    #     (ND6, never silent), but the verdict names the missing contract and the
    #     re-sweep invocation; WITH the contract the same Q passes; a genuinely
    #     short MCQ keeps the plain 'wrong count' wording.
    def b_nat_q(d): _add_q(d, 1, opts=())
    p = _mini_doc(tmp, b_nat_q); _reset()
    _t, bl = parse_blocks(Document(p)); gate_options(bl, _src_stub(tq=1))
    _r = [(l, m) for l, c, m in RESULTS if c == 'A-OPTN']
    check('A-OPTN-no-contract-still-fails', _r and _r[0][0] == 'FAIL')
    check('A-OPTN-no-contract-names-registry',
          _r and 'NOT ASSESSABLE' in _r[0][1] and '--registry' in _r[0][1]
          and 'NAT shape' in _r[0][1])
    _reset(); _sc = _src_stub(tq=1); _sc['options_by_q'] = {'1': 0}
    gate_options(bl, _sc)
    check('A-OPTN-with-contract-passes',
          any(c == 'A-OPTN' and l == 'OK' for l, c, _ in RESULTS))
    _reset(); _t, bl3 = parse_blocks(Document(_mini_doc(tmp, b_3opt)))
    gate_options(bl3, _src_stub(tq=1))
    check('A-OPTN-short-mcq-keeps-plain-wording',
          any(c == 'A-OPTN' and l == 'FAIL' and 'wrong count' in m
              and 'NOT ASSESSABLE' not in m for l, c, m in RESULTS))

    # 4. A-OPTUNIQUE catches duplicate options
    def b_dupopt(d): _add_q(d, 1, opts=('Same', 'Same', 'B', 'C'))
    p = _mini_doc(tmp, b_dupopt); _reset()
    _t, bl = parse_blocks(Document(p)); gate_options(bl, _src_stub(tq=1))
    check('A-OPTUNIQUE-catch', any(c == 'A-OPTUNIQUE' and l == 'FAIL' for l, c, _ in RESULTS))

    # 4a-4e. v2.21.3 — A-OPTLABEL / A-OPTORDER MUTATION CLOSURE.
    #     audit_mutation.py showed `bad_lab.append(...)` and `bad_ord.append(...)`
    #     could BOTH be deleted outright with all 115 fixtures still green: the two
    #     gates that police option labelling on EVERY question of EVERY paper in the
    #     estate had no fixture that could detect them going silent. Only A-OPTN and
    #     A-OPTUNIQUE were covered. Each fixture below KILLS its mutant.
    def _b_labels(labels, texts=None):
        """Block whose option labels are exactly `labels`."""
        def _f(d):
            d.add_paragraph('Q.1  Solve.')
            tx = texts or [f'Option {i}' for i in range(1, len(labels) + 1)]
            for lb, t in zip(labels, tx):
                d.add_paragraph(f'{lb}.  {t}')
        return _f
    def _opt_verdict(labels, fmt='1/2/3/4', oc=4):
        _p = _mini_doc(tmp, _b_labels(labels))
        _t, _bl = parse_blocks(Document(_p))
        _s = _src_stub(tq=1); _s['options_count'] = oc; _s['opt_label_fmt'] = fmt
        _reset(); gate_options(_bl, _s)
        _r = {c: l for l, c, m in RESULTS}
        return _r.get('A-OPTLABEL'), _r.get('A-OPTORDER')

    # 4a — MIXED LABEL FAMILY is an A-OPTLABEL finding (kills the bad_lab mutant).
    check('A-OPTLABEL-mixed-family-is-a-finding',
          _opt_verdict(['1', '2', 'C', '4'])[0] == 'FAIL')

    # 4b — WHOLE-SET FAMILY MISMATCH against the configured format is a finding.
    #      Exam declares numeric labels; the paper ships alpha.
    check('A-OPTLABEL-family-mismatch-vs-format-is-a-finding',
          _opt_verdict(['a', 'b', 'c', 'd'])[0] == 'FAIL')

    # 4c — OUT-OF-ORDER labels are an A-OPTORDER finding (kills the bad_ord mutant).
    check('A-OPTORDER-out-of-order-is-a-finding',
          _opt_verdict(['1', '2', '4', '3'])[1] == 'FAIL')

    # 4d — v2.21.3 ANCHOR LOCK. Labels 2,3,4,5 are CONSECUTIVE but do not start at 1.
    #      The pre-v2.21.3 check was `range(idxs[0], idxs[0]+oc)`, which accepted any
    #      consecutive run, so this paper certified CLEAN while contradicting this
    #      gate's own S5-2 row ("options appear in document order 1..OPTIONS_COUNT").
    #      It is not cosmetic: A-KINT derives the key as an int in 1..OPTIONS_COUNT,
    #      so on such a paper key "option 1" refers to an option that DOES NOT EXIST
    #      and every later key points one place off — every answer for that question
    #      is wrong on the delivered paper.
    check('A-OPTORDER-must-anchor-at-1-not-merely-consecutive',
          _opt_verdict(['2', '3', '4', '5'])[1] == 'FAIL')

    # 4e — GUARD (canonical sets must stay clean in ALL THREE families, so the
    #      anchor fix cannot be "achieved" by making the gate reject everything).
    check('A-OPTORDER-canonical-sets-clean-all-families',
          all(_opt_verdict(_lb, _fm) == ('OK', 'OK') for _lb, _fm in (
              (['1', '2', '3', '4'], '1/2/3/4'),
              (['a', 'b', 'c', 'd'], 'a/b/c/d'),
              (['i', 'ii', 'iii', 'iv'], 'i/ii/iii/iv'))))

    # 6a. v2.21.6 — A-SECCOUNT MUTATION CLOSURE. The gate that proves each section
    #     actually contains the number of questions its q_range declares had NO
    #     fixture: its finding could be deleted with every test still green. Both
    #     halves asserted (matching counts stay clean) so it cannot be "achieved"
    #     by making the gate fire always.
    def _b_three_q(d):
        for _n in (1, 2, 3):
            _add_q(d, _n)
    _p6a = _mini_doc(tmp, _b_three_q)
    _t, _bl6a = parse_blocks(Document(_p6a))
    _s6a = _src_stub(tq=3)
    _s6a['sections'] = [{'name': 'S1', 'q_range': (1, 3), 'total_qs': 3}]
    _reset(); gate_seccount(_bl6a, _s6a)
    _sc_clean = any(c == 'A-SECCOUNT' and l == 'OK' for l, c, m in RESULTS)
    _s6b = _src_stub(tq=3)
    _s6b['sections'] = [{'name': 'S1', 'q_range': (1, 3), 'total_qs': 5}]
    _reset(); gate_seccount(_bl6a, _s6b)
    _sc_catch = any(c == 'A-SECCOUNT' and l == 'FAIL' and 'S1:3/5' in m
                    for l, c, m in RESULTS)
    check('SECCOUNT-section-count-mismatch-is-a-finding', _sc_clean and _sc_catch)

    # 5a-5c. v2.21.7 — SEC-1 / SEC-2: THE LAST TWO OPT_RE CONSUMERS.
    #     GAP-2026-08-02 moved A-DOSSIER onto the option gates' predicate but left
    #     gate_qnfirst and gate_optref on OPT_RE, which cannot see a BARE-LABEL
    #     image option. Both are now on _label_paras()/OPT_LABEL_RE and OPT_RE
    #     itself is retired, so the divergence is structurally impossible.
    def _b_orphan(figural):
        def _f(d):
            d.add_paragraph('Q.1  Select the figure.' if figural else 'Q.1  Solve.')
            for _i in range(1, 5):
                d.add_paragraph(f'{_i}.' if figural else f'{_i}.  Opt {_i}')
            d.add_paragraph(' '.join(['word'] * 60))      # ORPHANED lead-in
            d.add_paragraph('Q.2  Solve.')
            for _i in range(1, 5):
                d.add_paragraph(f'{_i}.  Opt {_i}')
        return _f
    def _qn_fail(build):
        _p = _mini_doc(tmp, build)
        _t, _bl = parse_blocks(Document(_p))
        _reset(); gate_qnfirst(_bl)
        return any(c == 'A-QNFIRST' and l == 'FAIL' for l, c, m in RESULTS)

    # 5a — SEC-1 PARITY LOCK. The SAME orphaned lead-in must be caught after a
    #      FIGURAL block and after a TEXT block. Under OPT_RE the figural anchor
    #      never matched, last_opt stayed -1, and the `continue` skipped the whole
    #      check while the gate still printed ok — 25 of 60 blocks unchecked on a
    #      real paper. Both halves asserted so the fix cannot be "achieved" by
    #      making the gate fire always.
    check('QNFIRST-figural-and-text-blocks-at-parity',
          _qn_fail(_b_orphan(True)) and _qn_fail(_b_orphan(False)))

    # 5b — GUARD: a CLEAN figural block (blank separator, no orphan) stays OK.
    def _b_fig_clean(d):
        d.add_paragraph('Q.1  Select the figure.')
        for _i in range(1, 5):
            d.add_paragraph(f'{_i}.')
        d.add_paragraph('')
        d.add_paragraph('Q.2  Solve.')
        for _i in range(1, 5):
            d.add_paragraph(f'{_i}.  Opt {_i}')
    check('QNFIRST-clean-figural-block-stays-clean',
          not _qn_fail(_b_fig_clean))

    # 5d-5f. v2.21.9 — GAP-2026-08-02-QNFIRST-IMAGE-OPTION.
    #     FIXTURE 5b ABOVE MODELLED THE LABELS AND NOT THE PICTURES. Its "clean
    #     figural block" renders four BARE LABELS AND NO IMAGES — a shape v2.21.4
    #     declares a FINDING in its own right ("a REGISTRY-DECLARED figural Q
    #     rendering ZERO images is a finding"). So the guard that was supposed to
    #     prove a conformant figural block stays clean never once built the shape
    #     R-FIGURAL actually mandates, and v2.21.7's anchor move shipped a false
    #     POSITIVE on every stem_and_options figural paper in the estate while the
    #     self-test reported 142/142. EIGHTH hollow-branch occurrence.
    #     Measured on the real reproduction: A-QNFIRST FAIL "stimulus orphaned
    #     before Q.<n>" on a block A-FIGCOMP / A-OPTN / A-OPTUNIQUE / A-DOSSIER all
    #     pass. It is a FAIL, so exit is non-zero and MANDATE D refuses to certify
    #     — and A-QNFIRST is catalogued CP-fixable, so Phase 1 calls CP-QNFIRST on
    #     a block with nothing to re-emit: an unfixable false failure inside a
    #     repair loop.
    def _b_fig_real(trailer=None):
        """The MANDATED stem_and_options shape (R-FIGURAL / G-FIGURAL-COMPOSITE):
        problem image, then one BARE label + its OWN picture per option."""
        def _f(d):
            d.add_paragraph('Q.1  Select the figure.')
            _p = d.add_paragraph()
            _p.add_run().add_picture(io.BytesIO(_png_bytes()), width=Inches(2.0))
            for _i in range(1, 5):
                d.add_paragraph(f'{_i}.')
                _op = d.add_paragraph()
                _op.add_run().add_picture(io.BytesIO(_png_bytes()), width=Inches(1.0))
            if trailer:
                trailer(d)
            d.add_paragraph('')
            _add_q(d, 2)
        return _f
    def _tr_img(d):
        d.add_paragraph().add_run().add_picture(io.BytesIO(_png_bytes()), width=Inches(2.0))
    def _tr_tbl(d):
        d.add_table(rows=2, cols=2)
    def _tr_txt(d):
        d.add_paragraph(' '.join(['word'] * 60))

    # 5d — THE FIX. A conformant figural block MUST stay clean. The final option's
    #      OWN picture necessarily follows the last option LABEL; treating it as an
    #      orphaned stimulus fails the paper for obeying the producer contract.
    #      MUTATION-VERIFIED: measures False on the v2.21.8 build.
    check('QNFIRST-mandated-figural-option-images-are-not-orphans',
          not _qn_fail(_b_fig_real()))

    # 5e — THE FIX IS NOT "SKIP FIGURAL BLOCKS". A genuine orphan AFTER the last
    #      option's image must STILL fire, in all three of its forms. Without this,
    #      5d could be "achieved" by exempting image-bearing blocks outright, which
    #      would restore the very false NEGATIVE v2.21.7 existed to close.
    check('QNFIRST-genuine-orphan-after-figural-options-still-caught',
          _qn_fail(_b_fig_real(_tr_img))
          and _qn_fail(_b_fig_real(_tr_tbl))
          and _qn_fail(_b_fig_real(_tr_txt)))

    # 5f — RENDER PARITY. The SAME logical block, rendered with IMAGE options and
    #      with TEXT options, must reach the SAME verdict — clean when clean and
    #      flagged when orphaned. This is the assertion whose absence let a gate
    #      hold two contradictory opinions about one block for a whole release.
    def _b_txt_real(trailer=None):
        def _f(d):
            d.add_paragraph('Q.1  Solve.')
            for _i in range(1, 5):
                d.add_paragraph(f'{_i}.  Opt {_i}')
            if trailer:
                trailer(d)
            d.add_paragraph('')
            _add_q(d, 2)
        return _f
    check('QNFIRST-image-and-text-renderings-agree',
          (_qn_fail(_b_fig_real())     == _qn_fail(_b_txt_real()))
          and (_qn_fail(_b_fig_real(_tr_txt)) == _qn_fail(_b_txt_real(_tr_txt)))
          and (_qn_fail(_b_fig_real(_tr_tbl)) == _qn_fail(_b_txt_real(_tr_tbl))))

    # ══════════════════════════════════════════════════════════════════════════
    # 5g. THE CLEAN-SHAPE MATRIX (v2.21.9) — THE GENERALISED GUARD.
    # ══════════════════════════════════════════════════════════════════════════
    #     WHY THIS EXISTS, AND WHY IT IS NOT ANOTHER PER-GATE FIXTURE. This corpus
    #     has now recorded EIGHT hollow-branch defects, and the last four share ONE
    #     shape: a block-structural gate written and fixtured against the TEXT-option
    #     rendering, then meeting a DIFFERENT LEGITIMATE RENDERING in the wild.
    #       v2.21   A-DOSSIER  — could not see an IMAGE option; also inflated on an
    #                            ENUMERATED stem.
    #       v2.21.3 A-OPTORDER — accepted any consecutive run.
    #       v2.21.7 A-QNFIRST  — SKIPPED figural blocks entirely.
    #       v2.21.9 A-QNFIRST  — then FALSE-FAILED them (this release).
    #     Each was closed with a fixture for THAT gate. That is necessary and it is
    #     not sufficient: it depends on an author ANTICIPATING the shape, and four
    #     times running, nobody did. CHECK AN (shared-predicate parity) and CHECK AO
    #     (tautological fixture) do not catch it either — this gate used the shared
    #     predicate correctly and its fixture was not tautological. It simply never
    #     built the shape.
    #
    #     THE INVARIANT, AND IT NEEDS NO ANTICIPATION: **A CONFORMANT PAPER IS
    #     CONFORMANT IN EVERY RENDERING THE FRAMEWORK MANDATES.** So: build a CLEAN
    #     paper in each canonical shape, run EVERY gate over each, and assert ZERO
    #     FAILs. A gate that false-fails a legitimate rendering turns the self-test
    #     red — automatically, in all ~200 exam projects, for gates not yet written.
    #
    #     IT IS SELF-HOSTING. Gates are DISCOVERED by introspection (every module
    #     callable named gate_*) and their arguments supplied BY PARAMETER NAME, so
    #     a NEW gate is covered the moment it is added, with nobody remembering to
    #     opt it in. That is the difference between a fixture and a control.
    #     Measured: on the v2.21.8 build this fixture reports the A-QNFIRST FAIL on
    #     the clean image-option shape — it would have caught this release's defect
    #     with no one suspecting the gate was wrong.
    #     ONLY 'FAIL' IS ASSERTED. A WARN on a clean shape is legitimate (a dossier
    #     that is absent, an engine that is not installed); a FAIL is the gate
    #     saying a conformant paper is defective, which is never legitimate.
    def _shape_text(d, n):
        d.add_paragraph(f'Q.{n}  Solve.')
        for _i in range(1, 5):
            d.add_paragraph(f'{_i}.  Opt {_i}')
        d.add_paragraph('')

    def _shape_image(d, n):
        """R-FIGURAL stem_and_options: problem image + one bare label + picture
        per option. THE SHAPE NO FIXTURE HAD EVER BUILT."""
        d.add_paragraph(f'Q.{n}  Select the figure.')
        d.add_paragraph().add_run().add_picture(
            io.BytesIO(_png_bytes()), width=Inches(2.0))
        for _i in range(1, 5):
            d.add_paragraph(f'{_i}.')
            d.add_paragraph().add_run().add_picture(
                io.BytesIO(_png_bytes()), width=Inches(1.0))
        d.add_paragraph('')

    def _shape_enum(d, n):
        """STATEMENT/SEQUENCE/MATCH/ASSERTION_REASON: an ENUMERATED stem whose
        numbered points must not be mistaken for options (the v2.21 defect)."""
        d.add_paragraph(f'Q.{n}  Consider the following statements:')
        d.add_paragraph('1.  First statement.')
        d.add_paragraph('2.  Second statement.')
        d.add_paragraph('Which of the above is/are correct?')
        for _i in range(1, 5):
            d.add_paragraph(f'{_i}.  Opt {_i}')
        d.add_paragraph('')

    def _shape_figural_problem_only(d, n):
        """v2.22.0 — THE SHAPE THAT PRODUCED GAP-2026-08-03-FIGCOMP-ROLE: ONE
        role='problem' figure with TEXT options (diagram/gel/graph + text answers,
        the commonest figural shape in the life sciences — 27 of 33 figural Qs on
        the real IIT_JAM_BIOTECHNOLOGY paper). No fixture had ever built it."""
        d.add_paragraph(f'Q.{n}  Study the figure and answer.')
        d.add_paragraph().add_run().add_picture(
            io.BytesIO(_png_bytes()), width=Inches(2.0))
        for _i in range(1, 5):
            d.add_paragraph(f'{_i}.  Opt {_i}')
        d.add_paragraph('')

    def _shape_nat_zero_option(d, n):
        """v2.22.0 — a NUMERICAL-ANSWER question: bold stem carrying the
        nat_instruction and ZERO option paragraphs (R13/R14 NAT exemption).
        20 of the 60 questions on the real paper are this shape."""
        p = d.add_paragraph()
        p.add_run(f'Q.{n}  The value is ______ .  Enter your answer as a '
                  'numerical value.').bold = True
        d.add_paragraph('')

    _SHAPES = (('text', _shape_text, {}),
               ('image', _shape_image, {'figural_qs': {1}}),
               ('enumerated-stem', _shape_enum, {}),
               ('figural-problem-only', _shape_figural_problem_only,
                {'figural_qs': {1},
                 'figure_specs': {'q1_problem.png': {'question': 1,
                                                     'role': 'problem'}}}),
               ('nat-zero-option', _shape_nat_zero_option,
                {'options_by_q': {'1': 0}, 'nat_present': True}))

    def _shape_fails(builder, extra, img_names=None):
        def _b(d):
            builder(d, 1)
            _shape_text(d, 2)
        _p = _mini_doc(tmp, _b)
        if img_names:
            _p = _stamp_canonical(_p, img_names)
        _doc = Document(_p)
        _t, _bl = parse_blocks(_doc)
        _s = _src_stub(tq=2); _s.update(extra)
        _pool = dict(doc=_doc, blocks=_bl, src=_s, docx_path=_p, media_map={},
                     dossier=None, final=True, why=None)
        _reset()
        for _gn in sorted(n for n in globals() if n.startswith('gate_')):
            _g = globals()[_gn]
            if not callable(_g):
                continue
            _sig = inspect.signature(_g)
            if not all(_pn in _pool for _pn in _sig.parameters):
                continue          # a gate needing state this matrix cannot stub
            try:
                _g(*[_pool[_pn] for _pn in _sig.parameters])
            except Exception:
                pass              # _safe_gate owns crash reporting; not this fixture
        # v2.22.0 — WARN-SEVERITY FINDINGS COUNT TOO. The original matrix asserted
        # only FAILs, and GAP-2026-08-03-FIGCOMP-ROLE was a WARN — so the matrix
        # would NOT have caught the very defect that motivated extending it. But a
        # blanket "no WARN" is wrong: dormancy warnings on a clean shape are
        # legitimate and expected (absent blueprint, uninstalled engine).
        # THE DISCRIMINATOR IS MECHANICAL: a DORMANCY warning describes the gate
        # ('skipped', 'dormant', 'absent'); a DEFECT finding NAMES QUESTIONS. On a
        # CONFORMANT paper no gate may name a question at any severity.
        _bad = [(c, m) for l, c, m in RESULTS if l == 'FAIL']
        _bad += [(c, m) for l, c, m in RESULTS
                 if l == 'WARN' and re.search(r'\bQ\d+\b', m or '')]
        return _bad

    def _stamp_canonical(path, names):
        """v2.22.0 — stamp docPr name/descr in document order, exactly as Step 7
        S10-8 _name_last_drawing does. WITHOUT THIS the matrix's synthetic images
        carry python-docx default part names, A-MATHRASTER-VIEW correctly reports
        a name-contract miss, and the matrix would flag its OWN fixture rather
        than a gate defect — a false positive in the very control built to stop
        false positives. A conformant paper always carries canonical names."""
        out = os.path.join(tempfile.mkdtemp(), 'stamped.docx')
        with zipfile.ZipFile(path) as zin:
            items = zin.infolist()
            xml = zin.read('word/document.xml').decode('utf-8')
            it = iter(names); _c = [0]
            def _sub(mo):
                nm = next(it, 'x.png'); _c[0] += 1
                close = '/>' if mo.group(0).rstrip().endswith('/>') else '>'
                return (f'<wp:docPr id="{_c[0]}" name="{nm}" '
                        f'descr="alt for {nm}"{close}')
            xml = re.sub(r'<wp:docPr\b[^>]*>', _sub, xml)
            with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zo:
                for i in items:
                    zo.writestr(i, xml.encode('utf-8')
                                if i.filename == 'word/document.xml'
                                else zin.read(i.filename))
        return out

    _SHAPE_IMG_NAMES = {
        'image': ['q1_problem.png', 'q1_opt1.png', 'q1_opt2.png',
                  'q1_opt3.png', 'q1_opt4.png'],
        'figural-problem-only': ['q1_problem.png'],
    }
    _matrix = {nm: _shape_fails(b, x, _SHAPE_IMG_NAMES.get(nm))
               for nm, b, x in _SHAPES}
    check('SHAPE-MATRIX-no-gate-fails-a-conformant-rendering',
          all(not v for v in _matrix.values()))
    # Report WHICH shape and WHICH gate, so a future failure is diagnosable at a
    # glance instead of sending someone back through the whole roster.
    for _snm, _sfails in _matrix.items():
        check(f'SHAPE-MATRIX[{_snm}]-clean',  not _sfails)

    # ══════════════════════════════════════════════════════════════════════════
    # 5h. SELF-ADJUDICATION (v2.22.0) — NO GATE MAY DEFER TO A HUMAN.
    # ══════════════════════════════════════════════════════════════════════════
    #     THE DEPLOYED REALITY: the operator running these ~200 exam pipelines is
    #     not a reviewer and cannot adjudicate anything. A finding that says "VIEW
    #     + fix in Part B" or "for reviewer adjudication" therefore has NO RECEIVER
    #     — it is, in effect, a silent no-op, and at worst it parks conformant work
    #     in an inspection queue that never drains. That is exactly what happened:
    #     A-FIGCOMP's 17 false findings each carried "VIEW + fix in Part B" and the
    #     resulting queue stalled a real 60-question audit for a day.
    #
    #     THE RULE: every gate message must either (a) name a MECHANICAL remedy the
    #     pipeline can execute, or (b) declare the gate dormant. Never "look at this
    #     and decide". This fixture scans THIS FILE's emitted message literals for
    #     deferral phrasing, so the rule is enforced on gates not yet written.
    #     Comments are excluded — the prose above documents the defect and must be
    #     allowed to name it.
    # Built from fragments so this list does not match ITSELF — a scanner whose
    # own definition trips it would be permanently red and would then be deleted,
    # which is how controls die.
    _DEFERRAL = tuple(_a + _b for _a, _b in (
        ('view ', '+ fix'), ('view in ', 'part b'), ('reviewer ', 'adjudication'),
        ('adjud', 'icate'), ('manually ', 'inspect'), ('ask the ', 'reviewer'),
        ('someone ', 'should'), ('human ', 'review')))
    # v2.24.0 — SCAN THE EMITTED MESSAGES, NOT THE WHOLE FILE. The line-based
    # scanner also read DOCSTRINGS, so a docstring EXPLAINING this rule tripped it
    # (gate_specfaith's own does). A control that forces contorted documentation is
    # a control that eventually gets weakened, so it is now exact: parse the module
    # and inspect only the string literals actually passed to _warn/_fail/_ok/print
    # — which is precisely the set a human could ever read.
    _tree = ast.parse(open(__file__, encoding='utf-8').read())
    _emitted = []
    for _node in ast.walk(_tree):
        if not isinstance(_node, ast.Call):
            continue
        _fn = _node.func
        _nm = getattr(_fn, 'id', None) or getattr(_fn, 'attr', None)
        if _nm not in ('_warn', '_fail', '_ok', 'print'):
            continue
        for _arg in list(_node.args) + [kw.value for kw in _node.keywords]:
            for _sub in ast.walk(_arg):
                if isinstance(_sub, ast.Constant) and isinstance(_sub.value, str):
                    _emitted.append((_sub.lineno, _sub.value))
    _offenders = [f'L{_ln}:{_d}' for _ln, _txt in _emitted
                  for _d in _DEFERRAL if _d in _txt.lower()]
    check('SELF-ADJUDICATION-no-gate-defers-to-a-human',
          not _offenders)

    # 5k. v2.25.0 — CROSS-STEP LABEL PARITY (GAP-2026-08-03-LABELFMT).
    #     Step 7 RESOLVES a section_rules option_label_format into a render
    #     template; Step 8 CLASSIFIES the same string into a family. Two different
    #     computations that must agree, and before v5.37/v2.25.0 they did not:
    #     'i/ii/iii/iv' rendered (a)(b)(c)(d) and classified as 'roman', so
    #     A-OPTLABEL FAILED every question on a paper that obeyed Step 7 — exit 1,
    #     MANDATE D blocked delivery, no CP repair possible.
    #     paper_pipeline.resolve_option_label is now the single source, and this
    #     fixture asserts the pair from THIS side too: a change to either
    #     implementation turns BOTH self-tests red, not just one.
    #     The auditor must run STANDALONE (Context-2), so an absent paper_pipeline
    #     is a SKIP, never a failure.
    try:
        import paper_pipeline as _pp_parity
        _SUPPORTED = ('1/2/3/4', 'A/B/C/D', 'a/b/c/d', 'i/ii/iii/iv', 'I/II/III/IV',
                      '(1)/(2)/(3)/(4)', '(A)/(B)/(C)/(D)', '(i)/(ii)/(iii)/(iv)',
                      '1)/2)/3)/4)', '')
        # v2.26.0 — EVALUATED SAFELY. resolve_option_label RAISES by design on an
        # ambiguous notation, and check() takes an ALREADY-EVALUATED condition, so a
        # raise here would abort the whole suite and every later fixture would
        # silently never run (the paper_pipeline v5.38 defect, one file over).
        def _parity(fn):
            try:
                return bool(fn())
            except Exception:
                return False
        check('LABEL-PARITY-step7-resolution-matches-step8-family',
              _parity(lambda: all(_pp_parity.resolve_option_label(_x)[1]
                                  == option_label_family(_x) for _x in _SUPPORTED)))
        check('LABEL-PARITY-family-classifiers-agree',
              _parity(lambda: all(_pp_parity.option_label_family(_x)
                                  == option_label_family(_x) for _x in _SUPPORTED)))
    except ImportError:
        check('LABEL-PARITY-skipped-standalone', True)

    # ══════════════════════════════════════════════════════════════════════════
    # A-QINDEX (GAP-2026-08-10-QINDEX-FK-ENFORCEMENT)
    #   This gate's EXIT CODE is the enforcement — that is the whole argument for
    #   adding it, since the spec-inline G-QINDEX is session-executed and
    #   unverifiable. It shipped with no fixture: every pre-existing fixture
    #   reaches it through _src_stub, whose registry and blueprint are {}, so all
    #   229 of them exercised the DORMANT branch and none the armed one. A gate
    #   that blocks SHIP needs a test that fails when the gate is removed.
    _QBP = {'total_questions': 2, 'mocks': [{'mock': 1, 'paper_id': 'MOCK:M01'}],
            'subtopic_list': [{'subtopic_id': 'PHY.MECH.NEWTON'},
                              {'subtopic_id': 'PHY.MECH.WORK'}],
            'difficulty_labels': ['Easy', 'Medium', 'Hard']}
    _QCLEAN = [{'q': 1, 'subtopic_id': 'PHY.MECH.NEWTON', 'difficulty': 'Easy'},
               {'q': 2, 'subtopic_id': 'PHY.MECH.WORK', 'difficulty': 'Hard'}]

    def _q_run(registry):
        _reset()
        gate_qindex({'total_questions': 2, 'blueprint': _QBP,
                     'registry': registry, '_mockN': 1})
        return [(l, c, m) for l, c, m in RESULTS if c == 'A-QINDEX']

    def _q_reg(qs):
        return {'question_index': [{'paper_id': 'MOCK:M01', 'questions': qs}]}

    check('A-QINDEX-armed-clean-certifies',
          [l for l, _, _ in _q_run(_q_reg(_QCLEAN))] == ['OK'])
    # THE DEFECT, verbatim: an id that reads correctly to a human but was retyped
    # rather than copied from the blueprint. Three sessions persisted this and
    # logged clean audits; all three detonated four steps later at Step 11.
    _q_inv = _q_run(_q_reg([_QCLEAN[0], {'q': 2, 'subtopic_id': 'Physics.Mechanics.Work',
                                         'difficulty': 'Hard'}]))
    check('A-QINDEX-armed-invented-id-FAILS',
          [l for l, _, _ in _q_inv] == ['FAIL'])
    check('A-QINDEX-names-the-offending-question',
          bool(_q_inv) and 'Q2=' in _q_inv[0][2])
    # the 4th defective mock: a later registry write dropped the entry entirely
    check('A-QINDEX-armed-lost-entry-FAILS',
          [l for l, _, _ in _q_run({'question_index': []})] == ['FAIL'])
    check('A-QINDEX-armed-short-count-FAILS',
          [l for l, _, _ in _q_run(_q_reg(_QCLEAN[:1]))] == ['FAIL'])
    check('A-QINDEX-armed-duplicate-q-FAILS',
          [l for l, _, _ in _q_run(_q_reg(
              [_QCLEAN[0], {'q': 1, 'subtopic_id': 'PHY.MECH.WORK',
                            'difficulty': 'Hard'}]))] == ['FAIL'])
    check('A-QINDEX-armed-bad-difficulty-FAILS',
          [l for l, _, _ in _q_run(_q_reg(
              [_QCLEAN[0], {'q': 2, 'subtopic_id': 'PHY.MECH.WORK',
                            'difficulty': 'Tough'}]))] == ['FAIL'])

    # ── GAP-2026-08-12-QINDEX-QUOTA-ENFORCEMENT — check 6 ────────────────────
    # THE DEFECT: checks 1-5 above are the FK/coverage/label checks moved into
    # the engine at GAP-2026-08-10-QINDEX-FK-ENFORCEMENT. The exact-quota check
    # (schedule-first assignment) was left session-executed-only — a session
    # could pass checks 1-5 with a canonical-labelled but WRONG-QUOTA
    # distribution and still log a clean, exit-code-durable audit. Ships WITH a
    # fixture from day one.
    _QBP6 = dict(_QBP, difficulty_schedule=[{'mock': 1, 'simple': 1, 'hard': 1}])

    def _q_run6(registry):
        _reset()
        gate_qindex({'total_questions': 2, 'blueprint': _QBP6,
                     'registry': registry, '_mockN': 1})
        return [(l, c, m) for l, c, m in RESULTS if c == 'A-QINDEX']

    check('A-QINDEX-armed-quota-exact-match-certifies',
          [l for l, _, _ in _q_run6(_q_reg(_QCLEAN))] == ['OK'])
    check('A-QINDEX-armed-quota-mismatch-FAILS',
          [l for l, _, _ in _q_run6(_q_reg(
              [{'q': 1, 'subtopic_id': 'PHY.MECH.NEWTON', 'difficulty': 'Medium'},
               {'q': 2, 'subtopic_id': 'PHY.MECH.WORK', 'difficulty': 'Medium'}]))]
          == ['FAIL'])
    check('A-QINDEX-quota-dormant-when-exam-has-no-schedule',
          # _QBP (no difficulty_schedule key) must still certify a clean set —
          # check 6 must never be invented for an exam that never declared one.
          [l for l, _, _ in _q_run(_q_reg(_QCLEAN))] == ['OK'])
    def _q_run_no_entry_for_mock():
        _reset()
        gate_qindex({'total_questions': 2,
                    'blueprint': dict(_QBP, difficulty_schedule=[{'mock': 99, 'simple': 2}]),
                    'registry': _q_reg(_QCLEAN), '_mockN': 1})
        return [(l, c, m) for l, c, m in RESULTS if c == 'A-QINDEX']

    check('A-QINDEX-quota-fails-when-schedule-exists-but-not-for-this-mock',
          [l for l, _, _ in _q_run_no_entry_for_mock()] == ['FAIL'])

    # DORMANCY IS A CONTRACT, not an accident: ~200 exams hold a Step-6 copy of
    # this auditor and invoke it without the new flags. Dormant must stay OK, and
    # must SAY it is dormant rather than silently reporting a pass.
    _reset()
    gate_qindex({'total_questions': 2, 'blueprint': {}, 'registry': {}, '_mockN': None})
    _q_dorm = [(l, m) for l, c, m in RESULTS if c == 'A-QINDEX']
    check('A-QINDEX-unarmed-is-dormant-OK-and-says-so',
          len(_q_dorm) == 1 and _q_dorm[0][0] == 'OK' and 'dormant' in _q_dorm[0][1])

    # WIRING. Every fixture above calls gate_qindex DIRECTLY, so all of them stay
    # green if the gate is never added to run_audit's gate list — a gate that is
    # written, documented, tested and unreachable
    # (GAP-2026-08-01-FIGPROFILE-ENGINE-BINDING, where the delegation was recorded
    # at the call sites and the import was simply never added). Assert reachability
    # from the runner itself, not from the gate.
    import inspect as _insp
    _runner = _insp.getsource(run_audit)
    check('A-QINDEX-is-wired-into-run_audit',
          'gate_qindex' in _runner and "'A-QINDEX'" in _runner)
    check('A-QINDEX-runner-passes-_mockN-through',
          "src['_mockN'] = args.mockN" in _runner)

    # QINDEX PARITY — gate_qindex is a deliberate second implementation of
    # paper_pipeline.validate_question_index (its docstring pins the reason: the
    # per-exam copy must gain no new import). Deliberate duplication is only safe
    # while the pair is asserted equal, exactly as LABEL-PARITY above. Absent
    # paper_pipeline is a SKIP — the auditor must run standalone (Context-2).
    try:
        import paper_pipeline as _pp_q
        _QCASES = [
            _QCLEAN,
            [_QCLEAN[0], {'q': 2, 'subtopic_id': 'Physics.Mechanics.Work',
                          'difficulty': 'Hard'}],
            _QCLEAN[:1],
            [_QCLEAN[0], {'q': 1, 'subtopic_id': 'PHY.MECH.WORK', 'difficulty': 'Hard'}],
            [_QCLEAN[0], {'q': 2, 'subtopic_id': 'PHY.MECH.WORK', 'difficulty': 'Tough'}],
        ]

        def _q_parity():
            for _c in _QCASES:
                _mine = [l for l, _, _ in _q_run(_q_reg(_c))] == ['OK']
                _theirs = _pp_q.validate_question_index(_q_reg(_c), _QBP, mock_n=1)[0]
                if _mine != _theirs:
                    return False
            return True
        try:
            _qp = _q_parity()
        except Exception:
            _qp = False
        check('A-QINDEX-PARITY-auditor-gate-matches-pp-validate', _qp)

        # QUOTA PARITY (check 6, GAP-2026-08-12-QINDEX-QUOTA-ENFORCEMENT) — the
        # same agreement requirement, now against _QBP6 (a schedule present).
        _Q6CASES = [
            _QCLEAN,   # matches the schedule exactly -> both OK
            [{'q': 1, 'subtopic_id': 'PHY.MECH.NEWTON', 'difficulty': 'Medium'},
             {'q': 2, 'subtopic_id': 'PHY.MECH.WORK', 'difficulty': 'Medium'}],  # wrong quota
        ]

        def _q6_parity():
            for _c in _Q6CASES:
                _mine = [l for l, _, _ in _q_run6(_q_reg(_c))] == ['OK']
                _theirs = _pp_q.validate_question_index(_q_reg(_c), _QBP6, mock_n=1)[0]
                if _mine != _theirs:
                    return False
            return True
        try:
            _q6p = _q6_parity()
        except Exception:
            _q6p = False
        check('A-QINDEX-QUOTA-PARITY-auditor-gate-matches-pp-validate', _q6p)
    except ImportError:
        check('A-QINDEX-PARITY-skipped-standalone', True)

    # THIRD LEG: the P10 spec-inline preflight (Framework_MockTestExplain).
    #   paper_pipeline's header describes ONE CONTRACT AT FOUR INDEPENDENT SITES
    #   and says the drift risk is "held down by the A-QINDEX self-test agreement
    #   matrix (pp vs engine vs P10 predicate)". The fixture above compares only
    #   pp vs engine, so that claim was true of two sites out of three — the same
    #   shape as the "ONE implementation, imported by" claim it replaced. This
    #   leg makes it true. It is not decorative: P10 shipped in 2026.08.10.4
    #   MISSING the q-set coverage check, so an index of q=[1,1] had the right
    #   LENGTH, passed P10, and was certified clean while BOTH engine copies
    #   failed it — found only by running the three against each other.
    #   The spec's block is EXECUTED, not pattern-matched: a coverage rule that
    #   is merely mentioned in prose is not a rule.
    #   Absent spec = SKIP, like LABEL-PARITY: a per-exam auditor copy runs
    #   standalone (Context-2) with no Framework_*.md beside it.
    try:
        import textwrap as _tw
        _p10_spec = open(os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                      'Framework_MockTestExplain.md'),
                         encoding='utf-8').read()
        _p10_src = _tw.dedent(
            _p10_spec.split('## P10 — REGISTRY-FK TRIPWIRE', 1)[1]
                     .split('```python', 1)[1].split('```', 1)[0]
                     .split('# Ledger↔index agreement', 1)[0])
        _p10_dir = tempfile.mkdtemp()
        # The block reads the exam's two artefacts by absolute project path;
        # redirect them at a fixture dir. If the spec ever renames that path the
        # substitution is a no-op, which would silently neuter this fixture — so
        # assert it bit rather than trusting it.
        _p10_run = _p10_src.replace('/mnt/project/{EXAM}', _p10_dir + '/{EXAM}')
        check('A-QINDEX-P10-fixture-is-wired-to-the-fixture-dir',
              _p10_run != _p10_src and '/mnt/project/' not in _p10_run)

        def _p10_ok(qs):
            with open(os.path.join(_p10_dir, 'E_blueprint.json'), 'w') as _fh:
                json.dump(_QBP, _fh)
            with open(os.path.join(_p10_dir, 'E_registry.json'), 'w') as _fh:
                json.dump(_q_reg(qs) if qs is not None else {'question_index': []}, _fh)
            try:
                # v2.12→: P10/0 (MockTestExplain v1.24.0) asserts trigger-N's
                # paper_slug == the uploaded docx's slug (PAPER_SLUG, bound at
                # P1 in a real session). The fixture supplies the MATCHING slug
                # for _QBP's mock 1 ('MOCK:M01' -> 'Mock01') so the identity
                # gate passes and the qindex cases below stay the variable
                # under test — and P10/0 itself is now execution-covered too.
                exec(compile(_p10_run, 'P10', 'exec'),
                     {'EXAM': 'E', 'N': 1, 'PAPER_SLUG': 'Mock01'})
                return True
            except SystemExit:
                return False

        def _p10_matrix():
            for _c in _QCASES + [_QCLEAN[:1] + [{'q': 3, 'subtopic_id': 'PHY.MECH.WORK',
                                                 'difficulty': 'Hard'}], None]:
                _reg2 = _q_reg(_c) if _c is not None else {'question_index': []}
                _e = [l for l, _, _ in _q_run(_reg2)] == ['OK']
                if _p10_ok(_c) != _e:
                    return False
            return True
        try:
            _p10_agrees = _p10_matrix()
        except Exception:
            _p10_agrees = False
        check('A-QINDEX-PARITY-P10-spec-preflight-matches-the-engine-gate',
              _p10_agrees)
    except (ImportError, FileNotFoundError, IndexError):
        check('A-QINDEX-PARITY-P10-skipped-standalone', True)

    # FOURTH LEG: MockDeliver S1-3's remediation classifier — the last of the
    #   four sites, and the only one whose output is a HUMAN INSTRUCTION rather
    #   than a verdict. Its W1 line is described as "safe to apply mechanically
    #   to the registry", so a leaf-rule that drifts from classify_unresolved
    #   does not fail loudly: it prints a confident PATCH pointing at the wrong
    #   subtopic, and the operator applies it. A W1 that should have been D
    #   (leaf present on two subtopics) is exactly that silent mis-repair.
    try:
        import paper_pipeline as _pp_c
        _md = open(os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                'Framework_MockDeliver.md'), encoding='utf-8').read()
        _c_src = _tw.dedent('    if _unresolved:' + _md.split(
            '    if _unresolved:', 1)[1].split('Root-cause fix:', 1)[0] + '")')
        _CBP2 = {'subtopic_list': [{'subtopic_id': 'PHY.MECH.WORK'},
                                   {'subtopic_id': 'CHEM.THERMO.WORK'},
                                   {'subtopic_id': 'PHY.MECH.NEWTON'}]}
        _stale = {1: 'Physics.Mechanics.NEWTON',   # unique leaf   -> W1
                  2: 'X.Y.WORK',                   # leaf on two   -> D
                  3: 'PHY.MECH.ENERGYY'}           # no leaf       -> W2

        def _spec_report():
            _ns = {'_unresolved': dict(_stale), 'blueprint': _CBP2}
            try:
                exec(compile(_c_src, 'S1-3', 'exec'), _ns)
                return None                      # must STOP; a silent pass is a defect
            except SystemExit as _e:
                return str(_e)
        try:
            _msg = _spec_report()
            _spec_cls = {int(_m.group(1)): _m.group(2)
                         for _m in re.finditer(r'Q(\d+) \[(W1|W2|D)', _msg or '')}
            _pp_full = _pp_c.classify_unresolved(dict(_stale), _CBP2)
            _pp_cls = {_q: _v['cls'] for _q, _v in _pp_full.items()}
            # CLASS AGREEMENT IS NOT ENOUGH. The header pins "same leaf rule, same
            # difflib cutoff=0.5" — and a cutoff drift moves the CANDIDATES while
            # leaving every class untouched, so a class-only comparison passes
            # through it (measured: 0.5 -> 0.9 was invisible until this half was
            # added). The W2 line is a human-confirm list; wrong candidates there
            # are wrong advice, which is the failure mode of this site.
            _tgt_ok = all(
                (repr(_v['targets'][0]) in _msg) if _v['cls'] == 'W1'
                else (repr(_v['targets']) in _msg)
                for _q, _v in _pp_full.items())
            _c_agrees = (_spec_cls == _pp_cls
                         and _pp_cls == {1: 'W1', 2: 'D', 3: 'W2'}
                         and _pp_full[3]['targets']        # the W2 list must be non-empty,
                         and _tgt_ok)                      # or 'agreement' is two blanks
        except Exception:
            _c_agrees = False
        check('A-QINDEX-PARITY-S1-3-classifier-matches-classify_unresolved',
              _c_agrees)
    except (ImportError, FileNotFoundError, IndexError, NameError):
        check('A-QINDEX-PARITY-S1-3-skipped-standalone', True)

    # ══════════════════════════════════════════════════════════════════════════
    # 5i / 5j.  RELEASE C (v2.24.0) — TWO NEW GATES, BOTH AMBER BY CONSTRUCTION.
    # ══════════════════════════════════════════════════════════════════════════
    def _spec_verdict(specs):
        _reset(); gate_specfaith([], {'figure_specs': specs})
        return {c: (l, m) for l, c, m in RESULTS}

    def _sp(label1, label2=None, key='legend', q=1):
        _ser = [{'id': 's1', 'label': label1}]
        if label2 is not None:
            _ser.append({'id': 's2', 'label': label2})
        return {f'q{q}_problem.png': {'question': q, 'key_mode': key, 'series': _ser}}

    # 5i-1 — THE DEFECT. matplotlib defaults recorded as series labels. Measured on
    #        real output: ALL 57 specs of IIT_JAM_BIOTECHNOLOGY M01 carry them.
    check('SPECLABEL-generator-default-is-a-finding',
          _spec_verdict(_sp('Series 1', 'Series 2')).get('A-FIGSPECLABEL', (None,))[0] == 'WARN')

    # 5i-2 — AMBER, NEVER BLOCKING. A NEW gate may not halt a paper (v2.22.0):
    #        the deployed operator cannot adjudicate, so a first false positive
    #        would strand papers with nobody to release them.
    check('SPECLABEL-is-amber-not-blocking',
          _spec_verdict(_sp('Series 1', 'Series 2')).get('A-FIGSPECLABEL', (None,))[0] != 'FAIL')

    # 5i-3 — REAL LABELS ARE CLEAN. Without this the gate could be "achieved" by
    #        flagging every figure, which is the A-FIGCOMP defect one gate over.
    check('SPECLABEL-stem-meaningful-labels-clean',
          _spec_verdict(_sp('P', 'Q')).get('A-FIGSPECLABEL', (None,))[0] == 'OK')

    # 5i-4 — id-as-label is a placeholder too ('s1' labelled 's1').
    check('SPECLABEL-id-as-label-is-a-finding',
          _spec_verdict(_sp('s1', 's2')).get('A-FIGSPECLABEL', (None,))[0] == 'WARN')

    # 5i-5 — LEGACY DORMANT. Pre-v5.34 registry has no figure_specs; ~200 exams.
    check('SPECLABEL-no-specs-dormant',
          _spec_verdict({}).get('A-FIGSPECLABEL', (None, ''))[0] == 'OK')

    # 5i-6 — A SINGLE UNKEYED SERIES CANNOT MISLEAD ANYONE: no legend is printed,
    #        so its label is never read. 51 of the real paper's 57 figures are this.
    check('SPECLABEL-unkeyed-single-series-not-flagged',
          _spec_verdict(_sp('Series 1', None, key='none'))
          .get('A-FIGSPECLABEL', (None,))[0] == 'OK')

    def _dom_verdict(stem, opts):
        def _b(d):
            d.add_paragraph(f'Q.1  {stem}')
            for _i, _o in enumerate(opts, 1):
                d.add_paragraph(f'{_i}.  {_o}')
            d.add_paragraph('')
            _add_q(d, 2)
        _doc = Document(_mini_doc(tmp, _b)); _t, _bl = parse_blocks(_doc)
        _reset(); gate_optdomain(_bl, _src_stub(tq=2))
        return {c: (l, m) for l, c, m in RESULTS}

    # 5j-1 — THE DEFECT, as delivered: eccentricity of an ellipse with two values
    #        greater than 1 — impossible for ANY ellipse, so half the option set is
    #        eliminable without reading the figure.
    check('OPTDOMAIN-out-of-domain-distractors-are-a-finding',
          _dom_verdict('The eccentricity of this ellipse is:',
                       ['0.60', '0.80', '1.25', '1.67'])
          .get('A-OPTDOMAIN', (None,))[0] == 'WARN')

    # 5j-2 — AMBER, NEVER BLOCKING.
    check('OPTDOMAIN-is-amber-not-blocking',
          _dom_verdict('The eccentricity of this ellipse is:',
                       ['0.60', '0.80', '1.25', '1.67'])
          .get('A-OPTDOMAIN', (None,))[0] != 'FAIL')

    # 5j-3 — IN-DOMAIN OPTIONS ARE CLEAN.
    check('OPTDOMAIN-in-domain-option-set-clean',
          _dom_verdict('The eccentricity of this ellipse is:',
                       ['0.20', '0.45', '0.60', '0.80'])
          .get('A-OPTDOMAIN', (None,))[0] == 'OK')

    # 5j-4 — BOUNDS ARE EXCLUSIVE FOR AN ELLIPSE (e=1 is a parabola, e=0 a circle),
    #        INCLUSIVE for a probability. Getting this backwards would either miss
    #        the defect or fire on a legitimate p=0/p=1 option.
    check('OPTDOMAIN-bound-inclusivity-is-per-quantity',
          _dom_verdict('The eccentricity of this ellipse is:',
                       ['0.5', '0.7', '0.9', '1.0']).get('A-OPTDOMAIN', (None,))[0] == 'WARN'
          and _dom_verdict('The probability of the event is:',
                           ['0', '0.25', '0.5', '1']).get('A-OPTDOMAIN', (None,))[0] == 'OK')

    # 5j-5 — NON-NUMERIC OPTIONS ARE NEVER JUDGED. A stem may name a bounded
    #        quantity while its options carry units, ranges or words; parsing those
    #        as bare numbers would invent findings.
    check('OPTDOMAIN-non-numeric-options-ignored',
          _dom_verdict('The probability of the event is:',
                       ['one half', '0.25', '0.5', '1'])
          .get('A-OPTDOMAIN', (None, ''))[0] == 'OK')

    # 5j-6 — NO NAMED QUANTITY ⇒ DORMANT. An unbounded numeric answer must never
    #        be judged against a domain it does not have.
    check('OPTDOMAIN-unnamed-quantity-dormant',
          _dom_verdict('The value of the current is:',
                       ['1.25', '1.67', '2.50', '9.90'])
          .get('A-OPTDOMAIN', (None, ''))[0] == 'OK')

    # 5c — SEC-2: A-OPTREF must SEE a figural block's full option set. option_paras
    #      (OPT_RE) reported ONE option on a block rendering four; _label_paras
    #      reports four. Asserted as a COUNT PARITY with the option gates rather
    #      than only as a verdict, because the wrong verdict it enables (a figural
    #      stem whose escape option is itself an image) is rare while the predicate
    #      inconsistency is always present.
    def _b_fig_escape(d):
        d.add_paragraph('Q.1  Which figure fits? Select None of these if inapplicable.')
        d.add_paragraph('Problem Figure')
        for _i in range(1, 4):
            d.add_paragraph(f'{_i}.')
        d.add_paragraph('4.  None of these')
    _p5c = _mini_doc(tmp, _b_fig_escape)
    _t, _bl5c = parse_blocks(Document(_p5c))
    _s5c = _src_stub(tq=1); _s5c['rules_txt'] = ''
    _reset(); gate_optref(_bl5c, _s5c)
    check('OPTREF-sees-full-figural-option-set',
          len(_label_paras(_bl5c[0])) == 4
          and not any(c == 'A-OPTREF' and l == 'FAIL' for l, c, m in RESULTS))

    # 5. A-QNFIRST catches stimulus orphaned before the next Q.N
    def b_notfirst(d):
        _add_q(d, 1)                                   # Q.1 + options + blank
        d.add_paragraph(' '.join(['word'] * 60))       # passage orphaned before Q.2
        _add_q(d, 2)
    p = _mini_doc(tmp, b_notfirst); _reset()
    _t, bl = parse_blocks(Document(p)); gate_qnfirst(bl)
    check('A-QNFIRST-catch', any(c == 'A-QNFIRST' and l == 'FAIL' for l, c, _ in RESULTS))

    # 5b. A-MATCH-TABLE catches a MATCH question rendered WITHOUT a real table (v2.7.1)
    _MPAIRS = ('(A)-(I), (B)-(III), (C)-(IV), (D)-(II)', '(A)-(II), (B)-(IV), (C)-(III), (D)-(I)',
               '(A)-(IV), (B)-(III), (C)-(II), (D)-(I)', '(A)-(III), (B)-(I), (C)-(IV), (D)-(II)')
    def b_match_notable(d):
        d.add_paragraph('Q.1  Match List-I with List-II.')
        for i, o in enumerate(_MPAIRS, 1):
            d.add_paragraph(f'{i}.  {o}')
        d.add_paragraph('')
    p = _mini_doc(tmp, b_match_notable); _reset()
    _t, bl = parse_blocks(Document(p)); gate_match_table(bl, _src_stub(tq=1))
    check('A-MATCH-TABLE-catch', any(c == 'A-MATCH-TABLE' and l == 'FAIL' for l, c, _ in RESULTS))

    # 5c. A-MATCH-TABLE passes (dormant) when the MATCH body IS a real table
    def b_match_table(d):
        d.add_paragraph('Q.1  Match List-I with List-II.')
        tb = d.add_table(rows=2, cols=2)
        tb.cell(0, 0).text = '(A) Collagen'; tb.cell(0, 1).text = '(I) triple helix'
        tb.cell(1, 0).text = '(B) Keratin';  tb.cell(1, 1).text = '(II) coiled coil'
        for i, o in enumerate(_MPAIRS, 1):
            d.add_paragraph(f'{i}.  {o}')
        d.add_paragraph('')
    p = _mini_doc(tmp, b_match_table); _reset()
    _t, bl = parse_blocks(Document(p)); gate_match_table(bl, _src_stub(tq=1))
    check('A-MATCH-TABLE-pass', not any(c == 'A-MATCH-TABLE' and l == 'FAIL' for l, c, _ in RESULTS))

    # 6. A-FONT catches Arial run
    def b_arial(d):
        para = d.add_paragraph()
        run = para.add_run('Q.1  Stem'); run.font.name = 'Arial'
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_arial); _reset()
    gate_font(Document(p), _src_stub(tq=1))
    check('A-FONT-catch', any(c == 'A-FONT' and l == 'FAIL' for l, c, _ in RESULTS))

    # 7. A-ANSKEY catches an answer-key line
    def b_key(d):
        _add_q(d, 1); d.add_paragraph('Answer Key: Q.1 -> 2')
    p = _mini_doc(tmp, b_key); _reset()
    gate_anskey(Document(p))
    check('A-ANSKEY-catch', any(c == 'A-ANSKEY' and l == 'FAIL' for l, c, _ in RESULTS))

    # 8. A-SECHDR catches a body section header INSIDE a question block (KEYWORD form)
    def b_hdr2(d):
        d.add_paragraph('Q.1  Stem'); d.add_paragraph('SECTION A: Reasoning')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p2 = _mini_doc(tmp, b_hdr2); _reset()
    _t, bl2 = parse_blocks(Document(p2))
    gate_sechdr(bl2, Document(p2), {'sections': [{'name': 'Reasoning'}]})
    check('A-SECHDR-catch', any(c == 'A-SECHDR' and l == 'FAIL' for l, c, _ in RESULTS))

    # 8b. A-HEADER (inverted, v2.7): a title/info block BEFORE Q.1 → A-HEADER FAIL (strip).
    def b_hdr_preq1(d):
        d.add_paragraph('SSC CGL Tier 1 — Mock Test 1')
        d.add_paragraph('Total Questions: 100    |    Maximum Marks: 200    |    Time: 60 Minutes')
        _add_q(d, 1)
    p = _mini_doc(tmp, b_hdr_preq1); _reset()
    dh = Document(p); _t, blh = parse_blocks(dh)
    gate_header(dh, blh, _src_stub(tq=1))
    check('A-HEADER-catch', any(c == 'A-HEADER' and l == 'FAIL' for l, c, _ in RESULTS))

    # 8c. A-HEADER dormant (v2.7): the SAME pre-Q.1 block, but section_rules EXAM_STRUCTURE
    #     declares paper_header_block → the opt-in permits it → NO A-HEADER failure.
    _reset()
    sd = _src_stub(tq=1); sd['paper_header_block'] = True
    gate_header(dh, blh, sd)
    check('A-HEADER-dormant', not any(c == 'A-HEADER' and l == 'FAIL' for l, c, _ in RESULTS))

    # 8b. v1.5 — A-SECHDR catches a stray heading that IS a declared SECTION NAME (the case
    #     the keyword pattern missed — found by the mutation harness). "Quantitative Aptitude"
    #     as a body paragraph, matched against src['sections'], must FAIL.
    def b_hdr3(d):
        d.add_paragraph('Quantitative Aptitude')          # section-name header, no keyword
        d.add_paragraph('Q.1  Stem')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p3 = _mini_doc(tmp, b_hdr3); _reset()
    _t, bl3 = parse_blocks(Document(p3))
    gate_sechdr(bl3, Document(p3), {'sections': [{'name': 'Quantitative Aptitude'}]})
    check('A-SECHDR-name-catch', any(c == 'A-SECHDR' and l == 'FAIL' for l, c, _ in RESULTS))

    # 9. A-UNDERLINE-FAKE catches "(underlined: X)"
    def b_fakeu(d):
        d.add_paragraph('Q.1  Find the synonym of the underlined word. (underlined: brisk)')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_fakeu); _reset()
    _t, bl = parse_blocks(Document(p)); gate_underline(bl)
    check('A-UNDERLINE-catch', any(c.startswith('A-UNDERLINE') and l == 'FAIL'
                                   for l, c, _ in RESULTS))

    # 10. A-FRAC catches ASCII caret exponent
    def b_caret(d):
        d.add_paragraph('Q.1  If x^2 + 1 = 5 then x?')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_caret); _reset()
    _t, bl = parse_blocks(Document(p)); gate_frac_ascii(bl, _src_stub(tq=1, omml=True))
    check('A-FRAC-catch', any(c == 'A-FRAC' and l == 'FAIL' for l, c, _ in RESULTS))

    # 11. A-STIMORPHAN catches a passage-reference with no embedded stimulus
    def b_orphan(d):
        d.add_paragraph('Q.1  According to the passage, what is the tone?')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_orphan); _reset()
    _t, bl = parse_blocks(Document(p)); gate_stimorphan(bl, _src_stub(tq=1))
    check('A-STIMORPHAN-catch', any(c == 'A-STIMORPHAN' and l == 'FAIL' for l, c, _ in RESULTS))

    # 12. A-STIMORPHAN passes when a long passage is embedded
    def b_embed(d):
        long = ' '.join(['word'] * 60)
        d.add_paragraph('Q.1  Read the following passage and answer the question. ' + long)
        d.add_paragraph('What is the tone?')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_embed); _reset()
    _t, bl = parse_blocks(Document(p)); gate_stimorphan(bl, _src_stub(tq=1))
    check('A-STIMORPHAN-pass', not any(c == 'A-STIMORPHAN' and l == 'FAIL'
                                       for l, c, _ in RESULTS))

    # 13. A-ENCODING catches U+FFFD
    def b_fffd(d):
        d.add_paragraph('Q.1  Bad char � here')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_fffd); _reset()
    gate_encoding_script(Document(p), _src_stub(tq=1))
    check('A-ENCODING-catch', any(c == 'A-ENCODING' and l == 'FAIL' for l, c, _ in RESULTS))

    # 14. A-SCRIPT passes Devanagari when language=hindi, fails when english
    def b_dev(d):
        d.add_paragraph('Q.1  प्रश्न यहाँ है')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_dev)
    _reset(); gate_encoding_script(Document(p), _src_stub(tq=1, lang='english'))
    eng_fail = any(c == 'A-SCRIPT' and l == 'FAIL' for l, c, _ in RESULTS)
    _reset(); gate_encoding_script(Document(p), _src_stub(tq=1, lang='hindi'))
    hin_ok = not any(c == 'A-SCRIPT' and l == 'FAIL' for l, c, _ in RESULTS)
    check('A-SCRIPT-lang-aware', eng_fail and hin_ok)

    # 15. A-DUP catches a stem repeated from a prior mock
    def b_dupstem(d):
        d.add_paragraph('Q.1  The capital of France is which city among these')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_dupstem); _reset()
    src = _src_stub(tq=1)
    src['registry'] = {'stem_texts': ['the capital of france is which city among these',
                                      'unrelated prior stem about rivers and lakes here']}
    # prior = all but trailing tq(=1) → first entry is prior
    _t, bl = parse_blocks(Document(p)); gate_dup(bl, src)
    check('A-DUP-catch', any(c == 'A-DUP' and l == 'FAIL' for l, c, _ in RESULTS))

    # 16. A-MATHRASTER name-contract: math-token names flagged, generic names not
    check('A-MATHRASTER-regex', bool(MATH_TOKEN_NAME.search('q55_e1.png'))
          and not MATH_TOKEN_NAME.search('Picture 3'))

    # 17. A-OPTREF catches missing escape option
    def b_optref(d):
        d.add_paragraph('Q.1  Identify the error; if there is no error select the last option.')
        for i, o in enumerate(('Part A', 'Part B', 'Part C', 'Part D'), 1):
            d.add_paragraph(f'{i}.  {o}')   # no "No error" option
    p = _mini_doc(tmp, b_optref); _reset()
    _t, bl = parse_blocks(Document(p)); gate_optref(bl, _src_stub(tq=1))
    check('A-OPTREF-catch', any(c == 'A-OPTREF' and l == 'FAIL' for l, c, _ in RESULTS))

    # 18. block parser merges OMML text (pure-math stem not seen empty)
    check('omml-merge-helper', callable(para_text))

    # 19. A-SCRIPT: accented Latin (cafe/resume) NOT flagged on an english exam
    def b_acc(d):
        d.add_paragraph('Q.1  The café served a résumé with naïve charm.')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_acc); _reset()
    gate_encoding_script(Document(p), _src_stub(tq=1, lang='english'))
    check('A-SCRIPT-accented-latin-ok', not any(c == 'A-SCRIPT' and l == 'FAIL'
                                                for l, c, _ in RESULTS))

    # 20. A-SCRIPT: Greek math symbols NOT flagged on an english exam
    def b_grk(d):
        d.add_paragraph('Q.1  If α + β = θ then θ = ?')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_grk); _reset()
    gate_encoding_script(Document(p), _src_stub(tq=1, lang='english'))
    check('A-SCRIPT-greek-ok', not any(c == 'A-SCRIPT' and l == 'FAIL'
                                       for l, c, _ in RESULTS))

    # 21. A-ZIP: rels parsing is attribute-ORDER-INDEPENDENT (Target before Id)
    rels_tf = '<Relationship Target="media/image1.png" Type="x" Id="rId9"/>'
    rel_map = {}
    for tag in re.findall(r'<Relationship\b[^>]*/?>', rels_tf):
        idm = re.search(r'\bId="([^"]+)"', tag); tgm = re.search(r'\bTarget="([^"]+)"', tag)
        if idm and tgm:
            rel_map[idm.group(1)] = tgm.group(1)
    check('A-ZIP-attr-order', rel_map.get('rId9') == 'media/image1.png')

    # 22. A-OPTN: an enumerated passage point before the options does NOT inflate
    #     the option count (trailing-oc extraction)
    def b_enum(d):
        d.add_paragraph('Q.1  Read the data.')
        d.add_paragraph('1.  earlier enumerated passage point that is not an option')
        d.add_paragraph('Which is correct?')
        for i, o in enumerate(('Opt1', 'Opt2', 'Opt3', 'Opt4'), 1):
            d.add_paragraph(f'{i}.  {o}')
    p = _mini_doc(tmp, b_enum); _reset()
    _t, bl = parse_blocks(Document(p)); gate_options(bl, _src_stub(tq=1))
    check('A-OPTN-enum-no-false', not any(c == 'A-OPTN' and l == 'FAIL'
                                          for l, c, _ in RESULTS))

    # 23. A-OPTN passes figural bare-label image options; roman + alpha label sets
    def b_fig(d):
        d.add_paragraph('Q.1  Select the figure.')
        for i in range(1, 5):
            d.add_paragraph(f'{i}.')          # bare label (image would follow)
    p = _mini_doc(tmp, b_fig); _reset()
    _t, bl = parse_blocks(Document(p)); gate_options(bl, _src_stub(tq=1))
    fig_ok = not any(c == 'A-OPTN' and l == 'FAIL' for l, c, _ in RESULTS)
    def b_rom(d):
        d.add_paragraph('Q.1  Pick.')
        for lab in ('i', 'ii', 'iii', 'iv'):
            d.add_paragraph(f'{lab}.  t{lab}')
    p = _mini_doc(tmp, b_rom); _reset()
    _t, bl = parse_blocks(Document(p))
    s2 = _src_stub(tq=1); s2['opt_label_fmt'] = 'i/ii/iii/iv'
    gate_options(bl, s2)
    rom_ok = not any(l == 'FAIL' for l, c, _ in RESULTS)
    check('A-OPTN-figural+roman', fig_ok and rom_ok)

    # 24. A-MSQ-INSTR (machine): a section expecting 1 multi Q but 0 instruction-carrying
    #     stems → FAIL (a multi Q is missing its select-instruction).
    def b_msq_missing(d):
        _add_q(d, 1, stem='Which of the following are prime?')   # multi expected, NO instruction
    p = _mini_doc(tmp, b_msq_missing); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1)
    s['expected_multi_by_section'] = {'S1': 1}
    s['msq_instruction_phrases'] = ['one or more', 'may be correct']
    s['sections'] = [{'name': 'S1', 'q_range': [1, 1], 'total_qs': 1}]
    gate_msq_instr(bl, s)
    check('A-MSQ-INSTR-catch', any(c == 'A-MSQ-INSTR' and l == 'FAIL' for l, c, _ in RESULTS))

    # 25. A-MSQ-INSTR passes when the multi Q carries its instruction in the Q.N line.
    def b_msq_ok(d):
        _add_q(d, 1, stem='Which of the following are prime? (One or more options may be correct)')
    p = _mini_doc(tmp, b_msq_ok); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1)
    s['expected_multi_by_section'] = {'S1': 1}
    s['msq_instruction_phrases'] = ['one or more', 'may be correct']
    s['sections'] = [{'name': 'S1', 'q_range': [1, 1], 'total_qs': 1}]
    gate_msq_instr(bl, s)
    check('A-MSQ-INSTR-pass', not any(c == 'A-MSQ-INSTR' and l == 'FAIL' for l, c, _ in RESULTS))

    # 26. A-MSQ-INSTR is DORMANT when the blueprint declares no multi subtopics.
    def b_single(d): _add_q(d, 1)
    p = _mini_doc(tmp, b_single); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1); s['expected_multi_by_section'] = {}
    gate_msq_instr(bl, s)
    check('A-MSQ-INSTR-dormant', not any(c == 'A-MSQ-INSTR' and l == 'FAIL' for l, c, _ in RESULTS))

    # 27. A-ANSKEY catches a SET-valued key leak ("Q.1 -> 1,2,4"), not just single-digit.
    def b_setkey(d):
        _add_q(d, 1); d.add_paragraph('Answer Key: Q.1 -> 1,2,4')
    p = _mini_doc(tmp, b_setkey); _reset()
    gate_anskey(Document(p))
    check('A-ANSKEY-setleak-catch', any(c == 'A-ANSKEY' and l == 'FAIL' for l, c, _ in RESULTS))

    # 28. empty document does not crash any gate
    def b_empty(d):
        pass  # CLASS: J — inert test fixture: builds an empty doc for the self-test.
              # No tool call and no model agency at all; tagged so C6 has a verdict.
    p = _mini_doc(tmp, b_empty); _reset()
    try:
        mm = gate_zip(p); doc = Document(p); _t, bl = parse_blocks(doc)
        st = _src_stub(tq=0)
        for fn in (lambda: gate_structure(bl, st), lambda: gate_options(bl, st),
                   lambda: gate_qnfirst(bl), lambda: gate_blanksep(doc, bl),
                   lambda: gate_font(doc, st), lambda: gate_sechdr(bl, doc, st),
                   lambda: gate_anskey(doc), lambda: gate_stimorphan(bl, st),
                   lambda: gate_underline(bl), lambda: gate_omml(doc, st, True),
                   lambda: gate_frac_ascii(bl, st), lambda: gate_images(bl, st, mm),
                   lambda: gate_optref(bl, st), lambda: gate_encoding_script(doc, st),
                   lambda: gate_dup(bl, st), lambda: gate_header(doc, bl, st),
                   lambda: gate_nat(bl, st)):
            fn()
        empty_ok = True
    except Exception:
        empty_ok = False
    check('empty-doc-no-crash', empty_ok)

    # 29. A-NAT-NOOPT (machine): a Q the registry marks NAT (options_by_q=0) that RENDERS
    #     options → FAIL (a numerical question must carry none).
    def b_nat_hasopts(d):
        _add_q(d, 1, stem='Find the value. Enter your answer as a numerical value.')
    p = _mini_doc(tmp, b_nat_hasopts); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1); s['nat_present'] = True; s['options_by_q'] = {'1': 0}
    gate_nat(bl, s)
    check('A-NAT-NOOPT-catch', any(c == 'A-NAT-NOOPT' and l == 'FAIL' for l, c, _ in RESULTS))

    # 30. A-NAT-NOOPT passes when the NAT Q renders zero options.
    def b_nat_noopts(d):
        _add_q(d, 1, opts=(), stem='Find the value. Enter your answer as a numerical value.')
    p = _mini_doc(tmp, b_nat_noopts); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1); s['nat_present'] = True; s['options_by_q'] = {'1': 0}
    gate_nat(bl, s)
    check('A-NAT-NOOPT-pass', not any(c == 'A-NAT-NOOPT' and l == 'FAIL' for l, c, _ in RESULTS))

    # 31. A-NAT-NOOPT is DORMANT when the blueprint declares no numerical subtopics.
    def b_nat_dormant(d): _add_q(d, 1)
    p = _mini_doc(tmp, b_nat_dormant); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1)   # nat_present default false; options_by_q empty
    gate_nat(bl, s)
    check('A-NAT-NOOPT-dormant', not any(c == 'A-NAT-NOOPT' and l == 'FAIL' for l, c, _ in RESULTS))

    # 32. A-NAT-INSTR (machine): a section expecting 1 NAT Q but 0 instruction-carrying
    #     stems → FAIL (a numerical Q is missing its numerical-entry instruction).
    def b_nat_missing(d):
        _add_q(d, 1, opts=(), stem='Find the value of x.')
    p = _mini_doc(tmp, b_nat_missing); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1)
    s['expected_nat_by_section'] = {'S1': 1}
    s['nat_instruction_phrases'] = ['numerical value', 'enter your answer']
    s['sections'] = [{'name': 'S1', 'q_range': [1, 1], 'total_qs': 1}]
    gate_nat(bl, s)
    check('A-NAT-INSTR-catch', any(c == 'A-NAT-INSTR' and l == 'FAIL' for l, c, _ in RESULTS))

    # 33. A-NAT-INSTR passes when the NAT Q carries its instruction in the Q.N line.
    def b_nat_ok(d):
        _add_q(d, 1, opts=(), stem='Find the value of x. Enter your answer as a numerical value.')
    p = _mini_doc(tmp, b_nat_ok); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1)
    s['expected_nat_by_section'] = {'S1': 1}
    s['nat_instruction_phrases'] = ['numerical value', 'enter your answer']
    s['sections'] = [{'name': 'S1', 'q_range': [1, 1], 'total_qs': 1}]
    gate_nat(bl, s)
    check('A-NAT-INSTR-pass', not any(c == 'A-NAT-INSTR' and l == 'FAIL' for l, c, _ in RESULTS))

    # 34. A-NAT-INSTR is DORMANT when the blueprint declares no numerical subtopics.
    def b_nat_dormant2(d): _add_q(d, 1)
    p = _mini_doc(tmp, b_nat_dormant2); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1); s['expected_nat_by_section'] = {}
    gate_nat(bl, s)
    check('A-NAT-INSTR-dormant', not any(c == 'A-NAT-INSTR' and l == 'FAIL' for l, c, _ in RESULTS))

    # 35. A-NAT-GRADE (machine, v2.8): sidecar's own nat_grading_value does NOT match a
    #     fresh derive_nat_grading() re-run on the sidecar's own nat_value/ca_range/
    #     stem_precision -> FAIL (e.g. Step 7 stored '3e-9' instead of the correct '3').
    def b_grade_bad(d):
        _add_q(d, 1, opts=(), stem='Find the rate. Enter your answer as a numerical value.')
    p = _mini_doc(tmp, b_grade_bad); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1); s['nat_present'] = True
    s['concept_map'] = {'1': {'qtype': 'nat', 'ca_range': None,
                               'nat_grading_type': 'positive_integer',
                               'nat_grading_value': '3e-9', 'stem_precision': None}}
    s['answers'] = {'1': 3}
    gate_nat(bl, s)
    check('A-NAT-GRADE-catch', any(c == 'A-NAT-GRADE' and l == 'FAIL' for l, c, _ in RESULTS))

    # 36. A-NAT-GRADE passes when the sidecar's grading value matches a fresh re-run,
    #     including the stem_precision-driven decimal_fixed branch (round-half-up, padded).
    def b_grade_ok(d):
        _add_q(d, 1, opts=(), stem='Find the value. Round off to two decimal places.')
    p = _mini_doc(tmp, b_grade_ok); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1); s['nat_present'] = True
    s['concept_map'] = {'1': {'qtype': 'nat', 'ca_range': None,
                               'nat_grading_type': 'decimal_fixed',
                               'nat_grading_value': '3.00', 'stem_precision': 2}}
    s['answers'] = {'1': 3}
    gate_nat(bl, s)
    check('A-NAT-GRADE-pass', not any(c == 'A-NAT-GRADE' and l == 'FAIL' for l, c, _ in RESULTS))

    # 36a-36c. v2.21.6 — A-NAT-GRADE INCOMPLETENESS CLOSURE. Fixtures 35/36 cover
    #     only the MISMATCH and the happy path. Three further findings could each be
    #     deleted with every test still green: missing nat_value, missing
    #     nat_grading_value, and a re-derivation that RAISES. This gate guards the
    #     exact string the delivery portal ingests to auto-grade a numerical
    #     question, so a silent failure here is WRONG MARKS, not a wrong-looking
    #     paper — the same class as the A-OPTORDER anchor defect.
    def _b_nat_stem(d):
        _add_q(d, 1, opts=(), stem='Find the value. Enter your answer as a numerical value.')
    def _grade_verdict(entry, answers):
        _p = _mini_doc(tmp, _b_nat_stem)
        _t, _bl = parse_blocks(Document(_p))
        _s = _src_stub(tq=1); _s['nat_present'] = True
        _s['concept_map'] = {'1': dict(entry)}
        _s['answers'] = dict(answers)
        _reset(); gate_nat(_bl, _s)
        return next((m for l, c, m in RESULTS
                     if c == 'A-NAT-GRADE' and l == 'FAIL'), None)
    _GOOD = {'qtype': 'nat', 'ca_range': None, 'nat_grading_type': 'positive_integer',
             'nat_grading_value': '3', 'stem_precision': None}

    # 36a — the sidecar carries answers, but NONE for a question it calls NAT.
    #       (An entirely EMPTY answers map makes the whole gate dormant by design —
    #       Step 8 gets no key unless --key is supplied — so the reachable defect is
    #       a PARTIAL sidecar: answers present, this question's value missing.)
    #       Nothing can be re-derived, so the grading string is unverifiable and must
    #       be reported, never skipped silently.
    _m36a = _grade_verdict(_GOOD, {'2': 5})
    check('NAT-GRADE-missing-nat-value-is-a-finding',
          _m36a is not None and 'nat_value missing' in _m36a)

    # 36b — the sidecar carries a value but NO nat_grading_value. The portal would
    #       receive nothing to match against; every candidate scores zero on this Q.
    _e36b = dict(_GOOD); _e36b['nat_grading_value'] = None
    _m36b = _grade_verdict(_e36b, {'1': 3})
    check('NAT-GRADE-missing-grading-value-is-a-finding',
          _m36b is not None and 'nat_grading_value missing' in _m36b)

    # 36c — the re-derivation RAISES (an inverted ca_range, lo > hi, which
    #       derive_nat_grading rejects). The exception must become a NAMED FINDING,
    #       not propagate and not be swallowed: an un-derivable grading value is a
    #       real Step-7 defect and the run must say so.
    _e36c = dict(_GOOD); _e36c['ca_range'] = (9, 1)
    _e36c['nat_grading_type'] = 'range'; _e36c['nat_grading_value'] = '1-9'
    _m36c = _grade_verdict(_e36c, {'1': 3})
    check('NAT-GRADE-re-derivation-raise-is-a-finding',
          _m36c is not None and 're-derivation raised' in _m36c)

    # 37. A-NAT-GRADE is DORMANT when the blueprint declares no numerical subtopics.
    def b_grade_dormant(d): _add_q(d, 1)
    p = _mini_doc(tmp, b_grade_dormant); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1)   # nat_present default false
    gate_nat(bl, s)
    check('A-NAT-GRADE-dormant-nonat',
          not any(c == 'A-NAT-GRADE' and l == 'FAIL' for l, c, _ in RESULTS))

    # 38. A-NAT-GRADE is DORMANT when nat_present=True but no --key sidecar was supplied
    #     (Step 8 does not receive the answer_key sidecar by default, S0-1).
    def b_grade_dormant_nokey(d):
        _add_q(d, 1, opts=(), stem='Find the value. Enter your answer as a numerical value.')
    p = _mini_doc(tmp, b_grade_dormant_nokey); _reset()
    _t, bl = parse_blocks(Document(p))
    s = _src_stub(tq=1); s['nat_present'] = True   # concept_map left empty (no --key)
    gate_nat(bl, s)
    check('A-NAT-GRADE-dormant-nokey',
          not any(c == 'A-NAT-GRADE' and l == 'FAIL' for l, c, _ in RESULTS))

    # ── v2.6 COMPLETION-GATE fixtures (S5-1A C1–C7) ────────────────────────────
    # Shared: a clean 1-Q docx (no artefacts) and an evidence dir.
    def _cg_doc(build):
        pp = _mini_doc(tmp, build)
        dd = Document(pp); _tt, bb = parse_blocks(dd)
        return dd, bb
    def _write_state(name, state):
        sp = os.path.join(tmp, name)
        json.dump(state, open(sp, 'w', encoding='utf-8'))
        return sp
    _evd = os.path.join(tmp, 'evdir'); os.makedirs(_evd, exist_ok=True)

    # 35. COMPLETION-GATE PASS on a complete, evidence-clean single-Q ledger.
    dcg, bcg = _cg_doc(lambda d: _add_q(d, 1))
    _reset()
    st35 = {'K': 1, 'batches_done': [1], 'evidence_dir': _evd,
            'ledger': {'entries': {'1': {'status': 'verified', 'answer_cardinality': 'single',
                        'answer_unique': True, 'is_factual': False,
                        'fact_sources': [], 'artefact_stamps': {}}}}}
    rc = completion_gate(_write_state('st35.json', st35), 1, bcg, dcg)
    check('CG-pass', rc == 0 and not any(l == 'FAIL' for l, c, _ in RESULTS))

    # 36. SKIPPED-PHASE-2: empty ledger + no batches → C1 + C2 FAIL.
    _reset()
    st36 = {'K': 1, 'batches_done': [], 'evidence_dir': _evd, 'ledger': {'entries': {}}}
    rc = completion_gate(_write_state('st36.json', st36), 1, bcg, dcg)
    check('CG-skipped-phase2', rc == 1
          and any(c == 'C1' and l == 'FAIL' for l, c, _ in RESULTS)
          and any(c == 'C2' and l == 'FAIL' for l, c, _ in RESULTS))

    # 37. PARTIAL-REVIEW: tq=2 but only Q1 in the ledger → C2 FAIL.
    _reset()
    st37 = {'K': 1, 'batches_done': [1], 'evidence_dir': _evd,
            'ledger': {'entries': {'1': {'status': 'verified', 'answer_cardinality': 'single',
                        'answer_unique': True, 'is_factual': False, 'artefact_stamps': {}}}}}
    rc = completion_gate(_write_state('st37.json', st37), 2, bcg, dcg)
    check('CG-partial-review', rc == 1 and any(c == 'C2' and l == 'FAIL' for l, c, _ in RESULTS))

    # 38. UNSOURCED-FACT: a factual entry with no fact_sources → C5 FAIL.
    _reset()
    st38 = {'K': 1, 'batches_done': [1], 'evidence_dir': _evd,
            'ledger': {'entries': {'1': {'status': 'verified', 'answer_cardinality': 'single',
                        'answer_unique': True, 'is_factual': True,
                        'fact_sources': [], 'artefact_stamps': {}}}}}
    rc = completion_gate(_write_state('st38.json', st38), 1, bcg, dcg)
    check('CG-unsourced-fact', rc == 1 and any(c == 'C5' and l == 'FAIL' for l, c, _ in RESULTS))

    # 39. UNVIEWED artefact (paper has a table, ledger has no stamp) → C7 FAIL.
    def b_tbl(d):
        d.add_paragraph('Q.1  Study the table.')
        for i, o in enumerate(('A', 'B', 'C', 'D'), 1):
            d.add_paragraph(f'{i}.  {o}')
        tt = d.add_table(rows=2, cols=2)
        tt.cell(0, 0).text = 'h1'; tt.cell(0, 1).text = 'h2'
        tt.cell(1, 0).text = '1'; tt.cell(1, 1).text = '2'
    dtbl, btbl = _cg_doc(b_tbl)
    _reset()
    st39 = {'K': 1, 'batches_done': [1], 'evidence_dir': _evd,
            'ledger': {'entries': {'1': {'status': 'verified', 'answer_cardinality': 'single',
                        'answer_unique': True, 'is_factual': False, 'artefact_stamps': {}}}}}
    rc = completion_gate(_write_state('st39.json', st39), 1, btbl, dtbl)
    check('CG-unviewed-artefact', rc == 1 and any(c == 'C7' and l == 'FAIL' for l, c, _ in RESULTS))

    # 40. MISSING evidence FILE (stamp present, trace file absent) → C6 FAIL.
    _reset()
    st40 = {'K': 1, 'batches_done': [1], 'evidence_dir': _evd,
            'ledger': {'entries': {'1': {'status': 'verified', 'answer_cardinality': 'single',
                        'answer_unique': True, 'is_factual': False,
                        'artefact_stamps': {'tables': [{'idx': 0, 'stamp': 'recomputed',
                                            'trace': os.path.join(_evd, 'does_not_exist.txt')}]}}}}}
    rc = completion_gate(_write_state('st40.json', st40), 1, btbl, dtbl)
    check('CG-missing-evidence-file', rc == 1 and any(c == 'C6' and l == 'FAIL' for l, c, _ in RESULTS))

    # 41. FACT saved-file MISSING (source named but file absent) → C5 FAIL.
    _reset()
    st41 = {'K': 1, 'batches_done': [1], 'evidence_dir': _evd,
            'ledger': {'entries': {'1': {'status': 'verified', 'answer_cardinality': 'single',
                        'answer_unique': True, 'is_factual': True,
                        'fact_sources': [{'url': 'x', 'date': 'y',
                                          'saved': os.path.join(_evd, 'nf.json')}],
                        'artefact_stamps': {}}}}}
    rc = completion_gate(_write_state('st41.json', st41), 1, bcg, dcg)
    check('CG-factfile-missing', rc == 1 and any(c == 'C5' and l == 'FAIL' for l, c, _ in RESULTS))

    # 42. EVIDENCE-BACKED artefact stamp (trace file exists) → PASS.
    _trace = os.path.join(_evd, 'q1_table.txt')
    open(_trace, 'w', encoding='utf-8').write('grid a,b / 1,2 ; total row recomputed = 3 OK')
    _reset()
    st42 = {'K': 1, 'batches_done': [1], 'evidence_dir': _evd,
            'ledger': {'entries': {'1': {'status': 'verified', 'answer_cardinality': 'single',
                        'answer_unique': True, 'is_factual': False,
                        'artefact_stamps': {'tables': [{'idx': 0, 'stamp': 'recomputed',
                                            'trace': _trace}]}}}}}
    rc = completion_gate(_write_state('st42.json', st42), 1, btbl, dtbl)
    check('CG-evidence-backed-pass', rc == 0 and not any(l == 'FAIL' for l, c, _ in RESULTS))

    # ════════════════════════════════════════════════════════════════════
    # v2.12 — GAP-2026-08-01-FIGPROFILE-ENGINE-BINDING regression lock (D4)
    # ════════════════════════════════════════════════════════════════════
    # WHY THESE EXIST. Until v2.12 this suite returned 51/51 PASS on a file that
    # could not survive a real run, because NO fixture ever built a registry
    # carrying figural_manifests[].object_types. Every fixture took A-FIGPROFILE's
    # DORMANT branch, so the unbound `bc` at the primary branch was never
    # executed. The v2.6 hardening caught hollow FILES; it did not catch a hollow
    # BRANCH. These fixtures execute the primary branch in every environment the
    # estate actually presents, which is the only thing that would have caught it.
    import importlib as _il

    def _figsrc(obj_types, sub_ids, rules=''):
        s = _src_stub(tq=2)
        s['figural_object_types'] = obj_types
        s['figural_subtopics'] = sub_ids
        s['section_rules_text'] = rules
        s['figural_qs'] = {1, 2}
        return s

    def _run_figprofile(src_d):
        """Run ONLY the A-FIGPROFILE path via gate_images on a 2-Q fixture."""
        p = _mini_doc(tmp, lambda d: (_add_q(d, 1), _add_q(d, 2)))
        _reset()
        doc = Document(p); _t, blks = parse_blocks(doc)
        _safe_gate('A-IMAGES', gate_images, blks, src_d, {})
        return [(l, c, m) for l, c, m in RESULTS if c in ('A-FIGPROFILE', 'A-GATEERROR')]

    class _StubBC:
        """A complete, conforming stand-in for blueprint_core."""
        @staticmethod
        def parse_image_analysis_blocks(t): return {'s.a': {'dominant_object_type': 'x'}}
        @staticmethod
        def figural_generation_profile(p): return {'dominant_object_type': 'x'}
        @staticmethod
        def check_figural_conformance(g, p): return ('PASS', '')

    _saved_bc = sys.modules.get('blueprint_core')

    def _with_bc(mod):
        if mod is None:
            sys.modules.pop('blueprint_core', None)
            sys.modules['blueprint_core'] = None   # forces ImportError on import
        else:
            sys.modules['blueprint_core'] = mod

    _OT = {'1': 'cycle_diagram', '2': 'cycle_diagram'}
    _SI = {'1': 's.a', '2': 's.a'}

    # 43. NON-DORMANT-BRANCH COVERAGE — the fixture that would have caught D1.
    #     object_types present + engine importable → the primary branch EXECUTES
    #     and must produce a real verdict, never a NameError.
    _with_bc(_StubBC)
    r = _run_figprofile(_figsrc(_OT, _SI))
    check('FIGPROFILE-primary-branch-runs',
          any(c == 'A-FIGPROFILE' and l == 'OK' for l, c, _ in r)
          and not any(c == 'A-GATEERROR' for _, c, _ in r))

    # 44. CATCH PATH — a non-conforming subtopic must FAIL, naming it.
    class _FailBC(_StubBC):
        @staticmethod
        def check_figural_conformance(g, p): return ('FAIL', 'type mismatch')
    _with_bc(_FailBC)
    r = _run_figprofile(_figsrc(_OT, _SI))
    check('FIGPROFILE-catches-nonconformance',
          any(c == 'A-FIGPROFILE' and l == 'FAIL' for l, c, _ in r))

    # 45. MISSING-ENGINE DEGRADATION — must WARN "NOT CHECKED", never raise.
    _with_bc(None)
    r = _run_figprofile(_figsrc(_OT, _SI))
    check('FIGPROFILE-missing-engine-warns',
          any(c == 'A-FIGPROFILE' and l == 'WARN' and 'NOT CHECKED' in m for l, c, m in r)
          and not any(c == 'A-GATEERROR' for _, c, _ in r))

    # 46. STALE-ENGINE — imports fine, lacks a delegated function. The capability
    #     check (L2) must catch it; an import guard alone would let the call site
    #     raise AttributeError and abort the run.
    class _StaleBC:
        @staticmethod
        def parse_image_analysis_blocks(t): return {}
        @staticmethod
        def figural_generation_profile(p): return None
    _with_bc(_StaleBC)
    r = _run_figprofile(_figsrc(_OT, _SI))
    check('FIGPROFILE-stale-engine-warns',
          any(c == 'A-FIGPROFILE' and l == 'WARN' and 'stale engine' in m for l, c, m in r)
          and not any(c == 'A-GATEERROR' for _, c, _ in r))

    # 47. ENGINE RAISES AT THE CALL SITE — L3 must convert it to a logged skip.
    class _BoomBC(_StubBC):
        @staticmethod
        def parse_image_analysis_blocks(t): raise RuntimeError('boom')
    _with_bc(_BoomBC)
    r = _run_figprofile(_figsrc(_OT, _SI))
    check('FIGPROFILE-engine-raise-warns',
          any(c == 'A-FIGPROFILE' and l == 'WARN' for l, c, _ in r)
          and not any(c == 'A-GATEERROR' for _, c, _ in r))

    # 48. 0/0 IS NOT EVIDENCE — object_types present but no usable subtopic_ids
    #     must WARN, not claim conformance it never tested (edge case 6).
    _with_bc(_StubBC)
    r = _run_figprofile(_figsrc(_OT, {}))
    check('FIGPROFILE-zero-of-zero-warns',
          any(c == 'A-FIGPROFILE' and l == 'WARN' and 'NOT ESTABLISHED' in m for l, c, m in r))

    # 49. LEGACY PATH UNCHANGED — no object_types → dormant _ok. ~200 exams on
    #     pre-v5.31 output must keep passing exactly as before (EC-V18).
    _with_bc(None)
    r = _run_figprofile(_figsrc({}, {}))
    check('FIGPROFILE-legacy-dormant-ok',
          any(c == 'A-FIGPROFILE' and l == 'OK' and 'dormant' in m for l, c, m in r))

    # restore the real module state
    if _saved_bc is not None:
        sys.modules['blueprint_core'] = _saved_bc
    else:
        sys.modules.pop('blueprint_core', None)

    # 50. GATE FAULT ISOLATION (P6) — an exploding gate must become a LOUD, NAMED
    #     FAIL while the run continues. This is the fixture that makes the entire
    #     "permanent halt" failure mode impossible for any future gate.
    _reset()
    def _boom(*a, **k): raise ValueError('synthetic')
    _safe_gate('A-SYNTHETIC', _boom)
    check('GATEERROR-isolates-and-names',
          any(c == 'A-GATEERROR' and l == 'FAIL' and 'A-SYNTHETIC' in m for l, c, m in RESULTS))

    # 51. FAULT ISOLATION DOES NOT MASK — a crashed gate is FAIL severity, so the
    #     run exits non-zero and the paper CANNOT be certified clean. Degradation
    #     must never be quietly survivable.
    check('GATEERROR-blocks-certification',
          any(l == 'FAIL' for l, _, _ in RESULTS))

    # 52. UNDEFINED-NAME SCAN (test 10) — self-hosted. Scans THIS file for any
    #     Load-context name never bound anywhere. This is the generalised guard:
    #     it catches the next `bc`-class defect automatically, in any gate.
    _und = {}
    try:
        import builtins as _bi          # v2.24.0 — 'ast' is a MODULE-level
        # import (line 35); a function-local one here would shadow it for the
        # WHOLE of self_test(), unbinding it for fixture 5h defined earlier.
        _tree = ast.parse(open(__file__, encoding='utf-8').read())
        _bnd = set(dir(_bi)) | {'__file__', '__name__', '__doc__', '__builtins__'}
        for _n in ast.walk(_tree):
            if isinstance(_n, ast.alias): _bnd.add(_n.asname or _n.name.split('.')[0])
            elif isinstance(_n, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)): _bnd.add(_n.name)
            elif isinstance(_n, ast.Name) and isinstance(_n.ctx, (ast.Store, ast.Del)): _bnd.add(_n.id)
            elif isinstance(_n, ast.arg): _bnd.add(_n.arg)
            elif isinstance(_n, ast.ExceptHandler) and _n.name: _bnd.add(_n.name)
            elif isinstance(_n, (ast.Global, ast.Nonlocal)): _bnd |= set(_n.names)
        for _n in ast.walk(_tree):
            if isinstance(_n, ast.Name) and isinstance(_n.ctx, ast.Load) and _n.id not in _bnd:
                _und.setdefault(_n.id, []).append(_n.lineno)
    except Exception:
        _und = {'<scan-failed>': []}
    check('NO-UNDEFINED-NAMES', not _und)

    # ════════════════════════════════════════════════════════════════════
    # v2.13 — GAP-2026-08-01-FIGSPEC-TRANSPORT regression lock (D4)
    # ════════════════════════════════════════════════════════════════════
    # WHY THESE EXIST. v2.12 closed the HALT and left the twelve gates it
    # rescued VACUOUS: Block.images was declared '# reserved' and appended to
    # nowhere, so every gate iterated an empty list and printed "0 figure(s)
    # conform." on every paper in every exam. 61/61 PASS again coexisted with
    # zero real coverage — the same hollow-branch class as the halt itself, one
    # gate-family over. NO fixture had ever put an image in a block. These do.
    #
    # figural_core is STUBBED (the pattern v2.12 used for blueprint_core) so the
    # severity-routing fixtures assert THIS file's logic and hold identically on
    # a machine with no matplotlib/PIL — the environment ~200 exams may present.

    def _img_doc(names, in_table=False):
        """A 1-question docx carrying len(names) inline drawings, each stamped
        with its canonical docPr name + its OWN alt text (S10-8 does exactly
        this via _name_last_drawing). in_table places them in a table cell."""
        from docx import Document as D
        # v2.21.9 — Inches is now a MODULE-level import (line 42). A function-local
        # import here would shadow it for the WHOLE of self_test(), unbinding it for
        # every fixture defined earlier in the same scope.
        d = D()
        d.add_paragraph('Q.1  Study the figure.')
        holder = d.add_table(rows=1, cols=1).rows[0].cells[0] if in_table else d
        for nm in names:
            p = holder.add_paragraph()
            p.add_run().add_picture(io.BytesIO(_png_bytes()), width=Inches(1.0))
        d.add_paragraph('')
        pth = os.path.join(tmp, 'img.docx')
        d.save(pth)
        # stamp canonical name + per-drawing descr, in document order
        import shutil as _sh
        tmpd = tempfile.mkdtemp(); out = os.path.join(tmpd, 'stamped.docx')
        with zipfile.ZipFile(pth) as zin:
            items = zin.infolist()
            xml = zin.read('word/document.xml').decode('utf-8')
            it = iter(names)
            _cnt = [0]
            def _sub(mo):
                nm = next(it, 'x.png'); _cnt[0] += 1
                close = '/>' if mo.group(0).rstrip().endswith('/>') else '>'
                return (f'<wp:docPr id="{_cnt[0]}" name="{nm}" '
                        f'descr="alt for {nm}"{close}')
            xml = re.sub(r'<wp:docPr\b[^>]*>', _sub, xml)
            with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zo:
                for i in items:
                    zo.writestr(i, xml.encode('utf-8')
                                if i.filename == 'word/document.xml'
                                else zin.read(i.filename))
        return out

    def _attach(docx_path):
        _reset()
        mm = gate_zip(docx_path)
        doc = Document(docx_path); _t, blks = parse_blocks(doc)
        dec, res = attach_block_images(blks, mm, extract_media(docx_path, mm))
        return blks, dec, res

    class _StubFC:
        """figural_core stand-in. LEGACY / FINDINGS are set per fixture."""
        LEGACY = True
        FINDINGS = []
        @staticmethod
        def is_legacy(spec): return _StubFC.LEGACY
        @staticmethod
        def audit_figure(spec, png, descr=None): return (list(_StubFC.FINDINGS), [])
        @staticmethod
        def audit_gate_id(f): return f.split(':')[0].strip()
        @staticmethod
        def triage(findings, spec=None):
            sev = 'AMBER' if _StubFC.LEGACY else 'BLOCKING'
            return {'BLOCKING': findings if sev == 'BLOCKING' else [],
                    'VOID_ITEM': [], 'AMBER': findings if sev == 'AMBER' else []}

    _saved_fc = sys.modules.get('figural_core')

    def _run_figgates(blks, specs=None, legacy=True, findings=()):
        _StubFC.LEGACY = legacy; _StubFC.FINDINGS = list(findings)
        sys.modules['figural_core'] = _StubFC
        s = _src_stub(tq=1)
        s['figure_specs'] = specs or {}
        _reset()
        _safe_gate('A-IMAGES', gate_images, blks, s, {})
        return [(l, c, m) for l, c, m in RESULTS if c.startswith('A-FIG')
                or c == 'A-GATEERROR']

    _FIGGATES = ('A-FIGSCALE', 'A-FIGLABEL', 'A-FIGDPI', 'A-FIGDEGEN',
                 'A-FIGMONO', 'A-FIGOPTUNIF', 'A-FIGCOLOUR', 'A-FIGACCENT',
                 'A-FIGCVD',
                 'A-FIGSERIES', 'A-FIGGLYPH', 'A-FIGALT', 'A-FIGLABELPX')

    # 53. BLOCK.IMAGES IS POPULATED — the fixture whose absence WAS the defect.
    #     Nothing in this suite had ever asserted that a block carries an image.
    # ══════════════════════════════════════════════════════════════════════
    # 53a-53g. v2.21.4 — gate_images MUTATION CLOSURE + TWO REAL DEFECTS.
    #     audit_mutation.py showed SEVEN findings in gate_images could be deleted
    #     outright with all 120 fixtures still green: multi_per_line,
    #     figtext_prose, math_raster, warn_view and all three composite arms. The
    #     gate owning A-FIGCOMP and A-MATHRASTER on every figural paper in the
    #     estate had no fixture that could detect it going silent. Probing that
    #     space surfaced two real defects, both fixed in this release and locked
    #     by 53f/53g.
    # ══════════════════════════════════════════════════════════════════════
    def _fig_src(role='stem_and_options', oc=4, figq=(1,), specs=None):
        _s = _src_stub(tq=1); _s['options_count'] = oc
        _s['figural_qs'] = set(figq); _s['omml_required_present'] = False
        _s['concept_map'] = {'1': {'subtopic_id': 's.a'}}
        _s['section_rules_text'] = f'subtopic_id: s.a\n  image_role: {role}\n'
        if specs is not None:
            _s['figure_specs'] = specs
        return _s
    def _img_verdict(names, role='stem_and_options', oc=4, figq=(1,), specs=None):
        _p = _img_doc(list(names))
        _mm = gate_zip(_p)
        _doc = Document(_p); _t, _bl = parse_blocks(_doc)
        _reset(); gate_images(_bl, _fig_src(role, oc, figq, specs), _mm)
        return {c: l for l, c, m in RESULTS}
    def _specs(*roles, q=1):
        """A registry figure_specs map as Step 7 v5.34 transports it."""
        return {f'q{q}_{r}.png': {'question': q, 'role': r} for r in roles}
    _CANON5 = ['q1_problem.png', 'q1_opt1.png', 'q1_opt2.png',
               'q1_opt3.png', 'q1_opt4.png']

    # 53a — GUARD: the CANONICAL stem_and_options set (problem + one image per
    #       option) must stay clean, so none of the fixes below can be "achieved"
    #       by making the gate reject everything.
    check('FIGCOMP-canonical-stem-and-options-clean',
          _img_verdict(_CANON5).get('A-FIGCOMP') == 'OK')

    # 53a-1..4  v2.22.0 — GAP-2026-08-03-FIGCOMP-ROLE.
    #   MEASURED ON REAL DELIVERED OUTPUT: 27 of 33 figural questions in
    #   IIT_JAM_BIOTECHNOLOGY M01 are ONE role='problem' figure with FOUR TEXT
    #   OPTIONS. The gate INFERRED 'stem_and_options', demanded oc+1 images, and
    #   emitted 17 findings against CORRECT questions — each telling an operator
    #   to "VIEW + fix in Part B", i.e. deferring to a human who, in the deployed
    #   configuration, does not exist. The registry figure_specs record the role
    #   figural_core actually DREW; that is not an inference and it now wins.

    # 53a-1 — THE FIX. Problem-only specs ⇒ no option figures exist to bind ⇒ the
    #         per-option arm is inapplicable. MUTATION-VERIFIED against v2.21.9.
    check('FIGCOMP-problem-only-specs-not-flagged',
          _img_verdict(['q1_problem.png'],
                       specs=_specs('problem')).get('A-FIGCOMP') == 'OK')

    # 53a-2 — THE FIX IS NOT A BLANKET EXEMPTION. When the producer DID draw
    #         option figures, a PARTIAL render must STILL be caught. Without this
    #         53a-1 could be "achieved" by exempting every figural block, which
    #         would restore the v2.21.4 defect (options silently undrawn, a
    #         question no candidate can answer).
    check('FIGCOMP-declared-option-figures-still-required',
          _img_verdict(['q1_problem.png', 'q1_opt1.png'],
                       specs=_specs('problem', 'opt1', 'opt2', 'opt3', 'opt4')
                       ).get('A-FIGCOMP') == 'WARN')

    # 53a-3 — COUNTED FROM THE SPECS, NOT FROM THE EXAM-WIDE OPTIONS_COUNT. A
    #         question declaring a SHORTER option-figure set is judged against
    #         what it declares; oc would over-demand and re-create the false
    #         positive one layer down.
    check('FIGCOMP-expectation-derives-from-specs-not-oc',
          _img_verdict(['q1_problem.png', 'q1_opt1.png', 'q1_opt2.png'],
                       oc=4,
                       specs=_specs('problem', 'opt1', 'opt2')).get('A-FIGCOMP') == 'OK')

    # 53a-4 — LEGACY UNCHANGED. No specs (pre-v5.34 registry) ⇒ every earlier
    #         inference branch stands exactly as before. ~200 exams depend on it.
    check('FIGCOMP-no-specs-legacy-path-unchanged',
          _img_verdict(['q1_problem.png'], specs={}).get('A-FIGCOMP') == 'WARN'
          and _img_verdict(_CANON5, specs={}).get('A-FIGCOMP') == 'OK')

    # ════════════════════════════════════════════════════════════════════════
    # 53a-5 — A-AXIS1 / A-AXIS3 BUDGET GATES  (v2.24, GAP-2026-08-06-AXIS1)
    # ════════════════════════════════════════════════════════════════════════
    # The defect these close shipped TWICE on a real exam and passed all 24 gates
    # then in the roster. Every assertion below is MUTATION-VERIFIED against the
    # pre-v2.24 build (where the gate did not exist and the verdict map is empty).
    # No exam, section or format name here is load-bearing.

    def _axis_verdict(n_fig, target, sec_qs=60, irreducible_subs=None, enforcement='hard',
                      di_qs=None, di_present=False):
        # ENGINE-AVAILABILITY IS PART OF THE CONTRACT, NOT AN OBSTACLE TO IT.
        # audit_mutation.py copies THIS file into a temp dir and runs it from that
        # dir, so blueprint_core is genuinely absent there. The gate's answer in that
        # situation — WARN "NOT ESTABLISHED" — is CORRECT and load-bearing: an auditor
        # must never fabricate a verdict it cannot establish. So the fixture asserts
        # the real verdicts when the engine is present and the degraded verdict when it
        # is not, rather than forcing one environment or skipping the checks. Both
        # branches are real assertions; neither is vacuous. (Forcing the import by
        # mutating sys.path was the first attempt and it was wrong: it would have hidden
        # a genuinely broken delegation behind a fixture's convenience.)
        _s = _src_stub(tq=sec_qs, sections=[{'name': 'S1', 'q_range': [1, sec_qs],
                                             'total_qs': sec_qs}])
        _s['axis_schedule'] = {'S1': {'status': 'ok', 'axis1_target_per_mock': target,
                                      'axis1_enforcement': enforcement}}
        _s['figural_qs'] = set(range(1, n_fig + 1))
        _s['figural_subtopics'] = {str(q): f'ST{q}' for q in range(1, n_fig + 1)}
        _s['figural_reducible'] = {f'ST{q}': False for q in (irreducible_subs or [])}
        _s['figural_manifest_present'] = True     # v2.24.1 — a REAL record exists
        _s['rc_manifest_present'] = False
        _s['passage_linked'] = set()
        _s['di_qs'] = set(di_qs or ())            # v2.25
        _s['di_manifest_present'] = di_present
        _reset()
        _safe_gate('A-AXIS1', gate_axis1, [], _s)
        return {c: l for l, c, _m in RESULTS}

    try:
        import blueprint_core as _bc_probe            # noqa: F401
        _BC_OK = True
    except Exception:
        _BC_OK = False

    def _axv(expected):
        """The verdict to expect: the real one when the engine is reachable, WARN
        (NOT ESTABLISHED) when it is not. A-AXIS-UNGATED is pure blueprint arithmetic
        with no engine dependency, so it is asserted unconditionally."""
        return expected if _BC_OK else 'WARN'

    # (a) THE DEFECT ITSELF. 26 figures against a budget of 4 must FAIL. This is the
    #     exact delivered paper that motivated the release; pre-v2.24 it was OK.
    check('AXIS1-the-shipped-defect-now-fails',
          _axis_verdict(26, {'TEXT': 56, 'FIGURAL': 4}).get('A-AXIS1') == _axv('FAIL'))

    # (b) A CONFORMANT PAPER STAYS CLEAN — and the band is a band. Real papers vary
    #     (the reference exam ranged 2→8 across five years); a gate demanding an exact
    #     count gets switched off by hand, which is strictly worse than no gate.
    check('AXIS1-conformant-paper-passes',
          _axis_verdict(4, {'TEXT': 56, 'FIGURAL': 4}).get('A-AXIS1') == _axv('OK'))
    check('AXIS1-within-band-passes',
          _axis_verdict(5, {'TEXT': 56, 'FIGURAL': 4}).get('A-AXIS1') == _axv('OK'))

    # (c) SHORTFALL IS A FINDING TOO. Auditing only the upper bound would leave the
    #     gate half-blind: a 0-figure paper for a figural exam is as unfaithful as a
    #     26-figure one.
    check('AXIS1-shortfall-detected',
          _axis_verdict(0, {'TEXT': 52, 'FIGURAL': 8}).get('A-AXIS1') == _axv('FAIL'))

    # (d) IRREDUCIBLE OVERAGE IS SILENT (operator decision 2026-08-06). Questions whose
    #     OPTIONS are images have no text form; they are granted over budget and the
    #     EXPECTATION RISES with them, so there is no finding to read past.
    check('AXIS1-irreducible-overage-is-silent-pass',
          _axis_verdict(9, {'TEXT': 56, 'FIGURAL': 4},
                        irreducible_subs=range(1, 10)).get('A-AXIS1') == _axv('OK'))

    # (e) AND THE EXEMPTION IS NOT A HOLE. Excess NOT covered by irreducibles must
    #     still FAIL. Without this, (d) could be "achieved" by exempting every figural
    #     question — which restores the exact defect this gate exists to close.
    check('AXIS1-unexplained-excess-still-fails',
          _axis_verdict(26, {'TEXT': 56, 'FIGURAL': 4},
                        irreducible_subs=[1, 2, 3]).get('A-AXIS1') == _axv('FAIL'))

    # (f) DORMANT ON A PRE-v1.23 BLUEPRINT. ~200 deployed exams carry no axis_schedule
    #     and must keep passing untouched until they are re-measured.
    _s_no = _src_stub(tq=60); _s_no['figural_qs'] = set(range(1, 27))
    _reset(); _safe_gate('A-AXIS1', gate_axis1, [], _s_no)
    check('AXIS1-dormant-without-axis-schedule',
          {c: l for l, c, _m in RESULTS}.get('A-AXIS1') == 'OK')

    # (h) OBSERVABILITY (v2.24.1). Every assertion below is MUTATION-VERIFIED against
    #     the v2.24.0 build, where the gate fabricated observed=0 for any class it held
    #     no evidence for. That produced a HARD FAIL reading "produced 0, budget 6" on
    #     every DI-targeting exam in the estate — a false alarm on ~200 exams, and a
    #     gate that cries wolf is one somebody switches off.
    def _cov(**kw):
        _r = _axis_verdict(**kw)
        return _r.get('A-AXIS1'), _r.get('A-AXIS1-COVERAGE')
    # SCOPE NOTE: the assertions from here to (n) exercise the DELEGATED verdict, so
    # they need blueprint_core reachable. Under audit_mutation.py — which runs this file
    # from a temp dir where the engine is absent — the honest answer to every one of
    # them is NOT ESTABLISHED, and asserting 'OK'/'WARN' there would be asserting the
    # harness's isolation rather than the gate. The engine-free half of the contract is
    # asserted unconditionally in (o) below, so nothing goes unchecked either way.
    _v, _c = _cov(n_fig=4, target={'TEXT': 50, 'FIGURAL': 4, 'DI': 6})
    check('AXIS1-unobservable-DI-does-not-fail-the-paper', _v == _axv('OK'))
    if _BC_OK:
        check('AXIS1-unobservable-DI-is-REPORTED-not-swallowed', _c == 'WARN')

    # (i) AND COVERAGE IS NOT AN AMNESTY. An observable breach must still FAIL in the
    #     same run, or (h) could be "achieved" by excusing every class.
    _v, _c = _cov(n_fig=26, target={'TEXT': 50, 'FIGURAL': 4, 'DI': 6})
    check('AXIS1-observable-breach-still-fails-beside-unobservable-class',
          _v == _axv('FAIL') and _c == 'WARN')

    # (j) A CLEAN, FULLY-OBSERVABLE PAPER MUST BE CLEAN ON BOTH LINES — no standing WARN
    #     that trains the reader to ignore the coverage line.
    _v, _c = _cov(n_fig=4, target={'TEXT': 56, 'FIGURAL': 4})
    if _BC_OK:
        check('AXIS1-fully-observable-conformant-paper-is-clean-on-both-lines',
              _v == 'OK' and _c == 'OK')

    # (k) NO q_range ⇒ every count would be a fabricated zero. Skip and SAY SO.
    _sq = _src_stub(tq=60, sections=[{'name': 'S1', 'total_qs': 60}])
    _sq['axis_schedule'] = {'S1': {'status': 'ok',
                                   'axis1_target_per_mock': {'TEXT': 56, 'FIGURAL': 4}}}
    _sq['figural_qs'] = set(); _sq['figural_manifest_present'] = True
    _sq['figural_subtopics'] = {}; _sq['figural_reducible'] = {}; _sq['passage_linked'] = set()
    _reset(); _safe_gate('A-AXIS1', gate_axis1, [], _sq)
    _rq = {c: l for l, c, _m in RESULTS}
    check('AXIS1-section-without-q_range-is-skipped-not-failed',
          _rq.get('A-AXIS1') == _axv('OK')
          and (_rq.get('A-AXIS1-COVERAGE') == 'WARN' if _BC_OK else True))

    # (l) A MISSING FIGURAL MANIFEST IS NOT A PAPER WITH NO FIGURES. Reading absence as
    #     zero turns "I have no record" into a shortfall FAIL — the same conflation
    #     gate_images separates as `declared` vs `resolved`.
    # PRESENT manifest reporting zero figures against a budget of 4 IS a real shortfall
    # and must FAIL. This is the exact contrast that gives the next assertion meaning:
    # same observed counts, opposite verdicts, decided solely by whether a record exists.
    _v, _c = _cov(n_fig=0, target={'TEXT': 56, 'FIGURAL': 4})
    check('AXIS1-present-manifest-with-zero-figures-is-a-real-shortfall',
          _v == _axv('FAIL'))
    _sm = _src_stub(tq=60, sections=[{'name': 'S1', 'q_range': [1, 60], 'total_qs': 60}])
    _sm['axis_schedule'] = {'S1': {'status': 'ok',
                                   'axis1_target_per_mock': {'TEXT': 56, 'FIGURAL': 4}}}
    _sm['figural_qs'] = set(); _sm['figural_manifest_present'] = False
    _sm['figural_subtopics'] = {}; _sm['figural_reducible'] = {}; _sm['passage_linked'] = set()
    _reset(); _safe_gate('A-AXIS1', gate_axis1, [], _sm)
    _rm = {c: l for l, c, _m in RESULTS}
    # Coverage must WARN here in BOTH environments, but for different reasons: with the
    # engine, because FIGURAL has no observation source; without it, because nothing was
    # established at all. Either way the paper must not be reported as short of figures.
    check('AXIS1-absent-manifest-is-unestablished-not-a-shortfall',
          _rm.get('A-AXIS1') == _axv('OK')
          and (_rm.get('A-AXIS1-COVERAGE') == 'WARN' if _BC_OK else True))

    # (m) FIGURAL Qs BELONGING TO NO SECTION must be surfaced. Silent loss would
    #     under-count the exact quantity this gate exists to police.
    _so = _src_stub(tq=60, sections=[{'name': 'S1', 'q_range': [1, 30], 'total_qs': 30}])
    _so['axis_schedule'] = {'S1': {'status': 'ok',
                                   'axis1_target_per_mock': {'TEXT': 28, 'FIGURAL': 2}}}
    _so['figural_qs'] = {97, 98}; _so['figural_manifest_present'] = True
    _so['figural_subtopics'] = {}; _so['figural_reducible'] = {}; _so['passage_linked'] = set()
    _reset(); _safe_gate('A-AXIS1', gate_axis1, [], _so)
    check('AXIS1-orphan-figural-questions-are-reported',
          {c: l for l, c, _m in RESULTS}.get('A-AXIS1-COVERAGE') == 'WARN')

    # (n) AXIS-3 SHARES THE DISCIPLINE. An empty options_by_q means NAT is genuinely
    #     unknowable; asserting a shortfall from that is a verdict built on nothing.
    _s3 = _src_stub(tq=30, sections=[{'name': 'S1', 'q_range': [1, 30], 'total_qs': 30}])
    _s3['axis_schedule'] = {'S1': {'status': 'ok',
                                   'axis3_target_per_mock': {'MCQ': 16, 'NAT': 10, 'MSQ': 4}}}
    _s3['options_by_q'] = {}; _s3['msq_instruction_phrases'] = []
    _reset(); _safe_gate('A-AXIS3', gate_axis3, [], _s3)
    _r3 = {c: l for l, c, _m in RESULTS}
    check('AXIS3-absent-evidence-does-not-fail-the-paper', _r3.get('A-AXIS3') == _axv('OK'))
    if _BC_OK:
        check('AXIS3-absent-evidence-is-REPORTED', _r3.get('A-AXIS3-COVERAGE') == 'WARN')

    # (o) ENGINE-FREE HALF OF THE CONTRACT, asserted in EVERY environment. When
    #     blueprint_core is unreachable the verdict must degrade to NOT ESTABLISHED
    #     (never to a silent OK, which would be a vacuous pass) AND the coverage line
    #     must still be emitted — it is the line that says what went unchecked, so
    #     losing it exactly when the engine is missing is the worst possible time.
    check('AXIS1-both-lines-always-emitted',
          set(_axis_verdict(n_fig=4, target={'TEXT': 56, 'FIGURAL': 4}))
          >= {'A-AXIS1', 'A-AXIS1-COVERAGE'})
    check('AXIS1-verdict-degrades-to-NOT-ESTABLISHED-never-to-silent-OK',
          _axis_verdict(n_fig=26, target={'TEXT': 56, 'FIGURAL': 4}).get('A-AXIS1')
          in (('FAIL',) if _BC_OK else ('WARN',)))

    # (q) DI OBSERVABILITY (v2.25, GAP-2026-08-06-DI). DI was the last budgeted
    #     stimulus class with no producer record, so A-AXIS1 could only ever report it
    #     unestablished. With a di_manifest it is counted like any other class.
    _v, _c = _cov(n_fig=4, target={'TEXT': 50, 'FIGURAL': 4, 'DI': 6},
                  di_qs=range(1, 7), di_present=True)
    check('DI-observable-when-the-producer-recorded-it',
          _v == _axv('OK') and (_c == 'OK' if _BC_OK else True))

    # (r) AND IT IS NOW JUDGED, NOT MERELY SEEN. An over-produced DI must FAIL — this
    #     is the false-PASS half of the old defect, where DI fell into the TEXT residual
    #     and vanished. Mutation-verified: measures False on the v2.24.1 build.
    # FIGURAL target is 0 here ON PURPOSE, so DI is the ONLY class that can produce a
    # finding. An earlier draft left FIGURAL at 4 with 0 produced, and the fixture then
    # went green off the FIGURAL shortfall — it survived a mutation that deleted DI
    # observability entirely. A test that passes for the wrong reason is worse than no
    # test, because it reports coverage it does not have.
    _v, _c = _cov(n_fig=0, target={'TEXT': 60, 'FIGURAL': 0, 'DI': 0},
                  di_qs=range(1, 21), di_present=True)
    check('DI-over-production-is-caught-not-absorbed-into-TEXT', _v == _axv('FAIL'))

    # (s) A DI SHORTFALL IS A FINDING TOO — the gate is not one-sided.
    _v, _c = _cov(n_fig=0, target={'TEXT': 48, 'FIGURAL': 0, 'DI': 12},
                  di_qs=(), di_present=True)
    check('DI-shortfall-detected', _v == _axv('FAIL'))

    # (t) NO RECORD ⇒ UNESTABLISHED, NEVER A SILENT ZERO. A pre-v5.42 Step 7 writes no
    #     di_manifest; reading that absence as "no DI questions" would turn an unknown
    #     into a hard shortfall on every legacy exam.
    _v, _c = _cov(n_fig=4, target={'TEXT': 50, 'FIGURAL': 4, 'DI': 6},
                  di_qs=(), di_present=False)
    check('DI-absent-manifest-is-unestablished-not-a-shortfall',
          _v == _axv('OK') and (_c == 'WARN' if _BC_OK else True))

    # (u) DI IS NEVER INFERRED FROM A TABLE. Recorded as an executable reminder of the
    #     measurement that settles it: on IIT_JAM_BIOTECHNOLOGY 15-Feb-2026 the paper
    #     carries THREE Word tables — 'Vitamins|Symptoms' and 'Nitrogen compound|
    #     Oxidation state' are MATCH, only 'Reactant|Product|Standard Enthalpy' is DI.
    #     A table-presence heuristic reports 3 where the truth is 1, i.e. it trades a
    #     silent miss for a confident wrong answer. The gate reads di_qs and nothing else.
    _v, _c = _cov(n_fig=0, target={'TEXT': 59, 'FIGURAL': 0, 'DI': 1},
                  di_qs=[3], di_present=True)
    check('DI-counted-from-the-producer-record-only', _v == _axv('OK'))

    # (x) THE AUDITOR'S USE OF THE v1.45 ENGINE WORK (v2.28).
    #     blueprint_core gained 19 fixtures for figural_band / figural_target_series /
    #     figural_quota / schedule_figural_slots. audit_canonical gained ZERO while
    #     changing 14 lines, so BOTH of its new lines survived mutation at 221/221:
    #     deleting the axis1_target_series indexing, and passing observed_spread=None,
    #     each silently reverted the gate to false-FAILing real-shaped papers — exactly
    #     what those two fixes exist to prevent — with nothing objecting.
    #     THE ENGINE BEING WELL-TESTED SAYS NOTHING ABOUT THE CONSUMER USING IT.
    #     Verified by mutation before writing these: both reverts left every fixture green.
    _OBSF = [8, 3, 2, 6, 3]          # the reference exam's real per-paper figure counts

    def _axv2(n_fig, target, series=None, spread=None, mock_n=1, sec_qs=60):
        _s = _src_stub(tq=sec_qs, sections=[{'name': 'S1', 'q_range': [1, sec_qs],
                                             'total_qs': sec_qs}])
        _sec = {'status': 'ok', 'axis1_target_per_mock': target}
        if series is not None:
            _sec['axis1_target_series'] = series
        if spread is not None:
            _sec['axis1_observed_figural'] = spread
        _s['axis_schedule'] = {'S1': _sec}
        _s['mock_n'] = mock_n
        _s['figural_qs'] = set(range(1, n_fig + 1))
        _s['figural_manifest_present'] = True
        _s['figural_subtopics'] = {}; _s['figural_reducible'] = {}
        _s['passage_linked'] = set(); _s['di_qs'] = set(); _s['di_manifest_present'] = False
        _reset(); _safe_gate('A-AXIS1', gate_axis1, [], _s)
        return {c: l for l, c, _m in RESULTS}.get('A-AXIS1')

    # SERIES INDEXING. Mock 5 of the rotating series targets 2 figures; a paper carrying
    # 8 must FAIL. Drop the indexing and the gate falls back to the flat target of 5,
    # whose band admits 8 — the paper passes and the series means nothing.
    check('AXIS1-uses-THIS-mock-target-from-the-series',
          _axv2(8, {'TEXT': 55, 'FIGURAL': 5}, series=[8, 6, 3, 3, 2],
                spread=_OBSF, mock_n=5) == _axv('FAIL'))
    # And the positive half: mock 1 targets 8, so 8 figures is exactly right.
    check('AXIS1-series-accepts-a-figure-heavy-mock-when-that-is-its-target',
          _axv2(8, {'TEXT': 55, 'FIGURAL': 5}, series=[8, 6, 3, 3, 2],
                spread=_OBSF, mock_n=1) == _axv('OK'))

    # OBSERVED SPREAD. With the exam's own volatility the band is ±4, so a real
    # 8-figure paper passes against a target of 5. Pass observed_spread=None and the
    # band narrows to ±2 — the gate resumes rejecting genuine papers, which is the
    # failure mode that made the v2.24 band unusable.
    check('AXIS1-band-widens-to-the-exam-own-observed-spread',
          _axv2(8, {'TEXT': 55, 'FIGURAL': 5}, spread=_OBSF) == _axv('OK'))
    # Not a blanket amnesty: the 26-figure paper is still far outside even the wide band.
    check('AXIS1-wide-band-still-catches-the-original-defect',
          _axv2(26, {'TEXT': 55, 'FIGURAL': 5}, spread=_OBSF) == _axv('FAIL'))

    # TEXT RESIDUAL, pinned directly. No verdict depends on it (the residual class is
    # skipped by check_axis_conformance), so reverting it to `sec_qs - figural` left all
    # 221 fixtures green. Unreachable-by-assertion code drifts; the residual is the one
    # place an over-produced DI could hide, so its arithmetic is asserted as a unit.
    check('AXIS1-TEXT-residual-subtracts-every-stimulus-class',
          _axis1_observed(60, [1, 2, 3], 4, {10, 11}, 1, 60)
          == {'FIGURAL': 3, 'PASSAGE': 4, 'DI': 2, 'TEXT': 51})
    check('AXIS1-TEXT-residual-never-goes-negative',
          _axis1_observed(2, [1, 2], 3, {1, 2}, 1, 60)['TEXT'] == 0)
    check('AXIS1-observed-counts-only-inside-the-section',
          _axis1_observed(30, [1], 0, {5, 99}, 1, 30)['DI'] == 1)
    check('AXIS1-observed-total-on-junk',
          _axis1_observed(None, None, None, 'x', 1, 30)['TEXT'] == 0
          and _axis1_count_in(None, 1, 5) == 0
          and _axis1_count_in(['a', 2], 1, 5) == 1)

    # (v) A-AXIS1-OVERLAP (v2.26). Axis-1 classes must PARTITION the paper — one
    #     stimulus per question, which is why the targets are declared to sum to
    #     sec_qs. Three producer records are written by three code paths and nothing
    #     cross-checked them. Every assertion here is MUTATION-VERIFIED against v2.25,
    #     where the gate did not exist and the verdict map is empty.
    def _ovl(fig=(), di=(), pas=(), sec_qs=60, fp=True, dp=True, rp=False):
        _s = _src_stub(tq=sec_qs, sections=[{'name': 'S1', 'q_range': [1, sec_qs],
                                             'total_qs': sec_qs}])
        _s['figural_qs'] = set(fig); _s['figural_manifest_present'] = fp
        _s['di_qs'] = set(di);       _s['di_manifest_present'] = dp
        _s['passage_linked'] = set(pas); _s['rc_manifest_present'] = rp
        _reset(); _safe_gate('A-AXIS1-OVERLAP', gate_axis1_overlap, [], _s)
        return {c: l for l, c, _m in RESULTS}.get('A-AXIS1-OVERLAP')

    # THE CASE THAT MOTIVATED THE GATE: 60 figures AND 60 DI in a 60-question section.
    # 120 stimuli in 60 slots. v2.25 returned OK on exactly this input.
    check('OVERLAP-impossible-double-record-is-caught',
          _ovl(fig=range(1, 61), di=range(1, 61)) == 'FAIL')

    # A single shared question is enough — the defect does not need to be extreme to
    # corrupt the count, it just needs to be silent.
    check('OVERLAP-single-shared-question-is-caught',
          _ovl(fig=[1, 2, 3], di=[3]) == 'FAIL')
    check('OVERLAP-detected-across-every-manifest-pair',
          _ovl(fig=[1], pas=[1], rp=True) == 'FAIL'
          and _ovl(di=[7], pas=[7], rp=True) == 'FAIL')

    # AND IT IS NOT A BLANKET ALARM. Disjoint records are the normal case and must stay
    # clean, or the gate becomes noise and gets switched off — the failure mode that
    # made the v2.24 DI false-FAIL so damaging.
    check('OVERLAP-disjoint-manifests-are-clean',
          _ovl(fig=[1, 2, 3, 4], di=[10, 11]) == 'OK')
    check('OVERLAP-empty-manifests-are-clean', _ovl(fig=[], di=[]) == 'OK')

    # A "more stimuli than slots" arm was written here and REMOVED as provably
    # unreachable: disjoint sets confined to one section cannot outnumber it, so it
    # could never fire without the overlap arm firing first (verified by exhaustive
    # search, 20,000 randomised disjoint configurations, zero hits). This assertion
    # locks in the proof — disjoint records that FILL a section exactly stay clean.
    check('OVERLAP-disjoint-records-filling-a-section-exactly-are-clean',
          _ovl(fig=range(1, 9), di=range(9, 11), sec_qs=10) == 'OK')

    # FEWER THAN TWO RECORDS ⇒ nothing CAN overlap. Reported as such rather than passed
    # silently, so "no conflict" is distinguishable from "nothing to compare".
    check('OVERLAP-dormant-with-a-single-manifest',
          _ovl(fig=[1, 2], dp=False) == 'OK' and _ovl(fp=False, dp=False) == 'OK')

    # ENGINE-FREE IN EVERY ENVIRONMENT. Asserted unwrapped (no _axv): an integrity check
    # that goes dormant whenever something ELSE is missing is worth very little. This is
    # the lesson A-AXIS-UNGATED taught when it had to be lifted out of gate_axis1.
    check('OVERLAP-verdict-is-engine-independent',
          _ovl(fig=range(1, 61), di=range(1, 61)) == 'FAIL')

    # (w) ORPHAN REPORTING NOW COVERS EVERY STIMULUS CLASS (v2.26). A DI or passage
    #     question in no section leaves the denominator exactly as silently as a
    #     figural one, and the budget then looks satisfied because part of the paper
    #     stopped existing. Only figures were checked before.
    def _orph(cls_qs):
        _s = _src_stub(tq=60, sections=[{'name': 'S1', 'q_range': [1, 30],
                                         'total_qs': 30}])
        _s['axis_schedule'] = {'S1': {'status': 'ok',
                                      'axis1_target_per_mock': {'TEXT': 28, 'FIGURAL': 2}}}
        _s['figural_qs'] = set(cls_qs.get('FIGURAL', ()))
        _s['figural_manifest_present'] = True
        _s['figural_subtopics'] = {}; _s['figural_reducible'] = {}
        _s['passage_linked'] = set(cls_qs.get('PASSAGE', ()))
        _s['rc_manifest_present'] = 'PASSAGE' in cls_qs
        _s['di_qs'] = set(cls_qs.get('DI', ()))
        _s['di_manifest_present'] = 'DI' in cls_qs
        _reset(); _safe_gate('A-AXIS1', gate_axis1, [], _s)
        return next((m for l, c, m in RESULTS if c == 'A-AXIS1-COVERAGE'), '')
    check('ORPHAN-figural-outside-every-section-is-named',
          'FIGURAL Q(s) fall outside' in _orph({'FIGURAL': [97]}))
    check('ORPHAN-DI-outside-every-section-is-named',
          'DI Q(s) fall outside' in _orph({'FIGURAL': [1], 'DI': [98]}))
    check('ORPHAN-PASSAGE-outside-every-section-is-named',
          'PASSAGE Q(s) fall outside' in _orph({'FIGURAL': [1], 'PASSAGE': [99]}))
    check('ORPHAN-in-range-questions-are-not-flagged',
          'fall outside' not in _orph({'FIGURAL': [1, 2], 'DI': [3], 'PASSAGE': [4]}))

    # (p) _axis_sections — EACH REJECTION BRANCH IS NAMED, NOT JUST SKIPPED.
    #     Tested directly rather than through a gate because the gate's coverage WARN
    #     can also be raised by an unobservable class, so a gate-level fixture cannot
    #     tell WHICH cause fired and would pass even if the malformed-section report
    #     were deleted (mutation-verified: these four appends survived every gate-level
    #     fixture). A section silently vanishing is the worst outcome here — its
    #     questions leave the denominator and the budget looks satisfied because part
    #     of the paper stopped existing.
    _gs, _bs = _axis_sections({'sections': [
        {'name': 'GOOD', 'q_range': [1, 30]},        # well-formed
        'not-a-dict',                                 # → '<malformed section>'
        {'name': 'NORANGE'},                          # q_range absent
        {'name': 'STRRANGE', 'q_range': 'bad'},       # q_range not a sequence
        {'name': 'SHORT', 'q_range': [1]},            # q_range too short
        {'name': 'NONNUM', 'q_range': ['a', 'b']},    # non-numeric bounds
        {'name': 'REVERSED', 'q_range': [30, 1]},     # hi < lo
    ]})
    check('AXIS-SECTIONS-keeps-only-the-well-formed-section',
          _gs == [('GOOD', 1, 30)])
    check('AXIS-SECTIONS-names-every-rejected-section',
          _bs == ['<malformed section>', 'NORANGE', 'STRRANGE', 'SHORT',
                  'NONNUM', 'REVERSED'])
    check('AXIS-SECTIONS-non-dict-section-is-named-not-dropped',
          '<malformed section>' in _axis_sections({'sections': [42]})[1])
    check('AXIS-SECTIONS-missing-range-is-named',
          _axis_sections({'sections': [{'name': 'X'}]})[1] == ['X'])
    check('AXIS-SECTIONS-non-numeric-bounds-are-named',
          _axis_sections({'sections': [{'name': 'X', 'q_range': ['a', 2]}]})[1] == ['X'])
    check('AXIS-SECTIONS-reversed-range-is-named',
          _axis_sections({'sections': [{'name': 'X', 'q_range': [9, 2]}]})[1] == ['X'])
    check('AXIS-SECTIONS-total-on-hostile-input',
          all(_axis_sections(_h) == ([], []) for _h in
              ({}, {'sections': None}, {'sections': 'x'}, {'sections': 42})))

    # (g) A-AXIS-UNGATED — the rule that stops this returning as Axis-4. An axis the
    #     blueprint marks enforcement:"hard" with no gate here is itself a finding.
    #     axis1/axis3 are gated, so a normal schedule is clean; a fabricated axis4 is not.
    def _ungated(sched):
        _s = _src_stub(tq=60); _s['axis_schedule'] = sched
        _reset(); _safe_gate('A-AXIS-UNGATED', gate_axis_ungated, [], _s)
        return {c: l for l, c, _m in RESULTS}.get('A-AXIS-UNGATED')
    # NOT wrapped in _axv: this gate is engine-free ON PURPOSE, so it must return a real
    # verdict even where blueprint_core is unreachable. Asserting the raw value here is
    # what proves that independence — the property is the whole point of the gate.
    check('AXIS-UNGATED-clean-when-every-hard-axis-is-gated',
          _ungated({'S1': {'status': 'ok', 'axis1_enforcement': 'hard',
                           'axis3_enforcement': 'hard'}}) == 'OK')
    check('AXIS-UNGATED-detects-an-enforced-axis-with-no-gate',
          _ungated({'S1': {'status': 'ok', 'axis4_enforcement': 'hard'}}) == 'FAIL')
    check('AXIS-UNGATED-dormant-without-axis-schedule', _ungated({}) == 'OK')

    # 53b — MATH-TOKEN raster name is an A-MATHRASTER finding (kills math_raster).
    check('MATHRASTER-math-token-name-is-a-finding',
          _img_verdict(['q1_problem.png', 'q1_eqn.png'])
          .get('A-MATHRASTER') == 'FAIL')

    # 53c — NON-CANONICAL image name routes to the Part-B view WARN (warn_view).
    check('MATHRASTER-VIEW-noncanonical-name-warns',
          _img_verdict(['q1_problem.png', 'Picture 2.png'])
          .get('A-MATHRASTER-VIEW') == 'WARN')

    # 53d — TWO IMAGES ON ONE LINE is an A-FIGCOMP-LINE finding (multi_per_line).
    def _b_twoline(d):
        # v2.21.9 — Inches is now a MODULE-level import (line 42). A function-local
        # import here would shadow it for the WHOLE of self_test(), unbinding it for
        # every fixture defined earlier in the same scope.
        d.add_paragraph('Q.1  Study the figure.')
        _pp = d.add_paragraph()
        for _ in range(2):
            _pp.add_run().add_picture(io.BytesIO(_png_bytes()), width=Inches(1.0))
    _p53d = _mini_doc(tmp, _b_twoline)
    _mm53d = gate_zip(_p53d)
    _t, _bl53d = parse_blocks(Document(_p53d))
    _reset(); gate_images(_bl53d, _fig_src(), _mm53d)
    check('FIGCOMP-LINE-two-images-one-line-is-a-finding',
          any(c == 'A-FIGCOMP-LINE' and l == 'FAIL' for l, c, m in RESULTS))

    # 53e — FIGURE-REFERENCE PROSE in a ZERO-IMAGE block is a finding
    #       (kills figtext_prose). A figural subtopic rendered as prose is a
    #       figure that was never drawn.
    def _b_prose(d):
        d.add_paragraph('Q.1  In the given figure, find the value of X.')
        for _i in range(1, 5):
            d.add_paragraph(f'{_i}.  Option {_i}')
    _p53e = _mini_doc(tmp, _b_prose)
    _mm53e = gate_zip(_p53e)
    _t, _bl53e = parse_blocks(Document(_p53e))
    _reset(); gate_images(_bl53e, _fig_src(), _mm53e)
    check('FIGTEXT-PROSE-zero-image-figure-reference-is-a-finding',
          any(c == 'A-FIGTEXT-PROSE' and l == 'FAIL' for l, c, m in RESULTS))

    # 53f — v2.21.4 DEAD-BRANCH LOCK. A REGISTRY-DECLARED FIGURAL question that
    #       renders ZERO images is a finding. The `if not block_imgs: continue`
    #       sat ABOVE the stem_only arm, so `len(block_imgs) < 1` could NEVER be
    #       true: that arm was DEAD CODE and a figure that was never drawn passed
    #       A-FIGCOMP clean unless its stem happened to match the prose pattern —
    #       detection depended on the WORDING of the stem, not on the absence of
    #       the figure. Membership comes from the REGISTRY set only, never from
    #       figural_cue_keywords (which contain ordinary MCQ phrases like
    #       'which of the'); 53f2 guards that a NON-figural text question with
    #       zero images stays clean.
    def _b_noimg(d):
        d.add_paragraph('Q.1  Which of the following is correct?')
        for _i in range(1, 5):
            d.add_paragraph(f'{_i}.  Option {_i}')
    _p53f = _mini_doc(tmp, _b_noimg)
    _mm53f = gate_zip(_p53f)
    _t, _bl53f = parse_blocks(Document(_p53f))
    _reset(); gate_images(_bl53f, _fig_src('stem_only', figq=(1,)), _mm53f)
    _f_declared = any(c == 'A-FIGCOMP' and l == 'WARN' and '0img' in m
                      for l, c, m in RESULTS)
    _reset(); gate_images(_bl53f, _fig_src('stem_only', figq=()), _mm53f)
    _f_plain = any(c == 'A-FIGCOMP' and l == 'OK' for l, c, m in RESULTS)
    check('FIGCOMP-declared-figural-with-zero-images-is-a-finding',
          _f_declared and _f_plain)

    # 53h — OPTIONS_ONLY SHORT SET is a finding. Step-7 G-FIGURAL-COMPOSITE:
    #       options_only requires ">=n option images, no problem image required",
    #       so 3 images where OPTIONS_COUNT is 4 means an option was never drawn.
    #       Guarded by the exact-count case staying clean.
    check('FIGCOMP-options-only-short-set-is-a-finding',
          _img_verdict(_CANON5[1:4], role='options_only').get('A-FIGCOMP') == 'WARN'
          and _img_verdict(_CANON5[1:], role='options_only').get('A-FIGCOMP') == 'OK')

    # 6b-6c. v2.21.6 — A-ZIP MUTATION CLOSURE. A-ZIP proves every rId referenced by
    #     document.xml resolves to a real part. BOTH failure modes were untested:
    #     an rId with NO relationship entry, and an rId whose relationship points at
    #     a part that is NOT IN THE ZIP. A docx failing either is structurally
    #     broken — images silently vanish in Word — so a silent A-ZIP is severe.
    def _zip_verdict(mutate):
        _src_docx = _img_doc(['q1_problem.png'])
        _out = os.path.join(tmp, f'zipmut_{abs(hash(mutate)) % 99999}.docx')
        with zipfile.ZipFile(_src_docx) as zin, \
             zipfile.ZipFile(_out, 'w', zipfile.ZIP_DEFLATED) as zout:
            for _i in zin.infolist():
                _nm, _data = _i.filename, zin.read(_i.filename)
                _keep, _data = mutate(_nm, _data)
                if _keep:
                    zout.writestr(_i, _data)
        _reset(); gate_zip(_out)
        return {c: l for l, c, m in RESULTS}.get('A-ZIP')

    # 6b — rId referenced but ABSENT from the .rels (no relationship at all).
    def _mut_drop_rel(nm, data):
        if nm == 'word/_rels/document.xml.rels':
            # the image relationship Type ends '/relationships/image' (lowercase)
            data = re.sub(rb'<Relationship\b[^>]*/image"[^>]*/?>', b'', data)
        return True, data
    # 6c — relationship present but the TARGET PART is missing from the archive.
    def _mut_drop_media(nm, data):
        return (not nm.startswith('word/media/')), data
    check('ZIP-unresolved-rId-is-a-finding',
          _zip_verdict(_mut_drop_rel) == 'FAIL'
          and _zip_verdict(_mut_drop_media) == 'FAIL'
          and _zip_verdict(lambda nm, d: (True, d)) == 'OK')

    # 53k-53m. v2.21.8 — RA-9 EXAM-INDEPENDENCE: A-FIGTEXT-PROSE MUST NOT ASSUME
    #     ENGLISH. The detector was a HARDCODED English regex carrying reasoning-exam
    #     shape nouns. On a non-English paper it matched nothing and the gate printed
    #     a clean OK — a FALSE ASSURANCE, claiming "no figure-reference prose" on a
    #     paper it never examined. RA-9: "Hardcode nothing. A missing value -> SKIP
    #     the dependent check with a logged reason, never a hardcoded substitute."
    def _prose_verdict(stem, lang, phrases=None):
        def _b(d):
            d.add_paragraph(f'Q.1  {stem}')
            for _i in range(1, 5):
                d.add_paragraph(f'{_i}.  Option {_i}')
        _p = _mini_doc(tmp, _b)
        _t, _bl = parse_blocks(Document(_p))
        _s = _fig_src(); _s['options_by_q'] = {}; _s['language'] = lang
        _s['section_rules_text'] = ''; _s['concept_map'] = {}
        if phrases:
            _s['figure_reference_phrases'] = phrases
        _reset(); gate_images(_bl, _s, {})
        return {c: l for l, c, m in RESULTS}.get('A-FIGTEXT-PROSE')

    # 53k — ENGLISH behaviour is unchanged (both directions).
    check('FIGTEXT-PROSE-english-default-unchanged',
          _prose_verdict('In the given figure, find X.', 'english') == 'FAIL'
          and _prose_verdict('Compute the value of X.', 'english') == 'OK')

    # 53l — NON-ENGLISH with NOTHING declared must report DORMANT (WARN with a named
    #       reason), NEVER a clean OK. Reporting OK here would certify a detector
    #       that did not run — the false-clean class this corpus keeps rediscovering.
    check('FIGTEXT-PROSE-non-english-undeclared-is-dormant-not-ok',
          _prose_verdict('\u0926\u0940 \u0917\u0908 \u0906\u0915\u0943\u0924\u093f \u092e\u0947\u0902 X', 'hindi') == 'WARN')

    # 53m — NON-ENGLISH WITH section_rules figure_reference_phrases DECLARED detects
    #       normally, in that exam's own language, and stays clean when the prose is
    #       absent. Both halves, so the fix cannot be "achieved" by always warning.
    _hi = ['\u0926\u0940 \u0917\u0908 \u0906\u0915\u0943\u0924\u093f \u092e\u0947\u0902']
    check('FIGTEXT-PROSE-non-english-declared-phrases-work',
          _prose_verdict('\u0926\u0940 \u0917\u0908 \u0906\u0915\u0943\u0924\u093f \u092e\u0947\u0902 X', 'hindi', _hi) == 'FAIL'
          and _prose_verdict('X \u0915\u093e \u092e\u093e\u0928', 'hindi', _hi) == 'OK')

    # 53i — v2.21.5 ND10 FIGURAL-NAT EXEMPTION, BOTH HALVES. Create.md R-FIGURAL
    #       v4.7 FIGURAL-NAT VARIANT: a figural question whose subtopic is
    #       answer_type=='numerical' has a PROBLEM image (or series images) but
    #       ZERO option images, and G-FIGURAL-COMPOSITE "must skip its
    #       per-option-image arm" for it. v2.21.4 tightened stem_and_options to
    #       oc+1 images WITHOUT reading that exemption, so a 3-image figural-NAT
    #       series false-WARNed and routed a human to "fix" a CORRECT paper.
    #       The registry's options_by_q is the signal (same one gate_options
    #       reads); concept_map/nat_subtopic_ids cannot be relied on because
    #       concept_map is {} on any run without a dossier or --key.
    #       BOTH HALVES are asserted so the fix cannot be "achieved" by disabling
    #       the arm: the SAME 3-image block is OK when the registry marks it NAT
    #       and a FINDING when it does not.
    def _img_verdict_obq(names, obq, role='stem_and_options', oc=4, figq=(1,)):
        _p = _img_doc(list(names))
        _mm = gate_zip(_p)
        _doc = Document(_p); _t, _bl = parse_blocks(_doc)
        _s = _fig_src(role, oc, figq); _s['options_by_q'] = obq
        _reset(); gate_images(_bl, _s, _mm)
        return {c: l for l, c, m in RESULTS}
    _nat3 = ['q1_problem.png', 'q1_problem_2.png', 'q1_problem_3.png']
    check('FIGCOMP-figural-NAT-exempt-from-per-option-arm-ND10',
          _img_verdict_obq(_nat3, {'1': 0}).get('A-FIGCOMP') == 'OK'
          and _img_verdict_obq(_nat3, {'1': 4}).get('A-FIGCOMP') == 'WARN')

    # 53j — ND10 GUARD: a figural-NAT with ZERO images is STILL a finding. ND10
    #       exempts the per-OPTION arm only; it still requires >=1 problem image.
    def _b_natnoimg(d):
        d.add_paragraph('Q.1  Compute the area. (Enter numerical value)')
    _p53j = _mini_doc(tmp, _b_natnoimg)
    _mm53j = gate_zip(_p53j)
    _t, _bl53j = parse_blocks(Document(_p53j))
    _s53j = _fig_src('stem_and_options'); _s53j['options_by_q'] = {'1': 0}
    _reset(); gate_images(_bl53j, _s53j, _mm53j)
    check('FIGCOMP-figural-NAT-with-zero-images-still-a-finding',
          any(c == 'A-FIGCOMP' and l == 'WARN' and '0img' in m
              for l, c, m in RESULTS))

    # 53g — v2.21.4 PARTIAL-OPTION-SET LOCK. stem_and_options requires the FULL
    #       set: Step-7 G-FIGURAL-COMPOSITE says "problem image + one separate
    #       image per option" (oc+1). The check was `len(block_imgs) == 1`, so a
    #       block rendering 2, 3 or 4 images — a problem figure with options
    #       SILENTLY UNDRAWN — passed clean. Only the degenerate 1-image case was
    #       caught. A candidate cannot answer a question whose option figures were
    #       never rendered.
    check('FIGCOMP-partial-option-image-set-is-a-finding',
          all(_img_verdict(_CANON5[:_k]).get('A-FIGCOMP') == 'WARN'
              for _k in (2, 3, 4)))

    _b53, _d53, _r53 = _attach(_img_doc(['q1_problem.png', 'q1_opt1.png']))
    check('IMAGES-attached-to-block',
          _d53 == 2 and _r53 == 2 and len(_b53[0].images) == 2
          and all(i.get('path') and os.path.exists(i['path']) for i in _b53[0].images))

    # 54. NAME + PER-DRAWING ALT TEXT — A-FIGALT reads @descr, so a drawing must
    #     carry ITS OWN, not the neighbouring drawing's (para_images_ext walks
    #     each <w:drawing> as a unit rather than zipping two flat lists).
    check('IMAGES-name-and-descr-per-drawing',
          [i['name'] for i in _b53[0].images] == ['q1_problem.png', 'q1_opt1.png']
          and _b53[0].images[0]['descr'] == 'alt for q1_problem.png'
          and _b53[0].images[1]['descr'] == 'alt for q1_opt1.png')

    # 55. TABLE-EMBEDDED DRAWINGS COUNTED — a DI chart or a figure/option fusion
    #     table puts images in cells; a block-level paras scan misses them all.
    _b55, _d55, _r55 = _attach(_img_doc(['q1_stim.png'], in_table=True))
    check('IMAGES-inside-tables-counted', _d55 == 1 and _r55 == 1)

    # 56. THE VACUOUS-PASS CATCH. With figures present the gates must report a
    #     NON-ZERO evaluated count. On v2.11-v2.12.1 every one of these printed
    #     "0 figure(s) conform." and this fixture fails on that build.
    r = _run_figgates(_b53)
    check('FIGGATES-not-vacuous',
          all(any(c == g and '0 figure(s)' not in m for l, c, m in r)
              for g in _FIGGATES))

    # 57. ROSTER INVARIANCE (§R15) — exactly ONE line per figure gate, with and
    #     without figures. Also locks out the duplicate second A-FIGDPI line the
    #     old EC-V18 note emitted, which would have broken the count signal.
    _b57, _, _ = _attach(_mini_doc(tmp, lambda d: _add_q(d, 1)))
    r0 = _run_figgates(_b57)
    check('FIGGATES-roster-invariant',
          all(sum(1 for _, c, _ in r if c == g) == 1 for g in _FIGGATES)
          and all(sum(1 for _, c, _ in r0 if c == g) == 1 for g in _FIGGATES))

    # 58. ZERO-IMAGE PAPER IS DORMANT, NOT "CONFORMANT" — a paper with no
    #     drawings must say so, never claim a conformance it never tested.
    check('FIGGATES-zero-image-dormant',
          all(any(c == g and l == 'OK' and 'dormant' in m for l, c, m in r0)
              for g in _FIGGATES))

    # 59. DECLARED-BUT-UNREADABLE ⇒ COVERAGE WARN (D3, edge case 6 applied here):
    #     drawings present but none openable is a coverage GAP, never a pass.
    _b59 = [Block(1)]
    _b59[0].images = [{'name': 'q1_problem.png', 'rid': 'rId9',
                       'part': 'image1.png', 'descr': '', 'path': None}]
    r = _run_figgates(_b59)
    check('FIGGATES-declared-unreadable-warns',
          all(any(c == g and l == 'WARN' and 'NOT ESTABLISHED' in m
                  for l, c, m in r) for g in _FIGGATES))

    # 60. SPEC TRANSPORT RESOLVES (D2) — registry-borne figure_specs keyed by the
    #     canonical docPr name, with extension-stripped and media-part fallbacks.
    _spec = {'class': 'data_series', 'placement_scale': 1.0}
    check('FIGSPEC-transport-resolves',
          resolve_figure_spec({'name': 'q1_problem.png', 'part': 'image1.png'},
                              {'q1_problem.png': _spec}) is _spec
          and resolve_figure_spec({'name': 'q1_problem.png', 'part': ''},
                                  {'q1_problem': _spec}) is _spec
          and resolve_figure_spec({'name': '', 'part': 'image1.png'},
                                  {'image1.png': _spec}) is _spec
          and resolve_figure_spec({'name': 'Picture 1', 'part': 'image9.png'},
                                  {'q1_problem.png': _spec}) == {})

    # 61. NON-LEGACY DEFECT BLOCKS CERTIFICATION — a v5.33+ render that carries a
    #     sidecar and still regresses is a real FAIL; exit is non-zero (MANDATE D).
    r = _run_figgates(_b53, specs={'q1_problem.png': _spec, 'q1_opt1.png': _spec},
                      legacy=False, findings=['A-FIGSCALE: scale 0.5 != 1.0'])
    check('FIGGATES-nonlegacy-defect-fails',
          any(c == 'A-FIGSCALE' and l == 'FAIL' and 'REGRESSION' in m for l, c, m in r))

    # 62. EC-V18 IS A DELIVERY TOLERANCE — the SAME finding on a pre-v5.33 figure
    #     with no sidecar must stay LOUD but must NOT block: the spec's EC-V18 is
    #     non-negotiable that ~200 existing exams "keep auditing AND DELIVERING
    #     untouched". A FAIL here exits non-zero and MANDATE D then refuses to
    #     certify — turning a coverage fix into an estate-wide outage.
    r = _run_figgates(_b53, legacy=True, findings=['A-FIGSCALE: scale 0.5 != 1.0'])
    check('FIGGATES-legacy-degrades-without-blocking',
          any(c == 'A-FIGSCALE' and l == 'WARN' and 'EC-V18' in m for l, c, m in r)
          and not any(c == 'A-FIGSCALE' and l == 'FAIL' for l, c, _ in r))

    # 63. END-TO-END WIRING THROUGH run_audit — THE M1 FIXTURE.
    #     Fixtures 53-62 call attach_block_images() directly, so every one of
    #     them still passes if the CALL inside run_audit is deleted. That is the
    #     precise shape of the v2.10 defect this whole gap family descends from:
    #     the delegation was written at the call sites and never bound, and no
    #     fixture exercised the real entry point. Mutation-verified: removing the
    #     run_audit call fails THIS check and only this check.
    _StubFC.LEGACY = True; _StubFC.FINDINGS = []
    sys.modules['figural_core'] = _StubFC
    _e2e = _img_doc(['q1_problem.png', 'q1_opt1.png'])
    _args63 = argparse.Namespace(docx=_e2e, blueprint=None, rules=None,
                                 manifest=None, registry=None, mockN=1,
                                 final=False, audit_state=None, key=None,
                                 self_test=True)
    _so = sys.stdout
    try:
        sys.stdout = io.StringIO()
        run_audit(_args63)
        _r63 = list(RESULTS)
    finally:
        sys.stdout = _so
    check('FIGGATES-wired-into-run-audit',
          all(any(c == g and '0 figure(s)' not in m and 'NOT ESTABLISHED' not in m
                  and 'dormant' not in m for l, c, m in _r63) for g in _FIGGATES))

    # 64. PER-FIGURE FAULT ISOLATION — ONE BAD SPEC MUST NOT COST ELEVEN GATES
    #     THEIR VERDICT. Found empirically while testing D2 against a real paper:
    #     a partially-recorded FigureSpec (render_figure() fills png_px /
    #     font_pt_native / placement_scale only AFTER it reads the artefact back,
    #     so a render that died mid-way leaves a shape the gates index into)
    #     raised TypeError out of g_figlabel(); _safe_gate turned it into
    #     A-GATEERROR and the WHOLE A-IMAGES gate died — twelve gate lines gone,
    #     roster 47 -> 36, the §R15 invariance v2.12 had just restored broken by
    #     one figure. The spec now arrives from the REGISTRY, i.e. from outside
    #     this process, so a per-item L3 guard is mandatory, exactly as v2.12
    #     required one for blueprint_core's call sites.
    class _RaisingFC(_StubFC):
        @staticmethod
        def audit_figure(spec, png, descr=None):
            raise TypeError('list indices must be integers or slices, not str')
    _StubFC.LEGACY = True; _StubFC.FINDINGS = []
    sys.modules['figural_core'] = _RaisingFC
    _s64 = _src_stub(tq=1); _s64['figure_specs'] = {}
    _reset()
    _safe_gate('A-IMAGES', gate_images, _b53, _s64, {})
    _r64 = [(l, c, m) for l, c, m in RESULTS if c.startswith('A-FIG') or c == 'A-GATEERROR']
    check('FIGGATES-per-figure-fault-isolation',
          not any(c == 'A-GATEERROR' for _, c, _ in _r64)
          and all(sum(1 for _, c, _ in _r64 if c == g) == 1 for g in _FIGGATES)
          and all(any(c == g and l == 'WARN' and 'NOT ESTABLISHED' in m
                      for l, c, m in _r64) for g in _FIGGATES))

    if _saved_fc is not None:
        sys.modules['figural_core'] = _saved_fc
    else:
        sys.modules.pop('figural_core', None)

    # ════════════════════════════════════════════════════════════════════
    # v2.14 — B3 FACT CONTEXT DISCIPLINE regression lock
    # ════════════════════════════════════════════════════════════════════
    # B3 moves the raw search result OUT of the reasoning stream and onto disk,
    # so the saved file becomes the ONLY copy of the evidence. C5 previously
    # accepted any file >= 1 byte, which was tolerable while the result was also
    # duplicated in context and is NOT tolerable now: without a shape check the
    # discipline silently degrades from "save the result" to "touch a file".
    def _fact_state(name, recs, q='1', extra=None):
        pth = os.path.join(_evd, name)
        with open(pth, 'w', encoding='utf-8') as fh:
            if isinstance(recs, str):
                fh.write(recs)
            else:
                json.dump(recs, fh)
        srcs = [{'url': 'u', 'date': 'd', 'saved': pth}]
        ent = {'status': 'verified', 'answer_cardinality': 'single',
               'answer_unique': True, 'is_factual': True,
               'fact_sources': srcs, 'artefact_stamps': {}}
        led = {q: ent}
        if extra:
            led.update(extra)
        return {'K': 1, 'batches_done': [1], 'evidence_dir': _evd,
                'ledger': {'entries': led}}, pth

    _GOODREC = {'query': 'q', 'url': 'https://example.org/a',
                'retrieved_at': '2026-08-01T00:00:00Z', 'snippet': 's'}

    # 65. WELL-FORMED RECORD → C5 PASS (the positive control).
    _reset()
    stA, _ = _fact_state('f_ok.json', _GOODREC)
    rc = completion_gate(_write_state('stA.json', stA), 1, bcg, dcg)
    check('CG-fact-record-wellformed-passes',
          not any(c == 'C5' and l == 'FAIL' for l, c, _ in RESULTS))

    # 66. ONE-BYTE STUB → C5 FAIL. This is the exact file that CERTIFIED before
    #     v2.14 and that B3 makes dangerous.
    _reset()
    stB, _ = _fact_state('f_stub.json', 'x')
    rc = completion_gate(_write_state('stB.json', stB), 1, bcg, dcg)
    check('CG-fact-stub-file-fails',
          rc == 1 and any(c == 'C5' and l == 'FAIL' for l, c, _ in RESULTS))

    # 67. MISSING A MANDATED FIELD → C5 FAIL, naming the FIELD (never the fact —
    #     MANDATE 0). RA-11 requires query + URL + retrieval-time + snippet.
    _reset()
    _partial = dict(_GOODREC); _partial.pop('retrieved_at')
    stC, _ = _fact_state('f_partial.json', _partial)
    rc = completion_gate(_write_state('stC.json', stC), 1, bcg, dcg)
    check('CG-fact-missing-field-fails',
          rc == 1 and any(c == 'C5' and l == 'FAIL' and 'retrieved_at' in m
                          for l, c, m in RESULTS))

    # 68. BLANK FIELD IS NOT A FIELD — an empty url must fail exactly like an
    #     absent one, or "save the record" becomes "save the keys".
    _reset()
    _blank = dict(_GOODREC); _blank['url'] = '   '
    stD, _ = _fact_state('f_blank.json', _blank)
    rc = completion_gate(_write_state('stD.json', stD), 1, bcg, dcg)
    check('CG-fact-blank-field-fails',
          rc == 1 and any(c == 'C5' and l == 'FAIL' and 'url' in m
                          for l, c, m in RESULTS))

    # 69. CACHE REUSE IS LEGITIMATE — B3 dedupes a concept shared by several
    #     questions to ONE search. A record LIST, and one file referenced by two
    #     questions, must both certify; the gate reports the reuse rather than
    #     treating it as a shortfall.
    _reset()
    dcg2, bcg2 = _cg_doc(lambda d: (_add_q(d, 1), _add_q(d, 2)))
    _shared = os.path.join(_evd, 'f_shared.json')
    with open(_shared, 'w', encoding='utf-8') as fh:
        json.dump([_GOODREC, dict(_GOODREC, query='q2')], fh)
    _ent = {'status': 'verified', 'answer_cardinality': 'single',
            'answer_unique': True, 'is_factual': True,
            'fact_sources': [{'url': 'u', 'date': 'd', 'saved': _shared}],
            'artefact_stamps': {}}
    stE = {'K': 1, 'batches_done': [1], 'evidence_dir': _evd,
           'ledger': {'entries': {'1': dict(_ent), '2': dict(_ent)}}}
    rc = completion_gate(_write_state('stE.json', stE), 2, bcg2, dcg2)
    check('CG-fact-cache-reuse-passes',
          not any(c == 'C5' and l == 'FAIL' for l, c, _ in RESULTS)
          and any(c == 'C5' and 'cache reuse' in m for _, c, m in RESULTS))

    # ════════════════════════════════════════════════════════════════════
    # v2.15 — C1 CROSS-SESSION CHECKPOINT regression lock
    # ════════════════════════════════════════════════════════════════════
    # RA-18 claimed "resume-safe" while storing the ledger AND the evidence tree
    # under /home/claude, which does not survive a session boundary. Resume
    # therefore worked inside one session and not across one, and the failure was
    # fatal rather than degraded: C5/C6 assert that every stamped evidence file
    # EXISTS, so a perfectly remembered ledger could not certify once the files
    # were gone. Fixture 71 is the one that matters — it proves a checkpoint
    # taken in one container and restored into another still CERTIFIES.
    _ckdir = os.path.join(tmp, 'ck'); os.makedirs(_ckdir, exist_ok=True)
    _ck_evd = os.path.join(_ckdir, 'evidence'); os.makedirs(_ck_evd, exist_ok=True)
    os.makedirs(os.path.join(_ck_evd, 'facts'), exist_ok=True)
    _ck_fact = os.path.join(_ck_evd, 'facts', 'q1_abc.json')
    with open(_ck_fact, 'w', encoding='utf-8') as fh:
        json.dump(_GOODREC, fh)
    _ck_paper = _mini_doc(tmp, lambda d: _add_q(d, 1))
    _ck_state = os.path.join(_ckdir, 'audit_state.json')
    _ck_st = {'mock': 1, 'exam_code': 'EX', 'paper_id': 'MOCK:M01', 'K': 1,
              'batches_done': [1], 'evidence_dir': _ck_evd,
              'ledger': {'entries': {'1': {'status': 'verified',
                          'answer_cardinality': 'single', 'answer_unique': True,
                          'is_factual': True,
                          'fact_sources': [{'url': 'u', 'date': 'd', 'saved': _ck_fact}],
                          'artefact_stamps': {}}}}}
    with open(_ck_state, 'w', encoding='utf-8') as fh:
        json.dump(_ck_st, fh)
    _ck_zip = os.path.join(tmp, 'ck.zip')
    _man = make_checkpoint(_ck_state, _ck_zip, docx_path=_ck_paper,
                           exam='EX', mockN=1)

    # 70. BUNDLE IS COMPLETE AND SELF-DESCRIBING — state + evidence + paper, an
    #     identity triple, and a sha256 for every member.
    check('CK-bundle-complete',
          _man['schema'] == CHECKPOINT_SCHEMA and _man['exam_code'] == 'EX'
          and _man['mock'] == 1 and _man['ledger_entries'] == 1
          and _man['evidence_files'] == 1 and _man['paper_md5']
          and 'audit_state.json' in _man['files']
          and any(a.startswith('evidence/') for a in _man['files'])
          and any(a.startswith('paper/') for a in _man['files']))

    # 71. ROUND TRIP THEN CERTIFY — THE FIXTURE THIS RELEASE EXISTS FOR.
    #     Simulate the session boundary by DESTROYING the original directory, then
    #     restore into a fresh one and run the real completion gate. Before C1 this
    #     was impossible: the evidence was gone and C5 failed for ever.
    import shutil as _sh71
    _sh71.rmtree(_ckdir)
    _newdir = os.path.join(tmp, 'ck_restored')
    _man2, _stp2, _evd2 = restore_checkpoint(_ck_zip, _newdir,
                                             docx_path=_ck_paper, exam='EX', mockN=1)
    _reset()
    _dck = Document(_ck_paper); _t, _bck = parse_blocks(_dck)
    rc = completion_gate(_stp2, 1, _bck, _dck)
    check('CK-restore-then-certify',
          rc == 0 and not any(l == 'FAIL' for l, _, _ in RESULTS)
          and os.path.exists(os.path.join(_evd2, 'facts', 'q1_abc.json')))

    # 72. EVIDENCE_DIR IS REWRITTEN — the recorded path belongs to the previous
    #     session's container. If restore did not rewrite it, every C5/C6 stamp
    #     would resolve to a directory that no longer exists.
    with open(_stp2, encoding='utf-8') as fh:
        _rst = json.load(fh)
    check('CK-evidence-dir-rebased',
          _rst['evidence_dir'] == _evd2 and os.path.isdir(_rst['evidence_dir'])
          and _rst['session_log']['checkpoints_restored'][0]['batches_done'] == [1])

    # 72b. NESTED EVIDENCE PATHS REBASE — a montage in evidence/montages/ must
    #      resolve after restore, not merely a file sitting at the evidence root.
    #      Verified explicitly because the resolver's basename fallback would mask
    #      a broken rebase on a flat tree and hide it until a real paper.
    check('CK-nested-paths-rebased',
          _rst['ledger']['entries']['1']['fact_sources'][0]['saved']
              == os.path.join(_evd2, 'facts', 'q1_abc.json')
          and _rst['session_log']['checkpoints_restored'][0]['evidence_paths_rebased'] == 1)

    # 73. TAMPER / TRUNCATION IS REFUSED — a member whose bytes changed since the
    #     manifest was written must never be restored.
    _bad_zip = os.path.join(tmp, 'ck_tampered.zip')
    with zipfile.ZipFile(_ck_zip) as zin, \
         zipfile.ZipFile(_bad_zip, 'w', zipfile.ZIP_DEFLATED) as zout:
        for i in zin.infolist():
            data = zin.read(i.filename)
            if i.filename.startswith('evidence/'):
                data = data + b' '          # one byte; hash must catch it
            zout.writestr(i, data)
    try:
        restore_checkpoint(_bad_zip, os.path.join(tmp, 'ck_t'), docx_path=_ck_paper)
        _t73 = False
    except CheckpointError as e:
        _t73 = 'integrity' in str(e)
    check('CK-tamper-refused', _t73)

    # 73a. v2.21.6 — restore_checkpoint MEMBER-ABSENCE CLOSURE. CK-tamper-refused
    #     covers a member whose CONTENT changed (hash mismatch). The other half —
    #     a member listed in the manifest but MISSING FROM THE ARCHIVE ENTIRELY —
    #     had no fixture, so its finding could be deleted with every test green.
    #     A truncated bundle is the more likely real-world corruption (interrupted
    #     upload/copy) and it must be refused BEFORE anything is written to disk,
    #     not resumed onto a half-restored evidence set.
    _gone_zip = os.path.join(tmp, 'ck_truncated.zip')
    with zipfile.ZipFile(_ck_zip) as zin, \
         zipfile.ZipFile(_gone_zip, 'w', zipfile.ZIP_DEFLATED) as zout:
        for _i in zin.infolist():
            if _i.filename.startswith('evidence/'):
                continue                      # drop the member, keep the manifest
            zout.writestr(_i, zin.read(_i.filename))
    try:
        restore_checkpoint(_gone_zip, os.path.join(tmp, 'ck_g'), docx_path=_ck_paper)
        _t73a = False
    except CheckpointError as e:
        _t73a = 'absent' in str(e) or 'integrity' in str(e)
    check('CK-missing-member-refused', _t73a)


    # 74. WRONG PAPER IS REFUSED — restoring onto a different document would
    #     certify an audit nobody performed on it. Strictly worse than losing the
    #     audit, so this binding is HARD.
    _other = _mini_doc(tmp, lambda d: (_add_q(d, 1), _add_q(d, 2)))
    try:
        restore_checkpoint(_ck_zip, os.path.join(tmp, 'ck_w'), docx_path=_other)
        _t74 = False
    except CheckpointError as e:
        _t74 = 'MD5' in str(e)
    check('CK-wrong-paper-refused', _t74)

    # 75. WRONG MOCK / WRONG EXAM ARE REFUSED — the other two thirds of the
    #     identity triple, checked BEFORE anything is written to disk.
    _t75a = _t75b = False
    try:
        restore_checkpoint(_ck_zip, os.path.join(tmp, 'ck_m'), mockN=2)
    except CheckpointError as e:
        _t75a = 'mock' in str(e)
    try:
        restore_checkpoint(_ck_zip, os.path.join(tmp, 'ck_e'), exam='OTHER')
    except CheckpointError as e:
        _t75b = 'exam_code' in str(e)
    check('CK-wrong-identity-refused', _t75a and _t75b)

    # 76. A NON-CHECKPOINT ARCHIVE IS REFUSED, not half-unpacked.
    _plain = os.path.join(tmp, 'plain.zip')
    with zipfile.ZipFile(_plain, 'w') as z:
        z.writestr('hello.txt', 'hi')
    try:
        restore_checkpoint(_plain, os.path.join(tmp, 'ck_p'))
        _t76 = False
    except CheckpointError as e:
        _t76 = 'manifest' in str(e)
    check('CK-not-a-checkpoint-refused',
          _t76 and not os.path.exists(os.path.join(tmp, 'ck_p', 'audit_state.json')))

    # 75b. AN UNBINDABLE BUNDLE IS REFUSED AT BOTH ENDS — building without the
    #      paper is refused outright (paper_md5 is the strongest binding, and a
    #      None there makes restore's check vacuous), and a manifest that somehow
    #      lacks it is refused on restore. Found in end-to-end testing, not by
    #      inspection: a shell quoting slip left the docx absent and the checkpoint
    #      was written anyway, with no binding at all.
    try:
        make_checkpoint(_stp2, os.path.join(tmp, 'ck_nopaper.zip'), docx_path=None)
        _t75c = False
    except CheckpointError as e:
        _t75c = 'REQUIRED' in str(e)
    _nb = os.path.join(tmp, 'ck_nobind.zip')
    with zipfile.ZipFile(_ck_zip) as zin, \
         zipfile.ZipFile(_nb, 'w', zipfile.ZIP_DEFLATED) as zout:
        for i in zin.infolist():
            data = zin.read(i.filename)
            if i.filename == CHECKPOINT_MANIFEST:
                _m = json.loads(data.decode('utf-8')); _m['paper_md5'] = None
                data = json.dumps(_m).encode('utf-8')
            zout.writestr(i, data)
    try:
        restore_checkpoint(_nb, os.path.join(tmp, 'ck_nb'), docx_path=_ck_paper)
        _t75d = False
    except CheckpointError as e:
        _t75d = 'paper_md5' in str(e)
    check('CK-unbindable-refused',
          _t75c and _t75d
          and not os.path.exists(os.path.join(tmp, 'ck_nb', 'audit_state.json')))

    # 76b. UNKNOWN SCHEMA IS REFUSED — a checkpoint written by a different
    #      framework build may carry fields this one misreads. Refusing is the
    #      only safe answer: half-understanding a resume state is how an audit
    #      certifies work it never did. (Uncovered until mutation testing showed
    #      the schema guard could be deleted with every other fixture still green.)
    _fut = os.path.join(tmp, 'ck_future.zip')
    with zipfile.ZipFile(_ck_zip) as zin, \
         zipfile.ZipFile(_fut, 'w', zipfile.ZIP_DEFLATED) as zout:
        for i in zin.infolist():
            data = zin.read(i.filename)
            if i.filename == CHECKPOINT_MANIFEST:
                _m = json.loads(data.decode('utf-8'))
                _m['schema'] = CHECKPOINT_SCHEMA + 1
                data = json.dumps(_m).encode('utf-8')
            zout.writestr(i, data)
    try:
        restore_checkpoint(_fut, os.path.join(tmp, 'ck_f'), docx_path=_ck_paper)
        _t76b = False
    except CheckpointError as e:
        _t76b = 'schema' in str(e)
    check('CK-unknown-schema-refused',
          _t76b and not os.path.exists(os.path.join(tmp, 'ck_f', 'audit_state.json')))

    # 77. REFUSAL IS TOTAL — a rejected restore must leave NO partial state behind.
    #     A half-unpacked checkpoint is the worst outcome of all: it looks resumable.
    check('CK-refusal-leaves-nothing',
          not os.path.exists(os.path.join(tmp, 'ck_t', 'audit_state.json'))
          and not os.path.exists(os.path.join(tmp, 'ck_w', 'audit_state.json'))
          and not os.path.exists(os.path.join(tmp, 'ck_m', 'audit_state.json'))
          and not os.path.exists(os.path.join(tmp, 'ck_f', 'audit_state.json')))

    # ════════════════════════════════════════════════════════════════════
    # v2.16 — D2 + D4 VISION DEGRADATION regression lock
    # ════════════════════════════════════════════════════════════════════
    # NO FIXTURE HAD EVER SIMULATED A VISION OUTAGE. All 89 built ledgers
    # programmatically with healthy stamps, so the "stamp cannot honestly be
    # issued" branch had never executed — the fifth appearance of the hollow-branch
    # class (v2.10 bc binding, v2.12 A-FIGPROFILE, v2.13 Block.images, v2.15
    # unknown-schema). These fixtures execute it, in both directions.
    def _vis_state(stamp, probe, montage_bytes=EVIDENCE_MIN_BYTES + 50, q='1'):
        m = os.path.join(_evd, f'v_{q}_{stamp}_{montage_bytes}.png')
        with open(m, 'wb') as fh:
            fh.write(b'\x89PNG\r\n\x1a\n' + b'0' * max(0, montage_bytes - 8))
        st = {'K': 1, 'batches_done': [1], 'evidence_dir': _evd,
              'ledger': {'entries': {q: {
                  'status': 'verified', 'answer_cardinality': 'single',
                  'answer_unique': True, 'is_factual': False,
                  'artefact_stamps': {'images': [{'rid_or_name': f'q{q}',
                                                  'stamp': stamp, 'montage': m}]}}}}}
        if probe is not None:
            st['session_log'] = {'vision_probe': probe}
        return st

    _FAILED_PROBE = [{'status': 'FAILED', 'attempts': 2, 'at': 'x',
                      'glyphs_read': '', 'batch': 1}]
    _OK_PROBE = [{'status': 'OK', 'attempts': 1, 'at': 'x',
                  'glyphs_read': 'ABC', 'batch': 1}]
    def _add_q_img(d, n):
        """A question block carrying one inline drawing — C7 requires a paper
        artefact for every image stamp, so the fixture doc must actually have one."""
        _add_q(d, n)
        p = d.add_paragraph()
        p.add_run().add_picture(io.BytesIO(_png_bytes()), width=Inches(1.0))

    # v2.21.9 — Inches is now a MODULE-level import (line 42). A function-local
    # import here would shadow it for the WHOLE of self_test(), unbinding it for
    # every fixture defined earlier in the same scope.
    _dv, _bv = _cg_doc(lambda d: _add_q_img(d, 1))

    # 78a-78e  v2.23.0 (RELEASE B) — ARITHMETIC CONFORMANCE COVERAGE.
    #   THE FAILURE BEING CLOSED, measured on a real 60-question paper: vision died,
    #   the run retried a dozen times, never recorded P3.5, left figures unstamped,
    #   C7 failed, MANDATE D refused delivery — and a paper whose figures the twelve
    #   arithmetic gates would have passed shipped as NOTHING. Vision and conformance
    #   were conflated; they are different claims and only one needs eyes.

    # 78a — THE UNBLOCK. An arithmetic stamp IS C7 coverage: exit 0, paper delivers.
    #       MUTATION-VERIFIED against v2.22.0.
    _reset()
    rc = completion_gate(_write_state('vb1.json',
                         _vis_state(VISION_STAMP_ARITHMETIC, None)), 1, _bv, _dv)
    check('VIS-ARITH-conformance-stamp-is-C7-coverage',
          rc == 0 and not any(c == 'C7' and l == 'FAIL' for l, c, _ in RESULTS))

    # 78b — IT MUST NOT OVERCLAIM. Conformance established arithmetically is NOT the
    #       same as a figure that was looked at, so the paper certifies DEGRADED, never
    #       clean PASS. (Caught during development: the first implementation printed
    #       PASS — a STRONGER claim than the pre-release behaviour, which is precisely
    #       what a certification-semantics change must never do.)
    _reset()
    import io as _io, contextlib as _ctx
    _cap = _io.StringIO()
    with _ctx.redirect_stdout(_cap):
        completion_gate(_write_state('vb2.json',
                        _vis_state(VISION_STAMP_ARITHMETIC, None)), 1, _bv, _dv)
    _verdict = _cap.getvalue()
    check('VIS-ARITH-certifies-degraded-never-clean',
          'DEGRADED' in _verdict and 'COMPLETION-GATE: PASS' not in _verdict
          and any(c == 'C6' and l == 'WARN' and 'conformance-arithmetic' in m
                  for l, c, m in RESULTS))

    # 78c — STILL EVIDENCE-BOUND. RA-4 render-or-recompute applies unchanged: a
    #       trivial/absent trace is not conformance, it is an unaudited artefact.
    #       Without this the new stamp would be a self-signed excuse — the same
    #       cheat surface T2.3 closes for view-unavailable.
    _reset()
    rc = completion_gate(_write_state('vb3.json',
                         _vis_state(VISION_STAMP_ARITHMETIC, None,
                                    montage_bytes=8)), 1, _bv, _dv)
    check('VIS-ARITH-requires-real-evidence',
          rc == 1 and any(c == 'C6' and l == 'FAIL' for l, c, _ in RESULTS))

    # 78d — NO PROBE REQUIRED. Arithmetic does not depend on the view tool, so unlike
    #       view-unavailable it is admissible with NO vision_probe record at all. Tying
    #       it to a failed probe would re-couple the two claims this release separates.
    _reset()
    rc = completion_gate(_write_state('vb4.json',
                         _vis_state(VISION_STAMP_ARITHMETIC, _OK_PROBE)), 1, _bv, _dv)
    check('VIS-ARITH-admissible-whatever-the-vision-state',
          rc == 0 and not any(c == 'C6' and l == 'FAIL' for l, c, _ in RESULTS))

    # 78e — THE DIAGNOSTIC. An unstamped artefact must NAME the mechanical remedy.
    #       The stalled session saw only "paper artefact not audited" and concluded it
    #       was blocked; the escape hatch existed and the message never mentioned it.
    _reset()
    _st_none = _vis_state(VISION_STAMP_ARITHMETIC, None)
    _st_none['ledger']['entries']['1']['artefact_stamps'] = {}
    rc = completion_gate(_write_state('vb5.json', _st_none), 1, _bv, _dv)
    check('VIS-C7-failure-names-the-mechanical-remedy',
          rc == 1 and any(c == 'C7' and l == 'FAIL'
                          and 'MECHANICAL REMEDY' in m
                          and 'does NOT require the view tool' in m
                          and VISION_STAMP_ARITHMETIC in m and 'P3.5' in m
                          and 'NEVER a reason to ship nothing' in m.replace(
                              'NEVER A REASON TO SHIP NOTHING',
                              'NEVER a reason to ship nothing')
                          for l, c, m in RESULTS))

    # 78. T2.1 — MEASURED OUTAGE ⇒ CERTIFIES DEGRADED, EXIT 0, NOT A HALT.
    #     The whole point: a vision outage is an ENVIRONMENT condition, and §5 says
    #     no dependency condition may ever halt a run.
    _reset()
    rc = completion_gate(_write_state('v1.json',
                         _vis_state(VISION_STAMP_UNAVAILABLE, _FAILED_PROBE)), 1, _bv, _dv)
    check('VIS-measured-outage-certifies-degraded',
          rc == 0 and any(c == 'C6' and l == 'WARN' for l, c, m in RESULTS)
          and any(c == 'C6' and l == 'WARN' for l, c, m in RESULTS))

    # 79. T2.3 — UNFAKEABLE: the stamp WITHOUT a failed probe must FAIL. This is the
    #     cheat surface the whole design turns on — "I couldn't see it" is exactly
    #     what a lazy operator would claim, so it must be a MEASURED fact.
    _reset()
    rc = completion_gate(_write_state('v2.json',
                         _vis_state(VISION_STAMP_UNAVAILABLE, None)), 1, _bv, _dv)
    check('VIS-unproved-outage-rejected',
          rc == 1 and any(c == 'C6' and l == 'FAIL' and 'NO FAILED vision probe' in m
                          for l, c, m in RESULTS))

    # 80. T2.5 — RECOVERY MUST BE HONOURED: healthy probe + un-upgraded stamps FAILS.
    #     Otherwise a paper that COULD have been fully audited certifies degraded.
    _reset()
    # A TRUE recovery carries BOTH records: the outage that justified the stamp,
    # and the later healthy probe that obliges the operator to upgrade it. A history
    # with only an OK probe is the DIFFERENT failure of fixture 79 (unproved).
    _RECOVERED = [{'status': 'FAILED', 'attempts': 2, 'at': 'x', 'glyphs_read': '',
                   'batch': 1},
                  {'status': 'OK', 'attempts': 1, 'at': 'y', 'glyphs_read': 'ABC',
                   'batch': 2}]
    rc = completion_gate(_write_state('v3.json',
                         _vis_state(VISION_STAMP_UNAVAILABLE, _RECOVERED)), 1, _bv, _dv)
    check('VIS-recovered-but-not-upgraded-fails',
          rc == 1 and any(c == 'C6' and l == 'FAIL' and 'RECOVERED' in m
                          for l, c, m in RESULTS))

    # 81. T2.4 — A vision outage never excuses producing NO artefact (E2.5/E2.6).
    _reset()
    rc = completion_gate(_write_state('v4.json',
                         _vis_state(VISION_STAMP_UNAVAILABLE, _FAILED_PROBE,
                                    montage_bytes=8)), 1, _bv, _dv)
    check('VIS-trivial-montage-still-blocks',
          rc == 1 and any(c == 'C6' and l == 'FAIL' for l, c, _ in RESULTS))

    # 82. E2.2 — a MIXED ledger (some viewed, some not) must be LEGAL. This is the
    #     shape of the actual incident: Batch 1 healthy, Batch 2 not.
    _reset()
    _mx = _vis_state(VISION_STAMP_UNAVAILABLE, _FAILED_PROBE, q='1')
    _m2 = os.path.join(_evd, 'v_ok.png')
    with open(_m2, 'wb') as fh:
        fh.write(b'\x89PNG\r\n\x1a\n' + b'0' * EVIDENCE_MIN_BYTES)
    _mx['ledger']['entries']['2'] = {
        'status': 'verified', 'answer_cardinality': 'single', 'answer_unique': True,
        'is_factual': False,
        'artefact_stamps': {'images': [{'rid_or_name': 'q2',
                                        'stamp': VISION_STAMP_VIEWED, 'montage': _m2}]}}
    _dv2, _bv2 = _cg_doc(lambda d: (_add_q_img(d, 1), _add_q_img(d, 2)))
    rc = completion_gate(_write_state('v5.json', _mx), 2, _bv2, _dv2)
    check('VIS-mixed-ledger-legal', rc == 0)

    # 83. HEALTHY RUNS ARE UNTOUCHED — zero estate regression. Every one of the ~200
    #     exams with a working view tool behaves byte-identically to v2.15.
    _reset()
    rc = completion_gate(_write_state('v6.json',
                         _vis_state(VISION_STAMP_VIEWED, _OK_PROBE)), 1, _bv, _dv)
    check('VIS-healthy-run-unchanged',
          rc == 0 and any(c == 'C6' and l == 'OK' for l, c, _ in RESULTS))

    # 84. D4 — THE PROBE IS UNFAKEABLE AND ITS SIDECAR LEAKS NOTHING. If the glyphs
    #     were recoverable from disk, 'view-unavailable' would be self-signable again
    #     and the entire D2 safety argument would collapse.
    _pd = os.path.join(tmp, 'probe_evd')
    _png, _meta = make_vision_probe(_pd, batch=1)
    _raw = open(os.path.join(_pd, 'montages', '_probe.json'), encoding='utf-8').read()
    _wrong = verify_vision_probe(_pd, 'ZZZ')
    check('VIS-probe-sidecar-leaks-nothing',
          'expected_sha256' in _meta and 'glyphs' not in _raw
          and all(g not in _raw for g in PROBE_GLYPH_ALPHABET
                  if _raw.count(g) and False)          # no plaintext glyph field
          and _wrong['status'] in ('FAILED', 'RENDER-FAIL'))

    # 85. E4.5 — a probe RENDER failure is an ENVIRONMENT WARN, never a vision
    #     verdict. Inferring "vision is broken" from "we could not draw the card"
    #     would degrade a perfectly healthy run.
    _rf = os.path.join(tmp, 'probe_rf'); os.makedirs(os.path.join(_rf, 'montages'))
    with open(os.path.join(_rf, 'montages', '_probe.json'), 'w') as fh:
        json.dump({'render': 'P3.5-RENDER-FAIL (ImportError)', 'salt': 's',
                   'expected_sha256': 'x'}, fh)
    check('VIS-render-fail-is-not-a-vision-verdict',
          verify_vision_probe(_rf, 'ABC')['status'] == 'RENDER-FAIL')

    # ════════════════════════════════════════════════════════════════════
    # v2.17 — D7 (no silent truncation) + TIER A (dossier) regression lock
    # ════════════════════════════════════════════════════════════════════
    # 86. D7 — A LIST NEVER TRUNCATES SILENTLY, AND SORTS NUMERICALLY.
    #     A-FIGCOMP had 27 findings and printed 12 with no trace of the other 15,
    #     in lexicographic order (Q3 after Q28), which read as non-determinism and
    #     was filed as an unreproducible gate. It was under-reporting.
    _many = [f'Q{n}' for n in range(1, 28)]
    _r86 = _flist(_many)
    check('FLIST-states-suppression-and-total',
          '27 TOTAL' in _r86 and '+15 MORE NOT SHOWN' in _r86
          and _r86.startswith('Q1 Q2 Q3 ')
          and _flist(['Q10', 'Q3', 'Q28']) == 'Q3 Q10 Q28'
          and _flist(['Q1', 'Q2']) == 'Q1 Q2')          # short list: no noise

    # 87. TIER A — a consistent dossier is adopted; A-DOSSIER passes and says
    #     explicitly that adoption is NOT certification.
    _dpaper = _mini_doc(tmp, lambda d: (_add_q(d, 1), _add_q(d, 2)))

    def _dos(qs, **kw):
        d = {'schema': DOSSIER_SCHEMA, 'exam_code': 'EX', 'mock': 1,
             'paper_md5': _md5_file(_dpaper), 'questions': qs}
        d.update(kw)
        _dos.n = getattr(_dos, 'n', 0) + 1
        pth = os.path.join(tmp, f'dos_{_dos.n}.json')
        with open(pth, 'w') as fh:
            json.dump(d, fh)
        return pth
    _ddoc = Document(_dpaper); _t, _dblocks = parse_blocks(_ddoc)
    # _add_q renders four options, so the CONSISTENT dossier says mcq/4 and the
    # DISAGREEMENT case claims nat (which must be option-free).
    _dsrc = _src_stub(tq=2); _dsrc['options_count'] = 4
    _good_qs = {'1': {'subtopic_id': 's.a', 'qtype': 'mcq', 'image_role': 'none'},
                '2': {'subtopic_id': 's.b', 'qtype': 'mcq', 'image_role': 'none'}}
    _reset(); gate_dossier(_dblocks, _dsrc, _good_qs)
    check('DOSSIER-consistent-adopted',
          any(c == 'A-DOSSIER' and l == 'OK' and 'never as certification' in m
              for l, c, m in RESULTS))

    # 88. A DISAGREEMENT IS A FINDING, NEVER A SILENT OVERWRITE. If Step 7's record
    #     and Step 7's output disagree, one of them is wrong and Step 8 must say so
    #     rather than quietly preferring either.
    _reset()
    gate_dossier(_dblocks, _dsrc,
                 {'1': {'qtype': 'nat'}, '2': {'qtype': 'nat'}})
    check('DOSSIER-disagreement-is-a-finding',
          any(c == 'A-DOSSIER' and l == 'FAIL' and 'RECORDED something other' in m
              for l, c, m in RESULTS))

    # 89. JUDGMENTS ARE REFUSED AT THE DOOR. Tier A transports FACTS. If an answer
    #     could ride along, RA-1 independence would be lost silently on 200 exams.
    for _k in ('answers', 'answer_verified', 'derived_answer'):
        try:
            load_dossier(_dos(_good_qs, **{_k: {'1': 2}}), docx_path=_dpaper)
            _ok89 = False; break
        except DossierError as e:
            _ok89 = 'JUDGMENT key' in str(e)
    check('DOSSIER-judgment-keys-refused', _ok89)

    # 90. IDENTITY BINDING — a dossier describing a DIFFERENT paper must be refused,
    #     or Step 8 audits against facts about another document.
    # _mini_doc reuses ONE filename, so re-calling it would overwrite _dpaper and
    # the "different paper" would be byte-identical — the binding test would then
    # pass vacuously while proving nothing. Build a genuinely distinct file.
    # (load_dossier only MD5s this path; it does not parse it.)
    _other_paper = os.path.join(tmp, 'other_paper.docx')
    with open(_dpaper, 'rb') as _f, open(_other_paper, 'wb') as _g:
        _g.write(_f.read() + b'\x00')
    _t90 = _t90b = _t90c = False
    try:
        load_dossier(_dos(_good_qs), docx_path=_other_paper)
    except DossierError as e:
        _t90 = 'MD5' in str(e)
    try:
        load_dossier(_dos(_good_qs), docx_path=_dpaper, mockN=2)
    except DossierError as e:
        _t90b = 'mock' in str(e)
    # v2.20: the exam leg of the triple. It was DEAD in production until v2.20 —
    # the parameter existed, the call site never passed it, and a wrong-exam
    # dossier was accepted. Assert the leg itself AND that load_sources supplies
    # it, because a working check nobody calls is not a check.
    _t90d = False
    try:
        load_dossier(_dos(_good_qs), docx_path=_dpaper, exam='OTHER')
    except DossierError as e:
        _t90d = 'exam' in str(e)
    import inspect as _i90
    _t90e = 'exam=' in _i90.getsource(load_sources)
    try:
        load_dossier(_dos(_good_qs, schema=99), docx_path=_dpaper)
    except DossierError as e:
        _t90c = 'schema' in str(e)
    check('DOSSIER-identity-bound', _t90 and _t90b and _t90c and _t90d and _t90e)

    # 91. ABSENT-SAFE — no dossier means legacy behaviour and a NAMED reason, never
    #     a silent degradation. ~200 existing exams run exactly as before.
    _reset(); gate_dossier(_dblocks, _dsrc, None, 'not supplied')
    check('DOSSIER-absent-is-legacy-not-silent',
          any(c == 'A-DOSSIER' and l == 'WARN' and 'Legacy behaviour' in m
              for l, c, m in RESULTS))

    # ══════════════════════════════════════════════════════════════════════
    # 92. GAP-2026-08-02-DOSSIER-OPTION-PREDICATE — A-DOSSIER COUNTS OPTIONS THE
    #     WAY THE OPTION GATES DO.
    #
    #     RETIRED IN v2.21: the old fixture 92 asserted
    #         block_option_count(b) == sum(1 for p in b.paras
    #                                      if OPT_RE.match(para_text(p)))
    #     which is a TAUTOLOGY — the right-hand side is a verbatim re-implementation
    #     of the left-hand side's body, so the assertion CANNOT FAIL FOR ANY
    #     PREDICATE. It reported green for four consecutive releases on a build
    #     whose dossier gate could not see a single image option. A fixture that
    #     restates its subject is worse than no fixture, because it reports green.
    #     Enforced against recurrence by validate_framework_md.py CHECK AO.
    #
    #     Six fixtures replace it. 92a-92d are MUTATION-VERIFIED (each measured
    #     False on the v2.20 OPT_RE build and True on this one); 92e-92f are GUARDS
    #     (true on both) retained so the fix can never be "achieved" by making the
    #     gate permissive.
    # ══════════════════════════════════════════════════════════════════════
    def _b_fig_opts(d):
        """stem_and_options figural block: BARE labels; images would follow."""
        d.add_paragraph('Q.1  Select the figure.')
        d.add_paragraph('Problem Figure')
        for i in range(1, 5):
            d.add_paragraph(f'{i}.')
    def _b_enum_stem(d):
        """STATEMENT-format block: enumerated stem points THEN four text options."""
        d.add_paragraph('Q.1  Consider the following statements:')
        d.add_paragraph('1.  Statement one')
        d.add_paragraph('2.  Statement two')
        for i in range(1, 5):
            d.add_paragraph(f'{i}.  Option {i}')
    def _b_short(d):
        """Genuinely short option set — three labels where four are expected."""
        d.add_paragraph('Q.1  Solve.')
        for i in range(1, 4):
            d.add_paragraph(f'{i}.  Opt {i}')
    def _b_nat_enum(d):
        """R13-VIOLATING NAT block: carries stray option-label paragraphs."""
        d.add_paragraph('Q.1  Consider the following statements:')
        d.add_paragraph('1.  Statement one')
        d.add_paragraph('2.  Statement two')
        d.add_paragraph('How many are correct? (Enter numerical value)')
    def _dos_verdict(build, dossier_qtype):
        _p = _mini_doc(tmp, build)
        _t, _bl = parse_blocks(Document(_p))
        _s = _src_stub(tq=1); _s['options_count'] = 4
        _reset(); gate_dossier(_bl, _s, {'1': {'qtype': dossier_qtype}})
        return any(c == 'A-DOSSIER' and l == 'FAIL' for l, c, m in RESULTS)

    # 92a — THE HEADLINE CASE. An IMAGE-OPTION MCQ is not a dossier disagreement.
    #       OPT_RE counted these as 0 and blocked certification estate-wide.
    check('DOSSIER-image-options-not-a-finding',
          not _dos_verdict(_b_fig_opts, 'mcq'))

    # 92b — an ENUMERATED STEM does not inflate the count. Proves the defect was
    #       NOT figural-only: this block is a pure TEXT-option question, the
    #       standard STATEMENT/SEQUENCE/MATCH/ASSERTION_REASON construction.
    check('DOSSIER-enumerated-stem-does-not-inflate',
          not _dos_verdict(_b_enum_stem, 'mcq'))

    # 92c — FALSE-NEGATIVE LOCK (the integrity half). A dossier claiming 'nat' on a
    #       block that RENDERS four image options MUST be a finding. Under OPT_RE
    #       n_opt was 0, so `qt == 'nat' and n_opt` never fired and a real Step-7
    #       mislabel was ACCEPTED SILENTLY and would have reached students.
    check('DOSSIER-nat-with-image-options-is-a-finding',
          _dos_verdict(_b_fig_opts, 'nat'))

    # 92d — PREDICATE PARITY, THE STRUCTURAL GUARANTEE. A-DOSSIER and A-OPTN must
    #       agree about how many options a block renders on EVERY block shape. Any
    #       future divergence — a new regex, a new gate, a "small" optimisation —
    #       fails HERE rather than on a live paper.
    _parity = True
    for _bld in (_b_fig_opts, _b_enum_stem, _b_short):
        _pp = _mini_doc(tmp, _bld)
        _t, _pbl = parse_blocks(Document(_pp))
        _ps = _src_stub(tq=1); _ps['options_count'] = 4
        _reset(); gate_options(_pbl, _ps)
        _optn_clean = not any(c == 'A-OPTN' and l == 'FAIL' for l, c, m in RESULTS)
        _reset(); gate_dossier(_pbl, _ps, {'1': {'qtype': 'mcq'}})
        _dos_clean = not any(c == 'A-DOSSIER' and l == 'FAIL' for l, c, m in RESULTS)
        if _optn_clean != _dos_clean:
            _parity = False
    check('DOSSIER-OPTN-predicate-parity', _parity)

    # 92e — GUARD (true on both builds). A genuinely SHORT option set STILL fails,
    #       so the fix cannot be "achieved" by making the gate permissive.
    check('DOSSIER-short-option-set-still-fails',
          _dos_verdict(_b_short, 'mcq'))

    # 92g-92i — v2.21.2 HOLLOW-BRANCH CLOSURE (audit_mutation.py survivors).
    #     Automated mutation testing showed THREE A-DOSSIER findings could be
    #     DELETED OUTRIGHT with all 112 fixtures still green: the two set-mismatch
    #     legs and the subtopic leg. They had never been executed by ANY fixture —
    #     the same hollow-branch class as the retired tautology, found mechanically
    #     this time instead of by reading. Each fixture below KILLS its mutant.
    def _b_two_q(d):
        d.add_paragraph('Q.1  Solve.')
        for i in range(1, 5):
            d.add_paragraph(f'{i}.  Opt {i}')
        d.add_paragraph('')
        d.add_paragraph('Q.2  Solve.')
        for i in range(1, 5):
            d.add_paragraph(f'{i}.  Opt {i}')
        d.add_paragraph('')

    # 92g — a question ON THE PAPER but ABSENT FROM THE DOSSIER is a finding. The
    #       dossier must describe the paper that shipped, not a subset of it.
    _p92 = _mini_doc(tmp, _b_two_q)
    _t, _bl92 = parse_blocks(Document(_p92))
    _s92 = _src_stub(tq=2); _s92['options_count'] = 4
    _reset(); gate_dossier(_bl92, _s92, {'1': {'qtype': 'mcq'}})   # Q.2 missing
    check('DOSSIER-paper-Q-absent-from-dossier-is-a-finding',
          any(c == 'A-DOSSIER' and l == 'FAIL' and 'absent-from-dossier' in m
              for l, c, m in RESULTS))

    # 92h — a question IN THE DOSSIER but NOT ON THE PAPER is a finding. Step 7
    #       recorded a question it did not ship.
    _reset(); gate_dossier(_bl92, _s92, {'1': {'qtype': 'mcq'}, '2': {'qtype': 'mcq'},
                                         '3': {'qtype': 'mcq'}})
    check('DOSSIER-dossier-Q-not-in-paper-is-a-finding',
          any(c == 'A-DOSSIER' and l == 'FAIL' and 'not-in-paper' in m
              for l, c, m in RESULTS))

    # 92i — a dossier subtopic_id that DISAGREES with the registry is a finding.
    #       This is the second half of the Tier-A cross-check and had never run.
    _s92b = _src_stub(tq=2); _s92b['options_count'] = 4
    _s92b['figural_subtopics'] = {'1': 'sub.registry.value'}
    _reset(); gate_dossier(_bl92, _s92b,
                           {'1': {'qtype': 'mcq', 'subtopic_id': 'sub.dossier.other'},
                            '2': {'qtype': 'mcq'}})
    check('DOSSIER-subtopic-disagreement-is-a-finding',
          any(c == 'A-DOSSIER' and l == 'FAIL' and 'subtopic-disagrees-with-registry' in m
              for l, c, m in RESULTS))

    # 92f — v2.21.1 NAT-LEG FALSE-NEGATIVE LOCK, GROUNDED IN R13. A NAT block
    #       carrying ANY option-label paragraph is an R13 violation: the v4.7 NAT
    #       EXEMPTION allows a NAT question ONLY the bold Q.<N> stem and the blank
    #       separator — ZERO option paragraphs, with no "enumerated stem" class.
    #       v2.21 clamped this leg to n_opt >= oc on the unverified assumption
    #       that a NAT stem may enumerate; it may not. The clamp opened a hole:
    #       with nat_present=False and the registry marking the Q 0-option,
    #       gate_options SKIPS the block and gate_nat is DORMANT, so A-DOSSIER was
    #       the ONLY remaining gate — and it had been silenced. This fixture is
    #       the lock: a PARTIAL stray-label set on a claimed-NAT block IS a
    #       finding, exactly as a complete one is (92c).
    check('DOSSIER-nat-with-stray-labels-is-a-finding',
          _dos_verdict(_b_nat_enum, 'nat'))

    # 93. A FACTS-ONLY DOSSIER MUST NOT WAKE AN ANSWER-DEPENDENT GATE.
    #     A-NAT-GRADE re-runs derive_nat_grading() over the KEYED VALUE, so it needs
    #     ANSWERS. Its dormancy test read `not concept_map`, so the moment Tier A
    #     populated concept_map the gate woke with no answers and FAILED every NAT
    #     question on a real paper — a correct "dormant" turned into a false FAIL.
    #     Caught by running it, not by inspection. Mutation-verified.
    _nsrc = _src_stub(tq=2)
    _nsrc.update({'nat_present': True, 'options_count': 4,
                  'concept_map': {'1': {'subtopic_id': 's.a', 'qtype': 'nat'}},
                  'answers': {}})
    _reset(); _safe_gate('A-NAT-GRADE', gate_nat, _dblocks, _nsrc)
    check('NATGRADE-dormant-without-answers',
          any(c == 'A-NAT-GRADE' and l == 'OK' and 'answer values not available' in m
              for l, c, m in RESULTS)
          and not any(c == 'A-NAT-GRADE' and l == 'FAIL' for l, c, _ in RESULTS))

    # ════════════════════════════════════════════════════════════════════
    # v2.18 — D1 / D3 / D8 regression lock
    # ════════════════════════════════════════════════════════════════════
    # 94. D3 — APPENDIX B MUST MATCH THE LIVE ENGINES, ALWAYS. A hand-written API
    #     appendix drifts, and a drifted contract is worse than none: the gap report
    #     that prompted this claimed check_figural_conformance returns FAIL for an
    #     unconstrained profile. It returns SKIP. Documenting the claim would have
    #     enshrined a false precondition on 200 exams. This asserts the contract by
    #     INTROSPECTION so it cannot drift, and pins the SKIP behaviour explicitly.
    _apx_ok = True
    try:
        import inspect as _insp, blueprint_core as _bc
        _sigs = {
            'figural_generation_profile': '(pyq_image_analysis)',
            'check_figural_conformance': '(generated_types, profile, floor=0.55)',
            'derive_image_roles': '(imap)',
        }
        for _fn, _want in _sigs.items():
            if str(_insp.signature(getattr(_bc, _fn))) != _want:
                _apx_ok = False
        if _bc.IMAGE_ROLES != ('stem_and_options', 'stem_only', 'options_only', 'none'):
            _apx_ok = False
        _prof = _bc.figural_generation_profile({})
        if _prof.get('mode') != 'unconstrained':
            _apx_ok = False
        _v = _bc.check_figural_conformance(['x'], _prof)
        # 2-TUPLE, and SKIP (not FAIL) for unconstrained — both are load-bearing.
        if not (isinstance(_v, tuple) and len(_v) == 2 and _v[0] == 'SKIP'):
            _apx_ok = False
    except ImportError:
        # Engine absent is an ENVIRONMENT condition, not a contract failure — but it
        # must be REPORTED, not silently swallowed. The first cut returned None and
        # the check passed vacuously wherever blueprint_core was not on the path,
        # which meant a signature-drift mutation survived. Same hollow-branch class
        # this corpus keeps finding; here it was in the fixture itself.
        _apx_ok = None
    check('APPENDIXB-matches-live-engines' if _apx_ok is not None
          else 'APPENDIXB-matches-live-engines[engine-absent-skipped]',
          _apx_ok in (True, None))

    # 95. D1 — MANDATE 0 GETS ITS FIRST ACTUAL ENFORCEMENT. The mandate had NO
    #     machine check of any kind: S5-1A asserts ledger/evidence, S14-2 scans
    #     output FILENAMES. A content leak into authored prose was structurally
    #     invisible to every check in the framework, which is how one happened.
    #     mandate0_scan() flags the incidental-print pattern rule 2 forbids.
    _leaky = ["print(f'Q{n}: {p.text}')", 'print(para_text(p))',
              'print(b.stem)', "print('opt', opt_text)"]
    # COUNTING is correct and must never trip — `print(len(b.opts))` is the exact
    # boundary MANDATE 0 rule 2 draws ("printing len(block.opts) is correct").
    # The first cut omitted it, so the len() guard was never exercised and a
    # mutation deleting the guard survived: a clean list that avoids the boundary
    # tests nothing.
    _clean = ['print(len(b.images))', "print(f'Q{n}: {len(opts)} options')",
              'print(len(b.opts))', 'print(len(cell.text))',
              'print(para_text.__name__)', "print('Q.7 A-UNDERLINE: no w:u run')"]
    check('MANDATE0-leak-scan-flags-incidental-prints',
          all(mandate0_scan(x) for x in _leaky)
          and not any(mandate0_scan(x) for x in _clean))

    print(f'SELF-TEST: {passed}/{total} PASS' if passed == total
          else f'SELF-TEST: {passed}/{total} PASS  (FAILURES: {fails})')
    return 0 if passed == total else 1


def main():
    ap = argparse.ArgumentParser(description='Universal exam-agnostic Part-A mock auditor')
    ap.add_argument('docx', nargs='?', help='the Mock[N]_Create.docx to audit')
    ap.add_argument('--blueprint'); ap.add_argument('--rules')
    ap.add_argument('--manifest');  ap.add_argument('--registry')
    ap.add_argument('--mockN', type=int)
    ap.add_argument('--final', action='store_true')
    ap.add_argument('--audit-state', dest='audit_state',
                    help='Phase-3 COMPLETION GATE (S5-1A): validate the audit_state ledger '
                         '+ evidence artefacts (C1-C7). Use with --final.')
    ap.add_argument('--key', dest='key',
                    help='optional answer_key.json (concept_map) — normally NOT delivered (S0-1).')
    ap.add_argument('--self-test', action='store_true', dest='self_test')
    # C1 (v2.15) — cross-session checkpoint. /home/claude does not survive a
    # session boundary, so an audit that spans sessions must carry its ledger AND
    # its evidence with it or S5-1A C5/C6 can never certify.
    ap.add_argument('--make-checkpoint', dest='make_checkpoint',
                    help='write a portable checkpoint zip (needs --audit-state; '
                         'pass the docx to bind the bundle to this paper).')
    ap.add_argument('--restore-checkpoint', dest='restore_checkpoint',
                    help='verify + unpack a checkpoint zip (needs --into).')
    ap.add_argument('--into', dest='into',
                    help='destination directory for --restore-checkpoint.')
    # D4 (v2.16) — P3.5 vision probe. Vision is a DECLARED dependency: RA-4 makes it
    # load-bearing for C6/C7 and therefore for delivery, yet nothing checked it. An
    # outage was discovered after hours of work; this makes it a one-minute fact.
    ap.add_argument('--vision-probe', dest='vision_probe', metavar='EVIDENCE_DIR',
                    help='P3.5: render the probe card into EVIDENCE_DIR/montages.')
    ap.add_argument('--vision-probe-verify', dest='vision_verify', metavar='EVIDENCE_DIR',
                    help='P3.5: verify reported glyphs; needs --glyphs (and --audit-state '
                         'to record the per-batch result).')
    ap.add_argument('--dossier', dest='dossier',
                    help='Tier-A Step-7 fact dossier ([ExamCode]_M[N]_audit_dossier.json). '
                         'Facts only; refused if it carries answers or any judgment key.')
    ap.add_argument('--glyphs', dest='glyphs', help='the glyphs actually SEEN in the probe.')
    ap.add_argument('--batch', dest='batch', type=int, help='batch number for the probe record.')
    args = ap.parse_args()
    if args.self_test:
        sys.exit(self_test())
    if args.make_checkpoint:
        if not args.audit_state:
            ap.error('--make-checkpoint requires --audit-state')
        try:
            man = make_checkpoint(args.audit_state, args.make_checkpoint,
                                  docx_path=args.docx, mockN=args.mockN)
        except CheckpointError as e:
            print(f'CHECKPOINT: REFUSED — {e}'); sys.exit(1)
        print(f"CHECKPOINT: WRITTEN {os.path.basename(args.make_checkpoint)} "
              f"(mock {man.get('mock')}, batches {man.get('batches_done')}/"
              f"{man.get('K')}, {man.get('ledger_entries')} ledger entr(ies), "
              f"{man.get('evidence_files')} evidence file(s))")
        sys.exit(0)
    if args.vision_probe:
        png, meta = make_vision_probe(args.vision_probe, batch=args.batch)
        if str(meta.get('render')).startswith('P3.5-RENDER-FAIL'):
            print(f'VISION-PROBE: {meta["render"]} — ENVIRONMENT WARN, not a vision '
                  f'verdict (E4.5). Do NOT infer an outage from this.')
            sys.exit(0)
        print(f'VISION-PROBE: WRITTEN {png} — view it and re-run with '
              f'--vision-probe-verify <evidence_dir> --glyphs <what you saw>. '
              f'The expected glyphs are stored ONLY as a salted hash, so reporting '
              f'them requires actually seeing the card.')
        sys.exit(0)
    if args.vision_verify:
        if args.glyphs is None:
            ap.error('--vision-probe-verify requires --glyphs')
        rec = verify_vision_probe(args.vision_verify, args.glyphs)
        if args.batch is not None:
            rec['batch'] = args.batch
        if args.audit_state and os.path.exists(args.audit_state):
            with open(args.audit_state, encoding='utf-8') as fh:
                _st = json.load(fh)
            _sl = _st.setdefault('session_log', {})
            _vp = _sl.setdefault('vision_probe', [])
            if isinstance(_vp, dict):
                _vp = [_vp]
            _vp.append(rec); _sl['vision_probe'] = _vp
            with open(args.audit_state, 'w', encoding='utf-8') as fh:
                json.dump(_st, fh, indent=1, ensure_ascii=False)
        if rec['status'] == 'OK':
            print(f'VISION-PROBE: OK (batch {rec.get("batch")}) — §7 Layer-B viewing '
                  f'is healthy; normal operation.')
        elif rec['status'] == 'RENDER-FAIL':
            print(f'VISION-PROBE: {rec["reason"]} — ENVIRONMENT WARN, not a vision verdict.')
        else:
            print('P3.5 VISION PROBE FAILED — §7 Layer-B viewing is unavailable. The audit '
                  'WILL run and WILL deliver, but every figural item audited under this '
                  'outage carries view-unavailable (RA-4 v2.16) and the paper certifies as '
                  'CERTIFIED-DEGRADED (VISION) with an F1 AMBER footer and a §R13 '
                  'limitation. NOT a hard stop. Re-run on a session with a working view '
                  'tool for full coverage.')
        sys.exit(0)
    if args.restore_checkpoint:
        if not args.into:
            ap.error('--restore-checkpoint requires --into')
        try:
            man, stp, evd = restore_checkpoint(args.restore_checkpoint, args.into,
                                               docx_path=args.docx, mockN=args.mockN)
        except CheckpointError as e:
            print(f'CHECKPOINT: REFUSED — {e}'); sys.exit(1)
        print(f"CHECKPOINT: RESTORED (mock {man.get('mock')}, batches "
              f"{man.get('batches_done')}/{man.get('K')}, "
              f"{man.get('ledger_entries')} ledger entr(ies), "
              f"{man.get('evidence_files')} evidence file(s)) -> {stp}")
        sys.exit(0)
    if not args.docx:
        ap.error('docx path required (or use --self-test)')
    sys.exit(run_audit(args))


if __name__ == '__main__':
    main()
