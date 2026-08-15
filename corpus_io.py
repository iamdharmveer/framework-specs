"""
corpus_io.py v1.11 — I/O shell for PYQ corpus acquisition, image integrity,
                    document size governance, tables and the Analysis doc.

v1.11 — 2026-07-29 — CLUSTER I RECORDED + ITS MISSING FIXTURES (debt closure).

    Closes the three follow-ups recorded in the 2026.07.29 deployment audit:

    (1) THIS ENTRY. The ~300-line Cluster I table-structure addition
        (GAP-2026-07-29-TBL: _table_rows rewritten as a raw w:tr/w:tc walk that
        inspects vMerge/gridSpan, read_table_spec(), the TableSpec model,
        normalise_table_spec/place_cells/table_spec_spans, build_di_table with a
        font_name param, adjacent_table_pairs, legacy {'headers','rows'} DI
        payload accepted for ever) shipped with the header still claiming v1.10
        — the third un-bumped occurrence of an addition landing without its
        version line. Bumped and recorded here.

    (2) CLUSTER I SELF-TEST FIXTURE. A gridSpan-4 header + vMerge-2 label table
        whose anchor-once expectations the old row.cells implementation CANNOT
        produce (row.cells repeats the anchor across every covered grid
        position), plus: a flat-table byte-identity check pinning the legacy
        list[list[str]] output, the spans round-trip through
        read_table_spec/table_spec_spans, the legacy DI payload shape, and a
        full build_di_table -> raw-XML read-back round trip asserting spans
        survive with zero stray empty cells. Reverting _table_rows to row.cells
        now breaks this suite instead of passing green — the exact revert that
        would have sailed through the previous 273/273.

    (3) is_option FIXTURE (open since 2026.07.26.2): OPT_PATTERNS text markers,
        the six-and-beyond / prose rejections, bare-marker + image-option (blip)
        recognition, imageless bare-marker rejection, Paragraph-vs-element
        duality of para_has_image, and clean_option_text marker stripping.

    Self-test 273 -> 303. No behaviour change to any shipped function.

v1.10 — 2026-07-27 — `fresh=` PARAMETER: RUN-SCOPED CALLERS OPT OUT OF THE UNION.

    Found while verifying the v1.9 union against its CALLERS, prompted by the
    self-test regression review. Step 1 (PYQPrepare) shares the same fixed
    VISION_WORKDIR as Step 5 but wants per-run semantics: it processes ONE source
    per trigger and completes Phase A->B->C inside that trigger, so anything already
    in the workdir is contamination from a previous paper or a previous step.
    v1.8's overwrite masked the sharing; the v1.9 union surfaced it (measured: a
    second PYQPrepare run saw queued=3 where it queued 1 — Phase B re-views the
    prior paper's cells, Phase C counts them unobserved, footer goes amber with
    wrong counts). fresh=True clears queue + sheets + observations before building.
    Default False preserves v1.9 union semantics for Step 5. Also bumps the version
    line so the three byte-distinct states that briefly all claimed v1.9 are
    unambiguous: v1.9 = union, v1.9+isolated self-tests = superseded, v1.10 = this.

v1.9 — 2026-07-27 — build_vision_queue() IS NOW IDEMPOTENT (GAP-2026-07-27-B).

    It wrote vision_queue.json and vision_sheet_NNN.png — both FIXED names — and never
    read what was already in the workdir. Any caller invoking it twice against one
    workdir destroyed the first call's work. Framework_MockTestAnalyse <= v2.38 called
    it once per PAPER into a workdir its own spec defines as one-per-RUN, so within a
    3-paper batch only the LAST paper reached Phase B.

    MEASURED (IIT_JAM_BIOTECHNOLOGY, 22 papers): 153 figural questions; _meta.vision
    recorded queued=8. Batch 7 queued 24 and 3 survived. Five separate sessions hit
    this and each invented a different workaround; none converged.

    It now unions the submitted items with any queue already on disk, keyed by
    (paper_id, q_num), before tiling. Safe by construction — bc.vision_tag_map derives
    a tag from paper_id alone, so re-tiling a superset leaves prior tags intact
    (EC-V11/V12). A hash-width widening is the one case that re-tags everything; it is
    now detected and reported via stats['tag_generation_changed'] rather than silently
    orphaning the previous generation's observations.

    stats gains carried_in / submitted / tag_generation_changed so a caller can
    distinguish a fresh queue from a resumed one without diffing files.

    This is HALF the fix. The caller must ALSO stop calling it per paper — see
    Framework_MockTestAnalyse v2.39. Hoisting alone still loses prior sessions;
    idempotence alone still loses papers 1..N-1 within a batch.

v1.8 — 2026-07-26 — VISION-PROBE FAMILY RETIRED (GAP-2026-07-26-003 follow-up).

    Deleted: normalise_for_view(), make_vision_probe(), score_vision_probe(), and the
    exception classes ProbeObservationMissing and VisionUnavailable.

    WHY. Cluster V's three-phase bridge subsumes all of them:
      * normalise_for_view was a public wrapper over the same normalisation
        build_vision_queue now performs internally via _open_rgb. Two spellings of one
        rule is the drift the DELEGATION contract exists to prevent.
      * make_vision_probe / score_vision_probe answered "can this session see?" by
        minting a token image. PHASE B looks at REAL figures, so the answer now arrives
        free as a by-product of the work: any observation returned proves vision works,
        none returned proves it does not. A separate probe would be a SECOND mechanism
        answering one question, free to disagree with the first.
      * ProbeObservationMissing had no remaining raiser once score_vision_probe went;
        VisionUnavailable was raised nowhere even before this release.

    audit_callgraph C4 flagged the first three as public-but-unreached. Left in place
    they would report 2 findings forever, and a check that always shows findings gets
    skimmed and then ignored — "a check that can be shipped past is decoration"
    (CLAUDE.md). Retiring them returns C4 to zero so the next real finding is visible.

    The v1.5 entry below is HISTORY and stays: it records why the empty-string verdict
    was wrong, which is the reasoning that produced the three-phase design.

v1.7 — 2026-07-26 — CLUSTER V: VISION PHASE A / OBSERVATION I/O
    (GAP-2026-07-26-003).

    Viewing an image is a CLASS T operation — it needs a TOOL CALL, and a tool call
    can only happen BETWEEN model turns. A python process launched from bash cannot
    suspend mid-loop to make one. Every "vision function" the corpus wrote as python
    was therefore unreachable code returning a default forever, on every run, in
    silence. Measured on IIT_JAM_BIOTECHNOLOGY: 153/153 figural questions recorded
    vision_unavailable, 45/45 FIGURAL subtopics shipped an empty object-type profile,
    and QV-9 returned PASS on the result.

    Adds the DETERMINISTIC half of the materialise-then-inject bridge:
      build_vision_queue()        PHASE A — normalise -> compose -> tile -> queue.json
      load_vision_queue()         reader, tolerant
      load_vision_observations()  PHASE B -> C boundary, tolerant
      write_vision_observations() one spelling of the observations file

    NOTHING IN CLUSTER V RAISES FOR A VISION REASON. An unopenable figure, a missing
    observations file and an absent Pillow all degrade with a recorded reason, and the
    run completes. The defect being fixed was SILENCE, and the remedy for silence is
    visibility (QV-14 + vision_status), never a halt.

    Two defects were found by property fuzzing rather than by example, and both are
    now locked by self-tests:
      * EC-V27 — two items sharing (paper_id, q_num) printed the SAME tag on two
        cells, making a Phase-B observation unattributable. Items are now merged by
        key with their sources unioned.
      * mixed-type tags in a model-written observations file raised TypeError inside
        sorted() (blueprint_core side).

v1.6 — 2026-07-26 — is_option() BECOMES THE SINGLE SHARED OPTION PREDICATE
    (audit_deep [XSPEC-DRIFT]).

    It was defined three times — MockTestAnalyse, PYQSort, PYQAnalyse — each claiming
    in its docstring to be aligned with the others. v2.34/v2.35 gave Step 5 the
    image-option path; the siblings kept the text-only form and the alignment claim
    silently became false.

    Not cosmetic: PYQSort USES its copy to count options (_count_options_in_body), so
    for an image option ("1." with a picture, no text) the predicate returned False
    and the option was UNDERCOUNTED — the identical defect class the wave had just
    fixed in Step 5, left live in Step 1.

    OPT_PATTERNS, BARE_OPT_PATTERNS, para_has_image(), is_option() and
    clean_option_text() now live here; all three specs delegate. para_has_image()
    accepts a python-docx Paragraph OR a raw <w:p> element, because Step 5 passes the
    first and Steps 1/3 pass the second — a delegation that assumed one form would
    silently reinstate the undercount on the other.

v1.5 — 2026-07-26 — score_vision_probe() REFUSES TO SCORE A NON-OBSERVATION
    (GAP-2026-07-26-002 PART A).

    It returned False for an empty string, collapsing two categorically different
    states into one session-terminating value:

        ''  / None / whitespace  ->  "I did not look"           -> procedural error
        'ABC12345' (wrong)       ->  "I looked, could not read" -> session verdict

    Returning False for the first made a careless non-observation indistinguishable
    from a genuinely blind session. IMG-6 then halted a production Step-5 run at
    batch 1 having processed nothing, on a probe image that had rendered correctly
    and was legible — the token was recovered from it on review.

    Now raises ProbeObservationMissing (new, a CorpusError subclass) for the empty
    case. A scorer can only compare an observation that exists; it is not an oracle
    for whether one was made. Callers retry rather than treating it as a failure.

    BREAKING for any caller that relied on False-on-empty. Verified: there were no
    executable callers — every reference was in spec prose — so nothing else moved.
    Framework_MockTestAnalyse >= v2.34 and Framework_PYQPrepare >= v1.10 depend on
    this behaviour; an older engine silently re-opens the false-halt path.

v1.4 — 2026-07-26 — THE TAXONOMY COMES FROM JSON, NOT FROM A WORD DOCUMENT.
    GAP-2026-07-25-003, resolved at the cause rather than at the symptom.
    v1.2 taught the reader both ingest forms and works. It works by matching an
    undocumented, unversioned extraction grammar the framework does not control,
    observed from one sample — so the next change to that grammar stops all ~200
    exams at once. Loudly, but at once. That is a mitigation, not a resolution.
    reconcile_taxonomy v1.3 now persists the approved taxonomy INSIDE
    approval_record.json, which it has always held in the right shape and thrown
    away. New load_taxonomy() prefers it and falls back to the Analysis doc for
    records that predate schema 1.3, so no already-approved exam needs re-running.
    The record VALIDATES ITSELF: the fingerprint recorded beside the taxonomy was
    computed from it at Step 2c, so nothing new is trusted, and a record whose two
    halves disagree hard stops rather than being half-believed.
    _assemble() is extracted so the record path and the document path build the ONE
    return shape from one implementation — two constructions of one shape is how
    GAP-2026-07-25-002's four readers happened.
    Also stopped: silent de-duplication. A name repeated in the recorded taxonomy
    would collapse on assembly and the fingerprint would NOT notice, because it is
    computed over the same duplicated triples and therefore agrees with itself.

v1.3 — 2026-07-26 — THE TAXONOMY LOCK GATE, ONCE (GAP-2026-07-25-003 follow-up).
    read_analysis_doc() asserts the doc agrees with ITSELF; nothing in Cluster K
    asserted it is the doc PYQApprove LOCKED. Framework_PYQSort S1-0b made that
    second claim and NOTHING else did — Steps 5 and 6 read the taxonomy and
    allocated against it with no identity check at all. At Step 5 it was worse than
    absent: the reader's hard stops were caught and downgraded to a log line TWICE,
    once in _extract_taxonomy_tuples_from_analysis_doc and again in its caller, so a
    superseded or mis-parsed doc that halts loudly at Step 3 was accepted three steps
    later without a word. Every subtopic_id minted from it would be wrong, and ids
    are what Steps 6-11 match on.
    Adds discover_approval_record(), read_approval_record() and
    assert_taxonomy_lock(doc, ...). The gate lives here ONCE rather than being
    written a fifth time in a fifth spec — the same anti-drift rule that produced
    Cluster K, now covering the lock as well as the reader.
    SCOPE IS NARROW ON PURPOSE: this asserts IDENTITY only. Whether the lock was
    EARNED — status, schema attestation, checks.missing/vacuous — is PYQSort S1-0's
    claim, made once at Step 3 against the same record, and is not duplicated. One
    function, one claim.

v1.2.1 — 2026-07-26 — UNKNOWN-FORM DIAGNOSTIC COULD RAISE THE FAULT IT DIAGNOSED.
    read_analysis_doc()'s INGEST_UNKNOWN branch formatted its message with
    os.path.getsize(path). _ingest_form() returns UNKNOWN for an unreadable path
    BECAUSE getsize() raised there, so the diagnostic re-raised it uncaught and the
    operator received exactly the library traceback v1.2 was written to replace —
    on the one input where the named stop matters most. Reproducible with any
    unreadable or vanished path; found by re-running the delivered spec end to end
    rather than by review. The size is diagnostic and is now reported as
    "unreadable (<Error>: <detail>)" when it cannot be taken.

v1.2 — 2026-07-26 — CLUSTER K: INGEST FORMS (GAP-2026-07-25-003). The Analysis doc
    is uploaded to the project's Files section after Step 2c and read from there at
    Steps 3, 4, 5 and 6. The platform stores an uploaded .docx in project knowledge
    as EXTRACTED MARKDOWN TEXT while PRESERVING the .docx filename, so the text form
    is the ONLY form the runtime ever receives for this artefact — yet read_analysis_doc()
    opened it with python-docx and hard stopped with PackageNotFoundError on every
    project, severing the pipeline at Step 3. Measured: chat attachment = 40,882-byte
    OOXML package; the same file in project Files = 12,911 bytes of Markdown text.
    v1.1 anticipated this, but put the diagnostic on discover_analysis_doc()'s "no file
    found" branch, which the preserved filename makes unreachable: discovery SUCCEEDS
    and hands _scan() a text file.

    This is a REGRESSION BOUNDARY, not a latent bug. Before v1.1 the artefact was read
    by prose in Framework_Blueprint S2-2 ("instruct Claude to extract the structure
    itself"), which works on the text form by construction. v1.1 replaced prose with
    python-docx and thereby traded a working non-deterministic read for a deterministic
    broken one. The answer is NOT to restore prose — of the four pre-v1.1 readers only
    Blueprint's worked at all; MockTestAnalyse's two returned 0 subtopics against a
    truth of 131 on every exam, silently. ONE reader stands; it learns the second form.

    Four changes, in the order they matter:
      (a) INGEST FORMS. The container is classified from CONTENT, never from the
          extension — the .pdf in the same project is a Zip page-bundle under its .pdf
          name, so the extension is not evidence of the container. _scan() is renamed
          _scan_ooxml() (body untouched, still used by write_analysis_doc()'s output and
          the round-trip self-test) and _scan_text() joins it, emitting the IDENTICAL
          (subjects, order, declared, anomalies) 4-tuple. verify_analysis_doc(),
          taxonomy_fingerprint(), the return shape and every downstream consumer are
          therefore untouched, and the text form is admitted THROUGH both existing gates
          rather than around them. An unrecognised form is a named hard stop, never a
          best-effort parse, and verify=False is refused on a non-OOXML form.
      (b) RAGGED-ROW HARD STOP (the defect the ingest-form fix would otherwise have
          introduced). A markdown table is rectangular. A '|' inside a taxonomy name
          splits one row wider than its neighbours, and the surplus cell is swallowed
          as the count column: 'Transverse | Longitudinal' parsed as 'Transverse', with
          'Longitudinal' discarded and the SUBTOPIC COUNT UNCHANGED — so D1, D2 and D3
          all agree and verify_analysis_doc() returns clean. Reproduced by execution.
          That matters beyond Step 3: the fingerprint cross-check against the approval
          record exists ONLY in Framework_PYQSort S1-0b, so at Steps 5 and 6 the mis-parse
          has no gate at all and is silent end to end. Non-uniform cell width across a
          table's rows is now a hard stop naming the row, and write_analysis_doc() refuses
          to emit a name containing '|' — the character is unrepresentable in this ingest
          form, so it is rejected where it originates rather than where it corrupts.
      (c) GRAMMAR TOLERANCE, bounded. Bold is stripped whether the extractor emits ** or
          a full-line __wrap__, backslash escapes on markdown punctuation are removed, and
          cells split only on UNESCAPED pipes. Measured against a second, independent
          docx->markdown extractor on the same live document, which emits __bold__,
          escapes ('IIT\\_JAM\\_BIOTECHNOLOGY', 'Subtopic\\-wise') and flattens tables
          entirely. The first two are now handled; the third is not, and is not guessed
          at — it produces the loud "extraction grammar may have changed" stop. This is
          tolerance of variants that have been OBSERVED, not speculation about grammars
          that have not.
      (d) The stale hint at the tail of discover_analysis_doc() is corrected. After this
          change the text form is READ, not diagnosed.

    KNOWN RESIDUAL, recorded rather than hidden: this module now depends on an
    undocumented, unversioned, platform-controlled extraction grammar. It fails LOUDLY
    when that grammar changes, but it fails for every exam at once. The durable
    resolution is to stop recovering machine state from a rendered document at all:
    reconcile_taxonomy.build_approval_record() already receives final_taxonomy in the
    exact {section: {topic: [subtopic]}} shape, derives the fingerprint from it and then
    DISCARDS it, and syllabus_provenance.read_taxonomy_draft() already reads that shape
    from JSON, which the platform stores byte-for-byte. Persisting the approved taxonomy
    into the record makes this whole class of defect unreachable. Tracked separately.

v1.1 — 2026-07-25 — CLUSTER K: THE Analysis-doc reader, writer and verifier
    (GAP-2026-07-25-002). The Step-2c Analysis doc previously had FOUR readers
    across four specs, implementing four different structural conventions, and no
    writer at all (PYQAnalyse S4-2 was `pass`). Three of the four were measurably
    wrong on the first real exam's doc. One reader, one writer, one verifier; all
    heading recognition delegated to blueprint_core.parse_taxonomy_level so every
    label form works in every step; correctness asserted by round-trip over a
    GENERATED matrix of exam shapes rather than by review. See Cluster K header.


v1.0.3 — 2026-07-25 — optimize_docx gains `always`, which runs the tier ladder even when
    the input is already under budget. Requested so PYQCompress can be driven purely by
    what the operator attaches rather than by a size threshold. force_tier was NOT reused
    for this: it pins one tier and makes the closing `return True` unconditional, so an
    oversized paper would have stopped at that tier, still over budget, and reported
    success with the floor-exceeded WARN suppressed. Adds _no_gain_guard, a
    document-level counterpart to the existing per-part no-growth rule: the ZIP container
    can grow even when no part does, and CHECK 5 treats growth as a HARD STOP, so a small
    already-optimal document would have aborted the run instead of being left alone.

v1.0.2 — 2026-07-25 — MEDIA STEM COLLISION (silent figure loss). optimize_docx renamed
    every re-encoded part to .jpeg/.png and stored it in a dict keyed by that name, so
    two parts sharing a stem and taking the same encoder route (word/media/image1.png +
    word/media/image1.jpg, both photographic) both resolved to word/media/image1.jpeg
    and one OVERWROTE the other. The writer then emitted a duplicate zip member and Word
    rendered whichever it met first, silently replacing one figure with the other.
    Invisible to every existing gate: parts, media and refs all still matched, no
    relationship dangled, and the per-stem pixel check that would have caught it is
    skipped by design at T2-T4 (allow_resample=True) — so the corrupted file passed
    parity and was DELIVERED. Reproduced end to end; T1 caught it, T2-T4 did not.
    Four changes: (a) _target_media_name keeps a part's name when its extension already
    maps to the encoder's content type, so .jpg is no longer renamed to .jpeg and the
    common collision never forms; (b) a genuine remaining collision is disambiguated
    (__c2) rather than overwritten; (c) an explicit input->output name map replaces the
    reverse-lookup in the writer, which could not express a disambiguated name; (d) the
    relationship rewrite is a single regex pass — chained str.replace() with
    {image1.png -> image1.jpeg, image1.jpeg -> image1__c2.jpeg} rewrote the part it had
    just manufactured and pointed both references at one file. A part-count guard
    (len(parts) == len(names)) backstops all of it. assert_docx_parity now compares
    pixel dimensions as an unordered multiset: keying on the stem carried the v1.0.1
    false positive one level up, because a disambiguated name changes the stem.

v1.0.1 — 2026-07-25 — assert_docx_parity raised a FALSE-POSITIVE HARD STOP whenever the
    governor renamed a media part. Re-encoding a photographic PNG produces image1.jpeg
    from image1.png, and the pixel-dimension comparison was keyed on the full basename,
    so the renamed part had no counterpart in `after` and was reported as a dimension
    change on a document whose dimensions were identical (verified: 2400x1800 before and
    after). Dormant on the measured corpus, whose papers store .jpeg parts that the jpeg
    route leaves named as they were; it fires on any PNG-sourced photograph, i.e. exactly
    the papers the governor exists for. Comparison is now by extension-less part name.
    Found by the Framework_PYQSort v1.12 governor test, not by this module's self-test —
    a regression case has been added (io_parity_rename_*).

WHY THIS MODULE EXISTS (and why none of it lives in blueprint_core.py)
    blueprint_core.py declares a THIN-CORE INVARIANT: pure, no I/O, standard library
    only. paper_pipeline.py makes the same promise. Every concern below is inherently
    impure — Drive retrieval, ZIP surgery, image transcoding, disk writes — so hosting
    it in either module would break a stated design invariant. Worse, importing PIL or
    python-docx into blueprint_core.py would make the ALLOCATION core unimportable
    wherever those packages are absent, which is exactly the P0 recorded in
    Framework_MockTestAnalyse v2.26: a failed `import blueprint_core` aborted Step 5
    for EVERY exam.

    corpus_io.py is therefore the one home for impure corpus plumbing, exactly as
    explain_engine.py is the one home for impure explanation plumbing.

        blueprint_core Cluster H   DECIDES   (pure data in -> pure data out)
        corpus_io                  PERFORMS  (bytes, files, images)

    Every decision this module needs is delegated to Cluster H. This module contains
    no thresholds, no ladders and no classification rules of its own.

ANTI-DRIFT CONTRACT
    Steps 1, 2b, 3, 4, 5 and PYQCompress all consume this module. No spec may define a
    local copy of any function named here — enforced by validate_framework_md.py
    Check Z. A spec may only call it or define a thin forwarding adapter. Re-localising
    a shared function produces ZERO drift signal until the two copies disagree; that is
    the failure that required Framework_MockTestAnalyse v2.27 and
    Framework_PYQAnalyse v2.20.

DEPENDENCY POLICY
    Standard library at module scope. Pillow and python-docx are imported LAZILY inside
    the functions that need them, so Cluster H shell functions (path / JSON / ZIP logic)
    keep working without them, and a missing optional dependency surfaces as a named
    error at the point of use rather than an import failure at load time.

DEFECTS THIS MODULE EXISTS TO FIX (see Framework_PYQCompress §3 for the full register)
    B  Drive download was unguarded — zero try/except existed anywhere in the corpus
    G  legacy .doc accepted at enumeration though python-docx cannot open it
    H  native Google Docs silently skipped, removing a paper with no error
    I  images inside TABLES invisible: doc.paragraphs does not descend into tables
    J  legacy VML <v:imagedata> images invisible in every image-handling spec
    K  header / footer / footnote images never extracted
    L  pre-Q.1 images silently discarded rather than bucketed
    N  the real download envelope (context spill + double JSON parse) undocumented
    O  duplicate paper identity double-counted a year
"""

import base64
import io
import json
import os
import re
import shutil
import zipfile

import blueprint_core as bc


# ── exceptions — named and actionable; the CALLER decides what is fatal ──────────

class CorpusError(Exception):
    """Base for every failure raised by this module."""


class TransportFallback(CorpusError):
    """Bytes could not be fetched from Drive. NOT fatal — route to the upload lane.

    Raised for the size cap, a truncated payload, a permission failure, a malformed
    envelope, and any unrecognised connector error. The caller MUST catch this and
    request the paper by chat upload instead of stopping the run.

    This is the guarantee that a future change to the connector's cap cannot break the
    pipeline: correctness does not depend on the predicted partition being right, only
    on the fallback being taken.
    """


class EnumerationError(CorpusError):
    """A Drive folder listing cannot be turned into a usable corpus."""


class DuplicatePaperError(EnumerationError):
    """Two files resolve to the same paper identity — HARD STOP.

    Left unchecked this double-counts the paper's year in the §1-6 coverage gate, or
    drops one of the two depending on listing order.
    """


class IntegrityError(CorpusError):
    """A document failed a structural or image-integrity gate — HARD STOP."""


class DependencyMissing(CorpusError):
    """An optional package (Pillow / python-docx) is needed here but not installed."""


# ── constants that are pure I/O concerns (all POLICY lives in Cluster H) ─────────

A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
V_NS = 'urn:schemas-microsoft-com:vml'

MEDIA_RE = re.compile(r'^word/(media|embeddings)/', re.I)
RASTER_EXT = ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tif', '.tiff', '.webp')
VECTOR_EXT = ('.emf', '.wmf', '.svg')

# Every part that can carry a drawing. Footnotes and endnotes are included because a
# figure there is still a figure; omitting them makes the reference count silently low
# and would fire IMG-4 for a benign reason.
STORY_PARTS_RE = re.compile(
    r'^word/(document|header\d*|footer\d*|footnotes|endnotes)\.xml$')

DEFAULT_Q_PATTERN = r'^\s*Q\.?\s*(\d+)'
PREAMBLE = 'preamble'
DEFAULT_UPLOAD_DIR = '/mnt/user-data/uploads'


def _need(mod):
    """Import an optional dependency, or fail with a named, actionable error."""
    try:
        if mod == 'PIL':
            from PIL import Image  # noqa: F401
            return __import__('PIL.Image', fromlist=['Image'])
        if mod == 'docx':
            import docx  # noqa: F401
            return docx
    except ImportError as exc:
        raise DependencyMissing(
            f'{mod} is required for this operation but is not installed ({exc}). '
            f'Install with: pip install {"Pillow" if mod == "PIL" else "python-docx"} '
            f'--break-system-packages')


# ═══════════════════════════════════════════════════════════════════════════════
# CLUSTER H SHELL — ACQUISITION
# ═══════════════════════════════════════════════════════════════════════════════

def parse_drive_folder_id(url):
    """Extract the folder id from a Drive folder URL, or return the input if it already
    looks like a bare id. None when neither."""
    s = str(url or '').strip()
    m = re.search(r'drive\.google\.com/drive/(?:u/\d+/)?folders/([A-Za-z0-9_-]+)', s)
    if m:
        return m.group(1)
    if re.fullmatch(r'[A-Za-z0-9_-]{12,}', s):
        return s
    return None


def normalise_drive_listing(response):
    """Turn any shape of Drive listing response into uniform file records.

    Accepts the bare list, {'files': [...]} or {'items': [...]}, and tolerates both the
    'title' and 'name' spellings the connector has used.

    fileSize arrives as a STRING and is absent for folders and native Google types.
    Both are normalised here so no caller has to think about it. Capturing size at
    enumeration costs nothing — the listing already carries it — and is what allows the
    AUTO/UPLOAD partition to be computed BEFORE any download is attempted, instead of
    discovering the blocker three-quarters of the way through a run.
    """
    items = response
    if isinstance(response, dict):
        items = response.get('files', response.get('items', []))
    out = []
    for it in items or []:
        if not isinstance(it, dict):
            continue
        raw = it.get('fileSize', it.get('size'))
        try:
            size = int(raw) if raw is not None else None
        except (TypeError, ValueError):
            size = None
        out.append({
            'id': it.get('id'),
            'name': it.get('title') or it.get('name') or '',
            'mimeType': it.get('mimeType', ''),
            'fileSize': size,
            'parentId': it.get('parentId'),
            'source': 'gdrive',
        })
    return out


def collect_corpus_files(list_fn, folder_id, recurse=True, require_size=True,
                         _seen=None, _papers=None, _rejects=None, _depth=0):
    """Walk a Drive folder and return (papers, rejects).

    `list_fn(folder_id, page_token=None)` is injected so this is testable without Drive
    and so the caller owns the connector call. It must return whatever the connector
    returns; normalise_drive_listing handles the shape.

    Pagination is followed to exhaustion — without it a folder of more than one page
    silently loses papers.

    Sub-folders are visited in DESCENDING name order so year-named folders are walked
    newest-first, preserving the recency-first intent of the existing enumeration.

    Raises DuplicatePaperError when two usable files share a canonical identity.
    """
    if _seen is None:
        _seen, _papers, _rejects = {}, [], []
    if _depth > 24:
        raise EnumerationError('folder nesting exceeded 24 levels — possible cycle')

    subfolders, page_token = [], None
    while True:
        resp = list_fn(folder_id, page_token=page_token) if page_token else list_fn(folder_id)
        entries = normalise_drive_listing(resp)
        page_token = resp.get('nextPageToken') if isinstance(resp, dict) else None

        for e in entries:
            verdict, reason = bc.screen_drive_entry(
                e['name'], e['mimeType'], e['fileSize'], require_size=require_size)
            if verdict == 'folder':
                subfolders.append((e['name'], e['id']))
                continue
            if verdict == 'reject':
                _rejects.append(dict(e, reason=reason))
                continue
            key = bc.canonical_paper_key(e['name'])
            if key in _seen:
                prev = _seen[key]
                raise DuplicatePaperError(
                    'Two files resolve to the same paper identity — HARD STOP.\n'
                    f'  identity : {key}\n'
                    f'  file A   : {prev["name"]}  (id {prev.get("id")}, '
                    f'{prev.get("fileSize")} bytes)\n'
                    f'  file B   : {e["name"]}  (id {e.get("id")}, '
                    f'{e.get("fileSize")} bytes)\n'
                    'Remove or rename one in Drive. Leaving both makes this paper\'s '
                    'year count twice, or drops one at random depending on listing order.')
            rec = dict(e, paper_key=key)
            _seen[key] = rec
            _papers.append(rec)

        if not page_token:
            break

    if recurse:
        subfolders.sort(key=lambda x: x[0], reverse=True)
        for _, sub_id in subfolders:
            collect_corpus_files(list_fn, sub_id, recurse, require_size,
                                 _seen, _papers, _rejects, _depth + 1)
    return _papers, _rejects


_B64_CHARS_ONLY = re.compile(r'^[A-Za-z0-9+/=\s]+$')


def _decode_bare_base64(text):
    """GAP-2026-08-15-PYQCOUNT-DRIVE-ACQUISITION (P3b) — accept a BARE base64 string.

    decode_drive_payload used to accept bytes, a spill path, a dict, a list and a JSON
    string, but NOT a plain base64 string with no envelope around it. That shape is not
    hypothetical: any staging protocol that writes a `.b64` sidecar, or a model that
    hands over just the `content` field it read out of a spill file, produces exactly
    it, and every one of them was rejected with a message that named three shapes and
    not this one.

    Ambiguity against the other accepted shapes is resolved by DECODING, not guessing:
    a candidate is accepted only if it is long, contains base64 characters exclusively,
    and decodes to bytes beginning PK\x03\x04. A spill path fails the length/charset
    test (it contains '/' but also '.', and it exists on disk so it never reaches here);
    a JSON envelope fails the charset test on '{'. Returns None — never raises — so the
    caller keeps ownership of the error message.
    """
    if not isinstance(text, str):
        return None
    stripped = text.strip()
    if len(stripped) <= 64 or not _B64_CHARS_ONLY.match(stripped):
        return None
    try:
        raw = base64.b64decode(stripped)
    except Exception:
        return None
    return raw if raw[:4] == b'PK\x03\x04' else None


def decode_drive_payload(payload):
    """Turn whatever download_file_content returned into raw file bytes.

    The real contract is not what a one-line pseudocode wrapper suggests, and every
    execution of Steps 2b/4/5 has rediscovered it by trial and error:

      * the tool result MAY be spilled to a file on disk instead of returned in the
        turn, so `payload` may be a PATH. Whether it spills, and to which directory,
        is a property of the DEPLOYMENT, not of the file size — measured 2026-08-15
        on one 40,488-byte .docx: spilled in one deployment, returned inline in
        another. Never assume either; this function accepts both;
      * that spilled JSON is a LIST whose [0]['text'] is itself a JSON STRING;
      * parsing that string yields {'id','title','mimeType','content'} where 'content'
        holds the base64 payload.

    Accepts, in order: raw bytes, a spill-file path, the parsed list/dict, or the inner
    JSON string. Anything else raises TransportFallback so the caller routes to upload
    rather than dying.
    """
    if isinstance(payload, (bytes, bytearray)):
        return bytes(payload)

    obj = payload
    if isinstance(obj, str):
        if os.path.exists(obj):
            try:
                with open(obj, 'r', encoding='utf-8') as fh:
                    obj = json.load(fh)
            except (OSError, ValueError) as exc:
                raise TransportFallback(f'could not read spill file {obj}: {exc}')
        else:
            try:
                obj = json.loads(obj)
            except ValueError:
                bare = _decode_bare_base64(obj)
                if bare is not None:
                    return bare
                raise TransportFallback(
                    'download payload is neither bytes, a spill-file path, bare '
                    'base64, nor JSON')

    for _ in range(6):                                   # bounded unwrap; never loops
        if isinstance(obj, list):
            if not obj:
                raise TransportFallback('download payload was an empty list')
            obj = obj[0]
            continue
        if isinstance(obj, dict):
            content = obj.get('content')
            if isinstance(content, str):
                try:
                    return base64.b64decode(content)
                except Exception as exc:
                    raise TransportFallback(f'base64 content would not decode: {exc}')
            text = obj.get('text')
            if isinstance(text, str):
                try:
                    obj = json.loads(text)
                except ValueError:
                    raise TransportFallback(
                        'payload text was not the expected JSON envelope')
                continue
        break
    raise TransportFallback('could not locate base64 content in the download payload')


def verify_downloaded_bytes(raw, expected_size=None, name=''):
    """Prove the payload is a complete .docx before anything downstream trusts it.

    The byte-count check is not redundant with the magic check. A payload truncated at a
    ZIP member boundary can still open as a valid archive while presenting FEWER media
    parts — a silent image loss with no error anywhere. Comparing against the size Drive
    reported is the only thing that catches that case.
    """
    label = f' for {name}' if name else ''
    if not raw:
        raise TransportFallback(f'download returned no bytes{label}')
    if raw[:4] != b'PK\x03\x04':
        raise TransportFallback(
            f'payload{label} is not a .docx (magic {raw[:4]!r}) — it may be an error '
            'page, an HTML login redirect, or a native Google type')
    if expected_size is not None and len(raw) != expected_size:
        raise TransportFallback(
            f'size mismatch{label}: received {len(raw):,} bytes, Drive reported '
            f'{expected_size:,}. Treating as truncated.')
    return True


def stage_drive_payload(payload, paper, dest_dir):
    """Decode + verify + persist an ALREADY-MATERIALISED Drive payload. No callable.

    GAP-2026-08-15-PYQCOUNT-DRIVE-ACQUISITION (P3a). Before this function the ONLY
    verified route from a Drive payload to a file on disk was fetch_drive_docx, which
    demands a `download_fn` callable. A caller in PHASE B already HOLDS the payload —
    the model fetched it in its own turn — so satisfying that signature meant inventing
    a callable, and the cheapest-looking invention is to pass the CLASS T marker name
    itself. That is precisely the defect this GAP documents. Removing the incentive is
    a stronger fix than documenting the trap.

    Guarantees are identical to fetch_drive_docx: byte count equals the fileSize the
    listing reported, PK\x03\x04 magic, and TransportFallback (never SystemExit) on
    any failure so the caller degrades to the upload lane.

    `payload` is whatever PHASE A produced for this paper — raw bytes, a spill-file
    path in ANY directory, the parsed dict/list, an inner JSON string, or bare base64.
    Returns the local path written.
    """
    name = paper.get('name', '<unnamed>')
    raw = decode_drive_payload(payload)
    verify_downloaded_bytes(raw, paper.get('fileSize'), name)
    os.makedirs(dest_dir, exist_ok=True)
    dest = os.path.join(dest_dir, os.path.basename(name))
    with open(dest, 'wb') as fh:
        fh.write(raw)
    return dest


def fetch_drive_docx(download_fn, paper, dest_dir):
    """Fetch one paper to disk. ANY failure raises TransportFallback — never SystemExit.

    `download_fn(file_id)` is injected so the caller owns the connector call. It must be
    a RESOLVER over results PHASE A already materialised — a plain lookup that performs
    no tool call. Passing the name of a CLASS T marker instead is unreachable code: see
    stage_drive_payload's docstring, GAP-2026-08-15-PYQCOUNT-DRIVE-ACQUISITION, and
    audit_callgraph C6/C7 which fail the build for it.

    Papers already known to exceed the connector cap are not attempted at all: the
    request is guaranteed to fail and would only add latency. Everything else is
    attempted, because the partition is predictive and a marginal file fetches fine.

    This is a thin wrapper over stage_drive_payload: it adds the cap pre-check and the
    guarded call, then delegates decode/verify/persist so both routes cannot diverge.
    """
    size = paper.get('fileSize')
    name = paper.get('name', '<unnamed>')
    if size is not None and size > bc.DRIVE_CAP:
        raise TransportFallback(
            f'{name} is {size:,} bytes, above the {bc.DRIVE_CAP:,}-byte connector cap — '
            'routing to the upload lane without attempting a download')
    try:
        payload = download_fn(paper['id'])
    except TransportFallback:
        raise
    except Exception as exc:                              # connector / auth / network
        raise TransportFallback(f'download of {name} failed: {exc}')

    return stage_drive_payload(payload, paper, dest_dir)


def resolve_uploaded_papers(expected_keys, upload_dir=DEFAULT_UPLOAD_DIR):
    """Match chat-uploaded .docx files against the papers this batch asked for.

    Matching is by CANONICAL identity, never by exact filename and NEVER by recency.

      * exact filename fails because the browser appends " (1)" when the original is
        already in the operator's Downloads folder — which happens on every remediation
        round trip;
      * recency ("the newest three files") fails because uploads ACCUMULATE across
        turns, so by batch 3 the directory still holds batches 1 and 2 and the step
        would silently reprocess them.

    Returns {'matched': {key: path}, 'missing': [...], 'unexpected': [...]}.
    An unrecognised upload is reported, never guessed at and never processed.
    """
    want = set(expected_keys)
    matched, unexpected = {}, []
    if os.path.isdir(upload_dir):
        for fn in sorted(os.listdir(upload_dir)):
            if fn.startswith('~$') or not fn.lower().endswith('.docx'):
                continue
            path = os.path.join(upload_dir, fn)
            key = bc.canonical_paper_key(fn)
            if key in want:
                matched.setdefault(key, path)
            else:
                unexpected.append(path)
    return {'matched': matched,
            'missing': sorted(want - set(matched)),
            'unexpected': unexpected}


# ═══════════════════════════════════════════════════════════════════════════════
# CLUSTER I SHELL — IMAGE INTEGRITY
# ═══════════════════════════════════════════════════════════════════════════════

def _safe_document(path):
    """Open a .docx via python-docx, converting any package-level failure into a NAMED
    error.

    python-docx surfaces a missing or corrupt part as a bare KeyError/BadZipFile from
    deep inside its package reader. Allowing that to escape means a structurally broken
    document produces an unhandled traceback instead of a gate verdict, which is exactly
    the class of failure this module exists to eliminate.
    """
    docx = _need('docx')
    try:
        return docx.Document(path)
    except DependencyMissing:
        raise
    except Exception as exc:
        raise IntegrityError(
            f'{os.path.basename(path)} could not be opened as a Word package '
            f'({type(exc).__name__}: {exc}). The file is corrupt, truncated, or is '
            f'missing a part that its relationships still reference.')


def _open_zip(path_or_zip):
    """Open a .docx as a ZIP, converting archive-level corruption into a NAMED error.

    A .docx IS a ZIP, so a truncated or malformed download surfaces here first as a bare
    BadZipFile. Letting that escape produces an unhandled traceback instead of a gate
    verdict — the exact class of failure this module exists to eliminate.
    """
    if isinstance(path_or_zip, zipfile.ZipFile):
        return path_or_zip
    try:
        return zipfile.ZipFile(path_or_zip)
    except zipfile.BadZipFile as exc:
        raise IntegrityError(
            f'{os.path.basename(str(path_or_zip))} is not a readable Word package '
            f'({exc}). The file is corrupt or truncated.')


def _rels_for(z, part):
    """rId -> media basename, for the .rels belonging to `part`."""
    rel = f'{os.path.dirname(part)}/_rels/{os.path.basename(part)}.rels'
    out = {}
    if rel in z.namelist():
        txt = z.read(rel).decode('utf-8', 'ignore')
        for rid, tgt in re.findall(r'Id="([^"]+)"[^>]*Target="([^"]+)"', txt):
            out[rid] = os.path.basename(tgt)
    return out


def count_image_refs(path_or_zip, body_only=False):
    """Count EVERY image reference in the document. Returns (total, per_part, unresolved).

    Matches BOTH image mechanisms:
      * <a:blip r:embed>      DrawingML — covers inline AND floating (<wp:anchor>)
      * <v:imagedata r:id>    legacy VML — emitted by older Word, some PDF converters,
                              and pasted OLE / equation objects

    and scans every story part (document, headers, footers, footnotes, endnotes), not
    just document.xml.

    Deliberately does NOT use python-docx `inline_shapes`. That property sees only
    inline body drawings, so anchored figures, VML objects and header images are
    invisible to it. A count that can silently run low is worse than no count at all,
    because it makes a broken document look verified.
    """
    z = _open_zip(path_or_zip)
    total, per_part, unresolved = 0, {}, []
    for part in z.namelist():
        if not STORY_PARTS_RE.match(part):
            continue
        if body_only and part != 'word/document.xml':
            continue
        xml = z.read(part).decode('utf-8', 'ignore')
        rels = _rels_for(z, part)
        rids = re.findall(r'<a:blip[^>]*r:embed="([^"]+)"', xml)
        rids += re.findall(r'<v:imagedata[^>]*r:id="([^"]+)"', xml)
        for rid in rids:
            tgt = rels.get(rid)
            if not tgt:
                unresolved.append((part, rid))
                continue
            per_part[tgt] = per_part.get(tgt, 0) + 1
            total += 1
    return total, per_part, unresolved


def dangling_media_targets(path_or_zip):
    """Relationship targets that point at a media part which does not exist."""
    z = _open_zip(path_or_zip)
    names = set(z.namelist())
    out = []
    for n in z.namelist():
        if not n.endswith('.rels'):
            continue
        owner = os.path.dirname(os.path.dirname(n))
        for t in re.findall(r'Target="((?:\.\./)?media/[^"]+)"',
                            z.read(n).decode('utf-8', 'ignore')):
            cand = os.path.normpath(os.path.join(owner, t)).replace('\\', '/')
            if cand not in names:
                out.append(cand)
    return out


def extract_images(path, outdir):
    """Write every media part to disk as ORIGINAL BYTES and describe each one.

    Nothing is re-encoded, re-rendered or resized here — a .docx is a ZIP, so the stored
    bytes ARE the original image. Re-encoding happens only inside the governor, and only
    under assert_docx_parity().

    Vector parts (EMF/WMF/SVG) cannot be opened as rasters. They are labelled 'vector'
    and REPORTED rather than skipped, so a document that will not shrink has a stated
    reason instead of an unexplained result.
    """
    Image = _need('PIL')
    os.makedirs(outdir, exist_ok=True)
    z = _open_zip(path)
    out = {}
    for n in z.namelist():
        if not MEDIA_RE.match(n):
            continue
        raw = z.read(n)
        base = os.path.basename(n)
        fp = os.path.join(outdir, base)
        with open(fp, 'wb') as fh:
            fh.write(raw)
        rec = {'part': n, 'path': fp, 'bytes': len(raw), 'kind': None,
               'format': None, 'size': None, 'mode': None, 'note': ''}
        if os.path.splitext(base)[1].lower() in VECTOR_EXT:
            rec['kind'] = 'vector'
            rec['note'] = 'vector part — rasterise before view(); governor leaves it as-is'
        else:
            try:
                im = Image.open(io.BytesIO(raw))
                im.load()
                rec.update(kind='raster', format=im.format, size=im.size, mode=im.mode)
            except Exception as exc:
                rec['kind'] = 'unreadable'
                rec['note'] = str(exc)[:80]
        out[base] = rec
    return out


def map_images_to_questions(path, q_pattern=DEFAULT_Q_PATTERN):
    """Attach every body image to the question it appears under, in DOCUMENT ORDER.

    Iterates `doc.element.body.iter()`, NOT `doc.paragraphs`.

    This is not a stylistic preference. In python-docx `doc.paragraphs` returns only
    paragraphs that are direct children of the body; paragraphs inside table cells are
    excluded entirely. Any figure laid out in a table — the normal arrangement for
    match-the-following items, multi-panel figures and option grids — is therefore never
    visited and never mapped. The question is then classified TEXT rather than FIGURAL,
    which corrupts the format distribution that drives Step 7 generation, and no error
    is raised anywhere. body.iter() descends into tables and fixes this.

    Images appearing before the first question go to a 'preamble' bucket rather than
    being silently dropped, so IMG-4 does not fire for a benign reason.

    q_pattern is supplied by the caller — this module knows nothing about how any
    particular exam numbers its questions.
    """
    d = _safe_document(path)
    z = _open_zip(path)
    rels = _rels_for(z, 'word/document.xml')
    mapping, cur = {}, PREAMBLE
    for el in d.element.body.iter():
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter() if t.tag.split('}')[-1] == 't')
        m = re.match(q_pattern, txt)
        if m:
            cur = int(m.group(1))
        for b in el.iter(f'{{{A_NS}}}blip'):
            rid = b.get(f'{{{R_NS}}}embed')
            if rid:
                mapping.setdefault(cur, []).append(rels.get(rid, f'UNRESOLVED:{rid}'))
        for v in el.iter(f'{{{V_NS}}}imagedata'):
            rid = v.get(f'{{{R_NS}}}id')
            if rid:
                mapping.setdefault(cur, []).append(rels.get(rid, f'UNRESOLVED:{rid}'))
    return mapping


def verify_images(path, extracted=None, mapping=None, expected_size=None,
                  q_pattern=DEFAULT_Q_PATTERN, workdir=None):
    """Run gates IMG-1 .. IMG-5. Returns (verdicts, stats). Never raises.

    The caller decides what to do with a FAIL; this function only reports. Verdict
    evaluation is delegated to blueprint_core.image_gate_verdict so the pass/fail rules
    are unit-testable without any file present.
    """
    if extracted is None:
        extracted = extract_images(path, workdir or f'/tmp/pyq_img/{os.getpid()}')
        # extract_images raises IntegrityError on a corrupt archive; that is a HARD STOP
        # and is allowed to propagate — there is nothing left to verify.

    # The mapping needs python-docx, which fails hard on a structurally broken package.
    # A broken document must still produce GATE VERDICTS rather than a traceback, so the
    # failure is captured and reported through IMG-4 like any other integrity fault.
    map_error = None
    if mapping is None:
        try:
            mapping = map_images_to_questions(path, q_pattern)
        except (IntegrityError, DependencyMissing) as exc:
            mapping, map_error = {}, str(exc)

    z = _open_zip(path)
    all_refs, per_part, unresolved = count_image_refs(z)
    body_refs, _, _ = count_image_refs(z, body_only=True)
    preamble = len(mapping.get(PREAMBLE, []))
    mapped = sum(len(v) for k, v in mapping.items() if k != PREAMBLE)

    verdicts = bc.image_gate_verdict(
        actual_size=os.path.getsize(path),
        expected_size=expected_size,
        unresolved=len(unresolved),
        missing_on_disk=sorted(set(per_part) - set(extracted)),
        mapped=mapped,
        preamble=preamble,
        body_refs=body_refs,
        unreadable=[k for k, v in extracted.items() if v['kind'] == 'unreadable'])

    if map_error:
        verdicts['IMG-4'] = f'FAIL document could not be read for mapping — {map_error}'

    stats = {
        'refs_all': all_refs,
        'refs_body': body_refs,
        'media_on_disk': len(extracted),
        'mapped': mapped,
        'preamble': preamble,
        'header_footer': all_refs - body_refs,
        'vector': sum(1 for v in extracted.values() if v['kind'] == 'vector'),
        'unreadable': sum(1 for v in extracted.values() if v['kind'] == 'unreadable'),
        'questions_with_images': len([k for k in mapping if k != PREAMBLE]),
        'dangling': dangling_media_targets(z),
    }
    return verdicts, stats


def figural_consistency(mapping, q_formats, figural_value='FIGURAL', overrides=()):
    """Gate IMG-5b — cross-check image presence against format classification.

    Catches two different faults with one assertion:
      * a question that HAS an image but was not classified FIGURAL  -> image was lost
        or the classifier missed it;
      * a question classified FIGURAL with NO image                  -> misclassification,
        unless it is an explicit logged override (the INHERENTLY-VISUAL path, where a
        question is visual in nature without carrying an embedded figure).

    Returns {'image_not_figural': [...], 'figural_no_image': [...]}.
    """
    ov = set(overrides)
    with_img = {k for k in mapping if k != PREAMBLE and mapping[k]}
    image_not_figural = sorted(
        q for q in with_img
        if q in q_formats and q_formats[q] != figural_value and q not in ov)
    figural_no_image = sorted(
        q for q, f in q_formats.items()
        if f == figural_value and q not in with_img and q not in ov)
    return {'image_not_figural': image_not_figural,
            'figural_no_image': figural_no_image}


# ═══════════════════════════════════════════════════════════════════════════════
# CLUSTER V — VISION PHASE A / OBSERVATION I/O  (GAP-2026-07-26-003)
# ═══════════════════════════════════════════════════════════════════════════════
#
# Viewing an image is a CLASS T operation — it needs a tool call, and a tool call can
# only happen BETWEEN model turns. A python process cannot suspend mid-loop to make
# one. Everything here is the DETERMINISTIC half of the bridge:
#
#   PHASE A (here)                 normalise -> compose -> tile -> queue.json
#   PHASE B (model, prose)         view() each sheet -> observations.json
#   PHASE C (blueprint_core)       merge observations -> fields
#
# NOTHING HERE RAISES ON A VISION PROBLEM. An unopenable figure, a missing
# observations file and an absent Pillow all DEGRADE with a recorded reason. The
# defect this cluster fixes was silence; the remedy is visibility, not a halt.

VISION_QUEUE_NAME = 'vision_queue.json'
VISION_OBS_NAME = 'vision_observations.json'
VISION_SHEET_FMT = 'vision_sheet_{:03d}.png'


def _open_rgb(src, max_edge=1600):
    """Open one figure as bounded RGB, or return None. EC-V7 (CMYK) / EC-V8 (vector).

    Returns None rather than raising: a figure that cannot be rendered must still
    appear in the queue as an unobserved item, not vanish from the denominator.
    """
    Image = _need('PIL')
    try:
        im = Image.open(src)
        im.load()
    except Exception:
        return None
    if im.mode not in ('RGB', 'L'):
        try:
            im = im.convert('RGB')
        except Exception:
            return None
    if max(im.size) > max_edge:
        r = max_edge / float(max(im.size))
        im = im.resize((max(1, int(im.size[0] * r)), max(1, int(im.size[1] * r))),
                       Image.LANCZOS)
    return im


def _font(size):
    from PIL import ImageFont
    for loader in (
            lambda: ImageFont.truetype(
                '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', size),
            lambda: ImageFont.load_default(size=size),
            lambda: ImageFont.load_default()):
        try:
            return loader()
        except Exception:
            continue
    return None


def _compose_cell(srcs, box, max_edge):
    """Compose ONE question's figures into one cell, side by side.

    EC-V6. A 4-panel series is ONE question and must be judged as a whole, so its
    panels share a cell rather than being scattered across the sheet where their
    sequence would be lost.

    Returns (image, ok). ok is False when nothing could be rendered.
    """
    Image = _need('PIL')
    w, h = box
    canvas = Image.new('RGB', (w, h), (255, 255, 255))
    ims = [im for im in (_open_rgb(s, max_edge) for s in srcs) if im is not None]
    if not ims:
        return canvas, False
    gap = 6
    tot_w = sum(i.size[0] for i in ims) + gap * (len(ims) - 1)
    tot_h = max(i.size[1] for i in ims)
    scale = min(w / float(tot_w), h / float(tot_h), 1.0)
    if scale <= 0:
        return canvas, False
    x = int((w - tot_w * scale) / 2)
    for im in ims:
        nw = max(1, int(im.size[0] * scale))
        nh = max(1, int(im.size[1] * scale))
        canvas.paste(im.resize((nw, nh), Image.LANCZOS),
                     (x, int((h - nh) / 2)))
        x += nw + gap
    return canvas, True


# ── CLUSTER Q — ORIGINAL QUESTION POSITION IN THE DATE LABEL (v1.9) ──────────
# GAP-2026-07-27-E. Step 3 (PYQSort) renumbers questions into taxonomy order, which
# destroys the exam position. Step 5 (MockTestAnalyse) needs that position back to
# recognise a section-banded question type — an MSQ block at Q31-40 that carries no
# multi-select instruction phrase of its own. Measured on IIT_JAM_BIOTECHNOLOGY: 24 MSQ
# detected across 1,719 questions against a marking scheme reserving ~10/paper.
#
# ONE WRITER, ONE READER, ONE DEFINITION. The stamp is produced by Step 3 and consumed
# by Step 5, which is exactly the shape that produced the is_option drift v2.36 had to
# unwind: two specs each claiming in a docstring to match the other, and silently
# diverging. Both specs delegate here so the format cannot drift by construction.
#
# The field is OPTIONAL. A label written before v1.18 has no Q-part and parses to None,
# which callers MUST treat as "unknown" and never as position 0. No re-sort is forced.
ORIG_QNUM_RE = re.compile(r'\sQ(\d+)\]$')


def parse_original_q_num(date_label):
    """Return the original exam question number stamped in a date label, or None."""
    m = ORIG_QNUM_RE.search((date_label or '').strip())
    return int(m.group(1)) if m else None


def stamp_original_q_num(date_label, original_q_num):
    """Insert ' Q<N>' before a date label's closing bracket. Idempotent."""
    if original_q_num is None:
        return date_label
    lbl = (date_label or '').strip()
    if not lbl.endswith(']') or ORIG_QNUM_RE.search(lbl):
        return date_label if not lbl.endswith(']') else lbl
    return lbl[:-1] + f' Q{int(original_q_num)}]'


def question_type_for_position(original_q_num, marking_scheme):
    """Return the question_type whose q_range covers this position, or None.

    Exam-agnostic: the bands come from exam_config's marking_scheme at runtime. Returns
    None for an unknown position, an absent/empty scheme, or a position no band covers —
    every one of which means "no positional evidence", never "not this type".
    """
    if original_q_num is None or not marking_scheme:
        return None
    for band in marking_scheme:
        rng = (band or {}).get('q_range') or []
        if len(rng) == 2:
            try:
                if int(rng[0]) <= int(original_q_num) <= int(rng[1]):
                    return band.get('question_type')
            except (TypeError, ValueError):
                continue
    return None


def build_vision_queue(items, workdir, per_sheet=6, max_edge=1600,
                       cell=520, label_h=30, cols=3, fresh=False):
    """PHASE A. Turn figural questions into contact sheets plus a work queue.

    items: [{'paper_id', 'q_num', 'srcs':[path,...], 'subtopic'?, 'image_role'?}]

    Writes ``vision_queue.json`` and ``vision_sheet_NNN.png`` into workdir and returns
    the queue dict. Never raises for a vision reason.

    WHY TILE. One view() per figure does not scale: this corpus has 153 figures across
    22 papers, and a 200-exam fleet has tens of thousands. At per_sheet=6 that is 26
    view() calls instead of 153. per_sheet is a parameter so a sheet that renders badly
    can be re-tiled without touching the caller (EC-V10).

    EC-V1  no figural content -> empty queue, no sheets, valid output
    EC-V6  multi-image question -> ONE item, panels composed side by side
    EC-V8  unrenderable (vector/corrupt) -> queued, flagged, counted; never dropped
    EC-V9  Pillow absent -> degraded mode, one item per sheet, original file viewed
    EC-V15 key is (paper_id, q_num), never a bare q_num
    """
    os.makedirs(workdir, exist_ok=True)

    # ── v1.10 `fresh` — RUN-SCOPED CALLERS OPT OUT OF THE UNION ─────────────────
    # The v1.9 union is correct for Step 5, whose workdir is one-per-RUN spanning a
    # whole batch, and whose resumed sessions must not orphan prior sheets. It is
    # WRONG for Step 1 (PYQPrepare), which processes ONE source per trigger and runs
    # Phase A -> B -> C within that same trigger: anything already in the workdir is
    # contamination from a previous paper or a previous step. Under v1.8's overwrite
    # this sharing was invisible; the union surfaced it — measured: a second
    # PYQPrepare run saw queued=3 where it queued 1, so Phase B re-views the prior
    # paper's cells and Phase C counts them unobserved (amber footer, wrong counts).
    # fresh=True removes the prior queue, sheets AND observations before building,
    # restoring per-run semantics EXPLICITLY at the call site rather than by accident
    # of write mode. Default False preserves union semantics for every other caller.
    if fresh:
        import glob as _glob
        for _stale in ([os.path.join(workdir, VISION_QUEUE_NAME),
                        os.path.join(workdir, VISION_OBS_NAME)]
                       + _glob.glob(os.path.join(workdir, VISION_SHEET_FMT.format(0)
                                                 .replace('000', '*')))):
            try:
                os.remove(_stale)
            except OSError:
                pass

    # ── v1.9 IDEMPOTENT UNION (GAP-2026-07-27-B) ────────────────────────────────
    # This function WRITES vision_queue.json and vision_sheet_NNN.png, both fixed
    # names, and never read what was already there. Callers that invoked it more than
    # once per workdir therefore destroyed every earlier call's work. Measured on
    # IIT_JAM_BIOTECHNOLOGY: 153 figural questions across 22 papers, of which the
    # run-level record retained 8 — the last paper's alone.
    #
    # The caller-side fix (hoist to the batch boundary) is necessary but not
    # sufficient: a RESUMED session starts a new batch against a workdir that may
    # already hold sheets, and would orphan them the same way. Unioning here closes
    # that half. Both halves ship together.
    #
    # Safe by construction: bc.vision_tag_map derives a tag from paper_id alone, so a
    # key's tag does not depend on which OTHER keys are present, and re-tiling a
    # superset leaves every prior tag intact. The one exception is a hash-width
    # widening, which re-tags the WHOLE queue — detected and reported below.
    prior_items, prior_width = [], None
    _qpath = os.path.join(workdir, VISION_QUEUE_NAME)
    if os.path.exists(_qpath):
        try:
            with open(_qpath, encoding='utf-8') as _f:
                _prior = json.load(_f)
            prior_width = _prior.get('tag_width')
            for _pit in (_prior.get('items') or []):
                prior_items.append({'paper_id': _pit.get('paper_id'),
                                    'q_num': _pit.get('q_num'),
                                    'srcs': list(_pit.get('srcs') or []),
                                    'subtopic': _pit.get('subtopic'),
                                    'image_role': _pit.get('image_role')})
        except Exception:
            # An unreadable prior queue must not stop Phase A. Rebuild from the
            # incoming items alone; the sheets are regenerated either way.
            prior_items, prior_width = [], None

    # EC-V27. Collapse repeats of the SAME (paper_id, q_num) into one item, unioning
    # their sources. A key maps to exactly one tag, so two items sharing a key would
    # print the SAME label on two cells — and a Phase-B observation against that label
    # could not be attributed to either. One question, one tag, one cell (EC-V6).
    # Prior items are seeded FIRST so a re-run is stable; incoming items fill gaps.
    merged = {}
    for it in (prior_items + list(items or [])):
        if not isinstance(it, dict):
            continue
        try:
            key = (str(it['paper_id']), int(it['q_num']))
        except (KeyError, TypeError, ValueError):
            continue
        srcs = [s for s in (it.get('srcs') or []) if s]
        if key in merged:
            prev_srcs, prev_it = merged[key]
            for s in srcs:
                if s not in prev_srcs:
                    prev_srcs.append(s)
            for f in ('subtopic', 'image_role'):
                if prev_it.get(f) is None and it.get(f) is not None:
                    prev_it[f] = it[f]
        else:
            merged[key] = (list(srcs), dict(it))
    # Deterministic order: the queue must be reproducible byte for byte (EC-V12).
    rows = sorted(((k, v[0], v[1]) for k, v in merged.items()),
                  key=lambda r: (r[0][0], r[0][1]))

    # v1.9. Pass the prior width as a FLOOR. vision_tag_map widens on collision but
    # would otherwise also NARROW when the key set shrinks, re-tagging every surviving
    # question for no reason. Under the union the paper set grows and shrinks across
    # batches and resumed sessions, so pinning the floor removes the spurious half of
    # the re-tag risk entirely. A genuine widening remains possible — 4 hex chars is
    # ~65k codes, so a collision appears in a few percent of large corpora — and is
    # reported rather than absorbed.
    tag_map, tag_width = bc.vision_tag_map([r[0] for r in rows],
                                           width=(prior_width or 4))
    stats = {'queued': len(rows), 'sheets': 0, 'unrenderable': 0, 'degraded': False,
             # v1.9 — union provenance, so a caller can tell a fresh queue from a
             # resumed one without diffing files.
             'carried_in': len(prior_items), 'submitted': len(items or []),
             'tag_generation_changed': False}
    if prior_width is not None and prior_width != tag_width:
        # EC-V12. A paper_id hash collision widened the code for the WHOLE queue, so
        # every tag changed and observations written under the previous generation can
        # no longer be matched. Never silent: Phase C would otherwise report them as
        # simply unobserved and the operator would re-view sheets that were already done.
        stats['tag_generation_changed'] = True
        stats['prior_tag_width'] = prior_width

    degraded = False
    try:
        _need('PIL')
    except DependencyMissing:
        degraded = True
        stats['degraded'] = True

    queue_items, sheets = [], []
    if degraded:
        # EC-V9. One view() per figure against the ORIGINAL file. Slower, never silent.
        for i, (key, srcs, it) in enumerate(rows):
            queue_items.append({
                'tag': tag_map[key], 'paper_id': key[0], 'q_num': key[1],
                'subtopic': it.get('subtopic'), 'image_role': it.get('image_role'),
                'sheet': srcs[0] if srcs else None, 'cell': 0,
                'renderable': bool(srcs), 'srcs': srcs})
            if not srcs:
                stats['unrenderable'] += 1
    else:
        Image = _need('PIL')
        from PIL import ImageDraw
        n_cols = max(1, min(int(cols), int(per_sheet)))
        n_rows = max(1, -(-int(per_sheet) // n_cols))
        cw, ch = int(cell), int(cell) + int(label_h)
        font = _font(20)
        for start in range(0, len(rows), int(per_sheet)):
            chunk = rows[start:start + int(per_sheet)]
            idx = len(sheets) + 1
            name = VISION_SHEET_FMT.format(idx)
            sheet = Image.new('RGB', (n_cols * cw, n_rows * ch), (245, 245, 245))
            d = ImageDraw.Draw(sheet)
            for j, (key, srcs, it) in enumerate(chunk):
                cx, cy = (j % n_cols) * cw, (j // n_cols) * ch
                img, ok = _compose_cell(srcs, (cw - 8, ch - label_h - 8), max_edge)
                if not ok:
                    stats['unrenderable'] += 1
                sheet.paste(img, (cx + 4, cy + label_h + 4))
                d.rectangle([cx + 2, cy + 2, cx + cw - 2, cy + ch - 2],
                            outline=(0, 0, 0), width=2)
                label = tag_map[key] + ('' if ok else '  [UNRENDERABLE]')
                d.rectangle([cx + 2, cy + 2, cx + cw - 2, cy + label_h],
                            fill=(0, 0, 0))
                d.text((cx + 10, cy + 6), label, fill=(255, 255, 255), font=font)
                queue_items.append({
                    'tag': tag_map[key], 'paper_id': key[0], 'q_num': key[1],
                    'subtopic': it.get('subtopic'), 'image_role': it.get('image_role'),
                    'sheet': name, 'cell': j, 'renderable': ok, 'srcs': srcs})
            path = os.path.join(workdir, name)
            sheet.save(path, 'PNG', optimize=True)
            sheets.append(name)
        stats['sheets'] = len(sheets)

    queue = {'version': 1, 'tag_width': tag_width, 'per_sheet': int(per_sheet),
             'degraded': degraded, 'sheets': sheets, 'items': queue_items,
             'stats': stats}
    with open(os.path.join(workdir, VISION_QUEUE_NAME), 'w', encoding='utf-8') as f:
        json.dump(queue, f, indent=2, sort_keys=True)
    return queue


def load_vision_queue(workdir):
    """Read back a queue. Returns (queue_dict, reason). Never raises."""
    path = os.path.join(workdir, VISION_QUEUE_NAME)
    if not os.path.exists(path):
        return None, 'no queue file — Phase A has not run for this batch'
    try:
        with open(path, encoding='utf-8') as f:
            q = json.load(f)
    except Exception as exc:
        return None, f'queue file unreadable ({exc})'
    if not isinstance(q, dict) or not isinstance(q.get('items'), list):
        return None, 'queue file has no items list'
    return q, ''


def load_vision_observations(workdir, path=None):
    """PHASE B -> C boundary. Read observations. Returns (list, reason). Never raises.

    A missing or malformed file is NOT an error condition here — it is the ordinary
    'Phase B has not run yet' state, and it must flow through to Phase C as zero
    observations so the run completes and QV-14 reports the gap. Raising here would
    reintroduce the halt this design exists to avoid.

    Accepts a bare list, or a dict carrying 'observations'.
    """
    p = path or os.path.join(workdir, VISION_OBS_NAME)
    if not os.path.exists(p):
        return [], 'no observations file — Phase B has not run'
    try:
        with open(p, encoding='utf-8') as f:
            data = json.load(f)
    except Exception as exc:
        return [], f'observations file unreadable ({exc}) — treated as zero observations'
    if isinstance(data, dict):
        data = data.get('observations', [])
    if not isinstance(data, list):
        return [], 'observations file is neither a list nor {"observations": [...]}'
    return [d for d in data if isinstance(d, dict)], ''


def write_vision_observations(workdir, observations, path=None):
    """Write the Phase-B result. Provided so the protocol has ONE spelling of the file."""
    p = path or os.path.join(workdir, VISION_OBS_NAME)
    os.makedirs(os.path.dirname(os.path.abspath(p)) or '.', exist_ok=True)
    with open(p, 'w', encoding='utf-8') as f:
        json.dump({'observations': list(observations or [])}, f,
                  indent=2, sort_keys=True)
    return p


# ── SHARED OPTION PREDICATE (Cluster I) ─────────────────────────────────────────
# GAP-2026-07-26-002 follow-up, audit_deep [XSPEC-DRIFT]. is_option() was defined
# THREE times — Framework_MockTestAnalyse (Step 5), Framework_PYQSort (Step 1) and
# Framework_PYQAnalyse (Step 4) — each carrying a docstring claiming alignment with
# the others. Step 5 then gained the image-option path and the other two did not,
# so the alignment claim became false and the SAME defect the wave fixed in Step 5
# stayed live in Step 1: PYQSort counts options with this predicate, so an image
# option was silently UNDERCOUNTED.
#
# One definition here; all three steps delegate. Drift is now impossible by
# construction rather than by comment.

OPT_PATTERNS = [
    r'^([1-5])\.\s+(.+)',           # 1. 2. 3. 4. 5.  (up to 5 options)
    r'^([A-Ea-e])\.\s+(.+)',        # A. B. C. D. E.
    r'^\(([1-5])\)\s+(.+)',         # (1) (2) (3) (4) (5)
    r'^\(([A-Ea-e])\)\s+(.+)',      # (A) (B) (C) (D) (E) / (a)(b)(c)(d)(e)
    r'^([A-Ea-e])\)\s+(.+)',        # A) B) C) D) E) / a) b) c) d) e)
]

# A marker with NO text after it. Every OPT_PATTERNS entry requires \s+(.+), so an
# IMAGE OPTION — the paragraph is literally "1." and the picture follows — matched
# none of them.
BARE_OPT_PATTERNS = [
    r'^\(?([1-5])\)?\.?\s*$',       # 1.  1)  (1)  1
    r'^\(?([A-Ea-e])\)?\.?\s*$',    # A.  A)  (A)  A
]

BLIP_TAG          = f'{{{A_NS}}}blip'
VML_IMAGEDATA_TAG = f'{{{V_NS}}}imagedata'


def para_has_image(para):
    """True when this paragraph carries an embedded image (DrawingML blip or VML).

    Accepts EITHER form the corpus actually passes:
      * a python-docx Paragraph  — Step 5 walks doc.paragraphs and has `._p`
      * a raw <w:p> lxml element — Steps 1/3 walk body_elems and ARE the element

    That difference is not cosmetic. A delegation that assumed one form would raise
    AttributeError on the other, or — worse — silently return False and reinstate the
    exact undercount this function exists to prevent.
    """
    if para is None:
        return False
    el = getattr(para, '_p', para)          # Paragraph -> element; element -> itself
    it = getattr(el, 'iter', None)
    if it is None:
        return False
    for e in it():
        if getattr(e, 'tag', None) in (BLIP_TAG, VML_IMAGEDATA_TAG):
            return True
    return False


def para_is_option_content(para):
    """True when a bare option LABEL's paragraph carries the option's actual content.

    GAP-2026-08-05-001 (S-1). para_has_image() recognises only a DrawingML blip or a
    VML imagedata reference, so an option whose content is an OMML equation or an
    embedded OLE object was not recognised as an option at all — even though PYQSort
    CHECK 4 accepts "(c) inline drawing" AND "(d) OMML formula" as valid option forms.
    Producer and consumer disagreed about what an option IS.

    Measured on SSC_CGL_TIER1 09-Sep-2024 Shift 1, a delivered sorted paper: Q.75
    options 1 and 2 are bare labels "1." / "2." whose content is <m:oMath>. Both
    returned False, so that question silently lost half its option set at Step 5.

    include_autonumber=False is deliberate: on a bare label the rendered w:numPr number
    merely duplicates the "1." already present as text and is not option CONTENT, so an
    auto-numbered but otherwise empty paragraph must NOT be promoted to an option.
    """
    return bool(para_has_image(para) or
                bc.paragraph_is_content_bearing(para, include_autonumber=False))


def is_option(text, para=None):
    """THE option-line predicate. Every step delegates to this; nobody re-implements it.

    Pass `para` wherever the paragraph is available so an image-only option is
    recognised. Without it the behaviour is byte-identical to the historic text-only
    predicate, so delegating is safe even at a call site that has no paragraph — but
    a call site that HAS one and does not pass it keeps the undercount.
    """
    t = (text or '').strip()
    if any(re.match(p, t) for p in OPT_PATTERNS):
        return True
    if para is not None and para_is_option_content(para):
        return any(re.match(p, t) for p in BARE_OPT_PATTERNS)
    return False


W_T_TAG = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
M_T_TAG = '{http://schemas.openxmlformats.org/officeDocument/2006/math}t'


def text_of(para):
    """Canonical VISIBLE text of a paragraph — every <w:t> AND <m:t> descendant,
    concatenated in document order. (GAP-2026-08-07-OMML, remedy M3.)

    python-docx's p.text is <w:t>-only, so a paragraph whose payload is an
    <m:oMath> equation reads as a bare label ("3. ") or as empty. That is the
    exact producer/consumer disagreement S-1 fixed structurally for option
    DETECTION; text_of() closes the remaining half — any consumer that needs the
    words a reader actually sees (dedupe keys, taxonomy matching, QV text
    audits, PYQPrepare CHECK 9) reads text_of(), never p.text, whenever OMML may
    be present. is_option()/BARE_OPT_PATTERNS stay the structural predicates.

    Accepts EITHER a python-docx Paragraph or a raw <w:p> lxml element — the
    same duality, for the same reason, as para_has_image().
    """
    if para is None:
        return ''
    el = getattr(para, '_p', para)
    it = getattr(el, 'iter', None)
    if it is None:
        return ''
    out = []
    for e in it():
        if getattr(e, 'tag', None) in (W_T_TAG, M_T_TAG):
            out.append(e.text or '')
    return ''.join(out)


def clean_option_text(text):
    """Option text with its marker stripped. '' for a bare marker — the option content
    IS the image, and an all-empty option set is what resolves option_format to
    'image_only' downstream."""
    t = (text or '').strip()
    for p in OPT_PATTERNS:
        m = re.match(p, t)
        if m:
            return m.group(2).strip()
    for p in BARE_OPT_PATTERNS:
        if re.match(p, t):
            return ''
    return t


# ═══════════════════════════════════════════════════════════════════════════════
# CLUSTER J SHELL — SIZE GOVERNOR
# ═══════════════════════════════════════════════════════════════════════════════

def media_display_inches(path):
    """media basename -> widest display width in inches, across every story part.

    Read from each drawing's own <wp:extent cx=".."> (EMU), falling back to the VML
    style width in points.

    This is what makes downscaling CONTENT-AWARE. A figure is only over-resolved
    relative to the size it is actually displayed at: one measured figure in this corpus
    is 2546 px wide but appears 2.09 inches wide, i.e. 1216 DPI. A blanket pixel cap
    would degrade a small, correctly-sized figure while barely touching that one.
    """
    EMU_PER_INCH = 914400
    z = _open_zip(path)
    widths = {}
    for part in z.namelist():
        if not STORY_PARTS_RE.match(part):
            continue
        rels = _rels_for(z, part)
        if not rels:
            continue
        xml = z.read(part).decode('utf-8', 'ignore')
        for blk in re.findall(r'<w:drawing>.*?</w:drawing>', xml, re.S):
            cx = re.search(r'<wp:extent[^>]*\bcx="(\d+)"', blk)
            rid = re.search(r'<a:blip[^>]*r:embed="([^"]+)"', blk)
            if not (cx and rid):
                continue
            base = rels.get(rid.group(1))
            if base:
                widths[base] = max(widths.get(base, 0.0),
                                   int(cx.group(1)) / float(EMU_PER_INCH))
        for blk in re.findall(r'<w:pict>.*?</w:pict>', xml, re.S):
            rid = re.search(r'r:id="([^"]+)"', blk)
            st = re.search(r'style="[^"]*width:\s*([\d.]+)pt', blk)
            if not rid:
                continue
            base = rels.get(rid.group(1))
            if base:
                widths[base] = max(widths.get(base, 0.0),
                                   float(st.group(1)) / 72.0 if st else 0.0)
    return widths


def _is_line_art(im, sample_cap=250_000):
    """<=256 distinct colours => line art. Deterministic; work is bounded."""
    s = im.convert('RGB')
    if s.size[0] * s.size[1] > sample_cap:
        r = (sample_cap / float(s.size[0] * s.size[1])) ** 0.5
        s = s.resize((max(1, int(s.size[0] * r)), max(1, int(s.size[1] * r))))
    return len(set(s.getdata())) <= 256


def _normalise_mode(im):
    """Return (image, has_alpha) in a colour space that can actually be encoded.

    CMYK is the single largest cause of bloat in this corpus — 100% of measured figures
    were CMYK. It carries a fourth channel and is written by design tools at print
    quality with no chroma subsampling; a 544x568 diagram occupied 595 KB, about 1.9
    bytes per pixel where a correct RGB JPEG uses 0.15-0.3. Converting to RGB is
    visually lossless for screen use and is where nearly all of the reduction comes from.

    Palette ('P') and 16-bit modes are also normalised, because PIL cannot quantise or
    JPEG-encode them directly — attempting it raises "image has wrong mode".
    """
    if im.mode in ('RGBA', 'LA'):
        return im, True
    if im.mode == 'P':
        conv = im.convert('RGBA' if 'transparency' in im.info else 'RGB')
        return conv, conv.mode == 'RGBA'
    if im.mode in ('L', '1'):
        return im.convert('L'), False
    if im.mode in ('I;16', 'I', 'F'):
        return im.convert('L'), False
    return im.convert('RGB'), False


def _recode(raw, quality, dpi_ceiling, display_in):
    """Re-encode one media part. Returns (bytes, ext, how) or (None, None, why)."""
    Image = _need('PIL')
    try:
        im = Image.open(io.BytesIO(raw))
        im.load()
    except Exception:
        return None, None, 'unreadable or vector — left untouched'

    src_fmt = (im.format or '').upper()
    im, has_alpha = _normalise_mode(im)

    if dpi_ceiling and display_in and display_in > 0:
        effective_dpi = im.size[0] / float(display_in)
        if effective_dpi > dpi_ceiling:
            r = dpi_ceiling / effective_dpi
            im = im.resize((max(1, int(im.size[0] * r)), max(1, int(im.size[1] * r))),
                           Image.LANCZOS)

    route = bc.classify_media_route(
        src_fmt, has_alpha, False if has_alpha else _is_line_art(im))

    out = io.BytesIO()
    if route == 'png':
        im.save(out, 'PNG', optimize=True)
        return out.getvalue(), 'png', 'png/alpha'
    if route == 'png-lineart':
        im.convert('P', palette=Image.ADAPTIVE, colors=256).save(out, 'PNG', optimize=True)
        return out.getvalue(), 'png', 'png/lineart'
    im.save(out, 'JPEG', quality=quality, optimize=True, progressive=True, subsampling=2)
    return out.getvalue(), 'jpeg', ('jpeg/reencode' if src_fmt in ('JPEG', 'JPG')
                                    else 'jpeg/photo')


#: Extensions already carrying the right content type for each encoder route, so a
#: part using one needs no rename. Renaming gratuitously (image1.jpg -> image1.jpeg)
#: churns every relationship that references it and manufactures stem collisions that
#: would not otherwise exist.
_EXT_OK = {'jpeg': ('.jpeg', '.jpg'), 'png': ('.png',)}


def _target_media_name(name, ext):
    """Output name for one re-encoded media part. Keeps the name when the existing
    extension already maps to the encoder's content type."""
    stem, cur = os.path.splitext(name)
    if cur.lower() in _EXT_OK[ext]:
        return name
    return stem + ('.jpeg' if ext == 'jpeg' else '.png')


def _no_gain_guard(src, dst, orig, report):
    """v1.0.3 — never hand back a file LARGER than the one we were given.

    Per-PART growth is already prevented (_recode keeps the original bytes whenever a
    re-encode comes out bigger), but the ZIP CONTAINER can still grow: a package whose
    media is already optimally encoded gains nothing from re-deflation and can lose to
    the original writer's settings by a few bytes.

    Before `always` this was unreachable — a file under budget was never re-encoded at
    all — so no document-level rule was needed. It is now the normal case for small
    documents, and Framework_PYQCompress CHECK 5 treats growth as a HARD STOP, so a file
    with nothing to gain would have ABORTED THE RUN rather than simply being left alone.
    Restoring the original bytes makes "nothing to gain" a reportable outcome instead.
    """
    if os.path.getsize(dst) >= orig:
        shutil.copyfile(src, dst)
        return dict(report, bytes=orig, ratio=1.0, no_gain=True)
    return dict(report, no_gain=False)


def optimize_docx(src, dst, budget=None, force_tier=None, always=False):
    """Bring `src` under `budget` using blueprint_core's deterministic tier ladder.

    Stops at the FIRST tier that meets the budget, so the least invasive change that
    works is the one applied. On the measured corpus every paper cleared at T1, meaning
    no image was downscaled at all and every pixel dimension was preserved.

    Guarantees: never grows a part, never drops a part, never changes the part count.
    Returns (ok, report, log). ok=False means the ladder floor was reached and the file
    is STILL over budget — the caller should WARN and route to the upload lane, not halt,
    because a legitimately huge paper must not block delivery.

    The caller MUST still run assert_docx_parity(); this function does not self-certify.
    """
    budget = bc.SIZE_BUDGET if budget is None else budget
    orig = os.path.getsize(src)
    # v1.0.3 — `always` runs the ladder regardless of input size. The ladder itself needs
    # no change: a small file's T1 output is under budget, so it returns at T1, which is
    # the least invasive tier. NOTE for callers: force_tier is NOT a substitute — it pins
    # a single tier AND makes the final `return True` unconditional, so an oversized file
    # would stop at that tier, still over budget, and be reported as a success.
    if orig <= budget and not force_tier and not always:
        return True, {'tier': 'T0', 'bytes': orig, 'orig': orig, 'ratio': 1.0,
                      'quality': None, 'dpi_ceiling': None,
                      'note': 'already under budget — untouched'}, []

    display = media_display_inches(src)
    report, log = {}, []

    for tier, quality, dpi in bc.TIER_LADDER:
        if force_tier and tier != force_tier:
            continue
        zin = zipfile.ZipFile(src)
        names = zin.namelist()
        name_set = set(names)
        parts, renames, outname, log = {}, {}, {}, []

        for n in names:
            raw = zin.read(n)
            if MEDIA_RE.match(n):
                base = os.path.basename(n)
                data, ext, how = _recode(raw, quality, dpi, display.get(base, 0.0))
                if data is None or len(data) >= len(raw):
                    new, payload = n, raw                  # never grow a part
                    log.append((base, len(raw), len(raw), how + ' [kept]'))
                else:
                    new = _target_media_name(n, ext)
                    payload = data
                    log.append((base, len(raw), len(data), how))

                # v1.0.2 — a media part may NEVER overwrite another. Two parts that
                # share a stem and take the same encoder route (image1.png +
                # image1.bmp, both photographic) both want image1.jpeg; the old code
                # collapsed them into ONE zip member, so Word rendered whichever it
                # met first and the other figure was silently replaced. Every count
                # the parity assert checks — parts, media, refs — still matched, and
                # the per-stem pixel check is skipped at T2-T4, so the corrupted file
                # was delivered clean. Disambiguate instead of overwriting.
                if new in parts or (new != n and new in name_set):
                    stem, dot_ext = os.path.splitext(new)
                    k = 2
                    while (f'{stem}__c{k}{dot_ext}' in parts
                           or f'{stem}__c{k}{dot_ext}' in name_set):
                        k += 1
                    new = f'{stem}__c{k}{dot_ext}'

                parts[new] = payload
                outname[n] = new
                if os.path.basename(new) != base:
                    renames[base] = os.path.basename(new)
            else:
                parts[n] = raw
                outname[n] = n
        zin.close()

        # v1.0.2 — belt-and-braces: one member out for every member in. Catches any
        # future path that collapses two parts into one key, including a duplicate
        # member name in a damaged input package.
        if len(parts) != len(names) or len(outname) != len(names):
            raise IntegrityError(
                f'part collision in {os.path.basename(src)} at tier {tier}: '
                f'{len(names)} package member(s) in, {len(parts)} out. Two parts '
                'resolved to a single name — refusing to write a package in which '
                'one figure silently replaces another.')

        # Renaming a media part invalidates every reference to it. Rewrite the
        # relationship targets and guarantee the content-type defaults exist, or Word
        # reports the file as corrupt.
        #
        # v1.0.2 — SINGLE PASS. Sequential str.replace() chained: with renames
        # {'image1.png': 'image1.jpeg', 'image1.jpeg': 'image1__c2.jpeg'} the first
        # substitution manufactured a second 'image1.jpeg', which the second
        # substitution then rewrote as well, pointing both references at one part.
        rename_re = (re.compile('|'.join(re.escape(k) for k in
                                sorted(renames, key=len, reverse=True)))
                     if renames else None)
        for pn in list(parts):
            if pn.endswith('.rels') or pn == '[Content_Types].xml':
                txt = parts[pn].decode('utf-8')
                if rename_re is not None:
                    txt = rename_re.sub(lambda m: renames[m.group(0)], txt)
                if pn == '[Content_Types].xml':
                    for ext_, ctype in (('jpeg', 'image/jpeg'), ('png', 'image/png'),
                                        ('jpg', 'image/jpeg')):
                        if f'Extension="{ext_}"' not in txt:
                            txt = re.sub(
                                r'(<Types[^>]*>)',
                                rf'\1<Default Extension="{ext_}" ContentType="{ctype}"/>',
                                txt, count=1)
                parts[pn] = txt.encode('utf-8')

        os.makedirs(os.path.dirname(os.path.abspath(dst)) or '.', exist_ok=True)
        with zipfile.ZipFile(dst, 'w', zipfile.ZIP_DEFLATED, compresslevel=9) as zo:
            for n in names:
                # v1.0.2 — write through the explicit map built above. The previous
                # reverse-lookup guessed the output name by trying '.jpeg' then
                # '.png', which cannot express a disambiguated name and silently
                # resolved two inputs to one key.
                zo.writestr(outname[n], parts[outname[n]])

        size = os.path.getsize(dst)
        report = {'tier': tier, 'bytes': size, 'orig': orig,
                  'ratio': size / float(orig), 'quality': quality, 'dpi_ceiling': dpi}
        if size <= budget or force_tier:
            return True, _no_gain_guard(src, dst, orig, report), log

    return False, _no_gain_guard(src, dst, orig, report), log


def docx_invariants(path):
    """Every observable property that optimisation must leave untouched."""
    import hashlib
    import xml.dom.minidom as minidom
    Image = _need('PIL')

    z = _open_zip(path)
    d = _safe_document(path)
    xml = z.read('word/document.xml').decode('utf-8', 'ignore')
    media = sorted(n for n in z.namelist() if n.startswith('word/media/'))

    xmlerr = 0
    for n in z.namelist():
        if n.endswith(('.xml', '.rels')):
            try:
                minidom.parseString(z.read(n))
            except Exception:
                xmlerr += 1

    txt = '\n'.join(p.text for p in d.paragraphs)
    tbl = ''.join(c.text for t in d.tables for r in t.rows for c in r.cells)
    return {
        'parts': len(z.namelist()),
        'media': len(media),
        'paras': len(d.paragraphs),
        'tables': len(d.tables),
        'shapes': len(d.inline_shapes),
        'chars': len(txt),
        'tablechars': len(tbl),
        'omml': xml.count('<m:oMath'),
        'drawings': xml.count('<w:drawing'),
        'pict': xml.count('<w:pict'),
        'hyperlinks': xml.count('<w:hyperlink'),
        'texthash': hashlib.sha256(txt.encode()).hexdigest(),
        'refs': count_image_refs(z)[0],
        'xmlerr': xmlerr,
        'zipbad': z.testzip() or 'OK',
        'dangling': dangling_media_targets(z),
        'px': {os.path.basename(n): Image.open(io.BytesIO(z.read(n))).size
               for n in media if os.path.splitext(n)[1].lower() in RASTER_EXT},
    }


def assert_docx_parity(src, dst, allow_resample=False):
    """Raise IntegrityError unless every observable property survived optimisation.

    Deliberately includes the SHA256 of the extracted text, the OMML equation count and
    the per-image pixel dimensions.

    A governor that quietly dropped a figure would still produce a smaller, perfectly
    openable document — and because the pipeline reads question stems as TEXT, the loss
    would not announce itself anywhere downstream. Byte size and "it opens in Word" are
    not evidence of correctness.

    allow_resample=True permits pixel dimensions to change (tiers T2-T4 downscale by
    design) while still enforcing every other invariant.
    """
    before, after = docx_invariants(src), docx_invariants(dst)
    keys = ['parts', 'media', 'paras', 'tables', 'shapes', 'chars', 'tablechars',
            'omml', 'drawings', 'pict', 'hyperlinks', 'texthash', 'refs']
    failures = [f'{k}: {before[k]} -> {after[k]}' for k in keys if before[k] != after[k]]
    if after['xmlerr']:
        failures.append(f"{after['xmlerr']} malformed XML part(s)")
    if after['zipbad'] != 'OK':
        failures.append(f"corrupt zip member: {after['zipbad']}")
    if after['dangling']:
        failures.append(f"dangling relationship(s): {after['dangling'][:3]}")
    if not allow_resample:
        # v1.0.1 — the original comparison keyed on the full basename, so a legitimately
        # renamed part (image1.png -> image1.jpeg on the photo route) had no counterpart
        # in `after` and was reported as a dimension change on a document whose
        # dimensions were untouched — a FALSE-POSITIVE HARD STOP. Dormant on the
        # measured corpus, whose papers store .jpeg parts the jpeg route leaves named as
        # they were; it fired on any PNG-sourced photograph. Fixed by keying on the
        # extension-less name.
        #
        # v1.0.2 — keying on the STEM has the same defect one level up: the governor may
        # now emit a disambiguated name (image1__c2.jpeg) when two parts share a stem and
        # take the same route, which changes the stem and reproduced the identical false
        # positive. Pixel dimensions are therefore compared as an unordered MULTISET over
        # all raster parts, which is what the invariant actually asserts — "no figure's
        # pixel dimensions changed" — and is immune to any renaming the governor performs.
        b_px = sorted(before['px'].values())
        a_px = sorted(after['px'].values())
        if b_px != a_px:
            def _by_stem(px):
                out = {}
                for k, v in px.items():
                    out.setdefault(os.path.splitext(k)[0], []).append(v)
                return out

            b_stem, a_stem = _by_stem(before['px']), _by_stem(after['px'])
            named = [k for k, v in b_stem.items() if sorted(v) != sorted(a_stem.get(k, []))]
            failures.append(
                'pixel dimensions changed: '
                + (str(named[:3]) if named else f'{b_px[:3]} -> {a_px[:3]}'))
    if failures:
        raise IntegrityError(
            'Optimised document failed the parity assertion — HARD STOP.\n  '
            + '\n  '.join(failures))
    return {'before': before, 'after': after}


# ═══════════════════════════════════════════════════════════════════════════════
# CLUSTER K SHELL — ANALYSIS DOC  (the Step-2c taxonomy artefact)
# ═══════════════════════════════════════════════════════════════════════════════
# GAP-2026-07-25-002. Before this cluster the Analysis doc had FOUR independent
# readers implementing FOUR different structural conventions, and no writer at all
# (Framework_PYQAnalyse S4-2 was `pass`). Measured on the live IIT_JAM_BIOTECHNOLOGY
# doc, three of the four readers were wrong:
#
#   Framework_PYQSort S1-2          1 subject  / 26 topics / 131 subtopics  (truth 6/26/131)
#   MockTestAnalyse extract_...     1 section  /  0 subtopics
#   MockTestAnalyse _extract_..._   0 sections /  0 subtopics
#   Blueprint S2-2                  prose ("Claude extracts") — non-deterministic
#
# The first flattened every subject into the first one (a `not section_name` latch);
# the other two read doc.paragraphs while every subtopic lives in a TABLE, and keyed
# hierarchy off Word Heading styles that the generator does not emit (para.style is
# None for every paragraph of a real doc).
#
# This cluster is now THE reader, THE writer and THE verifier. Every consumer
# delegates. Heading recognition is delegated one level further, to
# blueprint_core.parse_taxonomy_level() — so all six level-1 label forms
# (Subject/Domain/Section/Part/Area) and all six level-2 forms
# (Topic/Chapter/Unit/Module/Block, numbered or colon-style, case-insensitive)
# work in EVERY step automatically, including any form the engine learns later.
#
# WHY A WRITER LIVES HERE TOO: correctness is asserted by round-trip,
# read(write(taxonomy)) == taxonomy, over a GENERATED matrix of exam shapes. The
# framework serves ~200 exams; it cannot be validated against 200 real corpora, so
# the self-test generates the shapes instead. That matrix is the correctness claim.

ANALYSIS_DOC_SUFFIX = '_PYQ_Analysis.docx'
ANALYSIS_SEARCH_DIRS = ('/mnt/project/', '/mnt/user-data/uploads/')

_NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Cells that are never a subtopic name. Compared case-insensitively after strip.
# A legitimate subtopic literally named "TOTAL" would be dropped here — that is
# SAFE rather than silent, because the per-topic "Subtopics: k" declaration then
# disagrees with the parse and verify_analysis_doc() HARD STOPS naming the topic.
_SKIP_CELLS = frozenset(('subtopic', 'topic', 'grand total', 'total', ''))

# Declaration patterns. The doc states its own shape THREE times per subject; the
# reader asserts its parse against all three. This is what makes a mis-parse loud
# without needing any second artefact.
_D1_RE = re.compile(r'(\d+)\s*Topics?\b.*?\b(\d+)\s*Subtopics?', re.IGNORECASE)
_D3_RE = re.compile(r'\bSubtopics?\s*:\s*(\d+)', re.IGNORECASE)
_GRAND_RE = re.compile(r'^GRAND\s+TOTAL$', re.IGNORECASE)


class AnalysisDocError(Exception):
    """Raised for every Analysis-doc fault. Callers surface .args[0] verbatim as a
    HARD STOP: the message always names the file, the fault and the operator action."""


def _para_text(child):
    """Text of a w:p element, runs joined. docx-js splits runs on '&' and on style
    changes, so joining is mandatory — reading only the first run truncates names."""
    return ''.join(t.text for t in child.iter('{%s}t' % _NS_W) if t.text).strip()


def _table_rows(child, doc):
    """Logical rows of a w:tbl — ONE entry per ANCHOR cell, never a repeat.

    GAP-2026-07-29-TBL. This used to read row.cells, which returns one entry per
    GRID COLUMN and repeats the anchor for every covered position: a 4-column
    merged header came back as ['Printers','Printers','Printers','Printers'] and a
    vertically merged label reappeared in every row it spanned. Harmless while the
    corpus contained only flat tables — and a data-corruption bug the moment Step 1
    started emitting the merges it always should have. Verified on python-docx
    1.2.0.

    A distinct-tc dedupe is NOT sufficient: row.cells substitutes the ANCHOR object
    for a vMerge continuation, so the anchor's text is returned again in the
    continuation row. Only a raw <w:tr>/<w:tc> walk that inspects w:vMerge can tell
    a continuation from a genuinely repeated value.

    Contract is UNCHANGED — list[list[str]] — and output on a flat table is
    byte-identical to the previous implementation.
    """
    W = '{%s}' % _NS_W
    rows = []
    for tr in child.findall(W + 'tr'):
        row = []
        for tc in tr.findall(W + 'tc'):
            tcPr = tc.find(W + 'tcPr')
            if tcPr is not None:
                v = tcPr.find(W + 'vMerge')
                if v is not None and (v.get(W + 'val') or 'continue') == 'continue':
                    continue                     # covered position, not an anchor
            row.append(''.join(t.text or '' for t in tc.iter(W + 't')).strip())
        rows.append(row)
    return rows


def read_table_spec(child, doc=None):
    """Read a w:tbl back into a TableSpec, spans included (GAP-2026-07-29-TBL).

    For consumers that need GEOMETRY, not just values — CHECK 17, Step 7 reuse of a
    PYQ table, and any future step that must rebuild a table it did not write.
    _table_rows deliberately stays value-only so its single legacy caller is
    untouched.
    """
    W = '{%s}' % _NS_W
    tr_list = child.findall(W + 'tr')
    grid, pending = [], {}
    for ri, tr in enumerate(tr_list):
        row, ci = [], 0
        for tc in tr.findall(W + 'tc'):
            tcPr = tc.find(W + 'tcPr')
            cs, vm = 1, None
            if tcPr is not None:
                g = tcPr.find(W + 'gridSpan')
                if g is not None:
                    try:
                        cs = int(g.get(W + 'val') or 1)
                    except (TypeError, ValueError):
                        cs = 1
                v = tcPr.find(W + 'vMerge')
                if v is not None:
                    vm = v.get(W + 'val') or 'continue'
            if vm == 'continue':
                for key in pending:
                    if key[1] == ci:
                        pending[key]['rs'] += 1
                        break
                ci += cs
                continue
            cell = {'t': ''.join(t.text or '' for t in tc.iter(W + 't')).strip()}
            if cs > 1:
                cell['cs'] = cs
            if vm == 'restart':
                cell['rs'] = 1
                pending[(ri, ci)] = cell
            row.append(cell)
            ci += cs
        grid.append(row)
    for cell in list(pending.values()):
        if cell.get('rs', 1) <= 1:
            cell.pop('rs', None)
    return {'grid': grid, 'header_rows': 1}


def discover_analysis_doc(search_dirs=ANALYSIS_SEARCH_DIRS):
    """Locate the ONE Analysis doc. Returns a path; raises AnalysisDocError otherwise.

    There is exactly one conforming name, [ExamCode]_PYQ_Analysis.docx. The
    per-subject form [ExamCode]_PYQ_Analysis_[Subject].docx was retired by
    PYQAnalyse v2.6 and no project has ever held one, so it is NOT a supported
    generation — it is a nonconforming file, and it is named as such rather than
    silently ignored (a silently-ignored file is how an operator ends up sorting
    against a taxonomy they did not intend).
    """
    seen, found, nonconforming, word97 = set(), [], [], []
    for d in search_dirs:
        if not os.path.isdir(d):
            continue
        for name in sorted(os.listdir(d)):
            if name.startswith('~$') or name.startswith('.'):
                continue                      # Word lock file / hidden artefact
            low = name.lower()
            if name in seen:
                continue                      # same name in both dirs: project wins
            if low.endswith(ANALYSIS_DOC_SUFFIX.lower()):
                seen.add(name)
                found.append(os.path.join(d, name))
            elif low.endswith('.docx') and '_pyq_analysis_' in low:
                seen.add(name)
                nonconforming.append(name)
            elif low.endswith('.doc') and '_pyq_analysis' in low:
                seen.add(name)
                word97.append(name)

    if len(found) > 1:
        raise AnalysisDocError(
            "HARD STOP: %d Analysis docs found — exactly one is expected.\n"
            "  %s\n"
            "One project holds one exam. Remove the ones that do not belong to this "
            "exam and re-run." % (len(found), '\n  '.join(os.path.basename(p) for p in found)))
    if found:
        return found[0]
    if word97:
        raise AnalysisDocError(
            "HARD STOP: %s is a Word 97 .doc, which cannot be read.\n"
            "Open it in Word and Save As .docx, then upload it as %s."
            % (word97[0], '[ExamCode]' + ANALYSIS_DOC_SUFFIX))
    if nonconforming:
        raise AnalysisDocError(
            "HARD STOP: found %s, which is not the expected deliverable name.\n"
            "PYQApprove (Step 2c) produces ONE merged doc named "
            "[ExamCode]%s containing every subject. The per-subject form was retired "
            "in PYQAnalyse v2.6.\n"
            "Re-run PYQApprove and upload the doc it delivers."
            % (nonconforming[0], ANALYSIS_DOC_SUFFIX))
    # v1.2 (GAP-2026-07-25-003): the extracted-text advice that used to sit here was
    # both unreachable — the platform preserves the filename, so the file is never
    # MISSING, only transformed — and, after this version, wrong: the text form is
    # READ, not diagnosed. No logic change; message only.
    raise AnalysisDocError(
        "HARD STOP: no Analysis doc found in %s.\n"
        "Expected [ExamCode]%s from PYQApprove (Step 2c).\n"
        "NEXT ACTION: upload the Analysis doc PYQApprove delivered to this project's "
        "Files section." % (' or '.join(search_dirs), ANALYSIS_DOC_SUFFIX))


# ══════════════════════════════════════════════════════════════════════════════
# INGEST FORMS (GAP-2026-07-25-003)
# ══════════════════════════════════════════════════════════════════════════════
# The Claude Projects platform applies a per-format transform to files uploaded to
# project knowledge and PRESERVES THE ORIGINAL FILENAME AND EXTENSION. A file named
# .docx in /mnt/project/ is therefore NOT an OOXML package: the platform stores
# EXTRACTED MARKDOWN TEXT. Confirmed independently by the .pdf in the same project,
# which is stored as a Zip page-bundle under its .pdf name; .xlsx and .json are
# byte-preserved, and so is anything attached to CHAT rather than uploaded to Files.
#
# The Analysis doc is uploaded to project Files after Step 2c and read from there at
# Steps 3, 4, 5 and 6, so INGEST_TEXT is the PRIMARY RUNTIME FORM for this artefact —
# not a fallback, not a degraded mode, and nothing to warn about. INGEST_OOXML is
# retained because write_analysis_doc() emits it and the round-trip self-test reads
# back what it writes.
#
# Consequence: the container MUST be detected from CONTENT. Inference from the
# extension is wrong on exactly the input the runtime actually receives.

INGEST_OOXML   = 'ooxml'      # genuine .docx package (writer output, self-test, chat)
INGEST_TEXT    = 'text'       # platform-extracted text (project Files) — PRIMARY
INGEST_UNKNOWN = 'unknown'    # neither — hard stop, never guess


def _ingest_form(path):
    """Classify the container of the Analysis doc by CONTENT.

    Returns INGEST_OOXML / INGEST_TEXT / INGEST_UNKNOWN. Never raises: the caller
    turns UNKNOWN into a named hard stop carrying an operator action.

    A truncated .docx and a non-Word Zip both fail the word/document.xml test. Both
    are genuinely unreadable, so both become UNKNOWN and receive a named stop instead
    of a library traceback — the same discipline as PYQSort S7-5's pre-flight, which
    inspects the package on the path before the library is allowed to raise.
    """
    import zipfile
    try:
        if os.path.getsize(path) == 0:
            return INGEST_UNKNOWN
        with open(path, 'rb') as fh:
            magic = fh.read(4)
        if magic[:2] == b'PK':
            try:
                with zipfile.ZipFile(path) as z:
                    names = set(z.namelist())
            except Exception:
                return INGEST_UNKNOWN
            return INGEST_OOXML if 'word/document.xml' in names else INGEST_UNKNOWN
        with open(path, 'rb') as fh:
            raw = fh.read()
        try:
            raw.decode('utf-8-sig')
        except UnicodeDecodeError:
            return INGEST_UNKNOWN
        return INGEST_TEXT
    except OSError:
        return INGEST_UNKNOWN


# ── markdown surface -> plain text ────────────────────────────────────────────
# Bounded tolerance: every variant handled here has been OBSERVED in a real
# docx->markdown extractor run against a real Analysis doc. Nothing is anticipated.
_MD_STAR  = re.compile(r'\*\*')                  # ** bold, incl. run-split ****
_MD_WRAP  = re.compile(r'^__(.+)__$', re.S)      # __bold__ as a WHOLE-LINE wrap only
_MD_SEP   = re.compile(r'^[\s|:\-]+$')           # markdown table separator row
_MD_ESC   = re.compile(r'\\([\\`*_{}\[\]()#+\-.!|~>])')
_MD_PIPE  = re.compile(r'(?<!\\)\|')             # cell split on UNESCAPED pipes only


def _md_plain(s):
    """Markdown inline markup -> the plain text the document author wrote.

    __ is unwrapped ONLY as a whole-line wrap, never stripped globally: a taxonomy
    name may legitimately contain underscores ('p_1 and p_2'), and a global strip
    would silently rewrite it. Escapes are removed only before markdown punctuation,
    so a backslash that is part of a name survives unless it precedes one of those.
    """
    s = _MD_STAR.sub('', s).strip()
    m = _MD_WRAP.match(s)
    if m:
        s = m.group(1).strip()
    return _MD_ESC.sub(r'\1', s).strip()


def _md_cells(line):
    """Cells of one pipe-table row. Splits on UNESCAPED pipes, then unescapes each
    cell — so an extractor that escapes '\\|' inside a name keeps the name intact."""
    s = line.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|') and not s.endswith('\\|'):
        s = s[:-1]
    return [_md_plain(c) for c in _MD_PIPE.split(s)]


def _assert_rectangular(rows, doc_name, line_no):
    """A markdown table is RECTANGULAR. HARD STOP when it is not.

    GAP-2026-07-25-003. A '|' inside a taxonomy name splits one row wider than its
    neighbours, and the surplus cell is swallowed as the count column: the name is
    truncated, the remainder is discarded, and the SUBTOPIC COUNT IS UNCHANGED — so
    D1, D2 and D3 all still agree and verify_analysis_doc() returns clean. Nothing
    downstream catches it either: the fingerprint cross-check against the approval
    record lives only in Framework_PYQSort S1-0b, so at Steps 5 and 6 there is no gate
    at all. Reproduced by execution, which is why this is a stop and not an anomaly.
    """
    widths = {len(r) for r in rows if r}
    if len(widths) <= 1:
        return
    bad = next(r for r in rows if len(r) == max(widths))
    raise AnalysisDocError(
        "HARD STOP: %s contains a table whose rows do not all have the same number "
        "of cells (widths %s), near line %d.\n"
        "  offending row: %s\n"
        "Most common cause: a subject, topic or subtopic name contains a '|' "
        "character. In the extracted-text ingest form a '|' is a cell separator, so "
        "the name is split and everything after the '|' is silently discarded while "
        "the declared counts still agree — which is why this is stopped here rather "
        "than left to a later check that cannot see it.\n"
        "NEXT ACTION: remove '|' from the taxonomy name(s) and re-run PYQApprove "
        "(Step 2c) to regenerate the Analysis doc."
        % (doc_name, sorted(widths), line_no, ' | '.join(bad[:6])))


def _scan_ooxml(doc_path):
    """Single pass over the body. Returns (subjects, order, declared, anomalies).

    ONE pass over doc.element.body, never doc.paragraphs, so paragraphs and tables
    are seen in true document order. The two-pass form the pre-fix readers used
    cannot attribute a table to a subject at all: by the time pass 2 runs, the
    subject boundaries from pass 1 are gone.
    """
    from docx import Document
    doc = Document(doc_path)

    subjects, order, declared, anomalies = {}, [], [], []
    cur_s = cur_t = None
    blk = None

    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            text = _para_text(child)
            if not text:
                continue
            level, name = bc.parse_taxonomy_level(text)

            if level == 1:
                cur_s, cur_t = name, None
                if cur_s not in subjects:
                    subjects[cur_s] = {}
                    order.append(cur_s)
                blk = {'subject': cur_s, 'd1': None, 'd2': [], 'd2_grand': None,
                       'd3': [], 'master_seen': False}
                declared.append(blk)
                continue

            if level == 2:
                if cur_s is None:
                    anomalies.append("topic heading %r appears before any subject "
                                     "heading" % text[:60])
                    continue
                cur_t = name
                subjects[cur_s].setdefault(cur_t, {})
                continue

            # level 3 => a decorative paragraph of the header/footer block.
            # Subtopics live in TABLES, never in paragraphs, so a level-3 paragraph
            # is never data. It may still carry a DECLARATION.
            if blk is None:
                continue
            m = _D1_RE.search(text)
            if m and blk['d1'] is None and not blk['master_seen']:
                blk['d1'] = (int(m.group(1)), int(m.group(2)))
                continue
            m = _D3_RE.search(text)
            if m and cur_t is not None:
                blk['d3'].append((cur_t, int(m.group(1))))
            continue

        if tag != 'tbl':
            continue

        rows = _table_rows(child, doc)
        data = [r for r in rows if r and r[0].strip()]
        is_master = bool(data) and any(
            bc.parse_taxonomy_level(r[0])[0] == 2 for r in data)

        if is_master:
            # MASTER SUMMARY TABLE — a declaration, never data.
            if blk is None:
                anomalies.append("a master summary table appears before any subject heading")
                continue
            blk['master_seen'] = True
            for r in data:
                lvl, tname = bc.parse_taxonomy_level(r[0])
                if lvl == 2 and len(r) >= 2 and r[1].strip().isdigit():
                    blk['d2'].append((tname, int(r[1].strip())))
                elif _GRAND_RE.match(r[0].strip()) and len(r) >= 2 and r[1].strip().isdigit():
                    blk['d2_grand'] = int(r[1].strip())
            continue

        if cur_s is None or cur_t is None:
            anomalies.append("a subtopic table at document position %d belongs to no "
                             "topic — the heading above it was not recognised"
                             % list(doc.element.body).index(child))
            continue

        bucket = subjects[cur_s][cur_t]
        for r in rows:
            if len(r) < 2:
                continue
            name = r[0].strip()
            if name.lower() in _SKIP_CELLS:
                continue
            if bc.parse_taxonomy_level(name)[0] == 2:
                continue                      # a stray topic row inside a data table
            if name not in bucket:
                bucket[name] = len(bucket)

    return subjects, order, declared, anomalies


def _scan_text(doc_path):
    """Scan the PLATFORM-EXTRACTED TEXT form — the primary runtime ingest form.

    Emits the IDENTICAL 4-tuple as _scan_ooxml(), so verify_analysis_doc(),
    taxonomy_fingerprint(), read_analysis_doc()'s return shape and every downstream
    consumer are untouched. Heading recognition is delegated to
    blueprint_core.parse_taxonomy_level(); the declaration regexes and the skip-cell
    set are the SAME module-level objects the OOXML scanner uses. There is no second
    structural convention anywhere in this module.
    """
    doc_name = os.path.basename(doc_path)
    with open(doc_path, encoding='utf-8-sig') as fh:
        text = fh.read()
    lines = text.replace('\r\n', '\n').replace('\r', '\n').split('\n')

    subjects, order, declared, anomalies = {}, [], [], []
    cur_s = cur_t = None
    blk = None
    i = 0
    while i < len(lines):
        raw = lines[i]

        if raw.lstrip().startswith('|'):                       # ── table block ──
            start = i + 1
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                if not _MD_SEP.match(lines[i]):
                    rows.append(_md_cells(lines[i]))
                i += 1
            _assert_rectangular(rows, doc_name, start)
            data = [r for r in rows if r and r[0].strip()]
            is_master = bool(data) and any(
                bc.parse_taxonomy_level(r[0])[0] == 2 for r in data)

            if is_master:
                # MASTER SUMMARY TABLE — a declaration, never data.
                if blk is None:
                    anomalies.append("a master summary table appears before any "
                                     "subject heading")
                    continue
                blk['master_seen'] = True
                for r in data:
                    lvl, tname = bc.parse_taxonomy_level(r[0])
                    if lvl == 2 and len(r) >= 2 and r[1].strip().isdigit():
                        blk['d2'].append((tname, int(r[1].strip())))
                    elif _GRAND_RE.match(r[0].strip()) and len(r) >= 2 \
                            and r[1].strip().isdigit():
                        blk['d2_grand'] = int(r[1].strip())
                continue

            if cur_s is None or cur_t is None:
                anomalies.append("a subtopic table at line %d belongs to no topic — "
                                 "the heading above it was not recognised" % start)
                continue

            bucket = subjects[cur_s][cur_t]
            for r in rows:
                if len(r) < 2:
                    continue
                name = r[0].strip()
                if name.lower() in _SKIP_CELLS:
                    continue
                if bc.parse_taxonomy_level(name)[0] == 2:
                    continue                      # a stray topic row inside a data table
                if name not in bucket:
                    bucket[name] = len(bucket)
            continue

        text_line = _md_plain(raw)                             # ── paragraph ──
        i += 1
        if not text_line:
            continue
        level, name = bc.parse_taxonomy_level(text_line)

        if level == 1:
            cur_s, cur_t = name, None
            if cur_s not in subjects:
                subjects[cur_s] = {}
                order.append(cur_s)
            blk = {'subject': cur_s, 'd1': None, 'd2': [], 'd2_grand': None,
                   'd3': [], 'master_seen': False}
            declared.append(blk)
            continue

        if level == 2:
            if cur_s is None:
                anomalies.append("topic heading %r appears before any subject "
                                 "heading" % text_line[:60])
                continue
            cur_t = name
            subjects[cur_s].setdefault(cur_t, {})
            continue

        # level 3 => decorative paragraph; never data, but may carry a DECLARATION.
        if blk is None:
            continue
        m = _D1_RE.search(text_line)
        if m and blk['d1'] is None and not blk['master_seen']:
            blk['d1'] = (int(m.group(1)), int(m.group(2)))
            continue
        m = _D3_RE.search(text_line)
        if m and cur_t is not None:
            blk['d3'].append((cur_t, int(m.group(1))))

    return subjects, order, declared, anomalies


def verify_analysis_doc(subjects, order, declared, anomalies, doc_name='Analysis doc'):
    """Assert the PARSE against the document's OWN declarations. Returns a list of
    plain-English failures; empty means the doc and the parse agree.

    The doc states its shape three times per subject:
      D1  header  "Total: — Questions | <T> Topics | <S> Subtopics"
      D2  master summary table: one row per topic (+ GRAND TOTAL)
      D3  per topic "Total PYQs: — | Subtopics: <k>"
    Nothing but a correct parse satisfies all three. This is what makes a
    mis-parse LOUD without any second artefact, on any exam, from day one.
    """
    fails = list(anomalies)
    decl_names = [b['subject'] for b in declared]

    dupes = sorted({s for s in decl_names if decl_names.count(s) > 1})
    if dupes:
        fails.append("the document declares the same subject more than once: %s — "
                     "each subject must appear exactly once" % dupes)
    if len(order) != len(declared):
        fails.append("SUBJECT COUNT: parsed %d, the document declares %d subject "
                     "heading(s) -> %s" % (len(order), len(declared), decl_names))

    for b in declared:
        s = b['subject']
        if s not in subjects:
            fails.append("[%s] declared in the document but absent from the parse" % s)
            continue
        topics = subjects[s]
        n_t = len(topics)
        n_s = sum(len(v) for v in topics.values())

        if b['d1'] and b['d1'] != (n_t, n_s):
            fails.append("[%s] header declares %d topics / %d subtopics; parsed %d / %d"
                         % (s, b['d1'][0], b['d1'][1], n_t, n_s))
        if b['d2']:
            if len(b['d2']) != n_t:
                fails.append("[%s] master summary table lists %d topics; parsed %d"
                             % (s, len(b['d2']), n_t))
            for tname, tcount in b['d2']:
                if tname not in topics:
                    fails.append("[%s] master summary lists topic %r, which is absent "
                                 "from the parse" % (s, tname))
                elif len(topics[tname]) != tcount:
                    fails.append("[%s] master summary declares %d subtopics under %r; "
                                 "parsed %d" % (s, tcount, tname, len(topics[tname])))
        if b['d2_grand'] is not None and b['d2_grand'] != n_s:
            fails.append("[%s] master summary GRAND TOTAL is %d; parsed %d"
                         % (s, b['d2_grand'], n_s))
        for tname, k in b['d3']:
            if tname in topics and len(topics[tname]) != k:
                fails.append("[%s] topic %r declares Subtopics: %d; parsed %d"
                             % (s, tname, k, len(topics[tname])))
        if not topics:
            fails.append("[%s] has no topics — an empty subject cannot be sorted "
                         "against" % s)
        for tname, subs in topics.items():
            if not subs:
                fails.append("[%s] topic %r has no subtopics" % (s, tname))

    for s in order:
        if s not in decl_names:
            fails.append("[%s] parsed but the document contains no matching subject "
                         "heading for it" % s)

    # Every name must survive the round trip through the sorted-PYQ heading parser.
    # bc.is_taxonomy_heading() rejects text at or above MAX_HEADING_LEN, so a longer
    # name stops being a heading at Step 4/5 and its questions are silently absorbed
    # by the PREVIOUS subtopic. Caught HERE, at the artefact, not four steps later.
    for s, topics in subjects.items():
        for label, value in [('subject', s)] + \
                [('topic', t) for t in topics] + \
                [('subtopic', x) for t in topics for x in topics[t]]:
            if len(value) >= bc.MAX_HEADING_LEN:
                fails.append("%s name is %d characters, at or above the %d-character "
                             "heading limit: %r. A name this long stops being "
                             "recognised as a heading in the sorted files and its "
                             "questions are silently attributed to the preceding "
                             "subtopic. Shorten it in the taxonomy and re-run "
                             "PYQApprove." % (label, len(value), bc.MAX_HEADING_LEN, value))
    return fails


def _assemble(order, subjects, path, exam_code, ingest_form, source):
    """Build THE return shape from (order, {subject: {topic: {name: idx}}}).

    v1.4 — extracted so the Analysis-doc path and the approval-record path cannot
    drift. Two constructions of one shape is how GAP-2026-07-25-002's four readers
    happened; there is exactly one here, and both callers go through it.

    topic_idx is POSITIONAL WITHIN THE SUBJECT — never derived from a printed
    "Topic N:" label, which restarts at 1 for every subject in a merged doc.
    """
    taxonomy, triples = {}, []
    for subject in order:
        topics_out = {}
        for t_i, (topic, subs) in enumerate(subjects[subject].items()):
            topics_out[topic] = {
                'topic_idx': t_i,
                'subtopics': {n: {'subtopic_idx': i} for n, i in subs.items()}}
            for n in subs:
                triples.append((subject, topic, n))
        taxonomy[subject] = {'subject_order': len(taxonomy), 'topics': topics_out}
    return {
        'path': path,
        'ingest_form': ingest_form,
        'source': source,           # 'analysis_doc' | 'approval_record'
        'exam_code': exam_code,
        'subjects': list(order),
        'taxonomy': taxonomy,
        'triples': triples,
        'counts': {'subjects': len(order),
                   'topics': sum(len(v['topics']) for v in taxonomy.values()),
                   'subtopics': len(triples)},
        'fingerprint': bc.taxonomy_fingerprint(triples),
    }


def load_taxonomy(exam_code=None, search_dirs=ANALYSIS_SEARCH_DIRS, step=None,
                  record=None, path=None):
    """THE taxonomy entry point for Steps 3-6. Returns the shape read_analysis_doc()
    returns, plus 'source'.

    SOURCE PRIORITY
      1. the approval record's "taxonomy" key   (schema >= 1.3)  -> source='approval_record'
      2. the Analysis doc                       (older records)  -> source='analysis_doc'

    WHY THE RECORD COMES FIRST. Every step from PYQSort onward needs a taxonomy, and
    until reconcile_taxonomy v1.3 the only place to get one was by PARSING A WORD
    DOCUMENT. That is what GAP-2026-07-25-003 severed on every exam at once: the
    platform stores an uploaded .docx in project knowledge as extracted Markdown
    text under its .docx name, and Cluster K opened it with python-docx. v1.2 taught
    the reader both ingest forms, which works — but it works by matching an
    undocumented, unversioned grammar the framework does not control, so the next
    change to that grammar stops all ~200 exams simultaneously, loudly.

    The record does not have that property. It is JSON, which the platform stores
    BYTE-FOR-BYTE, and it already carries the fingerprint that validates its own
    taxonomy, so nothing new is trusted. Reading it is not a workaround for the
    document; it is the removal of a dependency that should never have existed.
    The Analysis doc reverts to what it always should have been: a HUMAN deliverable.

    The fallback is not a courtesy. Every exam approved before schema 1.3 has a
    record with a fingerprint and no taxonomy, and those projects must keep working
    without being re-approved. They take path 2, which is fully gated by
    assert_taxonomy_lock() exactly as before.

    DISPLAY NAMES. The fingerprint is slug-normalised BY DESIGN, so it proves
    identity and cannot detect a display-byte change. On path 1 the names come
    straight from JSON and no such change is possible. On path 2 they survive a text
    extraction, and PYQAnalyse Task 2.5 remains the byte-identity check.
    """
    if record is None:
        record = read_approval_record(exam_code, search_dirs)
    rec_name = os.path.basename(record.get('_path', 'approval record'))
    where = (' at %s' % step) if step else ''
    tax = record.get('taxonomy')

    if tax is None:
        # ── path 2: pre-1.3 record. The doc is the only carrier of the content. ──
        doc = read_analysis_doc(path, search_dirs)
        assert_taxonomy_lock(doc, record=record, step=step)
        return doc

    # ── path 1: the record carries the taxonomy ──────────────────────────────────
    sections = tax.get('sections') if isinstance(tax, dict) else None
    if not isinstance(sections, dict) or not sections:
        raise AnalysisDocError(
            "HARD STOP%s: %s carries a 'taxonomy' key that is not a populated "
            "{'sections': {subject: {topic: [subtopic]}}} object.\n"
            "A record that carries a BROKEN taxonomy is worse than one that carries "
            "none: the absent case has a defined fallback and this one does not.\n"
            "NEXT ACTION: re-run PYQApprove with reconcile_taxonomy.py >= v1.3. That "
            "is RECONCILIATION, not re-derivation — it cannot change a locked "
            "taxonomy. Do NOT re-run PYQDraft." % (where, rec_name))

    order, subjects, dupes = [], {}, []
    for subject, topics in sections.items():
        if not isinstance(topics, dict) or not topics:
            raise AnalysisDocError(
                "HARD STOP%s: %s declares subject %r with no topics. An empty "
                "subject cannot be sorted or allocated against.\n"
                "NEXT ACTION: re-run PYQApprove." % (where, rec_name, subject))
        if subject in subjects:
            dupes.append('subject %r' % subject)
            continue
        subjects[subject] = {}
        order.append(subject)
        for topic, subs in topics.items():
            if not isinstance(subs, list) or not subs:
                raise AnalysisDocError(
                    "HARD STOP%s: %s declares topic %r under %r with no subtopics.\n"
                    "NEXT ACTION: re-run PYQApprove."
                    % (where, rec_name, topic, subject))
            bucket = subjects[subject].setdefault(topic, {})
            for name in subs:
                if name in bucket:
                    dupes.append('%r under %r / %r' % (name, subject, topic))
                    continue
                bucket[name] = len(bucket)

    if dupes:
        # Silent de-duplication would make the assembled taxonomy disagree with the
        # counts the record declares, and the fingerprint would NOT catch it — it is
        # computed over the same duplicated triples, so it agrees with itself.
        raise AnalysisDocError(
            "HARD STOP%s: %s declares the same name more than once:\n  %s\n"
            "NEXT ACTION: re-run PYQApprove."
            % (where, rec_name, '\n  '.join(dupes[:8])))

    out = _assemble(order, subjects, path=record.get('_path'),
                    exam_code=record.get('exam_code', exam_code),
                    ingest_form='json', source='approval_record')

    # SELF-VALIDATION. The record proves its own content: the fingerprint recorded
    # beside the taxonomy was computed FROM it at Step 2c. A disagreement means the
    # record was hand-edited or written by two different engines, and the taxonomy
    # half cannot be trusted just because the hash half looks well-formed.
    fp_locked = record.get('taxonomy_fingerprint')
    if not fp_locked:
        raise AnalysisDocError(
            "HARD STOP%s: %s carries a taxonomy but no taxonomy_fingerprint, so the "
            "taxonomy cannot validate itself.\n"
            "NEXT ACTION: re-run PYQApprove with reconcile_taxonomy.py >= v1.3."
            % (where, rec_name))
    if fp_locked != out['fingerprint']:
        raise AnalysisDocError(
            "HARD STOP%s: %s does not agree with itself — the fingerprint it records "
            "is not the fingerprint of the taxonomy it records.\n"
            "  recorded fingerprint : %s\n"
            "  taxonomy computes to : %s\n"
            "  counts declared      : %s\n"
            "  counts assembled     : %s\n"
            "The record was edited by hand, or written by two engines that disagree. "
            "Neither half can be trusted.\n"
            "NEXT ACTION: re-run PYQApprove. RECONCILIATION only — it cannot change "
            "a locked taxonomy."
            % (where, rec_name, fp_locked, out['fingerprint'],
               record.get('taxonomy_counts'), out['counts']))

    declared = record.get('taxonomy_counts')
    if declared and declared != out['counts']:
        raise AnalysisDocError(
            "HARD STOP%s: %s declares taxonomy_counts %s but the taxonomy it carries "
            "assembles to %s.\n"
            "NEXT ACTION: re-run PYQApprove."
            % (where, rec_name, declared, out['counts']))
    return out


def read_analysis_doc(path=None, search_dirs=ANALYSIS_SEARCH_DIRS, verify=True):
    """THE Analysis-doc reader. Returns a dict:

        {'path','exam_code','subjects','taxonomy','triples','counts','fingerprint'}

    taxonomy/triples carry exactly the shape PYQSort S1-2 has always returned, so
    downstream consumers are unchanged:
        {subject: {'subject_order': int,
                   'topics': {topic: {'topic_idx': int,
                                      'subtopics': {name: {'subtopic_idx': int}}}}}}

    topic_idx is POSITIONAL WITHIN THE SUBJECT, which is what Framework_PYQSort
    §L1111 has always specified ("position of topic within its section's Analysis
    doc"). The pre-fix code derived it from the printed "Topic N:" label, which
    restarts at 1 for every subject in a merged doc — six topics claiming index 0.
    """
    if path is None:
        path = discover_analysis_doc(search_dirs)
    base = os.path.basename(path)

    form = _ingest_form(path)

    if form == INGEST_UNKNOWN:
        # The size is DIAGNOSTIC, and formatting a diagnostic must never raise the
        # fault it is diagnosing. _ingest_form() returns UNKNOWN for an unreadable
        # path precisely BECAUSE os.path.getsize() raised there, so calling it again
        # here re-raises it uncaught and the operator gets the library traceback this
        # whole cluster exists to replace. Found by re-running the delivered spec
        # against a path that had gone away mid-session.
        try:
            size_txt = '%d bytes' % os.path.getsize(path)
        except OSError as _ex:
            size_txt = 'unreadable (%s: %s)' % (type(_ex).__name__, _ex)
        raise AnalysisDocError(
            "HARD STOP: %s is neither a readable .docx package nor readable "
            "extracted text.\n"
            "  size: %s\n"
            "This is not a recognised ingest form for the Analysis doc. Likely "
            "causes: a truncated or failed upload, a file saved under the wrong name, "
            "or a platform transform this framework version does not know about.\n"
            "NEXT ACTION: re-upload the Analysis doc that PYQApprove (Step 2c) "
            "delivered to the project's Files section. If it opens correctly in Word, "
            "report this as a NEW ingest form (GAP-2026-07-25-003 follow-up) — do NOT "
            "work around it." % (base, size_txt))

    if form == INGEST_TEXT and not verify:
        # Admitted THROUGH the invariants, never around them. The text form carries no
        # container-level integrity of its own, so the three self-declarations are the
        # only thing standing between a mis-parse and a silently mis-sorted corpus.
        verify = True

    try:
        if form == INGEST_OOXML:
            subjects, order, declared, anomalies = _scan_ooxml(path)
        else:
            subjects, order, declared, anomalies = _scan_text(path)
    except AnalysisDocError:
        raise
    except Exception as ex:
        raise AnalysisDocError(
            "HARD STOP: %s could not be read as %s (%s: %s).\n"
            "NEXT ACTION: re-upload the Analysis doc PYQApprove (Step 2c) delivered "
            "to the project's Files section."
            % (base, 'a .docx package' if form == INGEST_OOXML
               else 'platform-extracted text', type(ex).__name__, ex))

    if not order:
        raise AnalysisDocError(
            "HARD STOP: %s contains no subject heading.\n"
            "  ingest form: %s\n"
            "Expected a paragraph of the form 'Subject: <name>' (or Domain:/Section:/"
            "Part:/Area:) for each subject. Either this is not a PYQApprove Analysis "
            "doc, or the PYQAnalyse §6 HEADING FORMAT CONTRACT was violated%s."
            % (base, form,
               ", or the platform's text-extraction grammar has changed and "
               "_scan_text() must be updated — do NOT work around it, report it as a "
               "GAP-2026-07-25-003 follow-up" if form == INGEST_TEXT else ""))

    if verify:
        fails = verify_analysis_doc(subjects, order, declared, anomalies, base)
        if fails:
            raise AnalysisDocError(
                "HARD STOP: %s does not agree with itself — the taxonomy parsed out of "
                "it does not match the totals the document declares. Sorting against a "
                "taxonomy that fails this check silently mis-files questions.\n  %s\n"
                "Re-run PYQApprove to regenerate the Analysis doc."
                % (base, '\n  '.join(fails[:12]) +
                   ('\n  ... and %d more' % (len(fails) - 12) if len(fails) > 12 else '')))

    exam_code = base[:-len(ANALYSIS_DOC_SUFFIX)] if base.endswith(ANALYSIS_DOC_SUFFIX) else None
    return _assemble(order, subjects, path=path, exam_code=exam_code, ingest_form=form,
                     source='analysis_doc')


# ══════════════════════════════════════════════════════════════════════════════
# THE TAXONOMY LOCK GATE — one implementation, every step
# ══════════════════════════════════════════════════════════════════════════════
# read_analysis_doc() asserts the document agrees with ITSELF. It cannot assert
# that the document is the one PYQApprove LOCKED. Those are different claims, and
# GAP-2026-07-25-002 Defect B satisfied the first while violating the second.
#
# Framework_PYQSort S1-0b made the second claim from v1.14 — and NOTHING else did.
# Steps 5 and 6 read the taxonomy and allocated against it with no identity check at
# all, so a superseded or mis-parsed Analysis doc that halted loudly at Step 3 was
# accepted silently three steps later. Worse at Step 5, where the reader's hard stops
# were caught and downgraded to a log line twice over: once in the extractor and
# again in its caller.
#
# The gate therefore lives HERE, once, rather than being written a fifth time in a
# fifth spec. This is the same anti-drift rule that produced Cluster K: one reader,
# one writer, one verifier — and now one lock gate.
#
# SCOPE, deliberately narrow: this asserts IDENTITY only. Whether the lock was
# EARNED — status, schema attestation, checks.missing / checks.vacuous — is
# Framework_PYQSort S1-0's claim, made once at Step 3 against the same record, and
# it is not duplicated here. One function, one claim.

APPROVAL_RECORD_SUFFIX = '_approval_record.json'


def discover_approval_record(exam_code=None, search_dirs=ANALYSIS_SEARCH_DIRS):
    """Locate the ONE approval record. Returns a path; raises AnalysisDocError.

    Same de-duplication rule as discover_analysis_doc(): a name seen in an earlier
    search dir wins, so the project's copy beats a chat attachment of the same name.
    """
    want = ('%s%s' % (exam_code, APPROVAL_RECORD_SUFFIX)).lower() if exam_code else None
    seen, found = set(), []
    for d in search_dirs:
        if not os.path.isdir(d):
            continue
        for name in sorted(os.listdir(d)):
            if name.startswith('~$') or name.startswith('.'):
                continue
            low = name.lower()
            if not low.endswith(APPROVAL_RECORD_SUFFIX.lower()):
                continue
            if want and low != want:
                continue
            if name in seen:
                continue
            seen.add(name)
            found.append(os.path.join(d, name))
    if len(found) > 1:
        raise AnalysisDocError(
            "HARD STOP: %d approval records found — exactly one is expected.\n"
            "  %s\n"
            "One project holds one exam. Remove the ones that do not belong to this "
            "exam and re-run."
            % (len(found), '\n  '.join(os.path.basename(p) for p in found)))
    if found:
        return found[0]
    raise AnalysisDocError(
        "HARD STOP: no approval record found in %s.\n"
        "Expected %s from PYQApprove (Step 2c).\n"
        "The Analysis doc shows what the taxonomy IS; only the approval record shows "
        "that it was the one APPROVED. Allocating against an unverified taxonomy is "
        "how a mis-parse travels to delivery.\n"
        "NEXT ACTION: upload the approval record PYQApprove delivered to this "
        "project's Files section."
        % (' or '.join(search_dirs),
           ('%s%s' % (exam_code, APPROVAL_RECORD_SUFFIX)) if exam_code
           else '[ExamCode]' + APPROVAL_RECORD_SUFFIX))


def read_approval_record(exam_code=None, search_dirs=ANALYSIS_SEARCH_DIRS, path=None):
    """Load the approval record as a dict. Raises AnalysisDocError on any fault."""
    if path is None:
        path = discover_approval_record(exam_code, search_dirs)
    try:
        with open(path, encoding='utf-8-sig') as fh:
            rec = json.load(fh)
    except Exception as ex:
        raise AnalysisDocError(
            "HARD STOP: %s could not be read as JSON (%s: %s).\n"
            "NEXT ACTION: re-upload the approval record PYQApprove (Step 2c) "
            "delivered." % (os.path.basename(path), type(ex).__name__, ex))
    if not isinstance(rec, dict):
        raise AnalysisDocError(
            "HARD STOP: %s does not contain a JSON object.\n"
            "NEXT ACTION: re-run PYQApprove." % os.path.basename(path))
    rec['_path'] = path
    return rec


def assert_taxonomy_lock(doc, record=None, search_dirs=ANALYSIS_SEARCH_DIRS, step=None):
    """THE identity gate: is this the taxonomy that was APPROVED?

    `doc`    : the dict read_analysis_doc() returned.
    `record` : an already-loaded approval record, or None to discover one using
               doc['exam_code'].
    `step`   : caller name for the message ('PYQExtract', 'MockBlueprint', ...).

    Returns the record on success. Raises AnalysisDocError otherwise — callers must
    NOT downgrade this to a warning. A taxonomy that fails this check is not a
    degraded input; it is the wrong input, and every id minted from it is wrong.
    """
    where = (' at %s' % step) if step else ''
    if record is None:
        record = read_approval_record(doc.get('exam_code'), search_dirs)

    fp_locked = record.get('taxonomy_fingerprint')
    if not fp_locked:
        raise AnalysisDocError(
            "HARD STOP%s: the approval record carries no taxonomy_fingerprint, so "
            "the Analysis doc cannot be shown to be the approved one.\n"
            "  record: %s\n"
            "NEXT ACTION: re-run PYQApprove with reconcile_taxonomy.py >= v1.2. That "
            "is RECONCILIATION, not re-derivation — it reads the draft and rewrites "
            "only the record, and CANNOT change a locked taxonomy. Do NOT re-run "
            "PYQDraft." % (where, os.path.basename(record.get('_path', '?'))))

    if fp_locked != doc['fingerprint']:
        raise AnalysisDocError(
            "HARD STOP%s: the Analysis doc is NOT the taxonomy that was approved.\n"
            "  doc    : %s\n"
            "  locked : %s\n"
            "  loaded : %s\n"
            "  counts : %s\n"
            "The fingerprint is taken over slugified triples, so it is already "
            "invariant to case, spacing and dash variants — this is a real "
            "difference in the taxonomy itself, not a cosmetic one.\n"
            "NEXT ACTION: restore the Analysis doc PYQApprove (Step 2c) delivered, "
            "or re-run PYQApprove if the taxonomy genuinely changed. If it changed, "
            "every step from PYQSort onward must be re-run — ids minted from one "
            "taxonomy do not resolve against another."
            % (where, os.path.basename(doc.get('path', '?')), fp_locked,
               doc['fingerprint'], doc.get('counts')))
    return record


def write_analysis_doc(taxonomy, exam_code, subject_order=None, out_dir='/mnt/user-data/outputs',
                       counts=None):
    """THE Analysis-doc writer (PYQAnalyse S4-2, which was `pass` until v2.24).

    taxonomy : {subject: {topic: [subtopic, ...]}}   — insertion order is preserved
    counts   : optional {(subject, topic, subtopic): int} for Phase B (--counts).
               Absent => every PYQ Count cell is an em-dash, per S4-1.

    Emits exactly the structure read_analysis_doc() contracts on, including all
    three redundant declarations. One writer + one reader + a round-trip assertion
    is the whole correctness argument; a doc this function produces is a doc that
    function reads, by construction rather than by review.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_BREAK

    order = list(subject_order or taxonomy.keys())
    missing = [s for s in taxonomy if s not in order]
    order += missing

    over = [(lbl, v) for s in taxonomy
            for lbl, v in ([('subject', s)] +
                           [('topic', t) for t in taxonomy[s]] +
                           [('subtopic', x) for t in taxonomy[s] for x in taxonomy[s][t]])
            if len(v) >= bc.MAX_HEADING_LEN]
    if over:
        raise AnalysisDocError(
            "HARD STOP: %d taxonomy name(s) are at or above the %d-character heading "
            "limit and cannot be written — they would stop being recognised as "
            "headings in the sorted files.\n  %s"
            % (len(over), bc.MAX_HEADING_LEN,
               '\n  '.join('%s (%d chars): %r' % (l, len(v), v) for l, v in over[:5])))

    # GAP-2026-07-25-003. '|' is a CELL SEPARATOR in the extracted-text ingest form,
    # which is the form every runtime read of this document receives. A name carrying
    # one is split on read, the remainder is discarded, and the declared counts still
    # agree — so the corruption is invisible to all three self-declarations. Rejected
    # where it ORIGINATES rather than where it corrupts: this document is unwritable,
    # not merely unreadable.
    piped = [(lbl, v) for s in taxonomy
             for lbl, v in ([('subject', s)] +
                            [('topic', t) for t in taxonomy[s]] +
                            [('subtopic', x) for t in taxonomy[s] for x in taxonomy[s][t]])
             if '|' in v]
    if piped:
        raise AnalysisDocError(
            "HARD STOP: %d taxonomy name(s) contain a '|' character, which cannot be "
            "written to the Analysis doc.\n  %s\n"
            "'|' is the cell separator in the form the platform stores this document "
            "in, so such a name is silently truncated on every downstream read while "
            "the declared totals continue to agree.\n"
            "NEXT ACTION: rename the taxonomy entries (an en-dash, a slash or a comma "
            "all read naturally) and re-run."
            % (len(piped), '\n  '.join('%s: %r' % (l, v) for l, v in piped[:5])))

    doc = Document()

    def bold(text, size):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        return p

    def cell(v):
        return '—' if v is None else str(v)

    for i, subject in enumerate(order):
        topics = taxonomy[subject]
        if i:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        n_sub = sum(len(v) for v in topics.values())
        bold('%s — %s' % (exam_code, subject), 14)
        bold('Subject: %s' % subject, 12)                      # LEVEL 1 — §6 contract
        bold('PYQ Topic & Subtopic-wise Count', 11)
        bold('Total: %s Questions  |  %d Topics  |  %d Subtopics'      # D1
             % (cell(sum(counts.get((subject, t, x), 0)
                         for t in topics for x in topics[t]) if counts else None),
                len(topics), n_sub), 11)

        tbl = doc.add_table(rows=1, cols=3)                     # D2
        h = tbl.rows[0].cells
        h[0].text, h[1].text, h[2].text = 'Topic', 'Total Subtopics', 'Total PYQs'
        for t_i, (topic, subs) in enumerate(topics.items(), 1):
            c = tbl.add_row().cells
            c[0].text = 'Topic %d: %s' % (t_i, topic)
            c[1].text = str(len(subs))
            c[2].text = cell(sum(counts.get((subject, topic, x), 0) for x in subs)
                             if counts else None)
        c = tbl.add_row().cells
        c[0].text = 'GRAND TOTAL'
        c[1].text = str(n_sub)
        c[2].text = cell(sum(counts.get((subject, t, x), 0)
                             for t in topics for x in topics[t]) if counts else None)

        for t_i, (topic, subs) in enumerate(topics.items(), 1):
            bold('Topic %d: %s' % (t_i, topic), 12)             # LEVEL 2 — §6 contract
            bold('Total PYQs: %s  |  Subtopics: %d'             # D3
                 % (cell(sum(counts.get((subject, topic, x), 0) for x in subs)
                         if counts else None), len(subs)), 11)
            st = doc.add_table(rows=1, cols=2)
            hh = st.rows[0].cells
            hh[0].text, hh[1].text = 'Subtopic', 'PYQ Count'
            for x in subs:
                r = st.add_row().cells
                r[0].text = x
                r[1].text = cell(counts.get((subject, topic, x), 0) if counts else None)
            r = st.add_row().cells
            r[0].text = 'TOTAL'
            r[1].text = cell(sum(counts.get((subject, topic, x), 0) for x in subs)
                             if counts else None)

        bold('IFAS Edutech  —  %s %s PYQ Analysis' % (exam_code, subject), 10)

    os.makedirs(out_dir, exist_ok=True)
    out = os.path.join(out_dir, '%s%s' % (exam_code, ANALYSIS_DOC_SUFFIX))
    doc.save(out)
    return out


CK_LIVE_FIXTURE = 'IIT_JAM_BIOTECHNOLOGY_PYQ_Analysis.docx'


def _ck_shapes():
    """The exam-shape matrix. The framework serves ~200 exams and cannot be validated
    against 200 real corpora, so the shapes are GENERATED. This matrix IS the
    exam-agnostic correctness claim; it covers shapes no real exam has produced yet.
    """
    # ABSOLUTE length, never derived from bc.MAX_HEADING_LEN. Deriving it would make
    # this shape vacuous with respect to the constant: lowering the cap would shrink
    # the fixture with it and the test would stay green (INV-8, a check that measures
    # nothing). 150 characters is comfortably legal at 300 and illegal at the old 100.
    long_ok = 'Regulation of Gene Expression in Prokaryotes and Eukaryotes ' + 'x' * 90
    assert len(long_ok) == 150
    return [
        ('single subject', {'Physics': {'Mechanics': ['Kinematics', 'Newton Laws']}}),
        ('six subjects', {'S%d' % i: {'T%d' % j: ['x%d_%d' % (i, j)]
                                      for j in range(1, i + 2)} for i in range(1, 7)}),
        ('thirty subjects', {'Sub %02d' % i: {'Topic A': ['one', 'two']}
                             for i in range(1, 31)}),
        ('duplicate TOPIC name across subjects',
         {'Physics': {'Thermodynamics': ['Entropy', 'Carnot']},
          'Chemistry': {'Thermodynamics': ['Enthalpy', 'Free Energy']}}),
        ('duplicate SUBTOPIC name across topics',
         {'Biology': {'Cell': ['Introduction', 'Organelles'],
                      'Genetics': ['Introduction', 'Linkage']}}),
        ('punctuation in names',
         {'General Intelligence & Reasoning':
          {'Analogy — Verbal / Non-Verbal': ["Bloom's Taxonomy, Applied",
                                             'Ratio & Proportion']}}),
        ('unicode in names', {'\u0938\u093e\u092e\u093e\u0928\u094d\u092f \u0939\u093f\u0902\u0926\u0940':
                              {'\u0935\u094d\u092f\u093e\u0915\u0930\u0923': ['\u0938\u0902\u0927\u093f', '\u0938\u092e\u093e\u0938']}}),
        ('name at the length boundary', {'Physics': {'Mechanics': [long_ok]}}),
        ('single subtopic per topic', {'Maths': {'Algebra': ['Matrices']}}),
        ('many topics one subject',
         {'GS': {'Topic %02d' % i: ['a', 'b', 'c'] for i in range(1, 21)}}),
    ]


def _ck_selftest(check, tmp):
    """Cluster K assertions. `check(name, cond)` is the caller's harness hook."""
    import glob as _glob

    # ── the heading limit is real, and a 150-char name must survive it ─────────
    class _R:
        def __init__(self, t):
            self.text, self.bold = t, True
    class _P:
        def __init__(self, t):
            self.text, self.runs = t, [_R(t)]
    check('K MAX_HEADING_LEN admits a 150-character subtopic heading',
          bc.is_taxonomy_heading(_P('n' * 150), lambda t: False))
    check('K MAX_HEADING_LEN still rejects a pasted stem',
          not bc.is_taxonomy_heading(_P('n' * bc.MAX_HEADING_LEN), lambda t: False))

    # ── ROUND TRIP over the generated matrix ────────────────────────────────────
    for label, tax in _ck_shapes():
        d = os.path.join(tmp, 'rt_' + re.sub(r'[^a-z0-9]+', '_', label.lower()))
        os.makedirs(d, exist_ok=True)
        try:
            p = write_analysis_doc(tax, 'EXAM', out_dir=d)
            got = read_analysis_doc(p)
        except Exception as ex:
            # Report per shape instead of aborting the suite: when this cluster
            # regresses, the operator needs to see WHICH shapes broke, not just the
            # first one.
            for suffix in ('round-trip', 'subject order', 'counts', 'topic_idx positional'):
                check('K %s: %s [%s]' % (suffix, label, type(ex).__name__), False)
            continue
        rebuilt = {s: {t: list(v['subtopics']) for t, v in d2['topics'].items()}
                   for s, d2 in got['taxonomy'].items()}
        check('K round-trip: ' + label, rebuilt == tax)
        check('K subject order: ' + label, got['subjects'] == list(tax))
        n = sum(len(x) for t in tax.values() for x in t.values())
        check('K counts: ' + label,
              got['counts'] == {'subjects': len(tax),
                                'topics': sum(len(t) for t in tax.values()),
                                'subtopics': n})
        # topic_idx must be contiguous 0..n-1 WITHIN each subject and never collide
        ok = all(sorted(v['topic_idx'] for v in d2['topics'].values())
                 == list(range(len(d2['topics']))) for d2 in got['taxonomy'].values())
        check('K topic_idx positional: ' + label, ok)

    # ── every level-1 and level-2 label form the engine blesses ─────────────────
    tax = {'Alpha': {'Beta': ['gamma', 'delta']}}
    d = os.path.join(tmp, 'labels'); os.makedirs(d, exist_ok=True)
    src = write_analysis_doc(tax, 'EXAM', out_dir=d)
    from docx import Document as _D
    for l1 in ('Subject', 'Domain', 'Section', 'Part', 'Area'):
        for l2 in ('Topic 1', 'Chapter 1', 'Unit 1', 'Module 1', 'Block 1'):
            doc = _D(src)
            for para in doc.paragraphs:
                t = para.text.strip()
                if t.startswith('Subject:'):
                    para.runs[0].text = '%s: Alpha' % l1
                elif t.startswith('Topic 1:'):
                    para.runs[0].text = '%s: Beta' % l2
            alt = os.path.join(d, 'ALT%s' % ANALYSIS_DOC_SUFFIX)
            doc.save(alt)
            # the master-summary rows still say "Topic 1:", which is level 2 either way
            got = read_analysis_doc(alt, verify=False)
            check('K label form %s / %s' % (l1, l2),
                  got['subjects'] == ['Alpha']
                  and list(got['taxonomy']['Alpha']['topics']) == ['Beta'])

    # ── the self-check must FIRE on each corruption, not just pass on good input ─
    def _mutate(fn, name):
        d2 = os.path.join(tmp, 'mut_' + name); os.makedirs(d2, exist_ok=True)
        p = write_analysis_doc({'Physics': {'Mechanics': ['Kinematics', 'Newton Laws']},
                                'Chemistry': {'Bonding': ['Ionic']}}, 'EXAM', out_dir=d2)
        doc = _D(p)
        fn(doc)
        doc.save(p)
        try:
            read_analysis_doc(p)
            return False
        except AnalysisDocError:
            return True

    def drop_subject_line(doc):
        for para in doc.paragraphs:
            if para.text.strip() == 'Subject: Chemistry':
                para.runs[0].text = 'Chemistry notes'
                return

    def dup_subject(doc):
        for para in doc.paragraphs:
            if para.text.strip() == 'Subject: Chemistry':
                para.runs[0].text = 'Subject: Physics'
                return

    def wrong_declared_count(doc):
        for para in doc.paragraphs:
            if para.text.strip().startswith('Total PYQs:'):
                para.runs[0].text = 'Total PYQs: —  |  Subtopics: 7'
                return

    def wrong_grand_total(doc):
        for tbl in doc.tables:
            for row in tbl.rows:
                if row.cells[0].text.strip() == 'GRAND TOTAL':
                    row.cells[1].text = '99'
                    return

    def drop_a_subtopic_row(doc):
        for tbl in doc.tables:
            for row in tbl.rows:
                if row.cells[0].text.strip() == 'Newton Laws':
                    row.cells[0].text = ''
                    return

    check('K catches a missing subject heading', _mutate(drop_subject_line, 'dropsubj'))
    check('K catches a duplicated subject', _mutate(dup_subject, 'dupsubj'))
    check('K catches a wrong declared subtopic count', _mutate(wrong_declared_count, 'badcount'))
    check('K catches a wrong GRAND TOTAL', _mutate(wrong_grand_total, 'badgrand'))
    check('K catches a dropped subtopic row', _mutate(drop_a_subtopic_row, 'droprow'))

    # ── name length is refused at the PRODUCER, not discovered downstream ───────
    too_long = 'y' * bc.MAX_HEADING_LEN
    d3 = os.path.join(tmp, 'toolong'); os.makedirs(d3, exist_ok=True)
    try:
        write_analysis_doc({'Physics': {'Mechanics': [too_long]}}, 'EXAM', out_dir=d3)
        check('K refuses to write an over-length name', False)
    except AnalysisDocError:
        check('K refuses to write an over-length name', True)

    # ── discovery ──────────────────────────────────────────────────────────────
    dd = os.path.join(tmp, 'disc'); os.makedirs(dd, exist_ok=True)
    up = os.path.join(tmp, 'disc_up'); os.makedirs(up, exist_ok=True)
    dirs = (dd + os.sep, up + os.sep)

    def _clear():
        for d4 in (dd, up):
            for f in _glob.glob(os.path.join(d4, '*')):
                os.remove(f)

    def _raises(sub):
        try:
            discover_analysis_doc(dirs)
            return False
        except AnalysisDocError as ex:
            return sub in str(ex)

    _clear()
    check('K no doc -> named stop', _raises('no Analysis doc found'))

    _clear()
    open(os.path.join(dd, '~$EXAM' + ANALYSIS_DOC_SUFFIX), 'w').close()
    check('K ignores a Word lock file', _raises('no Analysis doc found'))

    _clear()
    open(os.path.join(dd, 'EXAM_PYQ_Analysis.doc'), 'w').close()
    check('K names the Word 97 .doc case', _raises('Word 97'))

    _clear()
    open(os.path.join(dd, 'EXAM_PYQ_Analysis_Physics.docx'), 'w').close()
    check('K names the retired per-subject form', _raises('not the expected deliverable name'))

    _clear()
    open(os.path.join(dd, 'A' + ANALYSIS_DOC_SUFFIX), 'w').close()
    open(os.path.join(dd, 'B' + ANALYSIS_DOC_SUFFIX), 'w').close()
    check('K stops on two Analysis docs', _raises('exactly one is expected'))

    _clear()
    write_analysis_doc({'Physics': {'Mechanics': ['Kinematics']}}, 'EXAM', out_dir=dd)
    write_analysis_doc({'Physics': {'Mechanics': ['DIFFERENT']}}, 'EXAM', out_dir=up)
    got = read_analysis_doc(search_dirs=dirs)
    check('K same name in both dirs -> project wins, loaded once',
          got['triples'] == [('Physics', 'Mechanics', 'Kinematics')])

    # ── fingerprint ────────────────────────────────────────────────────────────
    a = read_analysis_doc(os.path.join(dd, 'EXAM' + ANALYSIS_DOC_SUFFIX))
    b = read_analysis_doc(os.path.join(up, 'EXAM' + ANALYSIS_DOC_SUFFIX))
    check('K fingerprint distinguishes different taxonomies',
          a['fingerprint'] != b['fingerprint'])
    check('K fingerprint is stable across reads',
          a['fingerprint'] == read_analysis_doc(a['path'])['fingerprint'])

    # ══════════════════════════════════════════════════════════════════════════
    # INGEST FORMS (GAP-2026-07-25-003)
    # ══════════════════════════════════════════════════════════════════════════
    def _render_md(tax, exam_code, bold='**', escape=False):
        """Emit the platform's OBSERVED docx->markdown grammar for a taxonomy.

        DELIBERATE CAVEAT: this helper is framework-written, so every test built on
        it is STRUCTURAL — it proves _scan_text() is the inverse of the grammar we
        believe the platform emits, never that the platform emits it. The live
        fixture is the only field evidence and must never be replaced by this.
        """
        esc = (lambda s: re.sub(r'([\\`*_{}\[\]()#+\-.!|])', r'\\\1', s)) if escape \
            else (lambda s: s)
        b = lambda s: '%s%s%s' % (bold, s, bold)
        out = []
        for subject in tax:
            topics = tax[subject]
            n_sub = sum(len(v) for v in topics.values())
            out += [b('%s — %s' % (exam_code, subject)), '',
                    b('Subject: %s' % subject), '',
                    b('PYQ Topic %s&%s Subtopic-wise Count' % (bold + bold, bold + bold))
                    if bold == '**' else b('PYQ Topic & Subtopic-wise Count'), '',
                    b('Total: — Questions  |  %d Topics  |  %d Subtopics'
                      % (len(topics), n_sub)), '',
                    '| Topic | Total Subtopics | Total PYQs |', '| --- | --- | --- |']
            for t_i, (topic, subs) in enumerate(topics.items(), 1):
                out.append('| %s | %d | — |' % (esc('Topic %d: %s' % (t_i, topic)), len(subs)))
            out += ['| GRAND TOTAL | %d | — |' % n_sub, '']
            for t_i, (topic, subs) in enumerate(topics.items(), 1):
                out += [b('Topic %d: %s' % (t_i, topic)), '',
                        b('Total PYQs: —  |  Subtopics: %d' % len(subs)), '',
                        '| Subtopic | PYQ Count |', '| --- | --- |']
                for x in subs:
                    out.append('| %s | — |' % esc(x))
                out += ['| TOTAL | — |', '']
            out += [b('IFAS Edutech  —  %s %s PYQ Analysis' % (exam_code, subject)), '']
        return '\n'.join(out) + '\n'

    ing = os.path.join(tmp, 'ingest'); os.makedirs(ing, exist_ok=True)

    def _write_text(name, body):
        p = os.path.join(ing, name)
        with open(p, 'w', encoding='utf-8') as fh:
            fh.write(body)
        return p

    def _stops(fn, sub):
        try:
            fn()
            return False
        except AnalysisDocError as ex:
            return sub in str(ex)

    # ── T3 form classification: content, never extension ───────────────────────
    _t3tax = {'Physics': {'Mechanics': ['Kinematics', 'Newton Laws']}}
    ok_docx = write_analysis_doc(_t3tax, 'T3', out_dir=ing)
    txt_doc = _write_text('T3TEXT' + ANALYSIS_DOC_SUFFIX, _render_md(_t3tax, 'T3'))
    empty = os.path.join(ing, 'empty.docx'); open(empty, 'wb').close()
    trunc = os.path.join(ing, 'trunc.docx')
    with open(ok_docx, 'rb') as _f, open(trunc, 'wb') as _g:
        _g.write(_f.read(400))
    notword = os.path.join(ing, 'notword.docx')
    import zipfile as _zf
    with _zf.ZipFile(notword, 'w') as _z:
        _z.writestr('xl/workbook.xml', '<x/>')
    utf16 = os.path.join(ing, 'utf16.docx')
    with open(utf16, 'wb') as _f:
        _f.write('Subject: X'.encode('utf-16'))
    check('K T3 ingest form: real .docx -> ooxml', _ingest_form(ok_docx) == INGEST_OOXML)
    check('K T3 ingest form: extracted text -> text', _ingest_form(txt_doc) == INGEST_TEXT)
    check('K T3 ingest form: zero bytes -> unknown', _ingest_form(empty) == INGEST_UNKNOWN)
    check('K T3 ingest form: truncated zip -> unknown', _ingest_form(trunc) == INGEST_UNKNOWN)
    check('K T3 ingest form: non-Word zip -> unknown', _ingest_form(notword) == INGEST_UNKNOWN)
    check('K T3 ingest form: UTF-16 binary -> unknown', _ingest_form(utf16) == INGEST_UNKNOWN)

    # ── T4/T5 round trip THROUGH the text form, for every generated shape ──────
    for label, tax in _ck_shapes():
        tp = _write_text('rt_%s.docx' % re.sub(r'[^a-z0-9]+', '_', label.lower()),
                         _render_md(tax, 'EXAM'))
        try:
            got = read_analysis_doc(tp)
        except Exception as ex:
            for suffix in ('text round-trip', 'text counts', 'text fingerprint parity'):
                check('K %s: %s [%s]' % (suffix, label, type(ex).__name__), False)
            continue
        rebuilt = {s: {t: list(v['subtopics']) for t, v in d2['topics'].items()}
                   for s, d2 in got['taxonomy'].items()}
        check('K text round-trip: ' + label, rebuilt == tax)
        n = sum(len(x) for t in tax.values() for x in t.values())
        check('K text counts: ' + label,
              got['counts'] == {'subjects': len(tax),
                                'topics': sum(len(t) for t in tax.values()),
                                'subtopics': n})
        d = os.path.join(tmp, 'x_' + re.sub(r'[^a-z0-9]+', '_', label.lower()))
        os.makedirs(d, exist_ok=True)
        # T2 cross-form identity: the two ingest forms of ONE taxonomy agree exactly
        check('K text fingerprint parity: ' + label,
              got['fingerprint'] ==
              read_analysis_doc(write_analysis_doc(tax, 'EXAM', out_dir=d))['fingerprint'])

    # ── grammar variants OBSERVED in a second real extractor ───────────────────
    check('K text grammar: __bold__ whole-line wrap',
          read_analysis_doc(_write_text('UND' + ANALYSIS_DOC_SUFFIX,
                                        _render_md(_t3tax, 'UND', bold='__')))['triples']
          == [('Physics', 'Mechanics', 'Kinematics'), ('Physics', 'Mechanics', 'Newton Laws')])
    _esc_tax = {'Maths': {'Series': ['p_1 and p_2 tests', 'Big-O [worst case]']}}
    check('K text grammar: backslash-escaped cells unescape exactly',
          read_analysis_doc(_write_text('ESC' + ANALYSIS_DOC_SUFFIX,
                                        _render_md(_esc_tax, 'ESC', escape=True)))['triples']
          == [('Maths', 'Series', 'p_1 and p_2 tests'),
              ('Maths', 'Series', 'Big-O [worst case]')])
    check('K text grammar: underscores inside a name are NEVER stripped',
          _md_plain('p_1 and p__2') == 'p_1 and p__2')

    # ── the ragged-row stop: a '|' in a name must never parse silently ─────────
    _pipe_md = _render_md({'Physics': {'Waves': ['Transverse', 'Standing']}}, 'PIPE')
    _pipe_md = _pipe_md.replace('| Transverse | — |', '| Transverse | Longitudinal | — |')
    pipe_doc = _write_text('PIPE' + ANALYSIS_DOC_SUFFIX, _pipe_md)
    check('K ragged row (| in a name) HARD STOPS rather than truncating',
          _stops(lambda: read_analysis_doc(pipe_doc), 'same number of cells'))
    check('K ragged row names the likely cause',
          _stops(lambda: read_analysis_doc(pipe_doc), "contains a '|' character"))
    check('K writer REFUSES to emit a name containing a pipe',
          _stops(lambda: write_analysis_doc({'S': {'T': ['a | b']}}, 'PIPE', out_dir=ing),
                 "cannot be written"))
    check('K writer accepts an escaped-pipe-free taxonomy unchanged',
          bool(write_analysis_doc({'S': {'T': ['a / b']}}, 'OKP', out_dir=ing)))

    # ── T6 corruption of the text form fires, per subject, every time ──────────
    _base_md = _render_md({'A': {'T1': ['x', 'y']}, 'B': {'T2': ['p', 'q']}}, 'C')
    for tag, mutated in (
            ('subject heading deleted', _base_md.replace('**Subject: B**\n', '')),
            ('a data row deleted', _base_md.replace('| y | — |\n', '')),
            ('a D1 count corrupted', _base_md.replace('2 Subtopics', '9 Subtopics', 1)),
            ('every table deleted',
             '\n'.join(l for l in _base_md.split('\n') if not l.lstrip().startswith('|')))):
        mp = _write_text('C%s.docx' % re.sub(r'\W+', '', tag), mutated)
        check('K T6 text corruption fires: ' + tag,
              _stops(lambda p=mp: read_analysis_doc(p), 'HARD STOP'))

    # ── T7 verify=False is REFUSED on the text form (never on OOXML) ───────────
    _bad = _write_text('T7' + ANALYSIS_DOC_SUFFIX,
                       _base_md.replace('2 Subtopics', '9 Subtopics', 1))
    check('K T7 verify=False is refused on a text-form doc',
          _stops(lambda: read_analysis_doc(_bad, verify=False), 'does not agree with itself'))

    # ── T8 unknown form: named, sized, and never a library traceback ───────────
    for nm, p in (('zero bytes', empty), ('truncated zip', trunc),
                  ('non-Word zip', notword), ('UTF-16 binary', utf16)):
        check('K T8 unknown form is named: ' + nm,
              _stops(lambda p=p: read_analysis_doc(p), 'not a recognised ingest form'))
    check('K T8 unknown form never surfaces PackageNotFoundError',
          not _stops(lambda: read_analysis_doc(trunc), 'PackageNotFoundError'))
    check('K T8 a vanished/unreadable path still gets the NAMED stop, not a traceback',
          _stops(lambda: read_analysis_doc(os.path.join(ing, 'gone_PYQ_Analysis.docx')),
                 'not a recognised ingest form'))
    check('K T8 the unreadable path reports why the size is missing',
          _stops(lambda: read_analysis_doc(os.path.join(ing, 'gone_PYQ_Analysis.docx')),
                 'unreadable ('))
    check('K T8 a DIRECTORY given as the path gets the named stop too',
          _stops(lambda: read_analysis_doc(ing), 'not a recognised ingest form'))

    # ── T9 downstream shape is unchanged, plus the additive ingest_form key ────
    _t9 = read_analysis_doc(txt_doc)
    check('K T9 return shape unchanged by the text path',
          set(_t9) >= {'path', 'exam_code', 'subjects', 'taxonomy', 'triples',
                       'counts', 'fingerprint'})
    check('K T9 ingest_form reported as text', _t9['ingest_form'] == INGEST_TEXT)
    check('K T9 ingest_form reported as ooxml',
          read_analysis_doc(ok_docx)['ingest_form'] == INGEST_OOXML)
    check('K T9 topic_idx positional on the text path',
          all(sorted(v['topic_idx'] for v in d2['topics'].values())
              == list(range(len(d2['topics']))) for d2 in _t9['taxonomy'].values()))

    # ══════════════════════════════════════════════════════════════════════════
    # THE TAXONOMY LOCK GATE
    # ══════════════════════════════════════════════════════════════════════════
    lk = os.path.join(tmp, 'lock'); os.makedirs(lk, exist_ok=True)
    _lt = {'Physics': {'Mechanics': ['Kinematics', 'Newton Laws']},
           'Chemistry': {'Bonding': ['Ionic', 'Covalent']}}
    _lp = write_analysis_doc(_lt, 'LOCKEXAM', out_dir=lk)
    _ld = read_analysis_doc(_lp)

    def _rec(fp, name='LOCKEXAM' + APPROVAL_RECORD_SUFFIX):
        body = {'exam_code': 'LOCKEXAM', 'status': 'CLEAN'}
        if fp is not None:
            body['taxonomy_fingerprint'] = fp
        with open(os.path.join(lk, name), 'w', encoding='utf-8') as fh:
            json.dump(body, fh)

    _rec(_ld['fingerprint'])
    check('K LOCK matching fingerprint passes',
          assert_taxonomy_lock(_ld, search_dirs=(lk,))['exam_code'] == 'LOCKEXAM')

    _rec('v1:9:9:' + 'f' * 64)
    check('K LOCK mismatched fingerprint HARD STOPS',
          _stops(lambda: assert_taxonomy_lock(_ld, search_dirs=(lk,)),
                 'NOT the taxonomy that was approved'))
    check('K LOCK names the step it fired at',
          _stops(lambda: assert_taxonomy_lock(_ld, search_dirs=(lk,), step='PYQExtract'),
                 'HARD STOP at PYQExtract'))

    _rec(None)
    check('K LOCK absent fingerprint HARD STOPS with the re-run action',
          _stops(lambda: assert_taxonomy_lock(_ld, search_dirs=(lk,)),
                 'no taxonomy_fingerprint'))

    os.remove(os.path.join(lk, 'LOCKEXAM' + APPROVAL_RECORD_SUFFIX))
    check('K LOCK missing record HARD STOPS rather than passing',
          _stops(lambda: assert_taxonomy_lock(_ld, search_dirs=(lk,)),
                 'no approval record found'))

    _rec(_ld['fingerprint'])
    _rec(_ld['fingerprint'], 'OTHEREXAM' + APPROVAL_RECORD_SUFFIX)
    check('K LOCK exam_code selects the right record when two are present',
          assert_taxonomy_lock(_ld, search_dirs=(lk,))['exam_code'] == 'LOCKEXAM')
    check('K LOCK ambiguous record set (no exam_code) HARD STOPS',
          _stops(lambda: discover_approval_record(search_dirs=(lk,)),
                 'exactly one is expected'))

    with open(os.path.join(lk, 'BADEXAM' + APPROVAL_RECORD_SUFFIX), 'w') as fh:
        fh.write('{not json')
    check('K LOCK malformed record HARD STOPS by name',
          _stops(lambda: read_approval_record('BADEXAM', search_dirs=(lk,)),
                 'could not be read as JSON'))

    # the gate must survive the TEXT ingest form identically
    _ltxt = os.path.join(lk, 'LOCKEXAM' + ANALYSIS_DOC_SUFFIX)
    with open(_ltxt, 'w', encoding='utf-8') as fh:
        fh.write(_render_md(_lt, 'LOCKEXAM'))
    check('K LOCK holds across ingest forms (text doc, same record)',
          assert_taxonomy_lock(read_analysis_doc(_ltxt),
                               search_dirs=(lk,))['exam_code'] == 'LOCKEXAM')

    # ══════════════════════════════════════════════════════════════════════════
    # load_taxonomy() — record first, document as fallback (v1.4)
    # ══════════════════════════════════════════════════════════════════════════
    jd = os.path.join(tmp, 'json'); os.makedirs(jd, exist_ok=True)
    _jt = {'General Biology': {'Cell Biology': ['Membrane', 'Nucleus'],
                               'Genetics': ['Linkage']},
           'Physics': {'Mechanics': ['Kinematics']}}
    _jdoc = read_analysis_doc(write_analysis_doc(_jt, 'JX', out_dir=jd))

    def _jrec(**over):
        body = {'exam_code': 'JX', 'schema_version': '1.3', 'status': 'CLEAN',
                'taxonomy_fingerprint': _jdoc['fingerprint'],
                'taxonomy_counts': dict(_jdoc['counts']),
                'taxonomy': {'sections': {s: {t: list(x) for t, x in v.items()}
                                          for s, v in _jt.items()}}}
        body.update(over)
        with open(os.path.join(jd, 'JX' + APPROVAL_RECORD_SUFFIX), 'w',
                  encoding='utf-8') as fh:
            json.dump(body, fh)

    _jrec()
    _lt = load_taxonomy('JX', search_dirs=(jd,))
    check('K JSON source is the approval record', _lt['source'] == 'approval_record')
    check('K JSON ingest_form is json', _lt['ingest_form'] == 'json')
    check('K JSON triples identical to the document path',
          _lt['triples'] == _jdoc['triples'])
    check('K JSON taxonomy identical to the document path',
          _lt['taxonomy'] == _jdoc['taxonomy'])
    check('K JSON fingerprint identical to the document path',
          _lt['fingerprint'] == _jdoc['fingerprint'])
    check('K JSON subject order preserved', _lt['subjects'] == list(_jt))
    check('K JSON topic_idx positional and contiguous',
          all(sorted(v['topic_idx'] for v in d2['topics'].values())
              == list(range(len(d2['topics']))) for d2 in _lt['taxonomy'].values()))

    # the record must be believed only when it agrees with ITSELF
    _jrec(taxonomy_fingerprint='v1:9:9:' + 'f' * 64)
    check('K JSON record whose halves disagree HARD STOPS',
          _stops(lambda: load_taxonomy('JX', search_dirs=(jd,)),
                 'does not agree with itself'))
    _jrec(taxonomy_counts={'subjects': 9, 'topics': 9, 'subtopics': 9})
    check('K JSON declared counts vs assembled counts HARD STOPS',
          _stops(lambda: load_taxonomy('JX', search_dirs=(jd,)),
                 'assembles to'))
    _jrec(taxonomy={'sections': {}})
    check('K JSON empty sections HARD STOPS',
          _stops(lambda: load_taxonomy('JX', search_dirs=(jd,)),
                 'not a populated'))
    _jrec(taxonomy={'sections': {'S': {'T': []}}})
    check('K JSON topic with no subtopics HARD STOPS',
          _stops(lambda: load_taxonomy('JX', search_dirs=(jd,)), 'no subtopics'))
    _jrec(taxonomy={'sections': {'S': {}}})
    check('K JSON subject with no topics HARD STOPS',
          _stops(lambda: load_taxonomy('JX', search_dirs=(jd,)), 'no topics'))
    _jrec(taxonomy={'sections': {'S': {'T': ['a', 'a']}}})
    check('K JSON duplicate name HARD STOPS instead of collapsing silently',
          _stops(lambda: load_taxonomy('JX', search_dirs=(jd,)),
                 'more than once'))
    _jrec(taxonomy_fingerprint=None)
    check('K JSON taxonomy without a fingerprint HARD STOPS',
          _stops(lambda: load_taxonomy('JX', search_dirs=(jd,)),
                 'cannot validate itself'))

    # ── the pre-1.3 fallback: record has a fingerprint and NO taxonomy ──────────
    _jrec(taxonomy=None)
    _fb = load_taxonomy('JX', search_dirs=(jd,))
    check('K pre-1.3 record falls back to the Analysis doc',
          _fb['source'] == 'analysis_doc')
    check('K fallback yields the IDENTICAL taxonomy',
          _fb['triples'] == _jdoc['triples'] and
          _fb['fingerprint'] == _jdoc['fingerprint'])
    check('K fallback is still lock-gated',
          _stops(lambda: (_jrec(taxonomy=None,
                                taxonomy_fingerprint='v1:9:9:' + 'f' * 64),
                          load_taxonomy('JX', search_dirs=(jd,)))[1],
                 'NOT the taxonomy that was approved'))

    # ── and it works when the doc is in the TEXT ingest form ───────────────────
    with open(os.path.join(jd, 'JX' + ANALYSIS_DOC_SUFFIX), 'w',
              encoding='utf-8') as fh:
        fh.write(_render_md(_jt, 'JX'))
    _jrec(taxonomy=None)
    _fbt = load_taxonomy('JX', search_dirs=(jd,))
    check('K fallback works on the TEXT ingest form too',
          _fbt['ingest_form'] == 'text' and _fbt['fingerprint'] == _jdoc['fingerprint'])

    # ── and the record path needs NO document at all ───────────────────────────
    nd = os.path.join(tmp, 'norec'); os.makedirs(nd, exist_ok=True)
    with open(os.path.join(nd, 'JX' + APPROVAL_RECORD_SUFFIX), 'w',
              encoding='utf-8') as fh:
        json.dump({'exam_code': 'JX', 'schema_version': '1.3',
                   'taxonomy_fingerprint': _jdoc['fingerprint'],
                   'taxonomy': {'sections': _jt}}, fh)
    check('K JSON path needs no Analysis doc on disk at all',
          load_taxonomy('JX', search_dirs=(nd,))['fingerprint'] == _jdoc['fingerprint'])

    # ── the live fixture: the one shape that is field evidence, not generated ───
    fx = None
    for cand in (os.path.join(os.path.dirname(os.path.abspath(__file__)), CK_LIVE_FIXTURE),
                 os.path.join('fixtures', CK_LIVE_FIXTURE), CK_LIVE_FIXTURE):
        if os.path.exists(cand):
            fx = cand
            break
    if fx:
        live = read_analysis_doc(fx)
        check('K live fixture: 6 subjects / 26 topics / 131 subtopics',
              live['counts'] == {'subjects': 6, 'topics': 26, 'subtopics': 131})
        check('K live fixture: 6 distinct subjects across the triples',
              len({t[0] for t in live['triples']}) == 6)


# ═══════════════════════════════════════════════════════════════════════════════
# CLUSTER I — TABLE STRUCTURE  (GAP-2026-07-29-TBL)
# ═══════════════════════════════════════════════════════════════════════════════
# WHY THIS LIVES HERE AND NOT IN A SPEC. Two builders existed for one concept —
# Framework_PYQPrepare S4-3 build_di_table(rows_data) and Framework_MockTestCreate
# S8-4 build_di_table_styled(headers, rows) — and BOTH modelled a table as a
# rectangle of strings. Neither could express a merged cell, so a grouped header
# ("Printers" spanning L/M/N/O above "Days" spanning two rows) had exactly one
# representable form: squared into a grid and padded with empty strings. Measured
# on SSC_CGL_Tier1 09-Sep-2024 Shift 1 — Q.52 and Q.61 each lost a 4-column span
# and a 2-row span and gained 4 stray empty cells; the delivered file carried
# 0 gridSpan and 0 vMerge elements and passed 16/16 checks. One model, one owner,
# one implementation: Step 1 writes it, Step 7 rebuilds it, Steps 3/4/5 read it.
#
# THE DATA MODEL (TableSpec)
#   spec = {
#     'grid'        : [[Cell, ...], ...],   # REQUIRED, row-major, ANCHOR CELLS ONLY
#     'header_rows' : int,                  # OPTIONAL, default 1
#     'col_widths'  : [float, ...] | None,  # OPTIONAL, relative weights
#     'align'       : 'left'|'center'|'right',   # OPTIONAL, default 'center'
#     'note'        : str | None,           # OPTIONAL, footnote rendered below
#   }
#   Cell = str                              # == {'t': str, 'cs': 1, 'rs': 1}
#        | {'t': str, 'cs': int, 'rs': int, 'align': str|None, 'bold': bool}
#
# THE ONE RULE THAT REMOVES THE AMBIGUITY: a row declares only the cells that
# START in it. A position covered by a span from above or from the left is NOT
# declared. Padding a row with '' to square the grid is BANNED — under this model
# it cannot occur, which is what makes a genuinely blank cell ({'t': ''}, e.g. the
# classic empty top-left corner) distinguishable from padding STRUCTURALLY rather
# than by heuristic.

TABLE_ALIGNMENTS = ('left', 'center', 'right')


def normalise_table_spec(spec):
    """Accept a legacy list-of-lists OR a TableSpec dict. Return a TableSpec dict.

    The legacy shape stays valid for ever: every pre-GAP-2026-07-29-TBL call site
    keeps working unchanged, and a flat table has no spans to lose.
    """
    if isinstance(spec, (list, tuple)):
        spec = {'grid': list(spec)}
    elif isinstance(spec, dict) and 'grid' not in spec and 'headers' in spec:
        # Framework_MockTestCreate S8-4 legacy DI payload {'headers': [...],
        # 'rows': [[...]]}. One header row, no spans — accepted for ever so no
        # registry.json / blueprint payload has to be migrated.
        spec = {'grid': [list(spec['headers'])] + [list(r) for r in (spec.get('rows') or [])],
                'header_rows': 1}
    if not isinstance(spec, dict) or 'grid' not in spec:
        raise ValueError('TableSpec must be a dict with a "grid" key, a '
                         '{"headers","rows"} payload, or a list of rows')
    grid = []
    for row in spec['grid']:
        out = []
        for c in row:
            out.append(dict(c) if isinstance(c, dict) else {'t': '' if c is None else str(c)})
        grid.append(out)
    out_spec = dict(spec)
    out_spec['grid'] = grid
    return out_spec


def place_cells(grid):
    """HTML-table occupancy placement. Deterministic; no heuristics, no guessing.

    Returns (placements, width, nrows, errors) where placements is
    [(row, col, rowspan, colspan, cell), ...] in declaration order.

    A 'hole' means the transcription UNDER-declares a row; an 'overlap' means it
    OVER-declares. Both are transcription errors, and both are REPORTED rather
    than silently normalised — silent normalisation is the defect this whole
    cluster exists to end.
    """
    occ, placements, errors = {}, [], []
    for r, row in enumerate(grid):
        c = 0
        for cell in row:
            try:
                cs = int(cell.get('cs', 1))
                rs = int(cell.get('rs', 1))
            except (TypeError, ValueError):
                cs = rs = 1
                errors.append('row %d: non-integer cs/rs, coerced to 1' % r)
            if cs < 1 or rs < 1:
                errors.append('row %d: cs=%s rs=%s, both must be >= 1' % (r, cs, rs))
                cs, rs = max(cs, 1), max(rs, 1)
            while (r, c) in occ:
                c += 1
            for rr in range(r, r + rs):
                for cc in range(c, c + cs):
                    if (rr, cc) in occ:
                        errors.append('overlapping span at (%d,%d)' % (rr, cc))
                    occ[(rr, cc)] = True
            placements.append((r, c, rs, cs, cell))
            c += cs
    width = max((cc for (_rr, cc) in occ), default=-1) + 1
    nrows = len(grid)
    for r in range(nrows):
        for c in range(width):
            if (r, c) not in occ:
                errors.append('hole at (%d,%d) — row %d under-declares its cells' % (r, c, r))
    return placements, width, nrows, errors


def table_spec_spans(spec):
    """Declared spans as a sorted (row, col, rowspan, colspan) list. CHECK 17 input."""
    placements, _w, _n, _e = place_cells(normalise_table_spec(spec)['grid'])
    return sorted((r, c, rs, cs) for (r, c, rs, cs, _cell) in placements if rs > 1 or cs > 1)


def build_di_table(doc, spec, font_pt=9, default_align='center', render=None,
                   style='Table Grid', avail_emu=5734050, font_name='Arial'):
    """Build a native Word table from a TableSpec (or a legacy list of lists).

    ORDERING IS LOAD-BEARING AND IS THE EASIEST THING TO GET WRONG:
      1. write text into ANCHOR cells
      2. THEN merge
      3. THEN set per-column widths
    cell.merge() CONCATENATES the text of both cells ('A' + 'B' -> 'A\\nB',
    verified on python-docx 1.2.0), so text must never be written into a covered
    position and merging must never precede text placement.

    render : optional callable(paragraph, text, bold=bool). When supplied, cell
             text is emitted through the caller's own renderer, so OMML math and
             {{u}} underline markers work INSIDE table cells exactly as they do in
             stems and options (Framework_PYQPrepare S1-6 / S1-11). When omitted,
             plain text is written.
    avail_emu : usable text width. Default 6.27in (A4 minus 1in margins) in EMU.
    font_name : cell font family. Default 'Arial' = the Row-file contract (S1-1).
                Step 7 passes its configurable FONT_NAME.
    """
    from docx.shared import Pt, Emu
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    amap = {'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT}

    spec = normalise_table_spec(spec)
    placements, width, nrows, errors = place_cells(spec['grid'])
    if errors:
        raise ValueError('TableSpec geometry invalid: ' + '; '.join(errors[:5]))

    tbl = doc.add_table(rows=nrows, cols=width, style=style)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    table_align = spec.get('align') or default_align
    if table_align not in TABLE_ALIGNMENTS:
        table_align = 'center'
    header_rows = int(spec.get('header_rows', 1) or 0)

    # 1 — text into anchors only
    for (r, c, _rs, _cs, cell) in placements:
        tc = tbl.cell(r, c)
        tc.text = ''
        para = tc.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        cell_align = cell.get('align') or table_align
        para.alignment = amap.get(cell_align, amap['center'])
        bold = bool(cell.get('bold', False))
        text = cell.get('t', '')
        if render is not None:
            render(para, text, bold=bold)
        elif text:
            para.add_run(text).bold = bold
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_pt)

    # 2 — merges, only after every anchor already holds its text
    for (r, c, rs, cs, _cell) in placements:
        if rs > 1 or cs > 1:
            tbl.cell(r, c).merge(tbl.cell(r + rs - 1, c + cs - 1))

    # 3 — per-cell column widths. tblGrid alone is advisory and several renderers
    #     ignore it, so the width must be stamped on every cell of the column.
    weights = spec.get('col_widths') or [1.0] * width
    if len(weights) != width:
        weights = [1.0] * width
    total = float(sum(weights)) or 1.0
    for ci in range(width):
        w = Emu(int(avail_emu * (weights[ci] / total)))
        for ri in range(nrows):
            tbl.cell(ri, ci).width = w

    if spec.get('note'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(spec['note']))
        run.font.name = font_name
        run.font.size = Pt(font_pt)
        run.italic = True
    _ = header_rows          # reserved for header styling by callers that want it
    return tbl


def adjacent_table_pairs(doc):
    """Count pairs of adjacent block-level <w:tbl> siblings. CHECK 18 input.

    Two consecutive w:tbl siblings with no w:p between them are FUSED into one
    table by every Word engine. This is an OOXML block rule, not a renderer quirk,
    and it is invisible in the emitted file: python-docx reports N tables where
    Word shows one. Measured: a Row file written with 19 tables came back from a
    Word-engine round-trip with 7.
    """
    kinds = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag in ('p', 'tbl'):
            kinds.append(tag)
    pairs, run = 0, 0
    for k in kinds:
        run = run + 1 if k == 'tbl' else 0
        if run >= 2:
            pairs += 1
    return pairs


# ═══════════════════════════════════════════════════════════════════════════════
# SELF-TEST
# ═══════════════════════════════════════════════════════════════════════════════

def self_test():
    passed = total = 0
    fails = []

    def check(name, cond):
        nonlocal passed, total
        total += 1
        if cond:
            passed += 1
        else:
            fails.append(name)

    # ── folder id parsing ────────────────────────────────────────────────────
    check('io_folder_url', parse_drive_folder_id(
        'https://drive.google.com/drive/folders/13sdlvaEJjwLKmsvP-zdCpv4l4GhZbF9s')
        == '13sdlvaEJjwLKmsvP-zdCpv4l4GhZbF9s')
    check('io_folder_url_u0', parse_drive_folder_id(
        'https://drive.google.com/drive/u/0/folders/ABC123defGHI') == 'ABC123defGHI')
    check('io_folder_bare_id', parse_drive_folder_id('13sdlvaEJjwLKmsvP') == '13sdlvaEJjwLKmsvP')
    check('io_folder_none', parse_drive_folder_id('not a link') is None)

    # ── listing normalisation ────────────────────────────────────────────────
    raw = [{'id': '1', 'title': 'a.docx', 'mimeType': bc.DOCX_MIME, 'fileSize': '12694409'},
           {'id': '2', 'name': 'b.docx', 'mimeType': bc.DOCX_MIME, 'size': 100},
           {'id': '3', 'title': 'yr', 'mimeType': bc.FOLDER_MIME}]
    norm = normalise_drive_listing(raw)
    check('io_norm_size_str_to_int', norm[0]['fileSize'] == 12694409)
    check('io_norm_name_alias', norm[1]['name'] == 'b.docx')
    check('io_norm_missing_size_none', norm[2]['fileSize'] is None)
    check('io_norm_envelope', len(normalise_drive_listing({'files': raw})) == 3)
    check('io_norm_empty', normalise_drive_listing([]) == [])
    check('io_norm_junk_skipped', len(normalise_drive_listing(['junk', None])) == 0)

    # ── enumeration: rejects, duplicates, pagination, recursion ──────────────
    tree = {
        'root': {'items': [
            {'id': 'f1', 'title': '2010.docx', 'mimeType': bc.DOCX_MIME, 'fileSize': '900'},
            {'id': 'f2', 'title': 'old.doc', 'mimeType': 'application/msword', 'fileSize': '9'},
            {'id': 'f3', 'title': 'native', 'mimeType': bc.GDOC_MIME},
            {'id': 'sub', 'title': '2011', 'mimeType': bc.FOLDER_MIME}]},
        'sub': {'items': [
            {'id': 'f4', 'title': '2011.docx', 'mimeType': bc.DOCX_MIME, 'fileSize': '800'}]},
    }

    def fake_list(fid, page_token=None):
        return tree[fid]

    papers, rejects = collect_corpus_files(fake_list, 'root')
    check('io_enum_finds_nested', len(papers) == 2)
    check('io_enum_rejects_both', len(rejects) == 2)
    check('io_enum_reject_has_reason', all(r['reason'] for r in rejects))
    check('io_enum_size_captured', papers[0]['fileSize'] == 900)
    check('io_enum_key_present', 'paper_key' in papers[0])

    def fake_dup(fid, page_token=None):
        return {'items': [
            {'id': 'a', 'title': 'X_2010.docx', 'mimeType': bc.DOCX_MIME, 'fileSize': '1'},
            {'id': 'b', 'title': 'X_2010 (1).docx', 'mimeType': bc.DOCX_MIME, 'fileSize': '2'}]}
    try:
        collect_corpus_files(fake_dup, 'root')
        check('io_enum_duplicate_hard_stop', False)
    except DuplicatePaperError as exc:
        check('io_enum_duplicate_hard_stop', 'X_2010' in str(exc))

    pages = [{'items': [{'id': 'p1', 'title': 'a.docx', 'mimeType': bc.DOCX_MIME,
                         'fileSize': '1'}], 'nextPageToken': 'T'},
             {'items': [{'id': 'p2', 'title': 'b.docx', 'mimeType': bc.DOCX_MIME,
                         'fileSize': '1'}]}]

    def fake_paged(fid, page_token=None):
        return pages[1] if page_token else pages[0]

    check('io_enum_paginates', len(collect_corpus_files(fake_paged, 'r')[0]) == 2)

    # ── payload decoding: every documented shape ─────────────────────────────
    payload = base64.b64encode(b'PK\x03\x04hello').decode()
    inner = json.dumps({'id': 'x', 'title': 't', 'mimeType': 'm', 'content': payload})
    check('io_decode_raw_bytes', decode_drive_payload(b'PK\x03\x04') == b'PK\x03\x04')
    check('io_decode_inner_json', decode_drive_payload(inner) == b'PK\x03\x04hello')
    check('io_decode_spill_shape',
          decode_drive_payload([{'text': inner, 'type': 'text'}]) == b'PK\x03\x04hello')
    check('io_decode_dict', decode_drive_payload({'content': payload}) == b'PK\x03\x04hello')
    spill = '/tmp/_ci_spill.json'
    with open(spill, 'w') as fh:
        json.dump([{'text': inner, 'type': 'text'}], fh)
    check('io_decode_spill_path', decode_drive_payload(spill) == b'PK\x03\x04hello')
    for bad, label in [([], 'empty_list'), ('nonsense', 'garbage'), ({'k': 1}, 'no_content')]:
        try:
            decode_drive_payload(bad)
            check(f'io_decode_reject_{label}', False)
        except TransportFallback:
            check(f'io_decode_reject_{label}', True)

    # ── download verification ────────────────────────────────────────────────
    check('io_verify_ok', verify_downloaded_bytes(b'PK\x03\x04abc', 7))
    for args, label in [((b'', None), 'empty'), ((b'<html>', None), 'not_docx'),
                        ((b'PK\x03\x04abc', 99), 'truncated')]:
        try:
            verify_downloaded_bytes(*args)
            check(f'io_verify_reject_{label}', False)
        except TransportFallback:
            check(f'io_verify_reject_{label}', True)

    # ── fetch: EVERY failure degrades, never SystemExit ──────────────────────
    def boom(_):
        raise RuntimeError('connector exploded')
    try:
        fetch_drive_docx(boom, {'id': 'x', 'name': 'n.docx', 'fileSize': 10}, '/tmp/ci')
        check('io_fetch_error_is_fallback', False)
    except TransportFallback:
        check('io_fetch_error_is_fallback', True)
    except Exception:
        check('io_fetch_error_is_fallback', False)
    try:
        fetch_drive_docx(boom, {'id': 'x', 'name': 'big.docx',
                                'fileSize': bc.DRIVE_CAP + 1}, '/tmp/ci')
        check('io_fetch_oversize_not_attempted', False)
    except TransportFallback as exc:
        check('io_fetch_oversize_not_attempted', 'without attempting' in str(exc))

    def good(_):
        return base64.b64encode(b'PK\x03\x04data').decode()
    p = fetch_drive_docx(lambda i: json.dumps({'content': good(i)}),
                         {'id': 'x', 'name': 'ok.docx', 'fileSize': 8}, '/tmp/ci_ok')
    check('io_fetch_writes_file', os.path.exists(p) and os.path.getsize(p) == 8)

    # ── upload resolution: identity, not filename, not recency ──────────────
    up = '/tmp/ci_uploads'
    os.makedirs(up, exist_ok=True)
    for f in os.listdir(up):
        os.remove(os.path.join(up, f))
    for fn in ['E_02-May-2010_Sorted (1).docx', 'E_2011_Sorted.docx',
               'Unrelated_2099.docx', '~$lock.docx']:
        open(os.path.join(up, fn), 'w').write('x')
    want = [bc.canonical_paper_key('E_02May2010_Sorted.docx'),
            bc.canonical_paper_key('E_2012_Sorted.docx')]
    res = resolve_uploaded_papers(want, up)
    check('io_upload_matches_through_suffix', len(res['matched']) == 1)
    check('io_upload_reports_missing',
          res['missing'] == [bc.canonical_paper_key('E_2012_Sorted.docx')])
    check('io_upload_flags_unexpected', len(res['unexpected']) == 2)
    check('io_upload_ignores_lockfiles',
          not any('~$' in p for p in res['unexpected']))
    check('io_upload_missing_dir_safe',
          resolve_uploaded_papers(['k'], '/tmp/does_not_exist')['missing'] == ['k'])

    # ── figural cross-check ──────────────────────────────────────────────────
    mp = {1: ['i1.png'], 5: ['i2.png'], PREAMBLE: ['i0.png']}
    fc = figural_consistency(mp, {1: 'FIGURAL', 5: 'TEXT', 9: 'FIGURAL'})
    check('io_figural_flags_image_not_figural', fc['image_not_figural'] == [5])
    check('io_figural_flags_figural_no_image', fc['figural_no_image'] == [9])
    check('io_figural_preamble_ignored', 'preamble' not in str(fc))
    fc2 = figural_consistency(mp, {1: 'FIGURAL', 5: 'TEXT', 9: 'FIGURAL'}, overrides=(5, 9))
    check('io_figural_overrides_respected',
          not fc2['image_not_figural'] and not fc2['figural_no_image'])

    # ── structural corruption must ALWAYS surface as a NAMED gate ────────────
    # Regression guard: python-docx raises a bare KeyError and zipfile a bare
    # BadZipFile from deep inside their readers. Both previously escaped as raw
    # tracebacks instead of gate verdicts. Found by adversarial mutation testing.
    _bad = '/tmp/_ci_badzip.docx'
    with open(_bad, 'wb') as _fh:
        _fh.write(b'PK\x03\x04' + b'\x00' * 400)
    for _label, _fn in (
            ('refs', lambda: count_image_refs(_bad)),
            ('extract', lambda: extract_images(_bad, '/tmp/_ci_badx')),
            ('display', lambda: media_display_inches(_bad)),
            ('governor', lambda: optimize_docx(_bad, '/tmp/_ci_o.docx', 1)),
            ('invariants', lambda: docx_invariants(_bad))):
        try:
            _fn()
            check(f'io_corrupt_named_{_label}', False)
        except IntegrityError:
            check(f'io_corrupt_named_{_label}', True)
        except Exception:
            check(f'io_corrupt_named_{_label}', False)

    _missing_ok = True
    try:
        _safe_document(_bad)
        _missing_ok = False
    except IntegrityError:
        pass
    except Exception:
        _missing_ok = False
    check('io_safe_document_named', _missing_ok)

    # ── v1.0.1 parity must not fire on a legitimate rename ───────────────────
    _pdir = f'/tmp/corpus_io_parity/{os.getpid()}'
    os.makedirs(_pdir, exist_ok=True)
    try:
        import random as _rnd
        from PIL import Image as _Im
        import docx as _dx
        _rnd.seed(11)
        _png = os.path.join(_pdir, 'photo.png')
        _im = _Im.new('RGB', (600, 400))
        _im.putdata([(_rnd.randrange(256), _rnd.randrange(256), _rnd.randrange(256))
                     for _ in range(600 * 400)])
        _im.save(_png)                      # PNG source -> jpeg route -> RENAME
        _d = _dx.Document()
        _d.add_paragraph('Q.1 figure').add_run().add_picture(_png)
        _src = os.path.join(_pdir, 'src.docx')
        _dst = os.path.join(_pdir, 'dst.docx')
        _d.save(_src)
        _ok, _rep, _log = optimize_docx(_src, _dst, budget=os.path.getsize(_src) // 2)
        check('io_parity_rename_engaged', _rep['tier'] == 'T1')
        _renamed = ('word/media/photo.png' not in zipfile.ZipFile(_dst).namelist())
        check('io_parity_rename_happened', _renamed)
        _before = docx_invariants(_src)['px']
        _after = docx_invariants(_dst)['px']
        check('io_parity_rename_dims_unchanged',
              sorted(_before.values()) == sorted(_after.values()))
        try:
            assert_docx_parity(_src, _dst, allow_resample=False)
            check('io_parity_rename_no_false_stop', True)
        except IntegrityError:
            check('io_parity_rename_no_false_stop', False)
        # and a REAL dimension change must still be caught
        _ok2, _rep2, _ = optimize_docx(_src, _dst, budget=1, force_tier='T4')
        _shrunk = docx_invariants(_dst)['px']
        if sorted(_shrunk.values()) != sorted(_before.values()):
            try:
                assert_docx_parity(_src, _dst, allow_resample=False)
                check('io_parity_real_resize_still_caught', False)
            except IntegrityError:
                check('io_parity_real_resize_still_caught', True)
            check('io_parity_real_resize_allowed_when_permitted',
                  bool(assert_docx_parity(_src, _dst, allow_resample=True)))
    except DependencyMissing:
        pass                                # optional deps absent — nothing to assert

    # ── CLUSTER K — ANALYSIS DOC (GAP-2026-07-25-002) ──────────────────────
    import tempfile as _tf
    with _tf.TemporaryDirectory() as _tmp:
        _ck_selftest(check, _tmp)

    # ── CLUSTER V — VISION PHASE A / OBSERVATION I/O ────────────────────────
    with _tf.TemporaryDirectory() as _vtmp:
        # EC-V1 — a text-only exam produces a valid empty queue, not a failure.
        _q0 = build_vision_queue([], _vtmp)
        check('v_ec1_empty_build',
              _q0['items'] == [] and _q0['sheets'] == [] and _q0['stats']['queued'] == 0)
        check('v_queue_file_written',
              os.path.exists(os.path.join(_vtmp, VISION_QUEUE_NAME)))

        try:
            _Img = _need('PIL')
        except DependencyMissing:
            _Img = None

        if _Img is not None:
            def _mk(p, size=(120, 90), colour=(200, 30, 30)):
                _Img.new('RGB', size, colour).save(p, 'PNG')
                return p
            _src = [_mk(os.path.join(_vtmp, f'f{i}.png')) for i in range(14)]
            # EC-V15 — the SAME q_num in two different papers must stay distinct.
            _items = ([{'paper_id': 'P_A', 'q_num': i + 1, 'srcs': [_src[i]],
                        'subtopic': 'S1'} for i in range(7)] +
                      [{'paper_id': 'P_B', 'q_num': i + 1, 'srcs': [_src[7 + i]],
                        'subtopic': 'S2'} for i in range(7)])
            _q = build_vision_queue(_items, _vtmp, per_sheet=6)
            check('v_ec15_distinct_keys', len(_q['items']) == 14)
            check('v_ec15_distinct_tags',
                  len({i['tag'] for i in _q['items']}) == 14)

            # EC-V10 — THE fixture the design turns on: every queued item appears on
            # exactly one sheet, and every sheet named in an item actually exists.
            _seen = {}
            for _it in _q['items']:
                _seen.setdefault((_it['paper_id'], _it['q_num']), []).append(_it['sheet'])
            check('v_ec10_each_item_once',
                  len(_seen) == 14 and all(len(v) == 1 for v in _seen.values()))
            check('v_ec10_sheets_exist',
                  all(os.path.exists(os.path.join(_vtmp, s)) for s in _q['sheets']))
            check('v_ec10_sheet_count', _q['stats']['sheets'] == 3)   # 14 @ 6 -> 3
            check('v_ec10_cells_within_per_sheet',
                  all(0 <= i['cell'] < 6 for i in _q['items']))

            # EC-V10 — re-tiling at a different per_sheet keeps the invariant.
            _q2 = build_vision_queue(_items, _vtmp, per_sheet=4)
            check('v_ec10_retile', _q2['stats']['sheets'] == 4
                  and len(_q2['items']) == 14)
            # tags are stable across a re-tile: observations survive re-tiling
            check('v_tag_stable_across_retile',
                  {i['tag'] for i in _q['items']} == {i['tag'] for i in _q2['items']})

            # v1.9 — build_vision_queue() now UNIONS with the queue already in the
            # workdir (GAP-2026-07-27-B), so each single-build assertion below gets its
            # OWN sub-workdir. Sharing one directory would make every later assertion
            # see the accumulated set. The union itself is asserted explicitly further
            # down; these cases are about ONE build in isolation.
            def _iso(name):
                _d = os.path.join(_vtmp, '_iso_' + name)
                os.makedirs(_d, exist_ok=True)
                return _d

            # EC-V6 — a 3-panel question is ONE queue item.
            _q3 = build_vision_queue(
                [{'paper_id': 'P_C', 'q_num': 5, 'srcs': _src[:3]}], _iso('ec6'))
            check('v_ec6_multi_image_one_item',
                  len(_q3['items']) == 1 and _q3['items'][0]['renderable'])

            # EC-V8 — an unopenable part is queued and flagged, never dropped.
            _bad = os.path.join(_vtmp, 'broken.emf')
            with open(_bad, 'wb') as _f:
                _f.write(b'not an image')
            _q4 = build_vision_queue(
                [{'paper_id': 'P_D', 'q_num': 1, 'srcs': [_bad]}], _iso('ec8'))
            check('v_ec8_unrenderable_queued', len(_q4['items']) == 1)
            check('v_ec8_unrenderable_flagged',
                  _q4['items'][0]['renderable'] is False
                  and _q4['stats']['unrenderable'] == 1)

            # EC-V27 — the same (paper_id,q_num) twice collapses to ONE item with
            # both sources; two cells sharing a tag would make a Phase-B observation
            # unattributable. Found by property fuzzing.
            _q4b = build_vision_queue(
                [{'paper_id': 'P_E', 'q_num': 3, 'srcs': [_src[0]], 'subtopic': None},
                 {'paper_id': 'P_E', 'q_num': 3, 'srcs': [_src[1]], 'subtopic': 'S9'}],
                _iso('ec27'))
            check('v_ec27_dedup_one_item', len(_q4b['items']) == 1)
            check('v_ec27_dedup_srcs_unioned',
                  len(_q4b['items'][0]['srcs']) == 2)
            check('v_ec27_dedup_meta_filled',
                  _q4b['items'][0]['subtopic'] == 'S9')
            check('v_ec27_tags_unique',
                  len({i['tag'] for i in _q4b['items']}) == len(_q4b['items']))

            # malformed item dicts are skipped without crashing the build
            _q5 = build_vision_queue(
                [{'no_paper': 1}, {'paper_id': 'X', 'q_num': 'NaN'},
                 {'paper_id': 'X', 'q_num': 2, 'srcs': [_src[0]]}], _iso('malformed'))
            check('v_malformed_items_tolerated', len(_q5['items']) == 1)

            # ── END-TO-END: A -> B -> C -> profile ───────────────────────────
            _e2e = _iso('e2e')
            _qe = build_vision_queue(_items[:6], _e2e, per_sheet=6)
            _obs = [{'tag': i['tag'], 'figure_readable': True,
                     'object_type': 'micrograph', 'transformation_type': 'N/A',
                     'arrangement': 'single', 'complexity': 'Simple'}
                    for i in _qe['items']]
            write_vision_observations(_e2e, _obs)
            _read, _why = load_vision_observations(_e2e)
            check('v_obs_roundtrip', len(_read) == 6 and _why == '')
            _bk, _st = bc.merge_vision_observations(_qe['items'], _read)
            check('v_e2e_all_observed',
                  _st['vision_status'] == 'observed' and _st['missing'] == 0)
            _prof = bc.vision_profile(list(_bk.values()), min_dominant=5)
            check('v_e2e_profile_dominant',
                  _prof['object_types']['dominant'] == ['micrograph'])
            check('v_e2e_profile_status', _prof['vision_status'] == 'observed')

            # NEGATIVE TEST — Phase B skipped. Must NOT halt, must report.
            # v1.9: reads the SAME isolated dir the e2e block wrote to, so removing the
            # observations file leaves that queue's 6 items unobserved rather than
            # deleting a file the shared workdir never had.
            os.remove(os.path.join(_e2e, VISION_OBS_NAME))
            _none, _why2 = load_vision_observations(_e2e)
            check('v_neg_no_obs_file', _none == [] and 'Phase B' in _why2)
            _bk2, _st2 = bc.merge_vision_observations(_qe['items'], _none)
            check('v_neg_status_unavailable',
                  _st2['vision_status'] == 'unavailable' and _st2['observed'] == 0)
            _prof2 = bc.vision_profile(list(_bk2.values()))
            check('v_neg_profile_empty_but_shaped',
                  _prof2['object_types']['dominant'] == []
                  and _prof2['vision_status'] == 'unavailable'
                  and _prof2['images_unobserved'] == 6)

            # ── v1.9 IDEMPOTENT UNION (GAP-2026-07-27-B) ─────────────────────
            # The defect: build_vision_queue() wrote fixed filenames and never read
            # what was there, so a caller invoking it once per PAPER destroyed every
            # earlier paper's queue and sheets. Measured on IIT_JAM_BIOTECHNOLOGY:
            # 153 figural questions across 22 papers, _meta.vision recorded queued=8.
            _u = _iso('union')
            _u1 = build_vision_queue(
                [{'paper_id': 'PAPER_ALPHA', 'q_num': n, 'srcs': [_src[0]]}
                 for n in (1, 2, 3)], _u)
            _u2 = build_vision_queue(
                [{'paper_id': 'PAPER_BETA', 'q_num': n, 'srcs': [_src[0]]}
                 for n in (7, 8)], _u)
            check('v_union_accumulates', _u2['stats']['queued'] == 5)
            check('v_union_carried_in', _u2['stats']['carried_in'] == 3)
            check('v_union_submitted', _u2['stats']['submitted'] == 2)
            # Prior tags MUST survive, or Phase B observations already recorded against
            # them silently orphan (EC-V11/V12).
            # The invariant is CONDITIONAL, deliberately: a paper_id hash collision may
            # legitimately widen the code and re-tag the whole queue (EC-V12). What must
            # never happen is a re-tag that is not REPORTED — a silent one orphans every
            # Phase B observation already recorded. Asserting unconditional tag survival
            # would encode a hash outcome rather than the contract.
            check('v_union_prior_tags_survive_or_reported',
                  _u2['stats']['tag_generation_changed']
                  or {i['tag'] for i in _u1['items']} <= {i['tag'] for i in _u2['items']})
            # Realistic paper_ids (long and varied, as make_paper_id emits) do not
            # collide: verified across the 22 real ids of the reference corpus, width
            # stayed 4 with zero generation changes over an incremental 22-paper build.
            check('v_union_no_generation_change',
                  _u2['stats']['tag_generation_changed'] is False)
            # Re-submitting the same items is a no-op, so a resumed session that replays
            # a batch cannot inflate the queue.
            _u3 = build_vision_queue(
                [{'paper_id': 'PAPER_ALPHA', 'q_num': n, 'srcs': [_src[0]]}
                 for n in (1, 2, 3)], _u)
            check('v_union_resubmit_idempotent', _u3['stats']['queued'] == 5)
            check('v_union_empty_submit_preserves',
                  build_vision_queue([], _u)['stats']['queued'] == 5)
            # The width floor stops a SHRINKING key set from re-tagging survivors.
            _w = _iso('width')
            _wide = build_vision_queue(
                [{'paper_id': 'P1', 'q_num': 1, 'srcs': [_src[0]]},
                 {'paper_id': 'P2', 'q_num': 1, 'srcs': [_src[0]]}], _w)
            _narrow = build_vision_queue(
                [{'paper_id': 'P1', 'q_num': 1, 'srcs': [_src[0]]}], _w)
            check('v_union_width_floor_holds',
                  _narrow['tag_width'] == _wide['tag_width'])
            check('v_union_width_floor_no_retag',
                  {i['tag'] for i in _wide['items'] if i['paper_id'] == 'P1'}
                  == {i['tag'] for i in _narrow['items'] if i['paper_id'] == 'P1'})

            # ── v1.10 `fresh` — RUN-SCOPED OPT-OUT OF THE UNION ──────────────
            _fr = _iso('fresh')
            build_vision_queue(
                [{'paper_id': 'PAPER_ALPHA', 'q_num': n, 'srcs': [_src[0]]}
                 for n in (1, 2)], _fr)
            _f2 = build_vision_queue(
                [{'paper_id': 'PAPER_BETA', 'q_num': 3, 'srcs': [_src[0]]}],
                _fr, fresh=True)
            check('v_fresh_ignores_prior', _f2['stats']['queued'] == 1
                  and _f2['stats']['carried_in'] == 0)
            # observations from the discarded generation must not survive a fresh
            # build — Step 1's Phase C would otherwise merge a prior paper's
            # observations into this paper's queue.
            write_vision_observations(_fr, [{'tag': _f2['items'][0]['tag'],
                                             'figure_readable': True}])
            _f3 = build_vision_queue(
                [{'paper_id': 'PAPER_GAMMA', 'q_num': 1, 'srcs': [_src[0]]}],
                _fr, fresh=True)
            check('v_fresh_clears_observations',
                  load_vision_observations(_fr)[0] == [])
            # default remains the union: the very next call without fresh carries.
            _f4 = build_vision_queue(
                [{'paper_id': 'PAPER_DELTA', 'q_num': 1, 'srcs': [_src[0]]}], _fr)
            check('v_fresh_default_still_unions', _f4['stats']['queued'] == 2)

            # ── v1.9 CLUSTER Q — ORIGINAL POSITION CODEC (GAP-2026-07-27-E) ──
            check('q_stamp_basic',
                  stamp_original_q_num('[15-Feb-2026 Shift 1]', 37)
                  == '[15-Feb-2026 Shift 1 Q37]')
            check('q_roundtrip',
                  parse_original_q_num(stamp_original_q_num('[15-Feb-2026]', 5)) == 5)
            check('q_stamp_idempotent',
                  stamp_original_q_num('[15-Feb-2026 Q9]', 12) == '[15-Feb-2026 Q9]')
            check('q_none_num_noop',
                  stamp_original_q_num('[15-Feb-2026]', None) == '[15-Feb-2026]')
            # A pre-v1.18 label MUST read as UNKNOWN, never as position 0.
            check('q_pre_v118_label_none',
                  parse_original_q_num('[15-Feb-2026 Shift 1]') is None)
            check('q_empty_none', parse_original_q_num('') is None
                  and parse_original_q_num(None) is None)
            check('q_non_label_untouched',
                  stamp_original_q_num('not a label', 3) == 'not a label')
            _ms = [{'q_range': [1, 30], 'question_type': 'MCQ'},
                   {'q_range': [31, 40], 'question_type': 'MSQ'},
                   {'q_range': [41, 60], 'question_type': 'NAT'}]
            check('q_band_msq', question_type_for_position(37, _ms) == 'MSQ')
            check('q_band_edges', question_type_for_position(31, _ms) == 'MSQ'
                  and question_type_for_position(40, _ms) == 'MSQ'
                  and question_type_for_position(41, _ms) == 'NAT')
            check('q_band_uncovered_none', question_type_for_position(99, _ms) is None)
            check('q_band_no_scheme_none',
                  question_type_for_position(37, []) is None
                  and question_type_for_position(None, _ms) is None)
            check('q_band_malformed_skipped',
                  question_type_for_position(37, [{'q_range': ['x', 'y']}] + _ms) == 'MSQ')

            # queue round-trip
            _lq, _lwhy = load_vision_queue(_e2e)
            check('v_queue_roundtrip', _lq is not None and _lwhy == '')

        # malformed / missing observation files never raise
        _badobs = os.path.join(_vtmp, 'bad_obs.json')
        with open(_badobs, 'w', encoding='utf-8') as _f:
            _f.write('{not json')
        check('v_obs_malformed_safe',
              load_vision_observations(_vtmp, path=_badobs)[0] == []
              and 'unreadable' in load_vision_observations(_vtmp, path=_badobs)[1])
        with open(_badobs, 'w', encoding='utf-8') as _f:
            _f.write('"a string"')
        check('v_obs_wrong_shape_safe',
              load_vision_observations(_vtmp, path=_badobs)[0] == [])
        check('v_queue_missing_safe',
              load_vision_queue(os.path.join(_vtmp, 'nope'))[0] is None)

    # ── CLUSTER I — TABLE STRUCTURE FIXTURE (GAP-2026-07-29-TBL, v1.11) ──────
    # THE fixture the 2026.07.29 deployment shipped without. The old row.cells
    # implementation returns one entry per GRID COLUMN and substitutes the
    # ANCHOR for every covered position, so on this table it yields
    #   [['Region','Printers','Printers','Printers','Printers'],
    #    ['Region','Laser','Inkjet','Dot','3D'], ...]
    # and the anchor-once expectations below FAIL. A revert of _table_rows to
    # row.cells therefore breaks this suite instead of passing green.
    import xml.etree.ElementTree as _ET
    _W_ = '{%s}' % _NS_W  # noqa: F841 — kept for symmetry with the parser
    _tbl_merged = _ET.fromstring(
        ('<w:tbl xmlns:w="%s">'
         ' <w:tr>'
         '  <w:tc><w:tcPr><w:vMerge w:val="restart"/></w:tcPr>'
         '   <w:p><w:r><w:t>Region</w:t></w:r></w:p></w:tc>'
         '  <w:tc><w:tcPr><w:gridSpan w:val="4"/></w:tcPr>'
         '   <w:p><w:r><w:t>Print</w:t></w:r><w:r><w:t>ers</w:t></w:r></w:p></w:tc>'
         ' </w:tr>'
         ' <w:tr>'
         '  <w:tc><w:tcPr><w:vMerge/></w:tcPr><w:p/></w:tc>'
         '  <w:tc><w:p><w:r><w:t>Laser</w:t></w:r></w:p></w:tc>'
         '  <w:tc><w:p><w:r><w:t>Inkjet</w:t></w:r></w:p></w:tc>'
         '  <w:tc><w:p><w:r><w:t>Dot</w:t></w:r></w:p></w:tc>'
         '  <w:tc><w:p><w:r><w:t>3D</w:t></w:r></w:p></w:tc>'
         ' </w:tr>'
         ' <w:tr>'
         '  <w:tc><w:p><w:r><w:t>North</w:t></w:r></w:p></w:tc>'
         '  <w:tc><w:p><w:r><w:t>5</w:t></w:r></w:p></w:tc>'
         '  <w:tc><w:p><w:r><w:t>6</w:t></w:r></w:p></w:tc>'
         '  <w:tc><w:p><w:r><w:t>7</w:t></w:r></w:p></w:tc>'
         '  <w:tc><w:p><w:r><w:t>8</w:t></w:r></w:p></w:tc>'
         ' </w:tr>'
         '</w:tbl>') % _NS_W)
    _rows = _table_rows(_tbl_merged, None)
    # anchor once, never a repeat — this is exactly what row.cells cannot do
    check('tbl_anchor_once_gridspan_header', _rows[0] == ['Region', 'Printers'])
    check('tbl_anchor_once_vmerge_label', _rows[1] == ['Laser', 'Inkjet', 'Dot', '3D'])
    check('tbl_plain_body_row', _rows[2] == ['North', '5', '6', '7', '8'])
    check('tbl_no_padding_cells', all('' not in r for r in _rows))
    # geometry survives the read-back: cs=4 header span + rs=2 label span
    _spec = read_table_spec(_tbl_merged)
    check('tbl_spec_spans_read_back',
          table_spec_spans(_spec) == [(0, 0, 2, 1), (0, 1, 1, 4)])
    check('tbl_spec_geometry_clean',
          place_cells(normalise_table_spec(_spec)['grid'])[3] == [])
    # flat-table BYTE-IDENTITY: output on a spanless table is pinned to the exact
    # serialised bytes the legacy implementation produced
    _tbl_flat = _ET.fromstring(
        ('<w:tbl xmlns:w="%s">'
         '<w:tr><w:tc><w:p><w:r><w:t>a</w:t></w:r></w:p></w:tc>'
         '<w:tc><w:p><w:r><w:t>b</w:t></w:r></w:p></w:tc></w:tr>'
         '<w:tr><w:tc><w:p><w:r><w:t>c</w:t></w:r></w:p></w:tc>'
         '<w:tc><w:p><w:r><w:t>d</w:t></w:r></w:p></w:tc></w:tr>'
         '</w:tbl>') % _NS_W)
    check('tbl_flat_byte_identity',
          json.dumps(_table_rows(_tbl_flat, None)) == json.dumps([['a', 'b'], ['c', 'd']]))
    # legacy {'headers','rows'} DI payload — accepted for ever, never migrated
    _leg = normalise_table_spec({'headers': ['X', 'Y'], 'rows': [['1', '2']]})
    check('tbl_legacy_di_payload',
          _leg['grid'][0][0]['t'] == 'X' and _leg.get('header_rows') == 1)
    check('tbl_legacy_flat_no_spans', table_spec_spans(_leg) == [])
    # full builder round trip: build_di_table -> raw-XML read-back. Spans must
    # survive and NO stray empty cell may appear (the Q.52/Q.61 defect shape).
    try:
        import docx as _dxT
        _dT = _dxT.Document()
        _built = build_di_table(_dT, {'grid': [
            [{'t': 'Region', 'rs': 2}, {'t': 'Printers', 'cs': 4}],
            [{'t': 'Laser'}, {'t': 'Inkjet'}, {'t': 'Dot'}, {'t': '3D'}],
            [{'t': 'North'}, {'t': '5'}, {'t': '6'}, {'t': '7'}, {'t': '8'}]],
            'header_rows': 1})
        _rt_rows = _table_rows(_built._tbl, None)
        check('tbl_build_roundtrip_values',
              _rt_rows == [['Region', 'Printers'],
                           ['Laser', 'Inkjet', 'Dot', '3D'],
                           ['North', '5', '6', '7', '8']])
        check('tbl_build_roundtrip_spans',
              table_spec_spans(read_table_spec(_built._tbl)) == [(0, 0, 2, 1), (0, 1, 1, 4)])
        check('tbl_build_no_stray_empty',
              sum(r.count('') for r in _rt_rows) == 0)
    except ImportError:
        pass                                # python-docx absent — parse-side already covered

    # ── is_option / para_has_image FIXTURE (follow-up open since 2026.07.26.2) ─
    # bare-marker + OPT_PATTERNS cases, image-option path, Paragraph-vs-element
    # duality, clean_option_text stripping. The single shared predicate every
    # step delegates to finally has its own regression net.
    check('opt_num_dot', is_option('1. speed of light'))
    check('opt_alpha_dot', is_option('B. mitochondria'))
    check('opt_paren_num', is_option('(3) 42'))
    check('opt_paren_alpha_lower', is_option('(c) all of the above'))
    check('opt_half_paren', is_option('d) none of these'))
    check('opt_beyond_range_rejected', not is_option('6. beyond the option range'))
    check('opt_prose_rejected', not is_option('Consider the following statements'))
    check('opt_bare_no_para_rejected', not is_option('1.'))

    # ── text_of() — w:t + m:t canonical accessor (GAP-2026-08-07-OMML M3) ──
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    from xml.etree import ElementTree as _ET   # stdlib — Check AB-safe
    _p_omml = _ET.fromstring(
        ('<w:p xmlns:w="%s" xmlns:m="%s">'
         '<w:r><w:t xml:space="preserve">3. </w:t></w:r>'
         '<m:oMath><m:f><m:num><m:r><m:t>1</m:t></m:r></m:num>'
         '<m:den><m:r><m:t>2</m:t></m:r></m:den></m:f></m:oMath>'
         '</w:p>') % (_W, _M))
    check('textof_wt_plus_mt', text_of(_p_omml) == '3. 12')
    check('textof_none', text_of(None) == '')
    check('textof_non_element', text_of('junk') == '')

    _p_img = _ET.fromstring(
        '<w:p xmlns:w="%s" xmlns:a="%s"><w:r><w:drawing><a:blip/></w:drawing></w:r></w:p>'
        % (_NS_W, A_NS))
    _p_txt = _ET.fromstring(
        '<w:p xmlns:w="%s"><w:r><w:t>1.</w:t></w:r></w:p>' % _NS_W)
    check('opt_bare_with_image', is_option('1.', para=_p_img))
    check('opt_bare_alpha_with_image', is_option('(D)', para=_p_img))
    check('opt_bare_imageless_rejected', not is_option('1.', para=_p_txt))
    check('opt_text_option_para_irrelevant', is_option('2. tundra', para=_p_txt))

    # ── GAP-2026-08-05-001 (S-1) — an option's content may be an EQUATION ────
    # PYQSort CHECK 4 accepts "(d) OMML formula" as a valid option form; para_has_image()
    # sees only a blip/VML, so these returned False and the option was silently dropped.
    # Real instance: SSC_CGL_TIER1 09-Sep-2024 Shift 1, Q.75 options 1 and 2 (paras
    # 592-593) — a delivered sorted paper losing half its option set at Step 5.
    # These FAIL on the pre-fix predicate.
    _p_eqn = _ET.fromstring(
        '<w:p xmlns:w="%s" xmlns:m="%s"><w:r><w:t>1.</w:t></w:r><m:oMath/></w:p>'
        % (_NS_W, 'http://schemas.openxmlformats.org/officeDocument/2006/math'))
    _p_obj = _ET.fromstring(
        '<w:p xmlns:w="%s"><w:r><w:t>2.</w:t><w:object/></w:r></w:p>' % _NS_W)
    _p_num = _ET.fromstring(
        '<w:p xmlns:w="%s"><w:pPr><w:numPr/></w:pPr><w:r><w:t>1.</w:t></w:r></w:p>' % _NS_W)
    check('opt_bare_with_equation', is_option('1.', para=_p_eqn))
    check('opt_bare_with_embedded_object', is_option('2.', para=_p_obj))
    # an auto-number alone is NOT option content — the rendered "1." only duplicates
    # the label already present as text, so this must stay rejected.
    check('opt_bare_autonumber_only_rejected', not is_option('1.', para=_p_num))

    class _FakeParagraph:                   # python-docx Paragraph form (has ._p)
        pass
    _fp = _FakeParagraph()
    _fp._p = _p_img
    check('opt_paragraph_form_accepted', is_option('A)', para=_fp))
    check('opt_para_none_safe', not para_has_image(None))
    check('opt_para_junk_safe', not para_has_image('not a paragraph'))
    check('opt_clean_marker_stripped',
          clean_option_text('(c) all of the above') == 'all of the above')
    check('opt_clean_bare_empty', clean_option_text('B.') == '')
    check('opt_clean_plain_untouched',
          clean_option_text('not an option') == 'not an option')

    # ── GAP-2026-08-15-PYQCOUNT-DRIVE-ACQUISITION ────────────────────────────
    # P3b: the six-shape decode matrix. Shape B (bare base64) was REJECTED before
    # this release; every other shape must keep working byte-for-byte.
    import base64 as _b64, tempfile
    _docx = b'PK\x03\x04' + b'\x00' * 136
    _b = _b64.b64encode(_docx).decode()
    _paper = {'name': 'probe.docx', 'fileSize': len(_docx), 'id': 'FID'}
    check('t_decode_A_bytes', decode_drive_payload(_docx) == _docx)
    check('t_decode_B_bare_base64', decode_drive_payload(_b) == _docx)
    check('t_decode_C_dict_content', decode_drive_payload({'content': _b}) == _docx)
    check('t_decode_D_list_text',
          decode_drive_payload([{'text': '{"content":"%s"}' % _b}]) == _docx)
    check('t_decode_E_json_string',
          decode_drive_payload('{"content":"%s"}' % _b) == _docx)
    _spill = os.path.join(tempfile.mkdtemp(), 'spill.txt')
    with open(_spill, 'w', encoding='utf-8') as _fh:
        _fh.write('{"content":"%s","id":"FID","title":"probe.docx"}' % _b)
    check('t_decode_F_spill_path', decode_drive_payload(_spill) == _docx)
    # A spill path may live in ANY directory — the deployment chooses it, and the
    # 2026-08-15 measurement found two different directories and one inline channel.
    check('t_decode_F_any_directory', _spill.startswith(tempfile.gettempdir()))
    # bare-base64 acceptance must not swallow genuine garbage
    _bad = 0
    for _junk in ('', 'not base64 at all!!', 'AAAA', _b64.b64encode(b'no magic').decode()):
        try:
            decode_drive_payload(_junk)
        except TransportFallback:
            _bad += 1
    check('t_decode_rejects_non_docx', _bad == 4)

    # P3a: stage_drive_payload gives a verified route to disk with NO callable.
    _dest = tempfile.mkdtemp()
    check('t_stage_writes_bytes',
          open(stage_drive_payload({'content': _b}, _paper, _dest), 'rb').read() == _docx)
    check('t_stage_accepts_spill_path',
          open(stage_drive_payload(_spill, _paper, _dest), 'rb').read() == _docx)
    _tr = 0
    try:
        stage_drive_payload({'content': _b64.b64encode(_docx[:100]).decode()},
                            _paper, _dest)
    except TransportFallback:
        _tr = 1
    check('t_stage_truncation_raises_fallback', _tr == 1)
    # fetch_drive_docx must produce the IDENTICAL result through a resolver, because
    # it is now a thin wrapper — the two routes can never drift apart again.
    check('t_fetch_via_resolver_matches_stage',
          open(fetch_drive_docx(lambda fid: {'content': _b}, _paper, _dest),
               'rb').read() == _docx)
    _cap = 0
    try:
        fetch_drive_docx(lambda fid: _docx,
                         {'name': 'huge.docx', 'id': 'X',
                          'fileSize': bc.DRIVE_CAP + 1}, _dest)
    except TransportFallback:
        _cap = 1
    check('t_fetch_cap_precheck_still_fires', _cap == 1)

    print(f"SELF-TEST: {passed}/{total} PASS")
    if fails:
        print('FAILED: ' + ', '.join(fails))
    return passed == total


if __name__ == '__main__':
    import sys
    if '--self-test' in sys.argv:
        sys.exit(0 if self_test() else 1)
    print('corpus_io.py — corpus acquisition / image integrity / size governor. '
          'Run with --self-test.')
